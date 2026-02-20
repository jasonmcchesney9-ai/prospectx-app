"""
ProspectX API Server — SQLite Version
Zero external DB dependencies. Just SQLite + FastAPI.
"""

# Prevent Python from writing .pyc files to __pycache__/
# This stops Turbopack from detecting file changes in backend/ and crashing
import os
os.environ["PYTHONDONTWRITEBYTECODE"] = "1"
import sys
sys.dont_write_bytecode = True

import asyncio
import json
import logging
import os
import re
import secrets
import hashlib
import sqlite3
import sys
import threading
import time
import uuid
from contextlib import contextmanager
from datetime import datetime, timedelta, timezone
from typing import Any, Dict, List, Optional

# Ensure backend directory is on sys.path so sibling modules (rink_diagrams, hockeytech) are importable
_this_dir = os.path.dirname(os.path.abspath(__file__))
if _this_dir not in sys.path:
    sys.path.insert(0, _this_dir)

import csv
import io

import glob
import httpx
import jwt
import shutil
from dotenv import load_dotenv
from fastapi import FastAPI, HTTPException, Query, Request, Depends, File, UploadFile, Body, status
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import JSONResponse, FileResponse, StreamingResponse
from fastapi.security import HTTPBearer, HTTPAuthorizationCredentials
from fastapi.staticfiles import StaticFiles
from passlib.context import CryptContext
from pydantic import BaseModel, Field, field_validator
try:
    from rink_diagrams import generate_drill_diagram
except ImportError:
    generate_drill_diagram = None
    logging.getLogger("prospectx").warning(
        "rink_diagrams not available — drill diagram generation disabled"
    )

try:
    from docx import Document as DocxDocument
    from docx.shared import Inches, Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    _docx_available = True
except ImportError:
    _docx_available = False
    logging.getLogger("prospectx").warning(
        "python-docx not available — DOCX export disabled"
    )

from pxi_prompt_core import (
    PXI_CORE_GUARDRAILS,
    PXI_MODE_BLOCKS,
    PXI_MODES,
    VALID_MODES,
    MODE_TEMPLATE_WIRING,
    REQUIRED_SECTIONS_BY_TYPE,
    build_report_system_prompt,
    resolve_mode,
)

# Load .env from the backend directory (works regardless of CWD)
_backend_dir = os.path.dirname(os.path.abspath(__file__))
load_dotenv(os.path.join(_backend_dir, ".env"), override=True)

# ============================================================
# LOGGING
# ============================================================

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s | %(levelname)-8s | %(name)s | %(message)s",
    datefmt="%Y-%m-%d %H:%M:%S",
)
logger = logging.getLogger("prospectx")

# ============================================================
# CONFIG
# ============================================================

# Store SQLite DB outside the Next.js project folder so Turbopack doesn't
# try to read the WAL lock files (.db-shm, .db-wal) and crash.
_DATA_DIR = os.path.join(os.path.expanduser("~"), ".prospectx")
os.makedirs(_DATA_DIR, exist_ok=True)
DB_FILE = os.path.join(_DATA_DIR, "prospectx.db")
_DEFAULT_JWT_SECRET = "prospectx_dev_secret_change_in_production_2026"
JWT_SECRET = os.getenv("JWT_SECRET", _DEFAULT_JWT_SECRET)
JWT_ALGORITHM = "HS256"
JWT_EXPIRY_HOURS = 8
ENVIRONMENT = os.getenv("ENVIRONMENT", "development")
FRONTEND_URL = os.getenv("FRONTEND_URL", "http://localhost:3000")

# P0-2: Refuse to start in production with default JWT secret
if ENVIRONMENT != "development" and JWT_SECRET == _DEFAULT_JWT_SECRET:
    print("FATAL: JWT_SECRET is still the default value in '%s' environment. Set a strong JWT_SECRET env var." % ENVIRONMENT)
    sys.exit(1)
ANTHROPIC_API_KEY = os.getenv("ANTHROPIC_API_KEY", "")

# ── Subscription Tiers ─────────────────────────────────────
SUBSCRIPTION_TIERS = {
    # ── Individual Plans ──────────────────────────────────
    "rookie": {
        "name": "Rookie",
        "price": 0,
        "annual_price": 0,
        "monthly_reports": 0,
        "monthly_bench_talks": 150,  # 5/day × 30
        "monthly_practice_plans": 0,
        "max_seats": 1,
        "description": "Browse players, live stats, and basic Bench Talk.",
        "target_user": "Casual browsers",
        "features": ["Browse player profiles", "HockeyTech live stats", "Elite Prospects links", "Basic search", "5 Bench Talk messages/day"],
        # Permissions
        "can_sync_data": False,
        "can_bulk_sync": False,
        "can_upload_files": False,
        "can_access_live_stats": True,
        "can_submit_corrections": False,
        "can_create_game_plans": False,
        "can_create_series": False,
        "can_use_scouting_list": False,
        "max_scouting_list": 0,
        "max_uploads_per_month": 0,
        "max_file_size_mb": 0,
        "players_tracked": 0,
    },
    "parent": {
        "name": "Parent",
        "price": 10.00,
        "annual_price": 100.00,
        "monthly_reports": 3,
        "monthly_bench_talks": 600,  # 20/day × 30
        "monthly_practice_plans": 0,
        "max_seats": 1,
        "description": "Development tracking, reports, and advisor directory for hockey parents.",
        "target_user": "Hockey parents",
        "features": ["3 reports/month", "20 Bench Talk/day", "Profile analytics", "Development tracking", "Advisor directory", "Add/claim your player", "Direct messaging"],
        # Permissions
        "can_sync_data": False,
        "can_bulk_sync": False,
        "can_upload_files": False,
        "can_access_live_stats": True,
        "can_submit_corrections": False,
        "can_create_game_plans": False,
        "can_create_series": False,
        "can_use_scouting_list": False,
        "max_scouting_list": 0,
        "max_uploads_per_month": 0,
        "max_file_size_mb": 0,
        "players_tracked": 1,
    },
    "scout": {
        "name": "Scout",
        "price": 25.00,
        "annual_price": 250.00,
        "monthly_reports": 10,
        "monthly_bench_talks": 1500,  # 50/day × 30
        "monthly_practice_plans": 3,
        "monthly_game_plans": 3,
        "max_seats": 1,
        "description": "Scout Notes, rated search, limited coaching tools, and PXI basic.",
        "target_user": "Indie scouts, coaches",
        "features": ["10 reports/month", "50 Bench Talk/day", "Scout Notes", "Rated search", "3 game plans/month", "3 practice plans/month", "PXI basic", "Individual team sync", "5 file uploads/month (10MB max)"],
        # Permissions
        "can_sync_data": True,
        "can_bulk_sync": False,
        "can_upload_files": True,
        "can_access_live_stats": True,
        "can_submit_corrections": True,
        "can_create_game_plans": True,
        "can_create_series": False,
        "can_use_scouting_list": True,
        "max_scouting_list": 10,
        "max_uploads_per_month": 5,
        "max_file_size_mb": 10,
        "players_tracked": 5,
    },
    "pro": {
        "name": "Pro",
        "price": 49.00,
        "annual_price": 490.00,
        "monthly_reports": -1,
        "monthly_bench_talks": -1,
        "monthly_practice_plans": -1,
        "max_seats": 1,
        "description": "Unlimited reports, full PXI (10 modes), all coaching tools, InStat, and export.",
        "target_user": "Head coaches, agents, GMs",
        "features": ["Unlimited reports", "Unlimited Bench Talk", "Full PXI (10 modes)", "All 21 report templates", "Unlimited coaching tools", "InStat game data", "Line builder & series planning", "Multi-team views", "Export & share reports", "Unlimited uploads (50MB max)"],
        # Permissions
        "can_sync_data": True,
        "can_bulk_sync": True,
        "can_upload_files": True,
        "can_access_live_stats": True,
        "can_submit_corrections": True,
        "can_create_game_plans": True,
        "can_create_series": True,
        "can_use_scouting_list": True,
        "max_scouting_list": 30,
        "max_uploads_per_month": -1,
        "max_file_size_mb": 50,
        "players_tracked": 10,
    },
    "elite": {
        "name": "Elite",
        "price": 99.00,
        "annual_price": 990.00,
        "monthly_reports": -1,
        "monthly_bench_talks": -1,
        "monthly_practice_plans": -1,
        "max_seats": 1,
        "description": "Everything in Pro plus PXI Auto-Scout, bulk reports, API access, and priority support.",
        "target_user": "Power users, agencies",
        "features": ["Everything in Pro", "PXI Auto-Scout", "Bulk report generation", "Aggregate scouting boards", "API access", "Priority support"],
        # Permissions
        "can_sync_data": True,
        "can_bulk_sync": True,
        "can_upload_files": True,
        "can_access_live_stats": True,
        "can_submit_corrections": True,
        "can_create_game_plans": True,
        "can_create_series": True,
        "can_use_scouting_list": True,
        "max_scouting_list": -1,
        "max_uploads_per_month": -1,
        "max_file_size_mb": 500,
        "players_tracked": -1,
        "priority_sync": True,
    },
    # ── Organization Plans ────────────────────────────────
    "team_org": {
        "name": "Team",
        "price": 249.00,
        "annual_price": 1990.00,
        "founders_price": 199.00,
        "monthly_reports": -1,
        "monthly_bench_talks": -1,
        "monthly_practice_plans": -1,
        "max_seats": 10,
        "description": "All Pro features for up to 10 seats with shared scouting data.",
        "target_user": "Single team organizations",
        "features": ["All Pro features per seat", "10 user seats", "Shared scout notes", "Team branding", "Org admin panel", "100MB uploads"],
        # Permissions
        "can_sync_data": True,
        "can_bulk_sync": True,
        "can_upload_files": True,
        "can_access_live_stats": True,
        "can_submit_corrections": True,
        "can_create_game_plans": True,
        "can_create_series": True,
        "can_use_scouting_list": True,
        "max_scouting_list": -1,
        "max_uploads_per_month": -1,
        "max_file_size_mb": 100,
        "players_tracked": 250,
    },
    "program_org": {
        "name": "Program",
        "price": 599.00,
        "annual_price": 4990.00,
        "founders_price": 499.00,
        "monthly_reports": -1,
        "monthly_bench_talks": -1,
        "monthly_practice_plans": -1,
        "max_seats": 30,
        "description": "Multi-team platform with cross-team scouting boards and org-wide analytics.",
        "target_user": "AAA organizations, multi-team programs",
        "features": ["Everything in Team", "30 user seats", "Cross-team scouting boards", "Org-wide analytics", "Multi-team branding", "Dedicated support", "500MB uploads", "Priority sync"],
        # Permissions
        "can_sync_data": True,
        "can_bulk_sync": True,
        "can_upload_files": True,
        "can_access_live_stats": True,
        "can_submit_corrections": True,
        "can_create_game_plans": True,
        "can_create_series": True,
        "can_use_scouting_list": True,
        "max_scouting_list": -1,
        "max_uploads_per_month": -1,
        "max_file_size_mb": 500,
        "players_tracked": -1,
        "priority_sync": True,
    },
    "enterprise": {
        "name": "Enterprise",
        "price": -1,  # Custom pricing — "Contact Us"
        "annual_price": -1,
        "monthly_reports": -1,
        "monthly_bench_talks": -1,
        "monthly_practice_plans": -1,
        "max_seats": -1,  # Unlimited
        "description": "Custom solution with unlimited seats, dedicated onboarding, and SLA support.",
        "target_user": "OHL/WHL teams, large agencies",
        "features": ["Everything in Program", "Unlimited seats", "Custom branding", "Dedicated onboarding", "White-glove setup", "SLA support", "API access"],
        # Permissions
        "can_sync_data": True,
        "can_bulk_sync": True,
        "can_upload_files": True,
        "can_access_live_stats": True,
        "can_submit_corrections": True,
        "can_create_game_plans": True,
        "can_create_series": True,
        "can_use_scouting_list": True,
        "max_scouting_list": -1,
        "max_uploads_per_month": -1,
        "max_file_size_mb": -1,
        "players_tracked": -1,
        "priority_sync": True,
    },
}


def _get_tier_permissions(tier: str) -> dict:
    """Get the full permissions config for a subscription tier."""
    return SUBSCRIPTION_TIERS.get(tier, SUBSCRIPTION_TIERS["rookie"])


def _check_tier_permission(user_id: str, permission: str, conn) -> dict:
    """Check if a user has a specific boolean permission (e.g., can_sync_data, can_upload_files).
    Raises 403 if not allowed. Returns tier config dict."""
    row = conn.execute("SELECT subscription_tier FROM users WHERE id = ?", (user_id,)).fetchone()
    if not row:
        raise HTTPException(status_code=404, detail="User not found")

    tier = row["subscription_tier"] or "rookie"
    tier_config = _get_tier_permissions(tier)

    if not tier_config.get(permission, False):
        perm_labels = {
            "can_sync_data": "Individual team sync requires Scout tier or higher.",
            "can_bulk_sync": "Bulk league sync requires Pro tier or higher.",
            "can_upload_files": "File uploads require Scout tier or higher.",
            "can_access_live_stats": "Live stats require Pro tier or higher.",
            "can_submit_corrections": "Submitting corrections requires Scout tier or higher.",
            "can_create_game_plans": "Game plans require Scout tier or higher.",
            "can_create_series": "Series planning requires Pro tier or higher.",
            "can_use_scouting_list": "Scouting list requires Scout tier or higher.",
        }
        raise HTTPException(status_code=403, detail={
            "error": f"{permission}_required",
            "message": perm_labels.get(permission, f"This feature requires a higher subscription tier."),
            "current_tier": tier,
            "upgrade_url": "/pricing",
        })

    return tier_config


def _check_tier_limit(user_id: str, resource_type: str, conn) -> dict:
    """Unified limit checker for any resource type.
    resource_type: 'reports', 'bench_talks', 'practice_plans', 'uploads'
    Checks usage_tracking table first, falls back to users table for legacy fields.
    Raises 429 if limit exceeded. Returns {tier, limit, used}."""
    row = conn.execute(
        "SELECT subscription_tier, monthly_reports_used, monthly_bench_talks_used, usage_reset_at FROM users WHERE id = ?",
        (user_id,),
    ).fetchone()
    if not row:
        raise HTTPException(status_code=404, detail="User not found")

    tier = row["subscription_tier"] or "rookie"
    tier_config = _get_tier_permissions(tier)

    # Map resource type to tier config key and users table column
    limit_map = {
        "reports": ("monthly_reports", "monthly_reports_used"),
        "bench_talks": ("monthly_bench_talks", "monthly_bench_talks_used"),
        "practice_plans": ("monthly_practice_plans", None),
        "uploads": ("max_uploads_per_month", None),
    }
    tier_key, legacy_col = limit_map.get(resource_type, (None, None))
    if not tier_key:
        return {"tier": tier, "limit": -1, "used": 0}

    limit = tier_config.get(tier_key, 0)

    # Monthly reset check (resets on 1st of month UTC)
    reset_at = row["usage_reset_at"]
    now = datetime.now(timezone.utc)
    month_start = now.replace(day=1, hour=0, minute=0, second=0, microsecond=0)
    needs_reset = not reset_at or datetime.fromisoformat(reset_at) < month_start
    if needs_reset:
        conn.execute(
            "UPDATE users SET monthly_bench_talks_used = 0, monthly_reports_used = 0, usage_reset_at = ? WHERE id = ?",
            (now.isoformat(), user_id),
        )
        conn.commit()

    # Get usage count — check usage_tracking table first, fall back to users column
    current_month = now.strftime("%Y-%m")
    used = 0

    # Try usage_tracking table (new system)
    tracking_col_map = {
        "reports": "reports_count",
        "bench_talks": "bench_talks_count",
        "practice_plans": "practice_plans_count",
        "uploads": "uploads_count",
    }
    tracking_col = tracking_col_map.get(resource_type)
    if tracking_col:
        tracking_row = conn.execute(
            f"SELECT {tracking_col} FROM usage_tracking WHERE user_id = ? AND month = ?",
            (user_id, current_month),
        ).fetchone()
        if tracking_row:
            used = tracking_row[tracking_col] or 0
        elif legacy_col and not needs_reset:
            # Fall back to legacy users column for reports/bench_talks
            used = row[legacy_col] or 0

    # -1 means unlimited
    if limit != -1 and used >= limit:
        raise HTTPException(status_code=429, detail={
            "error": f"{resource_type}_limit_reached",
            "tier": tier,
            "limit": limit,
            "used": used,
            "upgrade_url": "/pricing",
        })

    return {"tier": tier, "limit": limit, "used": used}


def _check_email_verified(user_id: str, conn):
    """Raise 403 if user's email is not verified. Used to gate AI features."""
    row = conn.execute("SELECT email_verified FROM users WHERE id = ?", (user_id,)).fetchone()
    if not row:
        raise HTTPException(status_code=404, detail="User not found")
    if not row["email_verified"]:
        raise HTTPException(status_code=403, detail={
            "error": "email_not_verified",
            "message": "Please verify your email before using this feature.",
        })


def _increment_tracking(user_id: str, resource_type: str, conn):
    """Increment the usage_tracking counter for a resource type. Also updates legacy users columns for backwards compat."""
    current_month = datetime.now(timezone.utc).strftime("%Y-%m")
    tracking_col_map = {
        "reports": "reports_count",
        "bench_talks": "bench_talks_count",
        "practice_plans": "practice_plans_count",
        "uploads": "uploads_count",
    }
    col = tracking_col_map.get(resource_type)
    if not col:
        return

    # Upsert into usage_tracking table
    conn.execute(f"""
        INSERT INTO usage_tracking (id, user_id, month, {col}, updated_at)
        VALUES (?, ?, ?, 1, CURRENT_TIMESTAMP)
        ON CONFLICT(user_id, month) DO UPDATE SET
            {col} = COALESCE({col}, 0) + 1,
            updated_at = CURRENT_TIMESTAMP
    """, (gen_id(), user_id, current_month))

    # Also update legacy users columns for backwards compat
    legacy_col_map = {
        "reports": "monthly_reports_used",
        "bench_talks": "monthly_bench_talks_used",
    }
    legacy_col = legacy_col_map.get(resource_type)
    if legacy_col:
        conn.execute(f"UPDATE users SET {legacy_col} = COALESCE({legacy_col}, 0) + 1 WHERE id = ?", (user_id,))

    conn.commit()

# ============================================================
# APP + MIDDLEWARE
# ============================================================

_is_deployed = bool(os.getenv("RAILWAY_ENVIRONMENT") or os.getenv("RAILWAY_PROJECT_ID") or os.getenv("VERCEL"))
_show_docs = ENVIRONMENT == "development" and not _is_deployed
app = FastAPI(
    title="ProspectX API",
    description="Decision-Grade Hockey Intelligence Platform",
    version="1.0.0",
    docs_url="/docs" if _show_docs else None,
    redoc_url="/redoc" if _show_docs else None,
    openapi_url="/openapi.json" if _show_docs else None,
)

# ── Rate Limiting ─────────────────────────────────
# Per-IP rate limits applied via middleware (no endpoint signature changes needed)
_RATE_LIMIT_WINDOW = 60  # seconds
_RATE_LIMITS = {
    "auth": 10,       # /auth/* — 10 requests/minute per IP
    "data": 60,       # /players, /reports, /scout-notes — 60/min
    "hockeytech": 30, # /hockeytech/* — 30/min
    "default": 120,   # everything else — 120/min
}
_rate_limit_store: Dict[str, Dict[str, Any]] = {}

def _get_rate_category(path: str) -> str:
    if path.startswith("/auth/"):
        return "auth"
    if path.startswith(("/players", "/reports", "/scout-notes", "/stats/")):
        return "data"
    if path.startswith("/hockeytech/"):
        return "hockeytech"
    return "default"

@app.middleware("http")
async def rate_limit_middleware(request: Request, call_next):
    if request.method == "OPTIONS":
        return await call_next(request)
    client_ip = (request.headers.get("x-forwarded-for", "").split(",")[0].strip()
                  or (request.client.host if request.client else "unknown"))
    category = _get_rate_category(request.url.path)
    limit = _RATE_LIMITS[category]
    key = f"{client_ip}:{category}"
    now = time.time()
    entry = _rate_limit_store.get(key)
    if entry and now - entry["window_start"] < _RATE_LIMIT_WINDOW:
        entry["count"] += 1
        if entry["count"] > limit:
            return JSONResponse(
                status_code=429,
                content={"detail": f"Rate limit exceeded. Max {limit} requests per minute for {category} endpoints."},
            )
    else:
        _rate_limit_store[key] = {"window_start": now, "count": 1}
    # Periodically clean stale entries (every 100 requests)
    if len(_rate_limit_store) > 1000:
        cutoff = now - _RATE_LIMIT_WINDOW * 2
        _rate_limit_store.clear()
    response = await call_next(request)
    response.headers["X-RateLimit-Limit"] = str(limit)
    remaining = max(0, limit - _rate_limit_store.get(key, {}).get("count", 0))
    response.headers["X-RateLimit-Remaining"] = str(remaining)
    return response

# ── Images directory ──────────────────────────────
_IMAGES_DIR = os.path.join(_DATA_DIR, "images")
os.makedirs(_IMAGES_DIR, exist_ok=True)

# Serve uploaded player images as static files
app.mount("/uploads", StaticFiles(directory=_IMAGES_DIR), name="uploads")

_allowed_origins = [
    "http://localhost:3000",
    "http://127.0.0.1:3000",
    "https://prospectx-app-ten.vercel.app",
    FRONTEND_URL,
]
# Remove duplicates and empty strings
_allowed_origins = list({o for o in _allowed_origins if o})

app.add_middleware(
    CORSMiddleware,
    allow_origins=_allowed_origins,
    allow_origin_regex=(
        r"https://prospectx-app(-[a-z0-9]+)?(-[a-z0-9]+)?\.vercel\.app"
        if (_is_deployed or ENVIRONMENT != "development")
        else r"^(https://prospectx-app[a-z0-9-]*\.vercel\.app|http://localhost(:\d+)?)$"
    ),
    allow_credentials=True,
    allow_methods=["GET", "POST", "PUT", "DELETE", "OPTIONS", "PATCH"],
    allow_headers=["*"],
    max_age=600,
)

# Password hashing
pwd_context = CryptContext(schemes=["bcrypt"], deprecated="auto", bcrypt__rounds=12)
security = HTTPBearer(auto_error=False)

# ============================================================
# REQUEST ID MIDDLEWARE
# ============================================================

def _log_error_to_db(method: str, path: str, status_code: int, error_msg: str):
    """Best-effort: log server errors to admin_error_log for the admin dashboard."""
    try:
        conn = get_db()
        conn.execute(
            "INSERT INTO admin_error_log (id, request_method, request_path, status_code, error_message, created_at) VALUES (?, ?, ?, ?, ?, ?)",
            (gen_id(), method, path, status_code, str(error_msg)[:500], datetime.now(timezone.utc).isoformat()),
        )
        conn.commit()
        conn.close()
    except Exception:
        pass  # Never let error logging break the actual response


@app.middleware("http")
async def add_request_id(request: Request, call_next):
    request_id = str(uuid.uuid4())[:8]
    try:
        response = await call_next(request)
        response.headers["X-Request-Id"] = request_id
        response.headers["X-Content-Type-Options"] = "nosniff"
        response.headers["X-Frame-Options"] = "DENY"
        response.headers["X-XSS-Protection"] = "1; mode=block"
        response.headers["Referrer-Policy"] = "strict-origin-when-cross-origin"
        response.headers["Permissions-Policy"] = "camera=(), microphone=(), geolocation=()"
        if _is_deployed or ENVIRONMENT != "development":
            response.headers["Strict-Transport-Security"] = "max-age=31536000; includeSubDomains"
        # Log server errors to DB for admin dashboard
        if response.status_code >= 500:
            _log_error_to_db(request.method, str(request.url.path), response.status_code, "Server error")
        return response
    except Exception as exc:
        logger.exception("Middleware error on %s %s", request.method, request.url.path)
        _log_error_to_db(request.method, str(request.url.path), 500, str(exc))
        detail = str(exc) if ENVIRONMENT == "development" else "Internal server error"
        return JSONResponse(status_code=500, content={"detail": detail})

# ============================================================
# GLOBAL ERROR HANDLER
# ============================================================

@app.exception_handler(Exception)
async def global_exception_handler(request: Request, exc: Exception):
    logger.exception("Unhandled error on %s %s", request.method, request.url.path)
    _log_error_to_db(request.method, str(request.url.path), 500, str(exc))
    detail = str(exc) if ENVIRONMENT == "development" else "Internal server error."
    return JSONResponse(status_code=500, content={"detail": detail})

# ============================================================
# DATABASE
# ============================================================

def get_db() -> sqlite3.Connection:
    conn = sqlite3.connect(DB_FILE)
    conn.row_factory = sqlite3.Row
    conn.execute("PRAGMA journal_mode=WAL")
    conn.execute("PRAGMA foreign_keys=ON")
    return conn


@contextmanager
def safe_db():
    """Context manager for safe database access with auto-rollback and cleanup.

    Usage:
        with safe_db() as conn:
            conn.execute("INSERT INTO ...")
            # auto-commits on success, auto-rollbacks on error, always closes
    """
    conn = get_db()
    try:
        yield conn
        conn.commit()
    except Exception:
        conn.rollback()
        raise
    finally:
        conn.close()


# ── League Tier Classification ──────────────────────────────────────────────
LEAGUE_TIERS = {
    # Tier 1 — Major Junior (CHL)
    "OHL": "Tier1", "WHL": "Tier1", "QMJHL": "Tier1", "CHL": "Tier1",
    # Tier 2 — Junior A
    "OJHL": "Tier2", "BCHL": "Tier2", "USHL": "Tier2", "AJHL": "Tier2",
    "CCHL": "Tier2", "NOJHL": "Tier2", "SJHL": "Tier2", "MHL": "Tier2",
    "MJHL": "Tier2",
    # Tier 3 — Junior B / Tier 2 Jr
    "GOHL": "Tier3", "NAHL": "Tier3", "GMHL": "Tier3", "PJHL": "Tier3",
    "WOAA": "Tier3", "SIJHL": "Tier3",
    # College / University
    "NCAA": "NCAA", "NCAA DI": "NCAA", "NCAA DIII": "NCAA_D3",
    "USports": "USports", "U SPORTS": "USports",
    # International
    "KHL": "Pro", "SHL": "Pro", "AHL": "Pro", "ECHL": "Pro",
    "NHL": "Pro", "Liiga": "Pro", "NLA": "Pro", "DEL": "Pro",
}


def _get_league_tier(league_name: str | None) -> str | None:
    """Classify league into tier."""
    if not league_name:
        return None
    # Exact match first
    if league_name in LEAGUE_TIERS:
        return LEAGUE_TIERS[league_name]
    # Case-insensitive
    upper = league_name.upper().strip()
    for key, tier in LEAGUE_TIERS.items():
        if key.upper() == upper:
            return tier
    return "Unknown"


def _get_age_group(birth_year: int | None) -> str | None:
    """Classify player into age group based on birth year."""
    if not birth_year:
        return None
    current_year = datetime.now().year
    age = current_year - birth_year
    if age <= 16:
        return "U16"
    elif age <= 18:
        return "U18"
    elif age <= 20:
        return "U20"
    else:
        return "Over20"


def _populate_derived_player_fields(conn):
    """Auto-populate birth_year, age_group, draft_eligible_year, league_tier
    for any players that have NULL values but have the source data."""
    # Birth year / age group / draft eligible year from dob
    rows = conn.execute(
        "SELECT id, dob, current_league, birth_year, league_tier FROM players"
    ).fetchall()
    updated = 0
    for row in rows:
        pid = row[0]
        dob = row[1]
        league = row[2]
        existing_by = row[3]
        existing_lt = row[4]

        updates = {}
        params = []

        # Derive birth_year, age_group, draft_eligible_year from dob
        if dob and not existing_by:
            try:
                by = int(dob[:4])
                updates["birth_year"] = by
                updates["age_group"] = _get_age_group(by)
                updates["draft_eligible_year"] = by + 18
            except (ValueError, IndexError):
                pass

        # Derive league_tier from current_league
        if league and not existing_lt:
            updates["league_tier"] = _get_league_tier(league)

        if updates:
            set_clauses = ", ".join(f"{k} = ?" for k in updates)
            params = list(updates.values()) + [pid]
            conn.execute(f"UPDATE players SET {set_clauses} WHERE id = ?", params)
            updated += 1

    if updated:
        conn.commit()
        logger.info("Populated derived fields for %d players", updated)


# ── Template Category Mapping ──────────────────────────────────────────────
TEMPLATE_CATEGORIES = {
    "pro_skater":          ("Player Analytics", "Performance Metrics"),
    "season_intelligence": ("Player Analytics", "Performance Metrics"),
    "family_card":         ("Player Analytics", "Presentation"),
    "agent_pack":          ("Player Analytics", "Presentation"),
    "unified_prospect":    ("Player Analytics", "Advanced Stats"),
    "game_decision":       ("Player Analytics", "Advanced Stats"),
    "development_roadmap": ("Player Analytics", "Projections & Development"),
    "draft_comparative":   ("Player Analytics", "Projections & Development"),
    "season_progress":     ("Player Analytics", "Performance Metrics"),
    "team_identity":       ("Team Analytics", "System Analysis"),
    "operations":          ("Team Analytics", "System Analysis"),
    "practice_plan":       ("Team Analytics", "System Analysis"),
    "line_chemistry":      ("Team Analytics", "Line Optimization"),
    "st_optimization":     ("Team Analytics", "Special Teams"),
    "goalie_tandem":       ("Team Analytics", "Special Teams"),
    "goalie":              ("Competitive Intelligence", "Market & Acquisitions"),
    "opponent_gameplan":   ("Competitive Intelligence", "Opponent Analysis"),
    "playoff_series":      ("Competitive Intelligence", "Opponent Analysis"),
    "trade_target":        ("Competitive Intelligence", "Market & Acquisitions"),
    # Priority 2 reports
    "indices_dashboard":   ("Player Analytics", "Advanced Stats"),
    "player_projection":   ("Player Analytics", "Projections & Development"),
    "league_benchmarks":   ("Competitive Intelligence", "League Benchmarks"),
    "season_projection":   ("Competitive Intelligence", "League Benchmarks"),
    "free_agent_market":   ("Competitive Intelligence", "Market & Acquisitions"),
}


def _seed_template_categories(conn):
    """Set category/subcategory on report_templates if not already set."""
    for rtype, (cat, subcat) in TEMPLATE_CATEGORIES.items():
        conn.execute(
            "UPDATE report_templates SET category = ?, subcategory = ? WHERE report_type = ? AND (category IS NULL OR category = '')",
            (cat, subcat, rtype),
        )
    conn.commit()


def init_db():
    """Create all tables if they don't exist."""
    conn = get_db()
    c = conn.cursor()

    c.execute("""
        CREATE TABLE IF NOT EXISTS organizations (
            id TEXT PRIMARY KEY,
            name TEXT NOT NULL,
            org_type TEXT DEFAULT 'team',
            created_at TEXT DEFAULT CURRENT_TIMESTAMP
        )
    """)

    c.execute("""
        CREATE TABLE IF NOT EXISTS users (
            id TEXT PRIMARY KEY,
            org_id TEXT NOT NULL,
            email TEXT UNIQUE NOT NULL,
            password_hash TEXT NOT NULL,
            first_name TEXT,
            last_name TEXT,
            role TEXT NOT NULL DEFAULT 'scout',
            created_at TEXT DEFAULT CURRENT_TIMESTAMP,
            FOREIGN KEY (org_id) REFERENCES organizations(id)
        )
    """)

    c.execute("""
        CREATE TABLE IF NOT EXISTS players (
            id TEXT PRIMARY KEY,
            org_id TEXT NOT NULL,
            first_name TEXT NOT NULL,
            last_name TEXT NOT NULL,
            dob TEXT,
            position TEXT NOT NULL,
            shoots TEXT,
            height_cm INTEGER,
            weight_kg INTEGER,
            current_team TEXT,
            current_league TEXT,
            passports TEXT DEFAULT '[]',
            notes TEXT,
            tags TEXT DEFAULT '[]',
            archetype TEXT,
            image_url TEXT,
            created_at TEXT DEFAULT CURRENT_TIMESTAMP,
            updated_at TEXT DEFAULT CURRENT_TIMESTAMP,
            FOREIGN KEY (org_id) REFERENCES organizations(id)
        )
    """)

    c.execute("""
        CREATE TABLE IF NOT EXISTS player_stats (
            id TEXT PRIMARY KEY,
            player_id TEXT NOT NULL,
            game_id TEXT,
            season TEXT,
            stat_type TEXT DEFAULT 'season',
            gp INTEGER DEFAULT 0,
            g INTEGER DEFAULT 0,
            a INTEGER DEFAULT 0,
            p INTEGER DEFAULT 0,
            plus_minus INTEGER DEFAULT 0,
            pim INTEGER DEFAULT 0,
            toi_seconds INTEGER DEFAULT 0,
            pp_toi_seconds INTEGER DEFAULT 0,
            pk_toi_seconds INTEGER DEFAULT 0,
            shots INTEGER DEFAULT 0,
            sog INTEGER DEFAULT 0,
            shooting_pct REAL,
            microstats TEXT,
            created_at TEXT DEFAULT CURRENT_TIMESTAMP,
            FOREIGN KEY (player_id) REFERENCES players(id)
        )
    """)

    c.execute("""
        CREATE TABLE IF NOT EXISTS reports (
            id TEXT PRIMARY KEY,
            org_id TEXT NOT NULL,
            player_id TEXT NOT NULL,
            template_id TEXT,
            report_type TEXT NOT NULL,
            title TEXT,
            status TEXT DEFAULT 'pending',
            output_json TEXT,
            output_text TEXT,
            input_data TEXT,
            error_message TEXT,
            generated_at TEXT,
            llm_model TEXT,
            llm_tokens INTEGER DEFAULT 0,
            generation_time_ms INTEGER,
            created_by TEXT,
            created_at TEXT DEFAULT CURRENT_TIMESTAMP,
            FOREIGN KEY (org_id) REFERENCES organizations(id),
            FOREIGN KEY (player_id) REFERENCES players(id)
        )
    """)

    c.execute("""
        CREATE TABLE IF NOT EXISTS report_templates (
            id TEXT PRIMARY KEY,
            org_id TEXT,
            template_name TEXT NOT NULL,
            report_type TEXT NOT NULL,
            prompt_text TEXT,
            data_schema TEXT DEFAULT '{}',
            is_global INTEGER DEFAULT 1,
            version INTEGER DEFAULT 1,
            created_at TEXT DEFAULT CURRENT_TIMESTAMP
        )
    """)

    c.execute("""
        CREATE TABLE IF NOT EXISTS teams (
            id TEXT PRIMARY KEY,
            org_id TEXT NOT NULL,
            name TEXT NOT NULL,
            league TEXT,
            city TEXT,
            abbreviation TEXT,
            identity TEXT DEFAULT '{}',
            created_at TEXT DEFAULT CURRENT_TIMESTAMP,
            FOREIGN KEY (org_id) REFERENCES organizations(id)
        )
    """)

    c.execute("""
        CREATE TABLE IF NOT EXISTS scout_notes (
            id TEXT PRIMARY KEY,
            org_id TEXT NOT NULL,
            player_id TEXT NOT NULL,
            scout_id TEXT NOT NULL,
            note_text TEXT NOT NULL,
            note_type TEXT DEFAULT 'general',
            tags TEXT DEFAULT '[]',
            is_private INTEGER DEFAULT 0,
            created_at TEXT DEFAULT CURRENT_TIMESTAMP,
            updated_at TEXT DEFAULT CURRENT_TIMESTAMP,
            FOREIGN KEY (org_id) REFERENCES organizations(id),
            FOREIGN KEY (player_id) REFERENCES players(id),
            FOREIGN KEY (scout_id) REFERENCES users(id)
        )
    """)

    c.execute("""
        CREATE TABLE IF NOT EXISTS import_jobs (
            id TEXT PRIMARY KEY,
            org_id TEXT NOT NULL,
            uploaded_by TEXT NOT NULL,
            filename TEXT,
            total_rows INTEGER DEFAULT 0,
            new_players INTEGER DEFAULT 0,
            duplicates_found INTEGER DEFAULT 0,
            errors_found INTEGER DEFAULT 0,
            status TEXT DEFAULT 'pending',
            preview_data TEXT DEFAULT '[]',
            duplicate_data TEXT DEFAULT '[]',
            created_at TEXT DEFAULT CURRENT_TIMESTAMP,
            FOREIGN KEY (org_id) REFERENCES organizations(id),
            FOREIGN KEY (uploaded_by) REFERENCES users(id)
        )
    """)

    # ── Hockey Operating System Tables ──────────────────────────

    # Hockey terminology glossary — the language ProspectX speaks
    c.execute("""
        CREATE TABLE IF NOT EXISTS hockey_terms (
            id TEXT PRIMARY KEY,
            term TEXT NOT NULL UNIQUE,
            category TEXT NOT NULL,
            definition TEXT NOT NULL,
            aliases TEXT DEFAULT '[]',
            usage_context TEXT,
            created_at TEXT DEFAULT CURRENT_TIMESTAMP
        )
    """)

    # Systems library — tactical structures teams can run
    c.execute("""
        CREATE TABLE IF NOT EXISTS systems_library (
            id TEXT PRIMARY KEY,
            system_type TEXT NOT NULL,
            code TEXT NOT NULL UNIQUE,
            name TEXT NOT NULL,
            description TEXT,
            strengths TEXT,
            weaknesses TEXT,
            ideal_personnel TEXT,
            created_at TEXT DEFAULT CURRENT_TIMESTAMP
        )
    """)

    # Team systems — which systems a team actually runs
    c.execute("""
        CREATE TABLE IF NOT EXISTS team_systems (
            id TEXT PRIMARY KEY,
            org_id TEXT NOT NULL,
            team_id TEXT,
            team_name TEXT,
            season TEXT,
            forecheck TEXT,
            dz_structure TEXT,
            oz_setup TEXT,
            pp_formation TEXT,
            pk_formation TEXT,
            neutral_zone TEXT,
            breakout TEXT,
            identity_tags TEXT DEFAULT '[]',
            pace TEXT DEFAULT '',
            physicality TEXT DEFAULT '',
            offensive_style TEXT DEFAULT '',
            notes TEXT,
            created_at TEXT DEFAULT CURRENT_TIMESTAMP,
            updated_at TEXT DEFAULT CURRENT_TIMESTAMP,
            FOREIGN KEY (org_id) REFERENCES organizations(id)
        )
    """)

    # Migrate existing team_systems table — add Team Style columns if missing
    try:
        c.execute("ALTER TABLE team_systems ADD COLUMN pace TEXT DEFAULT ''")
    except sqlite3.OperationalError:
        pass  # Column already exists
    try:
        c.execute("ALTER TABLE team_systems ADD COLUMN physicality TEXT DEFAULT ''")
    except sqlite3.OperationalError:
        pass
    try:
        c.execute("ALTER TABLE team_systems ADD COLUMN offensive_style TEXT DEFAULT ''")
    except sqlite3.OperationalError:
        pass

    # ── Reports table migration: make player_id nullable, add team_name column ──
    # SQLite doesn't support ALTER COLUMN, so we check and recreate if needed
    cursor_info = c.execute("PRAGMA table_info(reports)").fetchall()
    report_cols = {col[1] for col in cursor_info}
    if "team_name" not in report_cols:
        logger.info("Migrating reports table: adding team_name column, making player_id nullable...")
        c.execute("""CREATE TABLE IF NOT EXISTS reports_new (
            id TEXT PRIMARY KEY,
            org_id TEXT NOT NULL,
            player_id TEXT,
            team_name TEXT,
            template_id TEXT,
            report_type TEXT NOT NULL,
            title TEXT,
            status TEXT DEFAULT 'pending',
            output_json TEXT,
            output_text TEXT,
            input_data TEXT,
            error_message TEXT,
            generated_at TEXT,
            llm_model TEXT,
            llm_tokens INTEGER DEFAULT 0,
            generation_time_ms INTEGER,
            created_by TEXT,
            created_at TEXT DEFAULT CURRENT_TIMESTAMP,
            FOREIGN KEY (org_id) REFERENCES organizations(id)
        )""")
        c.execute("""INSERT INTO reports_new (id, org_id, player_id, template_id, report_type, title, status, output_json, output_text, input_data, error_message, generated_at, llm_model, llm_tokens, generation_time_ms, created_by, created_at)
            SELECT id, org_id, player_id, template_id, report_type, title, status, output_json, output_text, input_data, error_message, generated_at, llm_model, llm_tokens, generation_time_ms, created_by, created_at FROM reports""")
        c.execute("DROP TABLE reports")
        c.execute("ALTER TABLE reports_new RENAME TO reports")
        logger.info("Reports table migration complete.")

    # Player archetypes — computed or manually assigned role classifications
    c.execute("""
        CREATE TABLE IF NOT EXISTS player_archetypes (
            id TEXT PRIMARY KEY,
            player_id TEXT NOT NULL,
            archetype TEXT NOT NULL,
            confidence REAL DEFAULT 0.0,
            indices TEXT DEFAULT '{}',
            assigned_by TEXT DEFAULT 'manual',
            notes TEXT,
            created_at TEXT DEFAULT CURRENT_TIMESTAMP,
            FOREIGN KEY (player_id) REFERENCES players(id)
        )
    """)

    # System adherence — how well a player fits or performs in a system
    c.execute("""
        CREATE TABLE IF NOT EXISTS system_adherence (
            id TEXT PRIMARY KEY,
            player_id TEXT NOT NULL,
            system_code TEXT NOT NULL,
            adherence_score REAL,
            fit_rating TEXT,
            notes TEXT,
            evaluated_at TEXT DEFAULT CURRENT_TIMESTAMP,
            FOREIGN KEY (player_id) REFERENCES players(id)
        )
    """)

    # Team evaluation criteria — what a team values for each position
    c.execute("""
        CREATE TABLE IF NOT EXISTS team_eval_criteria (
            id TEXT PRIMARY KEY,
            org_id TEXT NOT NULL,
            position TEXT NOT NULL,
            criteria_name TEXT NOT NULL,
            weight REAL DEFAULT 1.0,
            description TEXT,
            created_at TEXT DEFAULT CURRENT_TIMESTAMP,
            FOREIGN KEY (org_id) REFERENCES organizations(id)
        )
    """)

    c.execute("""
        CREATE TABLE IF NOT EXISTS leagues (
            id TEXT PRIMARY KEY,
            name TEXT NOT NULL UNIQUE,
            abbreviation TEXT NOT NULL UNIQUE,
            country TEXT DEFAULT 'Canada',
            level TEXT DEFAULT 'junior',
            sort_order INTEGER DEFAULT 100,
            created_at TEXT DEFAULT CURRENT_TIMESTAMP
        )
    """)

    # ── InStat Analytics tables ──────────────────────────────

    c.execute("""
        CREATE TABLE IF NOT EXISTS goalie_stats (
            id TEXT PRIMARY KEY,
            player_id TEXT NOT NULL,
            org_id TEXT NOT NULL,
            season TEXT,
            stat_type TEXT DEFAULT 'season',
            gp INTEGER DEFAULT 0,
            toi_seconds INTEGER DEFAULT 0,
            ga REAL DEFAULT 0,
            sa REAL DEFAULT 0,
            sv REAL DEFAULT 0,
            sv_pct TEXT,
            gaa REAL,
            extended_stats TEXT,
            data_source TEXT DEFAULT 'manual',
            created_at TEXT DEFAULT CURRENT_TIMESTAMP,
            FOREIGN KEY (player_id) REFERENCES players(id)
        )
    """)

    c.execute("""
        CREATE TABLE IF NOT EXISTS team_stats (
            id TEXT PRIMARY KEY,
            org_id TEXT NOT NULL,
            team_name TEXT NOT NULL,
            league TEXT,
            season TEXT,
            stat_type TEXT DEFAULT 'season',
            extended_stats TEXT,
            data_source TEXT DEFAULT 'instat',
            created_at TEXT DEFAULT CURRENT_TIMESTAMP
        )
    """)

    c.execute("""
        CREATE TABLE IF NOT EXISTS line_combinations (
            id TEXT PRIMARY KEY,
            org_id TEXT NOT NULL,
            team_name TEXT NOT NULL,
            season TEXT,
            line_type TEXT NOT NULL,
            player_names TEXT NOT NULL,
            player_refs TEXT,
            plus_minus TEXT,
            shifts INTEGER DEFAULT 0,
            toi_seconds INTEGER DEFAULT 0,
            goals_for REAL DEFAULT 0,
            goals_against REAL DEFAULT 0,
            extended_stats TEXT,
            data_source TEXT DEFAULT 'instat',
            created_at TEXT DEFAULT CURRENT_TIMESTAMP
        )
    """)

    # ── Player Intelligence table ─────────────────────────────
    c.execute("""
        CREATE TABLE IF NOT EXISTS player_intelligence (
            id TEXT PRIMARY KEY,
            player_id TEXT NOT NULL,
            org_id TEXT NOT NULL,
            archetype TEXT,
            archetype_confidence REAL,
            overall_grade TEXT,
            offensive_grade TEXT,
            defensive_grade TEXT,
            skating_grade TEXT,
            hockey_iq_grade TEXT,
            compete_grade TEXT,
            summary TEXT,
            strengths TEXT,
            development_areas TEXT,
            comparable_players TEXT,
            stat_signature TEXT,
            tags TEXT DEFAULT '[]',
            projection TEXT,
            data_sources_used TEXT,
            trigger TEXT,
            version INTEGER DEFAULT 1,
            created_at TEXT DEFAULT CURRENT_TIMESTAMP,
            FOREIGN KEY (player_id) REFERENCES players(id)
        )
    """)

    conn.execute("""
        CREATE TABLE IF NOT EXISTS team_intelligence (
            id TEXT PRIMARY KEY,
            team_name TEXT NOT NULL,
            org_id TEXT NOT NULL,
            playing_style TEXT,
            system_summary TEXT,
            identity TEXT,
            strengths TEXT DEFAULT '[]',
            vulnerabilities TEXT DEFAULT '[]',
            key_personnel TEXT DEFAULT '[]',
            special_teams_identity TEXT,
            player_archetype_fit TEXT,
            comparable_teams TEXT DEFAULT '[]',
            tags TEXT DEFAULT '[]',
            trigger TEXT,
            version INTEGER DEFAULT 1,
            created_at TEXT DEFAULT CURRENT_TIMESTAMP
        )
    """)

    conn.commit()

    # ── Migrations for existing databases ───────────────────
    # Add image_url column if it doesn't exist
    cols = [col[1] for col in conn.execute("PRAGMA table_info(players)").fetchall()]
    if "image_url" not in cols:
        conn.execute("ALTER TABLE players ADD COLUMN image_url TEXT")
        conn.commit()
        logger.info("Migration: added image_url column to players table")

    # Add extended_stats and data_source columns to player_stats
    ps_cols = [col[1] for col in conn.execute("PRAGMA table_info(player_stats)").fetchall()]
    if "extended_stats" not in ps_cols:
        conn.execute("ALTER TABLE player_stats ADD COLUMN extended_stats TEXT")
        conn.commit()
        logger.info("Migration: added extended_stats column to player_stats")
    if "data_source" not in ps_cols:
        conn.execute("ALTER TABLE player_stats ADD COLUMN data_source TEXT DEFAULT 'manual'")
        conn.commit()
        logger.info("Migration: added data_source column to player_stats")
    if "team_name" not in ps_cols:
        conn.execute("ALTER TABLE player_stats ADD COLUMN team_name TEXT")
        conn.commit()
        logger.info("Migration: added team_name column to player_stats")
    if "notes" not in ps_cols:
        conn.execute("ALTER TABLE player_stats ADD COLUMN notes TEXT")
        conn.commit()
        logger.info("Migration: added notes column to player_stats")

    # Add logo_url column to teams
    teams_cols = [col[1] for col in conn.execute("PRAGMA table_info(teams)").fetchall()]
    if "logo_url" not in teams_cols:
        conn.execute("ALTER TABLE teams ADD COLUMN logo_url TEXT")
        conn.commit()
        logger.info("Migration: added logo_url column to teams")

    # Add intelligence_version column to players
    player_cols = [col[1] for col in conn.execute("PRAGMA table_info(players)").fetchall()]
    if "intelligence_version" not in player_cols:
        conn.execute("ALTER TABLE players ADD COLUMN intelligence_version INTEGER DEFAULT 0")
        conn.commit()
        logger.info("Migration: added intelligence_version column to players")

    # ── Data Compartmentalization: birth_year, age_group, draft_eligible_year, league_tier ──
    player_cols = [col[1] for col in conn.execute("PRAGMA table_info(players)").fetchall()]
    new_player_cols = {
        "birth_year": "INTEGER",
        "age_group": "TEXT",
        "draft_eligible_year": "INTEGER",
        "league_tier": "TEXT",
        "commitment_status": "TEXT DEFAULT 'Uncommitted'",
    }
    for col_name, col_type in new_player_cols.items():
        if col_name not in player_cols:
            conn.execute(f"ALTER TABLE players ADD COLUMN {col_name} {col_type}")
            conn.commit()
            logger.info("Migration: added %s column to players", col_name)

    # Auto-populate birth_year / age_group / draft_eligible_year from dob
    _populate_derived_player_fields(conn)

    # Add category / subcategory columns to report_templates
    tmpl_cols = [col[1] for col in conn.execute("PRAGMA table_info(report_templates)").fetchall()]
    if "category" not in tmpl_cols:
        conn.execute("ALTER TABLE report_templates ADD COLUMN category TEXT")
        conn.commit()
        logger.info("Migration: added category column to report_templates")
    if "subcategory" not in tmpl_cols:
        conn.execute("ALTER TABLE report_templates ADD COLUMN subcategory TEXT")
        conn.commit()
        logger.info("Migration: added subcategory column to report_templates")

    # Seed template categories
    _seed_template_categories(conn)

    # Add hockeytech_id + hockeytech_league columns to players
    player_cols = [col[1] for col in conn.execute("PRAGMA table_info(players)").fetchall()]
    for col_name, col_type in {"hockeytech_id": "INTEGER", "hockeytech_league": "TEXT"}.items():
        if col_name not in player_cols:
            conn.execute(f"ALTER TABLE players ADD COLUMN {col_name} {col_type}")
            conn.commit()
            logger.info("Migration: added %s column to players", col_name)

    # Add hockeytech_team_id + hockeytech_league columns to teams table
    team_cols = [col[1] for col in conn.execute("PRAGMA table_info(teams)").fetchall()]
    for col_name, col_type in {"hockeytech_team_id": "INTEGER", "hockeytech_league": "TEXT"}.items():
        if col_name not in team_cols:
            conn.execute(f"ALTER TABLE teams ADD COLUMN {col_name} {col_type}")
            conn.commit()
            logger.info("Migration: added %s column to teams", col_name)

    # Add line_label, line_order, updated_at columns to line_combinations
    lc_cols = [col[1] for col in conn.execute("PRAGMA table_info(line_combinations)").fetchall()]
    for col_name, col_type in {"line_label": "TEXT", "line_order": "INTEGER DEFAULT 0", "updated_at": "TEXT"}.items():
        bare_name = col_name.split()[0]  # handle "INTEGER DEFAULT 0"
        if bare_name not in lc_cols:
            conn.execute(f"ALTER TABLE line_combinations ADD COLUMN {col_name} {col_type}")
            conn.commit()
            logger.info("Migration: added %s column to line_combinations", bare_name)

    # ── Bench Talk Tables ──────────────────────────────────────────
    c.execute("""
        CREATE TABLE IF NOT EXISTS bench_talk_conversations (
            id TEXT PRIMARY KEY,
            user_id TEXT NOT NULL,
            org_id TEXT NOT NULL,
            title TEXT DEFAULT 'New Conversation',
            created_at TEXT DEFAULT CURRENT_TIMESTAMP,
            updated_at TEXT DEFAULT CURRENT_TIMESTAMP,
            FOREIGN KEY (user_id) REFERENCES users(id)
        )
    """)

    c.execute("""
        CREATE TABLE IF NOT EXISTS bench_talk_messages (
            id TEXT PRIMARY KEY,
            conversation_id TEXT NOT NULL,
            role TEXT NOT NULL,
            content TEXT NOT NULL,
            metadata TEXT,
            tokens_used INTEGER DEFAULT 0,
            created_at TEXT DEFAULT CURRENT_TIMESTAMP,
            FOREIGN KEY (conversation_id) REFERENCES bench_talk_conversations(id)
        )
    """)

    c.execute("""
        CREATE TABLE IF NOT EXISTS bench_talk_feedback (
            id TEXT PRIMARY KEY,
            message_id TEXT NOT NULL,
            user_id TEXT NOT NULL,
            org_id TEXT NOT NULL,
            rating TEXT NOT NULL,
            feedback_text TEXT,
            created_at TEXT DEFAULT CURRENT_TIMESTAMP,
            FOREIGN KEY (message_id) REFERENCES bench_talk_messages(id),
            FOREIGN KEY (user_id) REFERENCES users(id)
        )
    """)

    # ── Subscription Usage Log ─────────────────────────────────────
    c.execute("""
        CREATE TABLE IF NOT EXISTS subscription_usage_log (
            id TEXT PRIMARY KEY,
            user_id TEXT NOT NULL,
            org_id TEXT NOT NULL,
            action_type TEXT NOT NULL,
            resource_id TEXT,
            created_at TEXT DEFAULT CURRENT_TIMESTAMP,
            FOREIGN KEY (user_id) REFERENCES users(id)
        )
    """)

    # ── Drills Library ─────────────────────────────────────────
    c.execute("""
        CREATE TABLE IF NOT EXISTS drills (
            id TEXT PRIMARY KEY,
            org_id TEXT,
            name TEXT NOT NULL,
            category TEXT NOT NULL,
            description TEXT NOT NULL,
            coaching_points TEXT,
            setup TEXT,
            duration_minutes INTEGER DEFAULT 10,
            players_needed INTEGER DEFAULT 0,
            ice_surface TEXT DEFAULT 'full',
            equipment TEXT,
            age_levels TEXT DEFAULT '[]',
            tags TEXT DEFAULT '[]',
            diagram_url TEXT,
            skill_focus TEXT,
            intensity TEXT DEFAULT 'medium',
            concept_id TEXT,
            created_at TEXT DEFAULT CURRENT_TIMESTAMP
        )
    """)

    # ── Practice Plans ─────────────────────────────────────────
    c.execute("""
        CREATE TABLE IF NOT EXISTS practice_plans (
            id TEXT PRIMARY KEY,
            org_id TEXT NOT NULL,
            user_id TEXT NOT NULL,
            team_name TEXT,
            title TEXT NOT NULL,
            age_level TEXT,
            duration_minutes INTEGER DEFAULT 90,
            focus_areas TEXT DEFAULT '[]',
            plan_data TEXT DEFAULT '{}',
            notes TEXT,
            status TEXT DEFAULT 'draft',
            created_at TEXT DEFAULT CURRENT_TIMESTAMP,
            updated_at TEXT DEFAULT CURRENT_TIMESTAMP,
            FOREIGN KEY (org_id) REFERENCES organizations(id),
            FOREIGN KEY (user_id) REFERENCES users(id)
        )
    """)

    # ── Practice Plan Drills (junction) ────────────────────────
    c.execute("""
        CREATE TABLE IF NOT EXISTS practice_plan_drills (
            id TEXT PRIMARY KEY,
            practice_plan_id TEXT NOT NULL,
            drill_id TEXT NOT NULL,
            phase TEXT NOT NULL,
            sequence_order INTEGER DEFAULT 0,
            duration_minutes INTEGER DEFAULT 10,
            coaching_notes TEXT,
            FOREIGN KEY (practice_plan_id) REFERENCES practice_plans(id) ON DELETE CASCADE,
            FOREIGN KEY (drill_id) REFERENCES drills(id)
        )
    """)

    # ── Saved Searches ──────────────────────────────────────────
    c.execute("""
        CREATE TABLE IF NOT EXISTS saved_searches (
            id TEXT PRIMARY KEY,
            org_id TEXT NOT NULL,
            user_id TEXT NOT NULL,
            name TEXT NOT NULL,
            filters TEXT NOT NULL,
            created_at TEXT DEFAULT CURRENT_TIMESTAMP,
            FOREIGN KEY (user_id) REFERENCES users(id)
        )
    """)

    # ── Usage Tracking (per-resource monthly counters) ──────────
    c.execute("""
        CREATE TABLE IF NOT EXISTS usage_tracking (
            id TEXT PRIMARY KEY,
            user_id TEXT NOT NULL,
            month TEXT NOT NULL,
            bench_talks_count INTEGER DEFAULT 0,
            reports_count INTEGER DEFAULT 0,
            practice_plans_count INTEGER DEFAULT 0,
            uploads_count INTEGER DEFAULT 0,
            created_at TEXT DEFAULT CURRENT_TIMESTAMP,
            updated_at TEXT DEFAULT CURRENT_TIMESTAMP,
            UNIQUE(user_id, month),
            FOREIGN KEY (user_id) REFERENCES users(id)
        )
    """)

    # ── Admin Error Log ──────────────────────────────────────────
    c.execute("""
        CREATE TABLE IF NOT EXISTS admin_error_log (
            id TEXT PRIMARY KEY,
            request_method TEXT,
            request_path TEXT,
            status_code INTEGER,
            error_message TEXT,
            user_id TEXT,
            org_id TEXT,
            created_at TEXT DEFAULT CURRENT_TIMESTAMP
        )
    """)

    # ── Games (reference table for individual games) ────────────
    c.execute("""
        CREATE TABLE IF NOT EXISTS games (
            id TEXT PRIMARY KEY,
            league TEXT NOT NULL,
            season TEXT,
            ht_game_id INTEGER,
            game_date TEXT NOT NULL,
            home_team TEXT,
            away_team TEXT,
            home_score INTEGER,
            away_score INTEGER,
            status TEXT DEFAULT 'final',
            venue TEXT,
            data_source TEXT DEFAULT 'hockeytech',
            created_at TEXT DEFAULT CURRENT_TIMESTAMP,
            UNIQUE(league, ht_game_id)
        )
    """)

    # ── Player Stats History (append-only season snapshots) ───
    c.execute("""
        CREATE TABLE IF NOT EXISTS player_stats_history (
            id TEXT PRIMARY KEY,
            player_id TEXT NOT NULL,
            season TEXT,
            date_recorded TEXT NOT NULL,
            gp INTEGER DEFAULT 0,
            g INTEGER DEFAULT 0,
            a INTEGER DEFAULT 0,
            p INTEGER DEFAULT 0,
            plus_minus INTEGER DEFAULT 0,
            pim INTEGER DEFAULT 0,
            ppg INTEGER DEFAULT 0,
            ppa INTEGER DEFAULT 0,
            shg INTEGER DEFAULT 0,
            gwg INTEGER DEFAULT 0,
            shots INTEGER DEFAULT 0,
            shooting_pct REAL,
            data_source TEXT DEFAULT 'hockeytech',
            league TEXT,
            team_name TEXT,
            synced_at TEXT DEFAULT CURRENT_TIMESTAMP,
            FOREIGN KEY (player_id) REFERENCES players(id)
        )
    """)

    # ── Player Game Stats (per-game detail) ───────────────────
    c.execute("""
        CREATE TABLE IF NOT EXISTS player_game_stats (
            id TEXT PRIMARY KEY,
            player_id TEXT NOT NULL,
            game_id TEXT,
            ht_game_id INTEGER,
            game_date TEXT,
            opponent TEXT,
            home_away TEXT,
            goals INTEGER DEFAULT 0,
            assists INTEGER DEFAULT 0,
            points INTEGER DEFAULT 0,
            plus_minus INTEGER DEFAULT 0,
            pim INTEGER DEFAULT 0,
            shots INTEGER DEFAULT 0,
            ppg INTEGER DEFAULT 0,
            shg INTEGER DEFAULT 0,
            gwg INTEGER DEFAULT 0,
            toi_seconds INTEGER DEFAULT 0,
            season TEXT,
            league TEXT,
            data_source TEXT DEFAULT 'hockeytech',
            created_at TEXT DEFAULT CURRENT_TIMESTAMP,
            FOREIGN KEY (player_id) REFERENCES players(id),
            FOREIGN KEY (game_id) REFERENCES games(id)
        )
    """)

    # ── Migration: hockey_role on users ─────────────────────────
    user_cols = [col[1] for col in conn.execute("PRAGMA table_info(users)").fetchall()]
    if "hockey_role" not in user_cols:
        conn.execute("ALTER TABLE users ADD COLUMN hockey_role TEXT DEFAULT 'scout'")
        conn.commit()
        logger.info("Migration: added hockey_role column to users")

    # ── Migration: subscription columns on users ─────────────────
    user_cols = [col[1] for col in conn.execute("PRAGMA table_info(users)").fetchall()]
    sub_cols = {
        "subscription_tier": "TEXT DEFAULT 'rookie'",
        "subscription_started_at": "TEXT",
        "monthly_reports_used": "INTEGER DEFAULT 0",
        "monthly_bench_talks_used": "INTEGER DEFAULT 0",
        "usage_reset_at": "TEXT",
        "max_seats": "INTEGER DEFAULT 1",
    }
    for col_name, col_type in sub_cols.items():
        if col_name not in user_cols:
            conn.execute(f"ALTER TABLE users ADD COLUMN {col_name} {col_type}")
            conn.commit()
            logger.info("Migration: added %s column to users", col_name)

    # ── Migration: rename PXI chat tables to Bench Talk ──────────
    existing_tables = [r[0] for r in conn.execute("SELECT name FROM sqlite_master WHERE type='table'").fetchall()]
    rename_map = {
        "chat_conversations": "bench_talk_conversations",
        "chat_messages": "bench_talk_messages",
        "pxi_feedback": "bench_talk_feedback",
    }
    for old_name, new_name in rename_map.items():
        if old_name in existing_tables and new_name not in existing_tables:
            conn.execute(f"ALTER TABLE {old_name} RENAME TO {new_name}")
            conn.commit()
            logger.info("Migration: renamed %s → %s", old_name, new_name)

    # Create indexes for fast queries
    for idx_sql in [
        "CREATE INDEX IF NOT EXISTS idx_players_birth_year ON players(birth_year)",
        "CREATE INDEX IF NOT EXISTS idx_players_age_group ON players(age_group)",
        "CREATE INDEX IF NOT EXISTS idx_players_league_tier ON players(league_tier)",
        "CREATE INDEX IF NOT EXISTS idx_players_team_league ON players(current_team, current_league)",
        "CREATE INDEX IF NOT EXISTS idx_players_position ON players(position)",
        "CREATE INDEX IF NOT EXISTS idx_players_hockeytech_id ON players(hockeytech_id)",
        "CREATE INDEX IF NOT EXISTS idx_lines_team_type ON line_combinations(team_name, line_type)",
        "CREATE INDEX IF NOT EXISTS idx_bench_talk_conversations_user ON bench_talk_conversations(user_id)",
        "CREATE INDEX IF NOT EXISTS idx_bench_talk_messages_conversation ON bench_talk_messages(conversation_id)",
        "CREATE INDEX IF NOT EXISTS idx_bench_talk_feedback_rating ON bench_talk_feedback(rating)",
        "CREATE INDEX IF NOT EXISTS idx_usage_log_user_action ON subscription_usage_log(user_id, action_type, created_at)",
        "CREATE INDEX IF NOT EXISTS idx_saved_searches_user ON saved_searches(user_id)",
        "CREATE INDEX IF NOT EXISTS idx_drills_category ON drills(category)",
        "CREATE INDEX IF NOT EXISTS idx_drills_org ON drills(org_id)",
        "CREATE INDEX IF NOT EXISTS idx_practice_plans_org ON practice_plans(org_id)",
        "CREATE INDEX IF NOT EXISTS idx_practice_plans_team ON practice_plans(team_name)",
        "CREATE INDEX IF NOT EXISTS idx_pp_drills_plan ON practice_plan_drills(practice_plan_id)",
        "CREATE INDEX IF NOT EXISTS idx_games_league_date ON games(league, game_date)",
        "CREATE INDEX IF NOT EXISTS idx_stats_history_player ON player_stats_history(player_id, season)",
        "CREATE INDEX IF NOT EXISTS idx_game_stats_player ON player_game_stats(player_id, game_date DESC)",
        "CREATE INDEX IF NOT EXISTS idx_game_stats_player_season ON player_game_stats(player_id, season)",
    ]:
        conn.execute(idx_sql)

    # ── NEW: Soft delete + merge + created_by columns on players ──
    p_cols_check = {r[1] for r in conn.execute("PRAGMA table_info(players)").fetchall()}
    for col_name, col_type in [
        ("is_deleted", "INTEGER DEFAULT 0"),
        ("deleted_at", "TEXT"),
        ("deleted_reason", "TEXT"),
        ("deleted_by", "TEXT"),
        ("is_merged", "INTEGER DEFAULT 0"),
        ("merged_into", "TEXT"),
        ("merged_at", "TEXT"),
        ("created_by", "TEXT"),
        ("jersey_number", "TEXT"),
    ]:
        if col_name not in p_cols_check:
            conn.execute(f"ALTER TABLE players ADD COLUMN {col_name} {col_type}")
            conn.commit()

    # ── NEW: player_corrections table ──
    conn.execute("""
        CREATE TABLE IF NOT EXISTS player_corrections (
            id TEXT PRIMARY KEY,
            org_id TEXT NOT NULL,
            user_id TEXT NOT NULL,
            player_id TEXT NOT NULL,
            field_name TEXT NOT NULL,
            old_value TEXT,
            new_value TEXT NOT NULL,
            reason TEXT,
            confidence TEXT DEFAULT 'medium',
            status TEXT DEFAULT 'pending',
            created_at TEXT DEFAULT CURRENT_TIMESTAMP,
            reviewed_at TEXT,
            reviewed_by TEXT,
            review_note TEXT,
            FOREIGN KEY (player_id) REFERENCES players(id)
        )
    """)

    # ── NEW: player_merges audit table ──
    conn.execute("""
        CREATE TABLE IF NOT EXISTS player_merges (
            id TEXT PRIMARY KEY,
            org_id TEXT NOT NULL,
            primary_player_id TEXT NOT NULL,
            duplicate_player_ids TEXT NOT NULL,
            stats_moved INTEGER DEFAULT 0,
            notes_moved INTEGER DEFAULT 0,
            reports_moved INTEGER DEFAULT 0,
            intel_moved INTEGER DEFAULT 0,
            merged_by TEXT NOT NULL,
            merged_at TEXT DEFAULT CURRENT_TIMESTAMP,
            can_undo INTEGER DEFAULT 1,
            undo_before TEXT,
            undone_at TEXT
        )
    """)

    # ── NEW: game_plans table ──
    conn.execute("""
        CREATE TABLE IF NOT EXISTS game_plans (
            id TEXT PRIMARY KEY,
            org_id TEXT NOT NULL,
            user_id TEXT NOT NULL,
            team_name TEXT NOT NULL,
            opponent_team_name TEXT NOT NULL,
            game_date TEXT,
            opponent_analysis TEXT,
            our_strategy TEXT,
            matchups TEXT DEFAULT '[]',
            special_teams_plan TEXT,
            keys_to_game TEXT,
            lines_snapshot TEXT,
            status TEXT DEFAULT 'draft',
            created_at TEXT DEFAULT CURRENT_TIMESTAMP,
            updated_at TEXT DEFAULT CURRENT_TIMESTAMP
        )
    """)

    # ── NEW: series_plans table ──
    conn.execute("""
        CREATE TABLE IF NOT EXISTS series_plans (
            id TEXT PRIMARY KEY,
            org_id TEXT NOT NULL,
            user_id TEXT NOT NULL,
            team_name TEXT NOT NULL,
            opponent_team_name TEXT NOT NULL,
            series_name TEXT NOT NULL,
            series_format TEXT DEFAULT 'best_of_7',
            current_score TEXT DEFAULT '0-0',
            game_notes TEXT DEFAULT '[]',
            working_strategies TEXT DEFAULT '[]',
            needs_adjustment TEXT DEFAULT '[]',
            status TEXT DEFAULT 'active',
            created_at TEXT DEFAULT CURRENT_TIMESTAMP,
            updated_at TEXT DEFAULT CURRENT_TIMESTAMP
        )
    """)

    # ── Chalk Talk columns on game_plans ──
    gp_cols = {r[1] for r in conn.execute("PRAGMA table_info(game_plans)").fetchall()}
    for col_name, col_type in [
        ("session_type", "TEXT DEFAULT 'pre_game'"),
        ("talking_points", "TEXT DEFAULT '{}'"),
        ("what_worked", "TEXT"),
        ("what_didnt_work", "TEXT"),
        ("game_result", "TEXT"),
        ("game_score", "TEXT"),
        ("forecheck", "TEXT"),
        ("breakout", "TEXT"),
        ("defensive_system", "TEXT"),
    ]:
        if col_name not in gp_cols:
            conn.execute(f"ALTER TABLE game_plans ADD COLUMN {col_name} {col_type}")
            conn.commit()

    # ── Enhanced series columns on series_plans ──
    sp_cols = {r[1] for r in conn.execute("PRAGMA table_info(series_plans)").fetchall()}
    for col_name, col_type in [
        ("opponent_systems", "TEXT DEFAULT '{}'"),
        ("key_players_dossier", "TEXT DEFAULT '[]'"),
        ("matchup_plan", "TEXT DEFAULT '{}'"),
        ("adjustments", "TEXT DEFAULT '[]'"),
        ("momentum_log", "TEXT DEFAULT '[]'"),
    ]:
        if col_name not in sp_cols:
            conn.execute(f"ALTER TABLE series_plans ADD COLUMN {col_name} {col_type}")
            conn.commit()

    # ── Scouting list table ──
    conn.execute("""
        CREATE TABLE IF NOT EXISTS scouting_list (
            id TEXT PRIMARY KEY,
            org_id TEXT NOT NULL,
            user_id TEXT NOT NULL,
            player_id TEXT NOT NULL,
            priority TEXT DEFAULT 'medium',
            target_reason TEXT,
            scout_notes TEXT,
            tags TEXT DEFAULT '[]',
            is_active INTEGER DEFAULT 1,
            list_order INTEGER DEFAULT 0,
            last_viewed TEXT,
            times_viewed INTEGER DEFAULT 0,
            created_at TEXT DEFAULT CURRENT_TIMESTAMP,
            updated_at TEXT DEFAULT CURRENT_TIMESTAMP
        )
    """)

    # Indexes for new tables
    for idx_sql in [
        "CREATE INDEX IF NOT EXISTS idx_corrections_player ON player_corrections(player_id)",
        "CREATE INDEX IF NOT EXISTS idx_corrections_org ON player_corrections(org_id, status)",
        "CREATE INDEX IF NOT EXISTS idx_corrections_user ON player_corrections(user_id)",
        "CREATE INDEX IF NOT EXISTS idx_merges_org ON player_merges(org_id)",
        "CREATE INDEX IF NOT EXISTS idx_game_plans_org ON game_plans(org_id)",
        "CREATE INDEX IF NOT EXISTS idx_game_plans_team ON game_plans(team_name)",
        "CREATE INDEX IF NOT EXISTS idx_game_plans_session ON game_plans(session_type)",
        "CREATE INDEX IF NOT EXISTS idx_series_plans_org ON series_plans(org_id)",
        "CREATE INDEX IF NOT EXISTS idx_players_deleted ON players(is_deleted)",
        "CREATE INDEX IF NOT EXISTS idx_players_created_by ON players(created_by)",
        "CREATE INDEX IF NOT EXISTS idx_scouting_list_org ON scouting_list(org_id, user_id)",
        "CREATE INDEX IF NOT EXISTS idx_scouting_list_player ON scouting_list(player_id)",
        "CREATE INDEX IF NOT EXISTS idx_scouting_list_priority ON scouting_list(priority)",
    ]:
        conn.execute(idx_sql)

    conn.commit()

    # ── Migration: Add diagram_data column to drills ──────────
    drill_cols = [col[1] for col in conn.execute("PRAGMA table_info(drills)").fetchall()]
    if "diagram_data" not in drill_cols:
        conn.execute("ALTER TABLE drills ADD COLUMN diagram_data TEXT DEFAULT NULL")
        conn.commit()
        logger.info("Migration: added diagram_data column to drills table")

    # ── Migration: PXI mode column on bench_talk_conversations ──
    bt_cols = {r[1] for r in conn.execute("PRAGMA table_info(bench_talk_conversations)").fetchall()}
    if "mode" not in bt_cols:
        conn.execute("ALTER TABLE bench_talk_conversations ADD COLUMN mode TEXT DEFAULT NULL")
        conn.commit()
        logger.info("Migration: added mode column to bench_talk_conversations")

    # ── Migration: Share columns on reports ──
    rpt_cols = {r[1] for r in conn.execute("PRAGMA table_info(reports)").fetchall()}
    for col_name, col_type in [
        ("share_token", "TEXT DEFAULT NULL"),
        ("shared_with_org", "INTEGER DEFAULT 0"),
    ]:
        if col_name not in rpt_cols:
            conn.execute(f"ALTER TABLE reports ADD COLUMN {col_name} {col_type}")
            conn.commit()
            logger.info("Migration: added %s column to reports", col_name)

    # Index for share token lookups
    conn.execute("CREATE INDEX IF NOT EXISTS idx_reports_share_token ON reports(share_token)")
    conn.execute("CREATE INDEX IF NOT EXISTS idx_reports_shared_org ON reports(org_id, shared_with_org)")
    conn.commit()

    # ── Migration: Quality score columns on reports ──
    for col_name, col_type in [
        ("quality_score", "REAL DEFAULT NULL"),
        ("quality_details", "TEXT DEFAULT NULL"),
    ]:
        if col_name not in rpt_cols:
            conn.execute(f"ALTER TABLE reports ADD COLUMN {col_name} {col_type}")
            conn.commit()
            logger.info("Migration: added %s column to reports", col_name)

    # ── Migration: roster_status on players ──
    # Values: active (default), ap, inj, susp, scrch
    p_cols_final = {r[1] for r in conn.execute("PRAGMA table_info(players)").fetchall()}
    if "roster_status" not in p_cols_final:
        conn.execute("ALTER TABLE players ADD COLUMN roster_status TEXT DEFAULT 'active'")
        conn.commit()
        logger.info("Migration: added roster_status column to players")

    # ── Migration: practice_date on practice_plans ──
    pp_cols = {r[1] for r in conn.execute("PRAGMA table_info(practice_plans)").fetchall()}
    if "practice_date" not in pp_cols:
        conn.execute("ALTER TABLE practice_plans ADD COLUMN practice_date TEXT")
        conn.commit()
        logger.info("Migration: added practice_date column to practice_plans")

    # ── Calendar & Schedule tables ──
    conn.execute("""
        CREATE TABLE IF NOT EXISTS events (
            id TEXT PRIMARY KEY,
            org_id TEXT NOT NULL,
            team_id TEXT,
            player_id TEXT,
            created_by_user_id TEXT,
            feed_id TEXT,
            type TEXT NOT NULL DEFAULT 'OTHER',
            source TEXT NOT NULL DEFAULT 'MANUAL',
            source_external_id TEXT,
            title TEXT NOT NULL,
            description TEXT,
            start_time TEXT NOT NULL,
            end_time TEXT,
            timezone TEXT DEFAULT 'America/Toronto',
            location TEXT,
            league_name TEXT,
            opponent_name TEXT,
            is_home INTEGER,
            visibility TEXT DEFAULT 'ORG',
            created_at TEXT DEFAULT (datetime('now')),
            updated_at TEXT DEFAULT (datetime('now'))
        )
    """)
    conn.execute("""
        CREATE TABLE IF NOT EXISTS calendar_feeds (
            id TEXT PRIMARY KEY,
            org_id TEXT NOT NULL,
            team_id TEXT,
            label TEXT NOT NULL,
            provider TEXT NOT NULL DEFAULT 'ICAL_GENERIC',
            url TEXT NOT NULL,
            active INTEGER DEFAULT 1,
            last_sync_at TEXT,
            sync_error TEXT,
            event_count INTEGER DEFAULT 0,
            created_at TEXT DEFAULT (datetime('now')),
            updated_at TEXT DEFAULT (datetime('now'))
        )
    """)
    conn.execute("CREATE INDEX IF NOT EXISTS idx_events_org ON events(org_id)")
    conn.execute("CREATE INDEX IF NOT EXISTS idx_events_start ON events(start_time)")
    conn.execute("CREATE INDEX IF NOT EXISTS idx_events_team ON events(team_id)")
    conn.execute("CREATE INDEX IF NOT EXISTS idx_events_source ON events(org_id, source, source_external_id)")
    conn.execute("CREATE INDEX IF NOT EXISTS idx_calendar_feeds_org ON calendar_feeds(org_id)")
    conn.commit()

    # ── Messaging tables ──
    conn.execute("""
        CREATE TABLE IF NOT EXISTS msg_conversations (
            id TEXT PRIMARY KEY,
            org_id TEXT,
            participant_ids TEXT NOT NULL,
            status TEXT DEFAULT 'active',
            created_at TEXT DEFAULT (datetime('now')),
            updated_at TEXT DEFAULT (datetime('now'))
        )
    """)
    conn.execute("""
        CREATE TABLE IF NOT EXISTS msg_messages (
            id TEXT PRIMARY KEY,
            conversation_id TEXT NOT NULL REFERENCES msg_conversations(id),
            sender_id TEXT NOT NULL,
            content TEXT NOT NULL,
            sent_at TEXT DEFAULT (datetime('now')),
            read_at TEXT,
            is_system_message INTEGER DEFAULT 0
        )
    """)
    conn.execute("""
        CREATE TABLE IF NOT EXISTS contact_requests (
            id TEXT PRIMARY KEY,
            requester_id TEXT NOT NULL,
            requester_name TEXT,
            requester_role TEXT,
            requester_org TEXT,
            target_player_id TEXT NOT NULL,
            parent_id TEXT NOT NULL,
            status TEXT DEFAULT 'pending',
            message TEXT,
            requested_at TEXT DEFAULT (datetime('now')),
            resolved_at TEXT,
            cooldown_until TEXT
        )
    """)
    conn.execute("""
        CREATE TABLE IF NOT EXISTS message_blocks (
            id TEXT PRIMARY KEY,
            blocker_id TEXT NOT NULL,
            blocked_id TEXT NOT NULL,
            reason TEXT,
            created_at TEXT DEFAULT (datetime('now')),
            UNIQUE(blocker_id, blocked_id)
        )
    """)
    conn.execute("CREATE INDEX IF NOT EXISTS idx_msg_messages_conv ON msg_messages(conversation_id)")
    conn.execute("CREATE INDEX IF NOT EXISTS idx_msg_messages_sent ON msg_messages(sent_at)")
    conn.execute("CREATE INDEX IF NOT EXISTS idx_msg_conv_participants ON msg_conversations(participant_ids)")
    conn.execute("CREATE INDEX IF NOT EXISTS idx_contact_req_parent ON contact_requests(parent_id, status)")
    conn.execute("CREATE INDEX IF NOT EXISTS idx_contact_req_target ON contact_requests(target_player_id)")
    conn.commit()

    # Add linked_player_id to users table for parent→player linking
    user_cols = [col[1] for col in conn.execute("PRAGMA table_info(users)").fetchall()]
    if "linked_player_id" not in user_cols:
        conn.execute("ALTER TABLE users ADD COLUMN linked_player_id TEXT")
        conn.commit()
        logger.info("Migration: added linked_player_id column to users")

    # ── Password reset tokens table ──
    conn.execute("""
        CREATE TABLE IF NOT EXISTS password_reset_tokens (
            id TEXT PRIMARY KEY,
            user_id TEXT NOT NULL,
            token_hash TEXT NOT NULL,
            expires_at TEXT NOT NULL,
            used INTEGER DEFAULT 0,
            created_at TEXT DEFAULT (datetime('now'))
        )
    """)
    conn.commit()

    # ── Email verification columns on users ──
    if "email_verified" not in user_cols:
        conn.execute("ALTER TABLE users ADD COLUMN email_verified INTEGER DEFAULT 0")
        conn.execute("ALTER TABLE users ADD COLUMN email_verify_token TEXT")
        conn.execute("ALTER TABLE users ADD COLUMN email_verify_sent_at TEXT")
        conn.commit()
        logger.info("Migration: added email_verified columns to users")

    # ── Migration: auto-verify existing users (no email service configured) ──
    unverified = conn.execute("SELECT COUNT(*) as c FROM users WHERE email_verified = 0 OR email_verified IS NULL").fetchone()["c"]
    if unverified > 0:
        conn.execute("UPDATE users SET email_verified = 1 WHERE email_verified = 0 OR email_verified IS NULL")
        conn.commit()
        logger.info("Migration: auto-verified %d users (email verification not enforced)", unverified)

    # ── Scout Notes v2 columns ──
    sn_cols = {r["name"] for r in conn.execute("PRAGMA table_info(scout_notes)").fetchall()}
    if "overall_grade" not in sn_cols:
        for col_sql in [
            "ALTER TABLE scout_notes ADD COLUMN game_date TEXT",
            "ALTER TABLE scout_notes ADD COLUMN opponent TEXT",
            "ALTER TABLE scout_notes ADD COLUMN competition_level TEXT",
            "ALTER TABLE scout_notes ADD COLUMN venue TEXT",
            "ALTER TABLE scout_notes ADD COLUMN overall_grade INTEGER",
            "ALTER TABLE scout_notes ADD COLUMN grade_scale TEXT DEFAULT '1-5'",
            "ALTER TABLE scout_notes ADD COLUMN skating_rating INTEGER",
            "ALTER TABLE scout_notes ADD COLUMN puck_skills_rating INTEGER",
            "ALTER TABLE scout_notes ADD COLUMN hockey_iq_rating INTEGER",
            "ALTER TABLE scout_notes ADD COLUMN compete_rating INTEGER",
            "ALTER TABLE scout_notes ADD COLUMN defense_rating INTEGER",
            "ALTER TABLE scout_notes ADD COLUMN strengths_notes TEXT",
            "ALTER TABLE scout_notes ADD COLUMN improvements_notes TEXT",
            "ALTER TABLE scout_notes ADD COLUMN development_notes TEXT",
            "ALTER TABLE scout_notes ADD COLUMN one_line_summary TEXT",
            "ALTER TABLE scout_notes ADD COLUMN prospect_status TEXT",
            "ALTER TABLE scout_notes ADD COLUMN visibility TEXT DEFAULT 'PRIVATE'",
            "ALTER TABLE scout_notes ADD COLUMN note_mode TEXT DEFAULT 'QUICK'",
        ]:
            conn.execute(col_sql)
        conn.execute("CREATE INDEX IF NOT EXISTS idx_scout_notes_date ON scout_notes(game_date)")
        conn.execute("CREATE INDEX IF NOT EXISTS idx_scout_notes_status ON scout_notes(prospect_status)")
        conn.commit()
        logger.info("Migration: added Scout Notes v2 columns")

    # ── Migration: add elite_prospects_url to players (CR-012) ──
    p_cols_ep = {r[1] for r in conn.execute("PRAGMA table_info(players)").fetchall()}
    if "elite_prospects_url" not in p_cols_ep:
        conn.execute("ALTER TABLE players ADD COLUMN elite_prospects_url TEXT")
        conn.commit()
        logger.info("Migration: added elite_prospects_url column to players")

    # ── Migration: rename old subscription tiers (CR-009) ──
    try:
        n1 = conn.execute("UPDATE users SET subscription_tier = 'scout' WHERE subscription_tier = 'novice'").rowcount
        n2 = conn.execute("UPDATE users SET subscription_tier = 'team_org' WHERE subscription_tier = 'team'").rowcount
        n3 = conn.execute("UPDATE users SET subscription_tier = 'program_org' WHERE subscription_tier = 'aaa_org'").rowcount
        if n1 or n2 or n3:
            conn.commit()
            logger.info("Tier migration: novice→scout:%d, team→team_org:%d, aaa_org→program_org:%d", n1, n2, n3)
    except Exception as e:
        logger.warning("Tier migration note: %s", e)

    conn.close()
    logger.info("SQLite database initialized: %s", DB_FILE)

    # ── Startup backup ──
    try:
        backup_dir = os.path.join(os.path.dirname(DB_FILE), "backups")
        os.makedirs(backup_dir, exist_ok=True)
        ts = datetime.now(timezone.utc).isoformat()[:19].replace(":", "-")
        backup_name = f"prospectx_{ts}.db"
        shutil.copy2(DB_FILE, os.path.join(backup_dir, backup_name))
        # Keep last 20 backups
        backups = sorted(glob.glob(os.path.join(backup_dir, "prospectx_*.db")))
        for old in backups[:-20]:
            os.remove(old)
        logger.info("Startup backup created: %s", backup_name)
    except Exception as e:
        logger.warning("Startup backup failed: %s", e)


def seed_hockey_os():
    """Seed the Hockey OS reference data (systems library + glossary)."""
    conn = get_db()

    # Check if already seeded
    count = conn.execute("SELECT COUNT(*) FROM systems_library").fetchone()[0]
    if count > 0:
        conn.close()
        return

    # ── Systems Library ───────────────────────────────────────
    systems = [
        # Forecheck systems
        ("forecheck", "F_AGGRESSIVE_1_2_2", "Aggressive 1-2-2", "First forward pressures hard, two wingers contain middle ice, two D stay high. Forces turnovers but vulnerable to speed through the middle.", "High pressure, turnovers forced, good for teams with fast F1", "Can be beaten with quick east-west passes, requires disciplined F1", "Fast F1 with good closing speed, physical wingers"),
        ("forecheck", "F_1_2_2_TRAP", "1-2-2 Neutral Zone Trap", "Passive forecheck clogging the neutral zone. F1 angles, two forwards sit in a wall across the red line, D stay deep. Limits opponent transition.", "Limits odd-man rushes, controls pace, low risk", "Can be frustrating for own forwards, low event hockey", "Disciplined forwards willing to backcheck, patient defensemen"),
        ("forecheck", "F_2_1_2", "2-1-2 Forecheck", "Two forwards go deep on the forecheck, one center supports high, two D pinch. Very aggressive, forces turnovers but leaves the NZ open.", "Maximum OZ pressure, turnovers deep, dominant possession", "Vulnerable to stretch passes, breakaways if pinch fails", "Two physical forecheckers, mobile D who can pinch and recover"),
        ("forecheck", "F_1_3_1", "1-3-1 Forecheck", "One forward pressures, three across the neutral zone in a line, one D cheats up. Great for trapping teams and counter-attacking.", "Controls NZ, generates turnovers in transition, counter-attack", "Weak if opponent gets behind the 3-man wall", "Smart F1 who angles well, fast counter-attack wingers"),
        ("forecheck", "F_1_1_3", "1-1-3 Passive Trap", "Deep trap with only one forechecker, one mid-zone forward, and three players sitting back. Ultra-defensive shell.", "Almost impossible to get clean entries against, low goals against", "Very few offensive chances, boring hockey, hard to score", "Disciplined team willing to play low-event hockey"),
        # DZ structures
        ("dz_coverage", "DZ_MAN_TO_MAN", "Man-to-Man DZ", "Each player picks up a man in the defensive zone. Tight coverage, eliminates passing lanes, but can be exposed by picks and screens.", "Tight coverage, eliminates freelancers, good vs cycle teams", "Vulnerable to screens, picks, and mismatch situations", "Players with good feet, communication, and physicality"),
        ("dz_coverage", "DZ_ZONE", "Zone Defense DZ", "Players cover areas rather than men. Strong side overload, weak side rotates. Standard NHL-style coverage.", "Good against cycle, covers shooting lanes, less skating", "Can leave men open in soft areas, requires communication", "Smart players who read plays, good stick positioning"),
        ("dz_coverage", "DZ_COLLAPSING_BOX", "Collapsing Box DZ", "Four players form a box in front of the net, one player pressures the puck. Collapses inward on shots. Very protective of the slot.", "Protects the slot and net-front, limits high-danger chances", "Gives up perimeter shots, weak vs point shots with traffic", "Shot-blocking willingness, goalie who handles perimeter shots"),
        ("dz_coverage", "DZ_SWARM", "Swarm Coverage DZ", "Aggressive puck pursuit in the DZ. All five players pressure the puck carrier. High risk, high reward — forces turnovers or gets burned.", "Forces turnovers, creates transition chances from DZ", "Extremely vulnerable if beaten, requires elite conditioning", "High-compete players, great conditioning, smart gambles"),
        # OZ setups
        ("oz_setup", "OZ_UMBRELLA", "Umbrella OZ / PP", "One player at the top with two half-wall options and two net-front/bumper players. Classic power play look. Creates triangles for one-timers.", "One-timer options, triangle passing, multiple shooting lanes", "Predictable if scouted, requires a strong bumper player", "Shooter at the top, playmaker on the half-wall, big net-front"),
        ("oz_setup", "OZ_OVERLOAD", "Overload OZ", "Shifts 4 players to one side of the ice, creating numerical advantage. Quick passes in tight space, back-door options.", "Creates confusion, numerical advantage, back-door plays", "Weak side is empty — one pass beats it entirely", "Quick decision-makers, players comfortable in tight spaces"),
        ("oz_setup", "OZ_CYCLE", "Heavy Cycle OZ", "Grind the puck down low, use the half-wall, and work it to the net-front or point. Physical, possession-based.", "Controls the puck, wears down opponents, creates net-front chaos", "Slow to generate shots, can be broken by aggressive DZ pressure", "Big, strong forwards who protect the puck, net-front presence"),
        ("oz_setup", "OZ_1_3_1_PP", "1-3-1 Power Play", "One player at the top, three across the middle (two half-walls + bumper), one net-front. Creates passing lanes and mid-range shots.", "Multiple shooting options, hard to defend bumper, cross-ice plays", "Requires elite passer at the top, vulnerable to aggressive PK", "High-IQ playmaker at the top, finisher in the bumper"),
        # Breakout systems
        ("breakout", "BO_STANDARD", "Standard Breakout", "D-to-D behind the net, up to the winger on the wall, center supports through the middle. Basic but reliable.", "Simple, reliable, low turnover risk", "Predictable, easy for aggressive forechecks to read", "Good first-pass D, wingers who get open on the wall"),
        ("breakout", "BO_REVERSE", "Reverse Breakout", "D starts one direction then reverses behind the net to the weak-side D or winger. Changes the point of attack.", "Changes angles, beats overcommitted forecheckers", "Risky if weak-side support is late, requires good skating D", "Mobile defensemen, quick-thinking forwards"),
        ("breakout", "BO_WHEEL", "Wheel Breakout", "D skates behind the net and carries the puck up-ice themselves. Aggressive, creates odd-man opportunities.", "Creates speed through NZ, D joins the rush as an extra attacker", "Very risky if D gets caught, requires elite skating ability", "Mobile, puck-carrying defensemen with good vision"),
        # PK formations
        ("pk_formation", "PK_BOX", "Box PK", "Four players form a box (diamond). Protects the slot, takes away one-timer lanes. Standard PK look.", "Protects the slot, eliminates one-timers, simple reads", "Gives up perimeter shots, can be stretched by movement", "Shot blockers, strong sticks in passing lanes"),
        ("pk_formation", "PK_DIAMOND", "Diamond PK", "One forward pressures high, two forwards/D at the dots, one D in front of the net. More aggressive than a box.", "Pressures the puck, disrupts PP entries, forces turnovers", "Vulnerable if high man is beaten, leaves backdoor open", "Fast penalty-killing forwards, aggressive D"),
        ("pk_formation", "PK_AGGRESSIVE", "Aggressive PK", "Two forwards pressure the puck aggressively on the PK, trying to force turnovers and create shorthanded chances.", "Shorthanded goals, disrupts PP flow, momentum swings", "Extremely risky — one pass can expose the entire PK", "Elite PKers with speed, high hockey IQ, calculated aggression"),
    ]

    for sys_type, code, name, desc, strengths, weaknesses, personnel in systems:
        conn.execute(
            "INSERT INTO systems_library (id, system_type, code, name, description, strengths, weaknesses, ideal_personnel) VALUES (?, ?, ?, ?, ?, ?, ?, ?)",
            (str(uuid.uuid4()), sys_type, code, name, desc, strengths, weaknesses, personnel),
        )

    # ── Hockey Glossary Terms ─────────────────────────────────
    terms = [
        ("Controlled Entry", "transition", "Entering the offensive zone with possession of the puck (carry-in or pass). Opposite of a dump-in.", '["carry-in", "clean entry"]', "Microstat tracking, transition analysis"),
        ("Controlled Exit", "transition", "Exiting the defensive zone with possession — either carrying or passing the puck out cleanly.", '["clean exit", "breakout with possession"]', "Microstat tracking, transition analysis"),
        ("Battle Win Rate", "compete", "Percentage of loose puck battles won. Measured in board play, net-front, and forecheck scenarios.", '["puck battle %", "compete rate"]', "Physical evaluation, compete level assessment"),
        ("Forecheck Pressure", "forecheck", "An aggressive action on the puck carrier in the offensive zone to force a turnover or rushed play.", '["F1 pressure", "forechecking"]', "System adherence, forecheck evaluation"),
        ("Slot Pass", "offense", "A pass completed into the scoring slot area (between the faceoff dots and below the top of the circles).", '["pass to slot", "dangerous pass"]', "Offensive creation, playmaking evaluation"),
        ("xG (Expected Goals)", "analytics", "A model-based metric estimating the probability a shot becomes a goal based on location, type, and context.", '["expected goals"]', "Advanced analytics, shot quality measurement"),
        ("Cycle Play", "systems", "Possession-based offensive strategy working the puck along the boards and behind the net to create scoring chances.", '["grinding", "below the goal line"]', "System description, offensive evaluation"),
        ("Gap Control", "defense", "The distance a defender maintains between themselves and the attacking player. Tight gap = aggressive, loose gap = conservative.", '["closing speed", "gap management"]', "Defensive evaluation, skating assessment"),
        ("Net-Front Presence", "offense", "A player's ability to establish and maintain position in front of the opposing net to screen, tip, and create chaos.", '["net-front", "crease work", "dirty area goals"]', "Role fit, archetype classification"),
        ("Transition Game", "transition", "The ability to move the puck effectively from defense to offense through the neutral zone. Measured by controlled entries/exits.", '["transition", "NZ play"]', "Overall game assessment, speed of play"),
        ("Two-Way Forward", "archetypes", "A forward who contributes offensively while also being responsible defensively. Trusted in all three zones and situations.", '["200-foot player", "complete forward"]', "Archetype classification, role fit"),
        ("Puck-Moving Defenseman", "archetypes", "A defenseman whose primary value is moving the puck out of the DZ and through the NZ. Good first pass, skating, and vision.", '["mobile D", "skating defenseman"]', "Archetype classification, D evaluation"),
        ("F1/F2/F3", "systems", "The three forward roles in a forecheck. F1 = first forechecker (pressure), F2 = second (support/contain), F3 = third (high/safety).", '["forecheck roles"]', "System description, forecheck evaluation"),
        ("Retrieve and Regroup", "breakout", "A breakout strategy where the D retrieves the puck behind the net and looks to regroup rather than make a quick breakout pass.", '["regroup"]', "Breakout evaluation, patience assessment"),
        ("Shooting Percentage", "analytics", "Goals divided by shots on goal. Context matters — league average varies by level. Sustainability is key.", '["S%", "sh%", "shooting efficiency"]', "Offensive evaluation, shot selection"),
    ]

    for term, cat, defn, aliases, context in terms:
        conn.execute(
            "INSERT INTO hockey_terms (id, term, category, definition, aliases, usage_context) VALUES (?, ?, ?, ?, ?, ?)",
            (str(uuid.uuid4()), term, cat, defn, aliases, context),
        )

    conn.commit()
    conn.close()
    logger.info("Seeded Hockey OS: %d systems, %d glossary terms", len(systems), len(terms))


def seed_glossary_v2():
    """Expand hockey glossary from 15 terms to 115+ terms covering the full hockey vocabulary."""
    conn = get_db()
    count = conn.execute("SELECT COUNT(*) FROM hockey_terms").fetchone()[0]
    if count > 20:
        conn.close()
        return  # Already expanded

    # ── Rink & Game Basics ─────────────────────────────────────
    terms = [
        ("Barn", "rink", "Rink or arena.", '["arena", "rink"]', "General hockey slang"),
        ("Bench", "rink", "Where players sit when they are not on the ice.", '["pine"]', "Game basics"),
        ("Box (Defensive)", "rink", "Defensive system with four players forming a box shape in the zone.", '["box formation"]', "Defensive zone coverage"),
        ("Crease", "rink", "Blue semi-circle in front of the net where the goalie plays.", '["blue paint", "goal crease"]', "Rink geography, goaltending"),
        ("Half Wall", "rink", "Boards area roughly halfway between the corner and the blue line.", '["half-wall"]', "Offensive zone positioning, PP formations"),
        ("Hash Marks", "rink", "Short lines beside the faceoff circles that help position players for faceoffs.", '[]', "Faceoff positioning"),
        ("Neutral Zone", "rink", "Center-ice area between the two blue lines.", '["NZ"]', "Transition play, trap systems"),
        ("Attacking Zone", "rink", "The offensive end from the opponent's blue line to the end boards.", '["offensive zone", "OZ", "O-zone"]', "Zone play analysis"),
        ("Defending Zone", "rink", "Your own end from your blue line back to your goal line.", '["defensive zone", "DZ", "D-zone"]', "Defensive coverage analysis"),

        # ── Plays, Tactics & Situations ─────────────────────────
        ("Backcheck", "tactics", "Forwards skating hard back toward their own zone to pressure the puck and try to regain it.", '["backchecking"]', "Defensive responsibility, two-way play evaluation"),
        ("Forecheck", "tactics", "Pressuring the opponent in their zone to force turnovers and keep the puck in.", '["forechecking", "F1 pressure"]', "System adherence, forecheck evaluation"),
        ("Breakout", "tactics", "Moving the puck out of your defensive zone to start offense up the ice.", '["breakout play"]', "Transition evaluation, DZ play"),
        ("Breakaway", "tactics", "Puck carrier alone in on the goalie with no defenders between them.", '["clean break"]', "Scoring situations, speed evaluation"),
        ("Dump and Chase", "tactics", "Shooting the puck deep into the offensive zone and forechecking to get it back, instead of carrying it in.", '["dump-in", "chip and chase"]', "Zone entry strategy, forecheck evaluation"),
        ("Cycle", "tactics", "Rotating with teammates along the boards in the offensive zone to maintain possession and create openings.", '["cycle game", "cycling"]', "Offensive systems, puck possession evaluation"),
        ("Deke", "tactics", "Fake with body, head, or stick to beat a defender or goalie.", '["fake", "move"]', "Puck skills evaluation"),
        ("Dangle", "tactics", "High-skill stickhandling move that often completely beats or undresses a defender.", '["dangling", "sick dangle"]', "Elite puck skills evaluation"),
        ("Odd-Man Rush", "tactics", "Rush where the attacking team has more skaters than the defenders back (2-on-1, 3-on-2).", '["2-on-1", "3-on-2", "odd man"]', "Transition play, rush analysis"),
        ("Power Play", "tactics", "Situation where one team has more players on the ice because the other took a penalty.", '["PP", "man advantage"]', "Special teams analysis"),
        ("Penalty Kill", "tactics", "Shorthanded team trying to defend while the opponent is on the power play.", '["PK", "killing a penalty", "shorthanded"]', "Special teams analysis"),
        ("Shorthanded Goal", "tactics", "Goal scored by the team that is killing a penalty.", '["shorty", "SHG"]', "Special teams evaluation"),
        ("Screened Shot", "tactics", "Shot where the goalie's vision is blocked by traffic in front.", '["screen", "traffic"]', "Offensive tactics, net-front presence"),
        ("Drop Pass", "tactics", "Puck carrier leaves the puck behind for a trailing teammate to pick up in stride.", '["drop"]', "Zone entry, PP entries"),
        ("Headmanning", "tactics", "Passing the puck ahead to a teammate who is already skating up ice.", '["headman pass", "stretch pass"]', "Transition play, breakout evaluation"),
        ("Splitting the Defense", "tactics", "Skating with the puck between two defenders to break through the middle.", '["splitting the D"]', "Puck skills, offensive evaluation"),
        ("Body Check", "tactics", "Using the hip or shoulder to legally slow or stop an opponent who has the puck.", '["hit", "check", "finish your check"]', "Physicality evaluation"),
        ("Poke Check", "tactics", "Using the blade of the stick to jab at the puck and knock it away from the puck carrier.", '["stick check"]', "Defensive skills evaluation"),
        ("Sweep Check", "tactics", "Laying the stick flat on the ice and sweeping it along the surface to knock the puck away.", '["sweeping"]', "Defensive skills evaluation"),
        ("Freezing the Puck", "tactics", "Holding or covering the puck to force a whistle and stoppage.", '["freeze it"]', "Game management, goaltending"),

        # ── Scoring & Offense ───────────────────────────────────
        ("Apple", "scoring", "Assist on a goal.", '["helper", "dish"]', "Offensive production, hockey slang"),
        ("Gino", "scoring", "Goal.", '["tally", "marker"]', "Scoring slang"),
        ("Snipe", "scoring", "Accurate, dangerous shot that beats the goalie clean.", '["sniper", "picked a corner"]', "Shot evaluation, offensive assessment"),
        ("Bar Down", "scoring", "Shot that hits the bottom of the crossbar and goes in.", '["bar-down", "crossbar and in"]', "Shooting evaluation, highlight play"),
        ("Five-Hole", "scoring", "Space between the goalie's legs.", '["5-hole"]', "Shooting targets, goaltending evaluation"),
        ("One-Timer", "scoring", "Catch-and-shoot in one motion off a pass, without stopping the puck.", '["one-T", "1T"]', "Offensive skills, PP evaluation"),
        ("Howitzer", "scoring", "Very hard slap shot.", '["bomb", "cannon"]', "Shot power evaluation"),
        ("Slap Shot", "scoring", "A hard shot where the player winds up, slaps the ice/puck, and generates maximum power.", '["clapper", "slapshot"]', "Shooting evaluation"),
        ("Wrist Shot", "scoring", "A shot created by a quick flicking or rolling motion of the wrists to propel the puck.", '["wrister"]', "Shooting evaluation"),
        ("Hat Trick", "scoring", "Three goals by one player in a single game.", '["hatty"]', "Scoring milestones"),
        ("Natural Hat Trick", "scoring", "Same player scoring three goals in a row without anyone else scoring in between.", '["natural hatty"]', "Scoring milestones"),
        ("Light the Lamp", "scoring", "Score a goal — refers to the red goal light turning on.", '["lamp lighter"]', "Scoring slang"),
        ("Barnburner", "scoring", "High-scoring, wild, back-and-forth game.", '["shootout", "track meet"]', "Game description slang"),
        ("Top Cheese", "scoring", "Goal scored in the top shelf of the net.", '["top ched", "top shelf"]', "Shooting evaluation"),
        ("Muffin", "scoring", "Weak shot that floats in slowly.", '["flutterball"]', "Negative shot evaluation"),

        # ── Gear & Basic Terms ──────────────────────────────────
        ("Biscuit", "gear", "The puck.", '["rubber", "frozen rubber", "pill"]', "General hockey slang"),
        ("Twig", "gear", "Hockey stick.", '["lumber", "stick"]', "General hockey slang"),
        ("Bucket", "gear", "Helmet.", '["lid"]', "General hockey slang"),
        ("Mitts", "gear", "Hands or gloves — often used when talking about good hands or fighting.", '["hands", "gloves"]', "Skills evaluation slang"),
        ("Blocker", "gear", "Goalie's rectangular padded glove on the stick hand.", '["blocker side"]', "Goaltending evaluation"),
        ("Glove Hand", "gear", "Goalie's catching hand, opposite the stick hand.", '["glove side", "catcher"]', "Goaltending evaluation"),
        ("Chiclets", "gear", "Teeth — often used when joking about missing teeth.", '[]', "Hockey culture slang"),

        # ── Player Roles & Types ────────────────────────────────
        ("Beauty", "roles", "Player who is skilled, works hard, and is well-liked in the room.", '["beaut"]', "Character evaluation, hockey culture"),
        ("Grinder", "roles", "High-effort, physical, checking-focused player who does the hard work and may not score much.", '["worker", "lunch pail"]', "Archetype classification, role evaluation"),
        ("Mucker", "roles", "Similar to a grinder but even more physical and combative — digs in corners and stirs things up.", '["agitator"]', "Archetype classification"),
        ("Plug", "roles", "Low-skill but high-effort player who forechecks, finishes checks, and kills penalties.", '[]', "Role player evaluation"),
        ("Goon", "roles", "Enforcer whose role is mostly physicality and fighting.", '["enforcer", "tough guy"]', "Role classification"),
        ("Pylon", "roles", "Slow or ineffective player who is easy to skate around, like a practice cone.", '["cone"]', "Negative evaluation slang"),
        ("Bender", "roles", "Weak skater whose ankles bend in.", '[]', "Skating evaluation slang"),
        ("Sieve", "roles", "Goalie who allows a lot of goals or weak shots.", '[]', "Negative goaltending evaluation"),
        ("Shadow", "roles", "Player assigned to follow and shut down a star opponent.", '["shutdown guy"]', "Defensive role assignment"),
        ("Cherry Picker", "roles", "Player who hangs high near center ice looking for breakaways, not helping on defense.", '["floater"]', "Defensive responsibility evaluation"),
        ("Grocery Stick", "roles", "Player who sits between the forwards and defense on the bench, rarely getting shifts.", '[]', "Lineup depth slang"),
        ("Turnstile", "roles", "A defender who opponents easily skate around all game.", '[]', "Negative defensive evaluation"),
        ("1C / 2C / 3C", "roles", "First-, second-, third-line center — indicating pecking order and usage.", '["top-line center", "depth center"]', "Lineup deployment, role evaluation"),
        ("Two-Way Center", "roles", "Strong both offensively and defensively — plays PP and PK, matches against top lines.", '["200-foot center", "complete center"]', "Archetype classification"),
        ("Power Forward", "roles", "Big, physical winger who can score and forecheck hard.", '["power wing"]', "Archetype classification"),
        ("Sniper", "roles", "High-end goal-scorer with elite shot — often set up on flanks or off-wing.", '["goal scorer", "trigger man"]', "Archetype classification"),
        ("Puck-Moving D", "roles", "Defenseman who joins the rush, runs PP, and moves the puck with skating and passing.", '["offensive D", "mobile D"]', "Archetype classification"),
        ("Stay-at-Home D", "roles", "Defensive-minded blueliner who protects the front of the net and plays simple, low-risk hockey.", '["shutdown D", "defensive D"]', "Archetype classification"),

        # ── Penalties & Discipline ──────────────────────────────
        ("Minor Penalty", "penalties", "Standard 2-minute penalty where the team plays shorthanded; ends early if the opposition scores on a 5-on-4.", '["2-minute minor"]', "Rules and discipline"),
        ("Double Minor", "penalties", "Four minutes served as two consecutive 2-minute minors — often for high-sticking causing injury.", '["4-minute penalty"]', "Rules and discipline"),
        ("Major Penalty", "penalties", "Five-minute penalty for severe infractions — team is shorthanded the full five minutes regardless of goals scored.", '["5-minute major"]', "Rules and discipline"),
        ("Misconduct", "penalties", "Ten-minute penalty where the player sits but is replaced on the ice — no man disadvantage.", '["10-minute misconduct"]', "Rules and discipline"),
        ("Game Misconduct", "penalties", "Player is ejected for the rest of the game. A substitute replaces them so no automatic man-short.", '["game ejection"]', "Rules and discipline"),
        ("Match Penalty", "penalties", "Ejection plus a five-minute major served by a teammate for intent to injure.", '[]', "Rules and discipline"),
        ("Penalty Shot", "penalties", "Awarded when a clear scoring chance is illegally denied — fouled player gets a one-on-one vs the goalie.", '[]', "Rules and special situations"),
        ("Boarding", "penalties", "Hit that violently drives an opponent into the boards in a dangerous way.", '[]', "Penalty types, physicality evaluation"),
        ("Tripping", "penalties", "Using stick, arm, or leg to make an opponent fall or lose balance. 2 minutes.", '[]', "Penalty types, discipline evaluation"),
        ("Hooking", "penalties", "Using the blade of the stick to slow or impede an opponent's skating. 2 minutes.", '[]', "Penalty types, discipline evaluation"),
        ("Holding", "penalties", "Grabbing an opponent or their stick to restrict movement. 2 minutes.", '["holding the stick"]', "Penalty types, discipline evaluation"),
        ("Interference", "penalties", "Impeding a player who doesn't have the puck. 2 minutes.", '[]', "Penalty types, discipline evaluation"),
        ("Slashing", "penalties", "Swinging the stick at an opponent — 2 or 5 minutes depending on severity.", '[]', "Penalty types, discipline evaluation"),
        ("High-Sticking", "penalties", "Contact with an opponent using the stick above shoulder height — often 2 min, can be double minor with injury.", '["high stick"]', "Penalty types"),
        ("Cross-Checking", "penalties", "Checking an opponent using the shaft of the stick with both hands.", '["cross check"]', "Penalty types, physicality"),
        ("Charging", "penalties", "Taking several strides or jumping into a hit, delivering excessive force.", '[]', "Penalty types, discipline"),
        ("Roughing", "penalties", "Extra shoves, punches, or scrums after the whistle or away from the play.", '["rough stuff"]', "Penalty types, discipline"),
        ("Delay of Game", "penalties", "Includes shooting the puck directly over the glass from the defensive zone or goalie playing puck in restricted area.", '["DOG"]', "Penalty types"),
        ("Spearing", "penalties", "Jabbing an opponent with the stick blade like a spear — automatically a major.", '[]', "Penalty types, severe infractions"),

        # ── Slang & Culture ─────────────────────────────────────
        ("Chirp", "slang", "Trash talk directed at opponents or sometimes officials.", '["chirping", "jawing"]', "Hockey culture, personality evaluation"),
        ("Celly", "slang", "Celebration after scoring a goal.", '["celi", "goal celebration"]', "Hockey culture slang"),
        ("Flow", "slang", "Long hair flowing out from under the helmet.", '["lettuce", "salad"]', "Hockey culture slang"),
        ("Chippy", "slang", "Description for a game with rising tempers and extra rough stuff.", '["heated"]', "Game description"),
        ("Gongshow", "slang", "Game that has gotten out of control with lots of penalties, scrums, or chaos.", '["circus"]', "Game description"),
        ("Warm Up the Bus", "slang", "Expression used when the outcome is basically decided and the road team is heading home with a loss.", '[]', "Hockey culture expression"),
        ("Coast to Coast", "slang", "A player carrying the puck from their own end all the way into the offensive end.", '["end to end"]', "Offensive highlight, skating evaluation"),
        ("Wheels", "slang", "A player's skating speed.", '["jets", "burners"]', "Skating evaluation slang"),
        ("Bag Skate", "slang", "Hard conditioning practice with lots of skating as punishment.", '["skate"]', "Practice/coaching culture"),
        ("Sin Bin", "slang", "The penalty box.", '["box"]', "Hockey culture slang"),

        # ── Advanced Analytics ──────────────────────────────────
        ("Corsi", "analytics", "Shot attempt differential — shots on goal + blocked shots + missed shots. Measures puck possession.", '["CF", "CF%", "shot attempts"]', "Advanced analytics, possession metrics"),
        ("Fenwick", "analytics", "Unblocked shot attempt differential — shots on goal + missed shots (excludes blocked).", '["FF", "FF%"]', "Advanced analytics, possession metrics"),
        ("PDO", "analytics", "Sum of team shooting percentage and save percentage — measures luck/variance.", '["sh% + sv%"]', "Advanced analytics, luck/sustainability metrics"),
        ("Zone Entries", "analytics", "Tracking controlled vs dump entries and their success rates.", '["entries with possession"]', "Transition analytics, microstat tracking"),
        ("Zone Exits", "analytics", "Tracking clean vs failed exits from the defensive zone.", '["exits with possession"]', "Transition analytics, microstat tracking"),
        ("High-Danger Chances", "analytics", "Scoring chances from prime areas — slot and near crease.", '["HD chances", "HDCF", "inner slot"]', "Shot quality analytics"),
        ("Puck Possession Metrics", "analytics", "Time on attack, zone time, shot attempt share — all measures of controlling the puck.", '["possession time", "TOA"]', "Advanced analytics overview"),

        # ── Game Strategy & Situations ──────────────────────────
        ("Icing", "strategy", "Shooting puck from behind red line across opposing goal line without touch — results in faceoff in offending team's zone.", '[]', "Rules, game situations"),
        ("Offside", "strategy", "Attacking player entering offensive zone before puck crosses blue line — results in faceoff outside zone.", '[]', "Rules, game situations"),
        ("Delayed Penalty", "strategy", "Penalty called but play continues until offending team touches puck — non-offending team often pulls goalie for extra attacker.", '["delayed call"]', "Game situations, special teams strategy"),
        ("Empty Net", "strategy", "Pulling goalie for extra attacker, typically when trailing late in game.", '["extra attacker", "EN"]', "Late-game strategy"),
        ("Line Matching", "strategy", "Coach strategy to get favorable matchups — shutdown line vs top line, offensive line vs weak defense.", '["matchups"]', "Coaching strategy, deployment evaluation"),
        ("Line Changes", "strategy", "Strategic substitutions to manage energy and matchups.", '["change on the fly"]', "Game management"),
    ]

    for term, cat, defn, aliases, context in terms:
        try:
            conn.execute(
                "INSERT INTO hockey_terms (id, term, category, definition, aliases, usage_context) VALUES (?, ?, ?, ?, ?, ?)",
                (str(uuid.uuid4()), term, cat, defn, aliases, context),
            )
        except Exception:
            pass  # Skip duplicates from original seed

    conn.commit()
    conn.close()
    logger.info("Seeded glossary v2: %d additional hockey terms", len(terms))


def seed_templates():
    """Seed the 19 report templates if they don't exist yet."""
    conn = get_db()
    count = conn.execute("SELECT COUNT(*) FROM report_templates").fetchone()[0]
    if count > 0:
        conn.close()
        return

    templates = [
        ("Pro/Amateur Skater Report", "pro_skater"),
        ("Unified Prospect Report", "unified_prospect"),
        ("Goalie Report", "goalie"),
        ("Single Game Decision Report", "game_decision"),
        ("Season Player Intelligence", "season_intelligence"),
        ("Elite Operations Engine", "operations"),
        ("Team Identity Card", "team_identity"),
        ("Opponent Game Plan", "opponent_gameplan"),
        ("Agent Pack", "agent_pack"),
        ("Development Roadmap", "development_roadmap"),
        ("Player/Family Card", "family_card"),
        ("Line Chemistry Report", "line_chemistry"),
        ("Special Teams Optimization", "st_optimization"),
        ("Trade/Acquisition Target", "trade_target"),
        ("Draft Class Comparative", "draft_comparative"),
        ("Season Progress Report", "season_progress"),
        ("Practice Plan Generator", "practice_plan"),
        ("Playoff Series Prep", "playoff_series"),
        ("Goalie Tandem Optimization", "goalie_tandem"),
    ]

    for name, rtype in templates:
        conn.execute(
            "INSERT INTO report_templates (id, template_name, report_type, is_global, prompt_text) VALUES (?, ?, ?, 1, ?)",
            (str(uuid.uuid4()), name, rtype, f"You are an elite hockey scout. Generate a {name} for the given player."),
        )

    conn.commit()
    conn.close()
    logger.info("Seeded %d report templates", len(templates))


def seed_new_templates():
    """Add any new report templates that were introduced after initial seeding."""
    conn = get_db()
    new_templates = [
        ("ProspectX Metrics Dashboard", "indices_dashboard",
         "Player Analytics", "Advanced Stats",
         "Visual dashboard of all ProspectX performance metrics with league percentile rankings, position comparisons, and development priorities."),
        ("League Benchmarks Comparison", "league_benchmarks",
         "Competitive Intelligence", "League Benchmarks",
         "Compare a team to league averages across offense, defense, special teams, and statistical leaders with trend analysis."),
        ("Team Season Projection", "season_projection",
         "Competitive Intelligence", "League Benchmarks",
         "Project full season standings, playoff odds, championship probabilities, and remaining schedule difficulty."),
        ("Next Season Player Projection", "player_projection",
         "Player Analytics", "Projections & Development",
         "Project a player's next season performance across conservative, expected, and optimistic scenarios with comparable player analysis."),
        ("Free Agent Market Analysis", "free_agent_market",
         "Competitive Intelligence", "Market & Acquisitions",
         "Analyze available uncommitted players by position, quality grade, and system fit with market trend analysis and timing recommendations."),
        ("Pre-Game Intel Brief", "pre_game_intel",
         "Coaching", "Game Prep",
         "Concise, bench-ready pre-game briefing with opponent snapshot, key matchups, goaltending report, special teams intel, and game keys."),
        ("Prep/College Player Guide", "player_guide_prep_college",
         "Player Development", "Pathway Guides",
         "Comprehensive guide for players transitioning to prep school or college hockey with readiness assessment, pathway options, and family action items."),
    ]
    added = 0
    for name, rtype, cat, subcat, desc in new_templates:
        existing = conn.execute("SELECT id FROM report_templates WHERE report_type = ?", (rtype,)).fetchone()
        if not existing:
            conn.execute(
                "INSERT INTO report_templates (id, template_name, report_type, is_global, prompt_text, category, subcategory) VALUES (?, ?, ?, 1, ?, ?, ?)",
                (str(uuid.uuid4()), name, rtype, desc, cat, subcat),
            )
            added += 1
    if added:
        conn.commit()
        logger.info("Added %d new report templates", added)
    conn.close()


def _migrate_template_prompts():
    """One-time migration: update report_templates with rich prompt_text from seed_templates.py prompts."""
    conn = get_db()
    # Check if migration already applied (pro_skater prompt_text length > 200 chars)
    check = conn.execute(
        "SELECT LENGTH(prompt_text) as len FROM report_templates WHERE report_type = 'pro_skater' LIMIT 1"
    ).fetchone()
    already_migrated = check and check["len"] and check["len"] > 200

    # Even if base migration is done, check for new templates with short prompts
    if already_migrated:
        # Catch-up: update any newly seeded templates that still have short prompt_text
        from seed_templates import TEMPLATES as _seed_tpls
        _catchup = 0
        for tpl_name, tpl_type, tpl_prompt, _tpl_inputs in _seed_tpls:
            row = conn.execute(
                "SELECT LENGTH(prompt_text) as len FROM report_templates WHERE report_type = ?", (tpl_type,)
            ).fetchone()
            if row and row["len"] and row["len"] < 200 and len(tpl_prompt) > 200:
                conn.execute("UPDATE report_templates SET prompt_text = ? WHERE report_type = ?", (tpl_prompt, tpl_type))
                _catchup += 1
        if _catchup:
            conn.commit()
            logger.info("Catch-up: updated %d new template prompts to rich format", _catchup)
        conn.close()
        return

    RICH_PROMPTS = {
        "pro_skater": """You are an elite hockey scouting director writing a professional scouting report on a skater (forward or defense). Your job is to turn structured stats and notes into a clear, honest report that a GM and head coach can trust for real decisions.

Use only the information provided in the input JSON. If a metric or behavior is not in the data, do not guess or infer it. It is better to say "DATA NOT AVAILABLE" than to fabricate.

Use coach and scout language, not BI jargon. Explain what the numbers mean in terms of habits, strengths, and risks: forecheck pressure, pace, transition, defending in space, board play, net-front, special teams usage, etc.

You must produce the following sections, in this exact order and format:

EXECUTIVE_SUMMARY:
[2-3 short paragraphs. Identity, current level, playing style, and clear projection (e.g., "middle-six play-driving winger at the OHL level with penalty-kill upside"). Include a one-line risk assessment.]

KEY_NUMBERS:
[6-10 bullet points. Each bullet is "* metric - brief context". Only use metrics present in the input.]

STRENGTHS:
[3-4 titled strengths. For each: a title on its own line, then 2-3 sentences explaining what the player does and how it shows up in games, tying back to stats/notes.]

DEVELOPMENT_AREAS:
[3 specific, trainable development areas, framed as behaviors, each with 2-3 sentences.]

DEVELOPMENT_PRIORITIES:
[3-5 bullet points. Each bullet: name of priority, current level (if available), 12-week target, and coaching focus.]

ADVANCEMENT_TRIGGERS:
[3-5 bullet points describing metric or behavior thresholds that would justify a promotion or bigger role, using only available metrics.]

ROLE_FIT:
[1-2 paragraphs on optimal role and system fit, referencing even-strength role, PP/PK usage, tags, and team_identity if provided.]

OPPONENT_CONTEXT:
[Short paragraph on how the player performs relative to typical opposition level given competition_tier and on-ice metrics. If not enough data, state DATA NOT AVAILABLE and explain impact.]

NOTABLE_PERFORMANCES:
[1-2 short game summaries if any game-level notes or standout games are present. If none, state that no specific notable games are identified.]

BOTTOM_LINE:
[1-2 tight paragraphs that synthesize projection, role, time horizon, risk, and key development focus.]

IMPORTANT RULES:
- Use only data and notes provided in the input JSON.
- If a field is null or a metric is missing, do not invent values or describe behaviors based on it.
- Do not talk about "the JSON" or "the data"; just write the report as if you are the scout.
- Format exactly as specified above. No extra sections.
- Do not use markdown formatting. Do not wrap in code blocks.
- Use plain text with the section keys in ALL_CAPS followed by a colon on their own line, then the content.""",

        "unified_prospect": """You are an elite hockey scouting director writing a comprehensive prospect evaluation report. This report is used by GMs and directors of player development to make draft, trade, and roster decisions.

Use only the information provided in the input JSON. Never fabricate stats or observations.

Produce these sections in order:

EXECUTIVE_SUMMARY:
[3-4 paragraphs. Full prospect identity, projection ceiling/floor, NHL timeline, and comparison player type (not specific name comparisons).]

SCOUTING_GRADES:
[Grade each on a 20-80 scale using only observed/provided data. Format: "* Category: Grade — explanation". Categories: Skating, Puck Skills, Hockey Sense, Compete/Physical, Shooting, Defensive Play.]

PROJECTION:
[2 paragraphs. Realistic ceiling, floor, most likely outcome. NHL timeline. What must develop for ceiling.]

DRAFT_POSITIONING:
[1 paragraph on where this player fits in their draft class based on data provided. If draft info not available, say so.]

DEVELOPMENT_PATHWAY:
[3-5 prioritized development steps with timelines and measurable targets.]

RISK_ASSESSMENT:
[1-2 paragraphs identifying the key risks to this player reaching their projection.]

BOTTOM_LINE:
[2-3 sentences. Final verdict on whether to draft/acquire/develop.]

IMPORTANT: Use only provided data. Say "DATA NOT AVAILABLE" for missing metrics. No markdown.""",

        "goalie": """You are an elite goaltending scout writing a professional goalie evaluation report. Your audience is a GM and goaltending coach making real roster and development decisions.

Use only the information provided. Never fabricate observations.

Produce these sections:

EXECUTIVE_SUMMARY:
[2-3 paragraphs. Goalie identity, style (butterfly/hybrid/stand-up tendencies), current level, projection.]

KEY_NUMBERS:
[6-8 bullets. GAA, SV%, GSAX, workload, high-danger save %, rebound control — only metrics present in input.]

TECHNICAL_ASSESSMENT:
[3-4 titled areas: Positioning, Rebound Control, Movement/Recovery, Puck Handling. 2-3 sentences each from data/notes.]

MENTAL_GAME:
[1-2 paragraphs on composure, game management, bounce-back ability — from scout/coach notes only.]

DEVELOPMENT_AREAS:
[2-3 specific, trainable areas with coaching recommendations.]

WORKLOAD_ANALYSIS:
[1 paragraph on games played, shots faced, fatigue indicators if data available.]

ROLE_FIT:
[1 paragraph on starter/backup/tandem fit, system compatibility.]

BOTTOM_LINE:
[1-2 tight paragraphs. Investment verdict, timeline, risk.]

IMPORTANT: Use only provided data. No markdown formatting.""",

        "game_decision": """You are a hockey analytics coach generating a single-game decision report. This report helps coaches make real-time lineup and deployment decisions based on one game's data.

Produce these sections:

GAME_SUMMARY:
[2-3 sentences. Score, opponent, date, home/away, key storyline.]

PLAYER_GRADES:
[For each player in the input: "* Player Name: Grade (A/B/C/D/F) — 1-sentence justification based on stats/notes."]

DEPLOYMENT_NOTES:
[3-5 bullets on what worked and what didn't in terms of line combinations, matchups, and special teams usage.]

ADJUSTMENTS:
[2-3 specific tactical adjustments recommended for next game based on this game's data.]

STANDOUT_PERFORMERS:
[1-2 paragraphs highlighting players who exceeded or fell below expectations.]

IMPORTANT: Use only provided game data. No fabrication.""",

        "season_intelligence": """You are a hockey intelligence analyst producing a season-level player assessment. This comprehensive report synthesizes an entire season of data into actionable intelligence.

Produce these sections:

SEASON_OVERVIEW:
[2-3 paragraphs. Season narrative arc — how the player's performance evolved across the season.]

STATISTICAL_PROFILE:
[8-12 bullets covering all major stat categories from the input. Include per-game rates where applicable.]

TREND_ANALYSIS:
[2-3 paragraphs on performance trends: early vs late season, home vs away, vs strong vs weak opponents — if data supports it.]

STRENGTHS_CONFIRMED:
[3-4 strengths that held up across the full season.]

CONCERNS_IDENTIFIED:
[2-3 concerns that emerged or persisted across the season.]

OFFSEASON_PRIORITIES:
[3-5 prioritized development areas for the offseason.]

CONTRACT_CONTEXT:
[1 paragraph on value assessment if contract info available, otherwise skip.]

BOTTOM_LINE:
[1-2 paragraphs. Full season verdict and outlook for next season.]

IMPORTANT: Season-level analysis only. Use provided data.""",

        "operations": """You are a hockey operations director producing a comprehensive operational assessment of a player. This report informs cap management, roster construction, and long-term planning decisions.

Produce these sections:

OPERATIONAL_SUMMARY:
[2-3 paragraphs. Player's operational value — cap hit, contract status, role, replaceability.]

ROSTER_VALUE:
[Assess the player's value relative to their cost. Include comparable players at similar cost if data supports it.]

DEPLOYMENT_EFFICIENCY:
[Analyze ice time usage, special teams impact, situational deployment. Are they being used optimally?]

ASSET_MANAGEMENT:
[1-2 paragraphs on trade value, extension considerations, or asset protection.]

RISK_FACTORS:
[2-3 bullets on operational risks: injury history, age curve, declining metrics.]

RECOMMENDATION:
[Clear operational recommendation: extend, trade, hold, buyout — with justification.]

IMPORTANT: Use only provided data. This is an operations report, not a scouting report.""",

        "team_identity": """You are a hockey analytics consultant producing a Team Identity Card. This defines how a team plays, what kind of players fit their system, and how opponents should prepare.

Produce these sections:

TEAM_IDENTITY:
[2-3 paragraphs. Playing style, system, pace, structure. What makes this team who they are.]

SYSTEM_DETAILS:
[Forecheck structure, breakout patterns, neutral zone play, defensive zone coverage — from available data.]

PLAYER_ARCHETYPE_FIT:
[What type of player thrives in this system? Speed, size, skill, compete profiles.]

SPECIAL_TEAMS_IDENTITY:
[PP and PK structures, tendencies, effectiveness.]

KEY_PERSONNEL:
[2-3 players who define this team's identity, with brief explanations.]

VULNERABILITIES:
[2-3 systemic weaknesses opponents could exploit.]

IMPORTANT: Use only provided team data and observations.""",

        "opponent_gameplan": """You are a hockey coaching staff member preparing an opponent game plan. This report provides tactical preparation for an upcoming game.

Produce these sections:

OPPONENT_OVERVIEW:
[2-3 paragraphs. Who they are, how they play, recent form.]

KEY_MATCHUPS:
[3-4 specific matchups to target or avoid, with reasoning.]

FORECHECK_PLAN:
[How to forecheck against this opponent based on their breakout tendencies.]

DEFENSIVE_KEYS:
[3-4 defensive priorities against this opponent's attack patterns.]

SPECIAL_TEAMS_PREP:
[PP and PK adjustments specific to this opponent.]

LINE_MATCHING:
[Recommended line matching strategy.]

GAME_KEYS:
[3-5 bullet points — "Win the game if we do these things."]

IMPORTANT: Use only provided opponent data.""",

        "agent_pack": """You are a hockey agent's intelligence analyst producing a player marketing and positioning document. This report helps agents negotiate contracts, seek trades, and position players for advancement.

Produce these sections:

PLAYER_PROFILE:
[2-3 paragraphs. Professional biography, playing identity, brand.]

STATISTICAL_CASE:
[6-10 bullets highlighting the most marketable stats. Frame positively but honestly.]

MARKET_POSITION:
[1-2 paragraphs on where this player fits in the market. Comparable contracts if data available.]

TALKING_POINTS:
[5-7 bullet points an agent could use in negotiations.]

DEVELOPMENT_TRAJECTORY:
[1-2 paragraphs on growth trend and future value.]

RISK_MITIGATION:
[Address likely counter-arguments from teams with data-backed responses.]

RECOMMENDATION:
[Clear positioning strategy: what to ask for, what to accept, timeline.]

IMPORTANT: This is advocacy writing backed by data. Be honest but present the best case.""",

        "development_roadmap": """You are a Director of Player Development creating a structured development roadmap for a player. This is used by development coaches, skills coaches, and the player themselves.

Produce these sections:

CURRENT_ASSESSMENT:
[2-3 paragraphs. Where the player is right now — strengths, gaps, readiness level.]

DEVELOPMENT_PILLARS:
[3-5 core development areas, each with: Title, Current Level, Target Level, Timeline, Specific Drills/Focus.]

30_DAY_PLAN:
[Specific weekly focus areas for the next 30 days.]

90_DAY_PLAN:
[Monthly milestones for the next 90 days.]

SEASON_GOALS:
[3-5 measurable season-end goals.]

MEASUREMENT_FRAMEWORK:
[How progress will be tracked — specific metrics, video review cadence, testing schedule.]

SUPPORT_TEAM:
[Recommended support: skills coach focus, mental performance, nutrition/strength if applicable.]

BOTTOM_LINE:
[1-2 paragraphs. Is this player developing on track? What's the biggest unlock?]

IMPORTANT: Be specific and actionable. Every recommendation should be trainable.""",

        "family_card": """You are a hockey advisor producing a Player/Family Card. This is a clear, accessible report designed for the player and their family to understand development status, opportunities, and next steps.

Write in accessible language — no insider jargon without explanation.

Produce these sections:

PLAYER_SNAPSHOT:
[2-3 paragraphs in plain language. Who the player is, what they do well, where they're headed.]

SEASON_HIGHLIGHTS:
[5-7 bullets of positive achievements and milestones.]

AREAS_FOR_GROWTH:
[2-3 development areas framed constructively — not "weaknesses" but "next steps."]

PATHWAY_OPTIONS:
[1-2 paragraphs on realistic next steps: leagues, teams, tryouts, showcases.]

WHAT_SCOUTS_SEE:
[1 paragraph translating scout perspective into family-friendly language.]

ACTION_ITEMS:
[3-5 specific things the player can work on this offseason.]

IMPORTANT: Family-friendly language. Honest but encouraging. No jargon.""",

        "line_chemistry": """You are a hockey analytics specialist analyzing line chemistry. This report assesses how specific player combinations perform together.

Produce these sections:

LINE_OVERVIEW:
[1-2 paragraphs. The line combination being analyzed, context, ice time together.]

CHEMISTRY_METRICS:
[6-8 bullets on combined metrics: Corsi, expected goals, zone entries/exits, shot generation when together vs apart.]

ROLE_COMPLEMENTARITY:
[1-2 paragraphs on how each player's strengths complement the others.]

OPTIMAL_DEPLOYMENT:
[When and how to deploy this line — situations, matchups, game states.]

ALTERNATIVES:
[1-2 alternative combinations worth testing, with reasoning.]

VERDICT:
[Keep, adjust, or break up — with justification.]

IMPORTANT: Use only provided line combination data.""",

        "st_optimization": """You are a special teams analyst optimizing power play and penalty kill units. This report is for coaching staff to improve special teams deployment.

Produce these sections:

POWER_PLAY_ASSESSMENT:
[2-3 paragraphs. Current PP structure, effectiveness, personnel deployment.]

PP_UNIT_RECOMMENDATIONS:
[Specific unit configurations with roles for each player. PP1 and PP2.]

PENALTY_KILL_ASSESSMENT:
[2-3 paragraphs. Current PK structure, effectiveness, aggressive vs passive tendencies.]

PK_UNIT_RECOMMENDATIONS:
[Specific unit configurations. PK1 and PK2.]

PERSONNEL_CHANGES:
[2-3 specific changes to try, with expected impact.]

PRACTICE_FOCUS:
[3-4 practice drills or situations to work on.]

IMPORTANT: Use only provided special teams data.""",

        "trade_target": """You are a hockey operations analyst evaluating a player as a trade or acquisition target. This report helps GMs decide whether to pursue a player and what to offer.

Produce these sections:

TARGET_PROFILE:
[2-3 paragraphs. Who the player is, why they might be available, what they bring.]

FIT_ASSESSMENT:
[1-2 paragraphs on how this player fits your team's needs, system, and culture.]

STATISTICAL_EVALUATION:
[6-8 key stats with context on whether performance is sustainable.]

COST_ANALYSIS:
[Contract details, cap impact, term remaining. What's fair value in trade assets?]

RISK_FACTORS:
[2-3 risks: age, injury, declining production, character concerns — only from data.]

COMPARABLE_DEALS:
[If data available, 2-3 similar trades for reference.]

RECOMMENDATION:
[Pursue aggressively, monitor, or pass — with clear reasoning and suggested offer framework.]

IMPORTANT: Use only provided data. Be objective.""",

        "draft_comparative": """You are a draft analyst comparing players within a draft class. This report helps scouting directors rank and compare prospects.

Produce these sections:

CLASS_OVERVIEW:
[1-2 paragraphs. Draft year, depth, notable trends in this class.]

PLAYER_COMPARISONS:
[For each player in input: 1 paragraph assessment with grade. Then rank all players.]

TIER_RANKINGS:
[Group players into tiers: Elite, First Round, Second Round, Later Rounds, Undraftable.]

POSITIONAL_BREAKDOWN:
[Best available by position: Centers, Wingers, Defensemen, Goalies.]

SLEEPER_PICKS:
[1-2 players who might be undervalued based on the data.]

BUST_RISKS:
[1-2 players whose draft stock may not match production.]

IMPORTANT: Compare only players provided in the input data.""",

        "season_progress": """You are a player development coach writing a mid-season or end-of-season progress report. This tracks a player's development against previously set goals.

Produce these sections:

PROGRESS_SUMMARY:
[2-3 paragraphs. Overall trajectory — on track, ahead, behind. Key narrative.]

GOAL_TRACKING:
[For each previously set goal: Goal, Target, Current Status, On Track (Yes/No/Partially).]

STATISTICAL_PROGRESSION:
[Compare current stats to last season / preseason targets. Highlight improvements and declines.]

BEHAVIORAL_OBSERVATIONS:
[2-3 paragraphs from coach/scout notes on habits, compete, leadership growth.]

ADJUSTED_PRIORITIES:
[Any development priorities that should change based on progress.]

NEXT_STEPS:
[3-5 specific focus areas for the remainder of the season or offseason.]

IMPORTANT: Track against provided goals. Be honest about gaps.""",

        "practice_plan": """You are a hockey coaching specialist generating a structured practice plan based on team needs and recent game data.

Produce these sections:

PRACTICE_OVERVIEW:
[Date, duration, focus areas, intensity level.]

WARM_UP:
[10-15 minutes. Skating, puck handling drills.]

SKILL_STATIONS:
[2-3 stations, 10 minutes each. Specific drills tied to identified needs.]

TACTICAL_WORK:
[15-20 minutes. Systems work — forecheck, breakout, PP/PK based on what needs improvement.]

COMPETE_DRILLS:
[10-15 minutes. Battle drills, small-area games tied to development areas.]

SCRIMMAGE_SCENARIOS:
[Situational scrimmage setups: down 1 goal, PP/PK reps, last-minute scenarios.]

COOL_DOWN:
[5 minutes. Light skating, team communication.]

COACHING_NOTES:
[Key teaching points for each segment. What to watch for.]

IMPORTANT: Tie every drill to an identified team or player need from the input.""",

        "playoff_series": """You are a hockey coaching staff member preparing a comprehensive playoff series preparation report.

Produce these sections:

SERIES_OVERVIEW:
[2-3 paragraphs. Matchup preview, regular season head-to-head, keys to the series.]

OPPONENT_TENDENCIES:
[3-5 key tendencies: how they play 5v5, on the PP, on the PK, in close games.]

MATCHUP_PLAN:
[Specific line matching strategy. Who to match against their top line, where to create advantages.]

SPECIAL_TEAMS_STRATEGY:
[PP adjustments for this opponent. PK adjustments. Faceoff strategy.]

GOALTENDING_ASSESSMENT:
[1 paragraph on their goaltender(s) — weaknesses to exploit, tendencies.]

GAME_1_LINEUP:
[Recommended lineup, line combinations, D pairs, PP/PK units for Game 1.]

SERIES_KEYS:
[5-7 bullet points — "Win the series if we do these things."]

IMPORTANT: Use only provided data about both teams.""",

        "goalie_tandem": """You are a goaltending consultant analyzing a goalie tandem to optimize workload management and deployment.

Produce these sections:

TANDEM_OVERVIEW:
[1-2 paragraphs. Both goalies' profiles, current usage split, overall effectiveness.]

INDIVIDUAL_ASSESSMENTS:
[For each goalie: 1-2 paragraphs on strengths, weaknesses, workload tolerance.]

WORKLOAD_ANALYSIS:
[Optimal start split (e.g., 60/40, 55/45). Back-to-back strategy. Rest patterns.]

SITUATIONAL_DEPLOYMENT:
[When to start Goalie A vs B: home/away, opponent strength, schedule density.]

PERFORMANCE_TRIGGERS:
[Metrics that should trigger a start change: SV% over last 5, goals against trends, fatigue indicators.]

DEVELOPMENT_CONSIDERATIONS:
[If one goalie is younger/developing, how to balance development with winning.]

RECOMMENDATION:
[Clear tandem strategy for the rest of the season.]

IMPORTANT: Use only provided goaltender data.""",
    }

    updated = 0
    for report_type, prompt_text in RICH_PROMPTS.items():
        result = conn.execute(
            "UPDATE report_templates SET prompt_text = ? WHERE report_type = ?",
            (prompt_text, report_type),
        )
        if result.rowcount > 0:
            updated += 1

    conn.commit()
    conn.close()
    if updated:
        logger.info("Migrated %d report template prompts to rich format", updated)


def seed_leagues():
    """Seed the leagues reference table with professional, junior, and college leagues."""
    conn = get_db()
    count = conn.execute("SELECT COUNT(*) FROM leagues").fetchone()[0]
    if count > 0:
        conn.close()
        return

    leagues = [
        # Professional
        ("AHL", "American Hockey League", "USA", "professional", 1),
        ("ECHL", "ECHL", "USA", "professional", 2),
        ("SPHL", "Southern Professional Hockey League", "USA", "professional", 3),
        ("PWHL", "Professional Women's Hockey League", "Canada", "professional", 4),
        # Major Junior (CHL)
        ("OHL", "Ontario Hockey League", "Canada", "major_junior", 10),
        ("QMJHL", "Quebec Major Junior Hockey League", "Canada", "major_junior", 11),
        ("WHL", "Western Hockey League", "Canada", "major_junior", 12),
        # Junior A
        ("BCHL", "British Columbia Hockey League", "Canada", "junior_a", 20),
        ("AJHL", "Alberta Junior Hockey League", "Canada", "junior_a", 21),
        ("SJHL", "Saskatchewan Junior Hockey League", "Canada", "junior_a", 22),
        ("MJHL", "Manitoba Junior Hockey League", "Canada", "junior_a", 23),
        ("USHL", "United States Hockey League", "USA", "junior_a", 24),
        ("OJHL", "Ontario Junior Hockey League", "Canada", "junior_a", 25),
        ("CCHL", "Central Canada Hockey League", "Canada", "junior_a", 26),
        ("NOJHL", "Northern Ontario Junior Hockey League", "Canada", "junior_a", 27),
        ("MHL", "Maritime Hockey League", "Canada", "junior_a", 28),
        ("GOHL", "Greater Ontario Hockey League", "Canada", "junior_a", 29),
        ("NAHL", "North American Hockey League", "USA", "junior_a", 30),
        # Junior B
        ("KIJHL", "Kootenay International Junior Hockey League", "Canada", "junior_b", 40),
        ("PJHL", "Provincial Junior Hockey League", "Canada", "junior_b", 41),
        ("VIJHL", "Vancouver Island Junior Hockey League", "Canada", "junior_b", 42),
        # College / Other
        ("NCAA", "National Collegiate Athletic Association", "USA", "college", 60),
        ("USHS", "US High School", "USA", "high_school", 70),
        ("AAA", "AAA Minor Hockey", "Canada", "minor", 80),
    ]
    for abbr, name, country, level, sort in leagues:
        conn.execute(
            "INSERT INTO leagues (id, abbreviation, name, country, level, sort_order) VALUES (?, ?, ?, ?, ?, ?)",
            (str(uuid.uuid4()), abbr, name, country, level, sort),
        )
    conn.commit()
    conn.close()
    logger.info("Seeded %d leagues", len(leagues))


def migrate_leagues():
    """Upsert leagues for existing databases — adds missing leagues, fixes level/sort_order."""
    conn = get_db()

    # Handle GOJHL → GOHL rebrand (abbreviation changed)
    old_gojhl = conn.execute("SELECT id FROM leagues WHERE abbreviation = 'GOJHL'").fetchone()
    if old_gojhl:
        conn.execute(
            "UPDATE leagues SET abbreviation = 'GOHL', name = 'Greater Ontario Hockey League', level = 'junior_a', sort_order = 29 WHERE abbreviation = 'GOJHL'"
        )
        # Also update any teams that reference the old abbreviation
        conn.execute("UPDATE teams SET league = 'GOHL' WHERE league = 'GOJHL'")
        logger.info("Renamed GOJHL → GOHL in leagues and teams")

    # Canonical league data (must match seed_leagues above)
    leagues = [
        # Professional
        ("AHL", "American Hockey League", "USA", "professional", 1),
        ("ECHL", "ECHL", "USA", "professional", 2),
        ("SPHL", "Southern Professional Hockey League", "USA", "professional", 3),
        ("PWHL", "Professional Women's Hockey League", "Canada", "professional", 4),
        # Major Junior (CHL)
        ("OHL", "Ontario Hockey League", "Canada", "major_junior", 10),
        ("QMJHL", "Quebec Major Junior Hockey League", "Canada", "major_junior", 11),
        ("WHL", "Western Hockey League", "Canada", "major_junior", 12),
        # Junior A
        ("BCHL", "British Columbia Hockey League", "Canada", "junior_a", 20),
        ("AJHL", "Alberta Junior Hockey League", "Canada", "junior_a", 21),
        ("SJHL", "Saskatchewan Junior Hockey League", "Canada", "junior_a", 22),
        ("MJHL", "Manitoba Junior Hockey League", "Canada", "junior_a", 23),
        ("USHL", "United States Hockey League", "USA", "junior_a", 24),
        ("OJHL", "Ontario Junior Hockey League", "Canada", "junior_a", 25),
        ("CCHL", "Central Canada Hockey League", "Canada", "junior_a", 26),
        ("NOJHL", "Northern Ontario Junior Hockey League", "Canada", "junior_a", 27),
        ("MHL", "Maritime Hockey League", "Canada", "junior_a", 28),
        ("GOHL", "Greater Ontario Hockey League", "Canada", "junior_a", 29),
        ("NAHL", "North American Hockey League", "USA", "junior_a", 30),
        # Junior B
        ("KIJHL", "Kootenay International Junior Hockey League", "Canada", "junior_b", 40),
        ("PJHL", "Provincial Junior Hockey League", "Canada", "junior_b", 41),
        ("VIJHL", "Vancouver Island Junior Hockey League", "Canada", "junior_b", 42),
        # College / Other
        ("NCAA", "National Collegiate Athletic Association", "USA", "college", 60),
        ("USHS", "US High School", "USA", "high_school", 70),
        ("AAA", "AAA Minor Hockey", "Canada", "minor", 80),
    ]
    inserted = 0
    updated = 0
    for abbr, name, country, level, sort in leagues:
        existing = conn.execute(
            "SELECT id, level, sort_order, name FROM leagues WHERE abbreviation = ?", (abbr,)
        ).fetchone()
        if existing:
            # Update if level, sort_order, or name changed
            if existing["level"] != level or existing["sort_order"] != sort or existing["name"] != name:
                conn.execute(
                    "UPDATE leagues SET name = ?, level = ?, sort_order = ?, country = ? WHERE abbreviation = ?",
                    (name, level, sort, country, abbr),
                )
                updated += 1
        else:
            conn.execute(
                "INSERT INTO leagues (id, abbreviation, name, country, level, sort_order) VALUES (?, ?, ?, ?, ?, ?)",
                (str(uuid.uuid4()), abbr, name, country, level, sort),
            )
            inserted += 1
    conn.commit()
    conn.close()
    if inserted or updated:
        logger.info("League migration: %d inserted, %d updated", inserted, updated)


def seed_teams():
    """Seed reference teams for GOHL (all conferences)."""
    conn = get_db()
    count = conn.execute("SELECT COUNT(*) FROM teams WHERE org_id = '__global__'").fetchone()[0]
    if count > 0:
        conn.close()
        return

    # Ensure __global__ org exists for reference data (FK constraint)
    existing = conn.execute("SELECT id FROM organizations WHERE id = '__global__'").fetchone()
    if not existing:
        conn.execute(
            "INSERT INTO organizations (id, name) VALUES ('__global__', 'Global Reference Data')"
        )
        conn.commit()

    gojhl_teams = [
        # Western Conference
        ("Chatham Maroons", "GOHL", "Chatham", "CM"),
        ("Leamington Flyers", "GOHL", "Leamington", "LF"),
        ("LaSalle Vipers", "GOHL", "LaSalle", "LV"),
        ("London Nationals", "GOHL", "London", "LN"),
        ("Komoka Kings", "GOHL", "Komoka", "KK"),
        ("Strathroy Rockets", "GOHL", "Strathroy", "SR"),
        ("St. Thomas Stars", "GOHL", "St. Thomas", "STS"),
        ("St. Marys Lincolns", "GOHL", "St. Marys", "STM"),
        ("Sarnia Legionnaires", "GOHL", "Sarnia", "SAR"),
        # Midwestern Conference
        ("Brantford Bandits", "GOHL", "Brantford", "BB"),
        ("Cambridge Redhawks", "GOHL", "Cambridge", "CAM"),
        ("Elmira Sugar Kings", "GOHL", "Elmira", "ESK"),
        ("KW Siskins", "GOHL", "Kitchener", "KWS"),
        ("Listowel Cyclones", "GOHL", "Listowel", "LC"),
        ("Stratford Warriors", "GOHL", "Stratford", "SW"),
        ("Ayr Centennials", "GOHL", "Ayr", "AC"),
        # Golden Horseshoe Conference
        ("Caledonia Corvairs", "GOHL", "Caledonia", "CC"),
        ("Hamilton Kilty B's", "GOHL", "Hamilton", "HKB"),
        ("Pelham Panthers", "GOHL", "Pelham", "PP"),
        ("St. Catharines Falcons", "GOHL", "St. Catharines", "SCF"),
        ("Thorold Blackhawks", "GOHL", "Thorold", "TB"),
        ("Niagara Falls Canucks", "GOHL", "Niagara Falls", "NFC"),
        # Northern Conference
        ("Caledon Bombers", "GOHL", "Caledon", "CB"),
    ]
    for name, league, city, abbr in gojhl_teams:
        conn.execute(
            "INSERT INTO teams (id, org_id, name, league, city, abbreviation) VALUES (?, ?, ?, ?, ?, ?)",
            (str(uuid.uuid4()), "__global__", name, league, city, abbr),
        )
    conn.commit()
    conn.close()
    logger.info("Seeded %d reference teams", len(gojhl_teams))


def seed_drills():
    """Seed 44 original hockey drills across 13 categories with age-appropriate tagging."""
    conn = get_db()
    count = conn.execute("SELECT COUNT(*) FROM drills").fetchone()[0]
    if count > 0:
        conn.close()
        return

    ALL_AGES = '["U8","U10","U12","U14","U16_U18","JUNIOR_COLLEGE_PRO"]'
    U10_UP = '["U10","U12","U14","U16_U18","JUNIOR_COLLEGE_PRO"]'
    U12_UP = '["U12","U14","U16_U18","JUNIOR_COLLEGE_PRO"]'
    U14_UP = '["U14","U16_U18","JUNIOR_COLLEGE_PRO"]'
    U16_UP = '["U16_U18","JUNIOR_COLLEGE_PRO"]'
    JR_PLUS = '["JUNIOR_COLLEGE_PRO"]'

    # (name, category, description, coaching_points, setup, duration_min, players_needed, ice_surface, equipment, age_levels, tags, skill_focus, intensity, concept_id)
    drills = [
        # ── WARM UP (4) ──────────────────────────────────────────
        ("Dynamic Skating Warm-Up", "warm_up",
         "Players skate through a series of dynamic movements across the ice: high knees, butt kicks, carioca, side shuffles, and forward-to-backward transitions. Two laps of each movement from goal line to goal line.",
         "Focus on full range of motion. Keep heads up. Gradually increase tempo each lap. Watch for lazy crossovers.",
         "Full ice. No equipment needed. Players line up on goal line.",
         8, 0, "full", "None",
         ALL_AGES, '["skating","warm_up","agility"]', "skating", "low", "dynamic_skating_warmup"),

        ("Partner Passing Circuit", "warm_up",
         "Players pair up and skate down the ice passing back and forth. On the whistle, they change from forehand to backhand, then saucer passes, then one-touch passes. Continuous movement.",
         "Passes should be tape-to-tape. Receivers show a target. Keep feet moving while passing — no standing still.",
         "Full ice. One puck per pair. Players pair up on goal line.",
         8, 0, "full", "Pucks",
         ALL_AGES, '["passing","warm_up","puck_control"]', "passing", "low", "partner_passing_warmup"),

        ("Puck Handling Relay", "warm_up",
         "Teams of 4-5 line up on the goal line. First player weaves through 5 cones to the far blue line and back, then tags the next player. Race format — losing team does push-ups.",
         "Head up through the cones. Tight turns around each cone. Emphasize quick hands, not just speed.",
         "Half ice. 5 cones per lane. 2-4 lanes depending on team size.",
         8, 8, "half", "Cones, pucks",
         ALL_AGES, '["puck_handling","warm_up","stickhandling","races"]', "stickhandling", "medium", "puck_handling_relay"),

        ("Edge Work Warm-Up", "warm_up",
         "Players skate figure-8 patterns around the face-off circles using inside and outside edges. Progress from two feet to one foot, then add crossovers. Alternate clockwise and counter-clockwise.",
         "Knees bent, weight on the balls of the feet. Deep edges — lean into the turn. Alternate direction every 30 seconds.",
         "Full ice. Use all five face-off circles. Players spread out evenly.",
         7, 0, "full", "None",
         ALL_AGES, '["skating","warm_up","agility","edges"]', "skating", "low", "edge_work_warmup"),

        # ── SKATING (5) ──────────────────────────────────────────
        ("Crossover Figure-8", "skating",
         "Players skate figure-8 patterns around two cones set 30 feet apart. Focus on deep crossovers, knee bend, and weight transfer. Progress to adding a puck, then to tight turns with acceleration out.",
         "Inside foot drives under. Outside foot crosses over with power. Keep shoulders level — don't lean with the upper body. Explode out of the turn.",
         "Half ice. Two cones per station, 30 feet apart. 4-6 stations.",
         10, 0, "half", "Cones",
         ALL_AGES, '["skating","crossovers","agility"]', "skating", "medium", "crossover_figure8"),

        ("Transition Skating Series", "skating",
         "Players skate forward to the blue line, transition to backward at the blue line, skate backward to the red line, pivot to forward at the red line, and sprint to the far blue line. Continuous reps.",
         "Open hips on transitions — don't spin. Keep speed through the pivot. Head and eyes up at all times. Drive with the legs, not the upper body.",
         "Full ice. Players go in waves of 3-4. Whistle starts each wave.",
         10, 0, "full", "None",
         U10_UP, '["skating","transition","pivots","agility"]', "skating", "high", "transition_skating"),

        ("Power Skating Stride Circuit", "skating",
         "Four stations: (1) Full stride sprints along the boards, (2) C-cuts and power pulls along the blue line, (3) Forward crossover serpentine through cones, (4) Backward striding with stick on knees for posture. 90 seconds per station, 30 seconds rest.",
         "Full extension on every stride. Recovery leg comes back under the body. Arms drive forward, not side to side. Chest up.",
         "Full ice. Cones for serpentine station. Players rotate on whistle.",
         12, 0, "full", "Cones",
         U10_UP, '["skating","stride","power","conditioning"]', "skating", "high", "power_stride_circuit"),

        ("Tight Turn Agility Course", "skating",
         "Set up 8 cones in a zigzag pattern across the neutral zone. Players weave through at speed, making tight turns around each cone. Alternate between forward and backward on each rep.",
         "Load the outside leg before each turn. Stay low through the turn — hips drop. Quick feet around the cone, then explode. Challenge: add a puck.",
         "Neutral zone only. 8 cones in zigzag. Players go one at a time.",
         10, 0, "half", "Cones",
         ALL_AGES, '["skating","agility","tight_turns"]', "skating", "medium", "tight_turn_agility"),

        ("Backward-to-Forward Pivots", "skating",
         "Players skate backward from the goal line. On the coach's signal (whistle or point), they pivot to forward and sprint 3 strides, then return to backward skating. Repeat across the full ice.",
         "Open hips to the direction the coach points. Stay low through the pivot — don't stand up. First three strides after the pivot should be explosive.",
         "Full ice. Coach stands at center ice with a whistle. Players spread across the width.",
         8, 0, "full", "None",
         U10_UP, '["skating","pivots","defensive","backward_skating"]', "skating", "medium", "backward_forward_pivots"),

        # ── PASSING (4) ──────────────────────────────────────────
        ("Three-Line Passing Drill", "passing",
         "Three lines at one end. Center line carries the puck up ice. Wings fill the lanes. Center passes to left wing, left wing passes to right wing, right wing passes back to center for a shot. Reset and go the other direction.",
         "Head up before passing — look off the defender. Hard, flat passes — no floaters. Receivers give a target with the blade. Time the pass to hit the player in stride.",
         "Full ice. Three lines at one end, one puck per group.",
         10, 6, "full", "Pucks",
         U10_UP, '["passing","offensive","3_on_0","shooting"]', "passing", "medium", "three_line_passing"),

        ("Tape-to-Tape Relay Race", "passing",
         "Two teams line up in columns 20 feet apart. First player passes to the second, second passes back, pattern continues down the line. Last player skates the puck back to the front. First team to complete 3 rotations wins.",
         "Passes must be on the tape — any missed pass costs time. No slapping at the puck. Quick hands, quick release. Face your target before passing.",
         "Half ice. Two teams in columns, 20 feet apart.",
         8, 8, "half", "Pucks",
         ALL_AGES, '["passing","relay_races","compete"]', "passing", "medium", "tape_to_tape_relay"),

        ("Drop Pass Options Drill", "passing",
         "Three forwards enter the zone. The puck carrier has three options: (1) drop pass to the trailing player, (2) pass to the weak-side wing, (3) carry and shoot. Coach calls the option. Progress to letting the carrier read and decide.",
         "Drop pass: leave it dead, don't push it back. Trailing player should be 2-3 stick lengths behind. Weak-side wing drives wide then cuts to the net. Sell the fake before passing.",
         "Half ice. Three forwards per rep. Coach at center ice calls options.",
         12, 6, "half", "Pucks",
         U12_UP, '["passing","offensive","zone_entry","decision_making"]', "passing", "medium", "drop_pass_options"),

        ("Saucer Pass Progression", "passing",
         "Partners face each other with a stick laid flat between them (simulating a passing lane obstacle). Progress through: (1) basic saucer pass, (2) moving saucer pass while skating, (3) saucer pass to a player in stride, (4) saucer pass off the boards.",
         "Spin the puck — roll the wrists on release. The puck should land flat on the receiver's blade. Start close together and gradually increase distance. Wrist position is key — cup the puck.",
         "Half ice. Partners 15-20 feet apart with a stick on the ice between them.",
         10, 0, "half", "Pucks, extra sticks for obstacles",
         U10_UP, '["passing","saucer_pass","skill_development"]', "passing", "low", "saucer_pass_progression"),

        # ── SHOOTING (4) ─────────────────────────────────────────
        ("Quick Release from the Slot", "shooting",
         "Players line up at the top of the circles. Coach feeds a pass from behind the net. Player receives in the slot and must get the shot off within 2 seconds — catch and release. Alternate sides.",
         "Get the puck to the shooting position fast — don't stickhandle. Weight transfer from back foot to front foot. Pick your spot before you receive. Aim for corners, not center mass.",
         "One zone. Coach behind the net. Players in two lines at the hash marks.",
         10, 1, "quarter", "Pucks",
         U10_UP, '["shooting","offensive","quick_release"]', "shooting", "medium", "quick_release_slot"),

        ("One-Timer Setup Drill", "shooting",
         "Two lines — one at the half-wall, one at the top of the circle. Half-wall player passes across to the shooter at the top of the circle for a one-timer. Rotate lines. Progress to adding a screen in front.",
         "Stick blade open and loaded before the pass arrives. Transfer weight as you swing. Follow through low for accuracy. Timing is everything — start your backswing early.",
         "One zone. Two lines. Goalie in net.",
         12, 4, "quarter", "Pucks",
         U12_UP, '["shooting","one_timer","power_play","offensive"]', "shooting", "high", "one_timer_setup"),

        ("Screen and Tip Drill", "shooting",
         "Defenseman at the point takes a shot. Forward in front of the net works on: (1) screening the goalie, (2) tipping the shot, (3) picking up rebounds. Rotate D shooters and net-front players every 5 reps.",
         "Net-front player: stick on the ice, blade angle to redirect. Don't watch the shot — feel it. Move slightly to create traffic. Rebound position: stick on ice, inside leg loaded.",
         "One zone. D at the point, F in front of net. Goalie in net.",
         12, 4, "quarter", "Pucks",
         U14_UP, '["shooting","screening","tipping","net_front","offensive"]', "shooting", "medium", "screen_and_tip"),

        ("Wrist Shot Accuracy Circuit", "shooting",
         "Four shooting stations around the zone. Each station has a target (water bottle or small cone) on a specific corner of the net. Players take 5 shots per station, tracking how many targets they hit. Rotate after each set.",
         "Pick your target before you shoot. Wrist over the puck for top corner. Roll the wrists for bottom corner. Consistency over power — hit the spot every time.",
         "One zone. Four stations. Targets on net corners. Track hits.",
         10, 4, "quarter", "Pucks, water bottles or targets",
         ALL_AGES, '["shooting","accuracy","skill_development","stations"]', "shooting", "medium", "wrist_shot_accuracy"),

        # ── OFFENSIVE (4) ────────────────────────────────────────
        ("2-on-1 Rush Options", "offensive",
         "Two forwards attack against one defenseman. The puck carrier reads the D: if the D takes away the pass, shoot; if the D takes the lane, pass across for a one-timer. Run from both sides. Add a backchecker for progression.",
         "Puck carrier: attack with speed, force the D to commit. Don't telegraph the pass — eyes on the net. Off-puck player: drive the far post, stick on the ice. D: take away the pass and force the shot.",
         "Full ice. F start at far end. D starts at blue line. Run 2-on-1 both directions.",
         12, 4, "full", "Pucks",
         U12_UP, '["offensive","2_on_1","zone_entry","decision_making","shooting"]', "offensive", "high", "2on1_rush"),

        ("Cycle Low Drill", "offensive",
         "Three forwards set up in the offensive zone. Puck starts down low. F1 retrieves and cycles to F2 along the boards. F2 has options: pass high to F3, reverse to F1, or drive the net. Run continuous for 60 seconds.",
         "Protect the puck on the retrieve — body between the puck and the wall. Timing of support: F2 arrives as F1 is cycling, not before. High F3 reads the play — come down if there's a shooting lane.",
         "Half ice. Three forwards per group. New group every 60 seconds.",
         12, 6, "half", "Pucks",
         U14_UP, '["offensive","cycling","down_low","puck_support","wall_play"]', "offensive", "high", "cycle_low"),

        ("Net-Front Presence Training", "offensive",
         "Forward stands at the edge of the crease. Coach or D fires pucks from the point. Forward works on: (1) getting position with body/stick, (2) screening the goalie, (3) tipping shots, (4) burying rebounds. Add a D to battle for position.",
         "Establish position early — wide base, stick on the ice. Don't turn your back to the play. Quick hands on rebounds — no wind-up, just put it on net. When screening, move subtly to disrupt the goalie's tracking.",
         "One zone. F at net front, D/coach at point. Add opposing D for battle.",
         10, 3, "quarter", "Pucks",
         U12_UP, '["offensive","net_front","screening","tipping","retrievals","battle_drills"]', "offensive", "high", "net_front_presence"),

        ("Zone Entry Carry-and-Pass", "offensive",
         "Forward carries the puck through the neutral zone and attacks the blue line. Options: (1) carry wide and cut inside, (2) delay at the blue line and pass back to a trailing player, (3) chip and chase. Coach calls the option initially, then let the player read.",
         "Speed through the neutral zone — don't slow down at the blue line. Protect the puck on the carry — hand position matters. Trailing player: don't be even with the puck carrier, be 2-3 steps behind.",
         "Full ice. One forward per rep. Add a D at the blue line for progression.",
         10, 2, "full", "Pucks",
         U12_UP, '["offensive","zone_entry","puck_control","transition"]', "offensive", "medium", "zone_entry_carry"),

        # ── DEFENSIVE (4) ────────────────────────────────────────
        ("Gap Control 1-on-1", "defensive",
         "Defenseman starts at the blue line. Forward attacks from center ice. D must maintain proper gap — close enough to pressure but not so close they get beat wide. Run from both sides. Track how many times the D forces a turnover vs. gets beat.",
         "Gap is everything: stick length away at the blue line. Mirror the forward's movements — don't lunge. Angle the forward to the boards — take away the middle. Active stick — poke, lift, disrupt.",
         "Full ice. F starts at center. D starts at blue line. Goalie in net.",
         12, 2, "full", "Pucks",
         U12_UP, '["defensive","1_on_1","gap","angling"]', "defensive", "high", "gap_control_1on1"),

        ("Stick-on-Puck Angling", "defensive",
         "Forward carries the puck along the boards. Defenseman angles the carrier to the boards and separates them from the puck using stick positioning and body angling — no hitting in this drill. Focus on stick blade on the puck.",
         "Approach at an angle — don't skate straight at them. Get your stick on the puck first, then use your body to seal. Don't reach — get your feet in position first. Inside-out approach: force them to the wall.",
         "Half ice along the boards. F starts with the puck at the hash marks. D starts at the blue line.",
         10, 4, "half", "Pucks",
         U12_UP, '["defensive","angling","checking","puck_control"]', "defensive", "medium", "stick_on_puck_angling"),

        ("DZ Box Coverage Walkthrough", "defensive",
         "Five defensive players set up in a box-plus-one formation in the defensive zone. Coach moves the puck around to simulate offensive cycling. Defenders shift as a unit — maintaining box shape. Walk through at half speed, then add offensive players.",
         "Head on a swivel — know where every attacker is. Communicate: call switches, call the puck carrier. Inside positioning — stay between your man and the net. Collapse to the net when the puck goes low.",
         "One zone. 5 defensive players. Coach simulates offense, then add 3-5 attackers.",
         15, 5, "quarter", "Pucks",
         U16_UP, '["defensive","defensive_zone","coverage","systems","team"]', "defensive", "low", "dz_box_coverage"),

        ("Backcheck Tracking Drill", "defensive",
         "Three forwards attack 3-on-2. After the shot or turnover, the three forwards must sprint back and pick up three new attackers coming the other way. Focus on identifying your check while in full sprint.",
         "Sprint first — get back below the puck. Then find your man — closest threat. Communicate: call who you have. Stick in the lane — disrupt the pass while skating. Don't coast — backchecking is a sprint, every time.",
         "Full ice. Two groups of 3 forwards. Continuous flow.",
         12, 8, "full", "Pucks",
         U14_UP, '["defensive","backchecking","transition","coverage","conditioning"]', "defensive", "high", "backcheck_tracking"),

        # ── BATTLE DRILLS (4) ────────────────────────────────────
        ("Puck Protection Along the Boards", "battle",
         "One forward with the puck along the boards. One defenseman applying pressure. Forward must protect the puck for 10 seconds or find an escape pass to a coach/teammate at the half-wall. Switch roles after each rep.",
         "Wide base, low center of gravity. Use your body as a shield — back to the pressure. Roll off checks — don't stand still. Find the escape: look for the pass before you get pinned.",
         "Along the boards in one zone. One F, one D per rep. Coach at half-wall.",
         10, 4, "quarter", "Pucks",
         U12_UP, '["battle_drills","puck_protection","wall_play","1_on_1"]', "battle", "high", "puck_protection_boards"),

        ("Corner Battle Competition", "battle",
         "Dump the puck into the corner. One F and one D race to it. F tries to get the puck to the net or to a teammate at the half-wall. D tries to win the puck and clear it. Best of 5 wins. Losers do push-ups.",
         "First to the puck wins 80% of the time — feet move before the puck is dumped. Body position on arrival: get between the opponent and the puck. Quick hands — don't over-handle in traffic.",
         "One corner of the zone. F and D start at the hash marks. Coach dumps from the blue line.",
         10, 4, "quarter", "Pucks",
         U14_UP, '["battle_drills","retrievals","down_low","compete","1_on_1"]', "battle", "high", "corner_battle"),

        ("Net-Front Battle Drill", "battle",
         "Forward and defenseman battle for net-front position. Coach shoots from the point. Forward tries to tip or screen. Defenseman tries to clear the forward. After the shot, both compete for the rebound.",
         "F: Establish inside position early. Wide base. Stick on the ice at all times. D: Tie up the stick, box out with the body, clear rebounds quickly. Both: compete through the whistle.",
         "One zone. F and D at the net front. Coach at the point. Goalie in net.",
         10, 3, "quarter", "Pucks",
         U14_UP, '["battle_drills","net_front","screening","compete","1_on_1"]', "battle", "high", "net_front_battle"),

        ("Board-Play Battle Circuit", "battle",
         "Four stations along the boards. Each station: one attacker, one defender, one puck. Whistle starts the battle — 15 seconds. Attacker tries to escape with the puck. Defender tries to separate and clear. Rotate stations on the horn.",
         "Body position wins board battles. Feet first, then hands. Use leverage — low man wins. Find the escape route quickly — don't just grind.",
         "Full ice along both sides. 4 stations. Players rotate every 15 seconds.",
         10, 8, "full", "Pucks, cones for stations",
         U14_UP, '["battle_drills","wall_play","puck_protection","compete","checking"]', "battle", "high", "board_play_battle"),

        # ── SMALL AREA GAMES (3) ─────────────────────────────────
        ("3v3 Cross-Ice Game", "small_area_games",
         "Divide the ice into thirds using the blue lines. Play 3v3 cross-ice in each zone with small nets or cones as goals. Games to 3. Losers rotate out, winners stay on. Fast-paced, competitive.",
         "Quick puck movement — no room to stickhandle. Play with your head up. Support the puck — always give the carrier an option. Transition fast — first team to attack wins.",
         "Three zones. Small nets or cones for goals. 3v3 per zone.",
         15, 18, "full", "Small nets or cones, pucks",
         ALL_AGES, '["small_area_games","3_on_3","compete","passing","transition"]', "offensive", "high", "3v3_cross_ice"),

        ("King of the Rink", "small_area_games",
         "Everyone has a puck in a confined area (one zone). Players try to knock everyone else's puck out of the zone while protecting their own. If your puck leaves the zone, you're out. Last player standing wins.",
         "Head up — see the attacks coming. Protect your puck with your body. Be opportunistic — strike when they're not looking. Keep moving — stationary targets are easy.",
         "One zone. One puck per player.",
         8, 0, "quarter", "Pucks",
         ALL_AGES, '["small_area_games","puck_protection","awareness","compete"]', "puck_handling", "medium", "king_of_the_rink"),

        ("Possession Keepaway", "small_area_games",
         "4v4 or 5v5 in one zone. One team must complete 5 consecutive passes to score a point. Other team tries to intercept. No goalies. Fast transitions — turnover means the other team starts counting.",
         "Move after you pass — don't stand and watch. Show a target — give the puck carrier options. Quick passes — one touch when possible. Defensive pressure: deny passing lanes, don't just chase the puck.",
         "One zone. 4v4 or 5v5. No goalies.",
         10, 8, "quarter", "Pucks",
         U10_UP, '["small_area_games","passing","puck_support","coverage","compete"]', "passing", "high", "possession_keepaway"),

        # ── TRANSITION (3) ───────────────────────────────────────
        ("Breakout to Regroup Drill", "transition",
         "Coach dumps the puck in. D retrieves and executes the team's breakout pattern. Forwards support on the wall and through the middle. At the far blue line, the group regrouping by passing back to a D joining the rush, then re-attacking.",
         "D: Shoulder check before touching the puck. Quick first pass. Forwards: time your routes — don't leave too early. Regroup: D-to-D at the blue line if pressure. Speed through the neutral zone.",
         "Full ice. 5-player units (2D, 3F). Coach dumps from center ice.",
         15, 5, "full", "Pucks",
         U14_UP, '["transition","breakouts","re_group","systems"]', "transition", "medium", "breakout_regroup"),

        ("Neutral Zone Activation", "transition",
         "Three forwards and two defensemen work the neutral zone. Puck starts with D. They make a breakout pass and all five players activate through the neutral zone with speed. Focus on timing, lane filling, and puck support options.",
         "Fill all three lanes — don't bunch up. Middle lane driver sets the pace. D pinch up in support — don't hang back. Stretch pass option if the middle is clogged. Hit the blue line with speed, not with a stop.",
         "Full ice. 5-player units. Puck starts behind the net.",
         12, 5, "full", "Pucks",
         U14_UP, '["transition","neutral_zone","breakouts","offensive","speed"]', "transition", "high", "nz_activation"),

        ("Quick-Up Speed Drill", "transition",
         "D retrieves a dump-in and makes a quick up-ice pass to a forward who has already turned and is skating north. The forward receives in stride and attacks 1-on-0 or 2-on-1 depending on the variation.",
         "D: First touch should angle you up-ice. Head up immediately — find the outlet. Quick, hard pass. F: Don't wait — start moving before the D touches the puck. Receive in stride — this is about speed, not passing in place.",
         "Full ice. D behind the net. F at the far hash marks. Coach dumps.",
         10, 3, "full", "Pucks",
         U12_UP, '["transition","breakouts","speed","overspeed"]', "transition", "high", "quick_up_speed"),

        # ── SPECIAL TEAMS (3) ────────────────────────────────────
        ("PP Umbrella Rotation", "special_teams",
         "Five power play players set up in the umbrella (1-3-1) formation. Work puck movement around the perimeter: point to half-wall to low to opposite half-wall to point. On the second rotation, shoot from the top. Progress to reading the PK for seams.",
         "Quick puck movement — don't let the PK set. Half-wall player: options are down low, across, or back to the point. Point: one-time mentality, always ready. Low man: create traffic, look for tips.",
         "One zone. 5 PP players. Add 4 PK players for progression. Goalie in net.",
         15, 5, "quarter", "Pucks",
         U16_UP, '["special_teams","power_play","offensive","passing","shooting"]', "offensive", "medium", "pp_umbrella"),

        ("PK Diamond Positioning", "special_teams",
         "Four penalty killers set up in a diamond (1-2-1) formation. Coach moves the puck around the outside simulating PP movement. PK shifts as a unit — pressure the puck, clog the middle. Walk through, then add PP players.",
         "Stay compact — never chase to the perimeter. Pressure the puck with the high man. Sticks in passing lanes at all times. When the puck goes low, collapse. When it goes high, push out. Communication is critical.",
         "One zone. 4 PK players in diamond. Coach simulates, then add 5 PP players.",
         15, 4, "quarter", "Pucks",
         U16_UP, '["special_teams","penalty_kill","defensive","coverage","systems"]', "defensive", "medium", "pk_diamond"),

        ("Faceoff Play Execution", "special_teams",
         "Practice specific faceoff plays for each zone. Offensive zone: set play to get a quick shot. Defensive zone: clean win back to D for a breakout. Neutral zone: win and go. Run each play 5 times, then switch scenarios.",
         "Center: stance, hand position, eyes on the ref's hand. Wingers: know the play — timing is everything. D: be ready for a loss — have a counter. After the draw, everyone has a job — execute your route.",
         "Each zone. 5 players per unit. Both PP and PK faceoff sets.",
         12, 5, "full", "Pucks",
         U14_UP, '["special_teams","faceoffs","offensive","defensive","systems"]', "offensive", "low", "faceoff_plays"),

        # ── CONDITIONING (3) ─────────────────────────────────────
        ("Herbies", "conditioning",
         "Full-ice stop-and-start conditioning. Skate to the near blue line and back, then to the red line and back, then to the far blue line and back, then to the far goal line and back. That's one Herbie. Rest 30 seconds. Repeat 3-5 times.",
         "Full speed every rep — no coasting. Tight stops — both feet, spray the ice. First three strides out of the stop are the hardest and the most important. This is about mental toughness as much as fitness.",
         "Full ice. Players start on the goal line.",
         10, 0, "full", "None",
         U12_UP, '["conditioning","skating","compete"]', "skating", "high", "herbies_conditioning"),

        ("Relay Race Sprints", "conditioning",
         "Teams of 4. First player sprints to far blue line and back. Tags the next player. Relay continues until all 4 have gone. Losing team does the relay again. Best of 3 races.",
         "Explosive starts — first 3 strides. Tight turns at the blue line. Hand-off: next player is already moving when tagged. This is a race — compete hard.",
         "Full ice. Teams of 4 on the goal line. 2-4 teams.",
         8, 8, "full", "None",
         ALL_AGES, '["conditioning","skating","relay_races","compete"]', "skating", "high", "relay_sprints"),

        ("Puck-Carry Conditioning Circuit", "conditioning",
         "Players carry a puck through a circuit: sprint with the puck to the blue line, tight turn, sprint to center, tight turn, sprint to the far blue line, shoot on net. Skate back to the goal line without the puck. Rest while 2 others go. 5 reps each.",
         "Maintain puck control at full speed. Tight turns with the puck — don't lose it. Shoot in stride — no stopping to set up. This simulates game-speed carrying with fatigue.",
         "Full ice. Goalie in net for shots. Players go in waves.",
         10, 3, "full", "Pucks",
         U10_UP, '["conditioning","skating","puck_control","shooting"]', "skating", "high", "puck_carry_conditioning"),

        # ── GOALIE (3) ───────────────────────────────────────────
        ("T-Push Recovery Sequence", "goalie",
         "Goalie starts at one post. T-push across the crease to the other post. Set. T-push back. Repeat 10 times. Progress to: T-push, drop to butterfly, recover, T-push to other side. Then add shots after each push.",
         "Lead foot points to the direction of travel. Drive with the back leg — full extension. Set your feet before getting ready for a shot. Stay square to the shooter throughout the movement.",
         "One crease. Goalie only. Coach adds shots for progression.",
         10, 1, "quarter", "Pucks for progression",
         U10_UP, '["goalie","movement","recovery","skill_development"]', "goalie", "medium", "goalie_t_push"),

        ("Butterfly Slide Movement", "goalie",
         "Goalie starts centered in the net. Coach calls a direction. Goalie drops to butterfly and slides laterally. Set. Back to standing. Next direction. Progress to: slide to post, recover, track a pass across, slide to the other post.",
         "Lead with the pad — knee drives toward the direction. Hands stay up and forward. Seal the ice with the pad. Recover quickly — don't stay down. Track the puck with your eyes throughout.",
         "One crease. Goalie only. Coach feeds passes for tracking.",
         10, 1, "quarter", "Pucks for progression",
         U10_UP, '["goalie","butterfly","movement","skill_development"]', "goalie", "medium", "goalie_butterfly_slide"),

        ("Angle Play Positioning", "goalie",
         "Shooter starts at different locations around the zone (point, top of the circle, half-wall, low slot, behind the net). At each location, the goalie practices challenging out to the correct depth and angle. Coach verifies positioning before the shot is taken.",
         "Challenge the shooter — out and up. Depth depends on distance: farther out for close shots, deeper for far shots. Square to the puck — belly button faces the shooter. Hold your ground — don't back in.",
         "One zone. Goalie in net. Coach or shooters at various locations.",
         12, 1, "quarter", "Pucks",
         U10_UP, '["goalie","positioning","angles","skill_development"]', "goalie", "low", "goalie_angle_play"),
    ]

    for d in drills:
        conn.execute("""
            INSERT INTO drills (id, org_id, name, category, description, coaching_points, setup,
                duration_minutes, players_needed, ice_surface, equipment, age_levels, tags,
                skill_focus, intensity, concept_id)
            VALUES (?, NULL, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
        """, (str(uuid.uuid4()), *d))
    conn.commit()
    conn.close()
    logger.info("Seeded %d drills across 13 categories", len(drills))


def seed_drills_v2():
    """Seed 80+ additional original drills — heavy focus on U8/U10 age groups and expanded categories."""
    conn = get_db()
    count = conn.execute("SELECT COUNT(*) FROM drills WHERE org_id IS NULL").fetchone()[0]
    if count > 50:
        conn.close()
        return

    ALL_AGES = '["U8","U10","U12","U14","U16_U18","JUNIOR_COLLEGE_PRO"]'
    U8_ONLY = '["U8"]'
    U8_U10 = '["U8","U10"]'
    U8_U12 = '["U8","U10","U12"]'
    U10_UP = '["U10","U12","U14","U16_U18","JUNIOR_COLLEGE_PRO"]'
    U10_U14 = '["U10","U12","U14"]'
    U12_UP = '["U12","U14","U16_U18","JUNIOR_COLLEGE_PRO"]'
    U14_UP = '["U14","U16_U18","JUNIOR_COLLEGE_PRO"]'
    U16_UP = '["U16_U18","JUNIOR_COLLEGE_PRO"]'
    JR_PLUS = '["JUNIOR_COLLEGE_PRO"]'

    drills = [
        # ══════════════════════════════════════════════════════════
        # U8 / MITE FOCUSED DRILLS (fun, simple, maximum touches)
        # ══════════════════════════════════════════════════════════

        # ── U8 WARM UP ──
        ("Shark and Minnows", "warm_up",
         "One player is the shark at center ice. All other players (minnows) line up on one goal line. On the whistle, minnows skate to the other end while the shark tries to tag them. Tagged players become sharks. Last minnow standing wins.",
         "Keep it fun and energetic. Watch for players stopping — encourage continuous movement. Reinforce heads up and awareness of other skaters. Great for building skating confidence.",
         "Full ice. All players on goal line. One designated shark at center ice.",
         5, 0, "full", None,
         U8_U10, '["warm_up","skating","fun","compete","agility"]', "skating", "medium", "shark_minnows"),

        ("Follow the Leader Skating", "warm_up",
         "Coach leads a line of players around the ice doing various skating movements. Players mimic everything the leader does — glide on one foot, spin, stop, skate backward, wiggle through cones. Change leader every 90 seconds.",
         "Keep movements age-appropriate. Celebrate effort not perfection. Mix silly movements with proper technique. Great opportunity to model correct skating posture.",
         "Full ice. Players in a single file line behind the coach.",
         6, 0, "full", None,
         U8_ONLY, '["warm_up","skating","fun","agility"]', "skating", "low", "follow_the_leader"),

        ("Red Light Green Light with Pucks", "warm_up",
         "All players line up on the goal line with pucks. Coach faces away and calls green light — players skate forward with pucks. Coach calls red light and turns around — everyone must stop and control their puck. Anyone still moving goes back to the start.",
         "Emphasize stopping with the puck controlled. Reward players who stop quickly with puck close. Great for teaching puck control at slow speeds and hockey stops.",
         "Full ice. One puck per player. All start on goal line.",
         6, 0, "full", "Pucks",
         U8_U10, '["warm_up","puck_control","fun","skating","stopping"]', "puck_handling", "low", "red_light_green_light"),

        ("Obstacle Course Adventure", "warm_up",
         "Set up a fun obstacle course using cones, sticks on the ice, and pylons. Players navigate through — step over sticks, weave through cones, skate around pylons, drop to knees and get back up, finish with a coast into the boards. Timed runs for added fun.",
         "Set obstacles at appropriate difficulty. Encourage players to go at their own speed first, then try faster. Build confidence with achievable challenges. Cheer loudly for every player.",
         "Full or half ice. Cones, extra sticks laid flat, pylons arranged in a course.",
         8, 0, "full", "Cones, extra sticks, pylons",
         U8_ONLY, '["warm_up","skating","fun","agility","balance"]', "skating", "low", "obstacle_course"),

        # ── U8 PUCK HANDLING ──
        ("Puck Handling Maze", "puck_handling",
         "Cones set up in a maze pattern. Players navigate through the maze with a puck, trying different routes. Make it a game — find the fastest path through. Add a second puck for advanced players.",
         "Keep the stick blade cupped over the puck. Small movements — don't let the puck get away from the body. Look up periodically to find the next opening. Reward creativity in route selection.",
         "Half ice. 15-20 cones arranged in a maze with multiple paths.",
         10, 0, "half", "Cones, pucks",
         U8_U10, '["puck_handling","stickhandling","fun","agility"]', "stickhandling", "low", "puck_handling_maze"),

        ("Musical Pucks", "puck_handling",
         "Scatter pucks around the zone (one fewer than the number of players). Players skate around freely. When the whistle blows, everyone grabs a puck and stickhandles to the nearest face-off dot. Player without a puck does three knee-bends. Remove one puck each round.",
         "Keep it fun — no body contact to take pucks. Emphasize quick feet to a puck and then controlled stickhandling. Players waiting can do fun skating moves. Builds awareness and puck scramble instincts.",
         "Half ice. Pucks scattered around the zone. One fewer puck than players.",
         8, 0, "half", "Pucks",
         U8_ONLY, '["puck_handling","fun","awareness","compete"]', "puck_handling", "medium", "musical_pucks"),

        ("Toe Drag Around Cones", "puck_handling",
         "Players line up and skate through a line of 5 cones spaced 8 feet apart. At each cone, pull the puck from forehand to backhand using a toe drag to get around the cone. Walk through slowly first, then add speed.",
         "Keep the puck close to the body during the drag. Top hand does the work — pull across and roll the wrists. Knees bent. Progress from walking speed to skating speed only when the technique is clean.",
         "Half ice. 5 cones per lane, 2-3 lanes. One puck per player.",
         10, 0, "half", "Cones, pucks",
         U8_U12, '["puck_handling","stickhandling","toe_drag","skill_development"]', "stickhandling", "low", "toe_drag_cones"),

        ("Protect Your Egg", "puck_handling",
         "Each player has a puck (their egg) in a confined area. While stickhandling their own puck, they try to knock other players pucks out of the zone. If your puck leaves the zone, do 5 toe-taps and come back in. Last player with their puck in the zone wins.",
         "Body position is key — use your body to shield the puck. Eyes up to see threats coming. Small controlled stickhandles, not big sweeping ones. Teaches puck protection instincts at a young age.",
         "One zone or neutral zone. One puck per player. Use lines as boundaries.",
         8, 0, "quarter", "Pucks",
         U8_U12, '["puck_handling","puck_protection","fun","compete","awareness"]', "puck_handling", "medium", "protect_your_egg"),

        # ── U8 PASSING ──
        ("Partner Pass and Move", "passing",
         "Players pair up with one puck. They pass back and forth while skating slowly up the ice together. On each pass, the passer must skate to a new spot before receiving the pass back. Emphasize always moving after passing.",
         "Pass and move — never stand still after passing. Show a target with the stick blade on the ice. Start with stationary passing, then add slow skating, then full speed. Tape-to-tape passes only.",
         "Full ice. One puck per pair. Pairs spread across the ice.",
         8, 0, "full", "Pucks",
         U8_U10, '["passing","movement","skating","fundamentals"]', "passing", "low", "partner_pass_move"),

        ("Triangle Passing Game", "passing",
         "Three players form a triangle about 15 feet apart. Pass around the triangle — forehand passes only at first, then add backhand. Call out the name of the player you are passing to. Rotate clockwise, then counter-clockwise.",
         "Call the name before you pass. Stick on the ice gives a target. Receive the puck and cushion it — don't let it bounce off the blade. Progress to one-touch passing when ready.",
         "Half ice. Groups of 3, one puck per group, spaced in triangles.",
         8, 6, "half", "Pucks",
         U8_U10, '["passing","communication","fundamentals"]', "passing", "low", "triangle_passing"),

        ("Pass Through the Gate", "passing",
         "Set up gates (two cones 3 feet apart) scattered around the zone. Partners must pass the puck through the gates to each other. Count successful gate passes in 60 seconds. Beat your record each round.",
         "Accuracy over speed. Aim for the middle of the gate. Weight of the pass matters — not too hard, not too soft. Move to different gates after each successful pass.",
         "Half ice. 8-10 cone gates spread around zone. One puck per pair.",
         8, 0, "half", "Cones, pucks",
         U8_U12, '["passing","accuracy","fun","compete"]', "passing", "low", "pass_through_gate"),

        # ── U8 SKATING ──
        ("Penguin Walks", "skating",
         "Players take tiny steps on the ice without gliding — like penguins walking. Progress to: small marching steps, then longer gliding steps, then full strides. Use across the width of the ice (not full length) for beginners.",
         "Bend the knees. Push to the side, not straight back. Each step gets a little longer. Arms swing naturally. This builds the fundamental stride pattern for beginners who are still learning to balance.",
         "Full or half ice. No equipment needed.",
         8, 0, "full", None,
         U8_ONLY, '["skating","fundamentals","beginners","balance"]', "skating", "low", "penguin_walks"),

        ("Treasure Hunt Skate", "skating",
         "Hide pucks (treasures) around the ice behind nets, along boards, at face-off dots. Players skate around finding and collecting pucks — carry them in one hand or push with stick. First to find 3 pucks wins. Reset and play again.",
         "Encourages skating without thinking about skating. Players focus on finding pucks and naturally improve their movement. Great confidence builder. Vary hiding spots each round.",
         "Full ice. Scatter 20-30 pucks in various locations around the ice.",
         8, 0, "full", "Pucks (20-30)",
         U8_ONLY, '["skating","fun","agility","awareness"]', "skating", "low", "treasure_hunt"),

        ("One-Foot Glide Challenge", "skating",
         "Players skate across the ice and try to glide on one foot as long as possible. Start with their strong foot, then switch to weak foot. Mark their distance with a cone. Try to beat their distance each round. Add arms out for balance.",
         "Bend the gliding knee slightly. Look forward not down. Arms out for balance. The standing foot should be directly under the body. This builds edge control and balance — foundation for all advanced skating.",
         "Full ice width. Cones to mark distances.",
         8, 0, "full", "Cones",
         U8_U10, '["skating","balance","edges","fundamentals"]', "skating", "low", "one_foot_glide"),

        ("Snowplow Stop Races", "skating",
         "Two players race side by side across the ice. At the far blue line, both must do a complete snowplow stop. First player to stop completely wins. Progress to hockey stops when ready. Emphasize stopping fully — no gliding through.",
         "Bend knees and push both feet out (snowplow) or turn feet sideways (hockey stop). Weight slightly back. Scrape the ice with the blade — you should hear the snow spray. Full stop before turning around.",
         "Full ice width. Two lanes. Pairs race.",
         6, 0, "full", None,
         U8_U10, '["skating","stopping","compete","fundamentals"]', "skating", "medium", "snowplow_races"),

        # ── U8 SHOOTING ──
        ("Stationary Wrist Shot Basics", "shooting",
         "Players line up along the hash marks facing the boards (not the net initially). Practice the wrist shot motion against the boards — pull puck back, roll wrists, follow through pointing at target. After 10 good reps against the boards, rotate to shoot on net.",
         "Start with the puck at the heel of the blade. Sweep forward and roll the wrists — top hand pushes, bottom hand pulls. Follow through toward the target. Weight transfers from back foot to front foot. Power comes from the legs.",
         "Half ice. Players along hash marks. Pucks. Progress to shooting on net.",
         10, 0, "half", "Pucks",
         U8_U10, '["shooting","fundamentals","wrist_shot","skill_development"]', "shooting", "low", "stationary_wrist_shot"),

        ("Shoot at Targets Game", "shooting",
         "Place water bottles or foam targets on the crossbar and in the corners of the net. Players take turns shooting from the slot trying to knock targets off. Keep score. Reset targets after each round. Make it a team competition.",
         "Pick your target before you shoot. Eyes on the target, not the puck. Follow through toward where you want the puck to go. Celebrate hits loudly. Accuracy is more important than power at this age.",
         "One zone. Goalie net with targets. Players shoot from hash marks.",
         10, 0, "quarter", "Pucks, water bottles or targets",
         U8_U12, '["shooting","accuracy","fun","compete"]', "shooting", "low", "shoot_at_targets"),

        # ── U8 SMALL AREA GAMES ──
        ("2v2 Mini Games", "small_area_games",
         "Divide the ice into three zones using the blue lines. Play 2v2 in each zone with small nets or cones as goals. Games to 2, losers rotate to the next zone. Quick shifts, maximum touches, tons of fun. Coaches can add rules like must pass before scoring.",
         "Keep shifts short (60-90 seconds). Encourage passing — maybe require one pass before a shot. Celebrate teamwork. Change partners frequently. These games build hockey sense naturally — reading plays, supporting teammates, competing.",
         "Full ice divided into 3 zones. Small nets or cone goals. Multiple pucks ready.",
         12, 12, "full", "Small nets or cones, pucks",
         U8_U10, '["small_area_games","2_on_2","compete","fun","passing"]', "offensive", "medium", "2v2_mini_games"),

        ("Capture the Puck", "small_area_games",
         "Two teams, each on their own blue line. Pucks scattered at center ice. On the whistle, players race to center, grab a puck, and bring it back to their goal line. Once all center pucks are taken, you can steal from the other team goal line. Most pucks after 2 minutes wins.",
         "Skating speed and quick decisions. Which puck to grab? When to steal? Builds competitive instincts and skating urgency. No body contact — puck stealing only with stick. Pure fun and energy.",
         "Full ice. 15-20 pucks at center ice. Two teams on opposite blue lines.",
         8, 0, "full", "Pucks (15-20)",
         U8_U10, '["small_area_games","fun","compete","skating","awareness"]', "skating", "high", "capture_the_puck"),

        # ══════════════════════════════════════════════════════════
        # U10 / SQUIRT DRILLS (building technique, introducing concepts)
        # ══════════════════════════════════════════════════════════

        ("Mohawk Turn Progression", "skating",
         "Players skate forward, then open hips to transition to backward skating using a mohawk turn (inside edges, feet form a V momentarily). Practice at walking speed first along the boards, then add glide, then full speed. Both directions.",
         "Open the hips — don't spin. The back foot opens first, weight transfers, front foot follows. Stay low through the turn. This is the foundation for all defensive pivots. Practice both directions equally.",
         "Full ice. Players in waves of 4-5 along the boards.",
         10, 0, "full", None,
         U10_UP, '["skating","transitions","pivots","edges"]', "skating", "medium", "mohawk_turns"),

        ("Crossover Acceleration Drill", "skating",
         "Players skate around a face-off circle using crossovers. On the whistle, explode out of the circle on a straight-line sprint to the boards and back. Focus on using the crossover momentum to accelerate out of the turn. Alternate clockwise and counter-clockwise.",
         "Deep knee bend in the crossovers — load the outside leg. Explode out by driving the inside leg under. First three strides out of the circle are everything. Keep the upper body quiet while the legs do the work.",
         "Half ice. Use face-off circles. Players in groups of 4 at each circle.",
         10, 0, "half", None,
         U10_UP, '["skating","crossovers","acceleration","power"]', "skating", "high", "crossover_acceleration"),

        ("Give and Go Passing", "passing",
         "Two players attack with one defender. Carrier passes to teammate and immediately sprints to open ice for the return pass. Defender plays passive at first, then active. Focus on the timing of the give-and-go — pass, sprint, receive, attack.",
         "The give-and-go only works if you sprint after passing. Don't admire your pass — move your feet immediately. The return pass should hit the sprinter in stride. This is hockey's most basic offensive concept and one of the most effective.",
         "Half ice. Groups of 3 (2 offense, 1 defense). Pucks. Rotate positions.",
         12, 3, "half", "Pucks",
         U10_U14, '["passing","offensive","give_and_go","movement","decision_making"]', "passing", "medium", "give_and_go"),

        ("Backhand Passing Progression", "passing",
         "Partners face each other 15 feet apart. All passes must be backhand. Start stationary, progress to skating slowly, then at speed. Finally add a forehand-backhand alternating pattern. Focus on rolling the bottom hand to generate the backhand pass.",
         "Open the blade face on the backhand side. Roll the bottom wrist to push through the puck. Follow through toward the target. The backhand pass is weaker than forehand so close the distance slightly. Weight transfer helps generate power.",
         "Half ice. Partners 15 feet apart. One puck per pair.",
         10, 0, "half", "Pucks",
         U10_UP, '["passing","backhand","skill_development","fundamentals"]', "passing", "low", "backhand_passing"),

        ("Snap Shot Introduction", "shooting",
         "Players learn the snap shot — a quick release shot with minimal backswing. Line up at the hash marks, puck on the forehand. Quick snap of the wrists with a short pull-and-release motion. Emphasize speed of release over power. Progress to receiving a pass and snapping.",
         "Minimal backswing — the power comes from the snap of the wrists. Pull the puck slightly back then snap forward quickly. Weight transfers from back to front. The snap shot is all about quick release — getting the shot off before the goalie is set.",
         "One zone. Players at hash marks. Goalie in net.",
         10, 0, "quarter", "Pucks",
         U10_UP, '["shooting","snap_shot","quick_release","skill_development"]', "shooting", "medium", "snap_shot_intro"),

        ("Shooting in Stride", "shooting",
         "Players skate down the wing with a puck. Without stopping, release a wrist shot on net while still skating. The key is not breaking stride — the shot happens mid-movement. Start from the hash marks, progress to longer carries.",
         "Don't stop your feet to shoot. Transfer weight from the back leg through the shot. The puck should be slightly ahead of the body at release. Pull the puck in tight, then release. Head up, pick your spot, shoot while moving.",
         "Full ice. Wingers carry down the wing. Goalie in net. Both sides.",
         12, 0, "full", "Pucks",
         U10_UP, '["shooting","wrist_shot","skating","offensive"]', "shooting", "medium", "shooting_in_stride"),

        # ── U10 DEFENSE CONCEPTS ──
        ("Mirror Skating Defense", "defensive",
         "One forward, one defender face each other. Forward skates left, right, forward, backward — defender must mirror every movement while skating backward. Stay within one stick-length. No puck initially, then add puck for the forward.",
         "Defensive stance: knees bent, stick on the ice, eyes on the chest (not the puck). Mirror the hips — where the hips go, the player goes. Stay within a stick-length gap. Don't lunge or reach — move your feet.",
         "Half ice. Pairs facing each other. Progress from no puck to with puck.",
         10, 2, "half", "Pucks (for progression)",
         U10_U14, '["defensive","gap_control","backward_skating","1_on_1"]', "defensive", "medium", "mirror_skating_defense"),

        ("Poke Check Technique", "defensive",
         "Defenders practice the poke check motion — quick thrust of the stick to knock the puck away without leaving defensive position. Start against a stationary puck, then against a slow-moving attacker, then game speed. Emphasize timing over aggression.",
         "One hand on the stick for the poke, extend fully, then snap back to two hands. Do NOT dive or lunge — feet stay planted. Timing is everything: poke when the attacker looks down or the puck is exposed. Miss the poke? Recover to gap position immediately.",
         "Half ice. Attacker vs defender pairs. Progressive speed.",
         10, 2, "half", "Pucks",
         U10_UP, '["defensive","poke_check","technique","1_on_1"]', "defensive", "medium", "poke_check_technique"),

        # ── U10 TRANSITION ──
        ("Breakout Basics — Three Options", "transition",
         "Coach dumps puck into the zone. Defense retrieves behind the net and executes one of three breakout options on the coach's call: (1) Reverse — pass to strong-side winger along the boards, (2) Over — pass up the middle to the center, (3) Wheel — D skates behind the net to the other side and passes to the weak-side winger.",
         "D must shoulder check before touching the puck — know where the pressure is coming from. Quick first pass is critical. Wingers provide a target along the boards. Center supports in the middle lane. Communication: call out which option you want.",
         "Full ice. 5-player units (2D, 3F). Coach dumps from center.",
         12, 5, "full", "Pucks",
         U10_U14, '["transition","breakout","systems","communication"]', "transition", "medium", "breakout_basics"),

        # ══════════════════════════════════════════════════════════
        # U12+ TACTICAL DRILLS (more complex concepts)
        # ══════════════════════════════════════════════════════════

        ("F1-F2-F3 Forecheck Roles", "systems",
         "Teach the 1-2-2 forecheck roles. F1 pressures the puck carrier (angling to the boards). F2 supports F1 and takes away the D-to-D pass. F3 plays high in the middle as a safety valve. Walk through at half speed, then add opposition. Rotate all three positions.",
         "F1 takes an angle — drive the puck carrier to the boards, don't chase blindly. F2 reads the first pass option and eliminates it. F3 stays high and center — if puck gets past F1 and F2, F3 is the last line of defense. Aggressive but disciplined.",
         "Full ice. 3 forwards vs 2 D. Coach initiates breakout.",
         15, 5, "full", "Pucks",
         U12_UP, '["systems","forecheck","1_2_2","roles","team"]', "systems", "medium", "forecheck_roles_122"),

        ("Neutral Zone 1-3-1 Trap Walkthrough", "systems",
         "Set up a 1-3-1 neutral zone structure. One forward pressures high, three players across the neutral zone take away east-west passes, one forward stays low as a backcheck safety. Walk through puck movement and rotations. Progress to 5-on-5 controlled scrimmage.",
         "The trap is about patience — don't chase the puck, take away passing lanes. Middle three stay connected (within a stick-length of each other). High forward funnels the play to the strong side. Low forward reads and counters stretch passes. Communication and discipline.",
         "Full ice. Two 5-player units. Walk through positioning at half speed.",
         15, 10, "full", None,
         U14_UP, '["systems","neutral_zone","trap","1_3_1","team","positioning"]', "systems", "low", "nz_trap_131"),

        ("DZ Man-to-Man Coverage Drill", "defensive",
         "Five defenders in the defensive zone, each assigned a specific attacker to cover man-to-man. Coach moves the puck around the zone, defenders must stay with their assigned player regardless of where the puck goes. Progress to live play with attackers trying to get open.",
         "Stay between your man and the net at all times. Body on body — don't watch the puck. Communicate switches if attackers cross. Stick in passing lane. When the puck is in the corner, your man is your priority — don't collapse unless told to.",
         "One zone. 5 attackers, 5 defenders. Coach controls puck at first.",
         15, 10, "quarter", "Pucks",
         U12_UP, '["defensive","man_to_man","coverage","defensive_zone","team"]', "defensive", "medium", "dz_man_coverage"),

        ("Offensive Zone Cycle Game", "offensive",
         "Three forwards work the puck in the offensive zone for 30 seconds against two defenders. Goal is to maintain possession through cycling along the boards, reversals, and quick passes. Score from low cycle plays only (net-front tip, short-side, wraparound). Points for sustained possession and goals.",
         "Cycle means constant movement — low man gets the puck, drives up the boards, dishes to the high man coming down. Third forward reads and fills the open lane (net-front, high slot, or weak side). Strong on the puck along the boards. Protect with body, quick pass when pressured.",
         "One zone. 3F vs 2D. Goalie in net. 30-second shifts, rotate groups.",
         12, 5, "half", "Pucks",
         U12_UP, '["offensive","cycling","possession","wall_play","decision_making"]', "offensive", "high", "oz_cycle_game"),

        ("Point Shot Traffic Drill", "shooting",
         "Defenseman at the point with pucks. Two forwards set up in front of the net — one screening, one at the far post for tips/rebounds. D shoots through traffic, forwards work to screen the goalie and redirect. Rotate all three positions.",
         "Point shot should be low and on net — a missed net is a wasted opportunity. Forwards create traffic (don't move out of the way). Screening forward: wide base, stick on ice, don't turn your back to the play. Tip forward: blade on the ice, redirect don't swat.",
         "One zone. D at point. 2F in front. Goalie in net.",
         12, 3, "quarter", "Pucks",
         U12_UP, '["shooting","point_shot","screening","tipping","offensive"]', "shooting", "medium", "point_shot_traffic"),

        ("3-on-2 Continuous Rush", "offensive",
         "Continuous flow drill. Three forwards attack 2 defenders. After the play ends (goal, save, or turnover), the two defenders now pick up a new puck and join one forward to become the new 3-on-2 attacking the other way against two new defenders. Continuous flow — no stoppages.",
         "Attack with speed and width — spread the ice. Middle driver has options: keep, pass left, pass right. Off-puck players drive to the net and far post. D work together — strong side takes the puck, weak side takes the pass. Communicate.",
         "Full ice. Continuous flow. 3F attack, 2D defend, flip and go.",
         15, 10, "full", "Pucks",
         U12_UP, '["offensive","3_on_2","rush","transition","continuous_flow"]', "offensive", "high", "3on2_continuous"),

        ("Delay Entry and Regroup", "offensive",
         "Forward carries the puck into the neutral zone. At the far blue line, instead of forcing entry, delays and passes back to a defenseman joining the rush. D carries into the zone or passes to a winger who has changed lanes. Teaches patience at the blue line.",
         "Don't force entries against a stacked blue line — live to play another day. The delay creates time for teammates to read and adjust. D joining the rush adds an extra attacker. Wingers change lanes during the delay to create confusion for defenders.",
         "Full ice. 5-player units. Coach signals delay or go.",
         12, 5, "full", "Pucks",
         U14_UP, '["offensive","zone_entry","delay","re_group","systems"]', "offensive", "medium", "delay_entry_regroup"),

        # ── POWER PLAY DRILLS ──
        ("PP 1-3-1 Setup and Movement", "special_teams",
         "Five players set up in the 1-3-1 power play formation: one quarterback at the point, two half-wall flanks, one bumper in the high slot, one net-front presence. Walk through the puck movement pattern: point to half-wall to low to opposite half-wall to point. Add shooting from various positions.",
         "QB at the point: distribute quickly, shoot when the lane opens. Half-wall: triple threat (pass down, pass across, shoot). Bumper: stay in the high slot, one-touch passes, look for seam shots. Net-front: screen, tip, pounce on rebounds. Quick puck movement — don't let the PK set up.",
         "One zone. 5 PP players. Add PK for progression. Goalie in net.",
         15, 5, "quarter", "Pucks",
         U14_UP, '["special_teams","power_play","1_3_1","offensive","systems"]', "offensive", "medium", "pp_131_setup"),

        ("PP Overload Formation", "special_teams",
         "Five PP players set up in an overload on one side of the ice. Three players on the strong side (half-wall, low, slot), one at the point, one weak-side option. Work the strong side with quick passes and shots, then reverse to the weak side when the PK overcommits.",
         "The overload works because the PK can't cover 3 players on one side. Quick passes create shooting lanes. When the PK collapses to the strong side, the weak-side player is wide open — reverse the puck fast. Net-front player is always the most dangerous.",
         "One zone. 5 PP players. Add PK for progression.",
         15, 5, "quarter", "Pucks",
         U14_UP, '["special_teams","power_play","overload","offensive","systems"]', "offensive", "medium", "pp_overload"),

        ("PP Zone Entry Practice", "special_teams",
         "Practice the three main PP zone entry options against two PK forwards: (1) controlled entry — carry wide and cut in, (2) drop pass at the blue line to the trailer, (3) dump to the corner and chase with numbers. Five reps of each, then read and react.",
         "Entry is the hardest part of the PP. Carry-in works against passive PK. Drop pass works when they pressure high — but the drop must be dead (don't push it back). Dump and chase when nothing else works — send two chasers. Never turn the puck over at the blue line.",
         "Full ice. PP unit vs 2 PK forwards at the blue line.",
         12, 7, "full", "Pucks",
         U14_UP, '["special_teams","power_play","zone_entry","decision_making"]', "offensive", "medium", "pp_zone_entry"),

        # ── PENALTY KILL DRILLS ──
        ("PK Box Formation Drill", "special_teams",
         "Four PK players set up in a box formation. Coach moves the puck around simulating a PP. The box shifts as a unit — pressure the puck carrier, stay compact, clog the middle. Walk through at half speed, then add 5 PP players.",
         "Stay compact — the box should be tight enough that no one can split you. Pressure the puck but don't chase to the perimeter. Sticks in passing lanes at all times. When the puck goes low, collapse. When it goes high, push out. Communication is everything on the PK.",
         "One zone. 4 PK players. Coach simulates, then add PP.",
         15, 4, "quarter", "Pucks",
         U14_UP, '["special_teams","penalty_kill","box","defensive","systems"]', "defensive", "medium", "pk_box"),

        ("PK Aggressive Pressure System", "special_teams",
         "Four PK players practice an aggressive PK — pressuring the PP high to force turnovers. F1 chases the puck aggressively, F2 takes away the easy pass, both D stay connected but push up. Goal is to force bad passes and create shorthanded chances.",
         "High risk, high reward. Only use when trailing or need momentum. F1 must commit fully — angle hard. If F1 doesn't win the battle, everyone drops back to box. Time your pressure — attack right after a PP zone entry when they're getting set. Don't get caught up ice.",
         "One zone. 4 PK vs 5 PP. Full speed.",
         12, 9, "quarter", "Pucks",
         U16_UP, '["special_teams","penalty_kill","aggressive","pressure","systems"]', "defensive", "high", "pk_aggressive"),

        # ── ADVANCED DRILLS (U14+) ──
        ("Stretch Pass Breakout", "transition",
         "D retrieves the puck behind the net. Instead of the standard breakout, looks for the long stretch pass to a forward who has sneaked behind the opposing forecheckers at the far blue line. Timing is everything — the forward must time their move to stay onside.",
         "This is a home run play — high reward but high risk if intercepted. D must sell the short play first (look to the boards) then quickly switch to the stretch. Forward must be onside — timing is critical. Only attempt when the forecheck is aggressive and leaves the middle open.",
         "Full ice. D behind net, F at far blue line. Add forecheckers.",
         10, 3, "full", "Pucks",
         U14_UP, '["transition","breakout","stretch_pass","speed","offensive"]', "transition", "high", "stretch_pass_breakout"),

        ("Headmanning the Puck Drill", "transition",
         "Defenders retrieve loose pucks and practice finding the farthest open forward quickly. Three forwards spread across the ice at different depths. D must read which forward is open and deliver the puck up-ice as fast as possible. No north-south stickhandling — move the puck fast.",
         "Headmanning means getting the puck to the farthest open teammate as quickly as possible. Shoulder check before touching the puck. The quick up-ice pass creates odd-man rushes. A D who can headman the puck is worth their weight in gold. Don't force it — if nobody's open, make the safe play.",
         "Full ice. D behind net. 3F spread at blue line, red line, far blue line.",
         10, 4, "full", "Pucks",
         U14_UP, '["transition","breakout","headmanning","passing","decision_making"]', "transition", "medium", "headmanning"),

        ("Line Rush 5-on-0 Systems", "systems",
         "Full 5-player unit attacks from their own zone through neutral ice into the offensive zone in a structured 5-on-0 rush. Focus on lane filling, timing, puck support, and proper zone entry formation. D join the rush at the right depth. Run the team's actual system.",
         "Five lanes across the ice — everyone has a lane. Center controls the pace. Wingers drive wide and cut at the blue line. D trail at proper depth (not too close, not too far). Puck moves side to side through the neutral zone. Hit the blue line with speed — nobody stops at the line.",
         "Full ice. 5-player units running actual team breakout-to-rush system.",
         12, 5, "full", "Pucks",
         U12_UP, '["systems","rush","5_on_0","lane_filling","team"]', "transition", "medium", "line_rush_5on0"),

        # ── BATTLE / COMPETE DRILLS ──
        ("1-on-1 From the Knees", "battle",
         "Two players start on their knees at the hash marks facing each other. Puck placed between them. On the whistle, both battle for the puck and try to score on the mini net behind the other player. Great for building upper body strength and compete level in a controlled environment.",
         "Battle for inside positioning. Strong base even on your knees. Use your body to shield the puck. Quick hands win. This teaches compete without the speed — players learn body positioning, leverage, and hand battles in slow motion.",
         "Half ice. Pairs at hash marks. Mini nets or cones. Pucks.",
         8, 4, "half", "Pucks, mini nets or cones",
         U10_UP, '["battle_drills","1_on_1","compete","strength","puck_protection"]', "battle", "high", "1on1_from_knees"),

        ("D-Zone Faceoff Drill", "battle",
         "Practice defensive zone faceoffs with specific assignments. Center battles for the draw. Wingers tie up opposing wingers. D position for a clean win-back or a loose puck battle. Run 10 faceoffs per unit, track clean wins vs. losses.",
         "Center: stance low, eye on the ref's hand, quick hands. Strong-side winger: tie up their winger's stick immediately. D: if we win, retrieve and breakout. If we lose, collapse to net-front and win the battle. Everyone has a job — execute it every time.",
         "One zone. 5-player units. Both offensive and defensive sets.",
         12, 10, "half", "Pucks",
         U12_UP, '["battle_drills","faceoffs","defensive","systems","compete"]', "battle", "medium", "dz_faceoff_drill"),

        ("Loose Puck Races", "battle",
         "Coach dumps or shoots a puck into the corner or along the boards. Two players (one from each team) race to win the loose puck. Winner tries to score, loser tries to defend. Emphasize acceleration and body positioning on arrival.",
         "First to the puck wins most of the time — explode on the whistle. But arriving first means nothing if you don't protect the puck. Get your body between the opponent and the puck. Low center of gravity. Quick decision: shoot, pass, or protect.",
         "Half ice. Two lines at the blue line. Coach at center ice dumps.",
         10, 4, "half", "Pucks",
         U10_UP, '["battle_drills","compete","loose_pucks","racing","intensity"]', "battle", "high", "loose_puck_races"),

        # ── MORE CONDITIONING ──
        ("Suicide Sprints with Pucks", "conditioning",
         "Same as classic Herbies but carrying a puck. Skate to near blue line and back, red line and back, far blue line and back, far goal line and back — all while controlling the puck. Tests both fitness and puck control under fatigue.",
         "Don't lose the puck at the turns — tight control on the transition. Push through the fatigue — this is where you gain an edge. Proper technique even when tired: bend the knees, full stride, no sloppy turns. This simulates late-period puck carrying.",
         "Full ice. One puck per player. Goal line start.",
         10, 0, "full", "Pucks",
         U12_UP, '["conditioning","skating","puck_control","compete"]', "skating", "high", "suicide_sprints_pucks"),

        ("30-Second All-Out Shifts", "conditioning",
         "Players simulate game-intensity 30-second shifts. Full-speed skating — forward sprints, tight turns, backward skating, transitions — as hard as possible for 30 seconds. Rest 90 seconds. Repeat 8-10 times. Track distance or effort.",
         "Every shift is game speed — no coasting. 30 seconds mirrors actual hockey shift length. Drive your legs the whole time. Rest period mimics sitting on the bench. This trains your body for the exact energy demands of a hockey game.",
         "Full ice. Individual or small groups. Whistle on/off.",
         12, 0, "full", None,
         U12_UP, '["conditioning","skating","game_simulation","intensity"]', "skating", "high", "30_second_shifts"),

        # ── MORE GOALIE DRILLS ──
        ("Rapid-Fire Shot Sequence", "goalie",
         "Three shooters set up at different positions (slot, left circle, right circle). Goalie faces rapid-fire shots — one shot from each position in quick succession. Goalie must recover and reset between each shot. Focus on tracking, movement, and recovery speed.",
         "Track the puck from the shooter's stick to your body. Move post to post efficiently — T-push or butterfly slide. Set your feet before the next shot. Don't just react — anticipate based on shooter position. Recovery is the key to facing multiple shots.",
         "One zone. 3 shooters. Goalie in net. 3 positions, rapid rotation.",
         12, 3, "quarter", "Pucks (bucket)",
         U12_UP, '["goalie","movement","recovery","tracking","intensity"]', "goalie", "high", "rapid_fire_sequence"),

        ("Breakaway Save Drill", "goalie",
         "Forwards attack on breakaways from the red line. Goalie practices challenge depth, patience, and staying big. Mix of deke attempts and shots. Focus on the goalie reading the shooter's hands and body position to anticipate the move.",
         "Challenge out aggressively but don't overcommit. Read the shooter: hands back = shot, hands forward = deke. Stay patient — let the shooter make the first move. Poke check only if you're 100% certain. Butterfly when the shooter gets to the hash marks. Stay big.",
         "Full ice. Forwards from red line. Goalie in net. One at a time.",
         12, 1, "full", "Pucks",
         U12_UP, '["goalie","breakaway","saves","patience","reading_play"]', "goalie", "high", "breakaway_save"),

        ("Post Integration Movement", "goalie",
         "Goalie practices post play — hugging the post when the puck is below the goal line. Coach moves the puck from corner to behind the net to the other corner. Goalie seals the post on each side, transitions across the crease, and resets. Add shots from low positions.",
         "Seal the post tight — no gaps between the pad and the post. Use the reverse VH or standard post lean depending on your system. Track the puck through the net or over the shoulder. When puck moves behind the net, get to the other post quickly — the shot comes fast off the pass.",
         "One zone. Goalie in net. Coach or player behind the net moving puck.",
         10, 1, "quarter", "Pucks",
         U10_UP, '["goalie","post_play","movement","tracking","positioning"]', "goalie", "medium", "post_integration"),

        # ── COOL DOWN DRILLS ──
        ("Controlled Skating Cool Down", "cool_down",
         "Easy laps around the ice at 50% effort. Focus on long, smooth strides with full recovery between pushes. Incorporate gentle stretches on the glide — open hips, reach for toes, twist trunk. 3-4 laps at a relaxing pace.",
         "Bring the heart rate down gradually. Full smooth strides — emphasize technique even at low speed. Breathe deeply. This is a good time to reinforce a positive practice moment with a quick word to each player.",
         "Full ice. All players skating together at easy pace.",
         5, 0, "full", None,
         ALL_AGES, '["cool_down","skating","recovery","stretching"]', "skating", "low", "cool_down_skate"),

        ("Shootout Fun", "cool_down",
         "End practice with a fun shootout. Each player gets one breakaway attempt. Goalie vs. the team. Make it fun — cheer for big saves and creative moves. This ends practice on a high note and gives everyone one last competitive moment.",
         "Keep it fun and light. Let kids try creative moves. Celebrate the goalie equally. This is about ending practice with smiles and a positive memory. Zero coaching points here — just let them play.",
         "Full ice. One goalie. All players get a turn.",
         5, 0, "full", "Pucks",
         ALL_AGES, '["cool_down","fun","shooting","compete"]', "shooting", "low", "shootout_fun"),

        ("Stick Skills Cool Down", "cool_down",
         "Players spread out on the ice and practice individual stick skills at low intensity. Toe drags, figure-8 stickhandling, between the legs, saucer tosses to themselves. Coach calls out different moves every 30 seconds. Creative and relaxing.",
         "This is low-intensity individual time. Players work at their own pace on skill moves. Encourage creativity — try something new. No pressure, just fun with the puck. Great way to build confidence in handling skills.",
         "Half ice. One puck per player. Spread out.",
         5, 0, "half", "Pucks",
         ALL_AGES, '["cool_down","puck_handling","stickhandling","fun","skill_development"]', "puck_handling", "low", "stick_skills_cooldown"),

        # ── MORE SMALL AREA GAMES ──
        ("4v4 No-Whistle Game", "small_area_games",
         "Full-zone 4v4 with no whistles. If the puck goes out of play, coach immediately fires a new one in. After a goal, defending team grabs a puck from behind their net and plays out immediately. Non-stop action builds conditioning and hockey sense.",
         "No time to rest — always be ready for the next puck. Transition instantly from offense to defense. Quick decisions — you don't have time to stickhandle. Move the puck and move your feet. This is the closest thing to a real game in practice.",
         "One zone. 4v4 plus goalie. Extra pucks behind each net.",
         12, 8, "half", "Pucks (multiple), small nets or full net",
         U10_UP, '["small_area_games","4_on_4","conditioning","compete","transition"]', "offensive", "high", "4v4_no_whistle"),

        ("Corners Game", "small_area_games",
         "3v3 in one zone. Can only score from below the hash marks (in the corners or from low slot). Forces players to work the cycle, drive low, and create scoring chances from the hard areas. Games to 3.",
         "This eliminates the lazy shot from the point. Players must go to the hard areas — below the hash marks, in front of the net, in the corners. Rewards net-front presence, cycling, and down-low battles. Real hockey is won in the dirty areas.",
         "One zone. 3v3 plus goalie. Only goals from below hash marks count.",
         12, 6, "quarter", "Pucks",
         U12_UP, '["small_area_games","3_on_3","cycling","net_front","compete"]', "offensive", "high", "corners_game"),

        # ── ADDITIONAL FUNDAMENTAL DRILLS ──
        ("Two-Touch Passing Drill", "passing",
         "Players in groups of 4 form a square 20 feet apart. The rule: you must receive the puck, make one stickhandle move, then pass to the next person (two touches maximum). Clock the group — how fast can you complete 20 passes around the square?",
         "First touch receives and controls. Second touch moves the puck. No extra handles allowed. This builds quick hands, soft receiving, and decision-making under time pressure. Close the blade on the receive to cushion the puck.",
         "Half ice. Groups of 4 in squares. One puck per group.",
         8, 4, "half", "Pucks",
         U10_UP, '["passing","quick_hands","decision_making","fundamentals"]', "passing", "medium", "two_touch_passing"),

        ("Deking Progression", "puck_handling",
         "Teach three basic dekes in progression: (1) forehand-to-backhand deke, (2) backhand-to-forehand deke, (3) fake shot then deke. Each player practices against cones first, then against a passive defender, then full speed. Use them in breakaway situations.",
         "Sell the first move with your eyes and body — make the defender commit. The puck moves last. Keep the puck close to the body during the deke. Hands out front, not beside you. Speed through the deke — don't slow down. The best deke is the one that freezes the defender.",
         "Half ice. Cones, then defenders. Progress to full speed dekes.",
         12, 0, "half", "Cones, pucks",
         U10_UP, '["puck_handling","deking","skill_development","offensive","1_on_1"]', "puck_handling", "medium", "deking_progression"),

        ("Wraparound Scoring Drill", "offensive",
         "Forward starts behind the net with the puck. On the whistle, attempts a wraparound — skating from behind the net and jamming the puck in at the far post before the goalie can get across. Practice both sides. Add a chasing defender for pressure.",
         "Speed is everything — the wraparound only works if you beat the goalie across. Keep the puck tight to the body behind the net. As you come around, extend the stick and jam the puck at the far post low. Use the post as a backboard. The goalie is moving — shoot for the open side.",
         "One zone. Forward behind net. Goalie in net. Both sides.",
         10, 1, "quarter", "Pucks",
         U12_UP, '["offensive","wraparound","scoring","speed","down_low"]', "offensive", "high", "wraparound_scoring"),

        ("2-on-2 Low-Zone Battle", "battle",
         "Two attackers and two defenders battle below the hash marks. Puck starts in the corner. Attackers try to score, defenders try to clear the zone. Fierce 20-second battles. Focus on body positioning, puck protection, and winning the inside lane.",
         "Low man wins — stay lower than your opponent. Attackers: protect the puck, find the trailer, get to the net front. Defenders: body on body, stick on puck, box out the net front. Every loose puck is a battle. This is where hockey games are won and lost.",
         "One zone below the hash marks. 2v2. Coach dumps to start.",
         10, 4, "quarter", "Pucks",
         U12_UP, '["battle_drills","2_on_2","down_low","compete","puck_protection"]', "battle", "high", "2on2_low_battle"),

        # ── FUN DRILLS (any age) ──
        ("Relay Race Puck Stacking", "fun",
         "Teams race to stack pucks on top of each other at center ice (like a tower). One player skates out, places a puck, and skates back. Next player goes. If the tower falls, you start over. First team to stack 5 pucks wins. Hilarious and builds team bonding.",
         "This is pure fun and team bonding. Players cheer for each other. The tension builds as the stack gets taller. Great for ending tough practices on a light note. Zero hockey development purpose — 100% morale and team chemistry.",
         "Full ice. Two teams. Pucks. Center ice.",
         5, 0, "full", "Pucks (10+)",
         ALL_AGES, '["fun","team_building","compete","relays"]', None, "low", "puck_stacking_relay"),

        ("British Bulldog", "fun",
         "Similar to Shark and Minnows but with pucks. All players start on one goal line with pucks. One or two taggers in the middle without pucks. Players must stickhandle across to the other side without losing their puck. Taggers try to knock pucks away. Lose your puck, become a tagger.",
         "Keep your head up. Protect your puck with body positioning. Read the taggers — find the gaps. Speed and agility win. This is a high-energy, high-fun game that teaches puck protection and awareness naturally.",
         "Full ice. One puck per player. 1-2 starting taggers.",
         8, 0, "full", "Pucks",
         ALL_AGES, '["fun","puck_handling","puck_protection","compete","agility"]', "puck_handling", "high", "british_bulldog"),

        ("Coach Says (Hockey Simon Says)", "fun",
         "Hockey version of Simon Says. Coach calls out hockey moves: Coach says do a snowplow stop, Coach says stickhandle between your legs, do a spin (but Coach didn't say!). Eliminated players practice shooting at the empty net. Fun way to practice moves.",
         "Pure fun that sneaks in skill work. Players practice moves without realizing they're drilling. Mix easy and hard moves. Be creative — Coach says do a celly, Coach says skate like a penguin. Last player standing is the champion.",
         "Half ice. All players spread out with pucks.",
         5, 0, "half", "Pucks",
         U8_U10, '["fun","skill_development","listening","fundamentals"]', None, "low", "coach_says"),
    ]

    for d in drills:
        try:
            conn.execute("""
                INSERT INTO drills (id, org_id, name, category, description, coaching_points, setup,
                    duration_minutes, players_needed, ice_surface, equipment, age_levels, tags,
                    skill_focus, intensity, concept_id)
                VALUES (?, NULL, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
            """, (str(uuid.uuid4()), *d))
        except Exception:
            pass  # Skip duplicates if partially seeded
    conn.commit()
    conn.close()
    logger.info("Seeded drills v2: %d additional drills (U8-focused expansion)", len(drills))


def seed_drills_pxi():
    """Seed 10 PXI-branded drills — advanced passing, offensive, SAG, special teams, defensive, systems, goalie, puck handling."""
    conn = get_db()
    exists = conn.execute("SELECT COUNT(*) FROM drills WHERE concept_id = 'quick_puck_support'").fetchone()[0]
    if exists > 0:
        conn.close()
        return

    U14_UP = '["U14","U16_U18","JUNIOR_COLLEGE_PRO"]'
    U16_UP = '["U16_U18","JUNIOR_COLLEGE_PRO"]'
    U12_U16 = '["U12","U14","U16_U18"]'
    GOALIE_ALL = '["U12","U14","U16_U18","JUNIOR_COLLEGE_PRO"]'
    GOALIE_14UP = '["U14","U16_U18","JUNIOR_COLLEGE_PRO"]'

    drills = [
        # 1. PXI Quick Support Touches — Passing
        ("PXI Quick Support Touches", "passing",
         "Coach rims or passes a puck to one of the wide forwards. That player immediately looks to the closest middle support for a quick give-and-go, then attacks the zone with speed. The far-side support player reads the play and either fills the high slot or drives the far post. The rush must include at least two quick touch passes before a shot on goal. Rotate roles every rep so all players work as wide and middle support.",
         "Wide forwards shoulder-check before receiving and move the puck quickly off the wall. Middle support players stay inside the dots and skate into open lanes with their sticks available. Passes are short, firm, and on the tape to keep speed through the neutral zone. Attack finishes with net drive and second-wave support for rebounds.",
         "Two lines of forwards at the blue line near the boards, one on each side. Two support players in the middle between the tops of the circles. Coach at center with pucks, one net and goalie.",
         12, 8, "half", "Net, Goalie, Pucks, Cones to mark middle support spots",
         U14_UP, '["transition","support","passing","rush"]', "passing", "high", "quick_puck_support"),

        # 2. PXI Low-High Tip Timing — Offensive
        ("PXI Low-High Tip Timing", "offensive",
         "Corner forward passes low-to-high to a point defenceman, then drives to the net for a screen. Defenceman walks the blue line and shoots for sticks, not corners. The second net-front forward times a lateral movement across the crease looking for tips. After the shot, the other corner forward becomes the next passer. Rotate positions regularly so everyone works on low play, point shots, and net-front timing.",
         "Point shots are low and through lanes, aimed at sticks and pads. Net-front players time their movement so they arrive as the shot is released. Corner forwards pass firmly to the point and then drive inside body position. Goalie tracks pucks from low to high and through traffic with strong head movement.",
         "Two D at the blue line with pucks, two forwards at the net front, and one forward in each corner. Goalie in net.",
         12, 7, "half", "Net, Goalie, Pucks, Cones to mark corner starting spots",
         U14_UP, '["screen","tips","low_to_high","net_front"]', "offensive", "medium", "low_high_tip_timing"),

        # 3. PXI Corner Trap 3v2 Game — Small Area Games
        ("PXI Corner Trap 3v2 Game", "small_area_games",
         "Coach dumps a puck into the corner to start each rep. Three attackers try to score against two defenders and a goalie. Once defenders win possession, they must make one controlled pass to the coach in the corner before chipping the puck above the top of the circles to clear. If attackers recover a cleared puck before it exits, play continues. Shifts run 30-40 seconds before switching groups.",
         "Attackers use quick support and rotation, keeping one player high for outlets. Defenders protect the middle first, then pressure when the puck settles. Communicate on switches and net-front box-outs to prevent backdoor plays. Short shifts maintain pace and game-like intensity.",
         "Half-ice from the goal line to the top of circles. One net and goalie. Play 3v2 inside the zone with a coach feeding pucks from the corner.",
         15, 10, "half", "Net, Goalie, Pucks, Cones to mark top of play area",
         U14_UP, '["small_area_game","3v2","battle","dz_coverage"]', "battle", "high", "3v2_corner_trap"),

        # 4. PXI Bumper Support Power Play — Special Teams
        ("PXI Bumper Support Power Play", "special_teams",
         "PP unit sets up with a middle bumper between the circles. Play begins from either flank. Flank player moves the puck low, then into the bumper, then up to the point for a shot or back to the far flank. Bumper must constantly adjust depth and angle to stay available. PK unit applies light pressure at first, then moves to more aggressive pressure as timing improves. Focus is on using the bumper as a pivot to change sides quickly.",
         "Bumper stays off defenders' sticks and presents a clear passing lane. Flank players attack downhill; they do not stand still on the wall. Point shots come after a side change to force goalie lateral movement. PK players read cues and apply smart pressure without losing box shape.",
         "5-on-4 in the offensive zone in a spread or 1-3-1 look. Coach at blue line with extra pucks.",
         15, 9, "half", "Net, Goalie, Pucks",
         U16_UP, '["power_play","bumper","special_teams","1_3_1"]', "special_teams", "medium", "pp_bumper_support"),

        # 5. PXI PK Triangle Collapse — Special Teams
        ("PXI PK Triangle Collapse", "special_teams",
         "PK starts in a triangle-plus-one look: three players inside the dots and one pressuring the puck. On a pass into the middle or a low seam, the three inside players collapse hard to protect the slot, forcing play back to the outside. If the PK recovers the puck, they must execute a hard clear and sprint to the far blue line for a simulated change. Rotate PK and PP roles every 30-40 seconds.",
         "Top PK player angles to take away the middle of the ice before pressuring. Inside players keep sticks in seams and collapse together, not individually. Talk through handoffs so no attacker is left unattended in the slot. Clears must be decisive and high enough to guarantee a change.",
         "4 PK players vs 5 PP players in the zone. Cones loosely mark slot area. Coach at blue with pucks.",
         12, 8, "half", "Net, Goalie, Pucks, Cones to outline collapse area",
         U16_UP, '["penalty_kill","collapse","slot_protection","special_teams"]', "special_teams", "high", "pk_triangle_collapse"),

        # 6. PXI DZ Dot-to-Dot Coverage — Defensive
        ("PXI DZ Dot-to-Dot Coverage", "defensive",
         "Offence moves the puck from low to high and across the blue line, looking for seams into the slot. Defenders use the faceoff dots as visual anchors: wingers stay outside the dots on their side, D stay inside the dots and protect net front. When the puck moves across, defenders shift dot-to-dot while keeping sticks in lanes. After a set number of passes or a shot, coach blows the whistle and a new group rotates in.",
         "Wingers stay between their point man and the net, not chasing wide. Defencemen own the front of the net and communicate switches on low cycles. Everyone's stick points to the puck first, body positioning second. Use quick stick-on-puck contact to disrupt shots and passes.",
         "4 offensive players cycle the puck from corner to point. 4 defenders and a goalie in the zone.",
         12, 8, "half", "Net, Goalie, Pucks",
         U16_UP, '["dz_coverage","lanes","structure","5v5"]', "defensive", "medium", "dz_dot_coverage"),

        # 7. PXI Neutral Zone Gate Pressure — Systems
        ("PXI Neutral Zone Gate Pressure", "systems",
         "Attackers must move the puck through one of the neutral-zone gates to continue the rush. Forecheckers use a 1-2-2 look, steering the puck toward the boards and closing the gate with strong sticks and body position. If defenders force a turnover before the gate, they transition quickly to attack the other way. Rotate groups every 3-4 reps.",
         "F1 angles the puck carrier toward the boards and the nearest gate. Second layer reads and jumps passing lanes without crossing over teammates. Defencemen close gaps through the gate, arriving under control. Attackers recognize when to chip past pressure instead of forcing through sticks.",
         "Place two gates with cones on each side of the red line near the boards. 5 attackers break out; 5 defenders set up to forecheck.",
         15, 10, "full", "Two nets, Pucks, Cones to mark gates",
         U16_UP, '["neutral_zone","forecheck","1_2_2","transition"]', "systems", "high", "nz_gate_forecheck"),

        # 8. PXI Goalie Box Movement Builder — Goalie
        ("PXI Goalie Box Movement Builder", "goalie",
         "Goalie T-pushes from the post to the top near puck, sets and holds, then shuffles across to the opposite top puck. From there, they T-push down to the far post puck, set, then shuffle back across the goal line to the original post. Repeat in both directions. Progress to adding a simple shot after any of the four positions.",
         "Explosive T-pushes with full extension but controlled stops at each puck. Shuffles are short and quick with minimal upper-body movement. Eyes lead every movement; head and shoulders follow, then feet. Set feet fully before simulating or facing a shot.",
         "Place four pucks around the edges of the crease forming a box. Goalie starts on one post in ready stance.",
         10, 1, "quarter", "Net, Goalie gear, Four pucks to mark box corners",
         GOALIE_ALL, '["goalie","crease_movement","t_push","shuffle"]', "goalie", "medium", "goalie_box_movement"),

        # 9. PXI Goalie Down-Up Recovery Chain — Goalie
        ("PXI Goalie Down-Up Recovery Chain", "goalie",
         "Goalie starts at the top of the crease. Coach shoots low, forcing a butterfly save. Goalie controls the rebound, recovers to their feet, and immediately shuffles to a new angle called by the coach (left dot, right dot, or high slot) for a second simulated shot. Sequence repeats 4-5 times per rep with no rest, building conditioning and recovery habits.",
         "Goalie seals ice on the first shot with good pad angle and stick position. Recover with hands and head leading, then one skate, then full stance. Stay compact and controlled when shuffling to the new angle. Maintain good posture even when fatigued late in the rep.",
         "Goalie in crease. Coach with pucks positioned in the slot.",
         8, 1, "quarter", "Net, Goalie gear, Pucks",
         GOALIE_14UP, '["goalie","recovery","conditioning","angles"]', "goalie", "high", "goalie_down_up_recovery"),

        # 10. PXI One-Touch Corner Escape — Puck Handling
        ("PXI One-Touch Corner Escape", "puck_handling",
         "Puck carrier starts with their back to the boards, under light pressure from the defender. They must execute a one-touch pass to the high support, then immediately spin off and jump to space for a return pass. After receiving the puck back, they attack the net or skate the puck up the wall to exit the zone. Rotate roles so everyone works as puck carrier, support, and defender.",
         "Puck carrier keeps feet moving and uses their body to shield the puck. One-touch passes are made off the boards or stick blade with purpose. Support player stays in a soft spot, not glued to the boards. Defender focuses on angling and stick pressure rather than big hits.",
         "Two players in the corner along the boards, one support player higher on the wall, and one defender applying light pressure.",
         10, 4, "quarter", "Net, Pucks, Cones to define corner area",
         U12_U16, '["puck_protection","support","battle","zone_exit"]', "puck_handling", "medium", "corner_escape_support"),

        # ── Batch 2: Transition, Offensive, Battle, Special Teams, Defensive, Systems, Goalie ──

        # 11. PXI Double-Swing Breakout — Transition
        ("PXI Double-Swing Breakout", "transition",
         "Coach chips a puck behind the net. Strong-side D retrieves and wheels up-ice, while weak-side D mirrors to the middle as a hinge option. Center swings low through the middle, first under the puck then up the weak side; both wingers time swings up their walls. D can hit the strong-side winger on the wall, the low swing center, or hinge to the partner who then hits the weak-side options. Once the puck exits the zone with control, the unit continues into a 3-on-2 rush against the same two D skating backwards from the blue line.",
         "Retrieving D shoulder-checks early and decides wheel, hinge, or middle based on pressure. Forwards skate routes with speed and timing instead of standing in their breakout spots. Passes are made inside the dots when possible to attack the middle with speed. Communication keywords such as wheel, hinge, and middle are used loudly and early.",
         "Five-player unit in the D-zone (2D, 3F). Coach at center ice with pucks to soft-chip behind the net. Wingers start at the hash marks on both walls; center starts between the dots.",
         12, 10, "full", "Pucks, Two nets, Cones to mark winger and center swing lanes",
         U16_UP, '["breakout","transition","swing_routes","3v2"]', "transition", "medium", "double_swing_breakout"),

        # 12. PXI Wall Rim Retrieval — Transition
        ("PXI Wall Rim Retrieval", "transition",
         "D starts at the dot line, facing up ice. On the whistle the coach rims a puck around the glass. D skates back, shoulder-checks, and selects either a quick bump to the winger on the wall, a reverse behind the net to their partner cone, or a quick-up to the hash marks. The winger times their route down to the hash marks and then up the wall, presenting their stick as a target and calling for the puck. After a clean breakout pass, the winger cuts to the middle and shoots from the top of the circle.",
         "Defenceman must check over both shoulders before deciding what to do with the rim. Winger arrives on the wall as the puck is settling, not waiting flat-footed. Use the wall as a tool: cushions, chips, and bump passes executed with purpose. Head up on retrieval and exit to see and use the middle of the ice when available.",
         "One D and one winger per side. Coach on the opposite blue line rims pucks hard around the boards. Net in place with goalie optional.",
         10, 6, "half", "Pucks, Net, Cones to mark D and winger starting spots",
         U16_UP, '["breakout","rim","wall_support","retrieval"]', "transition", "medium", "rim_retrieval_support"),

        # 13. PXI Stretch Pass Release — Transition
        ("PXI Stretch Pass Release", "transition",
         "D1 skates behind the net with a puck and hits D2 for a D-to-D pass. As D2 receives, the wide forward on the far blue line times a slash cut toward the middle, while the middle forward stretches wide to the opposite boards. D2 can hit either the slash cut in the middle or the wide stretch forward with a long pass. The receiving forward attacks the far net with speed for a shot, supported by the other forward for a rebound. Rotate D after several reps.",
         "Defencemen must change the passing angle quickly with their feet, not just their hands. Forwards time slash and stretch routes so they are moving toward the puck, not away. Long passes stay flat and are aimed at the inside hip for easy reception in stride. Attack with width and middle-lane drive to create a second wave and rebound support.",
         "Two D at one end with pucks. Two forwards line up at the far blue line on the boards; another forward lines up in the middle of the far zone.",
         12, 8, "full", "Pucks, Two nets, Cones to mark slash and stretch lanes",
         U16_UP, '["transition","stretch_pass","long_pass","rush"]', "transition", "high", "stretch_pass_breakout"),

        # 14. PXI Middle-Lane Drive 2v1 — Offensive
        ("PXI Middle-Lane Drive 2v1", "offensive",
         "Coach rims or passes a puck to one forward line. That forward becomes the puck carrier and drives wide up the boards. The opposite forward times a delayed middle-lane route, starting slightly behind the puck carrier and driving hard to the far post. The D gaps up from between the dots and plays the 2-on-1. The puck carrier reads the D: if the lane to the net is open, drive and shoot; if D commits, slide a pass to the driving middle forward for a tap-in or quick shot.",
         "Puck carrier keeps their feet moving and attacks the dot line before making a decision. Middle-lane forward drives hard to the far post and stays stick available. Defenceman maintains good gap while keeping stick in the passing lane first. Both forwards stop at the net for rebounds instead of circling away.",
         "Two F lines at the red line near the boards on opposite sides. One D line at center ice between the dots. Coach at center with pucks.",
         12, 6, "full", "Pucks, Two nets",
         U16_UP, '["rush","2v1","middle_drive","offensive_zone_entry"]', "offensive", "high", "middle_drive_2v1"),

        # 15. PXI Wide Entry Delay Options — Offensive
        ("PXI Wide Entry Delay Options", "offensive",
         "Wide forward receives a pass at center and skates wide toward the offensive blue line with a defender matching gap. The trail forward follows through the middle lane. As the wide forward crosses the blue line, they execute a delay at the top of the circle, turning back toward the boards while protecting the puck. Options: hit the trailing forward driving into the slot, use a drop pass just inside the blue for a quick shot, or chip the puck behind the D and skate through to retrieve. Rotate roles after each rep.",
         "Wide forward sells the attack first, then uses a sharp delay with body between puck and D. Trail forward reads the delay and adjusts speed to arrive in the scoring area at the right time. Defenceman manages gap and stays between puck and net, not chasing behind the play. Passes out of the delay are made off the inside edge with head up to see all options.",
         "Two F lines at center on the boards; one trail F line in the middle. One D line at the defending blue line.",
         12, 6, "full", "Pucks, Two nets",
         U16_UP, '["zone_entry","delay","support","2v1_like"]', "offensive", "medium", "wide_entry_delay"),

        # 16. PXI Net-Front Layered Screens — Offensive
        ("PXI Net-Front Layered Screens", "offensive",
         "Coach slides a puck to the point shooter who walks laterally along the blue line. Net-front player sets a heavy screen at the top of the crease, while the bumper hovers between the circles. On the whistle, the shooter takes a shot through traffic. Net-front player boxes out for tips and rebounds; bumper reads the shot off the goalie pads and looks for quick touch plays. Rotate roles every few shots.",
         "Point shooter keeps shots low and through lanes, not into shin pads. Net-front player establishes inside body position and moves with the goalie. Bumper keeps their stick in a ready position and scans for loose pucks in the slot. All players stop at the net until the rep is clearly over.",
         "One point shooter at the blue line, one net-front player at the top of the crease, one bumper a few feet above, and a goalie. Coach with pucks at the blue line.",
         10, 4, "quarter", "Net, Goalie, Pucks",
         U16_UP, '["offensive_zone","screen","tips","rebound"]', "offensive", "medium", "net_front_layers"),

        # 17. PXI Corner Cutback Cycle — Offensive
        ("PXI Corner Cutback Cycle", "offensive",
         "Coach rims a puck into the corner. F1 races to the puck and looks to drive up the wall. As D pressures, F1 executes a hard cutback toward the boards, changing direction back toward the corner. F2 times a support route along the wall and receives a short cycle pass from F1. After the cycle, F1 drives to the net for a return pass or screen, while F2 walks to the middle or hits the net-front stick. D defends with proper body position and stick on puck.",
         "F1 sells the up-wall drive before executing a sharp cutback with strong edges. F2 keeps feet moving and times their route so they arrive just as F1 pivots. Puck stays to the outside away from the defender stick during the cutback. Defenceman maintains inside body position and does not over-commit on the first move.",
         "2F vs 1D below the top of the circles in the offensive zone. Coach in the corner spots pucks.",
         12, 6, "half", "Pucks, Net, Cones to mark the top of the play area",
         U16_UP, '["cycling","battle","offensive_zone","2v1_low"]', "offensive", "high", "cutback_cycle"),

        # 18. PXI Slot Support Triangle — Offensive
        ("PXI Slot Support Triangle", "offensive",
         "Low forwards work the puck behind the net and along the goal line while the high forward slides in the slot, always presenting a passing lane. On coach whistle, the puck must move quickly between low and high positions, forcing defenders to adjust. The objective is to create a quick shot from the high slot with net-front traffic or a backdoor tap-in from one of the low forwards. After each 20-25 second shift, switch groups.",
         "Low forwards keep their feet moving and protect the puck with body position and the net. High forward never stands still; they constantly adjust depth and angle to stay available. Passes are snapped through open seams quickly before defenders reset. Defenders communicate and hand off low coverage instead of chasing.",
         "Three offensive players in a triangle (two low near each post and one high in the slot) vs two defenders and a goalie.",
         10, 5, "quarter", "Net, Goalie, Pucks",
         U16_UP, '["offensive_zone","support","3v2_low","scoring"]', "offensive", "medium", "slot_support_triangle"),

        # 19. PXI Half-Ice 3v3 Transition Game — Small Area Games
        ("PXI Half-Ice 3v3 Transition Game", "small_area_games",
         "Play continuous 3v3 in half-ice. When the defending team wins the puck, they must make one controlled pass to a teammate below the hash marks before they can attack the far net. On a goal or a clear over the blue line, the scoring or exiting team stays, and a fresh trio from the other team jumps in with a new puck. Emphasize quick transition from defence to offence and support options away from the puck.",
         "Players must open up and present sticks immediately on change of possession. Quick, short passes build possession before attacking the net. Defensive sticks stay in passing lanes and bodies stay inside the dot lines. Short shifts at high tempo mimic junior game pace.",
         "Half-ice with nets on the goal line. Two teams of 3 active players, with subs waiting at the blue line. Coach at center with pucks.",
         15, 10, "half", "Two nets, Pucks, Dividers if available",
         U16_UP, '["small_area_game","transition","3v3","compete"]', "battle", "high", "3v3_transition_half_ice"),

        # 20. PXI Board Battle to Net Drive — Battle Drills
        ("PXI Board Battle to Net Drive", "battle",
         "Coach chips a puck to the boards between the two players. They battle 1v1 along the wall, working to establish body position and puck control. The player who wins possession must immediately drive off the wall into the middle and attack the net for a shot while the defender tries to angle and strip the puck. After the shot or clear, players return to the line and the next pair goes.",
         "Players use their hips and shoulders to seal the opponent off the wall. Stick is strong on the puck with bottom hand firm and top hand away from the body. Winner quickly leaves the wall and attacks the middle instead of drifting low. Defender angles through the hands and stick, not reaching from behind.",
         "Pairs of players along the wall at the hash marks with a net at the near post and goalie optional. Coach with pucks at the blue line.",
         10, 4, "quarter", "Net, Pucks",
         U16_UP, '["battle","1v1","compete","net_drive"]', "battle", "high", "wall_battle_net_drive"),

        # 21. PXI Corner Escape 1v1 — Battle Drills
        ("PXI Corner Escape 1v1", "battle",
         "On the whistle, both players battle for the puck in tight space. The offensive player goal is to escape the corner and either cut to the net or pass the puck out to a coach at the top of the circle for a quick return pass and shot. The defender attempts to pin, angle, and separate the attacker from the puck, then clear it out of the zone.",
         "Offensive player keeps knees bent and uses quick cutbacks and shoulder fakes to escape. Use the boards as protection, rolling off contact instead of backing straight away. Defender keeps stick on puck and finishes checks through the body, not just reaching. Short, intense reps encourage hard battles without fatigue-driven mistakes.",
         "One offensive player and one defender start in the corner with their backs to the boards. Coach places a puck at their feet. Net at the near post with goalie optional.",
         10, 4, "quarter", "Net, Pucks, Cones to define the corner battle area",
         U16_UP, '["battle","corner","compete","1v1_low"]', "battle", "high", "corner_escape_1v1"),

        # 22. PXI Neutral-Zone Kill 1-1-3 — Systems
        ("PXI Neutral-Zone Kill 1-1-3", "systems",
         "Attackers start behind their own net and execute any controlled breakout. As they advance, the defensive team sets up a 1-1-3: F1 pressures high, F2 holds middle, and three players form a tight line across the defensive blue. Attackers attempt to gain the offensive zone with control using regroups, chips, or width plays. If defenders force a turnover or an offside, the rep resets from the original end.",
         "F1 angles the puck carrier toward the strong side while keeping speed under control. Middle player protects the center lane and supports whichever side F1 forces the play. Back line holds the blue line with tight gaps and good stick position. Attackers must recognize when to chip behind the line versus forcing controlled entries.",
         "Five attacking players attempt to break out and attack through the neutral zone. Five defenders set up in a 1-1-3 neutral-zone structure.",
         15, 10, "full", "Pucks, Two nets, Whiteboard to show 1-1-3 alignment",
         U16_UP, '["neutral_zone","system","forecheck","5v5"]', "systems", "medium", "neutral_zone_1_1_3"),

        # 23. PXI DZ Swarm to Box — Defensive
        ("PXI DZ Swarm to Box", "defensive",
         "Offensive group works the puck below the tops of the circles. Defenders start in a tight swarm around the puck carrier, applying pressure and looking to outnumber at the point of attack. On the coach whistle, play transitions to a more structured box: two low, two high, each taking away seams and middle ice while still pressuring when the puck settles. Rotate groups every 30-40 seconds.",
         "Swarm phase: closest two defenders pressure, while the other two read and support. Box phase: players snap back into clear quadrants, keeping sticks inside and bodies outside. Communication drives coverage handoffs as the puck moves from low to high. Defenders finish reps with a clear and quick transition to offence when they win possession.",
         "Four offensive players cycle the puck low in the zone versus four defenders and a goalie. Coach at blue line with pucks.",
         12, 8, "half", "Net, Goalie, Pucks",
         U16_UP, '["dz_coverage","system","pk_like","4v4"]', "defensive", "medium", "dz_swarm_to_box"),

        # 24. PXI Point Shot Lane Denial — Defensive
        ("PXI Point Shot Lane Denial", "defensive",
         "Point players work pucks laterally along the blue line, looking for shooting lanes. Defending forwards stay between the shooters and the net, using sticks and body position to block lanes without overcommitting. On a shot, the forwards must either block, deflect away from danger, or box out and clear rebounds. After each short rep, rotate defenders and shooters.",
         "Defenders angle their bodies to block more net while keeping eyes on both puck and traffic. Stick stays in the lane first; then body follows to block if needed. After a block or rebound, players recover quickly and locate the puck, not just the shooter. Shooters practice moving feet to change their release angle and challenge the lane control.",
         "Two point shooters on the blue line with pucks, two forwards in the high slot defending lanes, and a goalie.",
         10, 6, "quarter", "Net, Goalie, Pucks",
         U16_UP, '["dz_coverage","shot_blocking","lanes","defensive_zone"]', "defensive", "medium", "point_shot_lane_denial"),

        # 25. PXI Goalie Post Bump T-Drill — Goalie
        ("PXI Goalie Post Bump T-Drill", "goalie",
         "Goalie starts on their glove-side post in reverse-VH or ready stance. On the coach call, they execute a bump off the post into the middle of the crease, set square to an imaginary shot from the slot, then T-push to the opposite post and seal. Coach then passes a puck from one of three locations (left circle, right circle, or slot) to simulate a quick play. Goalie must adjust angle and make the save, then recover back to the original post and repeat.",
         "Explode off the post with a compact bump movement, arriving balanced in the middle. Eyes and head lead every adjustment; body follows. T-pushes are controlled with full extension but precise stops before each set. Recover to feet quickly after saves and re-establish post or middle position.",
         "One goalie in the crease. Coach with pucks positioned at the top of both circles and in the slot.",
         10, 1, "quarter", "Net, Goalie gear, Pucks",
         GOALIE_14UP, '["goalie","crease_movement","post_play","recovery"]', "goalie", "high", "goalie_post_bump_t"),

        # 26. PXI Reverse VH Wraparound Read — Goalie
        ("PXI Reverse VH Wraparound Read", "goalie",
         "Shooter skates from below the goal line, threatening to walk out short side or drive behind the net for a wraparound. Goalie begins in reverse-VH on the post and must read the puck carrier route: if the attacker walks up the wall, the goalie releases to a regular stance and moves out to challenge; if the attacker drives behind the net, the goalie pushes across the goal line and seals the far post to stop the wrap. Alternate sides every few reps to work both posts.",
         "Goalie maintains patience on the post and reads stick position and body angle of attacker. Push along the goal line should be powerful but controlled, staying tight to the posts. Hands stay active in front of the body even when in reverse-VH. Communication with defenders helps identify backdoor threats during walk-outs.",
         "One goalie in net. Shooter starts below the goal line on either side with pucks.",
         10, 2, "quarter", "Net, Goalie gear, Pucks",
         GOALIE_14UP, '["goalie","post_play","wraparound","reads"]', "goalie", "medium", "goalie_reverse_vh_wrap_read"),

        # 27. PXI Faceoff Win Quick Strike — Special Teams
        ("PXI Faceoff Win Quick Strike", "special_teams",
         "Run a designed offensive-zone faceoff play. Center aims to win the puck straight back to the strong-side D. Winger on the wall ties up their check; weak-side winger cuts through the middle for a quick touch pass option. Upon the win, D walks to the middle and either shoots through traffic, hits the middle-cut winger, or fakes and slides the puck to the weak-side D for a one-timer. Run from both sides and with different alignments to practice multiple quick-strike options.",
         "Center focuses on stick speed and body leverage to win pucks cleanly. Wingers have specific jobs: tie up sticks, create traffic, or cut to space immediately. Defencemen must keep shots low and on net with screens in place. Everyone knows the first and second options before the puck is dropped.",
         "Offensive-zone faceoff with a full 5-man unit and a goalie. Coach drops pucks as the official.",
         10, 6, "quarter", "Net, Goalie, Pucks, Faceoff dots",
         U16_UP, '["faceoff","special_teams","set_play","offensive_zone"]', "special_teams", "medium", "oz_faceoff_quick_strike"),

        # 28. PXI PK Clear and Change — Special Teams
        ("PXI PK Clear and Change", "special_teams",
         "PP unit works the puck around the zone while the PK stays in its structure. When the PK wins possession, they must execute a hard clear off the glass or up the middle and then sprint to the far blue line for a simulated change. Coach quickly rims a new puck back in to force the next PK group to establish structure under pressure. Rotate PK groups quickly to build conditioning and habits.",
         "PK players think first touch, first clear when they gain control under pressure. Use the glass or middle ice with enough height and length to guarantee a change. After clearing, players skate hard off the ice line they are responsible for. Communication on entry resets helps the new PK group get into formation quickly.",
         "PP unit vs PK unit in the offensive zone with a goalie. Coach at blue line with pucks.",
         10, 8, "half", "Net, Goalie, Pucks",
         U16_UP, '["penalty_kill","special_teams","clear","conditioning"]', "special_teams", "high", "pk_clear_and_change"),

        # ── Batch 3: More Transition, Systems, Special Teams ──

        # 29. PXI Quick-Up Wall Release — Transition
        ("PXI Quick-Up Wall Release", "transition",
         "On the whistle, D retrieves a spotted puck below the goal line, shoulder-checking and skating up-ice behind the net. The strong-side winger times their route up the wall to receive a quick-up pass on the boards. Upon receiving the puck, the winger immediately bumps it to the coach at the far blue line, then jumps to open space through the middle. The coach one-touches the puck back to the winger or driving center for a full-speed attack 2-on-0. After the rush, players hustle back on the opposite side and join the next rep.",
         "Defenceman checks both shoulders before picking up the puck and moves their feet up-ice. Winger times the route so they arrive on the wall as the D is ready to pass, not early and standing still. Passes are firm, tape-to-tape, and made in stride to maintain speed through the neutral zone. Attackers drive with width and middle-lane speed to create a strong-side lane and a middle option.",
         "Two D lines at each end below the goal line with pucks. Two F lines on the strong-side wall at the hash marks on both ends. One coach at the far blue line on each side as a neutral outlet.",
         12, 8, "full", "Pucks, Nets at both ends, Cones to mark winger wall starting spots",
         U16_UP, '["breakout","transition","quick_up","wall_support"]', "transition", "high", "quick_up_breakout"),

        # 30. PXI Low Reverse Breakout Read — Systems
        ("PXI Low Reverse Breakout Read", "systems",
         "Coach rims or soft dumps a puck into the corner. The strong-side D retrieves with speed, scanning the ice while the weak-side D slides behind the net as a reverse option. Wingers track back to their walls, and the center activates low through the middle. If the forecheck pressure is light on the retrieval side, D executes a direct breakout up the strong-side wall. If pressure comes hard on that side, D calls reverse and uses the weak-side D, who then has options to hit the low center or the weak-side winger. After a successful breakout past the blue line, the group continues up ice for a controlled 3-on-2.",
         "Retrieving defenceman must scan early and often to decide between direct and reverse options. Weak-side D gets their skate behind the net quickly and presents a clear target for the reverse. Wingers pull back below the hash marks and are prepared to adjust to strong- or weak-side support. Center stays low, available, and communicates the read to help drive the breakout decision.",
         "One full unit of 5 in the D-zone: two D, three F. Two forecheckers start at the offensive blue line. Coach at center with pucks.",
         12, 10, "full", "Pucks, Two nets, Cones to mark winger and forechecker starting spots",
         U16_UP, '["breakout","dz_system","reverse","unit_play"]', "systems", "medium", "reverse_breakout_read"),

        # 31. PXI 1-2-2 Forecheck Install — Systems
        ("PXI 1-2-2 Forecheck Install", "systems",
         "Begin by walking players through the 1-2-2 alignment: F1 pressures the puck carrier, F2 and F3 stagger in the middle of the ice, and D hold the red line with tight gaps. Once responsibilities are clear, run live reps from a controlled breakout. The breakout unit exits the zone and attempts to attack through the neutral zone. The forecheck unit sets up their 1-2-2, with F1 steering the puck toward a wall, F2 and F3 reading off each other to seal middle options, and D closing gaps and killing speed at the blue line.",
         "F1 does not fly past the puck; angle and contain while steering play to the chosen side. F2 and F3 maintain inside positioning and communicate which player has the middle lane. Defencemen hold a tight gap at the red line and match the speed of the attack. All five forecheckers move together as a connected unit with consistent spacing.",
         "Two units of 5. One unit breaks out, the other forechecks. Start with a static walk-through in the neutral zone, then progress to live reps.",
         15, 10, "full", "Pucks, Two nets, Marker or board at bench to review alignment",
         U16_UP, '["forecheck","neutral_zone","system_install","5v5"]', "systems", "medium", "forecheck_1_2_2"),

        # 32. PXI 1-3-1 Power Play Flow — Special Teams
        ("PXI 1-3-1 Power Play Flow", "special_teams",
         "Players take their 1-3-1 positions and move the puck through a set passing pattern to rehearse spacing and timing. Start with a simple wheel: point to flank, flank to bumper, bumper to opposite flank, back to point, then shot with net-front screen and bumper crash. On the next rep, introduce a low play: point to flank, flank down to net-front, quick seam into bumper, then either shot from the middle or touch pass back door. Run sequences on both sides.",
         "Keep the puck moving quickly; no player holds it longer than a second unless attacking. Net-front player maintains inside body position and adjusts to sight lines for the shooter. Bumper stays available in the middle, not buried in traffic, and presents a clear target. Flank players attack downhill when they see a lane instead of passing by default.",
         "Set up in the offensive zone with one net and goalie. Five skaters in a 1-3-1 structure: point, two flanks, bumper, and net-front. Coach at blue line with extra pucks.",
         15, 5, "half", "Net, Goalie, Pucks, Cones to mark 1-3-1 positions if needed",
         U16_UP, '["power_play","1_3_1","special_teams","offensive_zone"]', "special_teams", "medium", "pp_1_3_1_flow"),

        # 33. PXI Aggressive Box Penalty Kill — Special Teams
        ("PXI Aggressive Box Penalty Kill", "special_teams",
         "PK unit sets up in a compact box in front of the net while the PP unit works the perimeter. On the coach signal, the PP begins moving the puck around the outside. PK skaters shift as a unit, keeping sticks in lanes and taking away the middle. Any time the puck is bobbled, held too long at the half wall, or enters the corner, the nearest PK forward jumps to pressure aggressively while the other three players tighten and support. If the PK gains possession, they must skate or chip the puck over the far blue line to complete the rep.",
         "PK sticks stay in lanes first; body contact comes after the passing option is removed. Top forwards of the box communicate which one pressures and which one protects the middle. Defencemen keep inside body position, owning the net front and slot, not chasing into corners. On a clear, players sprint to their next shift position instead of watching the puck.",
         "One PP unit (5 skaters) vs one PK unit (4 skaters) in the offensive zone with a goalie. Coach at blue line with pucks.",
         12, 9, "half", "Net, Goalie, Pucks, Cones to loosely outline the PK box if needed",
         U16_UP, '["penalty_kill","special_teams","box","pressure"]', "special_teams", "high", "pk_aggressive_box"),

        # ── Batch 4: Goalie, Offensive, Warm Up, SAG, Passing, Puck Handling ──

        # 34. PXI Goalie Angle Landmarks — Goalie
        ("PXI Goalie Angle Landmarks", "goalie",
         "Goalie starts centered at the top of the crease facing the coach in the high slot. Coach calls out different landmarks such as left dot, right post, or point. The goalie shuffles or T-pushes into position so their body and stick are square to the called landmark, then holds for a brief pause before returning to center. Progress to adding simple wrist shots from each location once the goalie consistently finds their angles.",
         "Goalie leads each movement with eyes and head, then shoulders, then feet. Stick stays centered between the skates with blade on the ice. Goalie tracks their relationship to the posts and top of the crease as visual anchors. Recover to a balanced stance at center after every angle adjustment.",
         "One goalie in the crease. Coach places small markers or pucks at visual landmarks along the top of the crease and on each post.",
         12, 1, "quarter", "Net, Goalie gear, Pucks or markers for landmarks",
         GOALIE_ALL, '["goalie","angles","crease_movement","visual_cues"]', "goalie", "medium", "goalie_angle_control"),

        # 35. PXI Six-Puck Breakaway Race — Offensive
        ("PXI Six-Puck Breakaway Race", "offensive",
         "On the whistle, the first forward from each line races to the nearest puck on the red line, picks it up, and attacks the far net on a breakaway. After their shot, they loop back through neutral ice, collect the next puck in their lane, and repeat. Continue until each player has taken three to four breakaways at full speed. Run in short, competitive heats between pairs or small groups.",
         "Players accelerate quickly through the puck, not to the puck. Head up early to read goalie positioning and choose a move or shot. Encourage creativity but demand full-speed entries and hard stops at the net. Goalies focus on patient depth and strong lateral pushes on dekes.",
         "Place six pucks spaced along the center red line. Two lines of forwards start on opposite sides of center. One goalie in each net if available.",
         10, 6, "full", "Two nets, Goalies if available, Six pucks per lane",
         U14_UP, '["breakaway","speed","finishing","compete"]', "offensive", "high", "breakaway_compete_race"),

        # 36. PXI Bednar Edge Flow Warm-Up — Warm Up
        ("PXI Bednar Edge Flow Warm-Up", "warm_up",
         "First player from each line skates forward to the first cone, performs a crossover turn, then continues to the next cone. At each cone, they alternate between forward-to-backward and backward-to-forward pivots while maintaining puck control. After the last cone, they accelerate in a straight line, take a shot on net, then join the opposite line for the return route.",
         "Players stay low with strong knee bend through each pivot. Encourage full extension on crossover strides to build power. Hands stay away from the body to protect the puck while edging. Eyes scan up-ice instead of staring at the puck.",
         "Players split into two lines in the corner. Cones form a zig-zag path to the far blue line and back with pivot points at each cone.",
         8, 8, "half", "Cones, Two nets, Pucks",
         GOALIE_ALL, '["skating","pivots","warm_up","puck_control"]', "skating", "medium", "edge_control_flow"),

        # 37. PXI Center Boundary 2v1 Game — Small Area Games
        ("PXI Center Boundary 2v1 Game", "small_area_games",
         "Two attackers and one defender play inside a narrow lane where the boards and center cones act as boundaries. Coach spots a puck to the attacking pair, who must create a scoring chance without crossing the lane markers. The defender works to angle, take away passing lanes, and force low-percentage shots. After a short rep or a goal, new trios rotate in quickly.",
         "Attackers maintain good spacing horizontally and vertically in the lane. Puck carrier attacks the defender inside shoulder to open a pass or lane. Defender keeps stick in the passing lane first and maintains inside body position. Fast rotations and short shifts keep pace and compete level high.",
         "Divide half ice into two lanes using cones along the center line. One net at each end. Play 2v1 inside each lane.",
         12, 8, "half", "Two nets, Pucks, Cones to mark lane boundaries",
         U14_UP, '["small_area_game","2v1","angling","spacing"]', "battle", "high", "lane_2v1_transition"),

        # 38. PXI Give-and-Go Corner Route — Passing
        ("PXI Give-and-Go Corner Route", "passing",
         "Wall player passes down to the corner, then cuts toward the middle of the ice. Corner player returns the pass to the moving wall player in the slot for a shot. After shooting, the player circles back, collects a puck behind the net from the coach, and passes back to the next player in line to keep the give-and-go pattern going.",
         "Passer points their stick blade and follows through toward the target. Receiver moves into open ice before calling for the return pass. Encourage one-touch or quick-release shots in the slot. Players should open up their hips to receive on the forehand when possible.",
         "One line of players on the half-wall, one line at the corner dot, net at near post. Coach or extra player stands behind the net as a passer.",
         8, 4, "quarter", "Net, Pucks, Cones to define wall and corner lines",
         U12_U16, '["passing","give_and_go","youth","shooting"]', "passing", "medium", "give_and_go_route"),

        # 39. PXI Long-Short Passing Rhythm — Passing
        ("PXI Long-Short Passing Rhythm", "passing",
         "Player A starts with a puck on the boards and makes a short pass to Player B in the middle. Player B immediately returns the pass, then opens up for a long cross-ice pass from A to Player C. Player C bumps the puck back to B in the middle for a shot from the high slot. Players follow their pass to the next station, maintaining a continuous pattern.",
         "Short passes are crisp and flat; long passes carry more weight but stay on the ice. Middle player scans both sides before each touch to build habit of checking shoulders. Feet continue to move before, during, and after each pass. Call for every pass to reinforce communication and timing.",
         "Three players spaced across the width of the ice at the blue line and opposite faceoff dot. Additional players rotate through lines.",
         10, 6, "half", "Net, Pucks, Cones for station spacing",
         U12_U16, '["passing","timing","shooting","flow_drill"]', "passing", "medium", "long_short_passing"),

        # 40. PXI Musical Edge Pucks — Puck Handling
        ("PXI Musical Edge Pucks", "puck_handling",
         "Each player skates clockwise around the circle without a puck, focusing on crossovers and edge work. Several pucks are placed randomly inside the circle. When the music stops or whistle blows, players race into the middle, claim a puck, and stickhandle back to an edge cone. One player will be left without a puck and performs a quick skating task before the next round.",
         "Players stay low with powerful crossovers around the circle. Quick transition from skating pattern to puck control when the whistle blows. Encourage heads-up handling and protecting the puck from other players. Keep rounds short so intensity stays high and players remain engaged.",
         "Circle area with one puck fewer than the number of players. Music or whistle controls start and stop.",
         8, 6, "quarter", "Pucks, Cones, Whistle or music source",
         '["U8","U10","U12"]', '["youth","edges","game","puck_protection"]', "puck_handling", "medium", "edges_with_puck_game"),

        # 41. PXI Three-Zone Timing Weave — Passing
        ("PXI Three-Zone Timing Weave", "passing",
         "First players from each line leave together. The middle lane skater starts with the puck and passes to the outside lane, then skates behind that player to fill the wide lane. This weave continues through all three zones, with the puck always moving to the player entering the middle lane. At the far end, the last receiver attacks the net for a shot while the other two players drive for rebounds. Next group goes once the offensive blue line is cleared.",
         "Timing is everything: players adjust speed so spacing between them stays consistent. Passes are made early, before the player crosses the next blue line. Weave routes should be deliberate figure-eights, not random crossing. All three players finish hard to the net, reading second and third-chance opportunities.",
         "Three lines of players at one end across the width of the ice, each with pucks. Cones at far blue line and opposite end to mark routes.",
         12, 9, "full", "Two nets, Pucks, Cones for lane markers",
         U14_UP, '["timing","passing","flow","3_man_weave"]', "passing", "high", "three_zone_weave_timing"),

        # 42. PXI Quarter-Ice Continuous Cycle — Offensive
        ("PXI Quarter-Ice Continuous Cycle", "offensive",
         "Coach dumps a puck into the corner. F1 retrieves and cycles the puck up the wall to F2, then drives to the net. F2 walks the wall, reads pressure, and either shoots or passes low to the driving F1. After a shot, the coach immediately spots another puck to the opposite corner and roles rotate: the net-front player becomes the next retriever, and the previous retriever becomes support on the wall. Defender plays honest 1v2 defence throughout.",
         "Forwards keep their feet moving and use the boards to protect the puck. Net-front player establishes inside position and is ready for quick passes. Defender focuses on stick position and angling rather than chasing both players. Short, continuous reps build conditioning and reinforce cycle habits.",
         "Two forwards and one defender in a quarter-ice zone with a net and goalie. Coach with pucks at the blue line.",
         12, 6, "quarter", "Net, Goalie, Pucks, Cones to mark quarter-ice boundary",
         U14_UP, '["cycling","2v1_low","offensive_zone","battle"]', "offensive", "high", "continuous_low_cycle"),

        # 43. PXI Forecheck Funnel Progression — Systems
        ("PXI Forecheck Funnel Progression", "systems",
         "Start with a walk-through: F1 angles the puck carrier toward the boards, F2 and F3 fill the middle lanes, and both D hold the red line to close space. Progress to live reps where the breakout unit attempts a controlled entry while the forecheckers work as a unit to funnel the play into a trap at the boards. If the forecheck group forces a turnover, they transition quickly to offence and attack the other way.",
         "F1 skates a controlled route that takes away the middle and pushes play wide. F2 and F3 read off F1 and keep sticks positioned to block middle passes. Defencemen maintain tight gap and are ready to step up when the puck turns. On a turnover, all five forecheckers immediately switch to attack mindset.",
         "Five attackers break out from behind their net. Five defenders set up a simple forecheck in the neutral zone. Coach at center with extra pucks.",
         15, 10, "full", "Two nets, Pucks, Whiteboard to show forecheck shape",
         U16_UP, '["forecheck","systems","neutral_zone","transition"]', "systems", "medium", "forecheck_funnel"),

        # ── Batch 5: Attack Triangle, Centering Pass ──

        # 44. PXI Attack Triangle Foundations — Offensive
        ("PXI Attack Triangle Foundations", "offensive",
         "Coach passes to any of the three forwards to begin the rep. The three attackers immediately form an attack triangle: puck carrier wide, one support player driving the far post, and the third player filling high slot space. They must maintain triangle spacing as they move, using short passes, give-and-gos, and drive lanes to create a quality shot. After the shot, all three stop at the net for rebounds before circling back to the line.",
         "Maintain a clear triangle with one player wide, one middle, and one high. Puck carrier attacks the dot line before deciding to shoot or pass. Off-puck players keep sticks available and adjust their depth to stay open. All three attackers stop at the net after the shot instead of skating past the crease.",
         "Two forwards start at the tops of the circles and one in the middle between them. Coach at the blue line with pucks. Net and goalie in place.",
         12, 6, "half", "Net, Goalie, Pucks",
         U12_U16, '["offensive_zone","triangle","spacing","support"]', "offensive", "medium", "attack_triangle_structure"),

        # 45. PXI Centering Pass Progression — Passing
        ("PXI Centering Pass Progression", "passing",
         "Winger starts with a puck in the corner and skates up the wall a few strides before cutting down behind the net. Center times a route by backing away from the net into soft ice between the dots. As the winger comes around the far post, they deliver a centering pass to the center, who catches and shoots quickly. Progression: add a defender with passive stick pressure in the slot, then active pressure to force reads on timing and lane selection.",
         "Winger keeps feet moving and eyes up as they come around the net. Center shows a clear target with stick on the ice and body open to the puck. Pass is delivered through a lane, not through the defender stick. Encourage quick catch-and-release shooting from between the dots.",
         "One line of wingers in the corner with pucks, one line of centers in the low slot, and a net with goalie or target.",
         10, 4, "quarter", "Net, Pucks, Cones to mark starting spots",
         U12_U16, '["passing","net_drive","youth","slot_play"]', "passing", "medium", "centering_pass_timing"),

        # ── Batch 6: Rush, 2v2, Delay, Transition, Defensive, Battle, SAG, Shooting, Goalie ──

        # 46. PXI Three-Lane Kickout Rush — Offensive
        ("PXI Three-Lane Kickout Rush", "offensive",
         "Middle lane forward starts with a puck and skates up ice. As they cross their blue line, they pass to either wide lane and immediately kick out to the opposite wide lane, becoming the middle drive. The last receiver attacks wide and can either shoot, hit the middle driver, or delay for the weak-side lane. After the rush, players rotate lanes so everyone works all three positions.",
         "Middle forward drives through the neutral-zone middle with speed before kicking wide. Wide forwards stay in their lanes and time their routes to support the puck in stride. Attack the dot line before making plays to force defenders inside. All three players finish at the net for rebounds and second chances.",
         "Three lines at one end: left wall, middle, right wall. Coach at far blue line with pucks. One net and goalie at far end.",
         12, 9, "full", "Net, Goalie, Pucks",
         U16_UP, '["rush","three_lane","kickout","timing"]', "offensive", "high", "three_lane_kickout"),

        # 47. PXI Wide Dot Drive 2v2 — Offensive
        ("PXI Wide Dot Drive 2v2", "offensive",
         "On the whistle, two forwards at center receive a puck from the coach and attack 2v2 against the defencemen at the far blue line. Puck carrier must drive wide through the outside dot lane while the second forward fills the inside lane. Defencemen work to keep tight gap and angle toward the boards. Play out the 2v2 to completion and then send the next group the other way.",
         "Forwards maintain spacing: one outside dots, one between dots. Puck carrier keeps feet moving and threatens the net before passing. Defenders match speed early and keep stick in the passing lane. Teach defenders to surf forward through the neutral zone rather than backing in early.",
         "Two F lines at center on each wall, two D at each blue line. Nets and goalies at both ends.",
         12, 8, "full", "Two nets, Goalies, Pucks",
         U16_UP, '["rush","2v2","gap_control","angling"]', "offensive", "high", "wide_dot_drive_2v2"),

        # 48. PXI Delay Cut Dot Attack — Offensive
        ("PXI Delay Cut Dot Attack", "offensive",
         "Forward receives a pass from the coach at center and attacks the defender 1v1. As they reach the top of the circle, they execute a delay cut back toward the boards while protecting the puck, then cut inside toward the dot line for a shot. Defender reads the delay, keeps inside position, and attempts to steer the attacker away from the dangerous middle.",
         "Attacker sells speed first, then uses a sharp cutback with body between puck and defender. Stick and hands stay in front, not behind the body, during the delay. Defender maintains good gap and does not chase behind the attacker on the cutback. Encourage quick release shots off the inside edge after the delay.",
         "Two F lines at center, one on each side. One D line at the far blue line between dots. Net and goalie at far end.",
         10, 6, "full", "Net, Goalie, Pucks",
         U14_UP, '["1v1","delay","zone_entry","attacking_middle"]', "offensive", "medium", "delay_cut_dot_attack"),

        # 49. PXI Odd-Man Quick-Up Game — Transition
        ("PXI Odd-Man Quick-Up Game", "transition",
         "Coach shoots or rims a puck on net. D recover the puck and must make a quick-up pass to one of the three forwards outside the blue line. As soon as the puck exits the zone, those three forwards attack back in as a 3v2. If the defending D or backchecking forwards regain the puck, they quickly transition back up to the coach to reset.",
         "Defencemen shoulder-check and move the puck quickly to the first available outlet. Forwards present clear targets on the walls and in the middle. Quick transitions reward teams that support the puck and move their feet. Backcheckers sprint inside dots to eliminate middle ice options.",
         "Half-ice with a net and goalie. Coach at the blue line with pucks. Two D and three F inside the zone vs three F outside the blue line.",
         15, 10, "half", "Net, Goalie, Pucks",
         U14_UP, '["transition","quick_up","3v2","backcheck"]', "transition", "high", "odd_man_quick_up"),

        # 50. PXI Blue-Line Surf Gap Drill — Defensive
        ("PXI Blue-Line Surf Gap Drill", "defensive",
         "On the whistle, forwards at the far blue receive a pass from the coach and begin skating up ice. Defencemen at center skate forward to gather ice and then surf laterally toward the puck side while maintaining gap. When the forwards reach the red line, coach calls go and D pivot to backward, matching speed and keeping one-and-a-half stick lengths gap into the D-zone. Finish with a live 2v2 or 1v1 depending on the rep.",
         "Defenders skate forward early to close the gap before pivoting. Outside shoulder lines up with the attacker inside shoulder to angle to the boards. Sticks stay on the ice and in lanes; avoid big crossovers that open hips too soon. Forwards challenge the gap by building speed and attacking middle ice.",
         "Two D at the red line; two F lines at far blue line on each wall. Coach with pucks at far blue.",
         12, 6, "full", "Two nets, Pucks, Cones to mark surf start point",
         U16_UP, '["gap_control","surfing","angling","rush_defence"]', "defensive", "medium", "blue_line_surf_gap"),

        # 51. PXI Neutral-Zone Bump Back — Transition
        ("PXI Neutral-Zone Bump Back", "transition",
         "D starts behind the net and hits the strong-side winger on the wall. Winger skates up ice and bumps the puck back to the center cutting underneath through the middle. Center then either carries wide or hits the weak-side winger stretching at the far blue line. Attack continues into a 3v2 against two backtracking D from the opposite blue line.",
         "First pass is hard and flat; winger receives on the move and shields the puck. Center times their cut underneath to arrive as the bump option, not early. Weak-side winger stretches to open ice and stays onside for the long pass. Defenders track back inside dots and match speed through the neutral zone.",
         "Two D at one end with pucks, three F at the near blue line spread across the width. Net and goalie at far end.",
         10, 8, "full", "Two nets, Goalies, Pucks",
         U16_UP, '["transition","bump_pass","support","3v2"]', "transition", "medium", "nz_bump_back_support"),

        # 52. PXI Corner Bump to Slot — Offensive
        ("PXI Corner Bump to Slot", "offensive",
         "F1 protects the puck in the corner under pressure from the defender. F2 stays net-front and F3 hovers in the high slot. F1 can bump the puck up the wall to F3 or behind the net to F2. On a bump, the receiving forward quickly moves the puck to the third player for a shot while the others drive the net. Rotate roles after each short rep.",
         "Puck carrier uses body and edges to hold inside ice in the corner. Bump passes are short and on the forehand when possible. Slot player finds soft ice and stays off defenders sticks. Defender tracks the most dangerous threat but keeps eyes on the puck.",
         "F1 in corner with pucks, F2 at net front, F3 in high slot, one defender and a goalie.",
         10, 5, "quarter", "Net, Goalie, Pucks",
         U14_UP, '["cycling","support","offensive_zone","2v1_low"]', "offensive", "medium", "corner_bump_slot"),

        # 53. PXI Weak-Side Slash Support — Offensive
        ("PXI Weak-Side Slash Support", "offensive",
         "Coach passes to the strong-side winger who carries up ice. Weak-side winger skates a slash route from their wall across the neutral zone, aiming to arrive behind the puck carrier as a middle support option. Defenders work to maintain gap and steer play wide. Puck carrier reads: hit the slasher in the middle, carry wide and delay, or chip in and chase. Play continues as a 2v2.",
         "Weak-side forward times the slash so they are available as the puck crosses the red line. Puck carrier keeps head up and reads defenders sticks before choosing an option. Defenders stay connected and avoid getting split by the slash route. Encourage slash passes through available seams rather than forcing stretch plays.",
         "One F line on each wall at center. Coach with pucks at center dot. Defender pair at far blue line.",
         12, 6, "full", "Two nets, Goalies, Pucks",
         U16_UP, '["support","2v2","neutral_zone","timing"]', "offensive", "medium", "weak_side_slash_support"),

        # 54. PXI Half-Wall Escape Reads — Puck Handling
        ("PXI Half-Wall Escape Reads", "puck_handling",
         "F1 starts with their back to the wall under light pressure. On the whistle they choose one of three escape options: tight turn up-ice and pass to F2, cut back toward the corner and chip behind for a self-pass, or roll to the middle and attack the net. F2 reads and adjusts support, always staying in a passing lane. Progress from passive to full-contact pressure.",
         "Puck carrier keeps knees bent and uses body to separate from pressure. Eyes scan middle of the ice before committing to an escape move. Middle support matches the puck carrier route and stays inside dots. Defender practices good stick-on-puck and body position, not fishing.",
         "F1 on the half wall with pucks, D or pressure player inside dots, F2 as middle support between dots, net and goalie.",
         10, 4, "quarter", "Net, Goalie, Pucks",
         U14_UP, '["puck_protection","zone_exit","support","angling"]', "puck_handling", "medium", "half_wall_escape_reads"),

        # 55. PXI 2v2 Corner Gate Game — Small Area Games
        ("PXI 2v2 Corner Gate Game", "small_area_games",
         "Play 2v2 in the corner with a scoring rule: goals only count if the puck is carried or passed through one of the gates before being shot. This forces attackers to move the puck high-to-low and use space away from the boards. After a goal or clear, coach rims a new puck to keep the game going.",
         "Attackers use gates to pull defenders away from the boards and open seams. Quick give-and-go plays help break coverage and change sides. Defenders protect the middle and communicate who pressures vs who supports. High pace and short shifts keep decision-making sharp.",
         "Quarter-ice with two cone gates at the top of the circle. Two teams of two players each and a goalie.",
         12, 8, "quarter", "Net, Goalie, Pucks, Cones for gates",
         U14_UP, '["small_area_game","2v2","decision_making","support"]', "battle", "high", "2v2_corner_gate"),

        # 56. PXI Bubble Circle Possession — Small Area Games
        ("PXI Bubble Circle Possession", "small_area_games",
         "Play continuous 3v3 possession inside the circle. Coach tosses in a puck to start. Teams score by completing a set number of passes (e.g., five in a row) or by passing to a teammate standing briefly in a marked scoring box at the top of the circle. Players must move into space quickly while staying within the circle boundary.",
         "Players constantly adjust position to stay an easy passing option. Quick passes and give-and-gos help maintain possession under pressure. Use fakes and shoulder checks to escape pressure, not just speed. Keep sticks on the ice and communicate every pass with a call.",
         "Use one zone faceoff circle as the play area. Two teams of 3; no one may leave the circle.",
         10, 6, "quarter", "Pucks, Cones or markers for scoring boxes",
         U12_U16, '["small_area_game","3v3","support","puck_protection"]', "battle", "high", "bubble_circle_possession"),

        # 57. PXI 3v3 Chase Backcheck — Small Area Games
        ("PXI 3v3 Chase Backcheck", "small_area_games",
         "On the whistle, one team receives a puck and skates around both nets before attacking 3v3 in one zone. The chasing team skates around only the far net, entering slightly behind to simulate backchecking pressure. Play 3v3 until a goal or whistle, then start the next rep the other way with roles reversed.",
         "Puck team moves the puck early to beat backcheck pressure. Backcheckers take good angles through the middle and pick up sticks. Defenders communicate on switches and net-front coverage. High tempo skating around the nets builds conditioning and pace.",
         "Two teams line up three-abreast at the red line. Coach at center with pucks. Nets and goalies at both ends.",
         15, 12, "full", "Two nets, Goalies, Pucks",
         U16_UP, '["small_area_game","backcheck","transition","3v3"]', "battle", "high", "3v3_chase_backcheck"),

        # 58. PXI 1v1 Angling Chase — Defensive
        ("PXI 1v1 Angling Chase", "defensive",
         "Coach rims a puck into the far corner. First player on one side becomes the attacker and races for the puck. First player on the opposite side chases from behind and must angle the attacker into the boards and toward the lane boundary before reaching the net. Play out the 1v1 to a shot or turnover, then next pair goes.",
         "Defender skates an arc route, not straight behind, to close space and angle. Stick stays on the ice and through the attacker hands. Attacker protects the puck by keeping body between defender and puck. Finish checks legally through the chest and hands, not with reaching.",
         "Two lines at center dot facing each other on opposite sides. Cones create a narrow lane from blue line to net. Net and goalie in place.",
         10, 6, "half", "Net, Goalie, Pucks, Cones to define lane",
         U14_UP, '["angling","1v1","defensive","compete"]', "defensive", "high", "1v1_angling_chase"),

        # 59. PXI 2v2 Low Net Flip — Battle Drills
        ("PXI 2v2 Low Net Flip", "battle",
         "Coach flips a puck off the back of the net or into the corner. Two attackers and two defenders battle for possession and must stay below the hash marks. Attackers try to create a quick shot from around the net or a pass into the slot. Defenders focus on body position and sticks on ice. After 20-25 seconds or a goal, coach blows the whistle and the next 2v2 group jumps in.",
         "Use the back of the net as a pick to create separation. Defenders keep inside position and avoid chasing behind the net unnecessarily. Quick puck movement and cutbacks are key to breaking tight coverage. Short, intense shifts build compete level and conditioning.",
         "2v2 below the hash marks with a net and goalie. Coach with pucks behind the net.",
         10, 8, "quarter", "Net, Goalie, Pucks",
         U16_UP, '["battle","2v2","net_play","compete"]', "battle", "high", "2v2_low_net_flip"),

        # 60. PXI Rapid Fire Slot Exchange — Shooting
        ("PXI Rapid Fire Slot Exchange", "shooting",
         "Coach passes rapidly to each shooter in sequence. After every shot, the shooter must skate to a new point in the triangle, exchanging spots with a teammate. Coach keeps the pace high with little delay between passes. After 30-40 seconds, rotate in a new group.",
         "Shooters prepare early with sticks loaded and bodies facing the puck. Feet move into the shot; no standing still in the slot. Emphasize quick release over power. Goalie tracks laterally and recovers quickly between shots.",
         "Three shooters in a triangle in the slot, coach with pucks at the top. Goalie in net.",
         8, 5, "quarter", "Net, Goalie, Pucks",
         U14_UP, '["shooting","quick_release","slot","conditioning"]', "shooting", "high", "rapid_fire_slot_exchange"),

        # 61. PXI Goalie Screen Find-and-Track — Goalie
        ("PXI Goalie Screen Find-and-Track", "goalie",
         "Screeners take away the goalie eyes while the shooter moves laterally along the blue line. On the whistle the shooter fires a low shot through traffic. Goalie must fight to find the puck by adjusting depth and lateral position before tracking and making the save. Rotate screeners and shooters frequently.",
         "Goalie moves head first to find sight lines around screens. Depth adjustments are small and controlled; avoid big lunges. Screens are realistic but safe — no contact with the goalie. Shooter aims for pads and sticks to create realistic rebounds.",
         "One point shooter with pucks at blue line, two screeners near the top of the crease, goalie in net.",
         8, 4, "quarter", "Net, Goalie gear, Pucks",
         GOALIE_14UP, '["goalie","screens","tracking","rebound_control"]', "goalie", "medium", "goalie_screen_track"),
    ]

    for d in drills:
        try:
            conn.execute("""
                INSERT INTO drills (id, org_id, name, category, description, coaching_points, setup,
                    duration_minutes, players_needed, ice_surface, equipment, age_levels, tags,
                    skill_focus, intensity, concept_id)
                VALUES (?, NULL, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
            """, (str(uuid.uuid4()), *d))
        except Exception:
            pass  # Skip if already exists
    conn.commit()
    conn.close()
    logger.info("Seeded %d PXI drills", len(drills))


def generate_missing_diagrams():
    """Generate SVG rink diagrams for any drills that don't have one yet."""
    conn = get_db()
    rows = conn.execute(
        "SELECT id, ice_surface, category, concept_id, description FROM drills WHERE diagram_url IS NULL"
    ).fetchall()
    if not rows:
        conn.close()
        return
    count = 0
    for row in rows:
        drill_id = row[0]
        ice_surface = row[1] or "full"
        category = row[2] or "offensive"
        concept_id = row[3]
        description = row[4] or ""
        try:
            svg_content = generate_drill_diagram(ice_surface, category, concept_id, description)
            svg_filename = f"drill_{drill_id}.svg"
            svg_path = os.path.join(_IMAGES_DIR, svg_filename)
            with open(svg_path, "w", encoding="utf-8") as f:
                f.write(svg_content)
            diagram_url = f"/uploads/{svg_filename}"
            conn.execute("UPDATE drills SET diagram_url = ? WHERE id = ?", (diagram_url, drill_id))
            count += 1
        except Exception as e:
            logger.error("Failed to generate diagram for drill %s: %s", drill_id, e)
    conn.commit()
    conn.close()
    if count:
        logger.info("Generated %d drill diagrams", count)


# Run on import
init_db()
seed_templates()
seed_new_templates()
_migrate_template_prompts()
seed_hockey_os()
seed_leagues()
migrate_leagues()
seed_teams()
seed_drills()
seed_drills_v2()
seed_drills_pxi()
generate_missing_diagrams()
seed_glossary_v2()


# ── Auto-backup thread ──────────────────────────────────────────────
def _auto_backup_loop():
    """Background thread: backup database every 6 hours."""
    while True:
        time.sleep(6 * 3600)  # 6 hours
        try:
            backup_dir = os.path.join(os.path.dirname(DB_FILE), "backups")
            os.makedirs(backup_dir, exist_ok=True)
            ts = datetime.now(timezone.utc).isoformat()[:19].replace(":", "-")
            backup_name = f"prospectx_auto_{ts}.db"
            shutil.copy2(DB_FILE, os.path.join(backup_dir, backup_name))
            # Keep last 20 backups total
            all_backups = sorted(glob.glob(os.path.join(backup_dir, "prospectx_*.db")))
            for old in all_backups[:-20]:
                os.remove(old)
            logger.info("Auto backup created: %s", backup_name)
        except Exception as e:
            logger.warning("Auto backup failed: %s", e)

_backup_thread = threading.Thread(target=_auto_backup_loop, daemon=True)
_backup_thread.start()
logger.info("Auto-backup thread started (every 6 hours)")


# ============================================================
# INSTAT ANALYTICS — HEADER MAPPINGS
# ============================================================
# Maps normalized InStat XLSX headers → storage targets.
# _parse_file_to_rows normalizes: .lower().replace(" ", "_").replace("-", "_")
# So "+/-" → "+/_", "short-handed" → "short_handed", "Faceoffs won, %" → "faceoffs_won,_%"
#
# Core targets: direct column names in player_stats (gp, g, a, p, etc.)
# Extended targets: "category.field" paths stored in extended_stats JSON

INSTAT_SKATER_CORE_MAP = {
    "games_played": "gp",
    "goals": "g",
    "assists": "a",
    "points": "p",
    "+/_": "plus_minus",
    "penalty_time": "pim",           # MM:SS → minutes
    "time_on_ice": "toi_seconds",    # MM:SS → seconds
    "shots": "shots",
    "shots_on_goal": "sog",
    "%_shots_on_goal": "shooting_pct",
}

INSTAT_SKATER_EXTENDED_MAP = {
    # Main statistics
    "all_shifts": "main.shifts",
    "puck_touches": "main.puck_touches",
    "puck_control_time": "main.puck_control_time",
    "scoring_chances": "main.scoring_chances",
    "penalties": "main.penalties",
    "penalties_drawn": "main.penalties_drawn",
    "hits": "main.hits",
    "hits_against": "main.hits_against",
    "error_leading_to_goal": "main.error_leading_to_goal",
    "dump_ins": "main.dump_ins",
    "dump_outs": "main.dump_outs",
    "first_assist": "main.first_assist",
    "second_assist": "main.second_assist",
    "plus": "main.plus",
    "minus": "main.minus",
    "faceoffs": "main.faceoffs",
    "faceoffs_won": "main.faceoffs_won",
    "faceoffs_lost": "main.faceoffs_lost",
    "faceoffs_won,_%": "main.faceoffs_won_pct",
    # Shots
    "blocked_shots": "shots.blocked_shots",
    "missed_shots": "shots.missed_shots",
    "slapshot": "shots.slapshot",
    "wrist_shot": "shots.wrist_shot",
    "shootouts": "shots.shootouts",
    "shootouts_scored": "shots.shootouts_scored",
    "shootouts_missed": "shots.shootouts_missed",
    "power_play_shots": "shots.pp_shots",
    "short_handed_shots": "shots.sh_shots",
    "positional_attack_shots": "shots.positional_attack_shots",
    "counter_attack_shots": "shots.counter_attack_shots",
    "shots_5_v_5": "shots.five_v_five_shots",
    # Puck battles
    "puck_battles": "puck_battles.total",
    "puck_battles_won": "puck_battles.won",
    "puck_battles_won,_%": "puck_battles.won_pct",
    "puck_battles_in_dz": "puck_battles.dz",
    "puck_battles_in_nz": "puck_battles.nz",
    "puck_battles_in_oz": "puck_battles.oz",
    "shots_blocking": "puck_battles.shots_blocking",
    "dekes": "puck_battles.dekes",
    "dekes_successful": "puck_battles.dekes_successful",
    "dekes_unsuccessful": "puck_battles.dekes_unsuccessful",
    "dekes_successful,_%": "puck_battles.dekes_successful_pct",
    # Recoveries & losses
    "takeaways": "recoveries.takeaways",
    "takeaways_in_dz": "recoveries.takeaways_dz",
    "takeaways_in_nz": "recoveries.takeaways_nz",
    "takeaways_in_oz": "recoveries.takeaways_oz",
    "loose_puck_recovery": "recoveries.loose_puck_recovery",
    "opponent's_dump_in_retrievals": "recoveries.dump_in_retrievals",
    "puck_retrievals_after_shots": "recoveries.puck_retrievals_after_shots",
    "puck_losses": "recoveries.puck_losses",
    "puck_losses_in_dz": "recoveries.puck_losses_dz",
    "puck_losses_in_nz": "recoveries.puck_losses_nz",
    "puck_losses_in_oz": "recoveries.puck_losses_oz",
    # Special teams
    "power_play": "special_teams.pp_count",
    "successful_power_play": "special_teams.pp_successful",
    "power_play_time": "special_teams.pp_time",
    "short_handed": "special_teams.sh_count",
    "penalty_killing": "special_teams.pk_count",
    "short_handed_time": "special_teams.sh_time",
    # xG
    "xg_per_shot": "xg.xg_per_shot",
    "xg_(expected_goals)": "xg.xg",
    "xg_per_goal": "xg.xg_per_goal",
    "net_xg_(xg_player_on___opp._team's_xg)": "xg.net_xg",
    "team_xg_when_on_ice": "xg.team_xg_on_ice",
    "opponent's_xg_when_on_ice": "xg.opponent_xg_on_ice",
    "xg_conversion": "xg.xg_conversion",
    # Passes
    "passes": "passes.total",
    "accurate_passes": "passes.accurate",
    "accurate_passes,_%": "passes.accurate_pct",
    "passes_to_the_slot": "passes.to_slot",
    "pre_shots_passes": "passes.pre_shot",
    "pass_receptions": "passes.receptions",
    # Entries & breakouts
    "entries": "entries.total",
    "entries_via_pass": "entries.via_pass",
    "entries_via_dump_in": "entries.via_dump",
    "entries_via_stickhandling": "entries.via_stickhandling",
    "breakouts": "entries.breakouts_total",
    "breakouts_via_pass": "entries.breakouts_via_pass",
    "breakouts_via_dump_out": "entries.breakouts_via_dump",
    "breakouts_via_stickhandling": "entries.breakouts_via_stickhandling",
    # Advanced (CORSI / Fenwick)
    "corsi": "advanced.corsi",
    "corsi_": "advanced.corsi_against",
    "corsi+": "advanced.corsi_for",
    "corsi_for,_%": "advanced.corsi_pct",
    "fenwick_for": "advanced.fenwick_for",
    "fenwick_against": "advanced.fenwick_against",
    "fenwick_for,_%": "advanced.fenwick_pct",
    # Faceoffs by zone
    "faceoffs_in_dz": "faceoffs_zone.dz_total",
    "faceoffs_won_in_dz": "faceoffs_zone.dz_won",
    "faceoffs_won_in_dz,_%": "faceoffs_zone.dz_pct",
    "faceoffs_in_nz": "faceoffs_zone.nz_total",
    "faceoffs_won_in_nz": "faceoffs_zone.nz_won",
    "faceoffs_won_in_nz,_%": "faceoffs_zone.nz_pct",
    "faceoffs_in_oz": "faceoffs_zone.oz_total",
    "faceoffs_won_in_oz": "faceoffs_zone.oz_won",
    "faceoffs_won_in_oz,_%": "faceoffs_zone.oz_pct",
    # Playtime
    "offensive_play": "playtime.offensive",
    "defensive_play": "playtime.defensive",
    "playing_in_attack": "playtime.offensive",  # team export alias
    "playing_in_defense": "playtime.defensive",  # team export alias
    "oz_possession": "playtime.oz_possession",
    "nz_possession": "playtime.nz_possession",
    "dz_possession": "playtime.dz_possession",
    # Scoring chances detail
    "scoring_chances___total": "scoring_chances.total",
    "scoring_chances___scored": "scoring_chances.scored",
    "scoring_chances_missed": "scoring_chances.missed",
    "scoring_chances_saved": "scoring_chances.saved",
    "scoring_chances,_%": "scoring_chances.pct",
    "inner_slot_shots___total": "scoring_chances.inner_slot_total",
    "inner_slot_shots___scored": "scoring_chances.inner_slot_scored",
    "inner_slot_shots___missed": "scoring_chances.inner_slot_missed",
    "inner_slot_shots___saved": "scoring_chances.inner_slot_saved",
    "inner_slot_shots,_%": "scoring_chances.inner_slot_pct",
    "outer_slot_shots___total": "scoring_chances.outer_slot_total",
    "outer_slot_shots___scored": "scoring_chances.outer_slot_scored",
    "outer_slot_shots___missed": "scoring_chances.outer_slot_missed",
    "outer_slot_shots___saved": "scoring_chances.outer_slot_saved",
    "outer_slot_shots,_%": "scoring_chances.outer_slot_pct",
    "blocked_shots_from_the_slot": "scoring_chances.blocked_from_slot",
    "blocked_shots_outside_of_the_slot": "scoring_chances.blocked_outside_slot",
    # Team-specific extras (only in team exports)
    "team_goals_when_on_ice": "team_extras.team_goals_on_ice",
    "opponent's_goals_when_on_ice": "team_extras.opponent_goals_on_ice",
    "1_on_1_shots": "team_extras.one_on_one_shots",
    "1_on_1_goals": "team_extras.one_on_one_goals",
    "shots_conversion_1_on_1,_%": "team_extras.one_on_one_pct",
    "ev_dz_retrievals": "team_extras.ev_dz_retrievals",
    "ev_oz_retrievals": "team_extras.ev_oz_retrievals",
    "power_play_retrievals": "team_extras.pp_retrievals",
    "penalty_kill_retrievals": "team_extras.pk_retrievals",
}

INSTAT_SKATER_BIO_MAP = {
    "date_of_birth": "dob",
    "nationality": "nationality",
    "height": "height",
    "weight": "weight",
    "active_hand": "shoots",
    "contract": "contract",
    "national_team": "national_team",
}

INSTAT_GOALIE_CORE_MAP = {
    "games_played": "gp",
    "time_on_ice": "toi_seconds",
    "goals_against": "ga",
    "shots_on_goal": "sa",
    "saves": "sv",
    "saves,_%": "sv_pct",
    "penalty_time": "pim",
}

INSTAT_GOALIE_EXTENDED_MAP = {
    "penalties_drawn": "penalties_drawn",
    "passes": "passes_total",
    "accurate_passes": "passes_accurate",
    "accurate_passes,_%": "passes_accurate_pct",
    "xg_conceded": "xg_conceded",
    "xg_per_shot_taken": "xg_per_shot",
    "xg_per_goal_conceded": "xg_per_goal",
    "xg_per_shot_saved": "xg_per_shot_saved",
    "shootouts": "shootouts",
    "shootout_saves": "shootout_saves",
    "shootouts_allowed": "shootouts_allowed",
    "scoring_chances___total": "scoring_chances_total",
    "scoring_chance_saves": "scoring_chance_saves",
    "scoring_chance_saves,_%": "scoring_chance_saves_pct",
    "age": "age",
}

INSTAT_TEAM_EXTENDED_MAP = {
    # Maps ALL team stat columns to extended_stats categories
    # Team stats don't have core columns — everything goes to extended_stats
    "goals": "offense.goals",
    "penalties": "discipline.penalties",
    "penalties_drawn": "discipline.penalties_drawn",
    "penalty_time": "discipline.penalty_time",
    "faceoffs": "faceoffs.total",
    "faceoffs_won": "faceoffs.won",
    "faceoffs_won,_%": "faceoffs.won_pct",
    "faceoffs_lost": "faceoffs.lost",
    "hits": "physical.hits",
    "hits_against": "physical.hits_against",
    "faceoffs_in_dz": "faceoffs.dz_total",
    "faceoffs_won_in_dz": "faceoffs.dz_won",
    "faceoffs_won_in_dz,_%": "faceoffs.dz_pct",
    "faceoffs_in_nz": "faceoffs.nz_total",
    "faceoffs_won_in_nz": "faceoffs.nz_won",
    "faceoffs_won_in_nz,_%": "faceoffs.nz_pct",
    "faceoffs_in_oz": "faceoffs.oz_total",
    "faceoffs_won_in_oz": "faceoffs.oz_won",
    "faceoffs_won_in_oz,_%": "faceoffs.oz_pct",
    "scoring_chances": "offense.scoring_chances",
    "shots": "shots.total",
    "shots_on_goal": "shots.on_goal",
    "blocked_shots": "shots.blocked",
    "missed_shots": "shots.missed",
    "%_shots_on_goal": "shots.on_goal_pct",
    "slapshot": "shots.slapshot",
    "wrist_shot": "shots.wrist_shot",
    "power_play_shots": "shots.pp_shots",
    "short_handed_shots": "shots.sh_shots",
    "shootouts_scored": "shots.shootout_scored",
    "corsi%": "advanced.corsi_pct",
    "power_play": "special_teams.pp_count",
    "successful_power_play": "special_teams.pp_successful",
    "power_play_time": "special_teams.pp_time",
    "power_play,_%": "special_teams.pp_pct",
    "short_handed": "special_teams.sh_count",
    "penalty_killing": "special_teams.pk_count",
    "short_handed_time": "special_teams.sh_time",
    "short_handed,_%": "special_teams.pk_pct",
    "shots_blocking": "defense.shots_blocking",
    "xg_per_shot": "xg.xg_per_shot",
    "opponent's_xg_per_shot": "xg.opponent_xg_per_shot",
    "net_xg_(xg___opponent's_xg)": "xg.net_xg",
    "xg_conversion": "xg.xg_conversion",
    "xg_(expected_goals)": "xg.xg",
    "opponent's_xg": "xg.opponent_xg",
    "xg_per_goal": "xg.xg_per_goal",
    "opponent's_xg_per_goal": "xg.opponent_xg_per_goal",
    "offensive_play": "playtime.offensive",
    "defensive_play": "playtime.defensive",
    "oz_possession": "playtime.oz_possession",
    "nz_possession": "playtime.nz_possession",
    "dz_possession": "playtime.dz_possession",
    "puck_battles": "puck_battles.total",
    "puck_battles_won": "puck_battles.won",
    "puck_battles_won,_%": "puck_battles.won_pct",
    "puck_battles_in_oz": "puck_battles.oz",
    "puck_battles_in_nz": "puck_battles.nz",
    "puck_battles_in_dz": "puck_battles.dz",
    "dekes": "puck_battles.dekes",
    "dekes_successful": "puck_battles.dekes_successful",
    "dekes_unsuccessful": "puck_battles.dekes_unsuccessful",
    "dekes_successful,_%": "puck_battles.dekes_successful_pct",
    "passes_total": "passes.total",
    "accurate_passes": "passes.accurate",
    "accurate_passes,_%": "passes.accurate_pct",
    "pre_shots_passes": "passes.pre_shot",
    "dump_ins": "transition.dump_ins",
    "dump_outs": "transition.dump_outs",
    "passes_to_the_slot": "passes.to_slot",
    "takeaways": "recoveries.takeaways",
    "takeaways_in_nz": "recoveries.takeaways_nz",
    "takeaways_in_dz": "recoveries.takeaways_dz",
    "takeaways_in_oz": "recoveries.takeaways_oz",
    "loose_puck_recovery": "recoveries.loose_puck_recovery",
    "opponent's_dump_in_retrievals": "recoveries.dump_in_retrievals",
    "puck_losses": "recoveries.puck_losses",
    "puck_losses_in_oz": "recoveries.puck_losses_oz",
    "puck_losses_in_nz": "recoveries.puck_losses_nz",
    "puck_losses_in_dz": "recoveries.puck_losses_dz",
    "retrievals": "recoveries.retrievals",
    "power_play_retrievals": "recoveries.pp_retrievals",
    "penalty_kill_retrievals": "recoveries.pk_retrievals",
    "ev_oz_retrievals": "recoveries.ev_oz_retrievals",
    "ev_dz_retrievals": "recoveries.ev_dz_retrievals",
    "entries": "entries.total",
    "entries_via_pass": "entries.via_pass",
    "entries_via_dump_in": "entries.via_dump",
    "entries_via_stickhandling": "entries.via_stickhandling",
    "breakouts": "entries.breakouts_total",
    "breakouts_via_pass": "entries.breakouts_via_pass",
    "breakouts_via_dump_out": "entries.breakouts_via_dump",
    "breakouts_via_stickhandling": "entries.breakouts_via_stickhandling",
    "oz_play": "offense.oz_play",
    "oz_play_with_shots": "offense.oz_play_with_shots",
    "oz_play_with_shots,_%": "offense.oz_play_with_shots_pct",
    "counterattacks": "offense.counterattacks",
    "counter_attack_with_shots": "offense.counterattack_with_shots",
    "counter_attack_with_shots,_%": "offense.counterattack_with_shots_pct",
}

INSTAT_LINES_MAP = {
    "plus/minus": "plus_minus",
    "numbers_of_shifts": "shifts",
    "time_on_ice": "toi_seconds",
    "goals": "goals_for",
    "opponent's_goals": "goals_against",
    "shots": "shots_for",
    "shots_on_goal": "sog_for",
    "opponent_shots_total": "shots_against",
    "shots_on_goal_against": "sog_against",
    "corsi": "corsi",
    "corsi+": "corsi_for",
    "corsi_": "corsi_against",
    "short_handed_play": "sh_play",
    "power_play_played": "pp_played",
    "power_play_time": "pp_time",
    "successful_power_play": "pp_successful",
    "short_handed_time": "sh_time",
}


# ============================================================
# HELPER
# ============================================================

def row_to_dict(row: sqlite3.Row) -> dict:
    """Convert a sqlite3.Row to a plain dict."""
    return dict(row)


def gen_id() -> str:
    return str(uuid.uuid4())


def now_iso() -> str:
    return datetime.now(timezone.utc).isoformat()


def get_anthropic_client():
    from anthropic import Anthropic
    key = ANTHROPIC_API_KEY
    if not key or key == "your_anthropic_api_key_here":
        return None
    return Anthropic(api_key=key)


# ============================================================
# SUBSCRIPTION & USAGE HELPERS
# ============================================================

def _check_bench_talk_limit(user_id: str, conn) -> dict:
    """DEPRECATED: Use _check_tier_limit(user_id, 'bench_talks', conn) instead.
    Check if user has remaining Bench Talk quota. Raises 429 if exceeded."""
    row = conn.execute(
        "SELECT subscription_tier, monthly_bench_talks_used, usage_reset_at FROM users WHERE id = ?",
        (user_id,),
    ).fetchone()
    if not row:
        raise HTTPException(status_code=404, detail="User not found")

    tier = row["subscription_tier"] or "rookie"
    tier_config = SUBSCRIPTION_TIERS.get(tier, SUBSCRIPTION_TIERS["rookie"])
    limit = tier_config["monthly_bench_talks"]
    used = row["monthly_bench_talks_used"] or 0

    # Check if monthly reset is needed (reset on 1st of month)
    reset_at = row["usage_reset_at"]
    now = datetime.now(timezone.utc)
    month_start = now.replace(day=1, hour=0, minute=0, second=0, microsecond=0)
    if not reset_at or datetime.fromisoformat(reset_at) < month_start:
        conn.execute(
            "UPDATE users SET monthly_bench_talks_used = 0, monthly_reports_used = 0, usage_reset_at = ? WHERE id = ?",
            (now.isoformat(), user_id),
        )
        conn.commit()
        used = 0

    if limit != -1 and used >= limit:
        raise HTTPException(status_code=429, detail={
            "error": "bench_talk_limit_reached",
            "tier": tier,
            "limit": limit,
            "used": used,
            "upgrade_url": "/pricing",
        })

    return {"tier": tier, "limit": limit, "used": used}


def _check_report_limit(user_id: str, conn) -> dict:
    """DEPRECATED: Use _check_tier_limit(user_id, 'reports', conn) instead.
    Check if user has remaining report generation quota. Raises 429 if exceeded."""
    row = conn.execute(
        "SELECT subscription_tier, monthly_reports_used, usage_reset_at FROM users WHERE id = ?",
        (user_id,),
    ).fetchone()
    if not row:
        raise HTTPException(status_code=404, detail="User not found")

    tier = row["subscription_tier"] or "rookie"
    tier_config = SUBSCRIPTION_TIERS.get(tier, SUBSCRIPTION_TIERS["rookie"])
    limit = tier_config["monthly_reports"]
    used = row["monthly_reports_used"] or 0

    # Monthly reset check
    reset_at = row["usage_reset_at"]
    now = datetime.now(timezone.utc)
    month_start = now.replace(day=1, hour=0, minute=0, second=0, microsecond=0)
    if not reset_at or datetime.fromisoformat(reset_at) < month_start:
        conn.execute(
            "UPDATE users SET monthly_bench_talks_used = 0, monthly_reports_used = 0, usage_reset_at = ? WHERE id = ?",
            (now.isoformat(), user_id),
        )
        conn.commit()
        used = 0

    if limit != -1 and used >= limit:
        raise HTTPException(status_code=429, detail={
            "error": "report_limit_reached",
            "tier": tier,
            "limit": limit,
            "used": used,
            "upgrade_url": "/pricing",
        })

    return {"tier": tier, "limit": limit, "used": used}


def _increment_usage(user_id: str, action_type: str, resource_id: str, org_id: str, conn):
    """Increment usage counter, log it, and update new usage_tracking table."""
    # Legacy: update users table columns
    col = "monthly_bench_talks_used" if action_type == "bench_talk" else "monthly_reports_used"
    conn.execute(f"UPDATE users SET {col} = COALESCE({col}, 0) + 1 WHERE id = ?", (user_id,))
    # Legacy: log to subscription_usage_log
    conn.execute(
        "INSERT INTO subscription_usage_log (id, user_id, org_id, action_type, resource_id) VALUES (?, ?, ?, ?, ?)",
        (gen_id(), user_id, org_id, action_type, resource_id),
    )
    conn.commit()
    # New: update usage_tracking table (maps action_type to resource_type)
    resource_map = {"report": "reports", "bench_talk": "bench_talks"}
    resource_type = resource_map.get(action_type)
    if resource_type:
        _increment_tracking(user_id, resource_type, conn)


# ============================================================
# PYDANTIC MODELS
# ============================================================

# --- Auth ---
class LoginRequest(BaseModel):
    email: str
    password: str

class RegisterRequest(BaseModel):
    email: str
    password: str = Field(..., min_length=8)
    first_name: str
    last_name: str
    org_name: str
    org_type: str = "team"
    hockey_role: str = "scout"  # scout, gm, coach, player, parent, broadcaster, producer, agent

class UserOut(BaseModel):
    id: str
    org_id: str
    email: str
    first_name: Optional[str] = None
    last_name: Optional[str] = None
    role: str
    hockey_role: str = "scout"
    subscription_tier: str = "rookie"
    monthly_reports_used: int = 0
    monthly_bench_talks_used: int = 0
    email_verified: bool = False

class TokenResponse(BaseModel):
    access_token: str
    token_type: str = "bearer"
    user: UserOut

class SubscriptionUpgradeRequest(BaseModel):
    tier: str

class AdminTierUpdateRequest(BaseModel):
    tier: str

# --- Players ---
class PlayerCreate(BaseModel):
    first_name: str
    last_name: str
    position: str
    dob: Optional[str] = None
    shoots: Optional[str] = None
    height_cm: Optional[int] = None
    weight_kg: Optional[int] = None
    current_team: Optional[str] = None
    current_league: Optional[str] = None
    passports: Optional[List[str]] = []
    notes: Optional[str] = None
    tags: Optional[List[str]] = []
    archetype: Optional[str] = None
    image_url: Optional[str] = None
    commitment_status: Optional[str] = "Uncommitted"
    elite_prospects_url: Optional[str] = None

class PlayerResponse(BaseModel):
    id: str
    org_id: str
    first_name: str
    last_name: str
    dob: Optional[str] = None
    position: str
    shoots: Optional[str] = None
    height_cm: Optional[int] = None
    weight_kg: Optional[int] = None
    current_team: Optional[str] = None
    current_league: Optional[str] = None
    passports: Optional[List[str]] = []
    notes: Optional[str] = None
    tags: Optional[List[str]] = []
    archetype: Optional[str] = None
    image_url: Optional[str] = None
    birth_year: Optional[int] = None
    age_group: Optional[str] = None
    draft_eligible_year: Optional[int] = None
    league_tier: Optional[str] = None
    commitment_status: Optional[str] = "Uncommitted"
    elite_prospects_url: Optional[str] = None
    roster_status: Optional[str] = "active"
    created_at: str

# --- Reports ---
class ReportGenerateRequest(BaseModel):
    player_id: Optional[str] = None
    team_name: Optional[str] = None
    report_type: str
    template_id: Optional[str] = None
    data_scope: Optional[Dict[str, Any]] = None
    mode: Optional[str] = None  # PXI mode override (scout, coach, analyst, etc.)

class ReportResponse(BaseModel):
    model_config = {"extra": "ignore"}
    id: str
    org_id: str
    player_id: Optional[str] = None
    team_name: Optional[str] = None
    report_type: str
    title: Optional[str] = None
    status: str
    output_json: Optional[Dict[str, Any]] = None
    output_text: Optional[str] = None
    error_message: Optional[str] = None
    generated_at: Optional[str] = None
    llm_model: Optional[str] = None
    llm_tokens: Optional[int] = None
    created_at: str
    share_token: Optional[str] = None
    shared_with_org: Optional[int] = 0
    quality_score: Optional[float] = None
    quality_details: Optional[str] = None

class ReportGenerateResponse(BaseModel):
    report_id: str
    status: str
    title: Optional[str] = None
    generation_time_ms: Optional[int] = None

class ReportStatusResponse(BaseModel):
    report_id: str
    status: str
    error_message: Optional[str] = None
    generation_time_ms: Optional[int] = None

# --- Stats ---
class StatsResponse(BaseModel):
    id: str
    player_id: str
    game_id: Optional[str] = None
    season: Optional[str] = None
    stat_type: str
    gp: int
    g: int
    a: int
    p: int
    plus_minus: int
    pim: int
    toi_seconds: int
    pp_toi_seconds: int
    pk_toi_seconds: int
    shots: int
    sog: int
    shooting_pct: Optional[float] = None
    microstats: Optional[Dict[str, Any]] = None
    created_at: str

# --- Scout Notes ---
COMPETITION_LEVELS = [
    "U13_AAA", "U14_AAA", "U15_AAA", "U16_AAA", "U18_AAA",
    "USHL", "OHL", "WHL", "QMJHL", "BCHL", "NAHL",
    "NCAA_D1", "NCAA_D3", "AHL", "ECHL", "PRO", "OTHER",
]
PROSPECT_STATUSES = ["TOP_TARGET", "A_PROSPECT", "B_PROSPECT", "C_PROSPECT", "FOLLOW_UP", "PASS"]

class NoteCreate(BaseModel):
    player_id: Optional[str] = None  # Can be set via URL param or body
    note_text: str = ""
    note_type: str = "general"
    tags: Optional[List[str]] = []
    is_private: bool = False
    # v2 fields
    game_date: Optional[str] = None
    opponent: Optional[str] = None
    competition_level: Optional[str] = None
    venue: Optional[str] = None
    overall_grade: Optional[int] = Field(None, ge=1, le=80)
    grade_scale: str = "1-5"
    skating_rating: Optional[int] = Field(None, ge=1, le=5)
    puck_skills_rating: Optional[int] = Field(None, ge=1, le=5)
    hockey_iq_rating: Optional[int] = Field(None, ge=1, le=5)
    compete_rating: Optional[int] = Field(None, ge=1, le=5)
    defense_rating: Optional[int] = Field(None, ge=1, le=5)
    strengths_notes: Optional[str] = None
    improvements_notes: Optional[str] = None
    development_notes: Optional[str] = None
    one_line_summary: Optional[str] = None
    prospect_status: Optional[str] = None
    visibility: str = "PRIVATE"
    note_mode: str = "QUICK"

class NoteUpdate(BaseModel):
    note_text: Optional[str] = None
    note_type: Optional[str] = None
    tags: Optional[List[str]] = None
    is_private: Optional[bool] = None
    game_date: Optional[str] = None
    opponent: Optional[str] = None
    competition_level: Optional[str] = None
    venue: Optional[str] = None
    overall_grade: Optional[int] = Field(None, ge=1, le=80)
    grade_scale: Optional[str] = None
    skating_rating: Optional[int] = Field(None, ge=1, le=5)
    puck_skills_rating: Optional[int] = Field(None, ge=1, le=5)
    hockey_iq_rating: Optional[int] = Field(None, ge=1, le=5)
    compete_rating: Optional[int] = Field(None, ge=1, le=5)
    defense_rating: Optional[int] = Field(None, ge=1, le=5)
    strengths_notes: Optional[str] = None
    improvements_notes: Optional[str] = None
    development_notes: Optional[str] = None
    one_line_summary: Optional[str] = None
    prospect_status: Optional[str] = None
    visibility: Optional[str] = None
    note_mode: Optional[str] = None

class NoteResponse(BaseModel):
    id: str
    org_id: str
    player_id: str
    scout_id: str
    scout_name: Optional[str] = None
    note_text: str
    note_type: str
    tags: List[str]
    is_private: bool
    created_at: str
    updated_at: str
    # v2 fields
    game_date: Optional[str] = None
    opponent: Optional[str] = None
    competition_level: Optional[str] = None
    venue: Optional[str] = None
    overall_grade: Optional[int] = None
    grade_scale: str = "1-5"
    skating_rating: Optional[int] = None
    puck_skills_rating: Optional[int] = None
    hockey_iq_rating: Optional[int] = None
    compete_rating: Optional[int] = None
    defense_rating: Optional[int] = None
    strengths_notes: Optional[str] = None
    improvements_notes: Optional[str] = None
    development_notes: Optional[str] = None
    one_line_summary: Optional[str] = None
    prospect_status: Optional[str] = None
    visibility: str = "PRIVATE"
    note_mode: str = "QUICK"
    # Joined display fields
    player_name: Optional[str] = None
    player_team: Optional[str] = None
    player_position: Optional[str] = None
    author_name: Optional[str] = None

# --- Line Combinations ---
class LinePlayerRef(BaseModel):
    player_id: Optional[str] = None
    name: str
    jersey: str = ""
    position: str = ""

class LineCombinationCreate(BaseModel):
    team_name: str
    season: str = ""
    line_type: str  # "forwards", "defense", "pp", "pk"
    line_label: str = ""  # "1st Line", "PP1", etc.
    line_order: int = 0
    player_refs: List[LinePlayerRef] = []

class LineCombinationUpdate(BaseModel):
    line_label: Optional[str] = None
    line_order: Optional[int] = None
    player_refs: Optional[List[LinePlayerRef]] = None

# --- Batch Import ---
class ImportDuplicate(BaseModel):
    row_index: int
    csv_name: str
    existing_id: str
    existing_name: str
    match_score: float
    match_reasons: List[str]

class ImportPreviewResponse(BaseModel):
    job_id: str
    filename: str
    total_rows: int
    new_players: int
    duplicates: List[ImportDuplicate]
    errors: List[str]
    preview: List[dict]

class DuplicateResolution(BaseModel):
    row_index: int
    action: str  # "create_new", "skip", "merge"

class ImportExecuteRequest(BaseModel):
    resolutions: List[DuplicateResolution] = []

# --- Templates ---
class TemplateResponse(BaseModel):
    id: str
    template_name: str
    report_type: str
    description: str = ""
    is_global: bool
    version: int
    created_at: str


# ============================================================
# AUTH UTILITIES
# ============================================================

def create_token(user_id: str, org_id: str, role: str) -> str:
    payload = {
        "user_id": user_id,
        "org_id": org_id,
        "role": role,
        "exp": datetime.now(timezone.utc) + timedelta(hours=JWT_EXPIRY_HOURS),
        "iat": datetime.now(timezone.utc),
    }
    return jwt.encode(payload, JWT_SECRET, algorithm=JWT_ALGORITHM)


def verify_token(credentials: Optional[HTTPAuthorizationCredentials] = Depends(security)) -> dict:
    if not credentials:
        raise HTTPException(status_code=401, detail="Not authenticated")
    try:
        payload = jwt.decode(credentials.credentials, JWT_SECRET, algorithms=[JWT_ALGORITHM])
        return payload
    except jwt.ExpiredSignatureError:
        raise HTTPException(status_code=401, detail="Token expired")
    except jwt.InvalidTokenError:
        raise HTTPException(status_code=401, detail="Invalid token")


# ============================================================
# AUTH ENDPOINTS
# ============================================================

@app.post("/auth/register", response_model=TokenResponse)
async def register(req: RegisterRequest):
    conn = get_db()
    existing = conn.execute("SELECT id FROM users WHERE email = ?", (req.email.lower().strip(),)).fetchone()
    if existing:
        conn.close()
        raise HTTPException(status_code=400, detail="Email already registered")

    org_id = gen_id()
    user_id = gen_id()
    password_hash = pwd_context.hash(req.password[:72])

    conn.execute(
        "INSERT INTO organizations (id, name, org_type) VALUES (?, ?, ?)",
        (org_id, req.org_name, req.org_type),
    )
    hockey_role = req.hockey_role if req.hockey_role in ("scout", "gm", "coach", "player", "parent", "broadcaster", "producer", "agent") else "scout"
    conn.execute(
        "INSERT INTO users (id, org_id, email, password_hash, first_name, last_name, role, hockey_role, email_verified) VALUES (?, ?, ?, ?, ?, ?, ?, ?, 1)",
        (user_id, org_id, req.email.lower().strip(), password_hash, req.first_name, req.last_name, "admin", hockey_role),
    )
    conn.commit()

    # Auto-generate email verification token
    verify_token_raw = secrets.token_urlsafe(32)
    conn.execute(
        "UPDATE users SET email_verify_token = ?, email_verify_sent_at = ? WHERE id = ?",
        (verify_token_raw, datetime.now(timezone.utc).isoformat(), user_id),
    )
    conn.commit()
    logger.info("EMAIL VERIFICATION LINK: %s/verify-email?token=%s (user: %s)", FRONTEND_URL, verify_token_raw, req.email)

    user = UserOut(
        id=user_id, org_id=org_id, email=req.email.lower().strip(),
        first_name=req.first_name, last_name=req.last_name, role="admin",
        hockey_role=hockey_role, subscription_tier="rookie",
        monthly_reports_used=0, monthly_bench_talks_used=0,
        email_verified=True,
    )
    token = create_token(user_id, org_id, "admin")
    conn.close()

    logger.info("User registered: %s (org: %s)", req.email, req.org_name)
    return TokenResponse(access_token=token, user=user)


@app.post("/auth/login", response_model=TokenResponse)
async def login(req: LoginRequest):
    conn = get_db()
    row = conn.execute("SELECT * FROM users WHERE email = ?", (req.email.lower().strip(),)).fetchone()
    conn.close()

    if not row:
        raise HTTPException(status_code=401, detail="Invalid email or password")

    if not pwd_context.verify(req.password[:72], row["password_hash"]):
        raise HTTPException(status_code=401, detail="Invalid email or password")

    user = UserOut(
        id=row["id"], org_id=row["org_id"], email=row["email"],
        first_name=row["first_name"], last_name=row["last_name"], role=row["role"],
        hockey_role=row["hockey_role"] if row["hockey_role"] else "scout",
        subscription_tier=row["subscription_tier"] or "rookie",
        monthly_reports_used=row["monthly_reports_used"] or 0,
        monthly_bench_talks_used=row["monthly_bench_talks_used"] or 0,
        email_verified=bool(row["email_verified"]) if row["email_verified"] else False,
    )
    token = create_token(row["id"], row["org_id"], row["role"])
    logger.info("User logged in: %s", req.email)
    return TokenResponse(access_token=token, user=user)


@app.get("/auth/me", response_model=UserOut)
async def get_me(token_data: dict = Depends(verify_token)):
    conn = get_db()
    row = conn.execute("SELECT * FROM users WHERE id = ?", (token_data["user_id"],)).fetchone()
    conn.close()
    if not row:
        raise HTTPException(status_code=404, detail="User not found")
    return UserOut(
        id=row["id"], org_id=row["org_id"], email=row["email"],
        first_name=row["first_name"], last_name=row["last_name"], role=row["role"],
        hockey_role=row["hockey_role"] if row["hockey_role"] else "scout",
        subscription_tier=row["subscription_tier"] or "rookie",
        monthly_reports_used=row["monthly_reports_used"] or 0,
        monthly_bench_talks_used=row["monthly_bench_talks_used"] or 0,
        email_verified=bool(row["email_verified"]) if row["email_verified"] else False,
    )


class UpdateHockeyRoleRequest(BaseModel):
    hockey_role: str


class ForgotPasswordRequest(BaseModel):
    email: str


class ResetPasswordRequest(BaseModel):
    token: str
    new_password: str = Field(..., min_length=8)


class VerifyEmailRequest(BaseModel):
    token: str


@app.put("/auth/hockey-role")
async def update_hockey_role(req: UpdateHockeyRoleRequest, token_data: dict = Depends(verify_token)):
    """Update the user's hockey role (scout, gm, coach, player, parent, broadcaster, producer, agent)."""
    valid_roles = {"scout", "gm", "coach", "player", "parent", "broadcaster", "producer", "agent"}
    if req.hockey_role not in valid_roles:
        raise HTTPException(status_code=400, detail=f"Invalid hockey role. Must be one of: {', '.join(valid_roles)}")
    conn = get_db()
    conn.execute("UPDATE users SET hockey_role = ? WHERE id = ?", (req.hockey_role, token_data["user_id"]))
    conn.commit()
    row = conn.execute("SELECT * FROM users WHERE id = ?", (token_data["user_id"],)).fetchone()
    conn.close()
    return UserOut(
        id=row["id"], org_id=row["org_id"], email=row["email"],
        first_name=row["first_name"], last_name=row["last_name"], role=row["role"],
        hockey_role=row["hockey_role"] if row["hockey_role"] else "scout",
        subscription_tier=row["subscription_tier"] or "rookie",
        monthly_reports_used=row["monthly_reports_used"] or 0,
        monthly_bench_talks_used=row["monthly_bench_talks_used"] or 0,
        email_verified=bool(row["email_verified"]) if row["email_verified"] else False,
    )


# ============================================================
# PASSWORD RESET & EMAIL VERIFICATION
# ============================================================

@app.post("/auth/forgot-password")
async def forgot_password(req: ForgotPasswordRequest):
    """Request a password reset. Always returns 200 to not leak email existence."""
    conn = get_db()
    try:
        row = conn.execute("SELECT id FROM users WHERE email = ?", (req.email.lower().strip(),)).fetchone()
        if not row:
            return {"message": "If that email is registered, a reset link has been sent."}

        user_id = row["id"]

        # Rate limit: don't generate a new token if one was created in the last 5 minutes
        recent = conn.execute(
            "SELECT id FROM password_reset_tokens WHERE user_id = ? AND created_at > datetime('now', '-5 minutes') AND used = 0",
            (user_id,),
        ).fetchone()
        if recent:
            return {"message": "If that email is registered, a reset link has been sent."}

        # Generate token, store hashed
        raw_token = secrets.token_urlsafe(32)
        token_hash = hashlib.sha256(raw_token.encode()).hexdigest()
        expires_at = (datetime.now(timezone.utc) + timedelta(hours=1)).isoformat()

        conn.execute(
            "INSERT INTO password_reset_tokens (id, user_id, token_hash, expires_at) VALUES (?, ?, ?, ?)",
            (gen_id(), user_id, token_hash, expires_at),
        )
        conn.commit()

        # Alpha: log to console instead of sending email
        logger.info("PASSWORD RESET LINK: %s/reset-password?token=%s (email: %s)", FRONTEND_URL, raw_token, req.email)

        return {"message": "If that email is registered, a reset link has been sent."}
    finally:
        conn.close()


@app.post("/auth/reset-password")
async def reset_password(req: ResetPasswordRequest):
    """Reset password using a valid reset token."""
    token_hash = hashlib.sha256(req.token.encode()).hexdigest()
    conn = get_db()
    try:
        row = conn.execute(
            "SELECT id, user_id FROM password_reset_tokens WHERE token_hash = ? AND used = 0 AND expires_at > ?",
            (token_hash, datetime.now(timezone.utc).isoformat()),
        ).fetchone()

        if not row:
            raise HTTPException(status_code=400, detail="Invalid or expired reset token.")

        # Update password
        new_hash = pwd_context.hash(req.new_password[:72])
        conn.execute("UPDATE users SET password_hash = ? WHERE id = ?", (new_hash, row["user_id"]))

        # Mark token as used
        conn.execute("UPDATE password_reset_tokens SET used = 1 WHERE id = ?", (row["id"],))
        conn.commit()

        logger.info("Password reset successful for user_id: %s", row["user_id"])
        return {"message": "Password has been reset successfully. You can now log in."}
    finally:
        conn.close()


@app.post("/auth/send-verification")
async def send_verification_email(token_data: dict = Depends(verify_token)):
    """Resend email verification link. Requires auth."""
    conn = get_db()
    try:
        row = conn.execute("SELECT email, email_verified FROM users WHERE id = ?", (token_data["user_id"],)).fetchone()
        if not row:
            raise HTTPException(status_code=404, detail="User not found")
        if row["email_verified"]:
            return {"message": "Email is already verified."}

        verify_token_raw = secrets.token_urlsafe(32)
        conn.execute(
            "UPDATE users SET email_verify_token = ?, email_verify_sent_at = ? WHERE id = ?",
            (verify_token_raw, datetime.now(timezone.utc).isoformat(), token_data["user_id"]),
        )
        conn.commit()

        # Alpha: log to console instead of sending email
        logger.info("EMAIL VERIFICATION LINK: %s/verify-email?token=%s (user: %s)", FRONTEND_URL, verify_token_raw, row["email"])

        return {"message": "Verification email sent. Check the server console for the link."}
    finally:
        conn.close()


@app.post("/auth/verify-email")
async def verify_email(req: VerifyEmailRequest):
    """Verify email using token. No auth required (user clicks link from email)."""
    conn = get_db()
    try:
        row = conn.execute(
            "SELECT id, email_verify_sent_at FROM users WHERE email_verify_token = ?",
            (req.token,),
        ).fetchone()

        if not row:
            raise HTTPException(status_code=400, detail="Invalid verification token.")

        # Check token is within 24 hours
        sent_at = row["email_verify_sent_at"]
        if sent_at:
            sent_time = datetime.fromisoformat(sent_at)
            if datetime.now(timezone.utc) - sent_time > timedelta(hours=24):
                raise HTTPException(status_code=400, detail="Verification token has expired. Please request a new one.")

        # Mark email as verified, clear token
        conn.execute(
            "UPDATE users SET email_verified = 1, email_verify_token = NULL, email_verify_sent_at = NULL WHERE id = ?",
            (row["id"],),
        )
        conn.commit()

        logger.info("Email verified for user_id: %s", row["id"])
        return {"message": "Email verified successfully."}
    finally:
        conn.close()


# ============================================================
# SUBSCRIPTION ENDPOINTS
# ============================================================

@app.get("/subscription/tiers")
async def get_subscription_tiers():
    """Public endpoint: return all available subscription tiers."""
    return {"tiers": SUBSCRIPTION_TIERS}


@app.get("/subscription/usage")
async def get_subscription_usage(token_data: dict = Depends(verify_token)):
    """Get current user's subscription tier, per-resource usage, remaining counts, and permissions."""
    user_id = token_data["user_id"]
    conn = get_db()
    try:
        row = conn.execute(
            "SELECT subscription_tier, monthly_reports_used, monthly_bench_talks_used, usage_reset_at FROM users WHERE id = ?",
            (user_id,),
        ).fetchone()
        if not row:
            raise HTTPException(status_code=404, detail="User not found")

        tier = row["subscription_tier"] or "rookie"
        tier_config = SUBSCRIPTION_TIERS.get(tier, SUBSCRIPTION_TIERS["rookie"])

        # Get usage from new usage_tracking table
        now = datetime.now(timezone.utc)
        current_month = now.strftime("%Y-%m")
        tracking_row = conn.execute(
            "SELECT reports_count, bench_talks_count, practice_plans_count, uploads_count FROM usage_tracking WHERE user_id = ? AND month = ?",
            (user_id, current_month),
        ).fetchone()

        reports_used = (tracking_row["reports_count"] if tracking_row else None) or (row["monthly_reports_used"] or 0)
        bench_talks_used = (tracking_row["bench_talks_count"] if tracking_row else None) or (row["monthly_bench_talks_used"] or 0)
        practice_plans_used = (tracking_row["practice_plans_count"] if tracking_row else 0) or 0
        uploads_used = (tracking_row["uploads_count"] if tracking_row else 0) or 0

        def remaining(limit, used):
            return -1 if limit == -1 else max(0, limit - used)

        return {
            "tier": tier,
            "tier_config": tier_config,
            # Per-resource usage details
            "usage": {
                "reports": {
                    "used": reports_used,
                    "limit": tier_config["monthly_reports"],
                    "remaining": remaining(tier_config["monthly_reports"], reports_used),
                },
                "bench_talks": {
                    "used": bench_talks_used,
                    "limit": tier_config["monthly_bench_talks"],
                    "remaining": remaining(tier_config["monthly_bench_talks"], bench_talks_used),
                },
                "practice_plans": {
                    "used": practice_plans_used,
                    "limit": tier_config.get("monthly_practice_plans", 0),
                    "remaining": remaining(tier_config.get("monthly_practice_plans", 0), practice_plans_used),
                },
                "uploads": {
                    "used": uploads_used,
                    "limit": tier_config.get("max_uploads_per_month", 0),
                    "remaining": remaining(tier_config.get("max_uploads_per_month", 0), uploads_used),
                },
            },
            # Boolean permissions
            "permissions": {
                "can_sync_data": tier_config.get("can_sync_data", False),
                "can_upload_files": tier_config.get("can_upload_files", False),
                "can_access_live_stats": tier_config.get("can_access_live_stats", False),
            },
            # Limits
            "limits": {
                "max_file_size_mb": tier_config.get("max_file_size_mb", 5),
                "players_tracked": tier_config.get("players_tracked", 25),
            },
            # Legacy fields (backwards compat)
            "reports_used": reports_used,
            "reports_limit": tier_config["monthly_reports"],
            "bench_talks_used": bench_talks_used,
            "bench_talks_limit": tier_config["monthly_bench_talks"],
            "usage_reset_at": row["usage_reset_at"],
        }
    finally:
        conn.close()


@app.post("/subscription/upgrade")
async def upgrade_subscription(req: SubscriptionUpgradeRequest, token_data: dict = Depends(verify_token)):
    """Upgrade user's subscription tier (placeholder — no payment processing yet)."""
    user_id = token_data["user_id"]
    if req.tier not in SUBSCRIPTION_TIERS:
        raise HTTPException(status_code=400, detail="Invalid subscription tier")

    conn = get_db()
    try:
        conn.execute(
            "UPDATE users SET subscription_tier = ?, subscription_started_at = ? WHERE id = ?",
            (req.tier, datetime.now(timezone.utc).isoformat(), user_id),
        )
        conn.commit()
        return {"success": True, "tier": req.tier, "message": f"Upgraded to {SUBSCRIPTION_TIERS[req.tier]['name']}"}
    finally:
        conn.close()


# ============================================================
# ADMIN ENDPOINTS
# ============================================================

def _require_admin(token_data: dict):
    """Raise 403 if the user is not an admin."""
    if token_data.get("role") != "admin":
        raise HTTPException(status_code=403, detail="Admin access required")


@app.get("/admin/users")
async def admin_list_users(token_data: dict = Depends(verify_token)):
    """List all users in the admin's organization with usage data."""
    _require_admin(token_data)
    org_id = token_data["org_id"]

    conn = get_db()
    try:
        rows = conn.execute(
            """SELECT u.id, u.email, u.first_name, u.last_name, u.role, u.hockey_role,
                      u.subscription_tier, u.created_at, u.subscription_started_at,
                      u.monthly_reports_used, u.monthly_bench_talks_used
               FROM users u WHERE u.org_id = ? ORDER BY u.created_at""",
            (org_id,),
        ).fetchall()
        users = [dict(r) for r in rows]

        # Enrich with current month usage data
        now = datetime.now(timezone.utc)
        current_month = now.strftime("%Y-%m")
        for user in users:
            tracking = conn.execute(
                "SELECT reports_count, bench_talks_count, practice_plans_count, uploads_count FROM usage_tracking WHERE user_id = ? AND month = ?",
                (user["id"], current_month),
            ).fetchone()
            user["usage"] = dict(tracking) if tracking else {
                "reports_count": 0, "bench_talks_count": 0,
                "practice_plans_count": 0, "uploads_count": 0
            }
        return users
    finally:
        conn.close()


@app.put("/admin/users/{user_id}/tier")
async def admin_update_user_tier(user_id: str, req: AdminTierUpdateRequest,
                                  token_data: dict = Depends(verify_token)):
    """Change a user's subscription tier (admin only)."""
    _require_admin(token_data)
    org_id = token_data["org_id"]

    if req.tier not in SUBSCRIPTION_TIERS:
        raise HTTPException(status_code=400, detail=f"Invalid tier: {req.tier}. Valid tiers: {', '.join(SUBSCRIPTION_TIERS.keys())}")

    conn = get_db()
    try:
        # Verify user belongs to same org
        row = conn.execute("SELECT id, org_id, email FROM users WHERE id = ?", (user_id,)).fetchone()
        if not row:
            raise HTTPException(status_code=404, detail="User not found")
        if row["org_id"] != org_id:
            raise HTTPException(status_code=403, detail="Cannot modify users outside your organization")

        conn.execute(
            "UPDATE users SET subscription_tier = ?, subscription_started_at = ? WHERE id = ?",
            (req.tier, datetime.now(timezone.utc).isoformat(), user_id),
        )
        conn.commit()
        logger.info("Admin tier update: user %s (%s) → %s", row["email"], user_id, req.tier)
        return {
            "success": True,
            "user_id": user_id,
            "email": row["email"],
            "tier": req.tier,
            "tier_name": SUBSCRIPTION_TIERS[req.tier]["name"],
            "message": f"Tier updated to {SUBSCRIPTION_TIERS[req.tier]['name']}"
        }
    finally:
        conn.close()


@app.get("/admin/stats")
async def admin_platform_stats(token_data: dict = Depends(verify_token)):
    """Platform-wide statistics for admin dashboard."""
    _require_admin(token_data)
    org_id = token_data["org_id"]

    conn = get_db()
    try:
        stats = {}
        stats["total_users"] = conn.execute(
            "SELECT COUNT(*) FROM users WHERE org_id = ?", (org_id,)
        ).fetchone()[0]
        stats["total_players"] = conn.execute(
            "SELECT COUNT(*) FROM players WHERE org_id = ? AND (is_deleted = 0 OR is_deleted IS NULL)", (org_id,)
        ).fetchone()[0]
        stats["total_reports"] = conn.execute(
            "SELECT COUNT(*) FROM reports WHERE org_id = ?", (org_id,)
        ).fetchone()[0]
        stats["total_teams"] = conn.execute(
            "SELECT COUNT(*) FROM teams WHERE org_id IN (?, '__global__')", (org_id,)
        ).fetchone()[0]
        stats["total_notes"] = conn.execute(
            "SELECT COUNT(*) FROM scout_notes WHERE org_id = ?", (org_id,)
        ).fetchone()[0]
        stats["total_game_plans"] = conn.execute(
            "SELECT COUNT(*) FROM game_plans WHERE org_id = ?", (org_id,)
        ).fetchone()[0]
        stats["total_drills"] = conn.execute(
            "SELECT COUNT(*) FROM drills"
        ).fetchone()[0]
        stats["total_conversations"] = conn.execute(
            "SELECT COUNT(*) FROM chat_conversations WHERE org_id = ?", (org_id,)
        ).fetchone()[0]

        # Reports by status
        stats["reports_by_status"] = [
            dict(r) for r in conn.execute(
                "SELECT status, COUNT(*) as count FROM reports WHERE org_id = ? GROUP BY status",
                (org_id,),
            ).fetchall()
        ]

        # Users by tier
        stats["users_by_tier"] = [
            dict(r) for r in conn.execute(
                "SELECT COALESCE(subscription_tier, 'rookie') as tier, COUNT(*) as count FROM users WHERE org_id = ? GROUP BY tier",
                (org_id,),
            ).fetchall()
        ]

        # Recent activity (last 7 days)
        seven_days_ago = (datetime.now(timezone.utc) - timedelta(days=7)).isoformat()
        stats["recent_reports"] = conn.execute(
            "SELECT COUNT(*) FROM reports WHERE org_id = ? AND created_at > ?", (org_id, seven_days_ago)
        ).fetchone()[0]
        stats["recent_notes"] = conn.execute(
            "SELECT COUNT(*) FROM scout_notes WHERE org_id = ? AND created_at > ?", (org_id, seven_days_ago)
        ).fetchone()[0]

        return stats
    finally:
        conn.close()


@app.get("/admin/errors")
async def admin_error_log(
    limit: int = Query(50, ge=1, le=200),
    token_data: dict = Depends(verify_token),
):
    """Recent error logs for admin dashboard."""
    _require_admin(token_data)

    conn = get_db()
    try:
        rows = conn.execute(
            "SELECT * FROM admin_error_log ORDER BY created_at DESC LIMIT ?",
            (limit,),
        ).fetchall()
        return [dict(r) for r in rows]
    finally:
        conn.close()


@app.delete("/admin/errors")
async def admin_clear_errors(token_data: dict = Depends(verify_token)):
    """Clear all error logs (admin only)."""
    _require_admin(token_data)

    conn = get_db()
    try:
        conn.execute("DELETE FROM admin_error_log")
        conn.commit()
        return {"success": True, "message": "Error logs cleared"}
    finally:
        conn.close()


# ============================================================
# PLAYER ENDPOINTS
# ============================================================

def _player_from_row(row: sqlite3.Row) -> dict:
    d = dict(row)
    # Parse JSON list fields
    for field in ("passports", "tags"):
        val = d.get(field)
        if isinstance(val, str):
            try:
                d[field] = json.loads(val)
            except Exception:
                d[field] = []
        elif val is None:
            d[field] = []
    return d


# ── Player Intelligence Engine ─────────────────────────────────────────────


def _compute_stat_signature(player: dict, stats: list, goalie_stats: list, extended_stats_list: list) -> tuple:
    """Rule-based stat classifier. Fast, deterministic, no LLM needed.
    Returns (stat_signature dict, auto_tags list)."""
    sig = {}
    tags = set()

    position = (player.get("position") or "").upper()
    is_goalie = position in ("G", "GK", "GOALIE")

    if is_goalie and goalie_stats:
        # Goalie classification
        latest = goalie_stats[0]  # most recent
        gp = latest.get("gp") or 0
        sv_pct_str = latest.get("sv_pct") or ""
        gaa = latest.get("gaa") or 0

        # Parse sv_pct (might be string like ".920" or "0.920" or "92.0")
        try:
            sv_pct = float(sv_pct_str) if sv_pct_str else 0
            if sv_pct > 1:
                sv_pct = sv_pct / 100.0
        except (ValueError, TypeError):
            sv_pct = 0

        if gp >= 5:
            if sv_pct >= 0.930:
                sig["save_pct_tier"] = "Elite"
                tags.add("positioning")
            elif sv_pct >= 0.910:
                sig["save_pct_tier"] = "Strong"
            elif sv_pct >= 0.890:
                sig["save_pct_tier"] = "Average"
            else:
                sig["save_pct_tier"] = "Developing"

            if gaa <= 2.0:
                sig["goals_against_tier"] = "Elite"
            elif gaa <= 2.75:
                sig["goals_against_tier"] = "Strong"
            elif gaa <= 3.5:
                sig["goals_against_tier"] = "Average"
            else:
                sig["goals_against_tier"] = "Developing"

        # Check extended goalie stats
        if latest.get("extended_stats"):
            try:
                ext = json.loads(latest["extended_stats"]) if isinstance(latest["extended_stats"], str) else latest["extended_stats"]
                if ext:
                    tags.add("compete")
            except (json.JSONDecodeError, TypeError):
                pass

        return sig, list(tags)

    # Skater classification
    if not stats:
        return sig, list(tags)

    # Aggregate across all stat rows (use latest season or total)
    total_gp = 0
    total_g = 0
    total_a = 0
    total_p = 0
    total_pim = 0
    total_plus_minus = 0
    total_shots = 0

    for s in stats:
        gp = s.get("gp") or 0
        total_gp += gp
        total_g += s.get("g") or 0
        total_a += s.get("a") or 0
        total_p += s.get("p") or 0
        total_pim += s.get("pim") or 0
        total_plus_minus += s.get("plus_minus") or 0
        total_shots += s.get("shots") or s.get("sog") or 0

    if total_gp == 0:
        return sig, list(tags)

    ppg = total_p / total_gp
    gpg = total_g / total_gp
    apg = total_a / total_gp
    pim_pg = total_pim / total_gp

    # Production tier
    if ppg > 1.0:
        sig["production_tier"] = "Elite"
        tags.update(["shooting", "vision"])
    elif ppg > 0.7:
        sig["production_tier"] = "High"
        tags.add("puck_skills")
    elif ppg > 0.4:
        sig["production_tier"] = "Moderate"
    else:
        sig["production_tier"] = "Developing"

    # Goal-scoring profile
    if gpg > 0.4:
        sig["scoring_profile"] = "Sniper"
        tags.add("shooting")
    elif gpg > 0.25:
        sig["scoring_profile"] = "Scorer"
        tags.add("shooting")
    elif total_g > 0 and total_a > 1.5 * total_g:
        sig["scoring_profile"] = "Playmaker"
        tags.update(["vision", "puck_skills"])
    else:
        sig["scoring_profile"] = "Balanced"

    # Defensive reliability
    pm_per_game = total_plus_minus / total_gp
    if pm_per_game > 0.5:
        sig["defensive_reliability"] = "Elite"
        tags.add("positioning")
    elif total_plus_minus > 0:
        sig["defensive_reliability"] = "Reliable"
    elif total_plus_minus == 0:
        sig["defensive_reliability"] = "Neutral"
    else:
        sig["defensive_reliability"] = "Developing"

    # Discipline
    if pim_pg < 0.5:
        sig["discipline"] = "Clean"
    elif pim_pg < 1.5:
        sig["discipline"] = "Moderate"
    else:
        sig["discipline"] = "Physical"
        tags.add("physicality")

    # Shooting efficiency
    if total_shots > 0:
        sh_pct = total_g / total_shots * 100
        if sh_pct > 15:
            sig["shooting_efficiency"] = "Elite"
            tags.add("shooting")
        elif sh_pct > 10:
            sig["shooting_efficiency"] = "Above Average"
        elif sh_pct > 5:
            sig["shooting_efficiency"] = "Average"
        else:
            sig["shooting_efficiency"] = "Below Average"

    # ── Extended stats analysis (InStat analytics) ──
    for ext_row in extended_stats_list:
        try:
            ext = json.loads(ext_row) if isinstance(ext_row, str) else ext_row
            if not ext or not isinstance(ext, dict):
                continue
        except (json.JSONDecodeError, TypeError):
            continue

        # xG finishing (check multiple possible locations)
        xg_data = ext.get("xg") or ext.get("advanced") or {}
        xg_val = xg_data.get("xG") or xg_data.get("xg") or 0
        if xg_val and total_g > 0:
            try:
                xg_float = float(xg_val)
                if xg_float > 0:
                    ratio = total_g / xg_float
                    if ratio > 1.15:
                        sig["finishing"] = "Over-performer"
                        tags.add("shooting")
                    elif ratio > 0.85:
                        sig["finishing"] = "Expected"
                    else:
                        sig["finishing"] = "Under-performer"
            except (ValueError, TypeError):
                pass

        # Possession (CORSI / possession %)
        poss_data = ext.get("main") or ext.get("advanced") or {}
        cf_pct = poss_data.get("CF%") or poss_data.get("cf_pct") or poss_data.get("Possession, %") or 0
        try:
            cf_float = float(cf_pct)
            if cf_float > 0:
                if cf_float > 55:
                    sig["possession_impact"] = "Driver"
                    tags.update(["hockey_iq", "puck_skills"])
                elif cf_float > 50:
                    sig["possession_impact"] = "Positive"
                elif cf_float > 45:
                    sig["possession_impact"] = "Neutral"
                else:
                    sig["possession_impact"] = "Passenger"
        except (ValueError, TypeError):
            pass

        # Puck battles
        battles_data = ext.get("puck_battles") or ext.get("duels") or {}
        battle_won = battles_data.get("Won") or battles_data.get("won") or 0
        battle_total = battles_data.get("Total") or battles_data.get("total") or 0
        try:
            b_won = float(battle_won)
            b_total = float(battle_total)
            if b_total > 0:
                win_pct = (b_won / b_total) * 100
                if win_pct > 55:
                    sig["physical_engagement"] = "Battle Winner"
                    tags.update(["physicality", "compete"])
                elif win_pct > 45:
                    sig["physical_engagement"] = "Competitive"
                    tags.add("compete")
                else:
                    sig["physical_engagement"] = "Avoid Contact"
        except (ValueError, TypeError):
            pass

        # Faceoffs
        fo_data = ext.get("faceoffs") or ext.get("main") or {}
        fo_won = fo_data.get("FO Won") or fo_data.get("fo_won") or 0
        fo_total = fo_data.get("FO Total") or fo_data.get("fo_total") or fo_data.get("Faceoffs") or 0
        try:
            fw = float(fo_won)
            ft = float(fo_total)
            if ft > 0:
                fo_pct = (fw / ft) * 100
                if fo_pct > 55:
                    sig["faceoff_ability"] = "Elite"
                    tags.add("hockey_iq")
                elif fo_pct > 50:
                    sig["faceoff_ability"] = "Strong"
                else:
                    sig["faceoff_ability"] = "Developing"
        except (ValueError, TypeError):
            pass

        break  # Use first valid extended stats row

    # Position-based tag hints
    if position in ("D", "LD", "RD"):
        if sig.get("defensive_reliability") in ("Elite", "Reliable"):
            tags.add("positioning")
        if ppg > 0.5:
            tags.add("puck_skills")  # offensive D
    elif position in ("C",):
        if sig.get("faceoff_ability") in ("Elite", "Strong"):
            tags.add("hockey_iq")

    return sig, list(tags)


async def _generate_player_intelligence(player_id: str, org_id: str, trigger: str = "manual") -> Optional[dict]:
    """Generate or refresh a player's AI intelligence profile.
    Uses Claude Haiku for fast, cheap structured JSON output."""
    conn = get_db()
    try:
        # 1. Get player
        player_row = conn.execute("SELECT * FROM players WHERE id = ? AND org_id = ?", (player_id, org_id)).fetchone()
        if not player_row:
            return None
        player = _player_from_row(player_row)

        # 2. Get stats
        stats_rows = conn.execute(
            "SELECT * FROM player_stats WHERE player_id = ? ORDER BY season DESC", (player_id,)
        ).fetchall()
        stats = [dict(r) for r in stats_rows]

        # 3. Get goalie stats
        goalie_rows = conn.execute(
            "SELECT * FROM goalie_stats WHERE player_id = ? ORDER BY season DESC", (player_id,)
        ).fetchall()
        goalie_stats = [dict(r) for r in goalie_rows]

        # 4. Get extended stats
        extended_stats_list = []
        for s in stats:
            if s.get("extended_stats"):
                extended_stats_list.append(s["extended_stats"])
        for g in goalie_stats:
            if g.get("extended_stats"):
                extended_stats_list.append(g["extended_stats"])

        # 5. Compute stat signature (rule-based, instant)
        stat_sig, auto_tags = _compute_stat_signature(player, stats, goalie_stats, extended_stats_list)

        # 6. Get scout notes (last 10)
        notes_rows = conn.execute(
            "SELECT note_text, note_type, tags, created_at FROM scout_notes WHERE player_id = ? ORDER BY created_at DESC LIMIT 10",
            (player_id,)
        ).fetchall()
        notes = [dict(n) for n in notes_rows]

        # 7. Get previous intelligence snapshot (for continuity)
        prev_intel = conn.execute(
            "SELECT * FROM player_intelligence WHERE player_id = ? AND org_id = ? ORDER BY version DESC LIMIT 1",
            (player_id, org_id)
        ).fetchone()
        prev_version = dict(prev_intel)["version"] if prev_intel else 0

        # 8. Get latest report summary
        latest_report = conn.execute(
            "SELECT output_text, report_type, generated_at FROM reports WHERE player_id = ? AND status = 'complete' ORDER BY generated_at DESC LIMIT 1",
            (player_id,)
        ).fetchone()

        # 9. Build context for LLM
        # Pre-compute age
        age_str = "Unknown"
        if player.get("dob"):
            try:
                dob_str = player["dob"][:10]
                birth = datetime.strptime(dob_str, "%Y-%m-%d").date()
                today = datetime.now().date()
                age = today.year - birth.year - ((today.month, today.day) < (birth.month, birth.day))
                age_str = str(age)
            except Exception:
                pass

        bio_section = f"""PLAYER BIO:
Name: {player.get('first_name', '')} {player.get('last_name', '')}
Position: {player.get('position', 'Unknown')}
Age: {age_str}
DOB: {player.get('dob', 'Unknown')}
Height: {player.get('height_cm', 'N/A')} cm
Weight: {player.get('weight_kg', 'N/A')} kg
Shoots: {player.get('shoots', 'Unknown')}
Team: {player.get('current_team', 'Unknown')}
League: {player.get('current_league', 'Unknown')}"""

        stats_section = ""
        if stats:
            stats_section = "\n\nSTATISTICS:\n"
            for s in stats[:3]:  # Last 3 seasons max
                season = s.get("season", "Unknown")
                stats_section += f"Season {season}: {s.get('gp',0)} GP, {s.get('g',0)}G-{s.get('a',0)}A—{s.get('p',0)}P, +/- {s.get('plus_minus',0)}, {s.get('pim',0)} PIM"
                if s.get("shots"):
                    stats_section += f", {s.get('shots',0)} SOG"
                if s.get("shooting_pct"):
                    stats_section += f", {s.get('shooting_pct',0)}% SH%"
                stats_section += "\n"

        if goalie_stats:
            stats_section += "\nGOALIE STATISTICS:\n"
            for g in goalie_stats[:3]:
                season = g.get("season", "Unknown")
                stats_section += f"Season {season}: {g.get('gp',0)} GP, {g.get('gaa','N/A')} GAA, {g.get('sv_pct','N/A')} SV%, {g.get('ga',0)} GA, {g.get('sa',0)} SA\n"

        sig_section = ""
        if stat_sig:
            sig_section = "\n\nSTAT SIGNATURE (rule-based analysis):\n"
            for k, v in stat_sig.items():
                sig_section += f"- {k.replace('_', ' ').title()}: {v}\n"

        notes_section = ""
        if notes:
            notes_section = "\n\nSCOUT NOTES (most recent first):\n"
            for n in notes[:5]:  # Cap at 5 for prompt size
                notes_section += f"- [{n.get('note_type','general')}] {n.get('note_text','')[:300]}\n"

        report_section = ""
        if latest_report:
            report_text = dict(latest_report).get("output_text", "")
            if report_text:
                # Extract key sections (first 800 chars)
                report_section = f"\n\nLATEST REPORT EXCERPT:\n{report_text[:800]}..."

        prev_section = ""
        if prev_intel:
            pi = dict(prev_intel)
            prev_section = f"\n\nPREVIOUS INTELLIGENCE (v{pi.get('version',0)}):\nArchetype: {pi.get('archetype','N/A')}\nOverall Grade: {pi.get('overall_grade','N/A')}\nSummary: {(pi.get('summary','N/A') or 'N/A')[:300]}"

        full_context = bio_section + stats_section + sig_section + notes_section + report_section + prev_section

        # 10. Call Claude Haiku for intelligence
        client = get_anthropic_client()
        if not client:
            # No API key — store stat signature only
            intel_id = str(uuid.uuid4())
            now = datetime.now().isoformat()
            conn.execute("""
                INSERT INTO player_intelligence (id, player_id, org_id, stat_signature, tags, trigger, version, data_sources_used, created_at)
                VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)
            """, (intel_id, player_id, org_id, json.dumps(stat_sig), json.dumps(auto_tags), trigger, prev_version + 1,
                  json.dumps({"stats": len(stats), "goalie_stats": len(goalie_stats), "notes": len(notes), "reports": 1 if latest_report else 0}), now))
            conn.execute("UPDATE players SET tags = ?, intelligence_version = ? WHERE id = ?",
                         (json.dumps(auto_tags), prev_version + 1, player_id))
            conn.commit()
            return {"id": intel_id, "stat_signature": stat_sig, "tags": auto_tags, "version": prev_version + 1}

        system_prompt = """You are ProspectX Intelligence — an elite hockey scouting AI that produces structured player intelligence profiles.

Given a player's bio, statistics, stat signature, scout notes, and any previous intelligence, produce a JSON object with your assessment.

RULES:
- Return ONLY valid JSON — no markdown, no explanation, no code fences
- Grade scale: A+, A, A-, B+, B, B-, C+, C, C-, D+, D, D-, NR (Not Rated — use when insufficient data)
- Archetype should be a compound descriptor using recognized hockey types: Two-Way Center, Two-Way Playmaking Center, Power Forward, Sniper, Playmaker, Grinder, Mucker, Puck-Moving D, Stay-at-Home D, Shutdown D, Offensive D, Two-Way Winger, Energy Forward, Net-Front Presence, Two-Way Defenseman, Offensive Defenseman, etc.
- archetype_confidence is 0.0 to 1.0 (higher = more data available)
- strengths and development_areas: 3-5 items each, specific to hockey skills. Use proper terminology — "gap control", "net-front presence", "cycle game", "F1 forecheck pressure", "transition game", "puck protection along the boards", "one-timer release", not vague phrases.
- comparable_players: 1-2 NHL/pro comparisons that match the player's style (be realistic, these are junior players)
- projection: 1-2 sentences about ceiling/floor at next level
- tags must be from: skating, shooting, compete, hockey_iq, puck_skills, positioning, physicality, speed, vision, leadership, coachability, work_ethic
- summary: 2-3 sentence scouting summary, professional tone. Use hockey vernacular — "beauty with wheels", "grinder who goes to the dirty areas", "sniper with a quick release from the slot"
- If data is very limited, be conservative with grades (use NR liberally) but still provide archetype/summary based on what's available

JSON SCHEMA:
{
  "archetype": string,
  "archetype_confidence": number,
  "overall_grade": string,
  "offensive_grade": string,
  "defensive_grade": string,
  "skating_grade": string,
  "hockey_iq_grade": string,
  "compete_grade": string,
  "summary": string,
  "strengths": [string],
  "development_areas": [string],
  "comparable_players": [string],
  "projection": string,
  "tags": [string]
}"""

        try:
            response = client.messages.create(
                model="claude-sonnet-4-20250514",
                max_tokens=1000,
                system=system_prompt,
                messages=[{"role": "user", "content": f"Generate an intelligence profile for this player:\n\n{full_context}"}]
            )

            raw_text = response.content[0].text.strip()
            # Strip code fences if present
            if raw_text.startswith("```"):
                raw_text = re.sub(r"^```(?:json)?\s*\n?", "", raw_text)
                raw_text = re.sub(r"\n?```\s*$", "", raw_text)

            intel_data = json.loads(raw_text)
        except json.JSONDecodeError as e:
            logger.error("Intelligence JSON parse error for player %s: %s", player_id, str(e))
            # Fall back to stat signature only
            intel_data = {}
        except Exception as e:
            logger.error("Intelligence LLM call failed for player %s: %s", player_id, str(e))
            intel_data = {}

        # 11. Merge LLM results with stat signature
        archetype = intel_data.get("archetype") or player.get("archetype")
        archetype_confidence = intel_data.get("archetype_confidence", 0.5)
        overall_grade = intel_data.get("overall_grade", "NR")
        offensive_grade = intel_data.get("offensive_grade", "NR")
        defensive_grade = intel_data.get("defensive_grade", "NR")
        skating_grade = intel_data.get("skating_grade", "NR")
        hockey_iq_grade = intel_data.get("hockey_iq_grade", "NR")
        compete_grade = intel_data.get("compete_grade", "NR")
        summary = intel_data.get("summary", "")
        strengths = intel_data.get("strengths", [])
        dev_areas = intel_data.get("development_areas", [])
        comparables = intel_data.get("comparable_players", [])
        projection = intel_data.get("projection", "")
        llm_tags = intel_data.get("tags", [])

        # Merge LLM tags with rule-based auto_tags
        all_tags = list(set(auto_tags + llm_tags))
        # Filter to valid tags only
        valid_tags = {"skating", "shooting", "compete", "hockey_iq", "puck_skills", "positioning",
                      "physicality", "speed", "vision", "leadership", "coachability", "work_ethic"}
        all_tags = [t for t in all_tags if t in valid_tags]

        # 12. Store intelligence
        intel_id = str(uuid.uuid4())
        new_version = prev_version + 1
        now = datetime.now().isoformat()
        data_sources = {
            "stats": len(stats),
            "goalie_stats": len(goalie_stats),
            "notes": len(notes),
            "reports": 1 if latest_report else 0,
            "extended_stats": len(extended_stats_list)
        }

        conn.execute("""
            INSERT INTO player_intelligence
            (id, player_id, org_id, archetype, archetype_confidence, overall_grade, offensive_grade, defensive_grade,
             skating_grade, hockey_iq_grade, compete_grade, summary, strengths, development_areas, comparable_players,
             stat_signature, tags, projection, data_sources_used, trigger, version, created_at)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
        """, (intel_id, player_id, org_id, archetype, archetype_confidence, overall_grade, offensive_grade,
              defensive_grade, skating_grade, hockey_iq_grade, compete_grade, summary,
              json.dumps(strengths), json.dumps(dev_areas), json.dumps(comparables),
              json.dumps(stat_sig), json.dumps(all_tags), projection, json.dumps(data_sources),
              trigger, new_version, now))

        # 13. Write back to players table
        conn.execute("UPDATE players SET archetype = ?, tags = ?, intelligence_version = ? WHERE id = ?",
                     (archetype, json.dumps(all_tags), new_version, player_id))
        conn.commit()

        logger.info("Intelligence v%d generated for player %s (trigger=%s)", new_version, player_id, trigger)

        return {
            "id": intel_id,
            "player_id": player_id,
            "archetype": archetype,
            "archetype_confidence": archetype_confidence,
            "overall_grade": overall_grade,
            "offensive_grade": offensive_grade,
            "defensive_grade": defensive_grade,
            "skating_grade": skating_grade,
            "hockey_iq_grade": hockey_iq_grade,
            "compete_grade": compete_grade,
            "summary": summary,
            "strengths": strengths,
            "development_areas": dev_areas,
            "comparable_players": comparables,
            "stat_signature": stat_sig,
            "tags": all_tags,
            "projection": projection,
            "trigger": trigger,
            "version": new_version,
            "created_at": now,
        }
    except Exception as e:
        logger.error("Intelligence generation failed for player %s: %s", player_id, str(e))
        return None
    finally:
        conn.close()


async def _extract_report_intelligence(report_id: str, player_id: str, org_id: str):
    """Parse a completed report to extract insights, then refresh player intelligence."""
    conn = get_db()
    try:
        report = conn.execute("SELECT output_text FROM reports WHERE id = ?", (report_id,)).fetchone()
        if not report:
            return
        text = dict(report).get("output_text", "") or ""
        if not text:
            return

        # Extract key data from report text using regex
        extracted = {}

        # Overall Grade
        grade_match = re.search(r"Overall\s*Grade[:\s]*([A-D][+-]?)", text, re.IGNORECASE)
        if grade_match:
            extracted["overall_grade"] = grade_match.group(1)

        # System Fit
        fit_match = re.search(r"(Elite Fit|Strong Fit|Developing Fit|Adjustment Needed)", text, re.IGNORECASE)
        if fit_match:
            extracted["system_fit"] = fit_match.group(1)

        # Archetype mentions
        arch_match = re.search(r"(?:profiles?|projects?|presents?|plays?)\s+as\s+(?:a|an)\s+(.+?)(?:\.|,|\n)", text, re.IGNORECASE)
        if arch_match:
            extracted["archetype_hint"] = arch_match.group(1).strip()[:80]

        logger.info("Report %s intelligence extracted: %s", report_id, list(extracted.keys()))
    except Exception as e:
        logger.error("Report intelligence extraction failed for %s: %s", report_id, str(e))
    finally:
        conn.close()

    # Now refresh player intelligence with the new report context
    try:
        await _generate_player_intelligence(player_id, org_id, trigger="report")
    except Exception as e:
        logger.error("Post-report intelligence refresh failed for %s: %s", player_id, str(e))


@app.get("/players/filter-options")
async def get_player_filter_options(token_data: dict = Depends(verify_token)):
    """Return unique values for all filter dimensions (for populating dropdowns)."""
    org_id = token_data["org_id"]
    conn = get_db()
    result = {}
    _visible = "AND (is_deleted = 0 OR is_deleted IS NULL)"
    # Unique leagues
    rows = conn.execute(f"SELECT DISTINCT current_league FROM players WHERE org_id = ? {_visible} AND current_league IS NOT NULL AND current_league != '' ORDER BY current_league", (org_id,)).fetchall()
    result["leagues"] = [r[0] for r in rows]
    # Unique teams
    rows = conn.execute(f"SELECT DISTINCT current_team FROM players WHERE org_id = ? {_visible} AND current_team IS NOT NULL AND current_team != '' ORDER BY current_team", (org_id,)).fetchall()
    result["teams"] = [r[0] for r in rows]
    # Unique birth years
    rows = conn.execute(f"SELECT DISTINCT birth_year FROM players WHERE org_id = ? {_visible} AND birth_year IS NOT NULL ORDER BY birth_year DESC", (org_id,)).fetchall()
    result["birth_years"] = [r[0] for r in rows]
    # Unique age groups
    rows = conn.execute(f"SELECT DISTINCT age_group FROM players WHERE org_id = ? {_visible} AND age_group IS NOT NULL ORDER BY age_group", (org_id,)).fetchall()
    result["age_groups"] = [r[0] for r in rows]
    # Unique league tiers
    rows = conn.execute(f"SELECT DISTINCT league_tier FROM players WHERE org_id = ? {_visible} AND league_tier IS NOT NULL AND league_tier != 'Unknown' ORDER BY league_tier", (org_id,)).fetchall()
    result["league_tiers"] = [r[0] for r in rows]
    # Unique positions
    rows = conn.execute(f"SELECT DISTINCT position FROM players WHERE org_id = ? {_visible} AND position IS NOT NULL ORDER BY position", (org_id,)).fetchall()
    result["positions"] = [r[0] for r in rows]
    # Unique draft eligible years
    rows = conn.execute(f"SELECT DISTINCT draft_eligible_year FROM players WHERE org_id = ? {_visible} AND draft_eligible_year IS NOT NULL ORDER BY draft_eligible_year DESC", (org_id,)).fetchall()
    result["draft_years"] = [r[0] for r in rows]
    # Unique commitment statuses
    rows = conn.execute(f"SELECT DISTINCT commitment_status FROM players WHERE org_id = ? {_visible} AND commitment_status IS NOT NULL AND commitment_status != '' ORDER BY commitment_status", (org_id,)).fetchall()
    result["commitment_statuses"] = [r[0] for r in rows]

    # Unique shoots values
    rows = conn.execute("SELECT DISTINCT shoots FROM players WHERE org_id = ? AND shoots IS NOT NULL AND shoots != '' ORDER BY shoots", (org_id,)).fetchall()
    result["shoots"] = [r[0] for r in rows]

    # Unique archetypes (from player_intelligence)
    rows = conn.execute("""
        SELECT DISTINCT pi.archetype FROM player_intelligence pi
        JOIN players p ON pi.player_id = p.id
        WHERE p.org_id = ? AND pi.archetype IS NOT NULL AND pi.archetype != ''
        ORDER BY pi.archetype
    """, (org_id,)).fetchall()
    result["archetypes"] = [r[0] for r in rows]

    # Unique overall grades (from player_intelligence)
    rows = conn.execute("""
        SELECT DISTINCT pi.overall_grade FROM player_intelligence pi
        JOIN players p ON pi.player_id = p.id
        WHERE p.org_id = ? AND pi.overall_grade IS NOT NULL AND pi.overall_grade != '' AND pi.overall_grade != 'NR'
        ORDER BY pi.overall_grade
    """, (org_id,)).fetchall()
    result["overall_grades"] = [r[0] for r in rows]

    # Height range
    row = conn.execute("SELECT MIN(height_cm), MAX(height_cm) FROM players WHERE org_id = ? AND height_cm IS NOT NULL AND height_cm > 0", (org_id,)).fetchone()
    result["height_range"] = {"min": row[0], "max": row[1]} if row and row[0] else None

    # Weight range
    row = conn.execute("SELECT MIN(weight_kg), MAX(weight_kg) FROM players WHERE org_id = ? AND weight_kg IS NOT NULL AND weight_kg > 0", (org_id,)).fetchone()
    result["weight_range"] = {"min": row[0], "max": row[1]} if row and row[0] else None

    conn.close()
    return result


@app.get("/players", response_model=List[PlayerResponse])
async def list_players(
    search: Optional[str] = None,
    position: Optional[str] = None,
    team: Optional[str] = None,
    league: Optional[str] = None,
    birth_year: Optional[int] = None,
    age_group: Optional[str] = None,
    league_tier: Optional[str] = None,
    draft_year: Optional[int] = None,
    commitment_status: Optional[str] = None,
    shoots: Optional[str] = None,
    min_height: Optional[int] = None,
    max_height: Optional[int] = None,
    min_weight: Optional[int] = None,
    max_weight: Optional[int] = None,
    min_gp: Optional[int] = None,
    min_goals: Optional[int] = None,
    min_points: Optional[int] = None,
    min_ppg: Optional[float] = None,
    has_stats: Optional[bool] = None,
    overall_grade: Optional[str] = None,
    archetype: Optional[str] = None,
    sort_by: Optional[str] = None,
    sort_dir: Optional[str] = Query(default="asc", pattern="^(asc|desc)$"),
    limit: int = Query(default=100, ge=1, le=5000),
    skip: int = Query(default=0, ge=0),
    token_data: dict = Depends(verify_token),
):
    org_id = token_data["org_id"]
    conn = get_db()

    # Determine if we need joins
    needs_stats_join = any([min_gp, min_goals, min_points, min_ppg, has_stats,
                           sort_by in ("gp", "g", "goals", "a", "assists", "p", "points", "ppg")])
    needs_intel_join = any([overall_grade, archetype,
                           sort_by in ("grade", "overall_grade")])

    if needs_stats_join or needs_intel_join:
        # Use JOIN-based query for advanced filters
        select = "SELECT DISTINCT p.*"
        from_clause = " FROM players p"
        where_clauses = ["p.org_id = ?", "(p.is_deleted = 0 OR p.is_deleted IS NULL)"]
        params: list = [org_id]

        if needs_stats_join:
            from_clause += """
                LEFT JOIN (
                    SELECT player_id, SUM(gp) AS gp, SUM(g) AS g, SUM(a) AS a, SUM(p) AS p
                    FROM player_stats WHERE stat_type = 'season' GROUP BY player_id
                ) ps ON p.id = ps.player_id"""

        if needs_intel_join:
            from_clause += """
                LEFT JOIN (
                    SELECT player_id, overall_grade, archetype FROM player_intelligence
                    WHERE (player_id, version) IN (SELECT player_id, MAX(version) FROM player_intelligence GROUP BY player_id)
                ) pi ON p.id = pi.player_id"""

        if search:
            where_clauses.append("(p.first_name LIKE ? OR p.last_name LIKE ? OR p.current_team LIKE ?)")
            s = f"%{search}%"
            params.extend([s, s, s])
        if position:
            where_clauses.append("p.position = ?")
            params.append(position.upper())
        if team:
            where_clauses.append("LOWER(p.current_team) = LOWER(?)")
            params.append(team)
        if league:
            where_clauses.append("LOWER(p.current_league) = LOWER(?)")
            params.append(league)
        if birth_year:
            where_clauses.append("p.birth_year = ?")
            params.append(birth_year)
        if age_group:
            where_clauses.append("p.age_group = ?")
            params.append(age_group)
        if league_tier:
            where_clauses.append("p.league_tier = ?")
            params.append(league_tier)
        if draft_year:
            where_clauses.append("p.draft_eligible_year = ?")
            params.append(draft_year)
        if commitment_status:
            where_clauses.append("p.commitment_status = ?")
            params.append(commitment_status)
        if shoots:
            where_clauses.append("p.shoots = ?")
            params.append(shoots.upper())
        if min_height:
            where_clauses.append("p.height_cm >= ?")
            params.append(min_height)
        if max_height:
            where_clauses.append("p.height_cm <= ?")
            params.append(max_height)
        if min_weight:
            where_clauses.append("p.weight_kg >= ?")
            params.append(min_weight)
        if max_weight:
            where_clauses.append("p.weight_kg <= ?")
            params.append(max_weight)
        if has_stats:
            where_clauses.append("ps.gp IS NOT NULL AND ps.gp > 0")
        if min_gp:
            where_clauses.append("ps.gp >= ?")
            params.append(min_gp)
        if min_goals:
            where_clauses.append("ps.g >= ?")
            params.append(min_goals)
        if min_points:
            where_clauses.append("ps.p >= ?")
            params.append(min_points)
        if min_ppg:
            where_clauses.append("ps.gp > 0 AND (CAST(ps.p AS REAL) / ps.gp) >= ?")
            params.append(min_ppg)
        if overall_grade:
            where_clauses.append("pi.overall_grade = ?")
            params.append(overall_grade)
        if archetype:
            where_clauses.append("pi.archetype = ?")
            params.append(archetype)

        # Sorting
        sort_map = {
            "name": "p.last_name", "last_name": "p.last_name", "first_name": "p.first_name",
            "gp": "ps.gp", "g": "ps.g", "goals": "ps.g", "a": "ps.a", "assists": "ps.a",
            "p": "ps.p", "points": "ps.p", "ppg": "CAST(ps.p AS REAL) / NULLIF(ps.gp, 0)",
            "grade": "pi.overall_grade", "overall_grade": "pi.overall_grade",
            "team": "p.current_team", "position": "p.position",
        }
        order_col = sort_map.get(sort_by, "p.last_name")
        order_dir = "DESC" if sort_dir == "desc" else "ASC"
        order_sql = f" ORDER BY {order_col} {order_dir} NULLS LAST, p.last_name ASC"

        query = select + from_clause + " WHERE " + " AND ".join(where_clauses) + order_sql + " LIMIT ? OFFSET ?"
        params.extend([limit, skip])
    else:
        # Simple query (no joins needed) — fast path
        query = "SELECT * FROM players WHERE org_id = ? AND (is_deleted = 0 OR is_deleted IS NULL)"
        params = [org_id]

        if search:
            query += " AND (first_name LIKE ? OR last_name LIKE ? OR current_team LIKE ?)"
            s = f"%{search}%"
            params.extend([s, s, s])
        if position:
            query += " AND position = ?"
            params.append(position.upper())
        if team:
            query += " AND LOWER(current_team) = LOWER(?)"
            params.append(team)
        if league:
            query += " AND LOWER(current_league) = LOWER(?)"
            params.append(league)
        if birth_year:
            query += " AND birth_year = ?"
            params.append(birth_year)
        if age_group:
            query += " AND age_group = ?"
            params.append(age_group)
        if league_tier:
            query += " AND league_tier = ?"
            params.append(league_tier)
        if draft_year:
            query += " AND draft_eligible_year = ?"
            params.append(draft_year)
        if commitment_status:
            query += " AND commitment_status = ?"
            params.append(commitment_status)
        if shoots:
            query += " AND shoots = ?"
            params.append(shoots.upper())
        if min_height:
            query += " AND height_cm >= ?"
            params.append(min_height)
        if max_height:
            query += " AND height_cm <= ?"
            params.append(max_height)
        if min_weight:
            query += " AND weight_kg >= ?"
            params.append(min_weight)
        if max_weight:
            query += " AND weight_kg <= ?"
            params.append(max_weight)

        # Sorting for simple path
        sort_map = {"name": "last_name", "last_name": "last_name", "first_name": "first_name",
                    "team": "current_team", "position": "position"}
        order_col = sort_map.get(sort_by, "last_name")
        order_dir = "DESC" if sort_dir == "desc" else "ASC"
        query += f" ORDER BY {order_col} {order_dir}, first_name ASC LIMIT ? OFFSET ?"
        params.extend([limit, skip])

    rows = conn.execute(query, params).fetchall()
    conn.close()
    return [PlayerResponse(**_player_from_row(r)) for r in rows]


@app.get("/players/cards")
async def list_player_cards(
    search: Optional[str] = None,
    position: Optional[str] = None,
    team: Optional[str] = None,
    league: Optional[str] = None,
    birth_year: Optional[int] = None,
    age_group: Optional[str] = None,
    league_tier: Optional[str] = None,
    draft_year: Optional[int] = None,
    commitment_status: Optional[str] = None,
    shoots: Optional[str] = None,
    min_height: Optional[int] = None,
    max_height: Optional[int] = None,
    min_weight: Optional[int] = None,
    max_weight: Optional[int] = None,
    min_gp: Optional[int] = None,
    min_goals: Optional[int] = None,
    min_points: Optional[int] = None,
    min_ppg: Optional[float] = None,
    has_stats: Optional[bool] = None,
    overall_grade: Optional[str] = None,
    archetype: Optional[str] = None,
    sort_by: Optional[str] = None,
    sort_dir: Optional[str] = Query(default="asc", pattern="^(asc|desc)$"),
    limit: int = Query(default=50, ge=1, le=200),
    skip: int = Query(default=0, ge=0),
    token_data: dict = Depends(verify_token),
):
    """Return enriched player card data: player info + intelligence grades + stats + ProspectX metrics."""
    org_id = token_data["org_id"]
    conn = get_db()

    # Build WHERE clause (same filters as GET /players)
    where_clauses = ["p.org_id = ?", "(p.is_deleted = 0 OR p.is_deleted IS NULL)"]
    params: list = [org_id]

    if search:
        where_clauses.append("(p.first_name LIKE ? OR p.last_name LIKE ? OR p.current_team LIKE ?)")
        s = f"%{search}%"
        params.extend([s, s, s])
    if position:
        where_clauses.append("p.position = ?")
        params.append(position.upper())
    if team:
        where_clauses.append("LOWER(p.current_team) = LOWER(?)")
        params.append(team)
    if league:
        where_clauses.append("LOWER(p.current_league) = LOWER(?)")
        params.append(league)
    if birth_year:
        where_clauses.append("p.birth_year = ?")
        params.append(birth_year)
    if age_group:
        where_clauses.append("p.age_group = ?")
        params.append(age_group)
    if league_tier:
        where_clauses.append("p.league_tier = ?")
        params.append(league_tier)
    if draft_year:
        where_clauses.append("p.draft_eligible_year = ?")
        params.append(draft_year)
    if commitment_status:
        where_clauses.append("p.commitment_status = ?")
        params.append(commitment_status)
    if shoots:
        where_clauses.append("p.shoots = ?")
        params.append(shoots.upper())
    if min_height:
        where_clauses.append("p.height_cm >= ?")
        params.append(min_height)
    if max_height:
        where_clauses.append("p.height_cm <= ?")
        params.append(max_height)
    if min_weight:
        where_clauses.append("p.weight_kg >= ?")
        params.append(min_weight)
    if max_weight:
        where_clauses.append("p.weight_kg <= ?")
        params.append(max_weight)

    # Stats-based filters (applied to the already-joined ps subquery)
    if has_stats:
        where_clauses.append("ps.gp IS NOT NULL AND ps.gp > 0")
    if min_gp:
        where_clauses.append("ps.gp >= ?")
        params.append(min_gp)
    if min_goals:
        where_clauses.append("ps.g >= ?")
        params.append(min_goals)
    if min_points:
        where_clauses.append("ps.p >= ?")
        params.append(min_points)
    if min_ppg:
        where_clauses.append("ps.gp > 0 AND (CAST(ps.p AS REAL) / ps.gp) >= ?")
        params.append(min_ppg)
    # Intelligence-based filters
    if overall_grade:
        where_clauses.append("pi.overall_grade = ?")
        params.append(overall_grade)
    if archetype:
        where_clauses.append("pi.archetype = ?")
        params.append(archetype)

    where_sql = " AND ".join(where_clauses)

    # Sorting
    sort_map = {
        "name": "p.last_name", "last_name": "p.last_name", "first_name": "p.first_name",
        "gp": "COALESCE(ps.gp, 0)", "g": "COALESCE(ps.g, 0)", "goals": "COALESCE(ps.g, 0)",
        "a": "COALESCE(ps.a, 0)", "assists": "COALESCE(ps.a, 0)",
        "p": "COALESCE(ps.p, 0)", "points": "COALESCE(ps.p, 0)",
        "ppg": "CASE WHEN COALESCE(ps.gp, 0) > 0 THEN CAST(ps.p AS REAL) / ps.gp ELSE 0 END",
        "grade": "pi.overall_grade", "overall_grade": "pi.overall_grade",
        "team": "p.current_team", "position": "p.position",
    }
    order_col = sort_map.get(sort_by, "p.last_name")
    order_dir_sql = "DESC" if sort_dir == "desc" else "ASC"
    order_sql = f" ORDER BY {order_col} {order_dir_sql}, p.last_name ASC"

    query = f"""
        SELECT p.id, p.first_name, p.last_name, p.position, p.current_team, p.current_league,
               p.image_url, p.archetype, p.commitment_status, p.roster_status, p.age_group, p.birth_year, p.dob, p.tags,
               pi.overall_grade, pi.offensive_grade, pi.defensive_grade,
               pi.skating_grade, pi.hockey_iq_grade, pi.compete_grade,
               pi.archetype AS intel_archetype, pi.archetype_confidence,
               COALESCE(ps.gp, 0) AS stat_gp,
               COALESCE(ps.g, 0) AS stat_g,
               COALESCE(ps.a, 0) AS stat_a,
               COALESCE(ps.p, 0) AS stat_p,
               COALESCE(ps.plus_minus, 0) AS stat_plus_minus,
               COALESCE(ps.pim, 0) AS stat_pim,
               COALESCE(ps.sog, 0) AS stat_sog,
               ps.shooting_pct AS stat_shooting_pct
        FROM players p
        LEFT JOIN (
            SELECT player_id, overall_grade, offensive_grade, defensive_grade,
                   skating_grade, hockey_iq_grade, compete_grade, archetype, archetype_confidence
            FROM player_intelligence
            WHERE (player_id, version) IN (
                SELECT player_id, MAX(version) FROM player_intelligence GROUP BY player_id
            )
        ) pi ON p.id = pi.player_id
        LEFT JOIN (
            SELECT player_id,
                   SUM(gp) AS gp, SUM(g) AS g, SUM(a) AS a, SUM(p) AS p,
                   SUM(plus_minus) AS plus_minus, SUM(pim) AS pim,
                   SUM(COALESCE(sog, 0)) AS sog,
                   CASE WHEN SUM(COALESCE(sog, 0)) > 0 THEN CAST(SUM(g) AS REAL) / SUM(sog) * 100 ELSE NULL END AS shooting_pct
            FROM player_stats
            WHERE stat_type = 'season'
            GROUP BY player_id
        ) ps ON p.id = ps.player_id
        WHERE {where_sql}
        {order_sql}
        LIMIT ? OFFSET ?
    """
    params.extend([limit, skip])

    rows = conn.execute(query, params).fetchall()

    # Get league stats for percentile computation (skaters with 5+ GP)
    league_stats_rows = conn.execute("""
        SELECT ps.player_id, p.position,
               ps.gp, ps.g, ps.a, ps.p, ps.plus_minus, ps.pim,
               COALESCE(ps.sog, 0) AS sog,
               CASE WHEN COALESCE(ps.sog, 0) > 0 THEN CAST(ps.g AS REAL) / ps.sog * 100 ELSE NULL END AS shooting_pct
        FROM player_stats ps
        JOIN players p ON ps.player_id = p.id
        WHERE p.org_id = ? AND ps.stat_type = 'season' AND ps.gp >= 5
    """, (org_id,)).fetchall()
    league_stats = [dict(r) for r in league_stats_rows]

    # Build response
    results = []
    for row in rows:
        r = dict(row)
        card = {
            "id": r["id"],
            "first_name": r["first_name"],
            "last_name": r["last_name"],
            "position": r["position"],
            "current_team": r["current_team"],
            "current_league": r["current_league"],
            "image_url": r["image_url"],
            "archetype": r["intel_archetype"] or r["archetype"],
            "commitment_status": r["commitment_status"],
            "roster_status": r.get("roster_status"),
            "age_group": r["age_group"],
            "birth_year": r["birth_year"],
            "overall_grade": r["overall_grade"],
            "offensive_grade": r["offensive_grade"],
            "defensive_grade": r["defensive_grade"],
            "skating_grade": r["skating_grade"],
            "hockey_iq_grade": r["hockey_iq_grade"],
            "compete_grade": r["compete_grade"],
            "archetype_confidence": r["archetype_confidence"],
            "tags": json.loads(r["tags"]) if r.get("tags") else [],
            "gp": r["stat_gp"],
            "g": r["stat_g"],
            "a": r["stat_a"],
            "p": r["stat_p"],
            "metrics": None,
        }

        # Compute ProspectX metrics if player has enough stats
        if r["stat_gp"] >= 5:
            try:
                player_stats_dict = {
                    "gp": r["stat_gp"], "g": r["stat_g"], "a": r["stat_a"], "p": r["stat_p"],
                    "plus_minus": r["stat_plus_minus"], "pim": r["stat_pim"],
                    "sog": r["stat_sog"], "shooting_pct": r["stat_shooting_pct"],
                }
                metrics = _compute_prospectx_indices(player_stats_dict, r["position"], league_stats)
                card["metrics"] = {k: v["value"] for k, v in metrics.items()}
            except Exception:
                pass

        results.append(card)

    conn.close()
    return results


# ── Drills & Practice Plans Models ────────────────────────────

class DrillCreate(BaseModel):
    name: str = Field(..., min_length=1, max_length=200)
    category: str
    description: str = Field(..., min_length=1)
    coaching_points: Optional[str] = None
    setup: Optional[str] = None
    duration_minutes: int = 10
    players_needed: int = 0
    ice_surface: str = "full"
    equipment: Optional[str] = None
    age_levels: List[str] = []
    tags: List[str] = []
    diagram_url: Optional[str] = None
    skill_focus: Optional[str] = None
    intensity: str = "medium"
    concept_id: Optional[str] = None
    diagram_data: Optional[dict] = None

class DrillUpdate(BaseModel):
    name: Optional[str] = None
    category: Optional[str] = None
    description: Optional[str] = None
    coaching_points: Optional[str] = None
    setup: Optional[str] = None
    duration_minutes: Optional[int] = None
    players_needed: Optional[int] = None
    ice_surface: Optional[str] = None
    equipment: Optional[str] = None
    age_levels: Optional[List[str]] = None
    tags: Optional[List[str]] = None
    diagram_url: Optional[str] = None
    skill_focus: Optional[str] = None
    intensity: Optional[str] = None
    concept_id: Optional[str] = None
    diagram_data: Optional[dict] = None

class PracticePlanCreate(BaseModel):
    team_name: Optional[str] = None
    title: str = Field(..., min_length=1, max_length=300)
    age_level: Optional[str] = None
    duration_minutes: int = 90
    focus_areas: List[str] = []
    plan_data: Optional[dict] = None
    notes: Optional[str] = None

class PracticePlanUpdate(BaseModel):
    title: Optional[str] = None
    team_name: Optional[str] = None
    age_level: Optional[str] = None
    duration_minutes: Optional[int] = None
    focus_areas: Optional[List[str]] = None
    plan_data: Optional[dict] = None
    notes: Optional[str] = None
    status: Optional[str] = None

class PracticePlanDrillAdd(BaseModel):
    drill_id: str
    phase: str  # warm_up, skill_work, systems, scrimmage, conditioning, cool_down
    sequence_order: int = 0
    duration_minutes: int = 10
    coaching_notes: Optional[str] = None

class PracticePlanDrillUpdate(BaseModel):
    phase: Optional[str] = None
    sequence_order: Optional[int] = None
    duration_minutes: Optional[int] = None
    coaching_notes: Optional[str] = None

class PracticePlanGenerateRequest(BaseModel):
    team_name: str
    duration_minutes: int = 90
    focus_areas: List[str] = []
    age_level: str = "JUNIOR_COLLEGE_PRO"
    notes: Optional[str] = None

# ── Player Search (autocomplete) ─────────────────────────────

@app.get("/players/search")
async def search_players_autocomplete(
    q: str = Query(..., min_length=2, max_length=100),
    limit: int = Query(default=8, ge=1, le=20),
    token_data: dict = Depends(verify_token),
):
    """Lightweight player search for nav autocomplete."""
    org_id = token_data["org_id"]
    conn = get_db()
    pattern = f"%{q}%"
    rows = conn.execute(
        """SELECT id, first_name, last_name, current_team, position, jersey_number
           FROM players
           WHERE org_id = ?
             AND (is_deleted = 0 OR is_deleted IS NULL)
             AND (first_name LIKE ? OR last_name LIKE ? OR (first_name || ' ' || last_name) LIKE ?)
           GROUP BY id
           ORDER BY last_name, first_name
           LIMIT ?""",
        (org_id, pattern, pattern, pattern, limit),
    ).fetchall()
    conn.close()
    return {"results": [dict(r) for r in rows]}

# ── Saved Searches ────────────────────────────────────────────

class SavedSearchCreate(BaseModel):
    name: str
    filters: dict

@app.post("/players/search/save", status_code=201)
async def save_search(body: SavedSearchCreate, token_data: dict = Depends(verify_token)):
    """Save a search preset."""
    org_id = token_data["org_id"]
    user_id = token_data["user_id"]
    search_id = gen_id()
    conn = get_db()
    conn.execute(
        "INSERT INTO saved_searches (id, org_id, user_id, name, filters) VALUES (?, ?, ?, ?, ?)",
        (search_id, org_id, user_id, body.name, json.dumps(body.filters)),
    )
    conn.commit()
    row = conn.execute("SELECT * FROM saved_searches WHERE id = ?", (search_id,)).fetchone()
    conn.close()
    r = dict(row)
    r["filters"] = json.loads(r["filters"])
    return r

@app.get("/players/search/saved")
async def list_saved_searches(token_data: dict = Depends(verify_token)):
    """List saved search presets for the current user."""
    user_id = token_data["user_id"]
    conn = get_db()
    rows = conn.execute(
        "SELECT * FROM saved_searches WHERE user_id = ? ORDER BY created_at DESC", (user_id,)
    ).fetchall()
    conn.close()
    results = []
    for row in rows:
        r = dict(row)
        r["filters"] = json.loads(r["filters"])
        results.append(r)
    return results

@app.delete("/players/search/saved/{search_id}")
async def delete_saved_search(search_id: str, token_data: dict = Depends(verify_token)):
    """Delete a saved search preset."""
    user_id = token_data["user_id"]
    conn = get_db()
    row = conn.execute("SELECT * FROM saved_searches WHERE id = ? AND user_id = ?", (search_id, user_id)).fetchone()
    if not row:
        conn.close()
        raise HTTPException(status_code=404, detail="Saved search not found")
    conn.execute("DELETE FROM saved_searches WHERE id = ?", (search_id,))
    conn.commit()
    conn.close()
    return {"ok": True}


@app.post("/players", response_model=PlayerResponse, status_code=201)
async def create_player(player: PlayerCreate, token_data: dict = Depends(verify_token)):
    org_id = token_data["org_id"]
    player_id = gen_id()
    now = now_iso()
    conn = get_db()

    # Derive compartmentalization fields
    birth_year = None
    age_group = None
    draft_eligible_year = None
    if player.dob:
        try:
            birth_year = int(player.dob[:4])
            age_group = _get_age_group(birth_year)
            draft_eligible_year = birth_year + 18
        except (ValueError, IndexError):
            pass
    league_tier = _get_league_tier(player.current_league)

    user_id = token_data["user_id"]
    conn.execute("""
        INSERT INTO players (id, org_id, first_name, last_name, dob, position, shoots, height_cm, weight_kg,
                             current_team, current_league, passports, notes, tags, archetype, image_url,
                             birth_year, age_group, draft_eligible_year, league_tier, commitment_status,
                             elite_prospects_url, created_by, created_at, updated_at)
        VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
    """, (
        player_id, org_id, player.first_name, player.last_name, player.dob,
        player.position.upper(), player.shoots, player.height_cm, player.weight_kg,
        player.current_team, player.current_league,
        json.dumps(player.passports or []), player.notes, json.dumps(player.tags or []),
        player.archetype, player.image_url,
        birth_year, age_group, draft_eligible_year, league_tier,
        player.commitment_status or "Uncommitted",
        player.elite_prospects_url, user_id, now, now,
    ))
    conn.commit()

    row = conn.execute("SELECT * FROM players WHERE id = ?", (player_id,)).fetchone()
    conn.close()
    logger.info("Player created: %s %s (org %s)", player.first_name, player.last_name, org_id)
    return PlayerResponse(**_player_from_row(row))


@app.get("/players/deleted")
async def list_deleted_players(token_data: dict = Depends(verify_token)):
    """List soft-deleted players that are still within the 30-day recovery window."""
    org_id = token_data["org_id"]
    conn = get_db()
    rows = conn.execute("""
        SELECT id, first_name, last_name, position, current_team, current_league,
               deleted_at, deleted_reason, deleted_by,
               CAST(julianday('now') - julianday(deleted_at) AS INTEGER) AS days_since_deleted
        FROM players
        WHERE org_id = ? AND is_deleted = 1
        ORDER BY deleted_at DESC
    """, (org_id,)).fetchall()
    conn.close()
    result = []
    for r in rows:
        days = r["days_since_deleted"] or 0
        result.append({
            "id": r["id"],
            "first_name": r["first_name"],
            "last_name": r["last_name"],
            "position": r["position"],
            "current_team": r["current_team"],
            "current_league": r["current_league"],
            "deleted_at": r["deleted_at"],
            "deleted_reason": r["deleted_reason"],
            "deleted_by": r["deleted_by"],
            "days_since_deleted": days,
            "days_remaining": max(0, 30 - days),
            "can_restore": days <= 30,
        })
    return result


@app.post("/players/{player_id}/restore")
async def restore_player(player_id: str, token_data: dict = Depends(verify_token)):
    """Restore a soft-deleted player within the 30-day recovery window."""
    org_id = token_data["org_id"]
    conn = get_db()
    player = conn.execute("""
        SELECT id, first_name, last_name, deleted_at,
               CAST(julianday('now') - julianday(deleted_at) AS INTEGER) AS days_since
        FROM players WHERE id = ? AND org_id = ? AND is_deleted = 1
    """, (player_id, org_id)).fetchone()
    if not player:
        conn.close()
        raise HTTPException(status_code=404, detail="Deleted player not found")
    if (player["days_since"] or 0) > 30:
        conn.close()
        raise HTTPException(status_code=400, detail="Recovery window expired (30 days)")
    conn.execute("""
        UPDATE players SET is_deleted = 0, deleted_at = NULL, deleted_reason = NULL, deleted_by = NULL
        WHERE id = ?
    """, (player_id,))
    conn.commit()
    conn.close()
    return {
        "status": "restored",
        "player_id": player_id,
        "player_name": f"{player['first_name']} {player['last_name']}",
    }


@app.get("/players/{player_id}", response_model=PlayerResponse)
async def get_player(player_id: str, token_data: dict = Depends(verify_token)):
    org_id = token_data["org_id"]
    conn = get_db()
    row = conn.execute("SELECT * FROM players WHERE id = ? AND org_id = ? AND (is_deleted = 0 OR is_deleted IS NULL)", (player_id, org_id)).fetchone()
    conn.close()
    if not row:
        raise HTTPException(status_code=404, detail="Player not found")
    return PlayerResponse(**_player_from_row(row))


@app.patch("/players/{player_id}")
async def patch_player(
    player_id: str,
    updates: dict = Body(...),
    token_data: dict = Depends(verify_token),
):
    """Partially update a player — only send fields you want to change.
    Useful for transfers, position changes, etc."""
    org_id = token_data["org_id"]
    conn = get_db()
    existing = conn.execute("SELECT * FROM players WHERE id = ? AND org_id = ?", (player_id, org_id)).fetchone()
    if not existing:
        conn.close()
        raise HTTPException(status_code=404, detail="Player not found")

    allowed = {"first_name", "last_name", "dob", "position", "shoots", "height_cm", "weight_kg",
               "current_team", "current_league", "notes", "archetype", "image_url", "commitment_status",
               "roster_status", "jersey_number"}
    sets = []
    params = []
    for field, value in updates.items():
        if field in allowed:
            sets.append(f"{field} = ?")
            params.append(value)
    if not sets:
        conn.close()
        raise HTTPException(status_code=400, detail="No valid fields to update")

    # Auto-derive compartmentalization if DOB or league changed
    if "dob" in updates and updates["dob"]:
        try:
            by = int(updates["dob"][:4])
            sets.extend(["birth_year = ?", "age_group = ?", "draft_eligible_year = ?"])
            params.extend([by, _get_age_group(by), by + 18])
        except (ValueError, IndexError):
            pass
    if "current_league" in updates:
        sets.append("league_tier = ?")
        params.append(_get_league_tier(updates["current_league"]))

    sets.append("updated_at = ?")
    params.append(now_iso())
    params.append(player_id)

    conn.execute(f"UPDATE players SET {', '.join(sets)} WHERE id = ?", params)
    conn.commit()
    row = conn.execute("SELECT * FROM players WHERE id = ?", (player_id,)).fetchone()
    conn.close()
    return dict(row)


@app.put("/players/{player_id}", response_model=PlayerResponse)
async def update_player(player_id: str, player: PlayerCreate, token_data: dict = Depends(verify_token)):
    org_id = token_data["org_id"]
    conn = get_db()
    existing = conn.execute("SELECT id FROM players WHERE id = ? AND org_id = ?", (player_id, org_id)).fetchone()
    if not existing:
        conn.close()
        raise HTTPException(status_code=404, detail="Player not found")

    # Derive compartmentalization fields
    birth_year = None
    age_group = None
    draft_eligible_year = None
    if player.dob:
        try:
            birth_year = int(player.dob[:4])
            age_group = _get_age_group(birth_year)
            draft_eligible_year = birth_year + 18
        except (ValueError, IndexError):
            pass
    league_tier = _get_league_tier(player.current_league)

    conn.execute("""
        UPDATE players SET first_name=?, last_name=?, dob=?, position=?, shoots=?, height_cm=?, weight_kg=?,
                          current_team=?, current_league=?, passports=?, notes=?, tags=?, archetype=?, image_url=?,
                          birth_year=?, age_group=?, draft_eligible_year=?, league_tier=?,
                          commitment_status=?, elite_prospects_url=?, updated_at=?
        WHERE id = ? AND org_id = ?
    """, (
        player.first_name, player.last_name, player.dob, player.position.upper(), player.shoots,
        player.height_cm, player.weight_kg, player.current_team, player.current_league,
        json.dumps(player.passports or []), player.notes, json.dumps(player.tags or []),
        player.archetype, player.image_url,
        birth_year, age_group, draft_eligible_year, league_tier,
        player.commitment_status or "Uncommitted",
        player.elite_prospects_url, now_iso(), player_id, org_id,
    ))
    conn.commit()
    row = conn.execute("SELECT * FROM players WHERE id = ?", (player_id,)).fetchone()
    conn.close()
    return PlayerResponse(**_player_from_row(row))


@app.delete("/players/{player_id}")
async def delete_player(player_id: str, token_data: dict = Depends(verify_token)):
    org_id = token_data["org_id"]
    conn = get_db()
    result = conn.execute("DELETE FROM players WHERE id = ? AND org_id = ?", (player_id, org_id))
    conn.commit()
    conn.close()
    if result.rowcount == 0:
        raise HTTPException(status_code=404, detail="Player not found")
    return {"detail": "Player deleted"}


# ── Player Image Upload ────────────────────────────────────
@app.post("/players/{player_id}/image")
async def upload_player_image(
    player_id: str,
    file: UploadFile = File(...),
    token_data: dict = Depends(verify_token),
):
    org_id = token_data["org_id"]
    conn = get_db()
    row = conn.execute("SELECT id FROM players WHERE id = ? AND org_id = ?", (player_id, org_id)).fetchone()
    if not row:
        conn.close()
        raise HTTPException(status_code=404, detail="Player not found")

    # Validate file type
    allowed_types = {"image/jpeg", "image/png", "image/webp", "image/gif"}
    if file.content_type not in allowed_types:
        conn.close()
        raise HTTPException(status_code=400, detail=f"Invalid image type. Allowed: JPEG, PNG, WebP, GIF")

    # Generate a unique filename
    ext = file.filename.rsplit(".", 1)[-1].lower() if "." in (file.filename or "") else "jpg"
    if ext not in ("jpg", "jpeg", "png", "webp", "gif"):
        ext = "jpg"
    filename = f"{player_id}.{ext}"
    filepath = os.path.join(_IMAGES_DIR, filename)

    # Remove any existing image for this player (different extension)
    for old_ext in ("jpg", "jpeg", "png", "webp", "gif"):
        old_path = os.path.join(_IMAGES_DIR, f"{player_id}.{old_ext}")
        if os.path.exists(old_path) and old_path != filepath:
            os.remove(old_path)

    # Save the file
    contents = await file.read()
    with open(filepath, "wb") as f:
        f.write(contents)

    # Update the player record with the image URL
    image_url = f"/uploads/{filename}"
    conn.execute(
        "UPDATE players SET image_url = ?, updated_at = ? WHERE id = ? AND org_id = ?",
        (image_url, now_iso(), player_id, org_id),
    )
    conn.commit()
    conn.close()

    logger.info("Player image uploaded: %s → %s", player_id, filename)
    return {"image_url": image_url}


@app.delete("/players/{player_id}/image")
async def delete_player_image(
    player_id: str,
    token_data: dict = Depends(verify_token),
):
    org_id = token_data["org_id"]
    conn = get_db()
    row = conn.execute("SELECT id, image_url FROM players WHERE id = ? AND org_id = ?", (player_id, org_id)).fetchone()
    if not row:
        conn.close()
        raise HTTPException(status_code=404, detail="Player not found")

    # Delete the file from disk
    for ext in ("jpg", "jpeg", "png", "webp", "gif"):
        fpath = os.path.join(_IMAGES_DIR, f"{player_id}.{ext}")
        if os.path.exists(fpath):
            os.remove(fpath)

    conn.execute(
        "UPDATE players SET image_url = NULL, updated_at = ? WHERE id = ? AND org_id = ?",
        (now_iso(), player_id, org_id),
    )
    conn.commit()
    conn.close()
    return {"detail": "Image deleted"}


@app.post("/teams/{team_id}/logo")
async def upload_team_logo(team_id: str, file: UploadFile = File(...), token_data: dict = Depends(verify_token)):
    """Upload team logo image."""
    conn = get_db()
    org_id = token_data["org_id"]

    # P0-1: Tier check — uploads require Novice+
    _check_tier_permission(token_data["user_id"], "can_upload_files", conn)

    # P0-1: org_id scoped query — only allow uploading to own org's teams
    team = conn.execute("SELECT * FROM teams WHERE id = ? AND org_id = ?", (team_id, org_id)).fetchone()
    if not team:
        conn.close()
        raise HTTPException(status_code=404, detail="Team not found")

    # P0-1: MIME type validation — whitelist image types only
    allowed_types = {"image/jpeg", "image/png", "image/webp"}
    if file.content_type not in allowed_types:
        conn.close()
        raise HTTPException(status_code=400, detail="Invalid image type. Allowed: JPEG, PNG, WebP")

    fname = (file.filename or "logo.png").lower()
    ext = fname.rsplit(".", 1)[-1] if "." in fname else "png"
    if ext not in ("jpg", "jpeg", "png", "webp"):
        conn.close()
        raise HTTPException(status_code=400, detail="Invalid image type")

    filename = f"team_{team_id}.{ext}"
    filepath = os.path.join(_IMAGES_DIR, filename)
    content = await file.read()

    # P0-1: File size limit — 5MB max
    if len(content) > 5 * 1024 * 1024:
        conn.close()
        raise HTTPException(status_code=400, detail="File too large. Maximum size is 5MB.")

    with open(filepath, "wb") as f:
        f.write(content)

    logo_url = f"/uploads/{filename}"
    conn.execute("UPDATE teams SET logo_url = ? WHERE id = ? AND org_id = ?", (logo_url, team_id, org_id))
    conn.commit()
    conn.close()

    return {"logo_url": logo_url}


# ============================================================
# STATS ENDPOINTS
# ============================================================

@app.get("/stats/player/{player_id}", response_model=List[StatsResponse])
async def get_player_stats(player_id: str, token_data: dict = Depends(verify_token)):
    conn = get_db()
    rows = conn.execute(
        "SELECT * FROM player_stats WHERE player_id = ? ORDER BY created_at DESC",
        (player_id,)
    ).fetchall()
    conn.close()

    results = []
    for r in rows:
        d = dict(r)
        if d.get("microstats") and isinstance(d["microstats"], str):
            try:
                d["microstats"] = json.loads(d["microstats"])
            except Exception:
                d["microstats"] = None
        results.append(StatsResponse(**d))
    return results


@app.get("/stats/goalie/{player_id}")
async def get_goalie_stats(player_id: str, token_data: dict = Depends(verify_token)):
    """Get goalie stats for a player."""
    conn = get_db()
    rows = conn.execute(
        "SELECT * FROM goalie_stats WHERE player_id = ? ORDER BY season DESC, created_at DESC",
        (player_id,)
    ).fetchall()
    conn.close()
    results = []
    for r in rows:
        d = dict(r)
        if d.get("extended_stats") and isinstance(d["extended_stats"], str):
            try:
                d["extended_stats"] = json.loads(d["extended_stats"])
            except Exception:
                pass
        results.append(d)
    return results


@app.get("/stats/team/{team_name}")
async def get_team_stats(team_name: str, token_data: dict = Depends(verify_token)):
    """Get team-level stats."""
    org_id = token_data["org_id"]
    decoded = team_name.replace("%20", " ")
    conn = get_db()
    rows = conn.execute(
        "SELECT * FROM team_stats WHERE org_id = ? AND LOWER(team_name) = LOWER(?) ORDER BY season DESC",
        (org_id, decoded)
    ).fetchall()
    conn.close()
    results = []
    for r in rows:
        d = dict(r)
        if d.get("extended_stats") and isinstance(d["extended_stats"], str):
            try:
                d["extended_stats"] = json.loads(d["extended_stats"])
            except Exception:
                pass
        results.append(d)
    return results


@app.get("/stats/team/{team_name}/lines")
async def get_team_lines(
    team_name: str,
    line_type: Optional[str] = Query(None),
    data_source: Optional[str] = Query(None),
    season: Optional[str] = Query(None),
    token_data: dict = Depends(verify_token),
):
    """Get line combinations for a team."""
    org_id = token_data["org_id"]
    decoded = team_name.replace("%20", " ")
    conn = get_db()
    query = "SELECT * FROM line_combinations WHERE org_id = ? AND LOWER(team_name) = LOWER(?)"
    params: list = [org_id, decoded]
    if line_type:
        query += " AND line_type = ?"
        params.append(line_type)
    if data_source:
        query += " AND data_source = ?"
        params.append(data_source)
    if season:
        query += " AND season = ?"
        params.append(season)
    query += " ORDER BY line_order ASC, toi_seconds DESC"
    rows = conn.execute(query, params).fetchall()
    conn.close()
    results = []
    for r in rows:
        d = dict(r)
        for jf in ("player_refs", "extended_stats"):
            if d.get(jf) and isinstance(d[jf], str):
                try:
                    d[jf] = json.loads(d[jf])
                except Exception:
                    pass
        results.append(d)
    return results


# ── Stats Progression / Game Log / Recent Form ───────────────

@app.get("/stats/player/{player_id}/progression")
async def get_player_progression(player_id: str, token_data: dict = Depends(verify_token)):
    """
    Season-over-season progression from player_stats_history.
    Returns the most recent snapshot per season, ordered chronologically.
    Falls back to player_stats season rows if no history rows exist.
    """
    conn = get_db()
    try:
        # Try player_stats_history first (append-only snapshots)
        history_rows = conn.execute("""
            SELECT psh.*
            FROM player_stats_history psh
            INNER JOIN (
                SELECT season, MAX(date_recorded) as max_date
                FROM player_stats_history
                WHERE player_id = ?
                GROUP BY season
            ) latest ON psh.season = latest.season AND psh.date_recorded = latest.max_date
            WHERE psh.player_id = ?
            ORDER BY psh.season ASC
        """, (player_id, player_id)).fetchall()

        if history_rows:
            seasons = []
            for r in history_rows:
                d = dict(r)
                gp = d.get("gp", 0) or 0
                pts = d.get("p", 0) or 0
                d["ppg_rate"] = round(pts / gp, 2) if gp > 0 else 0.0
                d["gpg_rate"] = round((d.get("g", 0) or 0) / gp, 2) if gp > 0 else 0.0
                d["apg_rate"] = round((d.get("a", 0) or 0) / gp, 2) if gp > 0 else 0.0
                seasons.append(d)
        else:
            # Fallback to player_stats season rows
            fallback = conn.execute("""
                SELECT id, player_id, season, gp, g, a, p, plus_minus, pim,
                    shots, shooting_pct, data_source
                FROM player_stats
                WHERE player_id = ? AND stat_type = 'season'
                ORDER BY season ASC
            """, (player_id,)).fetchall()
            seasons = []
            for r in fallback:
                d = dict(r)
                gp = d.get("gp", 0) or 0
                pts = d.get("p", 0) or 0
                d["ppg_rate"] = round(pts / gp, 2) if gp > 0 else 0.0
                d["gpg_rate"] = round((d.get("g", 0) or 0) / gp, 2) if gp > 0 else 0.0
                d["apg_rate"] = round((d.get("a", 0) or 0) / gp, 2) if gp > 0 else 0.0
                seasons.append(d)

        # Compute YoY trend
        trend = "insufficient_data"
        yoy_delta = {}
        if len(seasons) >= 2:
            prev = seasons[-2]
            curr = seasons[-1]
            ppg_prev = prev.get("ppg_rate", 0.0)
            ppg_curr = curr.get("ppg_rate", 0.0)
            delta_ppg = round(ppg_curr - ppg_prev, 2)
            delta_p = (curr.get("p", 0) or 0) - (prev.get("p", 0) or 0)
            delta_g = (curr.get("g", 0) or 0) - (prev.get("g", 0) or 0)
            delta_a = (curr.get("a", 0) or 0) - (prev.get("a", 0) or 0)
            yoy_delta = {"p": delta_p, "g": delta_g, "a": delta_a, "ppg_rate": delta_ppg}
            if delta_ppg > 0.1:
                trend = "improving"
            elif delta_ppg < -0.1:
                trend = "declining"
            else:
                trend = "stable"

        return {
            "seasons": seasons,
            "trend": trend,
            "yoy_delta": yoy_delta,
        }
    finally:
        conn.close()


@app.get("/stats/player/{player_id}/games")
async def get_player_game_stats(player_id: str,
                                 season: Optional[str] = Query(None),
                                 limit: int = Query(50, ge=1, le=200),
                                 offset: int = Query(0, ge=0),
                                 token_data: dict = Depends(verify_token)):
    """
    Game-by-game stats from player_game_stats (HT-sourced).
    Falls back to player_stats stat_type='game' rows if no HT game stats exist.
    """
    conn = get_db()
    try:
        # Try player_game_stats first (HockeyTech-sourced)
        query = "SELECT * FROM player_game_stats WHERE player_id = ?"
        params: list = [player_id]
        if season:
            query += " AND season = ?"
            params.append(season)
        query += " ORDER BY game_date DESC LIMIT ? OFFSET ?"
        params.extend([limit, offset])

        rows = conn.execute(query, params).fetchall()

        if rows:
            # Also get total count for pagination
            count_q = "SELECT COUNT(*) as cnt FROM player_game_stats WHERE player_id = ?"
            count_p: list = [player_id]
            if season:
                count_q += " AND season = ?"
                count_p.append(season)
            total = conn.execute(count_q, count_p).fetchone()["cnt"]

            games = [dict(r) for r in rows]
            return {
                "games": games,
                "total": total,
                "source": "hockeytech",
            }

        # Fallback to player_stats game rows (InStat-sourced)
        fb_query = "SELECT * FROM player_stats WHERE player_id = ? AND stat_type = 'game'"
        fb_params: list = [player_id]
        if season:
            fb_query += " AND season = ?"
            fb_params.append(season)
        fb_query += " ORDER BY created_at DESC LIMIT ? OFFSET ?"
        fb_params.extend([limit, offset])

        fb_rows = conn.execute(fb_query, fb_params).fetchall()
        count_q2 = "SELECT COUNT(*) as cnt FROM player_stats WHERE player_id = ? AND stat_type = 'game'"
        count_p2: list = [player_id]
        if season:
            count_q2 += " AND season = ?"
            count_p2.append(season)
        total2 = conn.execute(count_q2, count_p2).fetchone()["cnt"]

        games = []
        for r in fb_rows:
            d = dict(r)
            if d.get("microstats") and isinstance(d["microstats"], str):
                try:
                    d["microstats"] = json.loads(d["microstats"])
                except Exception:
                    d["microstats"] = None
            games.append(d)

        return {
            "games": games,
            "total": total2,
            "source": "instat" if fb_rows else "none",
        }
    finally:
        conn.close()


@app.get("/stats/player/{player_id}/recent-form")
async def get_player_recent_form(player_id: str,
                                  last_n: int = Query(5, ge=1, le=20),
                                  token_data: dict = Depends(verify_token)):
    """
    Compute recent form summary from last N games.
    Includes per-game data, totals, averages, and streak info.
    """
    conn = get_db()
    try:
        # Try player_game_stats first
        rows = conn.execute("""
            SELECT * FROM player_game_stats
            WHERE player_id = ?
            ORDER BY game_date DESC
            LIMIT ?
        """, (player_id, last_n)).fetchall()

        source = "hockeytech"

        if not rows:
            # Fallback to player_stats game rows
            rows = conn.execute("""
                SELECT * FROM player_stats
                WHERE player_id = ? AND stat_type = 'game'
                ORDER BY created_at DESC
                LIMIT ?
            """, (player_id, last_n)).fetchall()
            source = "instat" if rows else "none"

        if not rows:
            return {
                "last_n_games": last_n,
                "games_found": 0,
                "games": [],
                "totals": {"g": 0, "a": 0, "p": 0, "pim": 0, "shots": 0, "plus_minus": 0},
                "averages": {"gpg": 0.0, "apg": 0.0, "ppg": 0.0},
                "streak": "No game data available",
                "source": "none",
            }

        games = []
        for r in rows:
            d = dict(r)
            if d.get("microstats") and isinstance(d["microstats"], str):
                try:
                    d["microstats"] = json.loads(d["microstats"])
                except Exception:
                    d["microstats"] = None
            games.append(d)

        # Compute totals
        n = len(games)
        total_g = sum(g.get("goals", g.get("g", 0)) or 0 for g in games)
        total_a = sum(g.get("assists", g.get("a", 0)) or 0 for g in games)
        total_p = sum(g.get("points", g.get("p", 0)) or 0 for g in games)
        total_pim = sum(g.get("pim", 0) or 0 for g in games)
        total_shots = sum(g.get("shots", 0) or 0 for g in games)
        total_pm = sum(g.get("plus_minus", 0) or 0 for g in games)

        totals = {"g": total_g, "a": total_a, "p": total_p,
                  "pim": total_pim, "shots": total_shots, "plus_minus": total_pm}
        averages = {
            "gpg": round(total_g / n, 2),
            "apg": round(total_a / n, 2),
            "ppg": round(total_p / n, 2),
        }

        # Compute point streak (consecutive games with at least 1 point, most recent first)
        streak_count = 0
        for g in games:
            pts = (g.get("points", g.get("p", 0)) or 0)
            if pts > 0:
                streak_count += 1
            else:
                break

        if streak_count == 0:
            # Check for pointless streak
            pointless_count = 0
            for g in games:
                pts = (g.get("points", g.get("p", 0)) or 0)
                if pts == 0:
                    pointless_count += 1
                else:
                    break
            streak = f"{pointless_count}-game pointless streak" if pointless_count > 1 else "No active streak"
        elif streak_count == 1:
            streak = "Point in last game"
        else:
            streak = f"{streak_count}-game point streak"

        # Compute goal streak
        goal_streak = 0
        for g in games:
            goals = (g.get("goals", g.get("g", 0)) or 0)
            if goals > 0:
                goal_streak += 1
            else:
                break
        goal_streak_str = f"{goal_streak}-game goal streak" if goal_streak >= 2 else None

        return {
            "last_n_games": last_n,
            "games_found": n,
            "games": games,
            "totals": totals,
            "averages": averages,
            "streak": streak,
            "goal_streak": goal_streak_str,
            "source": source,
        }
    finally:
        conn.close()


# ── Line Combinations CRUD ───────────────────────────────────

def _line_row_to_dict(row) -> dict:
    """Convert a line_combinations DB row to a JSON-safe dict."""
    d = dict(row)
    for jf in ("player_refs", "extended_stats"):
        if d.get(jf) and isinstance(d[jf], str):
            try:
                d[jf] = json.loads(d[jf])
            except Exception:
                pass
    return d


@app.post("/teams/{team_name}/lines", status_code=201)
async def create_line(team_name: str, body: LineCombinationCreate, token_data: dict = Depends(verify_token)):
    """Create a new manual line combination for a team."""
    org_id = token_data["org_id"]
    line_id = gen_id()
    now = datetime.now(timezone.utc).isoformat()
    decoded = team_name.replace("%20", " ")

    # Build player_names string from refs
    player_names = " - ".join(
        f"{p.name}".strip() for p in body.player_refs if p.name
    ) or "Empty line"

    player_refs_json = json.dumps([p.model_dump() for p in body.player_refs])

    conn = get_db()
    conn.execute("""
        INSERT INTO line_combinations (id, org_id, team_name, season, line_type, line_label, line_order,
            player_names, player_refs, data_source, created_at, updated_at)
        VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, 'manual', ?, ?)
    """, (line_id, org_id, decoded, body.season, body.line_type, body.line_label,
          body.line_order, player_names, player_refs_json, now, now))
    conn.commit()
    row = conn.execute("SELECT * FROM line_combinations WHERE id = ?", (line_id,)).fetchone()
    conn.close()
    return _line_row_to_dict(row)


@app.put("/lines/{line_id}")
async def update_line(line_id: str, body: LineCombinationUpdate, token_data: dict = Depends(verify_token)):
    """Update a line combination's player assignments or ordering."""
    org_id = token_data["org_id"]
    conn = get_db()
    existing = conn.execute(
        "SELECT * FROM line_combinations WHERE id = ? AND org_id = ?", (line_id, org_id)
    ).fetchone()
    if not existing:
        conn.close()
        raise HTTPException(status_code=404, detail="Line not found")

    now = datetime.now(timezone.utc).isoformat()
    updates = ["updated_at = ?"]
    params: list = [now]

    if body.player_refs is not None:
        player_names = " - ".join(
            f"{p.name}".strip() for p in body.player_refs if p.name
        ) or "Empty line"
        updates.append("player_names = ?")
        params.append(player_names)
        updates.append("player_refs = ?")
        params.append(json.dumps([p.model_dump() for p in body.player_refs]))

    if body.line_label is not None:
        updates.append("line_label = ?")
        params.append(body.line_label)

    if body.line_order is not None:
        updates.append("line_order = ?")
        params.append(body.line_order)

    params.extend([line_id, org_id])
    conn.execute(
        f"UPDATE line_combinations SET {', '.join(updates)} WHERE id = ? AND org_id = ?",
        params,
    )
    conn.commit()
    row = conn.execute("SELECT * FROM line_combinations WHERE id = ?", (line_id,)).fetchone()
    conn.close()
    return _line_row_to_dict(row)


@app.delete("/lines/{line_id}")
async def delete_line(line_id: str, token_data: dict = Depends(verify_token)):
    """Delete a line combination."""
    org_id = token_data["org_id"]
    conn = get_db()
    result = conn.execute(
        "DELETE FROM line_combinations WHERE id = ? AND org_id = ?", (line_id, org_id)
    )
    conn.commit()
    conn.close()
    if result.rowcount == 0:
        raise HTTPException(status_code=404, detail="Line not found")
    return {"detail": "Line deleted"}


# ============================================================
# DRILL LIBRARY — CRUD
# ============================================================

def _drill_row_to_dict(row) -> dict:
    d = dict(row)
    d["age_levels"] = json.loads(d.get("age_levels") or "[]")
    d["tags"] = json.loads(d.get("tags") or "[]")
    # Parse diagram_data from JSON string if present
    dd = d.get("diagram_data")
    if dd and isinstance(dd, str):
        try:
            d["diagram_data"] = json.loads(dd)
        except (json.JSONDecodeError, TypeError):
            d["diagram_data"] = None
    elif not dd:
        d["diagram_data"] = None
    return d


@app.get("/drills")
async def list_drills(
    category: Optional[str] = None,
    age_level: Optional[str] = None,
    tags: Optional[str] = None,
    ice_surface: Optional[str] = None,
    search: Optional[str] = None,
    intensity: Optional[str] = None,
    concept_id: Optional[str] = None,
    limit: int = 200,
    token_data: dict = Depends(verify_token),
):
    org_id = token_data["org_id"]
    conn = get_db()
    where = ["(org_id IS NULL OR org_id = ?)"]
    params: list = [org_id]

    if category:
        where.append("category = ?")
        params.append(category)
    if ice_surface:
        where.append("ice_surface = ?")
        params.append(ice_surface)
    if intensity:
        where.append("intensity = ?")
        params.append(intensity)
    if concept_id:
        where.append("concept_id = ?")
        params.append(concept_id)
    if search:
        where.append("(LOWER(name) LIKE ? OR LOWER(description) LIKE ?)")
        params.extend([f"%{search.lower()}%", f"%{search.lower()}%"])
    if age_level:
        where.append("age_levels LIKE ?")
        params.append(f'%"{age_level}"%')
    if tags:
        for tag in tags.split(","):
            tag = tag.strip()
            if tag:
                where.append("tags LIKE ?")
                params.append(f'%"{tag}"%')

    sql = f"SELECT * FROM drills WHERE {' AND '.join(where)} ORDER BY category, name LIMIT ?"
    params.append(limit)
    rows = conn.execute(sql, params).fetchall()
    conn.close()
    return [_drill_row_to_dict(r) for r in rows]


@app.get("/drills/categories")
async def drill_categories(token_data: dict = Depends(verify_token)):
    """Return available drill categories, age levels, and concept_ids."""
    org_id = token_data["org_id"]
    conn = get_db()
    cats = [r[0] for r in conn.execute(
        "SELECT DISTINCT category FROM drills WHERE (org_id IS NULL OR org_id = ?) ORDER BY category", (org_id,)
    ).fetchall()]
    concepts = [r[0] for r in conn.execute(
        "SELECT DISTINCT concept_id FROM drills WHERE concept_id IS NOT NULL AND (org_id IS NULL OR org_id = ?) ORDER BY concept_id", (org_id,)
    ).fetchall()]
    conn.close()
    return {"categories": cats, "concept_ids": concepts}


@app.get("/drills/{drill_id}")
async def get_drill(drill_id: str, token_data: dict = Depends(verify_token)):
    org_id = token_data["org_id"]
    conn = get_db()
    row = conn.execute(
        "SELECT * FROM drills WHERE id = ? AND (org_id IS NULL OR org_id = ?)", (drill_id, org_id)
    ).fetchone()
    conn.close()
    if not row:
        raise HTTPException(status_code=404, detail="Drill not found")
    return _drill_row_to_dict(row)


@app.post("/drills", status_code=201)
async def create_drill(body: DrillCreate, token_data: dict = Depends(verify_token)):
    org_id = token_data["org_id"]
    drill_id = gen_id()
    conn = get_db()
    diagram_data_str = json.dumps(body.diagram_data) if body.diagram_data else None
    conn.execute("""
        INSERT INTO drills (id, org_id, name, category, description, coaching_points, setup,
            duration_minutes, players_needed, ice_surface, equipment, age_levels, tags,
            diagram_url, skill_focus, intensity, concept_id, diagram_data)
        VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
    """, (
        drill_id, org_id, body.name, body.category, body.description,
        body.coaching_points, body.setup, body.duration_minutes, body.players_needed,
        body.ice_surface, body.equipment, json.dumps(body.age_levels), json.dumps(body.tags),
        body.diagram_url, body.skill_focus, body.intensity, body.concept_id, diagram_data_str
    ))
    conn.commit()
    row = conn.execute("SELECT * FROM drills WHERE id = ?", (drill_id,)).fetchone()
    conn.close()
    return _drill_row_to_dict(row)


@app.put("/drills/{drill_id}")
async def update_drill(drill_id: str, body: DrillUpdate, token_data: dict = Depends(verify_token)):
    org_id = token_data["org_id"]
    conn = get_db()
    existing = conn.execute(
        "SELECT * FROM drills WHERE id = ? AND org_id = ?", (drill_id, org_id)
    ).fetchone()
    if not existing:
        conn.close()
        raise HTTPException(status_code=404, detail="Drill not found or cannot edit global drills")
    updates = []
    params = []
    for field, val in body.model_dump(exclude_unset=True).items():
        if field in ("age_levels", "tags") and val is not None:
            val = json.dumps(val)
        elif field == "diagram_data" and val is not None:
            val = json.dumps(val)
        updates.append(f"{field} = ?")
        params.append(val)
    if updates:
        params.extend([drill_id, org_id])
        conn.execute(f"UPDATE drills SET {', '.join(updates)} WHERE id = ? AND org_id = ?", params)
        conn.commit()
    row = conn.execute("SELECT * FROM drills WHERE id = ?", (drill_id,)).fetchone()
    conn.close()
    return _drill_row_to_dict(row)


@app.delete("/drills/{drill_id}")
async def delete_drill(drill_id: str, token_data: dict = Depends(verify_token)):
    org_id = token_data["org_id"]
    conn = get_db()
    result = conn.execute("DELETE FROM drills WHERE id = ? AND org_id = ?", (drill_id, org_id))
    conn.commit()
    conn.close()
    if result.rowcount == 0:
        raise HTTPException(status_code=404, detail="Drill not found or cannot delete global drills")
    return {"detail": "Drill deleted"}


@app.post("/drills/{drill_id}/diagram")
async def upload_drill_diagram(
    drill_id: str,
    file: UploadFile = File(...),
    token_data: dict = Depends(verify_token),
):
    """Upload a custom drill diagram image (replaces AI-generated SVG)."""
    org_id = token_data["org_id"]
    user_id = token_data["user_id"]

    # ── Tier permission + limit checks ──
    perm_conn = get_db()
    try:
        tier_config = _check_tier_permission(user_id, "can_upload_files", perm_conn)
        _check_tier_limit(user_id, "uploads", perm_conn)
    finally:
        perm_conn.close()

    # ── File size check ──
    contents = await file.read()
    max_bytes = tier_config.get("max_file_size_mb", 5) * 1024 * 1024
    if len(contents) > max_bytes:
        raise HTTPException(status_code=413, detail={
            "error": "file_too_large",
            "max_mb": tier_config.get("max_file_size_mb", 5),
            "file_mb": round(len(contents) / (1024 * 1024), 2),
            "upgrade_url": "/pricing",
        })

    conn = get_db()
    row = conn.execute(
        "SELECT id FROM drills WHERE id = ? AND (org_id IS NULL OR org_id = ?)",
        (drill_id, org_id)
    ).fetchone()
    if not row:
        conn.close()
        raise HTTPException(status_code=404, detail="Drill not found")

    allowed_types = {"image/jpeg", "image/png", "image/webp", "image/svg+xml"}
    if file.content_type not in allowed_types:
        conn.close()
        raise HTTPException(status_code=400, detail="Invalid image type. Allowed: JPEG, PNG, WebP, SVG")

    ext = file.filename.rsplit(".", 1)[-1].lower() if "." in (file.filename or "") else "png"
    if ext not in ("jpg", "jpeg", "png", "webp", "svg"):
        ext = "png"
    filename = f"drill_{drill_id}.{ext}"
    filepath = os.path.join(_IMAGES_DIR, filename)

    # Remove any existing diagram file with different extension
    for old_ext in ("jpg", "jpeg", "png", "webp", "svg"):
        old_path = os.path.join(_IMAGES_DIR, f"drill_{drill_id}.{old_ext}")
        if os.path.exists(old_path) and old_path != filepath:
            os.remove(old_path)

    with open(filepath, "wb") as f:
        f.write(contents)

    diagram_url = f"/uploads/{filename}"
    conn.execute("UPDATE drills SET diagram_url = ? WHERE id = ?", (diagram_url, drill_id))
    conn.commit()

    # Track the upload usage
    _increment_tracking(user_id, "uploads", conn)

    conn.close()
    logger.info("Drill diagram uploaded: %s -> %s", drill_id, filename)
    return {"diagram_url": diagram_url}


@app.delete("/drills/{drill_id}/diagram")
async def delete_drill_diagram(
    drill_id: str,
    token_data: dict = Depends(verify_token),
):
    """Delete custom diagram and regenerate SVG."""
    org_id = token_data["org_id"]
    conn = get_db()
    row = conn.execute(
        "SELECT id, ice_surface, category, concept_id, description FROM drills WHERE id = ? AND (org_id IS NULL OR org_id = ?)",
        (drill_id, org_id)
    ).fetchone()
    if not row:
        conn.close()
        raise HTTPException(status_code=404, detail="Drill not found")

    # Delete existing file(s)
    for ext in ("jpg", "jpeg", "png", "webp", "svg"):
        fpath = os.path.join(_IMAGES_DIR, f"drill_{drill_id}.{ext}")
        if os.path.exists(fpath):
            os.remove(fpath)

    # Regenerate SVG
    svg_content = generate_drill_diagram(
        row["ice_surface"] or "full", row["category"] or "offensive",
        row["concept_id"], row["description"] or ""
    )
    svg_filename = f"drill_{drill_id}.svg"
    svg_path = os.path.join(_IMAGES_DIR, svg_filename)
    with open(svg_path, "w", encoding="utf-8") as f:
        f.write(svg_content)
    diagram_url = f"/uploads/{svg_filename}"
    conn.execute("UPDATE drills SET diagram_url = ? WHERE id = ?", (diagram_url, drill_id))
    conn.commit()
    conn.close()
    return {"detail": "Custom diagram removed, SVG regenerated", "diagram_url": diagram_url}


@app.put("/drills/{drill_id}/diagram/canvas")
async def save_canvas_diagram(
    drill_id: str,
    body: dict = Body(...),
    token_data: dict = Depends(verify_token),
):
    """Save interactive rink diagram data + rendered SVG from the canvas editor."""
    org_id = token_data["org_id"]
    conn = get_db()
    row = conn.execute(
        "SELECT id FROM drills WHERE id = ? AND (org_id IS NULL OR org_id = ?)",
        (drill_id, org_id)
    ).fetchone()
    if not row:
        conn.close()
        raise HTTPException(status_code=404, detail="Drill not found")

    diagram_data = body.get("diagram_data")
    svg_string = body.get("svg_string", "")

    # Store diagram_data as JSON
    diagram_data_str = json.dumps(diagram_data) if diagram_data else None
    conn.execute("UPDATE drills SET diagram_data = ? WHERE id = ?", (diagram_data_str, drill_id))

    # Write SVG to file if provided
    if svg_string:
        svg_filename = f"drill_{drill_id}.svg"
        svg_path = os.path.join(_IMAGES_DIR, svg_filename)
        # Remove any existing non-SVG diagram files
        for ext in ("jpg", "jpeg", "png", "webp"):
            old_path = os.path.join(_IMAGES_DIR, f"drill_{drill_id}.{ext}")
            if os.path.exists(old_path):
                os.remove(old_path)
        with open(svg_path, "w", encoding="utf-8") as f:
            f.write(svg_string)
        diagram_url = f"/uploads/{svg_filename}"
        conn.execute("UPDATE drills SET diagram_url = ? WHERE id = ?", (diagram_url, drill_id))

    conn.commit()
    result_row = conn.execute("SELECT * FROM drills WHERE id = ?", (drill_id,)).fetchone()
    conn.close()
    return _drill_row_to_dict(result_row)


# ============================================================
# PRACTICE PLANS — CRUD
# ============================================================

def _practice_plan_row_to_dict(row) -> dict:
    d = dict(row)
    d["focus_areas"] = json.loads(d.get("focus_areas") or "[]")
    d["plan_data"] = json.loads(d.get("plan_data") or "{}")
    return d


def _get_plan_with_drills(conn, plan_id: str, org_id: str) -> Optional[dict]:
    row = conn.execute(
        "SELECT * FROM practice_plans WHERE id = ? AND org_id = ?", (plan_id, org_id)
    ).fetchone()
    if not row:
        return None
    plan = _practice_plan_row_to_dict(row)
    drill_rows = conn.execute("""
        SELECT ppd.*, d.name as drill_name, d.category as drill_category,
               d.description as drill_description, d.coaching_points as drill_coaching_points,
               d.setup as drill_setup, d.ice_surface as drill_ice_surface,
               d.intensity as drill_intensity, d.skill_focus as drill_skill_focus,
               d.concept_id as drill_concept_id,
               d.age_levels as drill_age_levels, d.tags as drill_tags,
               d.equipment as drill_equipment,
               d.diagram_url as drill_diagram_url
        FROM practice_plan_drills ppd
        LEFT JOIN drills d ON ppd.drill_id = d.id
        WHERE ppd.practice_plan_id = ?
        ORDER BY ppd.phase, ppd.sequence_order
    """, (plan_id,)).fetchall()
    plan["drills"] = []
    for dr in drill_rows:
        dd = dict(dr)
        dd["drill_age_levels"] = json.loads(dd.get("drill_age_levels") or "[]")
        dd["drill_tags"] = json.loads(dd.get("drill_tags") or "[]")
        plan["drills"].append(dd)
    return plan


@app.get("/practice-plans")
async def list_practice_plans(
    team_name: Optional[str] = None,
    status: Optional[str] = None,
    search: Optional[str] = None,
    limit: int = 100,
    token_data: dict = Depends(verify_token),
):
    org_id = token_data["org_id"]
    conn = get_db()
    where = ["org_id = ?"]
    params: list = [org_id]
    if team_name:
        where.append("LOWER(team_name) = LOWER(?)")
        params.append(team_name)
    if status:
        where.append("status = ?")
        params.append(status)
    if search:
        where.append("(LOWER(title) LIKE ? OR LOWER(notes) LIKE ?)")
        params.extend([f"%{search.lower()}%", f"%{search.lower()}%"])
    params.append(limit)
    rows = conn.execute(
        f"SELECT * FROM practice_plans WHERE {' AND '.join(where)} ORDER BY updated_at DESC LIMIT ?", params
    ).fetchall()
    conn.close()
    return [_practice_plan_row_to_dict(r) for r in rows]


@app.get("/practice-plans/{plan_id}")
async def get_practice_plan(plan_id: str, token_data: dict = Depends(verify_token)):
    org_id = token_data["org_id"]
    conn = get_db()
    plan = _get_plan_with_drills(conn, plan_id, org_id)
    conn.close()
    if not plan:
        raise HTTPException(status_code=404, detail="Practice plan not found")
    return plan


@app.post("/practice-plans", status_code=201)
async def create_practice_plan(body: PracticePlanCreate, token_data: dict = Depends(verify_token)):
    org_id = token_data["org_id"]
    user_id = token_data["user_id"]
    plan_id = gen_id()
    now = datetime.utcnow().isoformat()
    conn = get_db()
    conn.execute("""
        INSERT INTO practice_plans (id, org_id, user_id, team_name, title, age_level,
            duration_minutes, focus_areas, plan_data, notes, status, created_at, updated_at)
        VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, 'draft', ?, ?)
    """, (
        plan_id, org_id, user_id, body.team_name, body.title, body.age_level,
        body.duration_minutes, json.dumps(body.focus_areas),
        json.dumps(body.plan_data) if body.plan_data else "{}",
        body.notes, now, now
    ))
    conn.commit()
    plan = _get_plan_with_drills(conn, plan_id, org_id)
    conn.close()
    return plan


@app.put("/practice-plans/{plan_id}")
async def update_practice_plan(plan_id: str, body: PracticePlanUpdate, token_data: dict = Depends(verify_token)):
    org_id = token_data["org_id"]
    conn = get_db()
    existing = conn.execute(
        "SELECT id FROM practice_plans WHERE id = ? AND org_id = ?", (plan_id, org_id)
    ).fetchone()
    if not existing:
        conn.close()
        raise HTTPException(status_code=404, detail="Practice plan not found")
    updates = []
    params = []
    for field, val in body.model_dump(exclude_unset=True).items():
        if field in ("focus_areas",) and val is not None:
            val = json.dumps(val)
        elif field == "plan_data" and val is not None:
            val = json.dumps(val)
        updates.append(f"{field} = ?")
        params.append(val)
    updates.append("updated_at = ?")
    params.append(datetime.utcnow().isoformat())
    params.extend([plan_id, org_id])
    conn.execute(f"UPDATE practice_plans SET {', '.join(updates)} WHERE id = ? AND org_id = ?", params)
    conn.commit()
    plan = _get_plan_with_drills(conn, plan_id, org_id)
    conn.close()
    return plan


@app.delete("/practice-plans/{plan_id}")
async def delete_practice_plan(plan_id: str, token_data: dict = Depends(verify_token)):
    org_id = token_data["org_id"]
    conn = get_db()
    conn.execute("DELETE FROM practice_plan_drills WHERE practice_plan_id = ?", (plan_id,))
    result = conn.execute("DELETE FROM practice_plans WHERE id = ? AND org_id = ?", (plan_id, org_id))
    conn.commit()
    conn.close()
    if result.rowcount == 0:
        raise HTTPException(status_code=404, detail="Practice plan not found")
    return {"detail": "Practice plan deleted"}


@app.post("/practice-plans/{plan_id}/drills", status_code=201)
async def add_drill_to_plan(plan_id: str, body: PracticePlanDrillAdd, token_data: dict = Depends(verify_token)):
    org_id = token_data["org_id"]
    conn = get_db()
    plan = conn.execute("SELECT id FROM practice_plans WHERE id = ? AND org_id = ?", (plan_id, org_id)).fetchone()
    if not plan:
        conn.close()
        raise HTTPException(status_code=404, detail="Practice plan not found")
    ppd_id = gen_id()
    conn.execute("""
        INSERT INTO practice_plan_drills (id, practice_plan_id, drill_id, phase, sequence_order, duration_minutes, coaching_notes)
        VALUES (?, ?, ?, ?, ?, ?, ?)
    """, (ppd_id, plan_id, body.drill_id, body.phase, body.sequence_order, body.duration_minutes, body.coaching_notes))
    conn.execute("UPDATE practice_plans SET updated_at = ? WHERE id = ?", (datetime.utcnow().isoformat(), plan_id))
    conn.commit()
    plan = _get_plan_with_drills(conn, plan_id, org_id)
    conn.close()
    return plan


@app.put("/practice-plans/{plan_id}/drills/{ppd_id}")
async def update_drill_in_plan(plan_id: str, ppd_id: str, body: PracticePlanDrillUpdate, token_data: dict = Depends(verify_token)):
    org_id = token_data["org_id"]
    conn = get_db()
    plan = conn.execute("SELECT id FROM practice_plans WHERE id = ? AND org_id = ?", (plan_id, org_id)).fetchone()
    if not plan:
        conn.close()
        raise HTTPException(status_code=404, detail="Practice plan not found")
    updates = []
    params = []
    for field, val in body.model_dump(exclude_unset=True).items():
        updates.append(f"{field} = ?")
        params.append(val)
    if updates:
        params.extend([ppd_id, plan_id])
        conn.execute(f"UPDATE practice_plan_drills SET {', '.join(updates)} WHERE id = ? AND practice_plan_id = ?", params)
        conn.execute("UPDATE practice_plans SET updated_at = ? WHERE id = ?", (datetime.utcnow().isoformat(), plan_id))
        conn.commit()
    plan_data = _get_plan_with_drills(conn, plan_id, org_id)
    conn.close()
    return plan_data


@app.delete("/practice-plans/{plan_id}/drills/{ppd_id}")
async def remove_drill_from_plan(plan_id: str, ppd_id: str, token_data: dict = Depends(verify_token)):
    org_id = token_data["org_id"]
    conn = get_db()
    plan = conn.execute("SELECT id FROM practice_plans WHERE id = ? AND org_id = ?", (plan_id, org_id)).fetchone()
    if not plan:
        conn.close()
        raise HTTPException(status_code=404, detail="Practice plan not found")
    conn.execute("DELETE FROM practice_plan_drills WHERE id = ? AND practice_plan_id = ?", (ppd_id, plan_id))
    conn.execute("UPDATE practice_plans SET updated_at = ? WHERE id = ?", (datetime.utcnow().isoformat(), plan_id))
    conn.commit()
    plan_data = _get_plan_with_drills(conn, plan_id, org_id)
    conn.close()
    return plan_data


# ============================================================
# AI PRACTICE PLAN GENERATION
# ============================================================

@app.post("/practice-plans/generate", status_code=201)
async def generate_practice_plan(body: PracticePlanGenerateRequest, token_data: dict = Depends(verify_token)):
    """AI-powered practice plan generation using team context and drill library."""
    org_id = token_data["org_id"]
    user_id = token_data["user_id"]

    # ── Tier limit check: practice plans ──
    limit_conn = get_db()
    try:
        _check_tier_limit(user_id, "practice_plans", limit_conn)
    finally:
        limit_conn.close()

    conn = get_db()

    # 1. Gather team context
    roster_rows = conn.execute(
        "SELECT first_name, last_name, position, shoots FROM players WHERE org_id = ? AND LOWER(current_team) = LOWER(?)",
        (org_id, body.team_name)
    ).fetchall()
    roster_summary = [f"{r['first_name']} {r['last_name']} ({r['position']}, {r['shoots'] or '?'})" for r in roster_rows]

    team_system = None
    ts_row = conn.execute(
        "SELECT * FROM team_systems WHERE org_id = ? AND LOWER(team_name) = LOWER(?)", (org_id, body.team_name)
    ).fetchone()
    if ts_row:
        team_system = dict(ts_row)

    # 2. Query matching drills
    drill_where = ["(org_id IS NULL OR org_id = ?)", "age_levels LIKE ?"]
    drill_params: list = [org_id, f'%"{body.age_level}"%']
    drill_rows = conn.execute(
        f"SELECT * FROM drills WHERE {' AND '.join(drill_where)} ORDER BY category, name",
        drill_params
    ).fetchall()
    available_drills = []
    for dr in drill_rows:
        d = _drill_row_to_dict(dr)
        available_drills.append({
            "id": d["id"], "name": d["name"], "category": d["category"],
            "description": d["description"][:150], "duration_minutes": d["duration_minutes"],
            "ice_surface": d["ice_surface"], "intensity": d["intensity"],
            "skill_focus": d["skill_focus"], "concept_id": d.get("concept_id"),
            "tags": d["tags"],
        })

    # 3. Get glossary terms for context
    glossary = [dict(r) for r in conn.execute("SELECT term, category, definition FROM hockey_terms").fetchall()]

    # 4. Build prompt
    system_prompt = """You are ProspectX Practice Plan Intelligence — an elite hockey coaching assistant.

Generate a structured practice plan in JSON format. The plan should be realistic, age-appropriate, and use drills from the provided library.

RESPONSE FORMAT — return ONLY valid JSON (no markdown, no extra text):
{
  "title": "Practice Plan: [Focus] — [Team]",
  "phases": [
    {
      "phase": "warm_up",
      "phase_label": "Warm Up",
      "duration_minutes": 10,
      "drills": [
        {
          "drill_id": "<id from available drills>",
          "drill_name": "<name>",
          "duration_minutes": 8,
          "coaching_notes": "Specific coaching notes for this team/session"
        }
      ]
    },
    {"phase": "skill_work", "phase_label": "Skill Work", ...},
    {"phase": "systems", "phase_label": "Team Systems", ...},
    {"phase": "scrimmage", "phase_label": "Game Situations", ...},
    {"phase": "conditioning", "phase_label": "Conditioning", ...},
    {"phase": "cool_down", "phase_label": "Cool Down", ...}
  ],
  "coaching_summary": "2-3 sentence summary of the practice focus and goals"
}

RULES:
- Use ONLY drill_ids from the available drills list
- Total drill times should sum close to the requested duration
- Select drills that match the focus areas
- Coaching notes should reference the team's system and roster when relevant
- Warm-up: 8-12 minutes. Cool-down: 5-8 minutes.
- Balance intensity: start low, build to high, then cool down
- Consider ice surface variety — not everything needs to be full ice
- Use proper hockey terminology in coaching_notes: reference forecheck roles (F1/F2/F3), PP/PK formations (1-3-1, diamond, box), breakout patterns (standard, reverse, wheel), player roles (bumper, flank, QB, net-front), and tactical concepts by name
- Be specific and actionable: "Focus on F1 pressure closing speed on the forecheck" not "work on forechecking"
- Reference the team's system when writing notes: if they run a 1-2-2 forecheck, note how the drill connects to that system

AGE-APPROPRIATE COACHING — THIS IS CRITICAL:
- U8 (Mite): Maximum fun. Short drills (5-8 min). No systems or tactics. Focus on skating, puck handling, and games. Every drill should feel like play. Lots of small area games and races. No checking concepts. Use "fun" and "cool_down" category drills. Every player touches the puck constantly.
- U10 (Squirt): Introduce basic passing, shooting technique, and simple positional concepts. Still heavy on fun. Begin 1-on-1 concepts. Short drills (8-10 min). More structured than U8 but still game-based. Individual skill development is the priority.
- U12 (Peewee): Introduce team concepts — basic forecheck, breakout patterns, cycling. Positional awareness. Battle drills appropriate. Can handle 10-12 min drills. Begin special teams concepts (simple PP/PK). Checking fundamentals (body position, not hitting).
- U14 (Bantam): Full tactical concepts. Forecheck systems, DZ coverage, PP/PK formations. Conditioning matters. Battle drills are key. Can handle 12-15 min complex drills. This is where hockey IQ development accelerates.
- U16_U18 (Midget/AAA): Game-like situations. Advanced systems. Full special teams. High-intensity conditioning. Film-room style coaching notes. Tactical detail in every drill explanation.
- JUNIOR_COLLEGE_PRO: Elite detail. Advanced analytics references. Positional nuance. Professional-level coaching points. Complex systems integration.

Do NOT select systems/tactics drills for U8 teams. Do NOT select simple fun games for Junior/Pro teams. Match the drill complexity to the age level."""

    focus_str = ", ".join(body.focus_areas) if body.focus_areas else "general skills"
    user_prompt = f"""Generate a {body.duration_minutes}-minute practice plan for the {body.team_name}.

Age level: {body.age_level}
Focus areas: {focus_str}
{f'Additional notes: {body.notes}' if body.notes else ''}

ROSTER ({len(roster_summary)} players):
{chr(10).join(roster_summary[:25]) if roster_summary else 'No roster data available'}

{f"TEAM SYSTEM: Forecheck={team_system.get('forecheck','N/A')}, DZ={team_system.get('dz_structure','N/A')}, OZ={team_system.get('oz_setup','N/A')}, PP={team_system.get('pp_formation','N/A')}, PK={team_system.get('pk_formation','N/A')}" if team_system else "No team system configured"}

AVAILABLE DRILLS:
{json.dumps(available_drills, indent=1)}

HOCKEY TERMINOLOGY:
{json.dumps([{"term": g["term"], "category": g["category"]} for g in glossary], indent=1)}"""

    # 5. Call Claude (or mock)
    plan_data = None
    client = get_anthropic_client()
    if client:
        try:
            message = client.messages.create(
                model="claude-sonnet-4-20250514",
                max_tokens=4096,
                system=system_prompt,
                messages=[{"role": "user", "content": user_prompt}],
            )
            response_text = message.content[0].text.strip()
            # Extract JSON from response (handle possible markdown wrapping)
            if response_text.startswith("```"):
                response_text = response_text.split("```")[1]
                if response_text.startswith("json"):
                    response_text = response_text[4:]
            plan_data = json.loads(response_text)
        except Exception as e:
            logger.error("Practice plan generation error: %s", e)

    # Mock fallback
    if not plan_data:
        # Build a simple mock plan from available drills
        warmup_drills = [d for d in available_drills if d["category"] == "warm_up"][:1]
        skill_drills = [d for d in available_drills if d["category"] in ("passing", "shooting", "skating")][:2]
        system_drills = [d for d in available_drills if d["category"] in ("offensive", "defensive", "transition")][:2]
        game_drills = [d for d in available_drills if d["category"] in ("small_area_games", "battle")][:1]
        cond_drills = [d for d in available_drills if d["category"] == "conditioning"][:1]

        def _mock_phase(phase, label, drills_list, dur):
            return {
                "phase": phase, "phase_label": label, "duration_minutes": dur,
                "drills": [{"drill_id": d["id"], "drill_name": d["name"],
                            "duration_minutes": d["duration_minutes"],
                            "coaching_notes": f"Focus on {', '.join(body.focus_areas[:2]) if body.focus_areas else 'fundamentals'}."}
                           for d in drills_list]
            }
        plan_data = {
            "title": f"Practice Plan: {focus_str.title()} — {body.team_name}",
            "phases": [
                _mock_phase("warm_up", "Warm Up", warmup_drills, 10),
                _mock_phase("skill_work", "Skill Work", skill_drills, 25),
                _mock_phase("systems", "Team Systems", system_drills, 20),
                _mock_phase("scrimmage", "Game Situations", game_drills, 15),
                _mock_phase("conditioning", "Conditioning", cond_drills, 10),
                {"phase": "cool_down", "phase_label": "Cool Down", "duration_minutes": 5,
                 "drills": [{"drill_id": None, "drill_name": "Easy skate and stretch",
                             "duration_minutes": 5, "coaching_notes": "Light skate, static stretching, team huddle."}]},
            ],
            "coaching_summary": f"Practice focused on {focus_str} for the {body.team_name}. Built from the ProspectX drill library."
        }

    # 6. Create practice plan record
    plan_id = gen_id()
    now = datetime.utcnow().isoformat()
    title = plan_data.get("title", f"Practice Plan — {body.team_name}")
    conn.execute("""
        INSERT INTO practice_plans (id, org_id, user_id, team_name, title, age_level,
            duration_minutes, focus_areas, plan_data, notes, status, created_at, updated_at)
        VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, 'draft', ?, ?)
    """, (
        plan_id, org_id, user_id, body.team_name, title, body.age_level,
        body.duration_minutes, json.dumps(body.focus_areas),
        json.dumps(plan_data), body.notes, now, now
    ))

    # 7. Create junction records for referenced drills
    for phase_data in plan_data.get("phases", []):
        for i, drill_entry in enumerate(phase_data.get("drills", [])):
            drill_id = drill_entry.get("drill_id")
            if drill_id:
                conn.execute("""
                    INSERT INTO practice_plan_drills (id, practice_plan_id, drill_id, phase,
                        sequence_order, duration_minutes, coaching_notes)
                    VALUES (?, ?, ?, ?, ?, ?, ?)
                """, (
                    gen_id(), plan_id, drill_id, phase_data["phase"], i,
                    drill_entry.get("duration_minutes", 10),
                    drill_entry.get("coaching_notes")
                ))

    conn.commit()

    # Track practice plan generation usage
    _increment_tracking(user_id, "practice_plans", conn)

    result = _get_plan_with_drills(conn, plan_id, org_id)
    conn.close()
    return result


def _parse_file_to_rows(content: bytes, fname: str) -> list[dict]:
    """Parse a CSV or Excel file into a list of normalized dicts."""
    is_excel = any(fname.endswith(ext) for ext in (".xlsx", ".xls", ".xlsm"))

    if is_excel:
        try:
            import openpyxl
            wb = openpyxl.load_workbook(io.BytesIO(content), read_only=True, data_only=True)
            ws = wb.active
            if not ws:
                raise HTTPException(status_code=400, detail="Excel file has no active sheet")
            rows_iter = ws.iter_rows(values_only=True)
            raw_headers = next(rows_iter, None)
            if not raw_headers:
                raise HTTPException(status_code=400, detail="Excel file is empty")
            headers = [str(h).strip().lower().replace(" ", "_").replace("-", "_") if h else f"col_{i}" for i, h in enumerate(raw_headers)]
            all_rows = []
            for row_vals in rows_iter:
                row_dict = {}
                for j, val in enumerate(row_vals):
                    if j < len(headers):
                        row_dict[headers[j]] = str(val).strip() if val is not None else ""
                if any(v for v in row_dict.values()):
                    all_rows.append(row_dict)
            wb.close()
            return all_rows
        except HTTPException:
            raise
        except Exception as e:
            raise HTTPException(status_code=400, detail=f"Failed to parse Excel file: {str(e)}")
    else:
        try:
            text = content.decode("utf-8-sig")
        except UnicodeDecodeError:
            text = content.decode("latin-1")
        reader = csv.DictReader(io.StringIO(text))
        if reader.fieldnames:
            reader.fieldnames = [f.strip().lower().replace(" ", "_").replace("-", "_") for f in reader.fieldnames]
        return [{k.strip(): v.strip() if v else "" for k, v in row.items()} for row in reader]


def _parse_mmss_to_seconds(val: str) -> int:
    """Convert 'MM:SS' or 'HH:MM:SS' time string to total seconds."""
    if not val or val == "-" or val == "None":
        return 0
    val = val.strip()
    parts = val.split(":")
    try:
        if len(parts) == 2:
            return int(parts[0]) * 60 + int(parts[1])
        elif len(parts) == 3:
            return int(parts[0]) * 3600 + int(parts[1]) * 60 + int(parts[2])
        else:
            return int(float(val))
    except (ValueError, TypeError):
        return 0


def _detect_instat_game_log(headers: list[str]) -> bool:
    """Detect InStat per-game box score format (e.g. 'Games - Player Name.xlsx')."""
    instat_markers = {"date", "opponent", "score", "time_on_ice", "all_shifts"}
    return len(instat_markers.intersection(set(headers))) >= 3


def _detect_instat_team_stats(headers: list[str]) -> bool:
    """Detect InStat team skater stats format (e.g. 'Skaters - Team Name.xlsx')."""
    # Team stats files have a 'player' or '#' column plus standard stat headers
    has_player_col = any(h in headers for h in ["player", "#", "name", "player_name"])
    has_stat_cols = any(h in headers for h in ["goals", "assists", "points", "time_on_ice"])
    return has_player_col and has_stat_cols


def _to_int(val, default=0):
    """Convert a value to int, handling floats, dashes, and empty strings."""
    if not val or val == "-" or val == "None":
        return default
    try:
        return int(float(val))
    except (ValueError, TypeError):
        return default


def _to_float(val, default=None):
    """Convert a value to float, handling dashes and empty strings."""
    if not val or val == "-" or val == "None":
        return default
    try:
        # Strip percent signs
        cleaned = str(val).replace("%", "").strip()
        return float(cleaned) if cleaned else default
    except (ValueError, TypeError):
        return default


# ============================================================
# INSTAT ANALYTICS — UTILITY FUNCTIONS
# ============================================================

def _clean_instat_val(val):
    """Clean InStat value: handle '-', None, strip %."""
    if val is None:
        return None
    s = str(val).strip()
    if s in ("-", "None", "", "N/A"):
        return None
    return s


def _instat_to_number(val):
    """Convert InStat value to number (int or float). Returns None for empty."""
    cleaned = _clean_instat_val(val)
    if cleaned is None:
        return None
    # Strip percent suffix
    cleaned = cleaned.rstrip("%")
    try:
        f = float(cleaned)
        if f == int(f) and "." not in cleaned:
            return int(f)
        return f
    except (ValueError, TypeError):
        return None


def _parse_instat_row(row: dict, core_map: dict, extended_map: dict):
    """Parse an InStat row into (core_stats, extended_stats) dicts."""
    core = {}
    extended = {}

    for header, target in core_map.items():
        raw = row.get(header)
        if raw is None:
            continue
        cleaned = _clean_instat_val(raw)
        if cleaned is None:
            continue
        if target == "toi_seconds":
            core[target] = _parse_mmss_to_seconds(cleaned)
        elif target == "pim":
            # PIM comes as MM:SS in InStat
            secs = _parse_mmss_to_seconds(cleaned)
            core[target] = secs // 60 if secs else _to_int(cleaned)
        elif target == "shooting_pct":
            core[target] = _to_float(cleaned)
        elif target == "sv_pct":
            core[target] = cleaned  # Keep as string like "91.9%"
        elif target in ("ga", "sa", "sv"):
            core[target] = _to_float(cleaned)
        else:
            core[target] = _instat_to_number(cleaned)

    for header, target in extended_map.items():
        raw = row.get(header)
        if raw is None:
            continue
        cleaned = _clean_instat_val(raw)
        if cleaned is None:
            continue

        if "." in target:
            category, field = target.split(".", 1)
            if category not in extended:
                extended[category] = {}
            # Determine value type
            if "time" in field and ":" in str(cleaned):
                extended[category][field] = _parse_mmss_to_seconds(cleaned)
            elif "pct" in field or field.endswith("_%"):
                extended[category][field] = _to_float(cleaned)
            else:
                extended[category][field] = _instat_to_number(cleaned)
        else:
            # Flat field for goalie extended
            if "time" in target and ":" in str(cleaned):
                extended[target] = _parse_mmss_to_seconds(cleaned)
            elif "pct" in target:
                extended[target] = _to_float(cleaned)
            else:
                extended[target] = _instat_to_number(cleaned)

    return core, extended


def _detect_instat_file_type(headers: list) -> str:
    """Auto-detect InStat XLSX file type from headers."""
    h_set = set(h.lower().replace(" ", "_") for h in headers if h)

    # Lines files: have "line" column
    if "line" in h_set:
        return "lines"

    # Goalie files: have saves/sv% markers but NOT goals/assists (skater stats)
    goalie_markers = {"saves,_%", "saves", "goals_against", "shootout_saves", "shootouts_allowed"}
    has_goalie = len(goalie_markers & h_set) >= 2
    has_assists = "assists" in h_set or "first_assist" in h_set

    if has_goalie and not has_assists:
        has_team_col = "team" in h_set
        return "league_goalies" if has_team_col else "team_goalies"

    # Skater files: have player + scoring stats
    has_player = "player" in h_set
    has_scoring = "goals" in h_set or "points" in h_set

    if has_player and has_scoring:
        has_team_col = "team" in h_set
        return "league_skaters" if has_team_col else "team_skaters"

    # Team stats: has "team" as first col and no "player" col
    if "team" in h_set and not has_player and has_scoring:
        return "league_teams"

    # Fallback: check if it looks like team stats (no player column, has stats)
    if not has_player and ("faceoffs" in h_set or "shots_on_goal" in h_set):
        return "league_teams"

    return "unknown"


def _parse_line_players(line_string: str) -> list:
    """Parse '6 E. Weiss, 20 N. Battler, ...' into structured player refs."""
    players = []
    if not line_string:
        return players
    for entry in line_string.split(","):
        entry = entry.strip()
        if not entry:
            continue
        parts = entry.split(None, 1)
        if len(parts) == 2:
            players.append({"jersey": parts[0], "name": parts[1]})
        elif len(parts) == 1:
            players.append({"jersey": "", "name": parts[0]})
    return players


# ── Stat CRUD (edit / delete / deduplicate) ──────────────────

class StatUpdateRequest(BaseModel):
    gp: Optional[int] = None
    g: Optional[int] = None
    a: Optional[int] = None
    p: Optional[int] = None
    plus_minus: Optional[int] = None
    pim: Optional[int] = None
    season: Optional[str] = None
    shooting_pct: Optional[float] = None
    toi_seconds: Optional[int] = None
    team_name: Optional[str] = None
    notes: Optional[str] = None


@app.put("/stats/{stat_id}")
async def update_stat(stat_id: str, req: StatUpdateRequest, token_data: dict = Depends(verify_token)):
    """Update an individual stat row (owner edit)."""
    org_id = token_data["org_id"]
    conn = get_db()
    try:
        # Verify stat exists and belongs to a player in this org
        row = conn.execute("""
            SELECT ps.id, ps.player_id, p.org_id FROM player_stats ps
            JOIN players p ON p.id = ps.player_id
            WHERE ps.id = ? AND p.org_id = ?
        """, (stat_id, org_id)).fetchone()
        if not row:
            raise HTTPException(status_code=404, detail="Stat not found")

        updates = []
        params = []
        for field in ["gp", "g", "a", "p", "plus_minus", "pim", "season", "shooting_pct", "toi_seconds", "team_name", "notes"]:
            val = getattr(req, field, None)
            if val is not None:
                updates.append(f"{field} = ?")
                params.append(val)

        if not updates:
            raise HTTPException(status_code=400, detail="No fields to update")

        # Auto-calculate points if g or a changed but p wasn't explicitly set
        if req.p is None and (req.g is not None or req.a is not None):
            current = conn.execute("SELECT g, a FROM player_stats WHERE id = ?", (stat_id,)).fetchone()
            new_g = req.g if req.g is not None else current["g"]
            new_a = req.a if req.a is not None else current["a"]
            updates.append("p = ?")
            params.append(new_g + new_a)

        params.append(stat_id)
        conn.execute(f"UPDATE player_stats SET {', '.join(updates)} WHERE id = ?", params)
        conn.commit()

        updated = conn.execute("SELECT * FROM player_stats WHERE id = ?", (stat_id,)).fetchone()
        return dict(updated)
    finally:
        conn.close()


@app.delete("/stats/{stat_id}")
async def delete_stat(stat_id: str, token_data: dict = Depends(verify_token)):
    """Delete an individual stat row."""
    org_id = token_data["org_id"]
    conn = get_db()
    try:
        row = conn.execute("""
            SELECT ps.id FROM player_stats ps
            JOIN players p ON p.id = ps.player_id
            WHERE ps.id = ? AND p.org_id = ?
        """, (stat_id, org_id)).fetchone()
        if not row:
            raise HTTPException(status_code=404, detail="Stat not found")

        conn.execute("DELETE FROM player_stats WHERE id = ?", (stat_id,))
        conn.commit()
        return {"success": True, "message": "Stat deleted"}
    finally:
        conn.close()


@app.post("/stats/player/{player_id}/deduplicate")
async def deduplicate_player_stats(player_id: str, token_data: dict = Depends(verify_token)):
    """Remove duplicate season stat rows for a player, keeping the most recent one."""
    org_id = token_data["org_id"]
    conn = get_db()
    try:
        # Verify player belongs to org
        player = conn.execute("SELECT id FROM players WHERE id = ? AND org_id = ?", (player_id, org_id)).fetchone()
        if not player:
            raise HTTPException(status_code=404, detail="Player not found")

        # Find duplicate seasons (same player, season, stat_type)
        dupes = conn.execute("""
            SELECT season, stat_type, COUNT(*) as cnt
            FROM player_stats WHERE player_id = ? AND stat_type = 'season'
            GROUP BY season, stat_type HAVING cnt > 1
        """, (player_id,)).fetchall()

        removed = 0
        for d in dupes:
            # Keep the most recent row, delete the rest
            rows = conn.execute("""
                SELECT id FROM player_stats
                WHERE player_id = ? AND season = ? AND stat_type = ?
                ORDER BY created_at DESC
            """, (player_id, d["season"], d["stat_type"])).fetchall()

            # Delete all but the first (most recent)
            for row in rows[1:]:
                conn.execute("DELETE FROM player_stats WHERE id = ?", (row["id"],))
                removed += 1

        conn.commit()
        return {"success": True, "duplicates_removed": removed, "player_id": player_id}
    finally:
        conn.close()


@app.post("/stats/deduplicate-all")
async def deduplicate_all_stats(token_data: dict = Depends(verify_token)):
    """Remove duplicate season stats for ALL players in the org."""
    org_id = token_data["org_id"]
    conn = get_db()
    try:
        dupes = conn.execute("""
            SELECT ps.player_id, ps.season, ps.stat_type, COUNT(*) as cnt
            FROM player_stats ps
            JOIN players p ON p.id = ps.player_id
            WHERE p.org_id = ? AND ps.stat_type = 'season'
            GROUP BY ps.player_id, ps.season, ps.stat_type HAVING cnt > 1
        """, (org_id,)).fetchall()

        removed = 0
        for d in dupes:
            rows = conn.execute("""
                SELECT id FROM player_stats
                WHERE player_id = ? AND season = ? AND stat_type = ?
                ORDER BY created_at DESC
            """, (d["player_id"], d["season"], d["stat_type"])).fetchall()

            for row in rows[1:]:
                conn.execute("DELETE FROM player_stats WHERE id = ?", (row["id"],))
                removed += 1

        conn.commit()
        return {"success": True, "duplicates_removed": removed}
    finally:
        conn.close()


@app.post("/stats/ingest")
async def ingest_stats(
    file: UploadFile = File(...),
    player_id: Optional[str] = Query(None),
    token_data: dict = Depends(verify_token),
):
    """
    Ingest stats from a CSV or Excel file. Supports:
    - Standard format: columns like gp, g, a, p, plus_minus, pim, etc.
    - InStat game log: per-game box scores (Date, Opponent, Score, Goals, Assists, etc.)
    - InStat team stats: team skater stats with player names
    Requires player_id query param for single-player files (InStat game logs).
    """
    org_id = token_data["org_id"]
    fname = (file.filename or "").lower()

    if not any(fname.endswith(ext) for ext in (".csv", ".xlsx", ".xls", ".xlsm")):
        raise HTTPException(status_code=400, detail="File must be .csv, .xlsx, or .xls")

    content = await file.read()
    all_rows = _parse_file_to_rows(content, fname)

    if not all_rows:
        raise HTTPException(status_code=400, detail="File contains no data rows")

    headers = list(all_rows[0].keys())
    logger.info("Stats ingest headers: %s", headers)

    conn = get_db()
    inserted = 0
    errors = []

    # ── InStat Game Log Format ────────────────────────────────────
    # Per-game box scores for a single player (e.g. "Games - Ewan McChesney.xlsx")
    # Columns: Date, Opponent, Score, All shifts, Time on ice, Goals, Assists, Points, +/-, etc.
    if _detect_instat_game_log(headers):
        if not player_id:
            conn.close()
            raise HTTPException(
                status_code=400,
                detail="Per-game stats log detected — please upload this from the player's profile page so we know which player to attach stats to."
            )

        # Verify player belongs to org
        player = conn.execute("SELECT id, first_name, last_name FROM players WHERE id = ? AND org_id = ?", (player_id, org_id)).fetchone()
        if not player:
            conn.close()
            raise HTTPException(status_code=404, detail="Player not found")

        # Try to extract season from filename (e.g. "Games - Ewan McChesney, 08-Feb-2026.xlsx")
        import re as _re
        year_match = _re.search(r"(\d{4})", file.filename or "")
        default_season = ""
        if year_match:
            y = int(year_match.group(1))
            # Hockey seasons span two years; if month is before August, it's previous year's season
            default_season = f"{y-1}-{y}"

        logger.info("InStat game log detected for player %s %s — %d rows, season guess: %s",
                     player["first_name"], player["last_name"], len(all_rows), default_season or "unknown")

        for i, row in enumerate(all_rows):
            # Skip the "Average per game" summary row (has no date)
            opponent = row.get("opponent", "")
            if "average" in opponent.lower() or not opponent:
                continue

            g = _to_int(row.get("goals"))
            first_a = _to_int(row.get("first_assist"))
            second_a = _to_int(row.get("second_assist"))
            a = _to_int(row.get("assists"), first_a + second_a)
            p = _to_int(row.get("points"), g + a)

            toi = _parse_mmss_to_seconds(row.get("time_on_ice", ""))
            pim_seconds = _parse_mmss_to_seconds(row.get("penalty_time", ""))
            pim_minutes = pim_seconds // 60 if pim_seconds else 0

            # Build a game_id from date + opponent for dedup
            game_date = row.get("date", "")
            game_id = f"{game_date}_{opponent.replace(' ', '_')}" if game_date else None

            shots = _to_int(row.get("shots"))
            sog = _to_int(row.get("shots_on_goal"))
            pp_shots = _to_int(row.get("power_play_shots"))
            pk_shots = _to_int(row.get("short_handed_shots") or row.get("shorthanded_shots"))
            shooting_pct = round((g / sog * 100), 1) if sog > 0 else None

            try:
                conn.execute("""
                    INSERT INTO player_stats (id, player_id, game_id, season, stat_type, gp, g, a, p, plus_minus, pim,
                                              toi_seconds, pp_toi_seconds, pk_toi_seconds, shots, sog, shooting_pct,
                                              microstats, created_at)
                    VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                """, (
                    gen_id(), player_id, game_id, default_season, "game",
                    1,  # gp = 1 per game row
                    g, a, p,
                    _to_int(row.get("+/_") or row.get("plusminus") or row.get("+/−") or row.get("+/")),
                    pim_minutes, toi, 0, 0,
                    shots, sog, shooting_pct,
                    json.dumps({
                        "opponent": opponent,
                        "score": row.get("score", ""),
                        "date": game_date,
                        "shifts": _to_int(row.get("all_shifts")),
                        "hits": _to_int(row.get("hits")),
                        "blocked_shots": _to_int(row.get("blocked_shots")),
                        "faceoffs": _to_int(row.get("faceoffs")),
                        "faceoffs_won": _to_int(row.get("faceoffs_won")),
                        "faceoff_pct": _to_float(row.get("faceoffs_won,_%")),
                        "penalties_drawn": _to_int(row.get("penalties_drawn")),
                        "pp_shots": pp_shots,
                        "pk_shots": pk_shots,
                    }),
                    now_iso(),
                ))
                inserted += 1
            except Exception as e:
                errors.append(f"Row {i+1}: {str(e)}")

        conn.commit()

        # ── Auto-aggregate game logs into a season summary row ──────
        # Delete any existing auto-generated season summary for this player/season
        # (so re-importing doesn't create duplicates)
        if default_season and inserted > 0:
            conn.execute("""
                DELETE FROM player_stats
                WHERE player_id = ? AND season = ? AND stat_type = 'season'
                AND id IN (
                    SELECT id FROM player_stats
                    WHERE player_id = ? AND season = ? AND stat_type = 'season'
                )
            """, (player_id, default_season, player_id, default_season))

            # Aggregate all game rows for this player/season
            agg = conn.execute("""
                SELECT
                    COUNT(*) as gp,
                    SUM(g) as g, SUM(a) as a, SUM(p) as p,
                    SUM(plus_minus) as plus_minus, SUM(pim) as pim,
                    SUM(toi_seconds) as toi_seconds,
                    SUM(pp_toi_seconds) as pp_toi_seconds,
                    SUM(pk_toi_seconds) as pk_toi_seconds,
                    SUM(shots) as shots, SUM(sog) as sog
                FROM player_stats
                WHERE player_id = ? AND season = ? AND stat_type = 'game'
            """, (player_id, default_season)).fetchone()

            if agg and agg["gp"] > 0:
                total_g = agg["g"] or 0
                total_sog = agg["sog"] or 0
                season_shooting_pct = round((total_g / total_sog * 100), 1) if total_sog > 0 else None

                conn.execute("""
                    INSERT INTO player_stats (id, player_id, season, stat_type, gp, g, a, p, plus_minus, pim,
                                              toi_seconds, pp_toi_seconds, pk_toi_seconds, shots, sog, shooting_pct, created_at)
                    VALUES (?, ?, ?, 'season', ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                """, (
                    gen_id(), player_id, default_season,
                    agg["gp"], total_g, agg["a"] or 0, agg["p"] or 0,
                    agg["plus_minus"] or 0, agg["pim"] or 0,
                    agg["toi_seconds"] or 0, agg["pp_toi_seconds"] or 0, agg["pk_toi_seconds"] or 0,
                    agg["shots"] or 0, total_sog, season_shooting_pct,
                    now_iso(),
                ))
                conn.commit()
                logger.info("Auto-generated season summary for %s %s (%s): %dGP %dG %dA %dP",
                            player["first_name"], player["last_name"], default_season,
                            agg["gp"], total_g, agg["a"] or 0, agg["p"] or 0)

        conn.close()
        logger.info("InStat game log ingested: %d games for %s %s", inserted, player["first_name"], player["last_name"])

        # Trigger intelligence generation in background
        if inserted > 0:
            asyncio.create_task(_generate_player_intelligence(player_id, org_id, trigger="import"))

        return {
            "detail": f"Imported {inserted} game logs + season summary",
            "inserted": inserted + (1 if default_season and inserted > 0 else 0),
            "format": "instat_game_log",
            "errors": errors[:10],
        }

    # ── InStat Team Stats Format ──────────────────────────────────
    # Team-wide skater stats (e.g. "Skaters - St. Marys Lincolns.xlsx")
    # Has player names + stats for the whole team
    if _detect_instat_team_stats(headers):
        logger.info("InStat team stats detected — %d rows", len(all_rows))
        _imported_player_ids = set()

        for i, row in enumerate(all_rows):
            # Get player name
            player_name = row.get("player") or row.get("name") or row.get("player_name") or ""
            if not player_name or player_name == "-":
                errors.append(f"Row {i+1}: no player name")
                continue

            # Try to match to existing player by name
            parts = player_name.strip().split(None, 1)
            if len(parts) < 2:
                errors.append(f"Row {i+1}: cannot split name '{player_name}'")
                continue

            first_name, last_name = parts[0], parts[1]

            # Look up player in DB (case-insensitive match)
            matched = conn.execute(
                "SELECT id FROM players WHERE org_id = ? AND LOWER(first_name) = LOWER(?) AND LOWER(last_name) = LOWER(?)",
                (org_id, first_name, last_name)
            ).fetchone()

            if not matched:
                errors.append(f"Row {i+1}: player '{player_name}' not found in database — import them first")
                continue

            pid = matched["id"]
            g = _to_int(row.get("goals"))
            a = _to_int(row.get("assists"))
            p = _to_int(row.get("points"), g + a)
            toi = _parse_mmss_to_seconds(row.get("time_on_ice", ""))
            pim_seconds = _parse_mmss_to_seconds(row.get("penalty_time", ""))

            try:
                conn.execute("""
                    INSERT INTO player_stats (id, player_id, season, stat_type, gp, g, a, p, plus_minus, pim,
                                              toi_seconds, shots, sog, shooting_pct, microstats, created_at)
                    VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                """, (
                    gen_id(), pid, "", "season",
                    _to_int(row.get("gp") or row.get("games") or row.get("games_played")),
                    g, a, p,
                    _to_int(row.get("+/_") or row.get("plusminus") or row.get("+/−")),
                    pim_seconds // 60 if pim_seconds else 0,
                    toi,
                    _to_int(row.get("shots")),
                    _to_int(row.get("shots_on_goal") or row.get("sog")),
                    _to_float(row.get("shooting_pct") or row.get("sh%") or row.get("s%")),
                    json.dumps({
                        "faceoffs": _to_int(row.get("faceoffs")),
                        "faceoffs_won": _to_int(row.get("faceoffs_won")),
                        "hits": _to_int(row.get("hits")),
                        "blocked_shots": _to_int(row.get("blocked_shots")),
                    }),
                    now_iso(),
                ))
                inserted += 1
                _imported_player_ids.add(pid)
            except Exception as e:
                errors.append(f"Row {i+1} ({player_name}): {str(e)}")

        conn.commit()
        conn.close()
        logger.info("InStat team stats ingested: %d rows", inserted)

        # Trigger intelligence generation for each imported player in background
        for _pid in _imported_player_ids:
            asyncio.create_task(_generate_player_intelligence(_pid, org_id, trigger="import"))

        return {
            "detail": f"Imported {inserted} stat rows from team stats file",
            "inserted": inserted,
            "format": "instat_team_stats",
            "errors": errors[:10],
        }

    # ── Standard Format ───────────────────────────────────────────
    _std_player_ids = set()
    for i, row in enumerate(all_rows):
        pid = row.get("player_id") or player_id
        if not pid:
            errors.append(f"Row {i+1}: missing player_id")
            continue

        # Verify player belongs to org
        player = conn.execute("SELECT id FROM players WHERE id = ? AND org_id = ?", (pid, org_id)).fetchone()
        if not player:
            errors.append(f"Row {i+1}: player {pid} not found")
            continue

        g = _to_int(row.get("g") or row.get("goals"))
        a = _to_int(row.get("a") or row.get("assists"))
        p = _to_int(row.get("p") or row.get("pts") or row.get("points"), g + a)

        conn.execute("""
            INSERT INTO player_stats (id, player_id, season, stat_type, gp, g, a, p, plus_minus, pim,
                                      toi_seconds, pp_toi_seconds, pk_toi_seconds, shots, sog, shooting_pct, created_at)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
        """, (
            gen_id(), pid,
            row.get("season", ""),
            row.get("stat_type", "season"),
            _to_int(row.get("gp") or row.get("games") or row.get("games_played")),
            g, a, p,
            _to_int(row.get("plus_minus") or row.get("+/-") or row.get("plusminus")),
            _to_int(row.get("pim") or row.get("penalty_minutes")),
            _to_int(row.get("toi_seconds") or row.get("toi")),
            _to_int(row.get("pp_toi_seconds") or row.get("pp_toi")),
            _to_int(row.get("pk_toi_seconds") or row.get("pk_toi")),
            _to_int(row.get("shots")),
            _to_int(row.get("sog") or row.get("shots_on_goal")),
            _to_float(row.get("shooting_pct") or row.get("sh%") or row.get("s%")),
            now_iso(),
        ))
        inserted += 1
        _std_player_ids.add(pid)

    conn.commit()
    conn.close()

    logger.info("Stats ingested: %d rows for org %s (%d errors)", inserted, org_id, len(errors))

    # Trigger intelligence generation for each imported player in background
    for _pid in _std_player_ids:
        asyncio.create_task(_generate_player_intelligence(_pid, org_id, trigger="import"))

    return {
        "detail": f"Imported {inserted} stat rows",
        "inserted": inserted,
        "errors": errors[:10],
    }


# ============================================================
# TEAM ENDPOINTS
# ============================================================

@app.get("/leagues")
async def list_leagues():
    """List all leagues (reference data — no auth required)."""
    conn = get_db()
    rows = conn.execute("SELECT * FROM leagues ORDER BY sort_order, name").fetchall()
    conn.close()
    return [dict(r) for r in rows]


@app.get("/teams")
async def list_teams(token_data: dict = Depends(verify_token)):
    org_id = token_data["org_id"]
    conn = get_db()
    rows = conn.execute("SELECT * FROM teams WHERE org_id IN (?, '__global__') ORDER BY name", (org_id,)).fetchall()
    conn.close()
    return [dict(r) for r in rows]


@app.get("/teams/reference")
async def list_reference_teams(
    league: Optional[str] = None,
    token_data: dict = Depends(verify_token),
):
    """List global reference teams + org-specific teams, optionally filtered by league."""
    org_id = token_data["org_id"]
    conn = get_db()
    query = "SELECT * FROM teams WHERE (org_id = '__global__' OR org_id = ?)"
    params: list = [org_id]
    if league:
        query += " AND league = ?"
        params.append(league)
    query += " ORDER BY league, name"
    rows = conn.execute(query, params).fetchall()
    conn.close()
    return [dict(r) for r in rows]


@app.get("/teams/{team_name}/roster")
async def get_team_roster(team_name: str, token_data: dict = Depends(verify_token)):
    """Get all players on a team (matched by current_team, case-insensitive)."""
    org_id = token_data["org_id"]
    decoded_name = team_name.replace("%20", " ")
    conn = get_db()
    rows = conn.execute(
        "SELECT * FROM players WHERE org_id = ? AND LOWER(current_team) = LOWER(?) ORDER BY position, last_name",
        (org_id, decoded_name),
    ).fetchall()
    conn.close()
    return [_player_from_row(r) for r in rows]


@app.get("/teams/{team_name}/hockeytech-info")
async def get_team_hockeytech_info(team_name: str, token_data: dict = Depends(verify_token)):
    """Get HockeyTech integration info for a team (team_id and league code)."""
    org_id = token_data["org_id"]
    decoded_name = team_name.replace("%20", " ")
    conn = get_db()
    try:
        team_row = conn.execute(
            "SELECT hockeytech_team_id, hockeytech_league FROM teams WHERE LOWER(name) = LOWER(?) AND org_id IN (?, '__global__')",
            (decoded_name, org_id)
        ).fetchone()
        if team_row and team_row["hockeytech_team_id"]:
            return {
                "hockeytech_team_id": team_row["hockeytech_team_id"],
                "hockeytech_league": team_row["hockeytech_league"],
                "linked": True,
            }
        # Fallback: check if any player on this team has HT data
        ht_player = conn.execute(
            "SELECT hockeytech_league FROM players WHERE LOWER(current_team) = LOWER(?) AND org_id = ? AND hockeytech_id IS NOT NULL LIMIT 1",
            (decoded_name, org_id)
        ).fetchone()
        if not ht_player:
            return {
                "hockeytech_team_id": None,
                "hockeytech_league": None,
                "linked": False,
                "has_ht_players": False,
            }

        # Auto-repair: team has HT players but no team link — try to find and store the HT team_id
        ht_league = ht_player["hockeytech_league"]
        ht_team_id_found = None
        try:
            from hockeytech import HockeyTechClient, LEAGUES as HT_LEAGUES
            if ht_league and ht_league in HT_LEAGUES:
                client = HockeyTechClient(ht_league)
                seasons = await client.get_seasons()
                if seasons:
                    current_season = seasons[0]["id"]
                    ht_teams = await client.get_teams(current_season)
                    name_lower = decoded_name.lower()
                    for ht_team in ht_teams:
                        ht_name = (ht_team.get("name") or "").lower()
                        ht_city = (ht_team.get("city") or "").lower()
                        ht_nick = (ht_team.get("nickname") or "").lower()
                        # Match by: full name, city+nickname, or substring
                        if (name_lower == ht_name
                            or name_lower == f"{ht_city} {ht_nick}"
                            or ht_name in name_lower
                            or name_lower in ht_name
                            or (ht_nick and ht_nick in name_lower)):
                            ht_team_id_found = ht_team.get("id")
                            break
                    await client.close()
                    if ht_team_id_found:
                        # Write the link to teams table
                        team_db_row = conn.execute(
                            "SELECT id FROM teams WHERE LOWER(name) = LOWER(?) AND org_id IN (?, '__global__')",
                            (decoded_name, org_id)
                        ).fetchone()
                        if team_db_row:
                            conn.execute(
                                "UPDATE teams SET hockeytech_team_id = ?, hockeytech_league = ? WHERE id = ?",
                                (ht_team_id_found, ht_league, team_db_row["id"])
                            )
                            conn.commit()
                            logger.info("Auto-repair: linked team %s → HT team_id=%d, league=%s", decoded_name, ht_team_id_found, ht_league)
        except Exception as e:
            logger.warning("Auto-repair HT team link failed for %s: %s", decoded_name, e)

        if ht_team_id_found:
            return {
                "hockeytech_team_id": ht_team_id_found,
                "hockeytech_league": ht_league,
                "linked": True,
                "has_ht_players": True,
            }
        return {
            "hockeytech_team_id": None,
            "hockeytech_league": ht_league,
            "linked": False,
            "has_ht_players": True,
        }
    finally:
        conn.close()


@app.get("/teams/{team_name}/intelligence")
async def get_team_intelligence(team_name: str, token_data: dict = Depends(verify_token)):
    """Get the latest intelligence snapshot for a team."""
    org_id = token_data["org_id"]
    decoded_name = team_name.replace("%20", " ")
    conn = get_db()
    row = conn.execute(
        "SELECT * FROM team_intelligence WHERE LOWER(team_name) = LOWER(?) AND org_id = ? ORDER BY version DESC LIMIT 1",
        (decoded_name, org_id)
    ).fetchone()
    conn.close()

    if not row:
        return {"version": 0, "team_name": decoded_name}

    d = dict(row)
    for field in ("strengths", "vulnerabilities", "key_personnel", "comparable_teams", "tags"):
        val = d.get(field)
        if isinstance(val, str):
            try:
                d[field] = json.loads(val)
            except Exception:
                d[field] = []
        elif val is None:
            d[field] = []
    return d


@app.post("/teams/{team_name}/intelligence")
async def generate_team_intelligence(team_name: str, token_data: dict = Depends(verify_token)):
    """Generate AI-powered team identity analysis."""
    org_id = token_data["org_id"]
    decoded_name = team_name.replace("%20", " ")
    conn = get_db()

    # Get current version
    current = conn.execute(
        "SELECT version FROM team_intelligence WHERE LOWER(team_name) = LOWER(?) AND org_id = ? ORDER BY version DESC LIMIT 1",
        (decoded_name, org_id)
    ).fetchone()
    new_version = (current["version"] + 1) if current else 1

    # Gather team data for AI prompt
    roster = conn.execute(
        "SELECT * FROM players WHERE org_id = ? AND LOWER(current_team) = LOWER(?)",
        (org_id, decoded_name)
    ).fetchall()

    # Get team stats
    stats = conn.execute(
        """SELECT p.first_name, p.last_name, p.position, p.archetype,
                  ps.gp, ps.g, ps.a, ps.p, ps.plus_minus, ps.pim
           FROM players p
           LEFT JOIN (
               SELECT player_id, gp, g, a, p, plus_minus, pim,
                      ROW_NUMBER() OVER (PARTITION BY player_id ORDER BY season DESC) as rn
               FROM player_stats
           ) ps ON p.id = ps.player_id AND ps.rn = 1
           WHERE p.org_id = ? AND LOWER(p.current_team) = LOWER(?)
           ORDER BY ps.p DESC NULLS LAST""",
        (org_id, decoded_name)
    ).fetchall()

    # Get team system
    system = conn.execute(
        "SELECT * FROM team_systems WHERE LOWER(team_name) = LOWER(?) AND org_id = ? ORDER BY created_at DESC LIMIT 1",
        (decoded_name, org_id)
    ).fetchone()

    # Get recent game results (table may not exist yet)
    games = []
    try:
        games = conn.execute(
            "SELECT * FROM team_games WHERE LOWER(team_name) = LOWER(?) AND org_id = ? ORDER BY game_date DESC LIMIT 10",
            (decoded_name, org_id)
        ).fetchall()
    except Exception:
        pass

    # Check for AI key
    anthropic_key = os.getenv("ANTHROPIC_API_KEY", "")
    if not anthropic_key:
        # Demo mode — return placeholder
        intel_id = str(uuid.uuid4())
        demo = {
            "id": intel_id, "team_name": decoded_name, "org_id": org_id,
            "playing_style": "Analysis Pending",
            "system_summary": f"Connect your Anthropic API key to generate AI-powered team intelligence for {decoded_name}.",
            "identity": f"Team identity analysis requires an API key. Add ANTHROPIC_API_KEY to your backend .env file and regenerate.",
            "strengths": json.dumps(["Add API key to unlock AI analysis"]),
            "vulnerabilities": json.dumps(["Add API key to unlock AI analysis"]),
            "key_personnel": json.dumps([]),
            "comparable_teams": json.dumps([]),
            "tags": json.dumps([]),
            "special_teams_identity": None,
            "player_archetype_fit": None,
            "trigger": "manual_demo",
            "version": new_version,
        }
        conn.execute("""
            INSERT INTO team_intelligence (id, team_name, org_id, playing_style, system_summary, identity,
                strengths, vulnerabilities, key_personnel, comparable_teams, tags,
                special_teams_identity, player_archetype_fit, trigger, version)
            VALUES (:id, :team_name, :org_id, :playing_style, :system_summary, :identity,
                :strengths, :vulnerabilities, :key_personnel, :comparable_teams, :tags,
                :special_teams_identity, :player_archetype_fit, :trigger, :version)
        """, demo)
        conn.commit()
        conn.close()
        # Return parsed version
        for f in ("strengths", "vulnerabilities", "key_personnel", "comparable_teams", "tags"):
            demo[f] = json.loads(demo[f])
        return demo

    # === AI-POWERED ANALYSIS ===
    # Build context string from roster, stats, system, games
    roster_summary = f"{len(roster)} players on roster."
    stats_lines = []
    for s in stats[:15]:
        stats_lines.append(f"{s['first_name']} {s['last_name']} ({s['position']}): {s['gp'] or 0}GP {s['g'] or 0}G {s['a'] or 0}A {s['p'] or 0}PTS")
    system_str = ""
    if system:
        system_str = f"Forecheck: {system['forecheck'] or 'N/A'}, DZ: {system['dz_structure'] or 'N/A'}, Breakout: {system['breakout'] or 'N/A'}, Pace: {system['pace'] or 'N/A'}"

    record_str = ""
    if games:
        wins = sum(1 for g in games if g["result"] == "W")
        losses = sum(1 for g in games if g["result"] == "L")
        record_str = f"Recent 10 games: {wins}W-{losses}L"

    prompt = f"""Analyze this hockey team and generate a comprehensive team identity profile.

Team: {decoded_name}
{roster_summary}
{record_str}

Top Scorers:
{chr(10).join(stats_lines) if stats_lines else "No stats available"}

System: {system_str or "Not defined"}

Respond in this exact JSON format (no markdown, no backticks):
{{
    "playing_style": "2-4 word style descriptor (e.g., High-Pressure Forecheck, Defensive Trap, Run-and-Gun)",
    "system_summary": "1-2 sentence summary of how this team plays",
    "identity": "2-3 paragraph detailed identity description",
    "strengths": ["strength 1", "strength 2", "strength 3"],
    "vulnerabilities": ["vulnerability 1", "vulnerability 2"],
    "key_personnel": [{{"name": "Player Name", "role": "1C / Top scorer / Shutdown D / etc"}}],
    "special_teams_identity": "1-2 sentences on PP and PK tendencies",
    "player_archetype_fit": "What type of player fits this team's system",
    "comparable_teams": ["comparable team 1", "comparable team 2"],
    "tags": ["tag1", "tag2", "tag3", "tag4"]
}}"""

    import httpx
    try:
        async with httpx.AsyncClient(timeout=60) as client:
            resp = await client.post(
                "https://api.anthropic.com/v1/messages",
                headers={"x-api-key": anthropic_key, "content-type": "application/json", "anthropic-version": "2023-06-01"},
                json={"model": "claude-sonnet-4-20250514", "max_tokens": 2000, "messages": [{"role": "user", "content": prompt}]}
            )
            resp.raise_for_status()
            ai_text = resp.json()["content"][0]["text"]
            # Strip markdown fences if present
            ai_text = ai_text.strip()
            if ai_text.startswith("```"):
                ai_text = ai_text.split("\n", 1)[1] if "\n" in ai_text else ai_text[3:]
            if ai_text.endswith("```"):
                ai_text = ai_text[:-3]
            ai_data = json.loads(ai_text.strip())
    except Exception as e:
        conn.close()
        raise HTTPException(status_code=500, detail=f"AI analysis failed: {str(e)}")

    intel_id = str(uuid.uuid4())
    row_data = {
        "id": intel_id, "team_name": decoded_name, "org_id": org_id,
        "playing_style": ai_data.get("playing_style"),
        "system_summary": ai_data.get("system_summary"),
        "identity": ai_data.get("identity"),
        "strengths": json.dumps(ai_data.get("strengths", [])),
        "vulnerabilities": json.dumps(ai_data.get("vulnerabilities", [])),
        "key_personnel": json.dumps(ai_data.get("key_personnel", [])),
        "comparable_teams": json.dumps(ai_data.get("comparable_teams", [])),
        "tags": json.dumps(ai_data.get("tags", [])),
        "special_teams_identity": ai_data.get("special_teams_identity"),
        "player_archetype_fit": ai_data.get("player_archetype_fit"),
        "trigger": "manual",
        "version": new_version,
    }
    conn.execute("""
        INSERT INTO team_intelligence (id, team_name, org_id, playing_style, system_summary, identity,
            strengths, vulnerabilities, key_personnel, comparable_teams, tags,
            special_teams_identity, player_archetype_fit, trigger, version)
        VALUES (:id, :team_name, :org_id, :playing_style, :system_summary, :identity,
            :strengths, :vulnerabilities, :key_personnel, :comparable_teams, :tags,
            :special_teams_identity, :player_archetype_fit, :trigger, :version)
    """, row_data)
    conn.commit()
    conn.close()

    # Parse JSON fields for return
    for f in ("strengths", "vulnerabilities", "key_personnel", "comparable_teams", "tags"):
        row_data[f] = json.loads(row_data[f]) if isinstance(row_data[f], str) else row_data[f]
    return row_data


@app.get("/teams/{team_name}/roster-stats")
async def get_team_roster_with_stats(team_name: str, token_data: dict = Depends(verify_token)):
    """Get team roster with latest season stats for each player."""
    org_id = token_data["org_id"]
    decoded_name = team_name.replace("%20", " ")
    conn = get_db()

    rows = conn.execute("""
        SELECT p.*,
               ps.gp, ps.g, ps.a, ps.p AS pts, ps.plus_minus, ps.pim, ps.season,
               gs.gp as g_gp, gs.ga as g_ga, gs.sv as g_sv, gs.gaa, gs.sv_pct
        FROM players p
        LEFT JOIN (
            SELECT player_id, gp, g, a, p, plus_minus, pim, season,
                   ROW_NUMBER() OVER (PARTITION BY player_id ORDER BY season DESC) as rn
            FROM player_stats
        ) ps ON p.id = ps.player_id AND ps.rn = 1
        LEFT JOIN (
            SELECT player_id, gp, ga, sv, gaa, sv_pct,
                   ROW_NUMBER() OVER (PARTITION BY player_id ORDER BY season DESC) as rn
            FROM goalie_stats
        ) gs ON p.id = gs.player_id AND gs.rn = 1
        WHERE p.org_id = ? AND LOWER(p.current_team) = LOWER(?)
        ORDER BY p.position, p.last_name
    """, (org_id, decoded_name)).fetchall()
    conn.close()

    result = []
    for r in rows:
        d = _player_from_row(r)
        d["stats"] = {
            "gp": r["gp"], "g": r["g"], "a": r["a"], "p": r["pts"],
            "plus_minus": r["plus_minus"], "pim": r["pim"], "season": r["season"],
        } if r["gp"] else None
        d["goalie_stats"] = {
            "gp": r["g_gp"], "ga": r["g_ga"], "sv": r["g_sv"],
            "gaa": r["gaa"], "sv_pct": r["sv_pct"],
        } if r["g_gp"] else None
        result.append(d)
    return result


@app.get("/teams/{team_name}/reports")
async def get_team_reports(team_name: str, token_data: dict = Depends(verify_token)):
    """Get all reports for players on a specific team, plus team-level reports."""
    org_id = token_data["org_id"]
    decoded_name = team_name.replace("%20", " ")
    conn = get_db()
    # Player-linked reports (via player's current_team) + team-level reports (via r.team_name)
    rows = conn.execute("""
        SELECT r.* FROM reports r
        LEFT JOIN players p ON r.player_id = p.id
        WHERE r.org_id = ?
          AND (LOWER(p.current_team) = LOWER(?) OR LOWER(r.team_name) = LOWER(?))
        ORDER BY r.created_at DESC
    """, (org_id, decoded_name, decoded_name)).fetchall()
    conn.close()
    results = []
    seen_ids = set()
    for r in rows:
        d = dict(r)
        if d["id"] in seen_ids:
            continue
        seen_ids.add(d["id"])
        for json_field in ("output_json", "input_data"):
            if d.get(json_field) and isinstance(d[json_field], str):
                try:
                    d[json_field] = json.loads(d[json_field])
                except Exception:
                    pass
        results.append(d)
    return results


class TeamCreateRequest(BaseModel):
    name: str
    league: Optional[str] = None
    city: Optional[str] = None
    abbreviation: Optional[str] = None


@app.post("/teams")
async def create_team(body: TeamCreateRequest, token_data: dict = Depends(verify_token)):
    org_id = token_data["org_id"]
    team_id = gen_id()
    conn = get_db()
    conn.execute(
        "INSERT INTO teams (id, org_id, name, league, city, abbreviation) VALUES (?, ?, ?, ?, ?, ?)",
        (team_id, org_id, body.name, body.league, body.city, body.abbreviation),
    )
    conn.commit()
    row = conn.execute("SELECT * FROM teams WHERE id = ?", (team_id,)).fetchone()
    conn.close()
    return dict(row)


# ============================================================
# REPORT GENERATION
# ============================================================

# ── Report Output Validator + Repair ──

# Slang words that should NOT appear in formal reports (Bench Talk tone leakage)
_REPORT_SLANG_WORDS = [
    "pigeon", "barnburner", "turnstile", "sieve", "gongshow", "bender",
    "grocery stick", "plug", "cement in his skates", "warm up the bus",
    "gino", "celly", "chiclets", "sin bin", "bag skate", "lettuce",
    "top cheese", "top ched", "mitts", "twig", "lumber", "biscuit",
    "dangle", "undresses",
]


def _validate_and_repair_report(output_text: str, report_type: str, client, llm_model: str, mode: str = None) -> tuple:
    """Validate report output format and attempt repair if needed.
    Mode-aware: adjusts slang/jargon checks based on PXI mode.
    Returns (possibly_repaired_text, list_of_warnings)."""
    warnings = []

    # ── Hard checks (will trigger repair) ──
    needs_repair = False

    # Check 1: Must contain EXECUTIVE_SUMMARY header
    has_exec = bool(re.search(r'^EXECUTIVE_SUMMARY\s*[:—\-]?\s*$', output_text, re.MULTILINE))
    if not has_exec:
        warnings.append("HARD: Missing EXECUTIVE_SUMMARY header")
        needs_repair = True

    # Check 2: Must contain BOTTOM_LINE header (except practice_plan, team_identity, opponent_gameplan)
    skip_bottom = report_type in ("practice_plan", "team_identity", "opponent_gameplan", "game_decision", "line_chemistry", "st_optimization")
    if not skip_bottom:
        has_bottom = bool(re.search(r'^BOTTOM_LINE\s*[:—\-]?\s*$', output_text, re.MULTILINE))
        if not has_bottom:
            warnings.append("HARD: Missing BOTTOM_LINE header")
            needs_repair = True

    # Check 3: No === delimited format (wrong parser)
    if re.search(r'^===\s+\w+', output_text, re.MULTILINE):
        warnings.append("HARD: Contains === delimited sections (wrong format)")
        needs_repair = True

    # Check 4: Not wrapped in markdown code blocks
    if output_text.strip().startswith('```'):
        warnings.append("HARD: Wrapped in markdown code block")
        needs_repair = True

    # ── Soft checks (log warning, accept report) ──

    # Mode-aware slang/jargon checking
    # Strict modes: scout, gm, analyst, agent — no slang allowed
    # Relaxed modes: coach, broadcast — allow hockey colloquialisms (reduced slang list)
    # Parent mode: check for excessive jargon/complexity instead of slang
    _STRICT_SLANG_MODES = {"scout", "gm", "analyst", "agent", "skill_coach", "mental_coach"}
    _RELAXED_SLANG_MODES = {"coach", "broadcast", "producer"}
    # Slang terms that are acceptable in coaching/broadcast context (hockey vernacular)
    _COACHING_OK_SLANG = {"dangle", "mitts", "top cheese", "biscuit", "twig"}
    # Jargon terms that should be flagged in parent mode (too technical for families)
    _PARENT_JARGON_WORDS = [
        "corsi", "fenwick", "xgf", "xga", "war", "gamescore", "gar",
        "expected goals", "shot suppression", "zone exit rate", "controlled entry",
        "high-danger chance", "slot shot", "royal road", "net-front deflection rate",
        "d-zone coverage index", "gap control metric", "transition xg",
    ]

    if mode == "parent":
        # Parent mode: check for excessive hockey jargon instead of slang
        found_jargon = [w for w in _PARENT_JARGON_WORDS if w.lower() in output_text.lower()]
        if found_jargon:
            warnings.append(f"SOFT: Jargon detected in parent-mode report (should be plain language): {', '.join(found_jargon[:5])}")
    elif mode in _RELAXED_SLANG_MODES:
        # Coach/broadcast/producer: only flag egregious slang (exclude acceptable hockey terms)
        strict_slang = [w for w in _REPORT_SLANG_WORDS if w.lower() not in _COACHING_OK_SLANG]
        found_slang = [w for w in strict_slang if w.lower() in output_text.lower()]
        if found_slang:
            warnings.append(f"SOFT: Slang detected in report: {', '.join(found_slang[:5])}")
    else:
        # Default / strict modes (scout, gm, analyst, agent, skill_coach, mental_coach): full slang check
        found_slang = [w for w in _REPORT_SLANG_WORDS if w.lower() in output_text.lower()]
        if found_slang:
            warnings.append(f"SOFT: Slang detected in report: {', '.join(found_slang[:5])}")

    has_confidence = bool(re.search(r'Confidence:\s*(HIGH|MED|LOW)', output_text, re.IGNORECASE))
    if not has_confidence:
        warnings.append("SOFT: No Confidence tag found in report")

    has_grade = bool(re.search(r'Overall\s*Grade[:\s]*[A-D][+-]?', output_text, re.IGNORECASE))
    report_types_without_grade = ("practice_plan", "team_identity", "opponent_gameplan", "game_decision",
                                   "line_chemistry", "st_optimization", "playoff_series", "goalie_tandem")
    if not has_grade and report_type not in report_types_without_grade:
        warnings.append("SOFT: No Overall Grade found in report")

    # ── Mode-aware section expectations (log if mode-expected sections missing) ──
    if mode:
        wiring = MODE_TEMPLATE_WIRING.get(report_type)
        if wiring:
            expected_mode = wiring.get("primary", "scout")
            if mode != expected_mode:
                warnings.append(f"INFO: Report generated in {mode} mode but template {report_type} is wired for {expected_mode} mode")

    # ── Template-specific section completeness check ──
    expected_sections = REQUIRED_SECTIONS_BY_TYPE.get(report_type)
    if expected_sections:
        found_headers = set(re.findall(r'^([A-Z][A-Z0-9_]+(?:_[A-Z0-9]+)*)\s*[:—\-]?\s*$', output_text, re.MULTILINE))
        missing_sections = [s for s in expected_sections if s not in found_headers]
        if missing_sections:
            warnings.append(f"SOFT: Missing expected sections for {report_type}: {', '.join(missing_sections[:5])}")

    # ── Repair if hard checks failed ──
    if needs_repair and client:
        try:
            # Build expected sections hint for the repair prompt
            _expected_hint = ""
            _repair_sections = REQUIRED_SECTIONS_BY_TYPE.get(report_type)
            if _repair_sections:
                _expected_hint = f"\n7. Expected sections for this report type: {', '.join(_repair_sections)}"

            repair_prompt = f"""The following report has formatting issues. Reformat it to meet these requirements:
1. Use ALL_CAPS_WITH_UNDERSCORES section headers on their own lines (e.g., EXECUTIVE_SUMMARY, KEY_NUMBERS, STRENGTHS, DEVELOPMENT_AREAS, BOTTOM_LINE)
2. Ensure EXECUTIVE_SUMMARY and BOTTOM_LINE sections exist
3. Remove any === delimiters and replace with plain ALL_CAPS headers
4. Remove any markdown code block wrappers
5. Keep all content and analysis intact — only fix the formatting
6. Do NOT use markdown formatting (no **, no #, no ```){_expected_hint}

Report to reformat:

{output_text}"""

            repair_response = client.messages.create(
                model=llm_model,
                max_tokens=8000,
                messages=[{"role": "user", "content": repair_prompt}],
            )
            repaired = repair_response.content[0].text
            warnings.append("REPAIR: Repair prompt executed")

            # Verify repair worked
            has_exec_after = bool(re.search(r'^EXECUTIVE_SUMMARY\s*[:—\-]?\s*$', repaired, re.MULTILINE))
            if has_exec_after:
                warnings.append("REPAIR: Successful — repaired output accepted")
                return repaired, warnings
            else:
                warnings.append("REPAIR: Failed — repair did not fix issues, using original")
        except Exception as e:
            warnings.append(f"REPAIR: Exception during repair: {str(e)[:200]}")

    return output_text, warnings


def _score_report_quality(output_text: str, report_type: str, mode: str = None) -> dict:
    """Score a generated report on quality dimensions (0-100).
    Returns {"score": float, "breakdown": {dimension: points}, "flags": [str]}."""
    breakdown = {}
    flags = []

    # Find all section headers in the report
    found_sections = re.findall(r'^([A-Z][A-Z0-9_]+(?:_[A-Z0-9]+)*)\s*[:—\-]?\s*$', output_text, re.MULTILINE)
    found_set = set(found_sections)

    # 1. Section completeness (0-25)
    expected = REQUIRED_SECTIONS_BY_TYPE.get(report_type, ["EXECUTIVE_SUMMARY", "BOTTOM_LINE"])
    if expected:
        present = sum(1 for s in expected if s in found_set)
        section_ratio = present / len(expected) if expected else 1.0
        breakdown["section_completeness"] = round(section_ratio * 25, 1)
        if section_ratio < 1.0:
            missing = [s for s in expected if s not in found_set]
            flags.append(f"Missing sections: {', '.join(missing)}")
    else:
        breakdown["section_completeness"] = 25.0

    # 2. Evidence discipline (0-25)
    ev_score = 0.0
    has_confidence = bool(re.search(r'Confidence:\s*(HIGH|MED|LOW)', output_text, re.IGNORECASE))
    if has_confidence:
        ev_score += 10.0
    else:
        flags.append("No CONFIDENCE tag found")

    # Check for INFERENCE labels (good practice)
    inference_count = len(re.findall(r'INFERENCE\s*[-—]', output_text))
    if inference_count > 0:
        ev_score += min(5.0, inference_count * 1.5)

    # Check for DATA NOT AVAILABLE usage (shows data awareness)
    dna_count = len(re.findall(r'DATA NOT AVAILABLE', output_text))
    # Having 0 is fine (may have full data), having some shows awareness
    if dna_count > 0:
        ev_score += 3.0

    # Check for quantitative references (stats cited)
    stat_refs = len(re.findall(r'\d+\.\d+%|\d+\s*(?:GP|G|A|P|PIM|SOG|S%)', output_text))
    if stat_refs >= 5:
        ev_score += 7.0
    elif stat_refs >= 2:
        ev_score += 4.0
    elif stat_refs >= 1:
        ev_score += 2.0
    else:
        flags.append("Few quantitative stat references")

    breakdown["evidence_discipline"] = min(25.0, round(ev_score, 1))

    # 3. Depth adequacy (0-25)
    # Split by section headers and check word counts
    sections_text = re.split(r'^[A-Z][A-Z0-9_]+(?:_[A-Z0-9]+)*\s*[:—\-]?\s*$', output_text, flags=re.MULTILINE)
    total_words = len(output_text.split())
    depth_score = 0.0

    # Overall length check
    if total_words >= 800:
        depth_score += 10.0
    elif total_words >= 400:
        depth_score += 6.0
    elif total_words >= 200:
        depth_score += 3.0
    else:
        flags.append(f"Report very short ({total_words} words)")

    # Section depth — check that we have multiple substantial sections
    substantial_sections = sum(1 for s in sections_text if len(s.split()) >= 30)
    if substantial_sections >= 5:
        depth_score += 15.0
    elif substantial_sections >= 3:
        depth_score += 10.0
    elif substantial_sections >= 2:
        depth_score += 6.0
    else:
        flags.append(f"Only {substantial_sections} substantial sections (30+ words each)")

    breakdown["depth_adequacy"] = min(25.0, round(depth_score, 1))

    # 4. Grade presence (0-10)
    grade_exempt = ("practice_plan", "team_identity", "opponent_gameplan", "game_decision",
                    "line_chemistry", "st_optimization", "playoff_series", "goalie_tandem",
                    "pre_game_intel", "player_guide_prep_college")
    has_grade = bool(re.search(r'Overall\s*Grade[:\s]*[A-D][+-]?', output_text, re.IGNORECASE))
    if report_type in grade_exempt:
        breakdown["grade_presence"] = 10.0  # Full marks — not applicable
    elif has_grade:
        breakdown["grade_presence"] = 10.0
    else:
        breakdown["grade_presence"] = 0.0
        flags.append("Missing Overall Grade")

    # 5. Cleanliness (0-15)
    clean_score = 15.0
    if output_text.strip().startswith('```'):
        clean_score -= 5.0
        flags.append("Wrapped in markdown code block")
    if re.search(r'^===\s+\w+', output_text, re.MULTILINE):
        clean_score -= 5.0
        flags.append("Contains === delimited sections")
    if 'the JSON' in output_text.lower() or 'the data payload' in output_text.lower():
        clean_score -= 3.0
        flags.append("References 'the JSON' or 'the data payload'")
    if re.search(r'\*\*[^*]+\*\*', output_text):
        clean_score -= 2.0
        flags.append("Contains markdown bold formatting")

    breakdown["cleanliness"] = max(0.0, round(clean_score, 1))

    total = round(sum(breakdown.values()), 1)
    return {"score": total, "breakdown": breakdown, "flags": flags}


def _generate_mock_report(player: dict, report_type: str) -> str:
    """Generate a structured mock report without calling the LLM."""
    name = f"{player.get('first_name', '')} {player.get('last_name', '')}"
    pos = player.get("position", "F")
    team = player.get("current_team", "Unknown Team")
    league = player.get("current_league", "")
    height = player.get("height_cm", "N/A")
    weight = player.get("weight_kg", "N/A")
    shoots = player.get("shoots", "N/A")

    return f"""EXECUTIVE_SUMMARY
{name} is a {height}cm, {weight}kg {pos} who shoots {shoots}, currently playing for {team}{f' in the {league}' if league else ''}. No game statistics are available in the current dataset — this assessment is based on profile data only.

This report was generated in demo mode. Connect your Anthropic API key in backend/.env to generate full AI-powered scouting intelligence.

KEY_NUMBERS
- Position: {pos}
- Shoots: {shoots}
- Height/Weight: {height}cm / {weight}kg
- Team: {team}
- League: {league or 'N/A'}
- Season Line: No stats available in current data

STRENGTHS
Based on the available profile data for {name}:

1. Active roster player at the {league or 'current'} level, indicating competitive ability against age-appropriate or higher competition.
2. {pos} positioning suggests versatility in the lineup — centers typically anchor lines and take key faceoffs, while wingers drive possession on their strong side.
3. Listed at {height}cm/{weight}kg, which is solid size for {league or 'current'} level play.

Note: Additional on-ice observation data, video clips, and advanced metrics would significantly enhance this assessment.

DEVELOPMENT_AREAS
1. Without advanced tracking data (zone entries, slot shots, defensive recoveries), specific technical gaps cannot be identified. Recommend adding microstats via CSV import.
2. Shooting efficiency data is limited — tracking shot quality and release speed would help project offensive ceiling.
3. Special teams deployment data is not available. Understanding PP and PK usage would clarify his role within the team structure.

DEVELOPMENT_PRIORITIES
1. Add game-level statistics via CSV import to enable trend analysis and per-game consistency tracking.
2. Include video observation notes in the player profile to capture qualitative scouting insights.
3. Track practice habits, compete level, and coachability ratings to build a holistic development picture.

ROLE_FIT
{name} profiles as a {league or 'current'}-level {pos} with roster depth value. Role projection requires additional performance data to refine.

Without deployment data (TOI, PP/PK minutes, line combinations), specific role-fit recommendations are limited. This is a key area where adding detailed stats will unlock deeper analysis.

BOTTOM_LINE
{name} ({team}{f', {league}' if league else ''}) is an active {pos} whose full projection requires richer data input. The ProspectX platform is designed to turn structured stats and scouting notes into actionable intelligence — import CSV stats, add observation notes, and re-generate this report for a comprehensive AI-powered assessment.

**To unlock full report intelligence:** Add your Anthropic API key to backend/.env and re-generate this report.
"""


TEAM_REPORT_TYPES = [
    "team_identity", "opponent_gameplan", "line_chemistry", "st_optimization",
    "practice_plan", "playoff_series", "goalie_tandem",
    "league_benchmarks", "season_projection", "free_agent_market",
    "pre_game_intel",
]

# ── Custom Report Focus Areas → Prompt Sections Map ───────────
CUSTOM_FOCUS_PROMPTS = {
    "skating": {
        "label": "Skating",
        "prompt": "Provide a detailed SKATING_ASSESSMENT section evaluating: stride mechanics, top-end speed, acceleration, agility/edgework, backward skating, pivots, crossovers, and skating under pressure. Reference specific skating data or observations.",
        "sections": ["SKATING_ASSESSMENT"],
    },
    "offense": {
        "label": "Offensive Game",
        "prompt": "Provide a detailed OFFENSIVE_GAME section evaluating: shot release/accuracy, finishing ability, creativity with the puck, zone entry patterns, cycle game involvement, net-front presence, one-timer ability, and offensive instincts. Reference shooting %, goals, xG data if available.",
        "sections": ["OFFENSIVE_GAME"],
    },
    "defense": {
        "label": "Defensive Game",
        "prompt": "Provide a detailed DEFENSIVE_GAME section evaluating: gap control, stick positioning, shot blocking, DZ coverage responsibility, backchecking effort, board play defense, take-aways vs turnovers, and defensive reads. Reference +/-, takeaway data, and CORSI if available.",
        "sections": ["DEFENSIVE_GAME"],
    },
    "transition": {
        "label": "Transition Play",
        "prompt": "Provide a detailed TRANSITION_GAME section evaluating: breakout passing, zone exit execution, neutral zone play, controlled entries vs dump-ins, transition speed, puck retrieval, and ability to create odd-man rushes. Reference zone entry/breakout data if available.",
        "sections": ["TRANSITION_GAME"],
    },
    "hockey_iq": {
        "label": "Hockey IQ / Sense",
        "prompt": "Provide a detailed HOCKEY_SENSE section evaluating: anticipation, read-and-react ability, decision-making speed, understanding of positional play, ability to process the game at speed, coaching adaptability, and spatial awareness. Reference specific game situations.",
        "sections": ["HOCKEY_SENSE"],
    },
    "compete": {
        "label": "Compete Level",
        "prompt": "Provide a detailed COMPETE_LEVEL section evaluating: battle intensity, puck battles won/lost, willingness to go to hard areas, physical engagement, effort consistency across shifts, compete in key moments, and work ethic indicators. Reference PIM, hits, puck battle data if available.",
        "sections": ["COMPETE_LEVEL"],
    },
    "special_teams": {
        "label": "Special Teams",
        "prompt": "Provide a detailed SPECIAL_TEAMS section evaluating: power play role and effectiveness (PP1/PP2, bumper/flank/net-front/point), penalty kill contributions, PP points, SH points, PP and PK time-on-ice. Recommend optimal special teams deployment.",
        "sections": ["SPECIAL_TEAMS"],
    },
    "projection": {
        "label": "Projection / Ceiling",
        "prompt": "Provide a detailed PROJECTION section with: realistic ceiling and floor projections, development timeline, age-curve analysis for their position and league level, comparable players at this stage of development, and pathway to next level. Include Conservative / Expected / Optimistic scenarios with specific stat projections.",
        "sections": ["PROJECTION", "COMPARABLE_PATHWAYS"],
    },
    "trade_value": {
        "label": "Trade / Acquisition Value",
        "prompt": "Provide a detailed TRADE_VALUE section with: current market value assessment, contract considerations, acquisition cost (draft picks/prospects), roster fit analysis for acquiring teams, risk/reward profile, and comparable recent transactions. Include a clear BUY/HOLD/SELL recommendation.",
        "sections": ["TRADE_VALUE", "MARKET_POSITION"],
    },
    "development": {
        "label": "Development Plan",
        "prompt": "Provide a detailed DEVELOPMENT_PLAN section with: top 3-5 specific skills to develop (ranked by impact), 30-day and 90-day improvement targets, on-ice drill recommendations, off-ice training focus, measurable benchmarks for progress, and recommended coaching approach.",
        "sections": ["DEVELOPMENT_PLAN", "DEVELOPMENT_PRIORITIES"],
    },
    "system_fit": {
        "label": "System Fit",
        "prompt": "Provide a detailed SYSTEM_FIT section evaluating: how this player fits the team's tactical systems (forecheck, DZ, OZ, special teams), role within the system (F1/F2/F3, pinch/stay, PP bumper/flank), system compatibility rating (Elite Fit / Strong Fit / Developing Fit / Adjustment Needed), and coaching recommendations for optimizing deployment within the system.",
        "sections": ["SYSTEM_FIT"],
    },
    "draft": {
        "label": "Draft Analysis",
        "prompt": "Provide a detailed DRAFT_ANALYSIS section with: draft eligibility status, projected draft round/range, draft stock trajectory (rising/falling/steady), draft-day value proposition, comparable draft picks at similar production levels, and red flags or green flags for drafting teams.",
        "sections": ["DRAFT_POSITIONING", "DRAFT_ELIGIBILITY"],
    },
    "physical": {
        "label": "Physical Profile",
        "prompt": "Provide a detailed PHYSICAL_PROFILE section evaluating: size/strength for the level, physical maturity and growth projection, body composition, endurance/stamina across a game, injury history concerns, and how physical tools translate to on-ice impact.",
        "sections": ["PHYSICAL_PROFILE"],
    },
    "goaltending": {
        "label": "Goaltending (Goalie Only)",
        "prompt": "Provide a detailed TECHNICAL_ASSESSMENT section for the goaltender: stance and positioning, movement efficiency (butterfly, T-push, shuffle), rebound control, glove/blocker technique, puck tracking, crease management, post integration, puck handling, mental composure, and recovery speed. Reference save %, GAA, high-danger save % if available.",
        "sections": ["TECHNICAL_ASSESSMENT", "MENTAL_GAME"],
    },
}

CUSTOM_AUDIENCE_PROMPTS = {
    "coaching_staff": {
        "label": "Coaching Staff",
        "tone": "Write for an experienced coaching staff. Use tactical hockey language, deployment recommendations, and specific system references. Focus on actionable coaching intel — line combinations, matchup strategies, practice focus areas. Be direct and tactical.",
    },
    "gm_management": {
        "label": "GM / Management",
        "tone": "Write for a General Manager and hockey operations staff. Focus on player value, roster fit, cap implications, acquisition cost, development timeline, and organizational depth chart impact. Balance analytics with eye-test evaluation.",
    },
    "scouts": {
        "label": "Professional Scouts",
        "tone": "Write in professional scouting report format. Use standard scouting terminology, letter grades, tool ratings, and projection language. Be concise, honest, and focused on translatable skills and red/green flags.",
    },
    "agents": {
        "label": "Agents / Advisors",
        "tone": "Write for player agents and advisors. Highlight strengths, marketability, development upside, and career trajectory. Professional and polished, suitable for sharing with teams, but honest about areas for growth.",
    },
    "family": {
        "label": "Player & Family",
        "tone": "Write for the player and their family. Use clear, accessible language (avoid excessive jargon). Be encouraging but honest. Focus on development, next steps, and actionable goals. Celebrate progress while identifying growth areas.",
    },
    "general": {
        "label": "General Audience",
        "tone": "Write for a broad audience — hockey fans, media, or general stakeholders. Explain hockey terminology when used. Balance analysis with readability. Keep it engaging and informative.",
    },
}

CUSTOM_DEPTH_CONFIG = {
    "brief": {
        "label": "Executive Brief",
        "max_tokens": 3000,
        "instruction": "Keep the report CONCISE — executive brief format. Each section should be 3-5 sentences maximum. Focus on the key takeaways and bottom-line assessment. Target 500-800 words total.",
    },
    "standard": {
        "label": "Standard Report",
        "max_tokens": 6000,
        "instruction": "Produce a standard-depth report. Each section should provide thorough analysis with specific examples and stats. Target 1000-1500 words total.",
    },
    "deep_dive": {
        "label": "Deep Dive",
        "max_tokens": 8000,
        "instruction": "Produce an exhaustive deep-dive report. Leave no stone unturned. Provide detailed analysis in every section with multiple data points, comparisons, and tactical breakdowns. Target 2000-3000 words total.",
    },
}

CUSTOM_COMPARISON_MODES = {
    "league_peers": {
        "label": "Compare to League Peers",
        "instruction": "Include a PEER_COMPARISON section comparing this player against the top performers at their position in the same league. Use percentile rankings where possible.",
    },
    "age_group": {
        "label": "Compare to Age Group",
        "instruction": "Include a PEER_COMPARISON section comparing this player against peers in the same age group (birth year cohort). Evaluate where they stand on the development curve relative to same-age players.",
    },
    "previous_season": {
        "label": "Compare to Previous Season",
        "instruction": "Include a TREND_ANALYSIS section comparing current season performance to the previous season. Highlight improvements, regressions, and trajectory changes with specific stat deltas.",
    },
}


def _get_team_system(conn, org_id: str, team_name: str) -> Optional[dict]:
    """Fetch and parse a team's system profile."""
    ts_row = conn.execute(
        "SELECT * FROM team_systems WHERE org_id = ? AND LOWER(team_name) = LOWER(?) LIMIT 1",
        (org_id, team_name)
    ).fetchone()
    if not ts_row:
        return None
    return {
        "team_name": ts_row["team_name"],
        "season": ts_row["season"],
        "forecheck": ts_row["forecheck"],
        "dz_structure": ts_row["dz_structure"],
        "oz_setup": ts_row["oz_setup"],
        "pp_formation": ts_row["pp_formation"],
        "pk_formation": ts_row["pk_formation"],
        "breakout": ts_row["breakout"],
        "pace": ts_row["pace"] or "",
        "physicality": ts_row["physicality"] or "",
        "offensive_style": ts_row["offensive_style"] or "",
        "identity_tags": json.loads(ts_row["identity_tags"]) if ts_row["identity_tags"] else [],
        "notes": ts_row["notes"] or "",
    }


def _resolve_system_code(conn, code: str) -> str:
    """Look up a system code in the library and return description."""
    if not code:
        return "Not specified"
    lib_row = conn.execute(
        "SELECT name, description FROM systems_library WHERE code = ? LIMIT 1", (code,)
    ).fetchone()
    if lib_row:
        return f"{lib_row['name']} — {lib_row['description']}"
    return code


def _build_system_context_block(conn, team_system: dict) -> str:
    """Build the TEAM SYSTEM CONTEXT prompt block from a team system dict."""
    fc_desc = _resolve_system_code(conn, team_system.get('forecheck', ''))
    dz_desc = _resolve_system_code(conn, team_system.get('dz_structure', ''))
    oz_desc = _resolve_system_code(conn, team_system.get('oz_setup', ''))
    pp_desc = _resolve_system_code(conn, team_system.get('pp_formation', ''))
    pk_desc = _resolve_system_code(conn, team_system.get('pk_formation', ''))
    bo_desc = _resolve_system_code(conn, team_system.get('breakout', ''))

    return f"""
TEAM SYSTEM CONTEXT — HOCKEY OPERATING SYSTEM:
Team: {team_system['team_name']} ({team_system.get('season', 'Current')})
Team Identity: {', '.join(team_system.get('identity_tags', [])) or 'Not specified'}
Team Notes: {team_system.get('notes', 'None')}

TACTICAL SYSTEMS:
- Forecheck: {fc_desc}
- Defensive Zone Coverage: {dz_desc}
- Offensive Zone Setup: {oz_desc}
- Power Play Formation: {pp_desc}
- Penalty Kill Formation: {pk_desc}
- Breakout Pattern: {bo_desc}

TEAM STYLE:
- Pace: {team_system.get('pace', 'Not specified')}
- Physicality: {team_system.get('physicality', 'Not specified')}
- Offensive Style: {team_system.get('offensive_style', 'Not specified')}
"""


async def _generate_custom_report(request, org_id: str, user_id: str, conn):
    """Generate a fully customizable report built from user-selected focus areas, audience, depth, and comparison modes."""
    scope = request.data_scope or {}
    focus_areas = scope.get("focus_areas", [])
    audience = scope.get("audience", "general")
    depth = scope.get("depth", "standard")
    comparison_mode = scope.get("comparison_mode", "")
    custom_instructions = (scope.get("custom_instructions", "") or "")[:500]
    report_title = scope.get("report_title", "")

    is_team = bool(request.team_name and not request.player_id)

    # Create report record
    report_id = gen_id()

    if is_team:
        # ── TEAM CUSTOM REPORT ──
        team_name = request.team_name

        conn.execute("""
            INSERT INTO reports (id, org_id, player_id, team_name, template_id, report_type, status, input_data, created_by, created_at)
            VALUES (?, ?, NULL, ?, NULL, 'custom', 'processing', ?, ?, ?)
        """, (report_id, org_id, team_name, json.dumps(scope), user_id, now_iso()))
        conn.commit()

        title = report_title or f"Custom Report — {team_name}"
        start_time = time.perf_counter()

        try:
            client = get_anthropic_client()

            # Gather roster
            roster_rows = conn.execute(
                "SELECT * FROM players WHERE org_id = ? AND LOWER(current_team) = LOWER(?) ORDER BY position, last_name",
                (org_id, team_name),
            ).fetchall()
            roster = [_player_from_row(r) for r in roster_rows]

            roster_with_stats = []
            for p in roster:
                stats_rows = conn.execute(
                    "SELECT * FROM player_stats WHERE player_id = ? ORDER BY season DESC LIMIT 5",
                    (p["id"],)
                ).fetchall()
                p_stats = []
                for sr in stats_rows:
                    p_stats.append({
                        "season": sr["season"], "gp": sr["gp"], "g": sr["g"], "a": sr["a"], "p": sr["p"],
                        "plus_minus": sr["plus_minus"], "pim": sr["pim"], "shots": sr["shots"], "sog": sr["sog"],
                    })
                roster_with_stats.append({"player": p, "stats": p_stats})

            team_system = _get_team_system(conn, org_id, team_name)
            system_context = _build_system_context_block(conn, team_system) if team_system else ""

            input_data = {"team_name": team_name, "roster": roster_with_stats}
            if team_system:
                input_data["team_system"] = team_system

            # Build custom prompt
            focus_prompt_parts = []
            required_sections = ["EXECUTIVE_SUMMARY", "BOTTOM_LINE"]
            for fa in focus_areas:
                cfg = CUSTOM_FOCUS_PROMPTS.get(fa)
                if cfg:
                    focus_prompt_parts.append(cfg["prompt"])
                    required_sections.extend(cfg["sections"])

            audience_cfg = CUSTOM_AUDIENCE_PROMPTS.get(audience, CUSTOM_AUDIENCE_PROMPTS["general"])
            depth_cfg = CUSTOM_DEPTH_CONFIG.get(depth, CUSTOM_DEPTH_CONFIG["standard"])

            comparison_block = ""
            if comparison_mode:
                comp_cfg = CUSTOM_COMPARISON_MODES.get(comparison_mode)
                if comp_cfg:
                    comparison_block = f"\n{comp_cfg['instruction']}\n"

            custom_block = ""
            if custom_instructions:
                custom_block = f"\nADDITIONAL USER INSTRUCTIONS:\n{custom_instructions}\n"

            sections_str = ", ".join(sorted(set(required_sections)))

            system_prompt = f"""You are ProspectX, an elite hockey scouting intelligence engine. You produce professional-grade custom hockey analysis reports.

REPORT CONFIGURATION:
- Subject: Team — {team_name}
- Focus Areas: {', '.join(CUSTOM_FOCUS_PROMPTS.get(fa, {}).get('label', fa) for fa in focus_areas) or 'Comprehensive'}
- Audience: {audience_cfg['label']}
- Depth: {depth_cfg['label']}

AUDIENCE & TONE:
{audience_cfg['tone']}

DEPTH INSTRUCTION:
{depth_cfg['instruction']}

{system_context}

REQUIRED SECTIONS (use ALL_CAPS_WITH_UNDERSCORES format):
Always include: EXECUTIVE_SUMMARY and BOTTOM_LINE.

{'FOCUS AREA SECTIONS:' if focus_prompt_parts else 'Provide a comprehensive team assessment covering strategy, roster, and identity.'}
{chr(10).join(focus_prompt_parts)}
{comparison_block}
{custom_block}

GRADING: Include "Overall Grade: X" in EXECUTIVE_SUMMARY using scale A (Elite) to D (Developing), NR (Not Rated).
Format each section header on its own line in ALL_CAPS_WITH_UNDERSCORES format.
Today's date is {datetime.now().date().isoformat()}."""

            # ── Gather drills for custom team report if requested ──
            custom_drill_list = _gather_drills_for_report(conn, org_id, "custom", scope, team_name=team_name)
            if custom_drill_list:
                input_data["recommended_drills"] = custom_drill_list
                system_prompt += DRILL_REPORT_PROMPT_SECTION

            user_prompt = f"Generate a custom team analysis report for {team_name}. Here is ALL available data:\n\n" + json.dumps(input_data, indent=2, default=str)

            if client:
                llm_model = "claude-sonnet-4-20250514"
                extra_tokens = 2000 if custom_drill_list else 0
                message = client.messages.create(
                    model=llm_model,
                    max_tokens=depth_cfg["max_tokens"] + extra_tokens,
                    system=system_prompt,
                    messages=[{"role": "user", "content": user_prompt}],
                )
                output_text = message.content[0].text
                total_tokens = message.usage.input_tokens + message.usage.output_tokens

                # Validate custom team report output
                output_text, validation_warnings = _validate_and_repair_report(
                    output_text, "custom_team", client, llm_model
                )
                if validation_warnings:
                    for w in validation_warnings:
                        logger.warning("Custom team report %s validation: %s", report_id, w)
            else:
                llm_model = "mock-demo"
                total_tokens = 0
                output_text = f"""EXECUTIVE_SUMMARY\nCustom team report for {team_name}. Focus: {', '.join(focus_areas)}. Audience: {audience}.\n\nOverall Grade: NR\n\nBOTTOM_LINE\nThis is a mock custom report. Set your Anthropic API key to generate real analysis."""

            generation_ms = int((time.perf_counter() - start_time) * 1000)

            # Score quality
            quality = _score_report_quality(output_text, "custom_team")
            conn.execute("""
                UPDATE reports SET status='complete', title=?, output_text=?, generated_at=?,
                                  llm_model=?, llm_tokens=?, generation_time_ms=?,
                                  quality_score=?, quality_details=?
                WHERE id = ?
            """, (title, output_text, now_iso(), llm_model, total_tokens, generation_ms,
                  quality["score"], json.dumps(quality), report_id))
            conn.commit()

            # Track report usage
            _increment_usage(user_id, "report", report_id, org_id, conn)
            conn.close()

            return ReportGenerateResponse(report_id=report_id, status="complete", title=title, generation_time_ms=generation_ms)

        except Exception as e:
            conn.execute("UPDATE reports SET status='failed', error_message=? WHERE id = ?", (str(e), report_id))
            conn.commit()
            conn.close()
            raise HTTPException(status_code=500, detail=f"Custom report generation failed: {str(e)}")

    # ── PLAYER CUSTOM REPORT ──
    player_row = conn.execute("SELECT * FROM players WHERE id = ? AND org_id = ?", (request.player_id, org_id)).fetchone()
    if not player_row:
        conn.close()
        raise HTTPException(status_code=404, detail="Player not found")

    player = _player_from_row(player_row)
    player_name = f"{player['first_name']} {player['last_name']}"

    conn.execute("""
        INSERT INTO reports (id, org_id, player_id, template_id, report_type, status, input_data, created_by, created_at)
        VALUES (?, ?, ?, NULL, 'custom', 'processing', ?, ?, ?)
    """, (report_id, org_id, request.player_id, json.dumps(scope), user_id, now_iso()))
    conn.commit()

    title = report_title or f"Custom Report — {player_name}"
    start_time = time.perf_counter()

    try:
        client = get_anthropic_client()

        def _row_get(row, key, default=None):
            try:
                return row[key]
            except (IndexError, KeyError):
                return default

        # Gather stats
        stats_rows = conn.execute(
            "SELECT * FROM player_stats WHERE player_id = ? ORDER BY season DESC, created_at DESC",
            (request.player_id,),
        ).fetchall()
        stats_list = []
        for sr in stats_rows:
            stat_entry = {
                "season": sr["season"], "stat_type": sr["stat_type"],
                "gp": sr["gp"], "g": sr["g"], "a": sr["a"], "p": sr["p"],
                "plus_minus": sr["plus_minus"], "pim": sr["pim"],
                "shots": sr["shots"], "sog": sr["sog"],
                "shooting_pct": sr["shooting_pct"],
                "toi_seconds": sr["toi_seconds"],
            }
            ext_raw = _row_get(sr, "extended_stats")
            if ext_raw:
                try:
                    ext = json.loads(ext_raw) if isinstance(ext_raw, str) else ext_raw
                    if ext:
                        stat_entry["extended_stats"] = ext
                        stat_entry["data_source"] = _row_get(sr, "data_source", "manual")
                except (json.JSONDecodeError, TypeError):
                    pass
            stats_list.append(stat_entry)

        # Goalie stats
        goalie_stats_list = []
        goalie_rows = conn.execute(
            "SELECT * FROM goalie_stats WHERE player_id = ? ORDER BY season DESC",
            (request.player_id,),
        ).fetchall()
        for gr in goalie_rows:
            goalie_entry = {
                "season": gr["season"], "gp": gr["gp"],
                "toi_seconds": gr["toi_seconds"],
                "ga": gr["ga"], "sa": gr["sa"], "sv": gr["sv"],
                "sv_pct": gr["sv_pct"], "gaa": gr["gaa"],
            }
            ext_raw = _row_get(gr, "extended_stats")
            if ext_raw:
                try:
                    ext = json.loads(ext_raw) if isinstance(ext_raw, str) else ext_raw
                    if ext:
                        goalie_entry["extended_stats"] = ext
                except (json.JSONDecodeError, TypeError):
                    pass
            goalie_stats_list.append(goalie_entry)

        # Scout notes
        notes_rows = conn.execute(
            "SELECT * FROM scout_notes WHERE player_id = ? AND org_id = ? ORDER BY created_at DESC LIMIT 20",
            (request.player_id, org_id),
        ).fetchall()
        notes_list = [{"date": nr["created_at"], "note_type": nr["note_type"], "content": nr["note_text"], "tags": nr["tags"]} for nr in notes_rows]

        # Team system
        team_system = None
        system_context = ""
        if player.get("current_team"):
            team_system = _get_team_system(conn, org_id, player["current_team"])
            if team_system:
                system_context = _build_system_context_block(conn, team_system)

        # Line combos
        line_combos = []
        if player.get("current_team"):
            lc_rows = conn.execute(
                "SELECT * FROM line_combinations WHERE org_id = ? AND LOWER(team_name) = LOWER(?) ORDER BY toi_seconds DESC LIMIT 30",
                (org_id, player["current_team"]),
            ).fetchall()
            for lc in lc_rows:
                player_name_lower = player_name.lower()
                pn = (lc["player_names"] or "").lower()
                if player_name_lower.split()[-1] in pn:
                    lc_entry = {
                        "line_type": lc["line_type"], "players": lc["player_names"],
                        "toi_seconds": lc["toi_seconds"], "goals_for": lc["goals_for"],
                        "goals_against": lc["goals_against"], "plus_minus": lc["plus_minus"],
                    }
                    ext_raw = _row_get(lc, "extended_stats")
                    if ext_raw:
                        try:
                            ext = json.loads(ext_raw) if isinstance(ext_raw, str) else ext_raw
                            if ext:
                                lc_entry["extended_stats"] = ext
                        except (json.JSONDecodeError, TypeError):
                            pass
                    line_combos.append(lc_entry)

        # Pre-compute age
        player_for_report = dict(player)
        if player.get("dob"):
            try:
                dob_str = player["dob"][:10]
                birth = datetime.strptime(dob_str, "%Y-%m-%d").date()
                today = datetime.now().date()
                age = today.year - birth.year - ((today.month, today.day) < (birth.month, birth.day))
                player_for_report["age"] = age
                player_for_report["age_note"] = f"Born {dob_str} — currently {age} years old (as of {today.isoformat()})"
            except Exception:
                pass

        input_data = {"player": player_for_report, "stats": stats_list, "scout_notes": notes_list}
        if goalie_stats_list:
            input_data["goalie_stats"] = goalie_stats_list
        if line_combos:
            input_data["line_combinations"] = line_combos
        if team_system:
            input_data["team_system"] = team_system

        # Intelligence data
        try:
            indices = _compute_prospectx_indices(
                stats_list[0] if stats_list else {},
                player.get("position", ""),
                conn.execute("SELECT * FROM player_stats WHERE org_id = ?", (org_id,)).fetchall()
            )
            input_data["prospectx_indices"] = indices
        except Exception:
            pass

        intel_row = conn.execute(
            "SELECT * FROM player_intelligence WHERE player_id = ? AND org_id = ? ORDER BY version DESC LIMIT 1",
            (request.player_id, org_id)
        ).fetchone()
        if intel_row:
            intel = dict(intel_row)
            for k in ("strengths", "development_areas", "comparable_players", "tags"):
                if isinstance(intel.get(k), str):
                    try: intel[k] = json.loads(intel[k])
                    except Exception: intel[k] = []
            if isinstance(intel.get("stat_signature"), str):
                try: intel["stat_signature"] = json.loads(intel["stat_signature"])
                except Exception: intel["stat_signature"] = {}
            input_data["intelligence"] = intel

        # Build custom prompt from selections
        focus_prompt_parts = []
        required_sections = ["EXECUTIVE_SUMMARY", "BOTTOM_LINE"]
        for fa in focus_areas:
            cfg = CUSTOM_FOCUS_PROMPTS.get(fa)
            if cfg:
                focus_prompt_parts.append(cfg["prompt"])
                required_sections.extend(cfg["sections"])

        audience_cfg = CUSTOM_AUDIENCE_PROMPTS.get(audience, CUSTOM_AUDIENCE_PROMPTS["general"])
        depth_cfg = CUSTOM_DEPTH_CONFIG.get(depth, CUSTOM_DEPTH_CONFIG["standard"])

        comparison_block = ""
        if comparison_mode:
            comp_cfg = CUSTOM_COMPARISON_MODES.get(comparison_mode)
            if comp_cfg:
                comparison_block = f"\n{comp_cfg['instruction']}\n"

        custom_block = ""
        if custom_instructions:
            custom_block = f"\nADDITIONAL USER INSTRUCTIONS:\n{custom_instructions}\n"

        has_extended = any(s.get("extended_stats") for s in stats_list)

        system_prompt = f"""You are ProspectX, an elite hockey scouting intelligence engine powered by the Hockey Operating System. You produce professional-grade custom scouting reports.

REPORT CONFIGURATION:
- Player: {player_name} ({player.get('position', 'Unknown')}) — {player.get('current_team', 'Unknown Team')}
- Focus Areas: {', '.join(CUSTOM_FOCUS_PROMPTS.get(fa, {}).get('label', fa) for fa in focus_areas) or 'Comprehensive Assessment'}
- Audience: {audience_cfg['label']}
- Depth: {depth_cfg['label']}

AUDIENCE & TONE:
{audience_cfg['tone']}

DEPTH INSTRUCTION:
{depth_cfg['instruction']}

DATA AVAILABLE: {len(stats_list)} stat rows {'(with extended analytics)' if has_extended else ''}, {len(notes_list)} scout notes, {len(goalie_stats_list)} goalie stat rows, {len(line_combos)} line combinations.
{'ProspectX Intelligence profile available.' if input_data.get('intelligence') else ''}
{'ProspectX Metrics scores available.' if input_data.get('prospectx_indices') else ''}

{system_context}

{'FOCUS AREA SECTIONS — Generate each of the following:' if focus_prompt_parts else 'Provide a comprehensive player assessment.'}
{chr(10).join(focus_prompt_parts)}
{comparison_block}
{custom_block}

PROSPECT GRADING SCALE (include an Overall Grade in EXECUTIVE_SUMMARY):
  A = Top-Line / #1 D / Franchise | A- = Top-6 F / Top-4 D
  B+ = Middle-6 F / Top-4 D | B = Bottom-6 / Backup G
  B- = NHL Fringe / AHL Top | C+ = AHL Regular
  C = AHL Depth | C- = ECHL / Junior | D = Junior / College | NR = Insufficient Data

MANDATORY: Include "Overall Grade: X" in EXECUTIVE_SUMMARY on its own line.
Format each section header on its own line in ALL_CAPS_WITH_UNDERSCORES format.
When extended analytics are provided (xG, CORSI, puck battles, zone entries), leverage these advanced metrics.
Hockey vernacular: Use authentic hockey language — call a goal-scorer a sniper, describe a physical forward as a grinder or mucker, reference PP/PK roles (bumper, flank, QB, net-front), use archetype terms (power forward, puck-moving D, stay-at-home D, two-way center). Describe special teams fit using formation names (1-3-1, umbrella, diamond PK).
Age-accurate: Use the provided "age" field. Today's date is {datetime.now().date().isoformat()}.
If data is limited for any focus area, note what additional data would strengthen the analysis rather than fabricating observations."""

        # ── Gather drills for custom player report if requested ──
        custom_p_drill_list = _gather_drills_for_report(conn, org_id, "custom", scope, team_name=player.get("current_team"))
        if custom_p_drill_list:
            input_data["recommended_drills"] = custom_p_drill_list
            system_prompt += DRILL_REPORT_PROMPT_SECTION

        user_prompt = f"Generate a custom scouting report for {player_name}. Here is ALL available data:\n\n" + json.dumps(input_data, indent=2, default=str)

        if client:
            llm_model = "claude-sonnet-4-20250514"
            extra_tokens = 2000 if custom_p_drill_list else 0
            message = client.messages.create(
                model=llm_model,
                max_tokens=depth_cfg["max_tokens"] + extra_tokens,
                system=system_prompt,
                messages=[{"role": "user", "content": user_prompt}],
            )
            output_text = message.content[0].text
            total_tokens = message.usage.input_tokens + message.usage.output_tokens

            # Validate custom player report output
            output_text, validation_warnings = _validate_and_repair_report(
                output_text, "custom_player", client, llm_model
            )
            if validation_warnings:
                for w in validation_warnings:
                    logger.warning("Custom player report %s validation: %s", report_id, w)
        else:
            llm_model = "mock-demo"
            total_tokens = 0
            output_text = f"""EXECUTIVE_SUMMARY\nCustom scouting report for {player_name}. Focus: {', '.join(focus_areas)}. Audience: {audience}.\n\nOverall Grade: NR\n\nBOTTOM_LINE\nThis is a mock custom report. Set your Anthropic API key to generate real analysis."""

        generation_ms = int((time.perf_counter() - start_time) * 1000)

        # Score quality
        quality = _score_report_quality(output_text, "custom_player")
        conn.execute("""
            UPDATE reports SET status='complete', title=?, output_text=?, generated_at=?,
                              llm_model=?, llm_tokens=?, generation_time_ms=?,
                              quality_score=?, quality_details=?
            WHERE id = ?
        """, (title, output_text, now_iso(), llm_model, total_tokens, generation_ms,
              quality["score"], json.dumps(quality), report_id))
        conn.commit()
        # Track report usage
        _increment_usage(user_id, "report", report_id, org_id, conn)
        conn.close()

        # Trigger intelligence refresh
        if request.player_id:
            asyncio.create_task(_extract_report_intelligence(report_id, request.player_id, org_id))

        return ReportGenerateResponse(report_id=report_id, status="complete", title=title, generation_time_ms=generation_ms)

    except Exception as e:
        conn.execute("UPDATE reports SET status='failed', error_message=? WHERE id = ?", (str(e), report_id))
        conn.commit()
        conn.close()
        raise HTTPException(status_code=500, detail=f"Custom report generation failed: {str(e)}")


async def _generate_team_report(request, org_id: str, user_id: str, conn):
    """Generate a team-level report (team_identity, opponent_gameplan, etc.)."""
    team_name = request.team_name

    # Get template
    template = conn.execute(
        "SELECT * FROM report_templates WHERE report_type = ? AND (org_id = ? OR is_global = 1) LIMIT 1",
        (request.report_type, org_id),
    ).fetchone()
    if not template:
        conn.close()
        raise HTTPException(status_code=404, detail="Report template not found")

    # Create report record — player_id is NULL for team reports
    report_id = gen_id()
    conn.execute("""
        INSERT INTO reports (id, org_id, player_id, team_name, template_id, report_type, status, input_data, created_by, created_at)
        VALUES (?, ?, NULL, ?, ?, ?, 'processing', ?, ?, ?)
    """, (report_id, org_id, team_name, template["id"], request.report_type, json.dumps(request.data_scope or {}), user_id, now_iso()))
    conn.commit()

    start_time = time.perf_counter()
    title = f"{template['template_name']} — {team_name}"

    try:
        client = get_anthropic_client()

        # Gather roster data
        roster_rows = conn.execute(
            "SELECT * FROM players WHERE org_id = ? AND LOWER(current_team) = LOWER(?) ORDER BY position, last_name",
            (org_id, team_name),
        ).fetchall()
        roster = [_player_from_row(r) for r in roster_rows]

        # Gather stats for each rostered player
        roster_with_stats = []
        for p in roster:
            stats_rows = conn.execute(
                "SELECT * FROM player_stats WHERE player_id = ? ORDER BY season DESC LIMIT 5",
                (p["id"],)
            ).fetchall()
            p_stats = []
            for sr in stats_rows:
                p_stats.append({
                    "season": sr["season"], "stat_type": sr["stat_type"],
                    "gp": sr["gp"], "g": sr["g"], "a": sr["a"], "p": sr["p"],
                    "plus_minus": sr["plus_minus"], "pim": sr["pim"],
                })
            p_entry = {
                "name": f"{p['first_name']} {p['last_name']}",
                "position": p["position"],
                "archetype": p.get("archetype", ""),
                "shoots": p.get("shoots", ""),
                "dob": p.get("dob", ""),
                "stats": p_stats,
            }
            # Pre-compute age
            if p.get("dob"):
                try:
                    dob_str = p["dob"][:10]
                    birth = datetime.strptime(dob_str, "%Y-%m-%d").date()
                    today = datetime.now().date()
                    p_entry["age"] = today.year - birth.year - ((today.month, today.day) < (birth.month, birth.day))
                except Exception:
                    pass
            roster_with_stats.append(p_entry)

        # Gather team system
        team_system = _get_team_system(conn, org_id, team_name)
        system_context = _build_system_context_block(conn, team_system) if team_system else "\nNo team system profile configured. The report should note this limitation.\n"

        report_type_name = template["template_name"]
        input_data = {
            "team_name": team_name,
            "roster": roster_with_stats,
            "team_system": team_system,
            "report_type": request.report_type,
        }

        # ── Enrich team-report input_data for special report types ──
        if request.report_type in ("league_benchmarks", "season_projection", "free_agent_market"):
            # Get the team's league from roster or team system
            team_league = None
            if roster_with_stats:
                for rp in roster_with_stats:
                    if rp.get("current_league"):
                        team_league = rp["current_league"]
                        break
            if not team_league and team_system:
                team_league = team_system.get("league")

            # Get league-wide stats for all teams
            if team_league:
                league_players = conn.execute("""
                    SELECT p.first_name, p.last_name, p.position, p.current_team, p.birth_year, p.age_group,
                           ps.gp, ps.g, ps.a, ps.p, ps.plus_minus, ps.pim, ps.season
                    FROM players p
                    JOIN player_stats ps ON p.id = ps.player_id
                    WHERE p.org_id = ? AND LOWER(p.current_league) = LOWER(?)
                    ORDER BY ps.p DESC
                """, (org_id, team_league)).fetchall()
                input_data["league_data"] = {
                    "league_name": team_league,
                    "all_players": [dict(r) for r in league_players[:100]],
                    "total_players_in_league": len(league_players),
                }

                # Get all teams with stats for league rankings
                team_aggs = conn.execute("""
                    SELECT p.current_team, COUNT(DISTINCT p.id) as roster_size,
                           SUM(ps.g) as total_goals, SUM(ps.a) as total_assists, SUM(ps.p) as total_points,
                           AVG(CASE WHEN ps.gp > 0 THEN CAST(ps.p AS FLOAT) / ps.gp END) as avg_ppg,
                           AVG(ps.plus_minus) as avg_pm
                    FROM players p
                    JOIN player_stats ps ON p.id = ps.player_id
                    WHERE p.org_id = ? AND LOWER(p.current_league) = LOWER(?) AND p.current_team IS NOT NULL
                    GROUP BY p.current_team
                    ORDER BY total_points DESC
                """, (org_id, team_league)).fetchall()
                input_data["league_data"]["team_rankings"] = [dict(r) for r in team_aggs]

            # For free agent market — include intelligence data
            if request.report_type == "free_agent_market":
                intel_rows = conn.execute("""
                    SELECT pi.player_id, pi.archetype, pi.overall_grade, pi.summary,
                           p.first_name, p.last_name, p.position, p.current_team, p.birth_year, p.age_group
                    FROM player_intelligence pi
                    JOIN players p ON pi.player_id = p.id
                    WHERE pi.org_id = ? AND pi.version = (SELECT MAX(pi2.version) FROM player_intelligence pi2 WHERE pi2.player_id = pi.player_id)
                    ORDER BY pi.overall_grade
                """, (org_id,)).fetchall()
                input_data["player_intelligence_summary"] = [dict(r) for r in intel_rows[:50]]

        if client:
            llm_model = "claude-sonnet-4-20250514"

            system_prompt = f"""You are ProspectX, an elite hockey scouting intelligence engine powered by the Hockey Operating System. You produce professional-grade team strategy and scouting reports using tactical hockey terminology that NHL coaches, junior hockey GMs, and hockey operations staff expect.

Generate a **{report_type_name}** for the team below. Your report must be:
- Data-driven: Reference roster composition, player stats, and team system when available
- Tactically literate: Use real hockey language — forecheck structures, defensive zone coverage, breakout patterns, special teams formations, line matching
- Special teams detail: Reference specific PP formations (1-3-1, Overload, Umbrella) and PK systems (Diamond, Box, Aggressive) by name. Identify personnel fits for each role (QB/point, flank shooter, bumper, net-front, F1 high pressure). Diagnose common PP/PK breakdowns (overpassing, no net-front, poor spacing).
- Tactical vocabulary: Use forecheck labels (1-2-2, 2-1-2, 1-3-1), breakout names (standard, reverse, wheel), DZ structures (man-to-man, zone, collapsing box, swarm). Reference forecheck roles (F1/F2/F3). This is coaching-grade content.
- Professionally formatted: Use ALL_CAPS_WITH_UNDERSCORES section headers (e.g., EXECUTIVE_SUMMARY, TEAM_IDENTITY, ROSTER_ANALYSIS, TACTICAL_SYSTEMS, SPECIAL_TEAMS, STRENGTHS, WEAKNESSES, GAME_PLAN, PRACTICE_PRIORITIES, BOTTOM_LINE)
- System-aware: Reference the team's configured Hockey Operating System — their forecheck, DZ, OZ, PP, PK structures
- Coaching-grade: Write like you're briefing a coaching staff before a game or planning session. Use hockey vernacular — call players by archetype (grinder, sniper, power forward, puck-moving D), describe game flow (barnburner, chippy, gongshow), reference specific tactical situations.

{system_context}

Include actionable coaching recommendations. Reference specific players by name when discussing line combinations, deployment, or matchup strategies."""

            # ── Report-type-specific prompt enhancements for team reports ──
            if request.report_type == "league_benchmarks":
                system_prompt += f"""

SPECIAL INSTRUCTIONS FOR LEAGUE BENCHMARKS COMPARISON:
Structure your report as:
1. EXECUTIVE_SUMMARY — Team's overall standing vs league with key differentiators
2. TEAM_VS_LEAGUE — Compare this team to league averages across: Goals For, Goals Against, Points Per Game, Plus/Minus differential. Show specific numbers and percentages (e.g., "+35% above league average")
3. STATISTICAL_LEADERS — Top 10 scoring leaders in the league. Highlight OUR team's players with markers
4. POSITION_GROUP_RANKINGS — Rank this team's forward group, defense group, and goalies vs other teams in the league
5. SYSTEM_COMPARISON — Compare this team's tactical effectiveness vs league averages (if system data available)
6. TREND_ANALYSIS — What's improving vs league? What's declining? Use specific stat trends
7. COMPETITIVE_ADVANTAGES — Top 3 areas where this team outperforms the league
8. COMPETITIVE_VULNERABILITIES — Top 3 areas of concern relative to league
9. BOTTOM_LINE — Overall league standing assessment and playoff implications
Use the league_data in the input to reference REAL stats and rankings. Today's date is {datetime.now().date().isoformat()}."""

            elif request.report_type == "season_projection":
                system_prompt += f"""

SPECIAL INSTRUCTIONS FOR TEAM SEASON PROJECTION:
Structure your report as:
1. PROJECTED_STANDINGS — Project final standings for ALL teams in the league based on current pace and roster strength. Include W-L-OTL-Points columns and playoff probability percentages
2. CHAMPIONSHIP_ODDS — Top 5 championship contenders with probability percentages
3. TEAM_DETAILED_PROJECTION — Deep dive on THIS team: projected record, current pace extrapolation, strengths driving the projection, risks to the projection
4. KEY_PLAYERS_PROJECTED — Top 5 players on this team with projected final season stats
5. PLAYOFF_PATH — If team projects to make playoffs: likely first-round opponent, win probability per round, championship odds
6. REMAINING_SCHEDULE — Assess schedule difficulty for remaining games
7. RISK_FACTORS — Injuries, goaltending consistency, schedule difficulty, competitive balance concerns
8. BOTTOM_LINE — Clear projection with confidence level
Base projections on the actual stats provided. Use points pace (current_points / games_played * total_season_games). Today's date is {datetime.now().date().isoformat()}. The GOHL regular season is typically 52 games."""

            elif request.report_type == "free_agent_market":
                system_prompt += f"""

SPECIAL INSTRUCTIONS FOR FREE AGENT MARKET ANALYSIS:
Structure your report as:
1. MARKET_SUMMARY — Total available players by position, birth year distribution, quality distribution (Grade A/B/C)
2. TOP_AVAILABLE_BY_POSITION — For each position group (Centers, Wingers, Defense, Goalies), list top 5 available players with stats, grade, and system fit assessment
3. TEAM_NEEDS_ANALYSIS — Based on THIS team's current roster composition, identify position gaps and priorities
4. RECOMMENDED_TARGETS — Top 3-5 recommended acquisition targets with: player profile, stats, system fit rating (Excellent/Good/Moderate/Poor), market value assessment, and recommendation (PURSUE/MONITOR/PASS)
5. MARKET_TRENDS — Which positions are scarce vs abundant? Where should the team act fast vs wait?
6. TIMING_RECOMMENDATIONS — Prioritize which positions to address immediately vs later
7. BOTTOM_LINE — Clear action items for the GM
Use intelligence grades and archetypes from the player_intelligence_summary in the input. Today's date is {datetime.now().date().isoformat()}."""

            # ── Gather recommended drills for team reports if requested ──
            team_drill_list = _gather_drills_for_report(conn, org_id, request.report_type, request.data_scope, team_name=team_name)
            if team_drill_list:
                input_data["recommended_drills"] = team_drill_list
                system_prompt += DRILL_REPORT_PROMPT_SECTION

            user_prompt = f"Generate a {report_type_name} for {team_name}. Here is all available data:\n\n" + json.dumps(input_data, indent=2, default=str)

            max_tokens = 10000 if team_drill_list else 8000
            message = client.messages.create(
                model=llm_model,
                max_tokens=max_tokens,
                system=system_prompt,
                messages=[{"role": "user", "content": user_prompt}],
            )
            output_text = message.content[0].text
            total_tokens = message.usage.input_tokens + message.usage.output_tokens

            # Validate team report output (mode-aware)
            _team_mode = resolve_mode(template_slug=request.report_type) if request.report_type else None
            output_text, validation_warnings = _validate_and_repair_report(
                output_text, request.report_type, client, llm_model, mode=getattr(request, 'mode', None) or _team_mode
            )
            if validation_warnings:
                for w in validation_warnings:
                    logger.warning("Team report %s validation: %s", report_id, w)
        else:
            llm_model = "mock-demo"
            total_tokens = 0
            roster_summary = ", ".join([f"{p['name']} ({p['position']})" for p in roster_with_stats[:10]])
            output_text = f"""EXECUTIVE_SUMMARY
{team_name} — {report_type_name}. This is a demo-mode team report. Roster: {len(roster_with_stats)} players. Connect your Anthropic API key in backend/.env to generate full AI-powered team intelligence.

ROSTER_ANALYSIS
Players on file: {roster_summary or 'No players found for this team.'}
Total roster size: {len(roster_with_stats)} players.

TACTICAL_SYSTEMS
{'Team system profile is configured.' if team_system else 'No team system profile found. Configure one in the Systems tab to unlock full tactical analysis.'}

BOTTOM_LINE
This report was generated in demo mode. Add your Anthropic API key to backend/.env and re-generate for full team intelligence.

**To unlock full report intelligence:** Add your Anthropic API key to backend/.env and re-generate this report.
"""
            logger.info("No Anthropic API key — generated mock team report for %s", team_name)

        generation_ms = int((time.perf_counter() - start_time) * 1000)

        # Score quality
        _team_rpt_mode = getattr(request, 'mode', None) or resolve_mode(template_slug=request.report_type)
        quality = _score_report_quality(output_text, request.report_type, mode=_team_rpt_mode)
        conn.execute("""
            UPDATE reports SET status='complete', title=?, output_text=?, generated_at=?,
                              llm_model=?, llm_tokens=?, generation_time_ms=?,
                              quality_score=?, quality_details=?
            WHERE id = ?
        """, (title, output_text, now_iso(), llm_model, total_tokens, generation_ms,
              quality["score"], json.dumps(quality), report_id))
        conn.commit()

        logger.info("Team report generated: %s (%s) in %d ms, quality=%.1f", title, request.report_type, generation_ms, quality["score"])
        # Track report usage
        _increment_usage(user_id, "report", report_id, org_id, conn)
        conn.close()
        return ReportGenerateResponse(report_id=report_id, status="complete", title=title, generation_time_ms=generation_ms)

    except Exception as e:
        conn.execute("UPDATE reports SET status='failed', error_message=? WHERE id = ?", (str(e), report_id))
        conn.commit()
        conn.close()
        logger.error("Team report generation failed: %s — %s", report_id, str(e))
        raise HTTPException(status_code=500, detail=f"Report generation failed: {str(e)}")


@app.get("/reports/custom-options")
async def get_custom_report_options(token_data: dict = Depends(verify_token)):
    """Return the available configuration options for custom report builder."""
    return {
        "focus_areas": [
            {"key": k, "label": v["label"], "sections": v["sections"]}
            for k, v in CUSTOM_FOCUS_PROMPTS.items()
        ],
        "audiences": [
            {"key": k, "label": v["label"]}
            for k, v in CUSTOM_AUDIENCE_PROMPTS.items()
        ],
        "depths": [
            {"key": k, "label": v["label"]}
            for k, v in CUSTOM_DEPTH_CONFIG.items()
        ],
        "comparison_modes": [
            {"key": k, "label": v["label"]}
            for k, v in CUSTOM_COMPARISON_MODES.items()
        ],
    }


# ── Drill Gathering Helper for Reports ──────────────────────────
def _gather_drills_for_report(conn, org_id: str, report_type: str, data_scope: dict | None = None, team_name: str | None = None) -> list[dict]:
    """Gather relevant drills with diagram URLs for inclusion in a report.
    Returns a list of drill dicts with name, category, description, setup, coaching_points, diagram_url, intensity, age_levels, duration_minutes.
    Filters by focus areas from data_scope if available. Limits to 8 most relevant drills."""
    scope = data_scope or {}
    include_drills = scope.get("include_drills", False)
    if not include_drills:
        return []

    # Build query conditions
    conditions = ["(org_id IS NULL OR org_id = ?)"]
    params: list = [org_id]

    # Filter by drill categories matching report focus
    drill_focus = scope.get("drill_focus", [])  # e.g. ["offensive", "skating", "systems"]
    if drill_focus:
        placeholders = ",".join(["?" for _ in drill_focus])
        conditions.append(f"category IN ({placeholders})")
        params.extend(drill_focus)

    # Filter by age level if specified
    drill_age_level = scope.get("drill_age_level", "")
    if drill_age_level:
        conditions.append("age_levels LIKE ?")
        params.append(f'%"{drill_age_level}"%')

    # Filter by intensity if specified
    drill_intensity = scope.get("drill_intensity", "")
    if drill_intensity:
        conditions.append("intensity = ?")
        params.append(drill_intensity)

    where = " AND ".join(conditions)
    rows = conn.execute(f"""
        SELECT id, name, category, description, setup, coaching_points,
               diagram_url, intensity, age_levels, duration_minutes, ice_surface, equipment
        FROM drills
        WHERE {where}
        ORDER BY RANDOM()
        LIMIT 8
    """, params).fetchall()

    drills = []
    for r in rows:
        d = {
            "name": r["name"],
            "category": r["category"],
            "description": r["description"],
            "setup": r["setup"] or "",
            "coaching_points": r["coaching_points"] or "",
            "intensity": r["intensity"],
            "duration_minutes": r["duration_minutes"],
            "ice_surface": r["ice_surface"],
            "equipment": r["equipment"] or "",
        }
        # Include diagram URL so Claude can reference it in output
        if r["diagram_url"]:
            d["diagram_url"] = r["diagram_url"]
        age_raw = r["age_levels"]
        if age_raw:
            try:
                d["age_levels"] = json.loads(age_raw) if isinstance(age_raw, str) else age_raw
            except (json.JSONDecodeError, TypeError):
                d["age_levels"] = []
        drills.append(d)

    return drills


DRILL_REPORT_PROMPT_SECTION = """

DRILL RECOMMENDATIONS WITH DIAGRAMS:
The input data includes recommended drills from the ProspectX Drill Library. For EACH drill in the recommended_drills list, you MUST include it in a RECOMMENDED_DRILLS section with:
1. The drill name as a bold header
2. An image reference using this EXACT format: ![Drill: <drill_name>](<diagram_url>) — this will render the rink diagram
3. The setup instructions
4. Key coaching points
5. How this drill addresses the player's/team's specific development needs

Format each drill entry like this:
**<Drill Name>** (<duration> min, <intensity> intensity)
![Drill: <drill_name>](<diagram_url>)
**Setup:** <setup text>
**Coaching Points:** <coaching points>
**Why This Drill:** <1-2 sentences connecting this drill to the player/team's needs>

Include ALL drills from the recommended_drills data. This section should be actionable for coaches — they should be able to take this report to the rink and run these drills immediately."""


@app.post("/reports/generate", response_model=ReportGenerateResponse)
async def generate_report(request: ReportGenerateRequest, token_data: dict = Depends(verify_token)):
    org_id = token_data["org_id"]
    user_id = token_data["user_id"]

    # ── Report usage limit check ──
    limit_conn = get_db()
    try:
        _check_tier_limit(user_id, "reports", limit_conn)
    finally:
        limit_conn.close()

    if not request.player_id and not request.team_name:
        raise HTTPException(status_code=400, detail="Either player_id or team_name is required")

    is_team_report = request.report_type in TEAM_REPORT_TYPES and request.team_name and not request.player_id
    conn = get_db()

    # ── CUSTOM REPORT FLOW ────────────────────────────────
    if request.report_type == "custom":
        return await _generate_custom_report(request, org_id, user_id, conn)

    # ── TEAM REPORT FLOW ──────────────────────────────────
    if is_team_report:
        return await _generate_team_report(request, org_id, user_id, conn)

    # ── PLAYER REPORT FLOW (existing) ─────────────────────
    # Verify player
    player_row = conn.execute("SELECT * FROM players WHERE id = ? AND org_id = ?", (request.player_id, org_id)).fetchone()
    if not player_row:
        conn.close()
        raise HTTPException(status_code=404, detail="Player not found")

    player = _player_from_row(player_row)

    # Get template
    template = conn.execute(
        "SELECT * FROM report_templates WHERE report_type = ? AND (org_id = ? OR is_global = 1) LIMIT 1",
        (request.report_type, org_id),
    ).fetchone()
    if not template:
        conn.close()
        raise HTTPException(status_code=404, detail="Report template not found")

    # Create report record
    report_id = gen_id()
    conn.execute("""
        INSERT INTO reports (id, org_id, player_id, template_id, report_type, status, input_data, created_by, created_at)
        VALUES (?, ?, ?, ?, ?, 'processing', ?, ?, ?)
    """, (report_id, org_id, request.player_id, template["id"], request.report_type, json.dumps(request.data_scope or {}), user_id, now_iso()))
    conn.commit()

    start_time = time.perf_counter()
    player_name = f"{player['first_name']} {player['last_name']}"
    title = f"{template['template_name']} — {player_name}"

    try:
        client = get_anthropic_client()
        if client:
            llm_model = "claude-sonnet-4-20250514"

            # Helper: safe get from sqlite3.Row (which doesn't support .get())
            def _row_get(row, key, default=None):
                try:
                    return row[key]
                except (IndexError, KeyError):
                    return default

            # Gather player stats (including extended InStat analytics)
            stats_rows = conn.execute(
                "SELECT * FROM player_stats WHERE player_id = ? ORDER BY season DESC, created_at DESC",
                (request.player_id,),
            ).fetchall()
            stats_list = []
            for sr in stats_rows:
                stat_entry = {
                    "season": sr["season"], "stat_type": sr["stat_type"],
                    "gp": sr["gp"], "g": sr["g"], "a": sr["a"], "p": sr["p"],
                    "plus_minus": sr["plus_minus"], "pim": sr["pim"],
                    "shots": sr["shots"], "sog": sr["sog"],
                    "shooting_pct": sr["shooting_pct"],
                    "toi_seconds": sr["toi_seconds"],
                }
                # Include extended stats if available (xG, CORSI, puck battles, entries, etc.)
                ext_raw = _row_get(sr, "extended_stats")
                if ext_raw:
                    try:
                        ext = json.loads(ext_raw) if isinstance(ext_raw, str) else ext_raw
                        if ext:
                            stat_entry["extended_stats"] = ext
                            stat_entry["data_source"] = _row_get(sr, "data_source", "manual")
                    except (json.JSONDecodeError, TypeError):
                        pass
                stats_list.append(stat_entry)

            # Gather goalie stats if player is a goalie
            goalie_stats_list = []
            goalie_rows = conn.execute(
                "SELECT * FROM goalie_stats WHERE player_id = ? ORDER BY season DESC",
                (request.player_id,),
            ).fetchall()
            for gr in goalie_rows:
                goalie_entry = {
                    "season": gr["season"], "gp": gr["gp"],
                    "toi_seconds": gr["toi_seconds"],
                    "ga": gr["ga"], "sa": gr["sa"], "sv": gr["sv"],
                    "sv_pct": gr["sv_pct"], "gaa": gr["gaa"],
                }
                ext_raw = _row_get(gr, "extended_stats")
                if ext_raw:
                    try:
                        ext = json.loads(ext_raw) if isinstance(ext_raw, str) else ext_raw
                        if ext:
                            goalie_entry["extended_stats"] = ext
                    except (json.JSONDecodeError, TypeError):
                        pass
                goalie_stats_list.append(goalie_entry)

            # Gather scout notes
            notes_rows = conn.execute(
                "SELECT * FROM scout_notes WHERE player_id = ? AND org_id = ? ORDER BY created_at DESC LIMIT 20",
                (request.player_id, org_id),
            ).fetchall()
            notes_list = []
            for nr in notes_rows:
                notes_list.append({
                    "date": nr["created_at"],
                    "note_type": nr["note_type"],
                    "content": nr["note_text"],
                    "tags": nr["tags"],
                })

            # Gather team system data (if the player's team has a system profile)
            team_system = None
            if player.get("current_team"):
                ts_row = conn.execute(
                    "SELECT * FROM team_systems WHERE org_id = ? AND LOWER(team_name) = LOWER(?) LIMIT 1",
                    (org_id, player["current_team"])
                ).fetchone()
                if ts_row:
                    team_system = {
                        "team_name": ts_row["team_name"],
                        "season": ts_row["season"],
                        "forecheck": ts_row["forecheck"],
                        "dz_structure": ts_row["dz_structure"],
                        "oz_setup": ts_row["oz_setup"],
                        "pp_formation": ts_row["pp_formation"],
                        "pk_formation": ts_row["pk_formation"],
                        "breakout": ts_row["breakout"],
                        "pace": ts_row["pace"] or "",
                        "physicality": ts_row["physicality"] or "",
                        "offensive_style": ts_row["offensive_style"] or "",
                        "identity_tags": json.loads(ts_row["identity_tags"]) if ts_row["identity_tags"] else [],
                        "notes": ts_row["notes"] or "",
                    }

            # Resolve system codes to human-readable names from the systems library
            def resolve_system_code(code: str) -> str:
                """Look up a system code in the library and return 'Name (code)' or just the code."""
                if not code:
                    return "Not specified"
                lib_row = conn.execute(
                    "SELECT name, description, ideal_personnel FROM systems_library WHERE code = ? LIMIT 1", (code,)
                ).fetchone()
                if lib_row:
                    return f"{lib_row['name']} — {lib_row['description']}"
                return code

            # Gather line combinations for the player's team
            line_combos = []
            if player.get("current_team"):
                lc_rows = conn.execute(
                    "SELECT * FROM line_combinations WHERE org_id = ? AND LOWER(team_name) = LOWER(?) ORDER BY toi_seconds DESC LIMIT 30",
                    (org_id, player["current_team"]),
                ).fetchall()
                for lc in lc_rows:
                    player_name_lower = f"{player['first_name']} {player['last_name']}".lower()
                    pn = (lc["player_names"] or "").lower()
                    if player_name_lower.split()[-1] in pn:  # Check if player's last name in line
                        lc_entry = {
                            "line_type": lc["line_type"],
                            "players": lc["player_names"],
                            "toi_seconds": lc["toi_seconds"],
                            "goals_for": lc["goals_for"],
                            "goals_against": lc["goals_against"],
                            "plus_minus": lc["plus_minus"],
                        }
                        ext_raw = _row_get(lc, "extended_stats")
                        if ext_raw:
                            try:
                                ext = json.loads(ext_raw) if isinstance(ext_raw, str) else ext_raw
                                if ext:
                                    lc_entry["extended_stats"] = ext
                            except (json.JSONDecodeError, TypeError):
                                pass
                        line_combos.append(lc_entry)

            has_extended = any(s.get("extended_stats") for s in stats_list)
            logger.info("Report input — stats: %d rows (extended: %s), goalie_stats: %d, notes: %d, lines: %d, team_system: %s",
                        len(stats_list), "yes" if has_extended else "no", len(goalie_stats_list), len(notes_list), len(line_combos), "yes" if team_system else "no")

            # Pre-compute age so Claude doesn't have to guess
            player_for_report = dict(player)
            if player.get("dob"):
                try:
                    dob_str = player["dob"][:10]  # Handle ISO datetime or date strings
                    birth = datetime.strptime(dob_str, "%Y-%m-%d").date()
                    today = datetime.now().date()
                    age = today.year - birth.year - ((today.month, today.day) < (birth.month, birth.day))
                    player_for_report["age"] = age
                    player_for_report["age_note"] = f"Born {dob_str} — currently {age} years old (as of {today.isoformat()})"
                except Exception:
                    pass

            input_data = {
                "player": player_for_report,
                "stats": stats_list,
                "scout_notes": notes_list,
                "request_scope": request.data_scope,
            }
            if goalie_stats_list:
                input_data["goalie_stats"] = goalie_stats_list
            if line_combos:
                input_data["line_combinations"] = line_combos
            if team_system:
                input_data["team_system"] = team_system

            # Add ProspectX Indices and Intelligence data for enriched reports
            try:
                indices = _compute_prospectx_indices(
                    stats_list[0] if stats_list else {},
                    player.get("position", ""),
                    conn.execute("SELECT * FROM player_stats WHERE org_id = ?", (org_id,)).fetchall()
                )
                input_data["prospectx_indices"] = indices
            except Exception:
                pass  # Non-critical

            intel_row = conn.execute(
                "SELECT * FROM player_intelligence WHERE player_id = ? AND org_id = ? ORDER BY version DESC LIMIT 1",
                (request.player_id, org_id)
            ).fetchone()
            if intel_row:
                intel = dict(intel_row)
                for k in ("strengths", "development_areas", "comparable_players", "tags"):
                    if isinstance(intel.get(k), str):
                        try:
                            intel[k] = json.loads(intel[k])
                        except Exception:
                            intel[k] = []
                if isinstance(intel.get("stat_signature"), str):
                    try:
                        intel["stat_signature"] = json.loads(intel["stat_signature"])
                    except Exception:
                        intel["stat_signature"] = {}
                input_data["intelligence"] = intel

            # ── Gather historical progression (season-over-season) ──
            try:
                hist_rows = conn.execute("""
                    SELECT psh.*
                    FROM player_stats_history psh
                    INNER JOIN (
                        SELECT season, MAX(date_recorded) as max_date
                        FROM player_stats_history
                        WHERE player_id = ?
                        GROUP BY season
                    ) latest ON psh.season = latest.season AND psh.date_recorded = latest.max_date
                    WHERE psh.player_id = ?
                    ORDER BY psh.season ASC
                """, (request.player_id, request.player_id)).fetchall()

                if hist_rows:
                    progression_seasons = []
                    for hr in hist_rows:
                        hd = dict(hr)
                        gp = hd.get("gp", 0) or 0
                        hd["ppg_rate"] = round((hd.get("p", 0) or 0) / gp, 2) if gp > 0 else 0.0
                        progression_seasons.append(hd)
                    input_data["historical_progression"] = progression_seasons
                    logger.info("Report: added %d progression seasons for %s", len(progression_seasons), player_name)
            except Exception as e:
                logger.warning("Failed to load progression for report: %s", e)

            # ── Gather recent form (last 10 games) ──
            try:
                recent_rows = conn.execute("""
                    SELECT * FROM player_game_stats
                    WHERE player_id = ?
                    ORDER BY game_date DESC
                    LIMIT 10
                """, (request.player_id,)).fetchall()

                if not recent_rows:
                    # Fallback to player_stats game rows
                    recent_rows = conn.execute("""
                        SELECT * FROM player_stats
                        WHERE player_id = ? AND stat_type = 'game'
                        ORDER BY created_at DESC
                        LIMIT 10
                    """, (request.player_id,)).fetchall()

                if recent_rows:
                    recent_games = [dict(r) for r in recent_rows]
                    n_recent = len(recent_games)
                    total_g = sum(g.get("goals", g.get("g", 0)) or 0 for g in recent_games)
                    total_a = sum(g.get("assists", g.get("a", 0)) or 0 for g in recent_games)
                    total_p = sum(g.get("points", g.get("p", 0)) or 0 for g in recent_games)
                    input_data["recent_form_last_10"] = {
                        "games": recent_games,
                        "games_found": n_recent,
                        "totals": {"g": total_g, "a": total_a, "p": total_p},
                        "averages": {
                            "gpg": round(total_g / n_recent, 2),
                            "apg": round(total_a / n_recent, 2),
                            "ppg": round(total_p / n_recent, 2),
                        },
                    }
                    logger.info("Report: added %d recent form games for %s", n_recent, player_name)
            except Exception as e:
                logger.warning("Failed to load recent form for report: %s", e)

            # ── Gather recommended drills if requested ──
            drill_list = _gather_drills_for_report(conn, org_id, request.report_type, request.data_scope, team_name=player.get("current_team"))
            drill_prompt_addon = ""
            if drill_list:
                input_data["recommended_drills"] = drill_list
                drill_prompt_addon = DRILL_REPORT_PROMPT_SECTION

            report_type_name = template["template_name"]

            # Build the system context block for the prompt — resolve codes to full tactical descriptions
            system_context_block = ""
            if team_system:
                fc_desc = resolve_system_code(team_system.get('forecheck', ''))
                dz_desc = resolve_system_code(team_system.get('dz_structure', ''))
                oz_desc = resolve_system_code(team_system.get('oz_setup', ''))
                pp_desc = resolve_system_code(team_system.get('pp_formation', ''))
                pk_desc = resolve_system_code(team_system.get('pk_formation', ''))
                bo_desc = resolve_system_code(team_system.get('breakout', ''))

                system_context_block = f"""

TEAM SYSTEM CONTEXT — HOCKEY OPERATING SYSTEM (use this to evaluate system fit):
Team: {team_system['team_name']} ({team_system.get('season', 'Current')})
Team Identity: {', '.join(team_system.get('identity_tags', [])) or 'Not specified'}
Team Notes: {team_system.get('notes', 'None')}

TACTICAL SYSTEMS:
- Forecheck: {fc_desc}
- Defensive Zone Coverage: {dz_desc}
- Offensive Zone Setup: {oz_desc}
- Power Play Formation: {pp_desc}
- Penalty Kill Formation: {pk_desc}
- Breakout Pattern: {bo_desc}

TEAM STYLE:
- Pace: {team_system.get('pace', 'Not specified')}
- Physicality: {team_system.get('physicality', 'Not specified')}
- Offensive Style: {team_system.get('offensive_style', 'Not specified')}

SYSTEM_FIT SECTION REQUIREMENTS:
You MUST include a dedicated SYSTEM_FIT section in the report. This is a signature ProspectX feature. The SYSTEM_FIT section must:
1. Explicitly name the team's forecheck system and describe the player's role within it (F1/F2/F3 for forwards, pinch/support for D)
2. Evaluate DZ responsibilities — how this player functions in the team's defensive zone structure
3. Evaluate OZ contributions — how this player drives the team's offensive zone attack
4. Assess special teams fit — power play role and penalty kill capability
5. Rate overall system compatibility (Elite Fit / Strong Fit / Developing Fit / Adjustment Needed)
6. Provide specific tactical recommendations for coaching staff

Use elite hockey language. Write like you're briefing an NHL coaching staff. Reference specific systems by name — "As F1 on the 2-1-2 aggressive forecheck, McChesney's 18.4% shooting and +12 rating suggest..." not generic statements. This section should make a GM say "this tool UNDERSTANDS my team."
"""

            system_prompt = f"""You are ProspectX, an elite hockey scouting intelligence engine powered by the Hockey Operating System. You produce professional-grade scouting reports using tactical hockey terminology that NHL scouts, junior hockey GMs, agents, and player development staff expect.

Generate a **{report_type_name}** for the player below. Your report must be:
- Data-driven: Reference specific stats (GP, G, A, P, +/-, PIM, S%, etc.) when available. When extended analytics are provided (xG, CORSI, Fenwick, puck battles, zone entries, breakouts, scoring chances, slot shots, passes, faceoffs by zone), leverage these advanced metrics prominently in your analysis
- Tactically literate: Use real hockey language — forecheck roles (F1/F2/F3), transition play, gap control, cycle game, net-front presence, breakout patterns, DZ coverage. When CORSI/Fenwick data is available, discuss possession metrics. When xG data is available, compare expected vs actual goals
- Professionally formatted: Use ALL_CAPS_WITH_UNDERSCORES section headers (e.g., EXECUTIVE_SUMMARY, KEY_NUMBERS, STRENGTHS, DEVELOPMENT_AREAS, SYSTEM_FIT, PROJECTION, BOTTOM_LINE)
- Specific to position: Tailor analysis to the player's position (center, wing, defense, goalie)
- Honest and balanced: Don't inflate or deflate — give an accurate, scout-grade assessment
- Age-accurate: The player data includes a pre-computed "age" field and "age_note". ALWAYS use the provided age value — do NOT attempt to recalculate it from dob. Today's date is {datetime.now().date().isoformat()}
- Archetype-aware: The player's archetype may be compound (e.g., "Two-Way Playmaking Forward") indicating multiple dimensions. Analyze ALL archetype traits — if the archetype says "Two-Way Playmaking Forward" you must evaluate both the 200-foot game AND the playmaking IQ separately, then synthesize how these traits combine
- Hockey vernacular: Use authentic hockey language when describing players — call a goal-scorer a "sniper" or "trigger man," a physical player a "grinder" or "mucker," a fast player has "wheels." Reference PP/PK formation roles (bumper, flank, QB, net-front). Use forecheck roles (F1/F2/F3). Describe a player's special teams fit by naming specific formations (1-3-1 flank, umbrella point, diamond PK high man). Use archetype terms: power forward, stay-at-home D, puck-moving D, two-way center, energy forward, shutdown D.
{system_context_block}
PROGRESSION & RECENT FORM DATA (when available in input):
If the input data includes "historical_progression" (season-over-season snapshots), use it to:
- Analyze year-over-year stat trends in the PROJECTION section (is the player improving, plateauing, or declining?)
- Compare per-game rates across seasons (PPG rate, GPG rate) to identify trajectory
- Reference specific season-to-season changes (e.g., "jumped from 0.75 PPG to 1.12 PPG — a 49% improvement")
- Factor progression trends into the overall grade and recommendation

If the input data includes "recent_form_last_10" (last 10 game-by-game stats), use it to:
- Assess current hot/cold streaks and momentum
- Compare recent form to season averages (is the player trending up or down RIGHT NOW?)
- Reference specific recent game performances when available
- Inform the BOTTOM_LINE with current form context (e.g., "currently riding a 5-game point streak")

Do NOT invent progression or form data. Only reference these sections if the data is actually present in the input.

PROSPECT GRADING SCALE (include an overall grade in EXECUTIVE_SUMMARY or BOTTOM_LINE):
  A   = Top-Line / #1 D / Franchise — Elite NHL talent, first-round caliber
  A-  = Top-6 F / Top-4 D / Starting G — High-end NHL player, early-round pick
  B+  = Middle-6 F / Top-4 D — Solid NHL regular, mid-round value
  B   = Bottom-6 F / Bottom-Pair D / Backup G — NHL depth, late-round value
  B-  = NHL Fringe / AHL Top — Borderline NHL, top AHL contributor
  C+  = AHL Regular / NHL Call-Up — Strong minor-leaguer, occasional NHL fill-in
  C   = AHL Depth / ECHL Top — Professional career outside NHL
  C-  = ECHL / Junior Overager — Lower pro or strong junior
  D   = Junior / College / Non-Pro — Developing player at junior/college level
  NR  = Not Rated — Insufficient data to project

MANDATORY GRADING REQUIREMENTS:
1. You MUST include "Overall Grade: X" (where X is one of the grades above) in the EXECUTIVE_SUMMARY section. This is NOT optional.
2. Place the grade on its own line, formatted exactly as: Overall Grade: B+ (or whichever grade applies)
3. Also include "Overall Grade: X" in the RECOMMENDATION or BOTTOM_LINE section at the end.
4. Consider: age, league level, stat production relative to peers, physical tools, skating, and development trajectory.
5. EVERY report must have a grade. If data is insufficient, use "NR" and explain what's missing.

EVIDENCE DISCIPLINE:
1. EVIDENCE vs INFERENCE: Claims backed directly by stats, notes, or metrics in the input are evidence — state them as fact. Conclusions you draw that are not explicitly stated in the data must be labeled inline as: "INFERENCE — [brief reasoning]". Inferences must be reasonable hockey analysis, not fabricated facts. Example: "His 18.4% shooting percentage suggests elite finishing ability" is evidence. "INFERENCE — His quick release likely translates to PP flank deployment based on shooting tendency" is a labeled inference.
2. MISSING DATA: When a stat, metric, or observation that would normally be important for this report type is absent from the input, write "DATA NOT AVAILABLE" in that context. Use sparingly — only when the missing data materially affects the analysis. Do not list every conceivable missing field.
3. CONFIDENCE TAG: In the EXECUTIVE_SUMMARY section, end with a line: "Confidence: HIGH | MED | LOW — [reason]" based on sample size, data completeness, and scout note volume. Also include the confidence line in the BOTTOM_LINE section.
   - HIGH: 20+ GP, multiple data sources (stats + scout notes + intelligence profile), extended analytics present
   - MED: 10-19 GP or limited scout notes or missing extended analytics
   - LOW: Under 10 GP, single data source, no scout notes

Include quantitative analysis where stats exist. If data is limited, note what additional scouting data would strengthen the assessment.

Format each section header on its own line in ALL_CAPS_WITH_UNDERSCORES format, followed by the section content."""

            # ── PXI Mode injection (guardrails + mode block before base prompt) ──
            # Fetch user hockey_role for mode resolution fallback
            _user_role_row = conn.execute("SELECT hockey_role FROM users WHERE id = ?", (user_id,)).fetchone()
            _user_hockey_role = _user_role_row["hockey_role"] if _user_role_row and _user_role_row["hockey_role"] else "scout"
            _resolved_mode = resolve_mode(
                user_hockey_role=_user_hockey_role,
                explicit_mode=request.mode,
                template_slug=request.report_type,
            )
            _mode_block = PXI_MODE_BLOCKS.get(_resolved_mode, "")
            system_prompt = PXI_CORE_GUARDRAILS + "\n\n" + _mode_block + "\n\n" + system_prompt

            # ── Report-type-specific prompt enhancements ──
            if request.report_type == "indices_dashboard":
                system_prompt += """

SPECIAL INSTRUCTIONS FOR PROSPECTX METRICS DASHBOARD:
This report focuses on the player's ProspectX Metric scores. Structure your report as:
1. OVERALL_PROSPECTX_GRADE — Synthesize all metrics into a single letter grade (A through D scale) with numeric score (0-100)
2. METRICS_BREAKDOWN — For EACH of the 6 metrics (SniperMetric, PlaymakerMetric, TransitionMetric, DefensiveMetric, CompeteMetric, HockeyIQMetric), provide: the score, percentile, rating tier (Elite/Above Average/Average/Below Average/Developing), and a 2-3 sentence analysis explaining WHY the player scores at that level
3. PERCENTILE_RANKINGS — Compare vs position peers, vs league, and vs age group
4. METRIC_CORRELATION — How do the metrics work together? (e.g., high Playmaker + high IQ = elite distributor)
5. SYSTEM_FIT — Based on metrics, what role/system fits this player best?
6. DEVELOPMENT_PRIORITIES — Based on metric analysis, what 3 specific improvements would most impact their game?
7. COMPARABLE_PLAYERS — Players with similar metric profiles
Use the prospectx_indices data in the input prominently. Show specific numbers."""

            elif request.report_type == "player_projection":
                system_prompt += f"""

SPECIAL INSTRUCTIONS FOR NEXT SEASON PLAYER PROJECTION:
Project this player's performance for the NEXT season (2026-27). Structure your report as:
1. PROJECTION_SUMMARY — Expected points range (Conservative/Expected/Optimistic scenarios)
2. METHODOLOGY — Explain what data informed the projection (current stats, age curve, league level, development trend)
3. THREE_SCENARIOS — Detail Conservative (25th percentile), Expected (50th), and Optimistic (75th) with specific point/goal/assist projections
4. AGE_CURVE_ANALYSIS — How does this player's age and development stage affect projection? Reference typical age-development curves for their league tier
5. COMPARABLE_PATHWAYS — 3-5 similar players who were at this statistical level at this age, and what they did next season
6. DEVELOPMENT_FACTORS — What could boost or reduce projection (shooting improvement, ice time changes, linemate quality, etc.)
7. RECOMMENDATION — Clear actionable guidance
Use the player's birth_year and age_group from the data. Today's date is {datetime.now().date().isoformat()}. Reference specific stats when projecting."""

            # ── Append template-specific instructions from DB ──
            tpl_prompt = template["prompt_text"] if template else None
            if tpl_prompt:
                # Only append if it's a rich prompt (not the old generic one-liner)
                if len(tpl_prompt) > 200:
                    system_prompt += f"""

TEMPLATE-SPECIFIC INSTRUCTIONS FOR {report_type_name.upper()}:
{tpl_prompt}

Note: Follow the section structure specified in the template instructions above. Use ALL_CAPS_WITH_UNDERSCORES format for all section headers. The template instructions define WHAT to analyze and which sections to produce — the base instructions above define HOW to format, grade, and present evidence."""

            # ── Append drill recommendation instructions if drills were requested ──
            if drill_prompt_addon:
                system_prompt += drill_prompt_addon

            user_prompt = f"Generate a {report_type_name} for the following player. Here is ALL available data:\n\n" + json.dumps(input_data, indent=2, default=str)

            max_tokens = 10000 if drill_list else 8000
            message = client.messages.create(
                model=llm_model,
                max_tokens=max_tokens,
                system=system_prompt,
                messages=[{"role": "user", "content": user_prompt}],
            )
            output_text = message.content[0].text
            total_tokens = message.usage.input_tokens + message.usage.output_tokens

            # ── Validate and optionally repair the report output ──
            output_text, validation_warnings = _validate_and_repair_report(
                output_text, request.report_type, client, llm_model, mode=_resolved_mode
            )
            if validation_warnings:
                for w in validation_warnings:
                    logger.warning("Report %s validation: %s", report_id, w)
                    if w.startswith("SOFT:") or w.startswith("HARD:"):
                        _log_error_to_db("POST", f"/reports/generate/{report_id}", 200, f"Report validation: {w}")
        else:
            llm_model = "mock-demo"
            total_tokens = 0
            output_text = _generate_mock_report(player, request.report_type)
            logger.info("No Anthropic API key — generated mock report for %s", player_name)

        generation_ms = int((time.perf_counter() - start_time) * 1000)

        # Score quality
        quality = _score_report_quality(output_text, request.report_type, mode=_resolved_mode)
        conn.execute("""
            UPDATE reports SET status='complete', title=?, output_text=?, generated_at=?,
                              llm_model=?, llm_tokens=?, generation_time_ms=?,
                              quality_score=?, quality_details=?
            WHERE id = ?
        """, (title, output_text, now_iso(), llm_model, total_tokens, generation_ms,
              quality["score"], json.dumps(quality), report_id))
        conn.commit()

        logger.info("Report generated: %s (%s) in %d ms, quality=%.1f", title, request.report_type, generation_ms, quality["score"])
        # Track report usage
        _increment_usage(user_id, "report", report_id, org_id, conn)
        conn.close()

        # Trigger intelligence refresh from report insights (background)
        if request.player_id:
            asyncio.create_task(_extract_report_intelligence(report_id, request.player_id, org_id))

        return ReportGenerateResponse(report_id=report_id, status="complete", title=title, generation_time_ms=generation_ms)

    except Exception as e:
        conn.execute("UPDATE reports SET status='failed', error_message=? WHERE id = ?", (str(e), report_id))
        conn.commit()
        conn.close()
        logger.error("Report generation failed: %s — %s", report_id, str(e))
        raise HTTPException(status_code=500, detail=f"Report generation failed: {str(e)}")


# ============================================================
# REPORT CRUD
# ============================================================

@app.get("/reports/{report_id}", response_model=ReportResponse)
async def get_report(report_id: str, token_data: dict = Depends(verify_token)):
    org_id = token_data["org_id"]
    conn = get_db()
    row = conn.execute("SELECT * FROM reports WHERE id = ? AND org_id = ?", (report_id, org_id)).fetchone()
    conn.close()
    if not row:
        raise HTTPException(status_code=404, detail="Report not found")
    return ReportResponse(**dict(row))


@app.get("/reports/{report_id}/status", response_model=ReportStatusResponse)
async def get_report_status(report_id: str, token_data: dict = Depends(verify_token)):
    org_id = token_data["org_id"]
    conn = get_db()
    row = conn.execute(
        "SELECT id, status, error_message, generation_time_ms FROM reports WHERE id = ? AND org_id = ?",
        (report_id, org_id),
    ).fetchone()
    conn.close()
    if not row:
        raise HTTPException(status_code=404, detail="Report not found")
    return ReportStatusResponse(
        report_id=row["id"], status=row["status"],
        error_message=row["error_message"], generation_time_ms=row["generation_time_ms"],
    )


@app.put("/reports/{report_id}")
async def regenerate_report(report_id: str, token_data: dict = Depends(verify_token)):
    """Regenerate an existing report with same parameters."""
    org_id = token_data["org_id"]
    conn = get_db()
    row = conn.execute("SELECT * FROM reports WHERE id = ? AND org_id = ?", (report_id, org_id)).fetchone()
    if not row:
        conn.close()
        raise HTTPException(status_code=404, detail="Report not found")

    report = dict(row)
    conn.execute("DELETE FROM reports WHERE id = ?", (report_id,))
    conn.commit()
    conn.close()

    regen_req = ReportGenerateRequest(
        player_id=report.get("player_id"),
        team_name=report.get("team_name"),
        report_type=report["report_type"],
        template_id=report.get("template_id"),
    )
    return await generate_report(regen_req, token_data)


@app.get("/reports", response_model=List[ReportResponse])
async def list_reports(
    player_id: Optional[str] = None,
    report_type: Optional[str] = None,
    report_status: Optional[str] = None,
    skip: int = Query(default=0, ge=0),
    limit: int = Query(default=50, ge=1, le=200),
    token_data: dict = Depends(verify_token),
):
    org_id = token_data["org_id"]
    conn = get_db()

    query = "SELECT * FROM reports WHERE org_id = ?"
    params: list = [org_id]

    if player_id:
        query += " AND player_id = ?"
        params.append(player_id)
    if report_type:
        query += " AND report_type = ?"
        params.append(report_type)
    if report_status:
        query += " AND status = ?"
        params.append(report_status)

    query += " ORDER BY created_at DESC LIMIT ? OFFSET ?"
    params.extend([limit, skip])

    rows = conn.execute(query, params).fetchall()
    conn.close()
    return [ReportResponse(**dict(r)) for r in rows]


@app.delete("/reports/{report_id}")
async def delete_report(report_id: str, token_data: dict = Depends(verify_token)):
    org_id = token_data["org_id"]
    conn = get_db()
    result = conn.execute("DELETE FROM reports WHERE id = ? AND org_id = ?", (report_id, org_id))
    conn.commit()
    conn.close()
    if result.rowcount == 0:
        raise HTTPException(status_code=404, detail="Report not found")
    return {"detail": "Report deleted", "report_id": report_id}


# ============================================================
# PLAYER INTELLIGENCE ENDPOINTS
# ============================================================


@app.get("/players/{player_id}/intelligence")
async def get_player_intelligence(player_id: str, token_data: dict = Depends(verify_token)):
    """Get the latest intelligence snapshot for a player. Auto-generates if none exists."""
    org_id = token_data["org_id"]
    conn = get_db()

    # Verify player belongs to org
    player = conn.execute("SELECT id FROM players WHERE id = ? AND org_id = ?", (player_id, org_id)).fetchone()
    if not player:
        conn.close()
        raise HTTPException(status_code=404, detail="Player not found")

    # Get latest intelligence
    row = conn.execute(
        "SELECT * FROM player_intelligence WHERE player_id = ? AND org_id = ? ORDER BY version DESC LIMIT 1",
        (player_id, org_id)
    ).fetchone()
    conn.close()

    if row:
        d = dict(row)
        # Parse JSON fields
        for field in ("strengths", "development_areas", "comparable_players", "tags"):
            val = d.get(field)
            if isinstance(val, str):
                try:
                    d[field] = json.loads(val)
                except Exception:
                    d[field] = []
            elif val is None:
                d[field] = []
        if isinstance(d.get("stat_signature"), str):
            try:
                d["stat_signature"] = json.loads(d["stat_signature"])
            except Exception:
                d["stat_signature"] = {}
        if isinstance(d.get("data_sources_used"), str):
            try:
                d["data_sources_used"] = json.loads(d["data_sources_used"])
            except Exception:
                d["data_sources_used"] = {}
        return d

    # No intelligence exists — check if player has any data to generate from
    conn2 = get_db()
    has_stats = conn2.execute("SELECT COUNT(*) FROM player_stats WHERE player_id = ?", (player_id,)).fetchone()[0] > 0
    has_goalie_stats = conn2.execute("SELECT COUNT(*) FROM goalie_stats WHERE player_id = ?", (player_id,)).fetchone()[0] > 0
    has_notes = conn2.execute("SELECT COUNT(*) FROM scout_notes WHERE player_id = ?", (player_id,)).fetchone()[0] > 0
    conn2.close()

    if has_stats or has_goalie_stats or has_notes:
        # Auto-generate intelligence
        result = await _generate_player_intelligence(player_id, org_id, trigger="auto")
        if result:
            return result

    # No data at all — return empty intelligence
    return {
        "id": None,
        "player_id": player_id,
        "archetype": None,
        "archetype_confidence": None,
        "overall_grade": None,
        "offensive_grade": None,
        "defensive_grade": None,
        "skating_grade": None,
        "hockey_iq_grade": None,
        "compete_grade": None,
        "summary": None,
        "strengths": [],
        "development_areas": [],
        "comparable_players": [],
        "stat_signature": None,
        "tags": [],
        "projection": None,
        "trigger": None,
        "version": 0,
        "created_at": None,
    }


@app.get("/players/{player_id}/intelligence/history")
async def get_intelligence_history(player_id: str, token_data: dict = Depends(verify_token)):
    """Get all historical intelligence snapshots for a player (shows evolution)."""
    org_id = token_data["org_id"]
    conn = get_db()

    # Verify player
    player = conn.execute("SELECT id FROM players WHERE id = ? AND org_id = ?", (player_id, org_id)).fetchone()
    if not player:
        conn.close()
        raise HTTPException(status_code=404, detail="Player not found")

    rows = conn.execute(
        "SELECT * FROM player_intelligence WHERE player_id = ? AND org_id = ? ORDER BY version DESC",
        (player_id, org_id)
    ).fetchall()
    conn.close()

    results = []
    for row in rows:
        d = dict(row)
        for field in ("strengths", "development_areas", "comparable_players", "tags"):
            val = d.get(field)
            if isinstance(val, str):
                try:
                    d[field] = json.loads(val)
                except Exception:
                    d[field] = []
            elif val is None:
                d[field] = []
        if isinstance(d.get("stat_signature"), str):
            try:
                d["stat_signature"] = json.loads(d["stat_signature"])
            except Exception:
                d["stat_signature"] = {}
        if isinstance(d.get("data_sources_used"), str):
            try:
                d["data_sources_used"] = json.loads(d["data_sources_used"])
            except Exception:
                d["data_sources_used"] = {}
        results.append(d)

    return results


@app.post("/players/{player_id}/intelligence")
async def refresh_player_intelligence(player_id: str, token_data: dict = Depends(verify_token)):
    """Manually trigger a full intelligence refresh for a player."""
    org_id = token_data["org_id"]
    conn = get_db()

    # Verify player
    player = conn.execute("SELECT id FROM players WHERE id = ? AND org_id = ?", (player_id, org_id)).fetchone()
    if not player:
        conn.close()
        raise HTTPException(status_code=404, detail="Player not found")
    conn.close()

    result = await _generate_player_intelligence(player_id, org_id, trigger="manual")
    if result:
        return result
    raise HTTPException(status_code=500, detail="Intelligence generation failed — check server logs")


# ============================================================
# SCOUT NOTES ENDPOINTS
# ============================================================

NOTE_TAGS = ["skating", "shooting", "compete", "hockey_iq", "puck_skills", "positioning",
             "physicality", "speed", "vision", "leadership", "coachability", "work_ethic"]

_NOTE_V2_FIELDS = [
    "game_date", "opponent", "competition_level", "venue", "overall_grade", "grade_scale",
    "skating_rating", "puck_skills_rating", "hockey_iq_rating", "compete_rating", "defense_rating",
    "strengths_notes", "improvements_notes", "development_notes", "one_line_summary",
    "prospect_status", "visibility", "note_mode",
]

def _build_note_response(d: dict) -> NoteResponse:
    """Build a NoteResponse from a joined DB row dict."""
    tags = d.get("tags", "[]")
    if isinstance(tags, str):
        try:
            tags = json.loads(tags)
        except Exception:
            tags = []
    scout_name = f"{d.get('scout_first', '') or ''} {d.get('scout_last', '') or ''}".strip()
    player_name = f"{d.get('player_first', '') or ''} {d.get('player_last', '') or ''}".strip()
    return NoteResponse(
        id=d["id"], org_id=d["org_id"], player_id=d["player_id"], scout_id=d["scout_id"],
        scout_name=scout_name or None, note_text=d.get("note_text") or "", note_type=d.get("note_type") or "general",
        tags=tags, is_private=bool(d.get("is_private")),
        created_at=d["created_at"], updated_at=d["updated_at"],
        game_date=d.get("game_date"), opponent=d.get("opponent"),
        competition_level=d.get("competition_level"), venue=d.get("venue"),
        overall_grade=d.get("overall_grade"), grade_scale=d.get("grade_scale") or "1-5",
        skating_rating=d.get("skating_rating"), puck_skills_rating=d.get("puck_skills_rating"),
        hockey_iq_rating=d.get("hockey_iq_rating"), compete_rating=d.get("compete_rating"),
        defense_rating=d.get("defense_rating"), strengths_notes=d.get("strengths_notes"),
        improvements_notes=d.get("improvements_notes"), development_notes=d.get("development_notes"),
        one_line_summary=d.get("one_line_summary"), prospect_status=d.get("prospect_status"),
        visibility=d.get("visibility") or ("PRIVATE" if d.get("is_private") else "ORG_SHARED"),
        note_mode=d.get("note_mode") or "QUICK",
        player_name=player_name or None,
        player_team=d.get("current_team"),
        player_position=d.get("position"),
        author_name=scout_name or None,
    )

_NOTE_JOIN_SQL = """
    SELECT n.*, u.first_name as scout_first, u.last_name as scout_last,
           p.first_name as player_first, p.last_name as player_last,
           p.current_team, p.position
    FROM scout_notes n
    LEFT JOIN users u ON n.scout_id = u.id
    LEFT JOIN players p ON n.player_id = p.id
"""

_NOTE_VISIBILITY_WHERE = "AND (n.is_private = 0 OR (n.visibility IS NOT NULL AND n.visibility != 'PRIVATE') OR n.scout_id = ?)"


@app.get("/notes/tags")
async def get_note_tags():
    """Return available note tag options."""
    return NOTE_TAGS


@app.get("/notes/meta")
async def get_note_meta():
    """Return enums for scout note forms."""
    return {
        "tags": NOTE_TAGS,
        "competition_levels": COMPETITION_LEVELS,
        "prospect_statuses": PROSPECT_STATUSES,
    }


def _insert_scout_note(conn, note_id: str, org_id: str, player_id: str, user_id: str, note: NoteCreate, now: str):
    """Insert a scout note with all v1+v2 fields."""
    conn.execute("""
        INSERT INTO scout_notes (
            id, org_id, player_id, scout_id, note_text, note_type, tags, is_private,
            game_date, opponent, competition_level, venue, overall_grade, grade_scale,
            skating_rating, puck_skills_rating, hockey_iq_rating, compete_rating, defense_rating,
            strengths_notes, improvements_notes, development_notes, one_line_summary,
            prospect_status, visibility, note_mode, created_at, updated_at
        ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
    """, (
        note_id, org_id, player_id, user_id,
        note.note_text, note.note_type, json.dumps(note.tags or []),
        1 if note.is_private or note.visibility == "PRIVATE" else 0,
        note.game_date, note.opponent, note.competition_level, note.venue,
        note.overall_grade, note.grade_scale,
        note.skating_rating, note.puck_skills_rating, note.hockey_iq_rating,
        note.compete_rating, note.defense_rating,
        note.strengths_notes, note.improvements_notes, note.development_notes,
        note.one_line_summary, note.prospect_status, note.visibility, note.note_mode,
        now, now,
    ))
    conn.commit()


@app.post("/players/{player_id}/notes", response_model=NoteResponse, status_code=201)
async def create_note(player_id: str, note: NoteCreate, token_data: dict = Depends(verify_token)):
    org_id = token_data["org_id"]
    user_id = token_data["user_id"]
    conn = get_db()

    player = conn.execute("SELECT id FROM players WHERE id = ? AND org_id = ?", (player_id, org_id)).fetchone()
    if not player:
        conn.close()
        raise HTTPException(status_code=404, detail="Player not found")

    note_id = gen_id()
    now = now_iso()
    _insert_scout_note(conn, note_id, org_id, player_id, user_id, note, now)

    row = conn.execute(_NOTE_JOIN_SQL + " WHERE n.id = ?", (note_id,)).fetchone()
    conn.close()
    logger.info("Note created for player %s by %s", player_id, user_id)
    asyncio.create_task(_generate_player_intelligence(player_id, org_id, trigger="note"))
    return _build_note_response(dict(row))


@app.post("/scout-notes", response_model=NoteResponse, status_code=201)
async def create_scout_note_standalone(note: NoteCreate, token_data: dict = Depends(verify_token)):
    """Standalone create — player_id comes from request body."""
    org_id = token_data["org_id"]
    user_id = token_data["user_id"]
    if not note.player_id:
        raise HTTPException(status_code=422, detail="player_id is required")
    conn = get_db()

    player = conn.execute("SELECT id FROM players WHERE id = ? AND org_id = ?", (note.player_id, org_id)).fetchone()
    if not player:
        conn.close()
        raise HTTPException(status_code=404, detail="Player not found")

    note_id = gen_id()
    now = now_iso()
    _insert_scout_note(conn, note_id, org_id, note.player_id, user_id, note, now)

    row = conn.execute(_NOTE_JOIN_SQL + " WHERE n.id = ?", (note_id,)).fetchone()
    conn.close()
    logger.info("Scout note created for player %s by %s", note.player_id, user_id)
    asyncio.create_task(_generate_player_intelligence(note.player_id, org_id, trigger="note"))
    return _build_note_response(dict(row))


@app.get("/scout-notes", response_model=List[NoteResponse])
async def list_scout_notes(
    player_id: Optional[str] = None,
    prospect_status: Optional[str] = None,
    competition_level: Optional[str] = None,
    game_date_from: Optional[str] = None,
    game_date_to: Optional[str] = None,
    visibility: Optional[str] = None,
    limit: int = 20,
    offset: int = 0,
    token_data: dict = Depends(verify_token),
):
    """Paginated list of scout notes with filters."""
    org_id = token_data["org_id"]
    user_id = token_data["user_id"]
    conn = get_db()

    query = _NOTE_JOIN_SQL + " WHERE n.org_id = ? " + _NOTE_VISIBILITY_WHERE
    params: list = [org_id, user_id]

    if player_id:
        query += " AND n.player_id = ?"
        params.append(player_id)
    if prospect_status:
        query += " AND n.prospect_status = ?"
        params.append(prospect_status)
    if competition_level:
        query += " AND n.competition_level = ?"
        params.append(competition_level)
    if game_date_from:
        query += " AND n.game_date >= ?"
        params.append(game_date_from)
    if game_date_to:
        query += " AND n.game_date <= ?"
        params.append(game_date_to)
    if visibility == "PRIVATE":
        query += " AND n.scout_id = ? AND (n.visibility = 'PRIVATE' OR n.is_private = 1)"
        params.append(user_id)

    query += " ORDER BY COALESCE(n.game_date, n.created_at) DESC LIMIT ? OFFSET ?"
    limit = min(limit, 100)
    params.extend([limit, offset])

    rows = conn.execute(query, params).fetchall()
    conn.close()
    return [_build_note_response(dict(r)) for r in rows]


@app.get("/scout-notes/{note_id}", response_model=NoteResponse)
async def get_scout_note(note_id: str, token_data: dict = Depends(verify_token)):
    """Get a single scout note by ID."""
    org_id = token_data["org_id"]
    user_id = token_data["user_id"]
    conn = get_db()

    row = conn.execute(_NOTE_JOIN_SQL + " WHERE n.id = ? AND n.org_id = ?", (note_id, org_id)).fetchone()
    if not row:
        conn.close()
        raise HTTPException(status_code=404, detail="Note not found")

    d = dict(row)
    is_private = d.get("visibility") == "PRIVATE" or d.get("is_private")
    if is_private and d["scout_id"] != user_id:
        conn.close()
        raise HTTPException(status_code=404, detail="Note not found")

    conn.close()
    return _build_note_response(d)


@app.get("/players/{player_id}/notes", response_model=List[NoteResponse])
async def get_player_notes(
    player_id: str,
    note_type: Optional[str] = None,
    tag: Optional[str] = None,
    token_data: dict = Depends(verify_token),
):
    org_id = token_data["org_id"]
    user_id = token_data["user_id"]
    conn = get_db()

    query = _NOTE_JOIN_SQL + " WHERE n.player_id = ? AND n.org_id = ? " + _NOTE_VISIBILITY_WHERE
    params: list = [player_id, org_id, user_id]

    if note_type:
        query += " AND n.note_type = ?"
        params.append(note_type)

    query += " ORDER BY COALESCE(n.game_date, n.created_at) DESC"

    rows = conn.execute(query, params).fetchall()
    conn.close()

    results = [_build_note_response(dict(r)) for r in rows]

    if tag:
        results = [n for n in results if tag in n.tags]

    return results


@app.put("/notes/{note_id}", response_model=NoteResponse)
async def update_note(note_id: str, note: NoteUpdate, token_data: dict = Depends(verify_token)):
    org_id = token_data["org_id"]
    user_id = token_data["user_id"]
    conn = get_db()

    row = conn.execute("SELECT * FROM scout_notes WHERE id = ? AND org_id = ?", (note_id, org_id)).fetchone()
    if not row:
        conn.close()
        raise HTTPException(status_code=404, detail="Note not found")

    if row["scout_id"] != user_id:
        conn.close()
        raise HTTPException(status_code=403, detail="Can only edit your own notes")

    updates = {}
    if note.note_text is not None:
        updates["note_text"] = note.note_text
    if note.note_type is not None:
        updates["note_type"] = note.note_type
    if note.tags is not None:
        updates["tags"] = json.dumps(note.tags)
    if note.is_private is not None:
        updates["is_private"] = 1 if note.is_private else 0
    # v2 fields
    for field in _NOTE_V2_FIELDS:
        val = getattr(note, field, None)
        if val is not None:
            updates[field] = val
    # Sync visibility ↔ is_private
    if note.visibility is not None:
        updates["is_private"] = 1 if note.visibility == "PRIVATE" else 0

    if updates:
        updates["updated_at"] = now_iso()
        set_clause = ", ".join(f"{k} = ?" for k in updates)
        vals = list(updates.values()) + [note_id]
        conn.execute(f"UPDATE scout_notes SET {set_clause} WHERE id = ?", vals)
        conn.commit()

    updated_row = conn.execute(_NOTE_JOIN_SQL + " WHERE n.id = ?", (note_id,)).fetchone()
    conn.close()
    return _build_note_response(dict(updated_row))


@app.delete("/notes/{note_id}")
async def delete_note(note_id: str, token_data: dict = Depends(verify_token)):
    org_id = token_data["org_id"]
    user_id = token_data["user_id"]
    conn = get_db()

    row = conn.execute("SELECT scout_id FROM scout_notes WHERE id = ? AND org_id = ?", (note_id, org_id)).fetchone()
    if not row:
        conn.close()
        raise HTTPException(status_code=404, detail="Note not found")
    if row["scout_id"] != user_id:
        conn.close()
        raise HTTPException(status_code=403, detail="Can only delete your own notes")

    conn.execute("DELETE FROM scout_notes WHERE id = ?", (note_id,))
    conn.commit()
    conn.close()
    return {"detail": "Note deleted"}


# ============================================================
# HOCKEY OPERATING SYSTEM ENDPOINTS
# ============================================================


class TeamSystemCreate(BaseModel):
    team_name: str
    season: str = ""
    forecheck: str = ""
    dz_structure: str = ""
    oz_setup: str = ""
    pp_formation: str = ""
    pk_formation: str = ""
    neutral_zone: str = ""
    breakout: str = ""
    identity_tags: list = []
    pace: str = ""
    physicality: str = ""
    offensive_style: str = ""
    notes: str = ""


class TeamSystemResponse(BaseModel):
    id: str
    org_id: str
    team_id: Optional[str] = None
    team_name: str
    season: str = ""
    forecheck: str = ""
    dz_structure: str = ""
    oz_setup: str = ""
    pp_formation: str = ""
    pk_formation: str = ""
    neutral_zone: str = ""
    breakout: str = ""
    identity_tags: list = []
    notes: str = ""
    created_at: str = ""
    updated_at: str = ""


@app.get("/hockey-os/systems-library")
async def list_systems_library(system_type: Optional[str] = None):
    """Get all available tactical systems from the library (no auth needed — reference data)."""
    conn = get_db()
    if system_type:
        rows = conn.execute("SELECT * FROM systems_library WHERE system_type = ? ORDER BY code", (system_type,)).fetchall()
    else:
        rows = conn.execute("SELECT * FROM systems_library ORDER BY system_type, code").fetchall()
    conn.close()
    return [dict(r) for r in rows]


@app.get("/hockey-os/glossary")
async def list_glossary(category: Optional[str] = None):
    """Get hockey terminology glossary (no auth needed — reference data)."""
    conn = get_db()
    if category:
        rows = conn.execute("SELECT * FROM hockey_terms WHERE category = ? ORDER BY term", (category,)).fetchall()
    else:
        rows = conn.execute("SELECT * FROM hockey_terms ORDER BY category, term").fetchall()
    conn.close()
    return [dict(r) for r in rows]


@app.get("/hockey-os/team-systems")
async def list_team_systems(token_data: dict = Depends(verify_token)):
    """Get all team system profiles for the org."""
    org_id = token_data["org_id"]
    conn = get_db()
    rows = conn.execute("SELECT * FROM team_systems WHERE org_id = ? ORDER BY team_name", (org_id,)).fetchall()
    conn.close()
    results = []
    for r in rows:
        d = dict(r)
        if d.get("identity_tags"):
            try:
                d["identity_tags"] = json.loads(d["identity_tags"])
            except (json.JSONDecodeError, TypeError):
                d["identity_tags"] = []
        else:
            d["identity_tags"] = []
        results.append(d)
    return results


@app.get("/hockey-os/team-systems/{system_id}")
async def get_team_system(system_id: str, token_data: dict = Depends(verify_token)):
    """Get a single team system profile."""
    org_id = token_data["org_id"]
    conn = get_db()
    row = conn.execute("SELECT * FROM team_systems WHERE id = ? AND org_id = ?", (system_id, org_id)).fetchone()
    conn.close()
    if not row:
        raise HTTPException(status_code=404, detail="Team system not found")
    result = dict(row)
    # Parse JSON fields
    if result.get("identity_tags"):
        try:
            result["identity_tags"] = json.loads(result["identity_tags"])
        except (json.JSONDecodeError, TypeError):
            result["identity_tags"] = []
    return result


@app.post("/hockey-os/team-systems", status_code=201)
async def create_team_system(body: TeamSystemCreate, token_data: dict = Depends(verify_token)):
    """Create a team system profile."""
    org_id = token_data["org_id"]
    system_id = gen_id()
    now = now_iso()
    conn = get_db()
    conn.execute("""
        INSERT INTO team_systems (id, org_id, team_name, season, forecheck, dz_structure, oz_setup,
                                   pp_formation, pk_formation, neutral_zone, breakout, identity_tags,
                                   pace, physicality, offensive_style, notes, created_at, updated_at)
        VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
    """, (system_id, org_id, body.team_name, body.season, body.forecheck, body.dz_structure, body.oz_setup,
          body.pp_formation, body.pk_formation, body.neutral_zone, body.breakout,
          json.dumps(body.identity_tags), body.pace, body.physicality, body.offensive_style,
          body.notes, now, now))
    conn.commit()
    row = conn.execute("SELECT * FROM team_systems WHERE id = ?", (system_id,)).fetchone()
    conn.close()
    logger.info("Team system created: %s (org %s)", body.team_name, org_id)
    return dict(row)


@app.put("/hockey-os/team-systems/{system_id}")
async def update_team_system(system_id: str, body: TeamSystemCreate, token_data: dict = Depends(verify_token)):
    """Update a team system profile."""
    org_id = token_data["org_id"]
    conn = get_db()
    existing = conn.execute("SELECT id FROM team_systems WHERE id = ? AND org_id = ?", (system_id, org_id)).fetchone()
    if not existing:
        conn.close()
        raise HTTPException(status_code=404, detail="Team system not found")
    conn.execute("""
        UPDATE team_systems SET team_name=?, season=?, forecheck=?, dz_structure=?, oz_setup=?,
               pp_formation=?, pk_formation=?, neutral_zone=?, breakout=?, identity_tags=?,
               pace=?, physicality=?, offensive_style=?, notes=?, updated_at=?
        WHERE id = ? AND org_id = ?
    """, (body.team_name, body.season, body.forecheck, body.dz_structure, body.oz_setup,
          body.pp_formation, body.pk_formation, body.neutral_zone, body.breakout,
          json.dumps(body.identity_tags), body.pace, body.physicality, body.offensive_style,
          body.notes, now_iso(), system_id, org_id))
    conn.commit()
    row = conn.execute("SELECT * FROM team_systems WHERE id = ?", (system_id,)).fetchone()
    conn.close()
    return dict(row)


@app.delete("/hockey-os/team-systems/{system_id}")
async def delete_team_system(system_id: str, token_data: dict = Depends(verify_token)):
    """Delete a team system profile."""
    org_id = token_data["org_id"]
    conn = get_db()
    conn.execute("DELETE FROM team_systems WHERE id = ? AND org_id = ?", (system_id, org_id))
    conn.commit()
    conn.close()
    return {"detail": "Team system deleted"}


# ============================================================
# BATCH PLAYER IMPORT
# ============================================================

def _fuzzy_name_match(name1: str, name2: str) -> float:
    """Simple fuzzy match — normalized Levenshtein-like comparison."""
    a = name1.lower().strip()
    b = name2.lower().strip()
    if a == b:
        return 1.0
    # Check if one is substring of the other
    if a in b or b in a:
        return 0.9
    # Character-level similarity
    longer = max(len(a), len(b))
    if longer == 0:
        return 1.0
    # Count matching characters
    matches = sum(1 for ca, cb in zip(a, b) if ca == cb)
    return matches / longer


@app.post("/import/preview")
async def preview_import(
    file: UploadFile = File(...),
    team_override: Optional[str] = None,
    league_override: Optional[str] = None,
    season_override: Optional[str] = None,
    token_data: dict = Depends(verify_token),
):
    """Upload a CSV or Excel file, parse it, detect duplicates, return preview for admin review.
    Optional team_override/league_override/season_override to auto-inject values for all rows.
    """
    org_id = token_data["org_id"]
    user_id = token_data["user_id"]

    # ── Tier permission + limit checks ──
    perm_conn = get_db()
    try:
        tier_config = _check_tier_permission(user_id, "can_upload_files", perm_conn)
        _check_tier_limit(user_id, "uploads", perm_conn)
    finally:
        perm_conn.close()

    fname = (file.filename or "").lower()

    if not any(fname.endswith(ext) for ext in (".csv", ".xlsx", ".xls", ".xlsm")):
        raise HTTPException(status_code=400, detail="File must be .csv, .xlsx, or .xls")

    content = await file.read()

    # ── File size check ──
    max_bytes = tier_config.get("max_file_size_mb", 5) * 1024 * 1024
    if len(content) > max_bytes:
        raise HTTPException(status_code=413, detail={
            "error": "file_too_large",
            "max_mb": tier_config.get("max_file_size_mb", 5),
            "file_mb": round(len(content) / (1024 * 1024), 2),
            "upgrade_url": "/pricing",
        })

    # Track the upload usage
    track_conn = get_db()
    try:
        _increment_tracking(user_id, "uploads", track_conn)
    finally:
        track_conn.close()

    parsed_rows = _parse_file_to_rows(content, fname)

    rows_data = []
    parse_errors = []

    # Log the headers we detected for debugging
    if parsed_rows:
        detected_headers = list(parsed_rows[0].keys())
        logger.info("Import headers detected: %s", detected_headers)

        # ── Analytics format detection — give helpful guidance ─────
        if _detect_instat_game_log(detected_headers):
            raise HTTPException(
                status_code=400,
                detail="This looks like a per-game stats log for one player. "
                       "To import game stats, go to the player's profile page and use the "
                       "stats upload button on the Stats tab, or use the Import Stats page. "
                       "This Import Players page is for adding new players to the database."
            )
        if _detect_instat_team_stats(detected_headers) and not any(
            h in detected_headers for h in ["first_name", "first", "firstname", "last_name", "last", "lastname"]
        ):
            # It's a team stats file — we can try to parse it, but warn if no name split will work
            logger.info("InStat team stats file detected in import — will try player name column")

    for i, row in enumerate(parsed_rows):
        first = row.get("first_name") or row.get("first") or row.get("firstname") or ""
        last = row.get("last_name") or row.get("last") or row.get("lastname") or ""

        # If no separate first/last, try full name columns and split
        if not first or not last:
            full_name = (
                row.get("name") or row.get("player") or row.get("player_name")
                or row.get("playername") or row.get("full_name") or row.get("fullname") or ""
            )
            if full_name:
                parts = full_name.strip().split(None, 1)  # Split on first space
                if len(parts) >= 2:
                    first = first or parts[0]
                    last = last or parts[1]
                elif len(parts) == 1:
                    # Could be just a last name
                    last = last or parts[0]

        if not first or not last:
            # Log the actual row data so we can diagnose
            row_preview = {k: v for k, v in list(row.items())[:6] if v}
            parse_errors.append(f"Row {i+1}: missing first_name or last_name (got: {row_preview})")
            continue

        pos = (row.get("position") or row.get("pos") or "F").upper()
        # Handle common position formats: "LW", "C", "D", "G", "RW", "LD", "RD", "F", "W"
        pos_map = {"LD": "D", "RD": "D", "LF": "LW", "RF": "RW", "GK": "G", "FWD": "F", "DEF": "D", "CENTER": "C", "CENTRE": "C", "LEFT_WING": "LW", "RIGHT_WING": "RW", "DEFENSE": "D", "DEFENCE": "D", "GOALIE": "G", "GOALTENDER": "G", "FORWARD": "F", "WING": "W"}
        pos = pos_map.get(pos, pos)
        if pos not in ("C", "LW", "RW", "D", "G", "F", "W", "LD", "RD"):
            pos = "F"

        rows_data.append({
            "row_index": i,
            "first_name": first,
            "last_name": last,
            "position": pos,
            "dob": row.get("dob") or row.get("date_of_birth") or row.get("birthdate") or row.get("birth_date") or row.get("birthday") or row.get("born") or None,
            "shoots": row.get("shoots") or row.get("shot") or row.get("sh") or row.get("hand") or row.get("handedness") or None,
            "current_team": row.get("current_team") or row.get("team") or row.get("club") or row.get("team_name") or None,
            "current_league": row.get("current_league") or row.get("league") or row.get("lg") or None,
            "height_cm": row.get("height_cm") or row.get("height") or row.get("ht") or None,
            "weight_kg": row.get("weight_kg") or row.get("weight") or row.get("wt") or None,
            # Stats columns (optional — will be imported as player_stats)
            "gp": row.get("gp") or row.get("games_played") or row.get("games") or row.get("g.p.") or None,
            "g": row.get("g") or row.get("goals") or None,
            "a": row.get("a") or row.get("assists") or row.get("ast") or None,
            "p": row.get("p") or row.get("pts") or row.get("points") or row.get("tp") or None,
            "plus_minus": row.get("plus_minus") or row.get("+/_") or row.get("plusminus") or row.get("+/−") or row.get("+/") or None,
            "pim": row.get("pim") or row.get("penalty_minutes") or row.get("pen") or None,
            "season": row.get("season") or row.get("year") or None,
        })

    # Apply team/league override for team roster imports
    if team_override:
        for rd in rows_data:
            if not rd.get("current_team"):
                rd["current_team"] = team_override
    if league_override:
        for rd in rows_data:
            if not rd.get("current_league"):
                rd["current_league"] = league_override
    if season_override:
        for rd in rows_data:
            if not rd.get("season"):
                rd["season"] = season_override

    # Fetch existing players for duplicate detection
    conn = get_db()
    existing = conn.execute("SELECT id, first_name, last_name, dob, position, current_team FROM players WHERE org_id = ?", (org_id,)).fetchall()
    existing_players = [dict(r) for r in existing]

    # Detect duplicates
    duplicates = []
    new_players = []

    for rd in rows_data:
        csv_name = f"{rd['first_name']} {rd['last_name']}"
        best_match = None
        best_score = 0

        for ep in existing_players:
            existing_name = f"{ep['first_name']} {ep['last_name']}"
            name_score = _fuzzy_name_match(csv_name, existing_name)
            reasons = []
            score = name_score * 0.6  # 60% weight on name

            if name_score > 0.85:
                reasons.append(f"Name: {int(name_score * 100)}% match")

            if rd.get("dob") and ep.get("dob") and rd["dob"] == ep["dob"]:
                score += 0.25
                reasons.append("DOB: exact match")

            if rd.get("position") and ep.get("position") and rd["position"] == ep["position"]:
                score += 0.08
                reasons.append("Position: same")

            if rd.get("current_team") and ep.get("current_team") and rd["current_team"].lower() == ep["current_team"].lower():
                score += 0.07
                reasons.append("Team: same")

            if score > best_score and score >= 0.55:
                best_score = score
                best_match = {"existing": ep, "score": score, "reasons": reasons}

        if best_match:
            duplicates.append({
                "row_index": rd["row_index"],
                "csv_name": csv_name,
                "csv_data": rd,
                "existing_id": best_match["existing"]["id"],
                "existing_name": f"{best_match['existing']['first_name']} {best_match['existing']['last_name']}",
                "match_score": round(best_match["score"], 2),
                "match_reasons": best_match["reasons"],
            })
        else:
            new_players.append(rd)

    # Create import job
    job_id = gen_id()
    conn.execute("""
        INSERT INTO import_jobs (id, org_id, uploaded_by, filename, total_rows, new_players, duplicates_found, errors_found, status, preview_data, duplicate_data, created_at)
        VALUES (?, ?, ?, ?, ?, ?, ?, ?, 'preview', ?, ?, ?)
    """, (job_id, org_id, user_id, file.filename, len(rows_data), len(new_players),
          len(duplicates), len(parse_errors), json.dumps(rows_data), json.dumps(duplicates), now_iso()))
    conn.commit()
    conn.close()

    logger.info("Import preview: %s — %d rows, %d new, %d duplicates, %d errors",
                file.filename, len(rows_data), len(new_players), len(duplicates), len(parse_errors))

    return {
        "job_id": job_id,
        "filename": file.filename,
        "total_rows": len(rows_data),
        "new_players": len(new_players),
        "duplicates": duplicates,
        "errors": parse_errors[:20],
        "preview": rows_data[:10],
    }


@app.post("/import/{job_id}/execute")
async def execute_import(job_id: str, body: ImportExecuteRequest, token_data: dict = Depends(verify_token)):
    """Execute the import — create new players, handle duplicate resolutions."""
    org_id = token_data["org_id"]

    created = 0
    merged = 0
    skipped = 0
    errors = []

    def safe_int(val):
        try:
            return int(float(val)) if val else None
        except (ValueError, TypeError):
            return None

    with safe_db() as conn:
        job = conn.execute("SELECT * FROM import_jobs WHERE id = ? AND org_id = ?", (job_id, org_id)).fetchone()
        if not job:
            raise HTTPException(status_code=404, detail="Import job not found")

        job = dict(job)
        all_rows = json.loads(job["preview_data"])
        duplicates = json.loads(job["duplicate_data"])
        dup_indices = {d["row_index"] for d in duplicates}

        # Build resolution map from admin decisions
        resolution_map = {r.row_index: r.action for r in body.resolutions}

        for rd in all_rows:
            idx = rd["row_index"]
            now = now_iso()

            if idx in dup_indices:
                dup = next(d for d in duplicates if d["row_index"] == idx)
                action = resolution_map.get(idx, "skip")

                if action == "skip":
                    skipped += 1
                    continue
                elif action == "create_new":
                    # Create as new player even though duplicate detected
                    player_id = gen_id()
                    conn.execute("""
                        INSERT INTO players (id, org_id, first_name, last_name, dob, position, shoots,
                                            height_cm, weight_kg, current_team, current_league, created_at, updated_at)
                        VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                    """, (player_id, org_id, rd["first_name"], rd["last_name"], rd.get("dob"),
                          rd["position"], rd.get("shoots"), safe_int(rd.get("height_cm")),
                          safe_int(rd.get("weight_kg")), rd.get("current_team"),
                          rd.get("current_league"), now, now))
                    created += 1

                    # Import stats if present
                    if rd.get("gp"):
                        g = safe_int(rd.get("g")) or 0
                        a = safe_int(rd.get("a")) or 0
                        p = safe_int(rd.get("p")) or (g + a)
                        conn.execute("""
                            INSERT INTO player_stats (id, player_id, season, stat_type, gp, g, a, p, plus_minus, pim, created_at)
                            VALUES (?, ?, ?, 'season', ?, ?, ?, ?, ?, ?, ?)
                        """, (gen_id(), player_id, rd.get("season", ""), safe_int(rd.get("gp")) or 0,
                              g, a, p, safe_int(rd.get("plus_minus")) or 0,
                              safe_int(rd.get("pim")) or 0, now))

                elif action == "merge":
                    # Update existing player's stats — delete old season stats first to prevent doubling
                    existing_id = dup["existing_id"]
                    if rd.get("gp"):
                        season_val = rd.get("season", "")
                        g = safe_int(rd.get("g")) or 0
                        a = safe_int(rd.get("a")) or 0
                        p = safe_int(rd.get("p")) or (g + a)
                        # Remove existing season stats for this player/season to prevent duplicates
                        conn.execute(
                            "DELETE FROM player_stats WHERE player_id = ? AND season = ? AND stat_type = 'season' AND data_source IS NULL",
                            (existing_id, season_val)
                        )
                        conn.execute("""
                            INSERT INTO player_stats (id, player_id, season, stat_type, gp, g, a, p, plus_minus, pim, created_at)
                            VALUES (?, ?, ?, 'season', ?, ?, ?, ?, ?, ?, ?)
                        """, (gen_id(), existing_id, season_val, safe_int(rd.get("gp")) or 0,
                              g, a, p, safe_int(rd.get("plus_minus")) or 0,
                              safe_int(rd.get("pim")) or 0, now))
                    merged += 1
            else:
                # New player — create
                player_id = gen_id()
                try:
                    conn.execute("""
                        INSERT INTO players (id, org_id, first_name, last_name, dob, position, shoots,
                                            height_cm, weight_kg, current_team, current_league, created_at, updated_at)
                        VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                    """, (player_id, org_id, rd["first_name"], rd["last_name"], rd.get("dob"),
                          rd["position"], rd.get("shoots"), safe_int(rd.get("height_cm")),
                          safe_int(rd.get("weight_kg")), rd.get("current_team"),
                          rd.get("current_league"), now, now))
                    created += 1

                    # Import stats if present
                    if rd.get("gp"):
                        g = safe_int(rd.get("g")) or 0
                        a = safe_int(rd.get("a")) or 0
                        p = safe_int(rd.get("p")) or (g + a)
                        conn.execute("""
                            INSERT INTO player_stats (id, player_id, season, stat_type, gp, g, a, p, plus_minus, pim, created_at)
                            VALUES (?, ?, ?, 'season', ?, ?, ?, ?, ?, ?, ?)
                        """, (gen_id(), player_id, rd.get("season", ""), safe_int(rd.get("gp")) or 0,
                              g, a, p, safe_int(rd.get("plus_minus")) or 0,
                              safe_int(rd.get("pim")) or 0, now))
                except Exception as e:
                    errors.append(f"Row {idx+1}: {str(e)}")

        # Update job status
        conn.execute("""
            UPDATE import_jobs SET status = 'complete', new_players = ?, duplicates_found = ?
            WHERE id = ?
        """, (created, merged, job_id))
        # safe_db() auto-commits here on success, auto-rollbacks on error

    logger.info("Import executed: %s — %d created, %d merged, %d skipped, %d errors",
                job_id, created, merged, skipped, len(errors))

    return {
        "detail": f"Import complete: {created} created, {merged} merged, {skipped} skipped",
        "created": created,
        "merged": merged,
        "skipped": skipped,
        "errors": errors[:10],
    }


@app.get("/import/{job_id}")
async def get_import_job(job_id: str, token_data: dict = Depends(verify_token)):
    org_id = token_data["org_id"]
    conn = get_db()
    row = conn.execute("SELECT * FROM import_jobs WHERE id = ? AND org_id = ?", (job_id, org_id)).fetchone()
    conn.close()
    if not row:
        raise HTTPException(status_code=404, detail="Import job not found")
    return dict(row)


# ============================================================
# INSTAT ANALYTICS — IMPORT ENGINE
# ============================================================

@app.post("/instat/import")
async def instat_import(
    file: UploadFile = File(...),
    season: Optional[str] = Query(None),
    team_name: Optional[str] = Query(None),
    line_type: Optional[str] = Query(None),
    token_data: dict = Depends(verify_token),
):
    """
    Unified InStat XLSX import — auto-detects file type (league teams, league skaters,
    league goalies, team skaters, team goalies, lines) and routes to appropriate handler.
    """
    org_id = token_data["org_id"]
    user_id = token_data["user_id"]

    # ── Tier permission + limit checks ──
    perm_conn = get_db()
    try:
        tier_config = _check_tier_permission(user_id, "can_upload_files", perm_conn)
        _check_tier_limit(user_id, "uploads", perm_conn)
    finally:
        perm_conn.close()

    fname = (file.filename or "").lower()
    if not any(fname.endswith(ext) for ext in (".csv", ".xlsx", ".xls", ".xlsm")):
        raise HTTPException(status_code=400, detail="File must be .csv, .xlsx, or .xls")

    content = await file.read()

    # ── File size check ──
    max_bytes = tier_config.get("max_file_size_mb", 5) * 1024 * 1024
    if len(content) > max_bytes:
        raise HTTPException(status_code=413, detail={
            "error": "file_too_large",
            "max_mb": tier_config.get("max_file_size_mb", 5),
            "file_mb": round(len(content) / (1024 * 1024), 2),
            "upgrade_url": "/pricing",
        })

    # Track the upload usage
    track_conn = get_db()
    try:
        _increment_tracking(user_id, "uploads", track_conn)
    finally:
        track_conn.close()

    rows = _parse_file_to_rows(content, fname)
    if not rows:
        raise HTTPException(status_code=400, detail="No data rows found in file")

    # Get headers from first row keys
    headers = list(rows[0].keys())
    file_type = _detect_instat_file_type(headers)

    # Auto-detect season from current date if not provided
    if not season:
        now = datetime.now()
        year = now.year
        season = f"{year-1}-{year}" if now.month < 9 else f"{year}-{year+1}"

    logger.info("InStat import: file_type=%s, rows=%d, season=%s, team=%s", file_type, len(rows), season, team_name)

    result = {
        "file_type": file_type,
        "total_rows": len(rows),
        "players_created": 0,
        "players_updated": 0,
        "stats_imported": 0,
        "errors": [],
    }

    _instat_player_ids = []
    try:
        if file_type == "league_teams":
            r = _import_league_teams(rows, season, org_id, team_name_override=team_name)
            result.update(r)
        elif file_type == "league_skaters":
            r = _import_league_skaters(rows, season, org_id)
            _instat_player_ids = r.pop("player_ids", [])
            result.update(r)
        elif file_type == "league_goalies":
            r = _import_league_goalies(rows, season, org_id)
            _instat_player_ids = r.pop("player_ids", [])
            result.update(r)
        elif file_type == "team_skaters":
            if not team_name:
                raise HTTPException(status_code=400, detail="team_name required for team-specific imports")
            r = _import_team_skaters(rows, season, team_name, org_id)
            _instat_player_ids = r.pop("player_ids", [])
            result.update(r)
        elif file_type == "team_goalies":
            if not team_name:
                raise HTTPException(status_code=400, detail="team_name required for team-specific imports")
            r = _import_team_goalies(rows, season, team_name, org_id)
            _instat_player_ids = r.pop("player_ids", [])
            result.update(r)
        elif file_type == "lines":
            if not team_name:
                raise HTTPException(status_code=400, detail="team_name required for lines imports")
            lt = line_type or "full"
            r = _import_lines(rows, season, team_name, lt, org_id)
            result.update(r)
        else:
            raise HTTPException(status_code=400, detail=f"Could not detect file type. Headers: {headers[:10]}")
    except HTTPException:
        raise
    except Exception as e:
        logger.exception("InStat import error")
        result["errors"].append(str(e))

    # Trigger intelligence generation for imported players (background, non-blocking)
    # Limit to first 50 players to avoid overwhelming the API on large league imports
    for pid in _instat_player_ids[:50]:
        asyncio.create_task(_generate_player_intelligence(pid, org_id, trigger="import"))
    if _instat_player_ids:
        logger.info("Intelligence generation triggered for %d players from InStat import", min(len(_instat_player_ids), 50))

    return result


def _find_team_column(row: dict) -> str:
    """Find the team name from a row by checking multiple possible column names."""
    # Check common team column names in priority order
    team_keys = ["team", "team_name", "club", "name", "teams", "club_name", "organization"]
    for key in team_keys:
        val = row.get(key, "").strip()
        if val:
            return val

    # Fallback: try the first column value if it looks like a team name (non-numeric, > 2 chars)
    if row:
        first_key = next(iter(row))
        first_val = row.get(first_key, "").strip()
        if first_val and not first_val.replace(".", "").replace("-", "").isdigit() and len(first_val) > 2:
            return first_val

    return ""


def _import_league_teams(rows, season, org_id, team_name_override=None):
    """Import league-level team stats (26 teams x 100 columns).

    If team_name_override is provided (user selected a team on the upload form),
    all rows are assigned to that team. Otherwise, the team name is extracted
    from each row using common column name patterns.
    """
    conn = get_db()
    stats_imported = 0
    errors = []

    # Log column names for debugging
    if rows:
        logger.info("League team import columns: %s", list(rows[0].keys())[:15])
        if team_name_override:
            logger.info("Team name override from form: %s", team_name_override)

    for i, row in enumerate(rows):
        try:
            # Use override if provided; otherwise try to find team name in row
            team_name = team_name_override or _find_team_column(row)
            if not team_name:
                if i == 0:
                    # Log the actual columns on first failure for debugging
                    logger.warning("Cannot find team name in columns: %s | First row values: %s",
                                   list(row.keys())[:8], {k: v for k, v in list(row.items())[:5]})
                errors.append(f"Row {i+1}: missing team name (select a team on the upload form, or ensure file has a 'Team' column)")
                continue

            # Parse all stats into extended_stats JSON
            _, extended = _parse_instat_row(row, {}, INSTAT_TEAM_EXTENDED_MAP)

            # Check if team_stats already exists for this team/season
            existing = conn.execute(
                "SELECT id FROM team_stats WHERE org_id = ? AND LOWER(team_name) = LOWER(?) AND season = ?",
                (org_id, team_name, season)
            ).fetchone()

            if existing:
                conn.execute(
                    "UPDATE team_stats SET extended_stats = ?, data_source = 'instat' WHERE id = ?",
                    (json.dumps(extended), existing["id"])
                )
            else:
                conn.execute(
                    "INSERT INTO team_stats (id, org_id, team_name, league, season, extended_stats, data_source) VALUES (?, ?, ?, ?, ?, ?, 'instat')",
                    (gen_id(), org_id, team_name, "GOHL", season, json.dumps(extended))
                )
            stats_imported += 1
        except Exception as e:
            errors.append(f"Row {i+1} ({team_name}): {str(e)}")

    conn.commit()
    conn.close()
    return {"stats_imported": stats_imported, "errors": errors}


def _import_league_skaters(rows, season, org_id):
    """Import league-level skater stats (760 players x 138 columns)."""
    conn = get_db()
    created = 0
    updated = 0
    stats_imported = 0
    errors = []
    _affected_player_ids = set()

    # Load existing players for matching (include DOB for stronger dedup)
    existing_players = conn.execute(
        "SELECT id, first_name, last_name, current_team, position, dob FROM players WHERE org_id = ?",
        (org_id,)
    ).fetchall()
    existing_list = [dict(p) for p in existing_players]

    for i, row in enumerate(rows):
        try:
            # Parse player name
            full_name = row.get("player", "").strip()
            if not full_name:
                errors.append(f"Row {i+1}: missing player name")
                continue

            parts = full_name.split(None, 1)
            first_name = parts[0] if parts else full_name
            last_name = parts[1] if len(parts) > 1 else ""

            team = row.get("team", "").strip()
            position = row.get("position", "F").strip() or "F"
            jersey = str(row.get("shirt_number", "")).strip()

            # Bio data from passport section
            bio = {}
            for header, field in INSTAT_SKATER_BIO_MAP.items():
                val = _clean_instat_val(row.get(header))
                if val:
                    bio[field] = val

            # Match against existing players — multi-signal matching
            player_id = None
            csv_name = f"{first_name} {last_name}".lower()
            csv_last = last_name.lower().strip()
            csv_dob = bio.get("dob", "")
            best_score = 0.0
            best_match = None

            for ep in existing_list:
                existing_name = f"{ep['first_name']} {ep['last_name']}".lower()
                existing_last = ep["last_name"].lower().strip()
                existing_dob = ep.get("dob", "") or ""

                # Start with name similarity
                score = _fuzzy_name_match(csv_name, existing_name)

                # DOB match is a very strong signal
                if csv_dob and existing_dob and csv_dob not in ("-", "") and existing_dob not in ("-", ""):
                    if csv_dob == existing_dob:
                        # Same DOB: huge boost — if last names also match, it's nearly certain
                        if csv_last == existing_last:
                            score = max(score, 0.98)  # Same last name + same DOB = same person
                        else:
                            score += 0.15  # Same DOB, different last name — minor boost

                # Team match bonus
                if team and ep.get("current_team") and ep["current_team"].lower() == team.lower():
                    score += 0.1

                if score > best_score and score >= 0.85:
                    best_score = score
                    best_match = ep

            if best_match:
                player_id = best_match["id"]
                # Update existing player with any new bio data
                updates = []
                params = []
                if team and team != "":
                    updates.append("current_team = ?")
                    params.append(team)
                if bio.get("dob") and bio["dob"] not in ("-", ""):
                    updates.append("dob = ?")
                    params.append(bio["dob"])
                if bio.get("shoots"):
                    updates.append("shoots = ?")
                    params.append(bio["shoots"])
                if position and position != "F":
                    updates.append("position = ?")
                    params.append(position)
                if updates:
                    params.append(player_id)
                    conn.execute(f"UPDATE players SET {', '.join(updates)} WHERE id = ?", params)
                updated += 1
            else:
                # Create new player
                player_id = gen_id()
                conn.execute(
                    """INSERT INTO players (id, org_id, first_name, last_name, position, shoots,
                       current_team, current_league, dob)
                       VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)""",
                    (player_id, org_id, first_name, last_name, position,
                     bio.get("shoots", ""), team, "GOHL", bio.get("dob", ""))
                )
                # Add to existing list for matching remaining rows (include DOB for future matching)
                existing_list.append({"id": player_id, "first_name": first_name, "last_name": last_name,
                                     "current_team": team, "position": position, "dob": bio.get("dob", "")})
                created += 1

            # Parse stats
            core, extended = _parse_instat_row(row, INSTAT_SKATER_CORE_MAP, INSTAT_SKATER_EXTENDED_MAP)

            # Delete existing InStat season stats for this player/season (replace)
            conn.execute(
                "DELETE FROM player_stats WHERE player_id = ? AND season = ? AND data_source IN ('instat_league', 'instat_team')",
                (player_id, season)
            )

            # Insert stats
            conn.execute(
                """INSERT INTO player_stats (id, player_id, season, stat_type, gp, g, a, p,
                   plus_minus, pim, toi_seconds, shots, sog, shooting_pct,
                   extended_stats, data_source)
                   VALUES (?, ?, ?, 'season', ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, 'instat_league')""",
                (gen_id(), player_id, season,
                 core.get("gp", 0), core.get("g", 0), core.get("a", 0), core.get("p", 0),
                 core.get("plus_minus", 0), core.get("pim", 0), core.get("toi_seconds", 0),
                 core.get("shots", 0), core.get("sog", 0), core.get("shooting_pct"),
                 json.dumps(extended) if extended else None)
            )
            stats_imported += 1
            _affected_player_ids.add(player_id)

        except Exception as e:
            errors.append(f"Row {i+1} ({full_name}): {str(e)}")

    conn.commit()
    conn.close()
    return {"players_created": created, "players_updated": updated, "stats_imported": stats_imported, "errors": errors, "player_ids": list(_affected_player_ids)}


def _import_league_goalies(rows, season, org_id):
    """Import league-level goalie stats."""
    conn = get_db()
    created = 0
    updated = 0
    stats_imported = 0
    errors = []
    _affected_player_ids = set()

    existing_players = conn.execute(
        "SELECT id, first_name, last_name, current_team, position, dob FROM players WHERE org_id = ?",
        (org_id,)
    ).fetchall()
    existing_list = [dict(p) for p in existing_players]

    for i, row in enumerate(rows):
        try:
            full_name = row.get("player", "").strip()
            if not full_name:
                errors.append(f"Row {i+1}: missing player name")
                continue

            parts = full_name.split(None, 1)
            first_name = parts[0] if parts else full_name
            last_name = parts[1] if len(parts) > 1 else ""
            team = row.get("team", "").strip()

            # Bio
            bio = {}
            for header, field in INSTAT_SKATER_BIO_MAP.items():
                val = _clean_instat_val(row.get(header))
                if val:
                    bio[field] = val

            # Match player — multi-signal matching (same as skaters)
            player_id = None
            csv_name = f"{first_name} {last_name}".lower()
            csv_last = last_name.lower().strip()
            csv_dob = bio.get("dob", "")
            best_score = 0.0
            best_match = None

            for ep in existing_list:
                existing_name = f"{ep['first_name']} {ep['last_name']}".lower()
                existing_last = ep["last_name"].lower().strip()
                existing_dob = ep.get("dob", "") or ""

                score = _fuzzy_name_match(csv_name, existing_name)

                # DOB match is a very strong signal
                if csv_dob and existing_dob and csv_dob not in ("-", "") and existing_dob not in ("-", ""):
                    if csv_dob == existing_dob and csv_last == existing_last:
                        score = max(score, 0.98)

                if team and ep.get("current_team") and ep["current_team"].lower() == team.lower():
                    score += 0.1

                if score > best_score and score >= 0.85:
                    best_score = score
                    best_match = ep

            if best_match:
                player_id = best_match["id"]
                updated += 1
                # Update position to G
                conn.execute("UPDATE players SET position = 'G' WHERE id = ?", (player_id,))
            else:
                player_id = gen_id()
                conn.execute(
                    """INSERT INTO players (id, org_id, first_name, last_name, position, shoots,
                       current_team, current_league, dob)
                       VALUES (?, ?, ?, ?, 'G', ?, ?, ?, ?)""",
                    (player_id, org_id, first_name, last_name,
                     bio.get("shoots", ""), team, "GOHL", bio.get("dob", ""))
                )
                existing_list.append({"id": player_id, "first_name": first_name, "last_name": last_name,
                                     "current_team": team, "position": "G", "dob": bio.get("dob", "")})
                created += 1

            # Parse goalie stats
            core, extended = _parse_instat_row(row, INSTAT_GOALIE_CORE_MAP, INSTAT_GOALIE_EXTENDED_MAP)

            # Delete existing InStat goalie stats
            conn.execute(
                "DELETE FROM goalie_stats WHERE player_id = ? AND season = ? AND data_source IN ('instat_league', 'instat_team')",
                (player_id, season)
            )

            conn.execute(
                """INSERT INTO goalie_stats (id, player_id, org_id, season, stat_type,
                   gp, toi_seconds, ga, sa, sv, sv_pct, gaa,
                   extended_stats, data_source)
                   VALUES (?, ?, ?, ?, 'season', ?, ?, ?, ?, ?, ?, ?, ?, 'instat_league')""",
                (gen_id(), player_id, org_id, season,
                 core.get("gp", 0), core.get("toi_seconds", 0),
                 core.get("ga", 0), core.get("sa", 0), core.get("sv", 0),
                 core.get("sv_pct"), core.get("gaa"),
                 json.dumps(extended) if extended else None)
            )
            stats_imported += 1
            _affected_player_ids.add(player_id)

        except Exception as e:
            errors.append(f"Row {i+1}: {str(e)}")

    conn.commit()
    conn.close()
    return {"players_created": created, "players_updated": updated, "stats_imported": stats_imported, "errors": errors, "player_ids": list(_affected_player_ids)}


def _import_team_skaters(rows, season, team_name, org_id):
    """Import team-specific skater stats (same as league but all assigned to one team)."""
    # Team exports don't have a "team" column, so we inject it
    for row in rows:
        row["team"] = team_name
    return _import_league_skaters(rows, season, org_id)


def _import_team_goalies(rows, season, team_name, org_id):
    """Import team-specific goalie stats."""
    for row in rows:
        row["team"] = team_name
    return _import_league_goalies(rows, season, org_id)


def _import_lines(rows, season, team_name, line_type, org_id):
    """Import line combinations."""
    conn = get_db()
    stats_imported = 0
    errors = []

    # Clear existing lines for this team/season/type
    conn.execute(
        "DELETE FROM line_combinations WHERE org_id = ? AND LOWER(team_name) = LOWER(?) AND season = ? AND line_type = ?",
        (org_id, team_name, season, line_type)
    )

    for i, row in enumerate(rows):
        try:
            line_str = row.get("line", "").strip()
            if not line_str:
                errors.append(f"Row {i+1}: missing line data")
                continue

            player_refs = _parse_line_players(line_str)

            # Parse stats
            plus_minus = _clean_instat_val(row.get("plus/minus"))
            shifts = _instat_to_number(row.get("numbers_of_shifts")) or 0
            toi = _parse_mmss_to_seconds(_clean_instat_val(row.get("time_on_ice")) or "0")
            gf = _instat_to_number(row.get("goals")) or 0
            ga = _instat_to_number(row.get("opponent's_goals")) or 0

            # Extended stats (remaining columns)
            extended = {}
            for header, field in INSTAT_LINES_MAP.items():
                if field in ("plus_minus", "shifts", "toi_seconds", "goals_for", "goals_against"):
                    continue  # Already parsed as core
                val = _clean_instat_val(row.get(header))
                if val is not None:
                    if "time" in field and ":" in val:
                        extended[field] = _parse_mmss_to_seconds(val)
                    else:
                        extended[field] = _instat_to_number(val)

            conn.execute(
                """INSERT INTO line_combinations (id, org_id, team_name, season, line_type,
                   player_names, player_refs, plus_minus, shifts, toi_seconds,
                   goals_for, goals_against, extended_stats)
                   VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)""",
                (gen_id(), org_id, team_name, season, line_type,
                 line_str, json.dumps(player_refs), plus_minus, shifts, toi,
                 gf, ga, json.dumps(extended) if extended else None)
            )
            stats_imported += 1

        except Exception as e:
            errors.append(f"Row {i+1}: {str(e)}")

    conn.commit()
    conn.close()
    return {"stats_imported": stats_imported, "errors": errors}


# ============================================================
# TEMPLATES
# ============================================================

@app.get("/templates")
async def list_templates(token_data: dict = Depends(verify_token)):
    org_id = token_data["org_id"]
    conn = get_db()
    rows = conn.execute("""
        SELECT id, template_name, report_type, description, is_global, version, created_at
        FROM report_templates
        WHERE org_id = ? OR is_global = 1
        ORDER BY template_name
    """, (org_id,)).fetchall()
    conn.close()

    results = []
    for r in rows:
        d = dict(r)
        d["is_global"] = bool(d.get("is_global", 0))
        results.append(d)
    return results


# ============================================================
# HEALTH
# ============================================================

@app.get("/health")
async def health_check():
    try:
        conn = get_db()
        conn.execute("SELECT 1")
        conn.close()
        db_status = "connected"
    except Exception:
        db_status = "error"

    return {
        "status": "healthy",
        "database": db_status,
        "db_type": "sqlite",
        "timestamp": now_iso(),
        "version": "1.0.0",
    }


@app.get("/admin/backup-db")
async def backup_database(token_data: dict = Depends(verify_token)):
    """Download the full SQLite database. Admin only."""
    conn = get_db()
    user = conn.execute("SELECT email FROM users WHERE id = ?", (token_data["user_id"],)).fetchone()
    conn.close()
    if not user or user["email"] != "jason@prospectx.com":
        raise HTTPException(status_code=403, detail="Admin only")

    backup_name = f"prospectx_backup_{now_iso()[:10]}.db"
    backup_path = os.path.join(os.path.dirname(DB_FILE), backup_name)
    shutil.copy2(DB_FILE, backup_path)
    return FileResponse(backup_path, filename=backup_name, media_type="application/octet-stream")


@app.post("/admin/backup")
async def create_server_backup(token_data: dict = Depends(verify_token)):
    """Create a server-side database backup. Available to all authenticated users."""
    backup_dir = os.path.join(os.path.dirname(DB_FILE), "backups")
    os.makedirs(backup_dir, exist_ok=True)
    ts = now_iso()[:19].replace(":", "-")
    backup_name = f"prospectx_manual_{ts}.db"
    backup_path = os.path.join(backup_dir, backup_name)
    shutil.copy2(DB_FILE, backup_path)

    # Keep last 20 backups (manual + auto + startup)
    all_backups = sorted(glob.glob(os.path.join(backup_dir, "prospectx_*.db")))
    for old in all_backups[:-20]:
        os.remove(old)

    size_mb = round(os.path.getsize(backup_path) / (1024 * 1024), 1)
    logger.info("Manual backup created: %s (%.1f MB)", backup_name, size_mb)
    return {"detail": f"Backup created: {backup_name}", "size_mb": size_mb, "file": backup_name}


# ============================================================
# ANALYTICS
# ============================================================

@app.get("/analytics/filters")
async def analytics_filters(token_data: dict = Depends(verify_token)):
    """Return available filter options: leagues, teams, positions for the org."""
    org_id = token_data["org_id"]
    conn = get_db()

    # Distinct leagues (from players table)
    leagues = conn.execute("""
        SELECT DISTINCT p.current_league as league FROM players p
        WHERE p.org_id = ? AND p.current_league IS NOT NULL AND p.current_league != ''
        ORDER BY p.current_league
    """, (org_id,)).fetchall()

    # Distinct teams (from players current_team)
    teams = conn.execute("""
        SELECT DISTINCT p.current_team, p.current_league as league
        FROM players p
        WHERE p.org_id = ? AND p.current_team IS NOT NULL AND p.current_team != ''
        ORDER BY p.current_league, p.current_team
    """, (org_id,)).fetchall()

    # Distinct positions
    positions = conn.execute("""
        SELECT DISTINCT p.position FROM players p
        WHERE p.org_id = ? AND p.position IS NOT NULL AND p.position != ''
        ORDER BY p.position
    """, (org_id,)).fetchall()

    conn.close()
    return {
        "leagues": [r["league"] for r in leagues],
        "teams": [{"name": r["current_team"], "league": r["league"] or ""} for r in teams],
        "positions": [r["position"] for r in positions],
    }


@app.get("/analytics/overview")
async def analytics_overview(
    league: str = None,
    team: str = None,
    position: str = None,
    token_data: dict = Depends(verify_token),
):
    """Platform overview stats: counts, averages, distributions. Optionally filter by league/team/position."""
    org_id = token_data["org_id"]
    conn = get_db()

    # Build player filter
    p_where = ["p.org_id = ?"]
    p_params: list = [org_id]
    if league:
        p_where.append("p.current_league = ?")
        p_params.append(league)
    if team:
        p_where.append("p.current_team = ?")
        p_params.append(team)
    if position:
        p_where.append("p.position = ?")
        p_params.append(position)

    p_filter = " AND ".join(p_where)

    # Total counts (filtered)
    total_players = conn.execute(f"SELECT COUNT(*) FROM players p WHERE {p_filter}", p_params).fetchone()[0]
    total_reports = conn.execute(
        f"SELECT COUNT(*) FROM reports r JOIN players p ON p.id = r.player_id WHERE {p_filter}",
        p_params
    ).fetchone()[0] if (league or team or position) else conn.execute(
        "SELECT COUNT(*) FROM reports WHERE org_id=?", (org_id,)
    ).fetchone()[0]
    total_notes = conn.execute(
        f"SELECT COUNT(*) FROM scout_notes sn JOIN players p ON p.id = sn.player_id WHERE {p_filter}",
        p_params
    ).fetchone()[0] if (league or team or position) else conn.execute(
        "SELECT COUNT(*) FROM scout_notes WHERE org_id=?", (org_id,)
    ).fetchone()[0]
    total_teams = conn.execute(
        f"SELECT COUNT(DISTINCT p.current_team) FROM players p WHERE {p_filter} AND p.current_team IS NOT NULL AND p.current_team != ''",
        p_params
    ).fetchone()[0]

    # Players with stats (at least 5 GP)
    players_with_stats = conn.execute(f"""
        SELECT COUNT(DISTINCT ps.player_id)
        FROM player_stats ps
        JOIN players p ON p.id = ps.player_id
        WHERE {p_filter} AND ps.gp >= 5
    """, p_params).fetchone()[0]

    # Players with intelligence
    players_with_intel = conn.execute(f"""
        SELECT COUNT(DISTINCT pi.player_id)
        FROM player_intelligence pi
        JOIN players p ON p.id = pi.player_id
        WHERE {p_filter}
    """, p_params).fetchone()[0]

    # Position breakdown
    positions_data = conn.execute(f"""
        SELECT p.position, COUNT(*) as count
        FROM players p WHERE {p_filter}
        GROUP BY p.position ORDER BY count DESC
    """, p_params).fetchall()

    # Reports by type (unfiltered — these are global stats)
    reports_by_type = conn.execute("""
        SELECT report_type, COUNT(*) as count
        FROM reports WHERE org_id=?
        GROUP BY report_type ORDER BY count DESC
    """, (org_id,)).fetchall()

    # Reports by status
    reports_by_status = conn.execute("""
        SELECT status, COUNT(*) as count
        FROM reports WHERE org_id=?
        GROUP BY status ORDER BY count DESC
    """, (org_id,)).fetchall()

    conn.close()
    return {
        "total_players": total_players,
        "total_reports": total_reports,
        "total_notes": total_notes,
        "total_teams": total_teams,
        "players_with_stats": players_with_stats,
        "players_with_intelligence": players_with_intel,
        "position_breakdown": [dict(r) for r in positions_data],
        "reports_by_type": [dict(r) for r in reports_by_type],
        "reports_by_status": [dict(r) for r in reports_by_status],
    }


@app.get("/analytics/scoring-leaders")
async def analytics_scoring_leaders(
    limit: int = 20,
    position: str = None,
    team: str = None,
    league: str = None,
    min_gp: int = 5,
    token_data: dict = Depends(verify_token),
):
    """Top scorers with per-game rates — league leaderboard."""
    org_id = token_data["org_id"]
    conn = get_db()

    where = ["p.org_id = ?", "ps.gp >= ?"]
    params: list = [org_id, min_gp]

    if position:
        where.append("p.position = ?")
        params.append(position)
    if team:
        where.append("p.current_team = ?")
        params.append(team)
    if league:
        where.append("p.current_league = ?")
        params.append(league)

    rows = conn.execute(f"""
        SELECT p.id, p.first_name, p.last_name, p.position, p.current_team,
               ps.season, ps.gp, ps.g, ps.a, ps.p, ps.plus_minus, ps.pim,
               ROUND(CAST(ps.p AS REAL) / ps.gp, 2) as ppg,
               ROUND(CAST(ps.g AS REAL) / ps.gp, 2) as gpg,
               ROUND(CAST(ps.a AS REAL) / ps.gp, 2) as apg
        FROM players p
        JOIN player_stats ps ON p.id = ps.player_id
        WHERE {" AND ".join(where)}
        ORDER BY ps.p DESC, ppg DESC
        LIMIT ?
    """, params + [limit]).fetchall()
    conn.close()
    return [dict(r) for r in rows]


@app.get("/analytics/top-prospects")
async def get_top_prospects(
    limit: int = Query(default=10, ge=1, le=50),
    token_data: dict = Depends(verify_token),
):
    """Top-graded players based on scout notes overall_grade."""
    org_id = token_data["org_id"]
    user_id = token_data["user_id"]
    conn = get_db()
    rows = conn.execute("""
        SELECT p.id, p.first_name, p.last_name, p.position, p.current_team, p.current_league,
               MAX(n.overall_grade) as top_grade,
               COUNT(n.id) as note_count,
               MAX(n.created_at) as last_noted
        FROM scout_notes n
        JOIN players p ON n.player_id = p.id
        WHERE n.org_id = ?
          AND n.overall_grade IS NOT NULL
          AND (p.is_deleted = 0 OR p.is_deleted IS NULL)
          AND (n.visibility != 'PRIVATE' OR n.scout_id = ?)
        GROUP BY p.id
        ORDER BY top_grade DESC, note_count DESC
        LIMIT ?
    """, (org_id, user_id, limit)).fetchall()
    conn.close()
    return [dict(r) for r in rows]


@app.get("/analytics/team-rankings")
async def analytics_team_rankings(
    league: str = None,
    team: str = None,
    position: str = None,
    token_data: dict = Depends(verify_token),
):
    """Team aggregate stats: total goals, points, avg PPG, player counts."""
    org_id = token_data["org_id"]
    conn = get_db()

    where = ["p.org_id = ?", "p.current_team IS NOT NULL", "p.current_team != ''"]
    params: list = [org_id]
    if league:
        where.append("p.current_league = ?")
        params.append(league)
    if team:
        where.append("p.current_team = ?")
        params.append(team)
    if position:
        where.append("p.position = ?")
        params.append(position)

    rows = conn.execute(f"""
        SELECT p.current_team as team,
               COUNT(DISTINCT p.id) as roster_size,
               COUNT(DISTINCT CASE WHEN ps.gp >= 5 THEN p.id END) as qualified_players,
               SUM(ps.gp) as total_gp,
               SUM(ps.g) as total_goals,
               SUM(ps.a) as total_assists,
               SUM(ps.p) as total_points,
               ROUND(AVG(CASE WHEN ps.gp >= 5 THEN CAST(ps.p AS REAL) / ps.gp END), 2) as avg_ppg,
               ROUND(AVG(ps.plus_minus), 1) as avg_plus_minus,
               SUM(ps.pim) as total_pim
        FROM players p
        JOIN player_stats ps ON p.id = ps.player_id
        WHERE {" AND ".join(where)}
        GROUP BY p.current_team
        HAVING qualified_players >= 1
        ORDER BY total_points DESC
    """, params).fetchall()
    conn.close()
    return [dict(r) for r in rows]


@app.get("/analytics/player-compare")
async def analytics_player_compare(
    player_ids: str,
    token_data: dict = Depends(verify_token),
):
    """Compare 2-5 players side-by-side. player_ids = comma-separated IDs."""
    org_id = token_data["org_id"]
    ids = [pid.strip() for pid in player_ids.split(",") if pid.strip()][:5]
    if not ids:
        raise HTTPException(status_code=400, detail="Provide at least one player_id")

    conn = get_db()
    results = []
    for pid in ids:
        player = conn.execute(
            "SELECT id, first_name, last_name, position, current_team, dob FROM players WHERE id=? AND org_id=?",
            (pid, org_id)
        ).fetchone()
        if not player:
            continue
        p = dict(player)

        # Best season stats
        stats = conn.execute("""
            SELECT season, gp, g, a, p, plus_minus, pim,
                   ROUND(CAST(p AS REAL) / NULLIF(gp, 0), 2) as ppg,
                   ROUND(CAST(g AS REAL) / NULLIF(gp, 0), 2) as gpg,
                   shooting_pct
            FROM player_stats WHERE player_id=? AND gp >= 5
            ORDER BY p DESC LIMIT 1
        """, (pid,)).fetchone()

        # Intelligence
        intel = conn.execute("""
            SELECT archetype, overall_grade, offensive_grade, defensive_grade,
                   skating_grade, hockey_iq_grade, compete_grade, strengths, development_areas
            FROM player_intelligence WHERE player_id=?
            ORDER BY version DESC LIMIT 1
        """, (pid,)).fetchone()

        p["stats"] = dict(stats) if stats else None
        if intel:
            i = dict(intel)
            for k in ("strengths", "development_areas"):
                try:
                    i[k] = json.loads(i[k]) if i[k] else []
                except (json.JSONDecodeError, TypeError):
                    i[k] = []
            p["intelligence"] = i
        else:
            p["intelligence"] = None

        results.append(p)

    conn.close()
    return results


@app.get("/analytics/position-stats")
async def analytics_position_stats(
    league: str = None,
    team: str = None,
    token_data: dict = Depends(verify_token),
):
    """Average stats by position for league benchmarking."""
    org_id = token_data["org_id"]
    conn = get_db()

    where = ["p.org_id = ?", "ps.gp >= 5"]
    params: list = [org_id]
    if league:
        where.append("p.current_league = ?")
        params.append(league)
    if team:
        where.append("p.current_team = ?")
        params.append(team)

    rows = conn.execute(f"""
        SELECT p.position,
               COUNT(DISTINCT p.id) as player_count,
               ROUND(AVG(ps.gp), 1) as avg_gp,
               ROUND(AVG(ps.g), 1) as avg_g,
               ROUND(AVG(ps.a), 1) as avg_a,
               ROUND(AVG(ps.p), 1) as avg_p,
               ROUND(AVG(CAST(ps.p AS REAL) / NULLIF(ps.gp, 0)), 2) as avg_ppg,
               ROUND(AVG(CAST(ps.g AS REAL) / NULLIF(ps.gp, 0)), 2) as avg_gpg,
               ROUND(AVG(ps.plus_minus), 1) as avg_plus_minus,
               ROUND(AVG(ps.pim), 1) as avg_pim,
               MAX(ps.g) as max_goals,
               MAX(ps.p) as max_points,
               MAX(CAST(ps.p AS REAL) / NULLIF(ps.gp, 0)) as max_ppg
        FROM players p
        JOIN player_stats ps ON p.id = ps.player_id
        WHERE {" AND ".join(where)}
        GROUP BY p.position
        ORDER BY avg_ppg DESC
    """, params).fetchall()
    conn.close()
    return [dict(r) for r in rows]


@app.get("/analytics/grade-distribution")
async def analytics_grade_distribution(token_data: dict = Depends(verify_token)):
    """Distribution of intelligence grades across all graded players."""
    org_id = token_data["org_id"]
    conn = get_db()

    grade_fields = ["overall_grade", "offensive_grade", "defensive_grade",
                    "skating_grade", "hockey_iq_grade", "compete_grade"]

    result = {}
    for field in grade_fields:
        rows = conn.execute(f"""
            SELECT pi.{field} as grade, COUNT(*) as count
            FROM player_intelligence pi
            JOIN players p ON p.id = pi.player_id
            WHERE p.org_id = ? AND pi.{field} IS NOT NULL AND pi.{field} != 'NR'
              AND pi.id IN (SELECT id FROM player_intelligence pi2
                            WHERE pi2.player_id = pi.player_id
                            ORDER BY pi2.version DESC LIMIT 1)
            GROUP BY pi.{field}
            ORDER BY count DESC
        """, (org_id,)).fetchall()
        result[field] = [dict(r) for r in rows]

    conn.close()
    return result


@app.get("/analytics/archetype-breakdown")
async def analytics_archetype_breakdown(
    league: str = None,
    team: str = None,
    position: str = None,
    token_data: dict = Depends(verify_token),
):
    """Count of players by AI-assigned archetype."""
    org_id = token_data["org_id"]
    conn = get_db()

    where = ["p.org_id = ?", "pi.archetype IS NOT NULL",
             "pi.id IN (SELECT id FROM player_intelligence pi2 WHERE pi2.player_id = pi.player_id ORDER BY pi2.version DESC LIMIT 1)"]
    params: list = [org_id]
    if league:
        where.append("p.current_league = ?")
        params.append(league)
    if team:
        where.append("p.current_team = ?")
        params.append(team)
    if position:
        where.append("p.position = ?")
        params.append(position)

    rows = conn.execute(f"""
        SELECT pi.archetype, COUNT(DISTINCT pi.player_id) as count,
               ROUND(AVG(pi.archetype_confidence), 2) as avg_confidence
        FROM player_intelligence pi
        JOIN players p ON p.id = pi.player_id
        WHERE {" AND ".join(where)}
        GROUP BY pi.archetype
        ORDER BY count DESC
    """, params).fetchall()
    conn.close()
    return [dict(r) for r in rows]


@app.get("/analytics/scoring-distribution")
async def analytics_scoring_distribution(
    min_gp: int = 5,
    league: str = None,
    team: str = None,
    position: str = None,
    token_data: dict = Depends(verify_token),
):
    """Points-per-game distribution for histogram/scatter charts."""
    org_id = token_data["org_id"]
    conn = get_db()

    where = ["p.org_id = ?", "ps.gp >= ?"]
    params: list = [org_id, min_gp]
    if league:
        where.append("p.current_league = ?")
        params.append(league)
    if team:
        where.append("p.current_team = ?")
        params.append(team)
    if position:
        where.append("p.position = ?")
        params.append(position)

    rows = conn.execute(f"""
        SELECT p.id, p.first_name, p.last_name, p.position, p.current_team,
               ps.gp, ps.g, ps.a, ps.p, ps.plus_minus,
               ROUND(CAST(ps.p AS REAL) / ps.gp, 3) as ppg,
               ROUND(CAST(ps.g AS REAL) / ps.gp, 3) as gpg
        FROM players p
        JOIN player_stats ps ON p.id = ps.player_id
        WHERE {" AND ".join(where)}
        ORDER BY ppg DESC
    """, params).fetchall()
    conn.close()
    return [dict(r) for r in rows]


@app.get("/analytics/tag-cloud")
async def analytics_tag_cloud(
    league: str = None,
    team: str = None,
    position: str = None,
    token_data: dict = Depends(verify_token),
):
    """Frequency of scout note tags and intelligence tags across all players."""
    org_id = token_data["org_id"]
    conn = get_db()

    # Build player filter for joins
    p_where = ["p.org_id = ?"]
    p_params: list = [org_id]
    if league:
        p_where.append("p.current_league = ?")
        p_params.append(league)
    if team:
        p_where.append("p.current_team = ?")
        p_params.append(team)
    if position:
        p_where.append("p.position = ?")
        p_params.append(position)

    p_filter = " AND ".join(p_where)

    # Scout note tags
    note_rows = conn.execute(f"""
        SELECT sn.tags FROM scout_notes sn
        JOIN players p ON p.id = sn.player_id
        WHERE {p_filter}
    """, p_params).fetchall()

    tag_counts = {}
    for row in note_rows:
        try:
            tags = json.loads(row["tags"]) if row["tags"] else []
        except (json.JSONDecodeError, TypeError):
            tags = []
        for tag in tags:
            tag_counts[tag] = tag_counts.get(tag, 0) + 1

    # Intelligence tags
    intel_rows = conn.execute(f"""
        SELECT pi.tags FROM player_intelligence pi
        JOIN players p ON p.id = pi.player_id
        WHERE {p_filter}
          AND pi.id IN (SELECT id FROM player_intelligence pi2
                        WHERE pi2.player_id = pi.player_id
                        ORDER BY pi2.version DESC LIMIT 1)
    """, p_params).fetchall()

    intel_tag_counts = {}
    for row in intel_rows:
        try:
            tags = json.loads(row["tags"]) if row["tags"] else []
        except (json.JSONDecodeError, TypeError):
            tags = []
        for tag in tags:
            intel_tag_counts[tag] = intel_tag_counts.get(tag, 0) + 1

    conn.close()
    return {
        "scout_note_tags": [{"tag": k, "count": v} for k, v in sorted(tag_counts.items(), key=lambda x: -x[1])],
        "intelligence_tags": [{"tag": k, "count": v} for k, v in sorted(intel_tag_counts.items(), key=lambda x: -x[1])],
    }


# ============================================================
# PROSPECTX METRICS ENGINE
# ============================================================
# Six proprietary metrics calculated from aggregate stats.
# Each metric is 0-100 scale with league percentile context.
# These are ProspectX's competitive moat — no competitor has these.

def _calc_percentile(value: float, all_values: list) -> int:
    """Calculate percentile rank (0-100) of value within all_values."""
    if not all_values or value is None:
        return 0
    below = sum(1 for v in all_values if v < value)
    return min(99, max(1, round((below / len(all_values)) * 100)))


def _compute_prospectx_indices(player_stats: dict, position: str, league_stats: list) -> dict:
    """
    Compute 6 ProspectX Metrics from aggregate stats.
    Returns dict with metric values (0-100) + percentiles.

    Metrics:
    1. SniperMetric      — Pure goal-scoring ability & finishing
    2. PlaymakerMetric   — Passing, assists, vision
    3. TransitionMetric  — Two-way play, offensive impact while being reliable
    4. DefensiveMetric   — Defensive reliability & impact
    5. CompeteMetric     — Physical engagement, discipline, toughness
    6. HockeyIQMetric    — Smart play indicators (efficiency, +/-, situation reads)
    """
    gp = max(player_stats.get("gp", 1), 1)
    g = player_stats.get("g", 0)
    a = player_stats.get("a", 0)
    p = player_stats.get("p", 0)
    pm = player_stats.get("plus_minus", 0)
    pim = player_stats.get("pim", 0)
    shots = player_stats.get("sog", 0) or player_stats.get("shots", 0) or 0
    shoot_pct = player_stats.get("shooting_pct", None)

    # Per-game rates
    gpg = g / gp
    apg = a / gp
    ppg = p / gp
    pm_pg = pm / gp
    pim_pg = pim / gp
    shots_pg = shots / gp if shots else 0

    # Calculate shooting % if not provided
    if shoot_pct is None and shots > 0:
        shoot_pct = (g / shots) * 100
    elif shoot_pct is None:
        shoot_pct = 0

    # Check for extended stats (InStat data)
    ext = player_stats.get("extended_stats", {}) or {}
    has_ext = bool(ext)
    puck_battles_won_pct = None
    takeaways = None
    entries_carry = None
    breakouts = None
    corsi_pct = None

    if has_ext:
        pb = ext.get("puck_battles", {})
        if pb:
            puck_battles_won_pct = pb.get("won_pct")
        rec = ext.get("recoveries", {})
        if rec:
            takeaways = rec.get("takeaways", 0)
        ent = ext.get("entries", {})
        if ent:
            entries_carry = ent.get("via_stickhandling", 0)
            breakouts = ent.get("breakouts_total", 0)
        adv = ext.get("advanced", {})
        if adv:
            corsi_pct = adv.get("corsi_pct")

    is_defense = position in ("D", "LD", "RD")

    # ── 1. SNIPER INDEX ────────────────────────────────────
    # Weights: GPG (35%), Shooting% (30%), Shots/GP (20%), Goal:Assist ratio (15%)
    goal_assist_ratio = g / max(a, 1) if g > 0 else 0
    sniper_raw = (
        min(gpg / 0.6, 1.0) * 35 +          # 0.6 GPG = max
        min(shoot_pct / 18, 1.0) * 30 +      # 18% = max
        min(shots_pg / 3.5, 1.0) * 20 +      # 3.5 SOG/GP = max
        min(goal_assist_ratio / 1.5, 1.0) * 15  # More goals than assists = sniper
    )
    if is_defense:
        sniper_raw = sniper_raw * 1.3  # Boost for D who score (harder to do)
    sniper_index = min(99, max(1, round(sniper_raw)))

    # ── 2. PLAYMAKER INDEX ─────────────────────────────────
    # Weights: APG (35%), Assist:Goal ratio (25%), PPG (20%), ext data (20%)
    assist_goal_ratio = a / max(g, 1) if a > 0 else 0
    playmaker_ext = 0
    if has_ext:
        passes = ext.get("passes", {})
        if passes:
            acc_pct = passes.get("accurate_pct", 0) or 0
            to_slot = passes.get("to_slot", 0) or 0
            playmaker_ext = min(acc_pct / 80, 1.0) * 10 + min(to_slot / (gp * 2), 1.0) * 10
    else:
        playmaker_ext = min(apg / 0.8, 1.0) * 20  # Fallback: just use APG more

    playmaker_raw = (
        min(apg / 1.0, 1.0) * 35 +             # 1.0 APG = max
        min(assist_goal_ratio / 2.5, 1.0) * 25 + # 2.5:1 A:G ratio = pure playmaker
        min(ppg / 1.2, 1.0) * 20 +               # Overall production matters
        playmaker_ext
    )
    playmaker_index = min(99, max(1, round(playmaker_raw)))

    # ── 3. TRANSITION INDEX ────────────────────────────────
    # Two-way impact — offensive contribution while maintaining defensive value
    # Weights: PPG (25%), +/- per game (30%), discipline (15%), ext data (30%)
    discipline_score = max(0, 1 - pim_pg / 3.0)  # 0 PIM = 1.0, 3 PIM/GP = 0.0

    transition_ext = 0
    if has_ext and entries_carry is not None and breakouts is not None:
        entries_pg = entries_carry / gp
        breakouts_pg = breakouts / gp
        transition_ext = (
            min(entries_pg / 3.0, 1.0) * 15 +    # Zone entry carry-ins
            min(breakouts_pg / 5.0, 1.0) * 15     # Breakout plays
        )
    elif corsi_pct and corsi_pct > 0:
        transition_ext = min(corsi_pct / 60, 1.0) * 30  # Corsi as proxy
    else:
        # Fallback: use +/- more heavily
        transition_ext = max(0, min((pm_pg + 0.5) / 1.0, 1.0)) * 30

    transition_raw = (
        min(ppg / 1.0, 1.0) * 25 +
        max(0, min((pm_pg + 0.5) / 1.5, 1.0)) * 30 +  # +/- centered at -0.5
        discipline_score * 15 +
        transition_ext
    )
    if is_defense:
        transition_raw = transition_raw * 1.1  # D bonus for transition play
    transition_index = min(99, max(1, round(transition_raw)))

    # ── 4. DEFENSIVE INDEX ─────────────────────────────────
    # Weights: +/- per game (35%), discipline (20%), ext data (45% if available)
    defense_ext = 0
    if has_ext:
        if puck_battles_won_pct and puck_battles_won_pct > 0:
            defense_ext += min(puck_battles_won_pct / 60, 1.0) * 20
        if takeaways and takeaways > 0:
            takeaways_pg = takeaways / gp
            defense_ext += min(takeaways_pg / 2.0, 1.0) * 15
        if corsi_pct and corsi_pct > 0:
            defense_ext += min(corsi_pct / 58, 1.0) * 10
    else:
        # Without extended stats, lean heavier on +/- and discipline
        defense_ext = max(0, min((pm_pg + 0.3) / 1.0, 1.0)) * 25 + discipline_score * 20

    defensive_raw = (
        max(0, min((pm_pg + 0.3) / 1.0, 1.0)) * 35 +
        discipline_score * 20 +
        defense_ext
    )
    if is_defense:
        defensive_raw = defensive_raw * 1.15  # D expected to grade higher
    defensive_index = min(99, max(1, round(defensive_raw)))

    # ── 5. COMPETE INDEX ───────────────────────────────────
    # Physical engagement, toughness, willingness to battle
    # Weights: PIM balance (30%), +/- (20%), ext data (50% if available)
    # Smart physicality: some PIMs are good (engaged), too many are bad
    pim_balance = 1.0 - abs(pim_pg - 0.8) / 2.0  # Optimal ~0.8 PIM/GP
    pim_balance = max(0, min(1, pim_balance))

    compete_ext = 0
    if has_ext:
        if puck_battles_won_pct and puck_battles_won_pct > 0:
            compete_ext += min(puck_battles_won_pct / 55, 1.0) * 25
        main = ext.get("main", {})
        hits = main.get("hits", 0) or 0
        if hits > 0:
            hits_pg = hits / gp
            compete_ext += min(hits_pg / 3.0, 1.0) * 15
        blocked = ext.get("shots", {}).get("shots_blocking", 0) or 0
        if blocked > 0:
            blocked_pg = blocked / gp
            compete_ext += min(blocked_pg / 2.0, 1.0) * 10
    else:
        # Without extended stats: use PIM as engagement proxy + production
        compete_ext = pim_balance * 30 + min(ppg / 0.8, 1.0) * 20

    compete_raw = (
        pim_balance * 30 +
        max(0, min((pm_pg + 0.3) / 0.8, 1.0)) * 20 +
        compete_ext
    )
    compete_index = min(99, max(1, round(compete_raw)))

    # ── 6. HOCKEY IQ INDEX ─────────────────────────────────
    # Smart player indicators: efficiency, +/- vs production, situation reads
    # High IQ = good +/- relative to ice time, efficient scoring, low turnovers
    efficiency = ppg / max(shots_pg, 0.5) if shots_pg > 0 else ppg  # Points per shot
    pm_vs_production = pm_pg / max(ppg, 0.1)  # +/- relative to offense

    iq_ext = 0
    if has_ext:
        passes = ext.get("passes", {})
        acc_pct = passes.get("accurate_pct", 0) or 0 if passes else 0
        rec = ext.get("recoveries", {})
        puck_losses = rec.get("puck_losses", 0) or 0 if rec else 0
        puck_losses_pg = puck_losses / gp if puck_losses > 0 else 0
        iq_ext = (
            min(acc_pct / 75, 1.0) * 15 +                      # Pass accuracy
            max(0, 1 - puck_losses_pg / 4.0) * 10 +            # Low turnovers
            (min(corsi_pct / 58, 1.0) * 10 if corsi_pct else 5)  # Possession
        )
    else:
        iq_ext = min(efficiency / 0.5, 1.0) * 20 + max(0, min(pm_vs_production / 2, 1.0)) * 15

    iq_raw = (
        min(efficiency / 0.4, 1.0) * 25 +           # Scoring efficiency
        max(0, min((pm_pg + 0.3) / 1.0, 1.0)) * 25 + # +/- as IQ proxy
        discipline_score * 15 +                       # Smart discipline
        iq_ext
    )
    iq_index = min(99, max(1, round(iq_raw)))

    # ── Compute league percentiles ─────────────────────────
    # Calculate same indices for all league players for percentile context
    league_sniper = []
    league_playmaker = []
    league_transition = []
    league_defensive = []
    league_compete = []
    league_iq = []

    for ls in league_stats:
        lgp = max(ls.get("gp", 1), 1)
        lg = ls.get("g", 0)
        la = ls.get("a", 0)
        lp = ls.get("p", 0)
        lpm = ls.get("plus_minus", 0)
        lpim = ls.get("pim", 0)
        lshots = ls.get("sog", 0) or ls.get("shots", 0) or 0
        lshoot = ls.get("shooting_pct", None)
        if lshoot is None and lshots > 0:
            lshoot = (lg / lshots) * 100
        elif lshoot is None:
            lshoot = 0

        lgpg = lg / lgp
        lapg = la / lgp
        lppg = lp / lgp
        lpm_pg = lpm / lgp
        lpim_pg = lpim / lgp
        lshots_pg = lshots / lgp if lshots else 0
        lgar = lg / max(la, 1) if lg > 0 else 0
        lagr = la / max(lg, 1) if la > 0 else 0
        ldisc = max(0, 1 - lpim_pg / 3.0)
        leff = lppg / max(lshots_pg, 0.5) if lshots_pg > 0 else lppg

        l_is_d = ls.get("position", "") in ("D", "LD", "RD")
        d_mult = 1.3 if l_is_d else 1.0
        d_mult2 = 1.1 if l_is_d else 1.0
        d_mult3 = 1.15 if l_is_d else 1.0

        s = min(99, max(1, round((min(lgpg/0.6,1)*35 + min(lshoot/18,1)*30 + min(lshots_pg/3.5,1)*20 + min(lgar/1.5,1)*15) * d_mult)))
        pm_val = min(99, max(1, round(min(lapg/1.0,1)*35 + min(lagr/2.5,1)*25 + min(lppg/1.2,1)*20 + min(lapg/0.8,1)*20)))
        t = min(99, max(1, round((min(lppg/1.0,1)*25 + max(0,min((lpm_pg+0.5)/1.5,1))*30 + ldisc*15 + max(0,min((lpm_pg+0.5)/1.0,1))*30) * d_mult2)))
        de = min(99, max(1, round((max(0,min((lpm_pg+0.3)/1.0,1))*35 + ldisc*20 + max(0,min((lpm_pg+0.3)/1.0,1))*25 + ldisc*20) * d_mult3)))
        lpim_b = max(0, min(1, 1-abs(lpim_pg-0.8)/2.0))
        co = min(99, max(1, round(lpim_b*30 + max(0,min((lpm_pg+0.3)/0.8,1))*20 + lpim_b*30 + min(lppg/0.8,1)*20)))
        iq = min(99, max(1, round(min(leff/0.4,1)*25 + max(0,min((lpm_pg+0.3)/1.0,1))*25 + ldisc*15 + min(leff/0.5,1)*20 + max(0,min(lpm_pg/max(lppg,0.1)/2,1))*15)))

        league_sniper.append(s)
        league_playmaker.append(pm_val)
        league_transition.append(t)
        league_defensive.append(de)
        league_compete.append(co)
        league_iq.append(iq)

    return {
        "sniper": {
            "value": sniper_index,
            "percentile": _calc_percentile(sniper_index, league_sniper),
            "label": "SniperMetric",
            "description": "Goal-scoring ability and finishing efficiency",
        },
        "playmaker": {
            "value": playmaker_index,
            "percentile": _calc_percentile(playmaker_index, league_playmaker),
            "label": "PlaymakerMetric",
            "description": "Passing, assists, and offensive vision",
        },
        "transition": {
            "value": transition_index,
            "percentile": _calc_percentile(transition_index, league_transition),
            "label": "TransitionMetric",
            "description": "Two-way impact and zone transition play",
        },
        "defensive": {
            "value": defensive_index,
            "percentile": _calc_percentile(defensive_index, league_defensive),
            "label": "DefensiveMetric",
            "description": "Defensive reliability and suppression",
        },
        "compete": {
            "value": compete_index,
            "percentile": _calc_percentile(compete_index, league_compete),
            "label": "CompeteMetric",
            "description": "Physical engagement and battle level",
        },
        "hockey_iq": {
            "value": iq_index,
            "percentile": _calc_percentile(iq_index, league_iq),
            "label": "HockeyIQMetric",
            "description": "Decision-making, efficiency, and smart play",
        },
    }


@app.get("/analytics/player-indices/{player_id}")
async def get_player_indices(player_id: str, token_data: dict = Depends(verify_token)):
    """Compute ProspectX Metrics for a single player with league percentiles."""
    org_id = token_data["org_id"]
    conn = get_db()

    # Get player
    player = conn.execute(
        "SELECT * FROM players WHERE id=? AND org_id=?", (player_id, org_id)
    ).fetchone()
    if not player:
        conn.close()
        raise HTTPException(status_code=404, detail="Player not found")
    player = dict(player)

    # Get player's best season stats
    stats_row = conn.execute("""
        SELECT * FROM player_stats
        WHERE player_id=? AND gp >= 5
        ORDER BY p DESC LIMIT 1
    """, (player_id,)).fetchone()

    if not stats_row:
        conn.close()
        raise HTTPException(status_code=400, detail="Player has insufficient stats (min 5 GP)")

    player_stats = dict(stats_row)
    # Parse extended stats
    if player_stats.get("extended_stats"):
        try:
            player_stats["extended_stats"] = json.loads(player_stats["extended_stats"])
        except (json.JSONDecodeError, TypeError):
            player_stats["extended_stats"] = {}

    # Get all league stats for percentile calculations
    league_rows = conn.execute("""
        SELECT p.position, ps.gp, ps.g, ps.a, ps.p, ps.plus_minus, ps.pim,
               ps.shots, ps.sog, ps.shooting_pct, ps.extended_stats
        FROM players p
        JOIN player_stats ps ON p.id = ps.player_id
        WHERE p.org_id = ? AND ps.gp >= 5
        ORDER BY ps.p DESC
    """, (org_id,)).fetchall()

    league_stats = []
    for r in league_rows:
        d = dict(r)
        if d.get("extended_stats"):
            try:
                d["extended_stats"] = json.loads(d["extended_stats"])
            except (json.JSONDecodeError, TypeError):
                d["extended_stats"] = {}
        league_stats.append(d)

    conn.close()

    indices = _compute_prospectx_indices(player_stats, player.get("position", "F"), league_stats)

    return {
        "player_id": player_id,
        "player_name": f"{player.get('first_name', '')} {player.get('last_name', '')}",
        "position": player.get("position", "F"),
        "season": player_stats.get("season"),
        "gp": player_stats.get("gp", 0),
        "indices": indices,
        "has_extended_stats": bool(player_stats.get("extended_stats")),
    }


@app.get("/analytics/league-indices")
async def get_league_indices(
    min_gp: int = 10,
    limit: int = 50,
    position: str = None,
    league: str = None,
    team: str = None,
    token_data: dict = Depends(verify_token),
):
    """Compute ProspectX Metrics for all qualified players — league-wide view."""
    org_id = token_data["org_id"]
    conn = get_db()

    where = ["p.org_id = ?", "ps.gp >= ?"]
    params: list = [org_id, min_gp]
    if position:
        where.append("p.position = ?")
        params.append(position)
    if league:
        where.append("p.current_league = ?")
        params.append(league)
    if team:
        where.append("p.current_team = ?")
        params.append(team)

    rows = conn.execute(f"""
        SELECT p.id, p.first_name, p.last_name, p.position, p.current_team,
               ps.gp, ps.g, ps.a, ps.p, ps.plus_minus, ps.pim,
               ps.shots, ps.sog, ps.shooting_pct, ps.extended_stats, ps.season
        FROM players p
        JOIN player_stats ps ON p.id = ps.player_id
        WHERE {" AND ".join(where)}
        ORDER BY ps.p DESC
        LIMIT ?
    """, params + [limit]).fetchall()

    all_stats = []
    for r in rows:
        d = dict(r)
        if d.get("extended_stats"):
            try:
                d["extended_stats"] = json.loads(d["extended_stats"])
            except (json.JSONDecodeError, TypeError):
                d["extended_stats"] = {}
        all_stats.append(d)

    conn.close()

    results = []
    for ps in all_stats:
        indices = _compute_prospectx_indices(ps, ps.get("position", "F"), all_stats)
        results.append({
            "player_id": ps["id"],
            "player_name": f"{ps.get('first_name', '')} {ps.get('last_name', '')}",
            "position": ps.get("position", "F"),
            "current_team": ps.get("current_team"),
            "gp": ps.get("gp", 0),
            "p": ps.get("p", 0),
            "indices": indices,
        })

    return results


# ============================================================
# PLAYER MANAGEMENT — Duplicates, Merge, Bulk Operations
# ============================================================

@app.get("/players/duplicates")
async def find_duplicate_players(token_data: dict = Depends(verify_token)):
    """Detect potential duplicate player profiles.
    Groups players by exact name match and by same-last-name + same-DOB.
    Returns groups of likely duplicates for admin review."""
    org_id = token_data["org_id"]
    conn = get_db()

    rows = conn.execute("""
        SELECT p.id, p.first_name, p.last_name, p.current_team, p.current_league,
               p.position, p.dob, p.shoots, p.created_at,
               (SELECT COUNT(*) FROM player_stats ps WHERE ps.player_id = p.id) as stat_count,
               (SELECT COUNT(*) FROM scout_notes sn WHERE sn.player_id = p.id) as note_count,
               (SELECT COUNT(*) FROM reports r WHERE r.player_id = p.id) as report_count,
               (SELECT COUNT(*) FROM player_intelligence pi WHERE pi.player_id = p.id) as intel_count
        FROM players p WHERE p.org_id = ? AND (p.is_deleted = 0 OR p.is_deleted IS NULL)
        ORDER BY p.last_name, p.first_name
    """, (org_id,)).fetchall()

    # Group 1: Exact first+last name matches
    from collections import defaultdict
    by_full_name = defaultdict(list)
    for r in rows:
        key = f"{r['first_name'].lower().strip()} {r['last_name'].lower().strip()}"
        by_full_name[key].append(dict(r))

    # Group 2: Same last name + same DOB (different first name — name variants)
    by_last_dob = defaultdict(list)
    for r in rows:
        dob = r["dob"] or ""
        if dob and dob not in ("-", ""):
            key = f"{r['last_name'].lower().strip()}|{dob}"
            by_last_dob[key].append(dict(r))

    duplicate_groups = []

    # Add exact name matches
    for name, group in by_full_name.items():
        if len(group) > 1:
            duplicate_groups.append({
                "match_type": "exact_name",
                "match_key": name,
                "confidence": "high",
                "players": group,
            })

    # Add DOB matches (only if not already covered by exact name)
    exact_ids = set()
    for g in duplicate_groups:
        for p in g["players"]:
            exact_ids.add(p["id"])

    for key, group in by_last_dob.items():
        if len(group) > 1:
            # Check if first names differ (otherwise it's already in exact matches)
            names = set(p["first_name"].lower().strip() for p in group)
            if len(names) > 1:
                # Only add if not all players are already in exact-name groups
                group_ids = set(p["id"] for p in group)
                if not group_ids.issubset(exact_ids):
                    duplicate_groups.append({
                        "match_type": "name_variant",
                        "match_key": key,
                        "confidence": "medium",
                        "players": group,
                    })

    # Sort by confidence (high first), then by player count
    duplicate_groups.sort(key=lambda g: (0 if g["confidence"] == "high" else 1, -len(g["players"])))

    conn.close()
    return {
        "total_groups": len(duplicate_groups),
        "total_duplicate_players": sum(len(g["players"]) for g in duplicate_groups),
        "groups": duplicate_groups,
    }


@app.post("/players/merge")
async def merge_players(
    request: dict = Body(...),
    token_data: dict = Depends(verify_token),
):
    """Merge duplicate player profiles into a single canonical record.

    Body: {
        "keep_id": "player-id-to-keep",
        "merge_ids": ["player-id-to-merge-1", "player-id-to-merge-2"],
        "update_fields": { optional fields to update on the kept player }
    }

    All stats, notes, reports, and intelligence from merge_ids are reassigned to keep_id.
    The merged player records are then deleted.
    """
    org_id = token_data["org_id"]
    keep_id = request.get("keep_id")
    merge_ids = request.get("merge_ids", [])
    update_fields = request.get("update_fields", {})

    if not keep_id or not merge_ids:
        raise HTTPException(status_code=400, detail="Provide keep_id and at least one merge_id")

    if keep_id in merge_ids:
        raise HTTPException(status_code=400, detail="keep_id cannot be in merge_ids")

    conn = get_db()

    # Verify all players exist and belong to this org
    all_ids = [keep_id] + merge_ids
    placeholders = ",".join(["?"] * len(all_ids))
    found = conn.execute(
        f"SELECT id FROM players WHERE id IN ({placeholders}) AND org_id = ?",
        all_ids + [org_id]
    ).fetchall()
    found_ids = set(r["id"] for r in found)

    missing = set(all_ids) - found_ids
    if missing:
        conn.close()
        raise HTTPException(status_code=404, detail=f"Players not found: {', '.join(missing)}")

    stats_moved = 0
    notes_moved = 0
    reports_moved = 0
    intel_moved = 0

    for mid in merge_ids:
        # Reassign stats
        result = conn.execute(
            "UPDATE player_stats SET player_id = ? WHERE player_id = ?",
            (keep_id, mid)
        )
        stats_moved += result.rowcount

        # Reassign scout notes
        result = conn.execute(
            "UPDATE scout_notes SET player_id = ? WHERE player_id = ?",
            (keep_id, mid)
        )
        notes_moved += result.rowcount

        # Reassign reports
        result = conn.execute(
            "UPDATE reports SET player_id = ? WHERE player_id = ?",
            (keep_id, mid)
        )
        reports_moved += result.rowcount

        # Reassign intelligence records
        result = conn.execute(
            "UPDATE player_intelligence SET player_id = ? WHERE player_id = ?",
            (keep_id, mid)
        )
        intel_moved += result.rowcount

        # Soft-delete the merged player record (instead of hard delete)
        conn.execute("""
            UPDATE players SET is_deleted = 1, is_merged = 1, merged_into = ?,
            merged_at = CURRENT_TIMESTAMP, deleted_at = CURRENT_TIMESTAMP,
            deleted_reason = 'Merged into another player'
            WHERE id = ?
        """, (keep_id, mid))

    # Optionally update fields on the kept player
    allowed_fields = {"first_name", "last_name", "position", "shoots", "dob",
                      "current_team", "current_league", "height_cm", "weight_kg",
                      "birth_year", "age_group", "draft_eligible_year", "league_tier"}
    updates = []
    params = []
    for field, value in update_fields.items():
        if field in allowed_fields:
            updates.append(f"{field} = ?")
            params.append(value)
    if updates:
        updates.append("updated_at = CURRENT_TIMESTAMP")
        params.append(keep_id)
        conn.execute(f"UPDATE players SET {', '.join(updates)} WHERE id = ?", params)

    # Insert audit record into player_merges
    import json as _json
    user_id = token_data["user_id"]
    merge_id = str(uuid.uuid4())
    conn.execute("""
        INSERT INTO player_merges (id, org_id, primary_player_id, duplicate_player_ids,
            stats_moved, notes_moved, reports_moved, intel_moved, merged_by, merged_at,
            can_undo, undo_before)
        VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, CURRENT_TIMESTAMP, 1,
            datetime('now', '+30 days'))
    """, (merge_id, org_id, keep_id, _json.dumps(merge_ids),
          stats_moved, notes_moved, reports_moved, intel_moved, user_id))

    conn.commit()
    conn.close()

    return {
        "status": "merged",
        "merge_id": merge_id,
        "kept_player_id": keep_id,
        "merged_player_ids": merge_ids,
        "stats_moved": stats_moved,
        "notes_moved": notes_moved,
        "reports_moved": reports_moved,
        "intel_moved": intel_moved,
    }


@app.get("/merges")
async def list_merges(token_data: dict = Depends(verify_token)):
    """List merge history for the organization."""
    org_id = token_data["org_id"]
    conn = get_db()
    rows = conn.execute("""
        SELECT m.*, p.first_name, p.last_name
        FROM player_merges m
        LEFT JOIN players p ON m.primary_player_id = p.id
        WHERE m.org_id = ?
        ORDER BY m.merged_at DESC
    """, (org_id,)).fetchall()
    conn.close()
    import json as _json
    result = []
    for r in rows:
        dup_ids = _json.loads(r["duplicate_player_ids"]) if r["duplicate_player_ids"] else []
        result.append({
            "id": r["id"],
            "primary_player_id": r["primary_player_id"],
            "primary_player_name": f"{r['first_name']} {r['last_name']}" if r["first_name"] else "Unknown",
            "duplicate_player_ids": dup_ids,
            "stats_moved": r["stats_moved"],
            "notes_moved": r["notes_moved"],
            "reports_moved": r["reports_moved"],
            "intel_moved": r["intel_moved"],
            "merged_by": r["merged_by"],
            "merged_at": r["merged_at"],
            "can_undo": bool(r["can_undo"]) and r["undone_at"] is None and (r["undo_before"] or "") >= datetime.now().isoformat(),
            "undo_before": r["undo_before"],
            "undone_at": r["undone_at"],
        })
    return result


@app.post("/merges/{merge_id}/undo")
async def undo_merge(merge_id: str, token_data: dict = Depends(verify_token)):
    """Undo a merge — restores the duplicate player records (data stays with primary)."""
    org_id = token_data["org_id"]
    conn = get_db()
    merge = conn.execute(
        "SELECT * FROM player_merges WHERE id = ? AND org_id = ?",
        (merge_id, org_id)
    ).fetchone()
    if not merge:
        conn.close()
        raise HTTPException(status_code=404, detail="Merge record not found")
    if merge["undone_at"]:
        conn.close()
        raise HTTPException(status_code=400, detail="This merge has already been undone")
    if not merge["can_undo"]:
        conn.close()
        raise HTTPException(status_code=400, detail="This merge cannot be undone")
    if (merge["undo_before"] or "") < datetime.now().isoformat():
        conn.close()
        raise HTTPException(status_code=400, detail="Undo window has expired (30 days)")

    import json as _json
    dup_ids = _json.loads(merge["duplicate_player_ids"]) if merge["duplicate_player_ids"] else []

    # Restore the duplicate player records (un-delete them)
    restored = 0
    for did in dup_ids:
        result = conn.execute("""
            UPDATE players SET is_deleted = 0, is_merged = 0, merged_into = NULL,
            merged_at = NULL, deleted_at = NULL, deleted_reason = NULL
            WHERE id = ? AND org_id = ?
        """, (did, org_id))
        restored += result.rowcount

    # Mark merge as undone
    conn.execute(
        "UPDATE player_merges SET undone_at = CURRENT_TIMESTAMP WHERE id = ?",
        (merge_id,)
    )

    conn.commit()
    conn.close()
    return {
        "status": "undone",
        "merge_id": merge_id,
        "players_restored": restored,
        "note": "Player records restored. Stats/notes/reports remain with the primary player.",
    }


@app.delete("/players/{player_id}")
async def delete_player(
    player_id: str,
    reason: Optional[str] = Query(default=None),
    token_data: dict = Depends(verify_token),
):
    """Soft-delete a player. Data is preserved and recoverable for 30 days."""
    org_id = token_data["org_id"]
    user_id = token_data["user_id"]
    conn = get_db()

    player = conn.execute(
        "SELECT id, first_name, last_name FROM players WHERE id = ? AND org_id = ? AND (is_deleted = 0 OR is_deleted IS NULL)",
        (player_id, org_id)
    ).fetchone()
    if not player:
        conn.close()
        raise HTTPException(status_code=404, detail="Player not found")

    conn.execute("""
        UPDATE players SET is_deleted = 1, deleted_at = CURRENT_TIMESTAMP,
        deleted_reason = ?, deleted_by = ? WHERE id = ?
    """, (reason or "Manual deletion", user_id, player_id))

    conn.commit()
    conn.close()

    return {
        "status": "deleted",
        "player_id": player_id,
        "player_name": f"{player['first_name']} {player['last_name']}",
        "recoverable_until": "30 days from now",
    }


@app.post("/players/{player_id}/corrections")
async def submit_correction(
    player_id: str,
    request: dict = Body(...),
    token_data: dict = Depends(verify_token),
):
    """Submit a correction for a player's data. Any authenticated user (novice+) can submit."""
    org_id = token_data["org_id"]
    user_id = token_data["user_id"]
    conn = get_db()

    _check_tier_permission(user_id, "can_submit_corrections", conn)

    # Verify player exists
    player = conn.execute(
        "SELECT id FROM players WHERE id = ? AND org_id = ? AND (is_deleted = 0 OR is_deleted IS NULL)",
        (player_id, org_id)
    ).fetchone()
    if not player:
        conn.close()
        raise HTTPException(status_code=404, detail="Player not found")

    field_name = request.get("field_name")
    new_value = request.get("new_value")
    reason = request.get("reason", "")
    confidence = request.get("confidence", "medium")

    CORRECTABLE_FIELDS = [
        "first_name", "last_name", "position", "shoots", "dob",
        "current_team", "current_league", "height_cm", "weight_kg",
        "commitment_status", "image_url"
    ]

    if not field_name or field_name not in CORRECTABLE_FIELDS:
        conn.close()
        raise HTTPException(status_code=400, detail=f"Invalid field. Must be one of: {', '.join(CORRECTABLE_FIELDS)}")

    # Get current value
    player_full = conn.execute("SELECT * FROM players WHERE id = ?", (player_id,)).fetchone()
    old_value = str(player_full[field_name]) if player_full[field_name] is not None else ""

    correction_id = str(uuid.uuid4())
    conn.execute("""
        INSERT INTO player_corrections (id, org_id, user_id, player_id, field_name,
            old_value, new_value, reason, confidence, status, created_at)
        VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, 'pending', CURRENT_TIMESTAMP)
    """, (correction_id, org_id, user_id, player_id, field_name,
          old_value, str(new_value), reason, confidence))

    conn.commit()
    conn.close()
    return {
        "id": correction_id,
        "status": "pending",
        "field_name": field_name,
        "old_value": old_value,
        "new_value": str(new_value),
    }


@app.get("/corrections")
async def list_corrections(
    status: Optional[str] = Query(default=None, pattern="^(pending|approved|rejected)$"),
    limit: int = Query(default=50, ge=1, le=200),
    skip: int = Query(default=0, ge=0),
    token_data: dict = Depends(verify_token),
):
    """List corrections for the organization. Filter by status."""
    org_id = token_data["org_id"]
    conn = get_db()
    query = """
        SELECT c.*, p.first_name, p.last_name, u.email AS submitter_email
        FROM player_corrections c
        LEFT JOIN players p ON c.player_id = p.id
        LEFT JOIN users u ON c.user_id = u.id
        WHERE c.org_id = ?
    """
    params: list = [org_id]
    if status:
        query += " AND c.status = ?"
        params.append(status)
    query += " ORDER BY c.created_at DESC LIMIT ? OFFSET ?"
    params.extend([limit, skip])
    rows = conn.execute(query, params).fetchall()

    # Total count for pagination
    count_query = "SELECT COUNT(*) FROM player_corrections WHERE org_id = ?"
    count_params: list = [org_id]
    if status:
        count_query += " AND status = ?"
        count_params.append(status)
    total = conn.execute(count_query, count_params).fetchone()[0]

    conn.close()
    return {
        "total": total,
        "corrections": [dict(r) for r in rows],
    }


@app.put("/corrections/{correction_id}/review")
async def review_correction(
    correction_id: str,
    request: dict = Body(...),
    token_data: dict = Depends(verify_token),
):
    """Approve or reject a correction. On approval, auto-applies the change to the player."""
    org_id = token_data["org_id"]
    user_id = token_data["user_id"]
    action = request.get("action")  # "approve" or "reject"
    review_note = request.get("review_note", "")

    if action not in ("approve", "reject"):
        raise HTTPException(status_code=400, detail="Action must be 'approve' or 'reject'")

    conn = get_db()
    correction = conn.execute(
        "SELECT * FROM player_corrections WHERE id = ? AND org_id = ?",
        (correction_id, org_id)
    ).fetchone()
    if not correction:
        conn.close()
        raise HTTPException(status_code=404, detail="Correction not found")
    if correction["status"] != "pending":
        conn.close()
        raise HTTPException(status_code=400, detail=f"Correction already {correction['status']}")

    new_status = "approved" if action == "approve" else "rejected"

    conn.execute("""
        UPDATE player_corrections SET status = ?, reviewed_at = CURRENT_TIMESTAMP,
        reviewed_by = ?, review_note = ? WHERE id = ?
    """, (new_status, user_id, review_note, correction_id))

    # If approved, auto-apply the correction to the player
    if action == "approve":
        field = correction["field_name"]
        new_val = correction["new_value"]
        player_id = correction["player_id"]

        # Handle numeric fields
        if field in ("height_cm", "weight_kg"):
            try:
                new_val = int(new_val)
            except (ValueError, TypeError):
                pass

        conn.execute(f"UPDATE players SET {field} = ?, updated_at = CURRENT_TIMESTAMP WHERE id = ?",
                     (new_val, player_id))

        # Auto-derive fields if DOB changed
        if field == "dob" and new_val:
            try:
                by = int(str(new_val)[:4])
                conn.execute("""
                    UPDATE players SET birth_year = ?, age_group = ?, draft_eligible_year = ?
                    WHERE id = ?
                """, (by, _get_age_group(by), by + 18, player_id))
            except (ValueError, IndexError):
                pass

        # If league changed, update tier
        if field == "current_league":
            conn.execute("UPDATE players SET league_tier = ? WHERE id = ?",
                         (_get_league_tier(str(new_val)), player_id))

    conn.commit()
    conn.close()
    return {
        "id": correction_id,
        "status": new_status,
        "applied": action == "approve",
    }


@app.get("/players/{player_id}/corrections")
async def get_player_corrections(player_id: str, token_data: dict = Depends(verify_token)):
    """Get all corrections for a specific player."""
    org_id = token_data["org_id"]
    conn = get_db()
    rows = conn.execute("""
        SELECT c.*, u.email AS submitter_email
        FROM player_corrections c
        LEFT JOIN users u ON c.user_id = u.id
        WHERE c.player_id = ? AND c.org_id = ?
        ORDER BY c.created_at DESC
    """, (player_id, org_id)).fetchall()
    conn.close()
    return [dict(r) for r in rows]


@app.post("/players/bulk-assign-league")
async def bulk_assign_league(
    request: dict = Body(...),
    token_data: dict = Depends(verify_token),
):
    """Assign league to players based on their team membership.
    Uses the reference teams table to look up the league for each team.

    Body: { "league": "GOHL" }  — optional, if not provided uses reference_teams table
    """
    org_id = token_data["org_id"]
    league = request.get("league")
    conn = get_db()

    if league:
        # Simple: set league for all players in the org that don't have one
        result = conn.execute(
            "UPDATE players SET current_league = ? WHERE org_id = ? AND (current_league IS NULL OR current_league = '')",
            (league, org_id)
        )
        updated = result.rowcount
    else:
        # Smart: look up league from reference_teams table by team name
        ref_teams = conn.execute(
            "SELECT name, league FROM reference_teams"
        ).fetchall()
        team_league_map = {r["name"].lower(): r["league"] for r in ref_teams}

        # Get all players without a league
        players = conn.execute(
            "SELECT id, current_team FROM players WHERE org_id = ? AND (current_league IS NULL OR current_league = '') AND current_team IS NOT NULL AND current_team != ''",
            (org_id,)
        ).fetchall()

        updated = 0
        for p in players:
            team = p["current_team"].lower()
            matched_league = team_league_map.get(team)
            if not matched_league:
                # Try partial match
                for ref_name, ref_league in team_league_map.items():
                    if ref_name in team or team in ref_name:
                        matched_league = ref_league
                        break
            if matched_league:
                conn.execute(
                    "UPDATE players SET current_league = ? WHERE id = ?",
                    (matched_league, p["id"])
                )
                updated += 1

    conn.commit()
    conn.close()
    return {"status": "updated", "players_updated": updated}


@app.post("/players/auto-assign-teams")
async def auto_assign_teams_from_stats(
    token_data: dict = Depends(verify_token),
):
    """Auto-assign current_team and current_league to players based on their most recent stats.

    For each player, finds the most recent stat row that has a data_source containing
    team information, and updates the player's current_team accordingly.
    Also uses the reference_teams table to assign leagues.
    """
    org_id = token_data["org_id"]
    conn = get_db()

    # Build team-to-league lookup from reference teams
    ref_teams = conn.execute("SELECT name, league FROM reference_teams").fetchall()
    team_league_map = {}
    for r in ref_teams:
        team_league_map[r["name"].lower()] = r["league"]

    # Find players and their teams from latest import data
    # The import functions set current_team during import, so look at which team
    # was last assigned based on stat import timestamp
    players = conn.execute("""
        SELECT p.id, p.first_name, p.last_name, p.current_team, p.current_league,
               (SELECT ps.data_source FROM player_stats ps
                WHERE ps.player_id = p.id ORDER BY ps.rowid DESC LIMIT 1) as latest_source
        FROM players p
        WHERE p.org_id = ?
    """, (org_id,)).fetchall()

    updated_team = 0
    updated_league = 0

    for p in players:
        pid = p["id"]
        needs_league = not p["current_league"] or p["current_league"] == ""

        if needs_league and p["current_team"]:
            team_lower = p["current_team"].lower()
            league = team_league_map.get(team_lower)
            if not league:
                for ref_name, ref_league in team_league_map.items():
                    if ref_name in team_lower or team_lower in ref_name:
                        league = ref_league
                        break
            if league:
                conn.execute(
                    "UPDATE players SET current_league = ? WHERE id = ?",
                    (league, pid)
                )
                updated_league += 1

    conn.commit()
    conn.close()
    return {
        "status": "updated",
        "teams_assigned": updated_team,
        "leagues_assigned": updated_league,
    }


# ============================================================
# HockeyTech Live League API
# ============================================================

try:
    from hockeytech import HockeyTechClient, LEAGUES as HT_LEAGUES
except ImportError:
    HockeyTechClient = None
    HT_LEAGUES = {}
    logging.getLogger("prospectx").warning(
        "hockeytech module not available — live league data disabled"
    )

@app.get("/hockeytech/leagues")
async def ht_list_leagues():
    """List all supported HockeyTech leagues."""
    return [
        {"code": code, "name": cfg["name"], "client_code": cfg["client_code"]}
        for code, cfg in HT_LEAGUES.items()
    ]

@app.get("/hockeytech/{league}/seasons")
async def ht_seasons(league: str, token_data: dict = Depends(verify_token)):
    """Get all seasons for a league."""
    try:
        client = HockeyTechClient(league)
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))
    return await client.get_seasons()

@app.get("/hockeytech/{league}/teams")
async def ht_teams(league: str, season_id: Optional[int] = None, token_data: dict = Depends(verify_token)):
    """Get teams for a league. If season_id is omitted, uses the current season."""
    try:
        client = HockeyTechClient(league)
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))
    if not season_id:
        season_id = await client.get_current_season_id()
        if not season_id:
            raise HTTPException(status_code=404, detail="No current season found")
    return await client.get_teams(season_id)

@app.get("/hockeytech/{league}/roster/{team_id}")
async def ht_roster(league: str, team_id: int, season_id: Optional[int] = None, token_data: dict = Depends(verify_token)):
    """Get team roster."""
    try:
        client = HockeyTechClient(league)
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))
    if not season_id:
        season_id = await client.get_current_season_id()
        if not season_id:
            raise HTTPException(status_code=404, detail="No current season found")
    return await client.get_roster(team_id, season_id)

@app.get("/hockeytech/{league}/stats/skaters")
async def ht_skater_stats(league: str, season_id: Optional[int] = None,
                           team_id: Optional[int] = None, limit: int = 100, token_data: dict = Depends(verify_token)):
    """Get skater stats for a league or team."""
    try:
        client = HockeyTechClient(league)
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))
    if not season_id:
        season_id = await client.get_current_season_id()
        if not season_id:
            raise HTTPException(status_code=404, detail="No current season found")
    return await client.get_skater_stats(season_id, team_id=team_id, limit=limit)

@app.get("/hockeytech/{league}/stats/leaders")
async def ht_top_scorers(league: str, season_id: Optional[int] = None, limit: int = 50, token_data: dict = Depends(verify_token)):
    """Get league-wide scoring leaders (cross-team)."""
    try:
        client = HockeyTechClient(league)
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))
    if not season_id:
        season_id = await client.get_current_season_id()
        if not season_id:
            raise HTTPException(status_code=404, detail="No current season found")
    return await client.get_top_scorers(season_id, limit=limit)

@app.get("/hockeytech/{league}/stats/goalies")
async def ht_goalie_stats(league: str, season_id: Optional[int] = None,
                           team_id: Optional[int] = None, limit: int = 50, token_data: dict = Depends(verify_token)):
    """Get goalie stats for a league or team."""
    try:
        client = HockeyTechClient(league)
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))
    if not season_id:
        season_id = await client.get_current_season_id()
        if not season_id:
            raise HTTPException(status_code=404, detail="No current season found")
    return await client.get_goalie_stats(season_id, team_id=team_id, limit=limit)

@app.get("/hockeytech/{league}/standings")
async def ht_standings(league: str, season_id: Optional[int] = None, token_data: dict = Depends(verify_token)):
    """Get league standings."""
    try:
        client = HockeyTechClient(league)
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))
    if not season_id:
        season_id = await client.get_current_season_id()
        if not season_id:
            raise HTTPException(status_code=404, detail="No current season found")
    return await client.get_standings(season_id)

@app.get("/hockeytech/{league}/scorebar")
async def ht_scorebar(league: str, days_back: int = 1, days_ahead: int = 3, token_data: dict = Depends(verify_token)):
    """Get recent and upcoming games."""
    try:
        client = HockeyTechClient(league)
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))
    return await client.get_scorebar(days_back=days_back, days_ahead=days_ahead)

@app.get("/hockeytech/{league}/player/{player_id}")
async def ht_player_profile(league: str, player_id: int, token_data: dict = Depends(verify_token)):
    """Get player profile from HockeyTech."""
    try:
        client = HockeyTechClient(league)
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))
    return await client.get_player_profile(player_id)

@app.get("/hockeytech/{league}/player/{player_id}/gamelog")
async def ht_player_gamelog(league: str, player_id: int, season_id: Optional[int] = None, token_data: dict = Depends(verify_token)):
    """Get player's game-by-game stats."""
    try:
        client = HockeyTechClient(league)
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))
    if not season_id:
        season_id = await client.get_current_season_id()
        if not season_id:
            raise HTTPException(status_code=404, detail="No current season found")
    return await client.get_player_game_log(player_id, season_id)

@app.get("/hockeytech/{league}/game/{game_id}")
async def ht_game_summary(league: str, game_id: int, token_data: dict = Depends(verify_token)):
    """Get full game summary."""
    try:
        client = HockeyTechClient(league)
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))
    return await client.get_game_summary(game_id)


# ── HockeyTech Roster Sync ──────────────────────────────────

def _parse_ht_height_to_cm(height_str: str) -> int | None:
    """Convert HockeyTech height like '5-11' or '6.01' or '5\\'11\"' to cm."""
    if not height_str:
        return None
    height_str = height_str.strip().replace("'", "-").replace('"', '').replace(".", "-")
    m = re.match(r"(\d+)-(\d+)", height_str)
    if m:
        feet, inches = int(m.group(1)), int(m.group(2))
        return round((feet * 12 + inches) * 2.54)
    return None

def _parse_ht_weight_to_kg(weight_str: str) -> int | None:
    """Convert HockeyTech weight (lbs string) to kg."""
    if not weight_str:
        return None
    try:
        lbs = int(re.sub(r"[^\d]", "", str(weight_str)))
        return round(lbs * 0.4536) if lbs > 0 else None
    except (ValueError, TypeError):
        return None

def _normalize_position(pos: str) -> str:
    """Map HockeyTech position strings to standard codes."""
    if not pos:
        return "F"
    p = pos.upper().strip()
    mapping = {
        "LEFT WING": "LW", "RIGHT WING": "RW", "CENTER": "C", "CENTRE": "C",
        "DEFENSE": "D", "DEFENCE": "D", "GOALIE": "G", "GOALTENDER": "G",
        "FORWARD": "F", "LW": "LW", "RW": "RW", "C": "C", "D": "D", "G": "G", "F": "F",
    }
    return mapping.get(p, "F")

def _normalize_dob(dob_str: str) -> str | None:
    """Normalize HockeyTech DOB to YYYY-MM-DD."""
    if not dob_str:
        return None
    # Already YYYY-MM-DD
    if re.match(r"\d{4}-\d{2}-\d{2}", dob_str):
        return dob_str[:10]
    # Try common formats
    for fmt in ("%B %d, %Y", "%b %d, %Y", "%m/%d/%Y", "%d/%m/%Y"):
        try:
            return datetime.strptime(dob_str.strip(), fmt).strftime("%Y-%m-%d")
        except ValueError:
            continue
    return dob_str


@app.post("/hockeytech/{league}/sync-roster/{team_id}")
async def ht_sync_roster(league: str, team_id: int, season_id: Optional[int] = None,
                          sync_stats: bool = False,
                          token_data: dict = Depends(verify_token)):
    """Sync a HockeyTech team roster into the ProspectX player database.

    For each player on the roster:
    - If already linked by hockeytech_id → update bio fields
    - If fuzzy name+DOB match found → link and update
    - Otherwise → create new player

    Returns summary of created, updated, and linked players.
    """
    org_id = token_data["org_id"]
    user_id = token_data["user_id"]

    # ── Tier permission check: sync requires Pro+ ──
    perm_conn = get_db()
    try:
        _check_tier_permission(user_id, "can_sync_data", perm_conn)
    finally:
        perm_conn.close()

    try:
        client = HockeyTechClient(league)
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))

    if not season_id:
        season_id = await client.get_current_season_id()
        if not season_id:
            raise HTTPException(status_code=404, detail="No current season found")

    roster = await client.get_roster(team_id, season_id)
    if not roster:
        raise HTTPException(status_code=404, detail="Empty roster returned from HockeyTech")

    # Get league display name for current_league field
    league_name = HT_LEAGUES.get(league, {}).get("name", league.upper())

    created, updated, skipped = 0, 0, 0
    results = []
    team_name_synced = ""
    logo_synced = False

    with safe_db() as conn:
        now = datetime.now(timezone.utc).isoformat()

        for rp in roster:
            ht_id = rp.get("id")
            first = (rp.get("first_name") or "").strip()
            last = (rp.get("last_name") or "").strip()
            if not first or not last:
                skipped += 1
                continue

            position = _normalize_position(rp.get("position", ""))
            dob = _normalize_dob(rp.get("dob", ""))
            height_cm = _parse_ht_height_to_cm(rp.get("height", ""))
            weight_kg = _parse_ht_weight_to_kg(rp.get("weight", ""))
            shoots = (rp.get("shoots") or "")[:1].upper()
            team_name = rp.get("team_name", "")
            photo_url = rp.get("photo", "")
            # Filter out HockeyTech "no photo" placeholders — treat as empty
            if photo_url and "nophoto" in photo_url.lower():
                photo_url = ""
            jersey = rp.get("jersey", "")

            # Derive fields
            birth_year = int(dob[:4]) if dob and len(dob) >= 4 else None
            age_group = _get_age_group(birth_year) if birth_year else None
            league_tier = _get_league_tier(league_name)

            # 1. Check by hockeytech_id first (exact match)
            existing = conn.execute(
                "SELECT id, first_name, last_name FROM players WHERE hockeytech_id = ? AND org_id = ?",
                (ht_id, org_id)
            ).fetchone()

            if existing:
                # Update bio fields
                conn.execute("""
                    UPDATE players SET
                        current_team = ?, current_league = ?, position = ?, shoots = ?,
                        height_cm = COALESCE(?, height_cm), weight_kg = COALESCE(?, weight_kg),
                        dob = COALESCE(?, dob), image_url = COALESCE(NULLIF(?, ''), image_url),
                        birth_year = COALESCE(?, birth_year), age_group = COALESCE(?, age_group),
                        league_tier = COALESCE(?, league_tier), hockeytech_league = ?,
                        updated_at = ?
                    WHERE id = ?
                """, (team_name, league_name, position, shoots, height_cm, weight_kg,
                      dob, photo_url, birth_year, age_group, league_tier, league, now, existing[0]))
                updated += 1
                results.append({"name": f"{first} {last}", "action": "updated", "player_id": existing[0]})
                continue

            # 2. Fuzzy match by name + DOB
            candidates = conn.execute(
                "SELECT id, first_name, last_name, dob FROM players WHERE org_id = ? AND LOWER(last_name) = LOWER(?)",
                (org_id, last)
            ).fetchall()

            matched_id = None
            for c in candidates:
                # Exact first name match (case insensitive)
                if c[1].lower() == first.lower():
                    matched_id = c[0]
                    break
                # First name starts with same letters (handle nicknames like "Mike" vs "Michael")
                if dob and c[3] == dob and (c[1].lower().startswith(first[:3].lower()) or first.lower().startswith(c[1][:3].lower())):
                    matched_id = c[0]
                    break

            if matched_id:
                # Link and update
                conn.execute("""
                    UPDATE players SET
                        hockeytech_id = ?, hockeytech_league = ?,
                        current_team = ?, current_league = ?, position = ?, shoots = ?,
                        height_cm = COALESCE(?, height_cm), weight_kg = COALESCE(?, weight_kg),
                        dob = COALESCE(?, dob), image_url = COALESCE(NULLIF(?, ''), image_url),
                        birth_year = COALESCE(?, birth_year), age_group = COALESCE(?, age_group),
                        league_tier = COALESCE(?, league_tier),
                        updated_at = ?
                    WHERE id = ?
                """, (ht_id, league, team_name, league_name, position, shoots, height_cm, weight_kg,
                      dob, photo_url, birth_year, age_group, league_tier, now, matched_id))
                updated += 1
                results.append({"name": f"{first} {last}", "action": "linked+updated", "player_id": matched_id})
                continue

            # 3. Create new player
            player_id = gen_id()
            conn.execute("""
                INSERT INTO players (id, org_id, first_name, last_name, dob, position, shoots,
                                    height_cm, weight_kg, current_team, current_league, image_url,
                                    hockeytech_id, hockeytech_league,
                                    birth_year, age_group, league_tier,
                                    tags, passports, created_at, updated_at)
                VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, '[]', '[]', ?, ?)
            """, (player_id, org_id, first, last, dob, position, shoots,
                  height_cm, weight_kg, team_name, league_name, photo_url,
                  ht_id, league, birth_year, age_group, league_tier, now, now))
            created += 1
            results.append({"name": f"{first} {last}", "action": "created", "player_id": player_id})

        # ── Helper: find team row in DB using multiple name strategies ──
        team_name_synced = roster[0].get("team_name", "") if roster else ""

        def _find_team_row_for_sync(conn_inner, names_to_try_inner, org_inner, cols="id"):
            """Try multiple name strategies to find a team in the DB."""
            for try_name in names_to_try_inner:
                if not try_name:
                    continue
                row = conn_inner.execute(
                    f"SELECT {cols} FROM teams WHERE LOWER(name) = LOWER(?) AND org_id IN (?, '__global__')",
                    (try_name, org_inner)
                ).fetchone()
                if row:
                    return row
            return None

        # Build list of names to try: HT team name, then synced players' current_team
        names_to_try = [team_name_synced]
        for r in results[:10]:
            pid = r.get("player_id")
            if pid:
                p_row = conn.execute("SELECT current_team FROM players WHERE id = ?", (pid,)).fetchone()
                if p_row and p_row["current_team"] and p_row["current_team"] not in names_to_try:
                    names_to_try.append(p_row["current_team"])

        # ── Sync team logo from HockeyTech if we don't have one locally ──
        if team_name_synced:
            try:
                ht_teams = await client.get_teams(season_id)
                for ht_team in ht_teams:
                    if ht_team.get("id") == team_id and ht_team.get("logo"):
                        ht_logo_url = ht_team["logo"]
                        # Check if this team exists in our DB and needs a logo
                        team_row = _find_team_row_for_sync(conn, names_to_try, org_id, "id, logo_url")
                        if team_row:
                            # Download the logo locally
                            async with httpx.AsyncClient(timeout=10) as dl_client:
                                img_resp = await dl_client.get(ht_logo_url)
                                if img_resp.status_code == 200:
                                    ext = "png" if "png" in ht_logo_url.lower() else "jpg"
                                    logo_filename = f"team_{team_row['id']}_ht.{ext}"
                                    logo_path = os.path.join(_IMAGES_DIR, logo_filename)
                                    with open(logo_path, "wb") as f:
                                        f.write(img_resp.content)
                                    local_logo_url = f"/uploads/{logo_filename}"
                                    conn.execute("UPDATE teams SET logo_url = ? WHERE id = ?",
                                                 (local_logo_url, team_row["id"]))
                                    logo_synced = local_logo_url
                                    logger.info("Synced team logo for %s from HockeyTech", team_name_synced)
                        break
            except Exception as e:
                logger.warning("Could not sync team logo for %s: %s", team_name_synced, e)

            # Store HockeyTech team_id and league on the teams table
            try:
                team_row_ht = _find_team_row_for_sync(conn, names_to_try, org_id)
                if team_row_ht:
                    conn.execute(
                        "UPDATE teams SET hockeytech_team_id = ?, hockeytech_league = ? WHERE id = ?",
                        (team_id, league, team_row_ht["id"])
                    )
                    logger.info("Updated HT team mapping: %s → team_id=%d, league=%s", team_name_synced, team_id, league)
            except Exception as e:
                logger.warning("Could not save HT team mapping for %s: %s", team_name_synced, e)

        # safe_db() auto-commits here on success, auto-rollbacks on error

    roster_result = {
        "synced": len(results),
        "created": created,
        "updated": updated,
        "skipped": skipped,
        "team_name": team_name_synced,
        "league": league_name,
        "logo_synced": logo_synced,
        "results": results,
    }

    # Optionally sync stats after roster
    if sync_stats:
        try:
            stats_result = await ht_sync_stats(league, team_id, season_id, token_data)
            roster_result["stats_sync"] = stats_result
        except Exception as e:
            roster_result["stats_sync_error"] = str(e)

    return roster_result


@app.post("/hockeytech/detect-transfers")
async def ht_detect_transfers(token_data: dict = Depends(verify_token)):
    """Detect player transfers by comparing ProspectX DB records against current HockeyTech rosters.

    Scans all players with a hockeytech_id, queries their current team in HockeyTech,
    and flags any whose team has changed. Optionally updates the player record.
    """
    org_id = token_data["org_id"]
    user_id = token_data["user_id"]

    # ── Tier permission check: sync requires Pro+ ──
    perm_conn = get_db()
    try:
        _check_tier_permission(user_id, "can_sync_data", perm_conn)
    finally:
        perm_conn.close()

    conn = get_db()

    # Get all players linked to HockeyTech
    linked = conn.execute("""
        SELECT id, first_name, last_name, current_team, current_league,
               hockeytech_id, hockeytech_league
        FROM players
        WHERE org_id = ? AND hockeytech_id IS NOT NULL AND hockeytech_league IS NOT NULL
    """, (org_id,)).fetchall()

    if not linked:
        conn.close()
        return {"transfers": [], "checked": 0, "message": "No HockeyTech-linked players found. Sync a roster first."}

    # Group by league to minimize API clients
    by_league: Dict[str, list] = {}
    for row in linked:
        lg = row[6]  # hockeytech_league
        if lg not in by_league:
            by_league[lg] = []
        by_league[lg].append(row)

    transfers = []
    errors = []
    checked = 0
    now = datetime.now(timezone.utc).isoformat()

    for league_code, players in by_league.items():
        try:
            client = HockeyTechClient(league_code)
        except ValueError:
            errors.append(f"Unknown league code: {league_code}")
            continue

        for p in players:
            pid, first, last, old_team, old_league, ht_id, ht_league = p
            try:
                profile = await client.get_player_profile(ht_id)
                checked += 1
            except Exception as e:
                errors.append(f"Failed to fetch {first} {last} (HT#{ht_id}): {str(e)}")
                continue

            if not profile:
                continue

            # Profile may be nested: SiteKit.Player (list or dict)
            player_data = profile.get("Player", profile)
            if isinstance(player_data, list) and player_data:
                player_data = player_data[0]
            if not isinstance(player_data, dict):
                continue

            # Get current team from profile
            new_team = (player_data.get("most_recent_team_name") or
                        player_data.get("team_name") or "")
            new_league_name = HT_LEAGUES.get(league_code, {}).get("name", league_code.upper())

            if new_team and old_team and new_team.lower() != old_team.lower():
                # Transfer detected!
                transfers.append({
                    "player_id": pid,
                    "name": f"{first} {last}",
                    "hockeytech_id": ht_id,
                    "old_team": old_team,
                    "new_team": new_team,
                    "league": new_league_name,
                })

                # Auto-update the player record
                conn.execute("""
                    UPDATE players SET current_team = ?, updated_at = ? WHERE id = ?
                """, (new_team, now, pid))

    conn.commit()

    # Also check across leagues — players who left one league for another
    # Look for players in our DB that might appear on rosters of OTHER leagues
    cross_league_transfers = []
    all_ht_ids = {row[5]: row for row in linked}  # ht_id → player row

    # For each league we support, check if any of our tracked players appear on a different league's roster
    # This is expensive so we only do it if user has players in multiple leagues
    leagues_used = set(row[6] for row in linked)
    all_league_codes = set(HT_LEAGUES.keys())
    other_leagues = all_league_codes - leagues_used

    # We skip the cross-league deep scan here to keep it fast.
    # The per-player profile check above already catches within-league moves.

    conn.close()

    return {
        "transfers": transfers,
        "checked": checked,
        "auto_updated": len(transfers),
        "errors": errors if errors else None,
        "message": f"Checked {checked} players. Found {len(transfers)} transfer(s)." +
                   (f" {len(errors)} errors." if errors else ""),
    }


# ============================================================
# HOCKEYTECH STATS SYNC — SEASON STATS + GAME LOGS
# ============================================================


@app.post("/hockeytech/{league}/sync-stats/{team_id}")
async def ht_sync_stats(league: str, team_id: int, season_id: Optional[int] = None,
                         token_data: dict = Depends(verify_token)):
    """Sync season stats for all HT-linked players on a team.

    For each matched player:
    - Upserts player_stats (current season, data_source='hockeytech')
    - Appends to player_stats_history (snapshot, never deleted)
    """
    org_id = token_data["org_id"]
    user_id = token_data["user_id"]

    # Tier permission check
    perm_conn = get_db()
    try:
        _check_tier_permission(user_id, "can_sync_data", perm_conn)
    finally:
        perm_conn.close()

    try:
        client = HockeyTechClient(league)
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))

    if not season_id:
        season_id = await client.get_current_season_id()
        if not season_id:
            raise HTTPException(status_code=404, detail="No current season found")

    # Derive season name from HT seasons list
    seasons = await client.get_seasons()
    season_name = ""
    for s in seasons:
        if s.get("id") == season_id:
            season_name = s.get("name", str(season_id))
            break

    # Fetch stats from HT
    skater_stats = await client.get_skater_stats(season_id, team_id=team_id, limit=200)
    goalie_stats_list = await client.get_goalie_stats(season_id, team_id=team_id, limit=50)

    conn = get_db()
    now = datetime.now(timezone.utc)
    now_iso = now.isoformat()
    today = now.strftime("%Y-%m-%d")

    synced_skaters = 0
    synced_goalies = 0
    snapshots_created = 0
    skipped = 0

    try:
        # Get all HT-linked players for this org
        ht_players = {}
        rows = conn.execute(
            "SELECT id, hockeytech_id FROM players WHERE org_id = ? AND hockeytech_id IS NOT NULL",
            (org_id,)
        ).fetchall()
        for r in rows:
            ht_players[r["hockeytech_id"]] = r["id"]

        # Sync skater stats
        for hs in skater_stats:
            ht_id = hs.get("player_id")
            if not ht_id or ht_id not in ht_players:
                skipped += 1
                continue

            player_id = ht_players[ht_id]
            gp = hs.get("gp", 0) or 0
            g = hs.get("goals", 0) or 0
            a = hs.get("assists", 0) or 0
            p = hs.get("points", 0) or 0
            plus_minus = hs.get("plus_minus", 0) or 0
            pim = hs.get("pim", 0) or 0
            shots = hs.get("shots", 0) or 0
            ppg = hs.get("ppg", 0) or 0
            ppa = hs.get("ppa", 0) or 0
            shg = hs.get("shg", 0) or 0
            gwg = hs.get("gwg", 0) or 0
            shooting_pct_str = hs.get("shooting_pct", "0")
            try:
                shooting_pct = float(str(shooting_pct_str).replace("%", "")) if shooting_pct_str else 0.0
            except (ValueError, TypeError):
                shooting_pct = 0.0

            # Upsert player_stats (current season)
            existing = conn.execute(
                "SELECT id FROM player_stats WHERE player_id = ? AND season = ? AND stat_type = 'season' AND data_source = 'hockeytech'",
                (player_id, season_name)
            ).fetchone()

            if existing:
                conn.execute("""
                    UPDATE player_stats SET gp = ?, g = ?, a = ?, p = ?, plus_minus = ?, pim = ?,
                        shots = ?, shooting_pct = ?, data_source = 'hockeytech'
                    WHERE id = ?
                """, (gp, g, a, p, plus_minus, pim, shots, shooting_pct, existing["id"]))
            else:
                conn.execute("""
                    INSERT INTO player_stats (id, player_id, season, stat_type, gp, g, a, p,
                        plus_minus, pim, shots, shooting_pct, data_source)
                    VALUES (?, ?, ?, 'season', ?, ?, ?, ?, ?, ?, ?, ?, 'hockeytech')
                """, (gen_id(), player_id, season_name, gp, g, a, p, plus_minus, pim, shots, shooting_pct))

            # Append to player_stats_history (skip if same player+season+date exists)
            already = conn.execute(
                "SELECT id FROM player_stats_history WHERE player_id = ? AND season = ? AND date_recorded = ?",
                (player_id, season_name, today)
            ).fetchone()
            if not already:
                conn.execute("""
                    INSERT INTO player_stats_history (id, player_id, season, date_recorded,
                        gp, g, a, p, plus_minus, pim, ppg, ppa, shg, gwg, shots, shooting_pct,
                        data_source, league, team_name, synced_at)
                    VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, 'hockeytech', ?, ?, ?)
                """, (gen_id(), player_id, season_name, today,
                      gp, g, a, p, plus_minus, pim, ppg, ppa, shg, gwg, shots, shooting_pct,
                      league.upper(), hs.get("team_name", ""), now_iso))
                snapshots_created += 1

            synced_skaters += 1

        # Sync goalie stats
        for gs in goalie_stats_list:
            ht_id = gs.get("player_id")
            if not ht_id or ht_id not in ht_players:
                continue

            player_id = ht_players[ht_id]
            gp = gs.get("gp", 0) or 0
            gaa_str = gs.get("gaa", "0")
            sv_pct_str = gs.get("save_pct", "0")
            wins = gs.get("wins", 0) or 0
            losses = gs.get("losses", 0) or 0
            shutouts = gs.get("shutouts", 0) or 0
            sa = gs.get("shots_against", 0) or 0
            sv = gs.get("saves", 0) or 0

            try:
                gaa = float(str(gaa_str)) if gaa_str else 0.0
            except (ValueError, TypeError):
                gaa = 0.0
            try:
                sv_pct = float(str(sv_pct_str)) if sv_pct_str else 0.0
            except (ValueError, TypeError):
                sv_pct = 0.0

            # Upsert goalie_stats
            existing_gs = conn.execute(
                "SELECT id FROM goalie_stats WHERE player_id = ? AND season = ? AND data_source = 'hockeytech'",
                (player_id, season_name)
            ).fetchone()

            if existing_gs:
                conn.execute("""
                    UPDATE goalie_stats SET gp = ?, ga = ?, sa = ?, sv = ?, sv_pct = ?, gaa = ?,
                        data_source = 'hockeytech'
                    WHERE id = ?
                """, (gp, gp * gaa if gaa else 0, sa, sv, str(sv_pct), gaa, existing_gs["id"]))
            else:
                conn.execute("""
                    INSERT INTO goalie_stats (id, player_id, org_id, season, stat_type, gp,
                        ga, sa, sv, sv_pct, gaa, data_source)
                    VALUES (?, ?, ?, ?, 'season', ?, ?, ?, ?, ?, ?, 'hockeytech')
                """, (gen_id(), player_id, org_id, season_name, gp,
                      gp * gaa if gaa else 0, sa, sv, str(sv_pct), gaa))

            synced_goalies += 1

        conn.commit()
    finally:
        conn.close()

    logger.info("HT stats sync: %d skaters, %d goalies, %d snapshots for team %d (%s)",
                synced_skaters, synced_goalies, snapshots_created, team_id, league)

    return {
        "synced_skaters": synced_skaters,
        "synced_goalies": synced_goalies,
        "snapshots_created": snapshots_created,
        "skipped": skipped,
        "season": season_name,
        "league": league.upper(),
    }


@app.post("/hockeytech/{league}/sync-gamelog/{player_id}")
async def ht_sync_gamelog(league: str, player_id: str, season_id: Optional[int] = None,
                           token_data: dict = Depends(verify_token)):
    """Sync game-by-game stats for a single HT-linked player."""
    org_id = token_data["org_id"]
    user_id = token_data["user_id"]

    # Tier permission check
    perm_conn = get_db()
    try:
        _check_tier_permission(user_id, "can_sync_data", perm_conn)
    finally:
        perm_conn.close()

    conn = get_db()
    try:
        player = conn.execute(
            "SELECT id, hockeytech_id, hockeytech_league, first_name, last_name FROM players WHERE id = ? AND org_id = ?",
            (player_id, org_id)
        ).fetchone()
        if not player:
            raise HTTPException(status_code=404, detail="Player not found")
        if not player["hockeytech_id"]:
            raise HTTPException(status_code=400, detail="Player is not linked to HockeyTech. Sync roster first.")

        ht_id = player["hockeytech_id"]
        ht_league = player["hockeytech_league"] or league

        try:
            client = HockeyTechClient(ht_league)
        except ValueError as e:
            raise HTTPException(status_code=400, detail=str(e))

        if not season_id:
            season_id = await client.get_current_season_id()
            if not season_id:
                raise HTTPException(status_code=404, detail="No current season found")

        # Derive season name
        seasons = await client.get_seasons()
        season_name = ""
        for s in seasons:
            if s.get("id") == season_id:
                season_name = s.get("name", str(season_id))
                break

        # Fetch parsed game log
        games = await client.get_parsed_game_log(ht_id, season_id)
        if not games:
            return {"games_synced": 0, "new_games": 0, "message": "No game log data returned from HockeyTech"}

        now_iso = datetime.now(timezone.utc).isoformat()
        games_synced = 0
        new_games = 0

        for g in games:
            ht_game_id = g.get("ht_game_id")

            # Upsert games table
            game_db_id = None
            if ht_game_id:
                existing_game = conn.execute(
                    "SELECT id FROM games WHERE league = ? AND ht_game_id = ?",
                    (ht_league, ht_game_id)
                ).fetchone()
                if existing_game:
                    game_db_id = existing_game["id"]
                else:
                    game_db_id = gen_id()
                    conn.execute("""
                        INSERT OR IGNORE INTO games (id, league, season, ht_game_id, game_date,
                            home_team, away_team, home_score, away_score, status, data_source)
                        VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, 'final', 'hockeytech')
                    """, (game_db_id, ht_league, season_name, ht_game_id, g["game_date"],
                          g.get("home_team", ""), g.get("away_team", ""),
                          g.get("home_score"), g.get("away_score")))

            # Upsert player_game_stats
            existing_pgs = None
            if ht_game_id:
                existing_pgs = conn.execute(
                    "SELECT id FROM player_game_stats WHERE player_id = ? AND ht_game_id = ?",
                    (player_id, ht_game_id)
                ).fetchone()

            if existing_pgs:
                conn.execute("""
                    UPDATE player_game_stats SET goals = ?, assists = ?, points = ?,
                        plus_minus = ?, pim = ?, shots = ?, ppg = ?, shg = ?, gwg = ?,
                        opponent = ?, home_away = ?, game_date = ?
                    WHERE id = ?
                """, (g["goals"], g["assists"], g["points"], g["plus_minus"],
                      g["pim"], g["shots"], g["ppg"], g["shg"], g["gwg"],
                      g["opponent"], g["home_away"], g["game_date"], existing_pgs["id"]))
            else:
                conn.execute("""
                    INSERT INTO player_game_stats (id, player_id, game_id, ht_game_id, game_date,
                        opponent, home_away, goals, assists, points, plus_minus, pim, shots,
                        ppg, shg, gwg, season, league, data_source)
                    VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, 'hockeytech')
                """, (gen_id(), player_id, game_db_id, ht_game_id, g["game_date"],
                      g["opponent"], g["home_away"], g["goals"], g["assists"], g["points"],
                      g["plus_minus"], g["pim"], g["shots"], g["ppg"], g["shg"], g["gwg"],
                      season_name, ht_league))
                new_games += 1

            games_synced += 1

        conn.commit()
    finally:
        conn.close()

    logger.info("HT gamelog sync: %d games (%d new) for player %s (%s)",
                games_synced, new_games, player_id, league)

    return {
        "games_synced": games_synced,
        "new_games": new_games,
        "player": f"{player['first_name']} {player['last_name']}",
        "season": season_name,
    }


@app.post("/hockeytech/{league}/sync-team-gamelogs/{team_id}")
async def ht_sync_team_gamelogs(league: str, team_id: int, season_id: Optional[int] = None,
                                  token_data: dict = Depends(verify_token)):
    """Batch sync game-by-game stats for all HT-linked players on a team."""
    org_id = token_data["org_id"]
    user_id = token_data["user_id"]

    # Tier permission check
    perm_conn = get_db()
    try:
        _check_tier_permission(user_id, "can_sync_data", perm_conn)
    finally:
        perm_conn.close()

    try:
        client = HockeyTechClient(league)
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))

    if not season_id:
        season_id = await client.get_current_season_id()
        if not season_id:
            raise HTTPException(status_code=404, detail="No current season found")

    # Get all HT-linked players for this team
    conn = get_db()
    try:
        ht_roster = conn.execute("""
            SELECT id, hockeytech_id, first_name, last_name
            FROM players
            WHERE org_id = ? AND hockeytech_id IS NOT NULL AND hockeytech_league = ?
        """, (org_id, league)).fetchall()
    finally:
        conn.close()

    if not ht_roster:
        return {"players_synced": 0, "total_games": 0,
                "message": "No HockeyTech-linked players found. Sync roster first."}

    players_synced = 0
    total_games = 0
    errors = []

    for player_row in ht_roster:
        try:
            # Call the single-player gamelog sync internally
            result = await ht_sync_gamelog(
                league=league,
                player_id=player_row["id"],
                season_id=season_id,
                token_data=token_data
            )
            total_games += result.get("games_synced", 0)
            players_synced += 1
        except Exception as e:
            errors.append(f"{player_row['first_name']} {player_row['last_name']}: {str(e)}")

        # Rate limit between players
        await asyncio.sleep(0.2)

    logger.info("HT team gamelogs sync: %d players, %d games for team %d (%s)",
                players_synced, total_games, team_id, league)

    return {
        "players_synced": players_synced,
        "total_games": total_games,
        "errors": errors if errors else None,
    }


@app.post("/hockeytech/{league}/sync-league")
async def ht_sync_league(league: str, season_id: Optional[int] = None,
                         sync_stats: bool = False,
                         token_data: dict = Depends(verify_token)):
    """Bulk sync all teams in a league. Fetches team list, then syncs each roster sequentially.

    Power-user feature: syncs every team in the league at once instead of one-at-a-time.
    Requires Pro tier or higher (can_bulk_sync permission).
    """
    org_id = token_data["org_id"]
    user_id = token_data["user_id"]

    # Tier permission check — bulk sync requires Pro+
    perm_conn = get_db()
    try:
        _check_tier_permission(user_id, "can_bulk_sync", perm_conn)
    finally:
        perm_conn.close()

    try:
        client = HockeyTechClient(league)
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))

    if not season_id:
        season_id = await client.get_current_season_id()
        if not season_id:
            raise HTTPException(status_code=404, detail="No current season found")

    # Get all teams in this league for the season
    all_teams = await client.get_teams(season_id)
    if not all_teams:
        raise HTTPException(status_code=404, detail="No teams found for this league/season")

    total_created = 0
    total_updated = 0
    total_skipped = 0
    teams_synced = 0
    teams_failed = 0
    team_results = []

    for team in all_teams:
        team_id = team.get("id")
        team_name = team.get("name", f"Team {team_id}")
        try:
            result = await ht_sync_roster(
                league=league,
                team_id=team_id,
                season_id=season_id,
                sync_stats=sync_stats,
                token_data=token_data
            )
            total_created += result.get("created", 0)
            total_updated += result.get("updated", 0)
            total_skipped += result.get("skipped", 0)
            teams_synced += 1
            team_results.append({
                "team_id": team_id,
                "team_name": result.get("team_name", team_name),
                "created": result.get("created", 0),
                "updated": result.get("updated", 0),
                "skipped": result.get("skipped", 0),
                "status": "success"
            })
        except Exception as e:
            teams_failed += 1
            team_results.append({
                "team_id": team_id,
                "team_name": team_name,
                "status": "failed",
                "error": str(e)
            })
            logger.warning("Bulk sync failed for team %s (%s): %s", team_name, league, str(e))

        # Rate limit between teams to avoid hammering HockeyTech API
        await asyncio.sleep(0.5)

    logger.info("Bulk league sync complete: %s — %d teams synced, %d failed, %d created, %d updated",
                league, teams_synced, teams_failed, total_created, total_updated)

    return {
        "league": league,
        "season_id": season_id,
        "teams_synced": teams_synced,
        "teams_failed": teams_failed,
        "total_created": total_created,
        "total_updated": total_updated,
        "total_skipped": total_skipped,
        "team_results": team_results,
    }


# ============================================================
# BENCH TALK — PROSPECTX AI CONVERSATION ENGINE
# ============================================================

BENCH_TALK_SYSTEM_PROMPT = """You are Bench Talk, the AI-powered hockey conversation engine behind the ProspectX intelligence platform.

# YOUR IDENTITY
- Name: Bench Talk
- Tagline: "Let's talk hockey."
- You're a hockey lifer — you talk like you've been around rinks your whole life
- You're the smartest person at the rink who also happens to be funny — you've coached, scouted, watched thousands of hours of tape, and you speak the language
- You know every formation, every penalty, every piece of slang, and every analytics term

# REPORT BOUNDARY RULE
- Bench Talk is for conversation, quick analysis, tool routing, and hockey discussion.
- If the user requests a formal scouting report (pro_skater, goalie, unified_prospect, or ANY of the 19+ ProspectX report types), you MUST use the start_report_generation tool. Do NOT write full multi-section reports inline in chat.
- If the user says "just write it here," "give me a quick report," or "skip the tool," explain that ProspectX professional reports use the report engine for consistent, decision-grade output with proper grading, evidence discipline, and system fit analysis. Then offer to generate one with the tool.
- You MAY provide quick 2-3 paragraph assessments, stat summaries, or comparison highlights in chat. The distinction is: structured multi-section reports with grades go through the engine; conversational analysis stays in chat.
- TONE ISOLATION: When a tool call produces a report or structured output, never inject hockey slang (pigeon, barnburner, turnstile, sieve, etc.) into the tool parameters or any structured data. Slang is for your conversational chat messages only.

# THE USER
- Name: {user_first_name}
- Role: {hockey_role_label}
- Address them by first name. You already know their role — no need to ask.

# PERSONALITY & TONE
You carry a natural hockey tone. You're not a robot reading stats — you're a hockey person who happens to have a database.

**Sprinkle in hockey language naturally:**
- A one-dimensional player? "He's a pigeon — sits on the hashmarks waiting for crumbs."
- A tough guy who can't skate? "The kid's got cement in his skates but he'll drop the mitts with anyone."
- A player who scores ugly goals? "He's got a nose for the net — lives in the blue paint like he's paying rent."
- A speedster? "He's got wheels — once he hits the redline, good luck catching him."
- A playmaker? "This kid sees the ice like he's got a drone view up there. Silky mitts."
- A gritty player? "He goes to the dirty areas. Not afraid to get his nose dirty."
- A soft player? "Needs to add some sandpaper to his game. A bit of a floater right now."
- A player who takes penalties? "Spends more time in the sin bin than on the ice."
- An inconsistent player? "Cherry picks when he feels like it — shows up when he wants to."
- A player who gets scored on? "He's a turnstile back there — guys are skating around him like a pylon."
- A goalie giving up goals? "Sieve mode — everything's going through."
- A player with great hands? "Silky mitts — dangles guys for fun. Undresses defenders."
- A big win? "That was a barnburner — gongshow from start to finish."
- A hard shot? "Absolute howitzer. Bar down, top cheese."
- A great assist? "What an apple — tape to tape, right on the money."
- A physical game? "Getting chippy out there. Bodies flying."
- Describing scoring? "Kid lights the lamp. Snipes corners. Got that quick release from the slot."
- Describing toughness? "A beauty — first guy in the corners, last guy off the ice."

**Keep it subtle and natural** — don't force slang every message. Mix it in like salt — enough to taste, not enough to overpower. When analyzing stats or giving serious assessments, be professional. When describing a player's style or tendencies, that's where the personality shines.

# HOCKEY VOCABULARY — USE THESE NATURALLY
**Scoring:** apple (assist), gino (goal), snipe (accurate shot), bar down (crossbar and in), top cheese/top ched (top shelf), muffin (weak shot), five-hole (between goalie's legs), one-T/one-timer, howitzer/clapper (hard shot), light the lamp, hat trick, natural hatty, barnburner (high-scoring game)
**Players:** beauty, grinder, mucker, plug, goon, pylon/cone (slow player), bender (weak skater), sieve (leaky goalie), shadow, cherry picker, turnstile, pigeon, grocery stick
**Roles:** 1C/2C/3C, two-way center (200-foot player), power forward, sniper, puck-moving D, stay-at-home D, shutdown D, energy forward, net-front presence
**Gear:** biscuit/rubber (puck), twig/lumber (stick), bucket (helmet), mitts (hands/gloves), chiclets (teeth)
**Slang:** chirp (trash talk), celly (celebration), flow/lettuce (long hair), chippy (rough game), gongshow (out of control), bag skate (punishment practice), wheels/jets (speed), coast to coast, sin bin (penalty box), warm up the bus

# TACTICAL SYSTEMS MASTERY — YOU KNOW ALL OF THIS COLD
**Power Play Formations:**
- **1-3-1:** QB at point, two flanks on dots (off-hand for one-timers), bumper in the slot, net-front. Best for teams with flank shooters. Think Ovechkin in the left circle.
- **Overload:** Three players overload one side (half-wall QB, low forward, bumper), point + weak-side support. Simpler reads, good for youth. Lots of short passes and 3-on-2s.
- **Umbrella:** Three high (point + two flanks), two low (net-front + slot). Spreads PK wide, opens point shots. Needs a bomb from the point.
- **5-on-3 Box+1:** Four players form loose box, one shooter in middle slot. Quick puck movement forces 3 killers to collapse.
- **PP Entries:** Drop-pass entry, wide lane carry, bump and kick. Read the blue-line gap — tight = chip to space, back in = walk it over.
- **Common PP Mistakes:** Overpassing on perimeter, not shooting when lanes open, unscreened point bombs, no net-front, poor spacing, forcing plays through sticks, blue-line turnovers.

**Penalty Kill Systems:**
- **Diamond PK:** One high on point, two on flanks, one low in front. Takes away point shots and circle one-timers. Vulnerable in the bumper area.
- **Box PK:** Two high, two low — maintains box shape as puck moves. Protects middle ice and net-front. Can be beaten by quick low plays.
- **Aggressive PK:** Two forwards pressure hard, trying for turnovers and shorthanded chances. High risk, high reward. Needs elite PKers with speed.
- **PK Counters vs 1-3-1:** Use diamond structure, wedge+1 on bumper, shot-lane denial, hard pressure on QB, jump bobbles, funnel to wall.

**Forecheck Systems:** 1-2-2 aggressive, 1-2-2 trap, 2-1-2 (aggressive deep), 1-3-1 (trap/counter), 1-1-3 passive trap, 1-4 (ultra-defensive)
**DZ Coverage:** Man-to-man, zone, collapsing box, swarm, hybrid
**Breakout Patterns:** Standard (D-to-D, up the wall), reverse (change point of attack), wheel (D carries), stretch (long pass)
**Forecheck Roles:** F1 (first pressure), F2 (support/contain), F3 (high safety)

# PENALTY & RULE EXPERTISE
You can explain any penalty or rule scenario:
- **Minor (2 min):** Tripping, hooking, holding, interference, slashing, roughing, delay of game, high-sticking. Ends early if scored on.
- **Double Minor (4 min):** Two consecutive 2-minute penalties — often high-sticking causing injury.
- **Major (5 min):** Fighting, boarding, spearing, dangerous fouls. Full 5 minutes regardless of goals.
- **Misconduct (10 min):** Player sits but team stays at full strength. Often for abuse of officials.
- **Game Misconduct:** Ejected for the game, substitute replaces.
- **Match Penalty:** Ejection + 5-min major served by teammate for intent to injure.
- **Penalty Shot:** Awarded when clear scoring chance illegally denied (breakaway from behind).
- **Icing:** Puck shot from behind red line past opposing goal line without touch. Faceoff in offender's zone.
- **Offside:** Attacking player enters zone before puck crosses blue line.
- **Delayed Penalty:** Play continues until offending team touches puck. Often pull goalie for extra attacker.

# ADVANCED ANALYTICS — SPEAK THESE FLUENTLY
- **Corsi (CF%):** Shot attempt differential (shots + blocks + misses). Gold standard for possession measurement.
- **Fenwick (FF%):** Unblocked shot attempts (shots + misses). Filters out shot-blocking variance.
- **PDO:** Shooting% + Save%. Measures luck/variance. 100 is average — much higher/lower tends to regress.
- **xG (Expected Goals):** Quality of scoring chances based on shot location, type, angle, and situation.
- **High-Danger Chances (HDCF):** Scoring chances from the slot and near-crease. Where real goals come from.
- **Zone Entries/Exits:** Controlled vs dump. Controlled entries create 2x more offense. Clean exits prevent sustained pressure.
- **Transition Metrics:** Speed from DZ to OZ, controlled breakout rate, neutral zone efficiency.
When analytics data is available, reference it. When it's not, note what additional data would help the analysis.

# ROLE-TAILORED APPROACH
{role_instructions}

# YOUR CAPABILITIES
You have access to the ProspectX database with:
- Junior hockey players across GOHL, OJHL, OHL and other leagues
- Full season stats, per-game breakdowns, and historical data
- ProspectX Intelligence profiles (grades, metrics, archetypes)
- 19 professional report templates
- Player comparison engine
- League leader rankings
- 60+ hockey drill library across 13 categories (skating, passing, shooting, offensive, defensive, battle, etc.)
- AI-powered practice plan generator (uses drills, roster, team systems, and glossary)
- Live league standings, schedules, and scores (via HockeyTech)
- Team line combinations, roster breakdowns, and tactical systems
- Active game plans, chalk talk sessions, and series strategies
- Scouting observations (notes, tags) across the organization
- Personal scouting watchlist with priority tracking

# TOOLS
1. **query_players** — Search with filters (position, league, team, stats)
2. **get_player_intelligence** — ProspectX grades, metrics, archetype
3. **compare_players** — Side-by-side comparison
4. **start_report_generation** — Generate a professional report (runs in the background — give them the link)
5. **league_leaders** — Top performers by stat category
6. **query_drills** — Search the drill library by category, age level, tags, intensity, or keyword
7. **generate_practice_plan** — Generate a complete AI practice plan for a team (warm-up through cool-down, using the drill library + team systems)
8. **get_player_recent_form** — Recent game log, streaks, season-over-season progression
9. **get_team_context** — Team's current lines, roster by position, tactical systems
10. **get_game_context** — Live standings, upcoming games, recent scores from HockeyTech
11. **get_coaching_prep** — Active game plans, chalk talk sessions, series strategies
12. **search_scout_notes** — Search scouting observations by player, tag, or type
13. **get_scouting_list** — User's personal scouting watchlist with priorities
14. **diagnose_player_struggles** — Analyze declining metrics, drought streaks, and potential causes for a struggling player

# TOOL ROUTING GUIDE
When the user asks about...
- Team lines, roster, systems → use `get_team_context`
- Standings, upcoming games, scores, schedule → use `get_game_context`
- Game plan, prep for tonight, series strategy, chalk talk → use `get_coaching_prep`
- Scouting notes, observations, what did we write about → use `search_scout_notes`
- Scouting list, who am I watching, priority targets → use `get_scouting_list`
- Player search, stats, who plays for a team → use `query_players`
- Deep player profile, grades, projection → use `get_player_intelligence`
- Side-by-side comparison → use `compare_players`
- Generate a report → use `start_report_generation`
- League leaders, top scorers → use `league_leaders`
- Find drills, practice activities → use `query_drills`
- Build a practice plan → use `generate_practice_plan`
- Recent form, hot streak, game log → use `get_player_recent_form`
- What's wrong with, why is struggling, diagnose player → use `diagnose_player_struggles`

# REPORT TYPES — suggest based on {hockey_role_label}'s needs:
- **Scouting:** pro_skater, unified_prospect, draft_comparative
- **Front Office:** operations, trade_target, season_intelligence
- **Coaching:** game_decision, line_chemistry, practice_plan, opponent_gameplan, st_optimization — PLUS you can generate interactive practice plans with the generate_practice_plan tool!
- **Development:** development_roadmap, season_progress
- **Family:** family_card (no jargon, encouraging, parent-friendly)
- **General:** goalie, playoff_series, goalie_tandem, team_identity, agent_pack

# PROSPECTX GRADES
A (Elite NHL) → B+ (Solid NHL) → B (Depth NHL) → B- (NHL Fringe/AHL Top) → C+ (AHL Regular) → C (AHL Depth) → D (Junior/College)

# HOW TO RESPOND
- ALWAYS use your tools — don't guess, pull the data
- Keep it conversational but backed by numbers
- 2-3 paragraphs max unless they ask for detail
- Use bullets when listing stats or comparisons
- Always include a "what next?" nudge — suggest a report, comparison, or deeper dive
- Cite "ProspectX Intelligence" when referencing grades
- When a coach asks about game prep, use get_coaching_prep to pull their actual plans
- When asked about standings or schedule, use get_game_context for live data

**NEVER:**
- Invent player data — use tools or say "I don't have that in the system"
- Generate reports without being asked
- Write full multi-section scouting reports inline — always use start_report_generation for formal reports
- Be condescending — a parent asking about their kid deserves the same respect as a pro scout
- Refuse off-topic questions — help out, then steer back to hockey

Current date: {current_date}
"""

# Role-specific instruction blocks — NOW SERVED BY pxi_prompt_core.PXI_MODE_BLOCKS
# The old BENCH_TALK_ROLE_INSTRUCTIONS dict has been replaced by 10 PXI mode blocks
# imported from pxi_prompt_core.py. See resolve_mode() and PXI_MODE_BLOCKS.


BENCH_TALK_TOOLS = [
    {
        "name": "query_players",
        "description": "Search for players in the ProspectX database with filters. Returns player profiles with basic info and stats.",
        "input_schema": {
            "type": "object",
            "properties": {
                "position": {
                    "type": "string",
                    "description": "Filter by position (C, LW, RW, D, G, F for any forward)"
                },
                "league": {
                    "type": "string",
                    "description": "Filter by league (GOHL, OJHL, OHL, etc.)"
                },
                "team": {
                    "type": "string",
                    "description": "Filter by team name (e.g. 'Chatham Maroons')"
                },
                "name": {
                    "type": "string",
                    "description": "Search by player name (partial match)"
                },
                "min_gp": {
                    "type": "integer",
                    "description": "Minimum games played"
                },
                "limit": {
                    "type": "integer",
                    "description": "Max results to return (default 10)",
                    "default": 10
                }
            }
        }
    },
    {
        "name": "get_player_intelligence",
        "description": "Get detailed ProspectX Intelligence profile for a player including grades, archetype, strengths, development areas, and stat signature.",
        "input_schema": {
            "type": "object",
            "properties": {
                "player_name": {
                    "type": "string",
                    "description": "Player's full or partial name"
                }
            },
            "required": ["player_name"]
        }
    },
    {
        "name": "compare_players",
        "description": "Compare two players side-by-side with stats and ProspectX Intelligence profiles.",
        "input_schema": {
            "type": "object",
            "properties": {
                "player1_name": {"type": "string", "description": "First player's name"},
                "player2_name": {"type": "string", "description": "Second player's name"}
            },
            "required": ["player1_name", "player2_name"]
        }
    },
    {
        "name": "start_report_generation",
        "description": "Queue a professional scouting report for generation. ONLY use when user explicitly asks for a report.",
        "input_schema": {
            "type": "object",
            "properties": {
                "player_name": {"type": "string", "description": "Player's name"},
                "report_type": {
                    "type": "string",
                    "description": "Report type to generate",
                    "enum": [
                        "pro_skater", "unified_prospect", "goalie", "game_decision",
                        "season_intelligence", "operations", "agent_pack",
                        "development_roadmap", "family_card", "trade_target",
                        "draft_comparative", "season_progress",
                        "team_identity", "opponent_gameplan", "st_optimization",
                        "line_chemistry", "playoff_series", "goalie_tandem", "practice_plan"
                    ]
                },
                "mode": {
                    "type": "string",
                    "description": "PXI mode for report tone and analytical focus. If not specified, uses the wiring table default for the report type.",
                    "enum": ["scout", "coach", "analyst", "gm", "agent", "parent", "skill_coach", "mental_coach", "broadcast", "producer"]
                }
            },
            "required": ["player_name", "report_type"]
        }
    },
    {
        "name": "league_leaders",
        "description": "Get top performers in a league by stat category (goals, assists, points, ppg).",
        "input_schema": {
            "type": "object",
            "properties": {
                "league": {
                    "type": "string",
                    "description": "League name (e.g. 'GOHL', 'OHL', 'OJHL')"
                },
                "stat": {
                    "type": "string",
                    "enum": ["goals", "assists", "points", "ppg"],
                    "default": "points",
                    "description": "Stat category to rank by"
                },
                "position": {
                    "type": "string",
                    "description": "Optional position filter"
                },
                "limit": {
                    "type": "integer",
                    "default": 10,
                    "description": "Number of results"
                }
            },
            "required": ["league"]
        }
    },
    {
        "name": "query_drills",
        "description": "Search the ProspectX drill library. Use when coaches ask about drills, practice activities, or skill development exercises.",
        "input_schema": {
            "type": "object",
            "properties": {
                "category": {
                    "type": "string",
                    "description": "Drill category: skating, passing, shooting, stickhandling, offensive, defensive, goalie, conditioning, battle, small_area_games, transition, special_teams, warm_up"
                },
                "age_level": {
                    "type": "string",
                    "description": "Age level filter: U8, U10, U12, U14, U16_U18, JUNIOR_COLLEGE_PRO"
                },
                "tags": {
                    "type": "string",
                    "description": "Comma-separated tags to search for (e.g., '1_on_1,battle_drills,compete')"
                },
                "intensity": {
                    "type": "string",
                    "description": "Intensity level: low, medium, high"
                },
                "search": {
                    "type": "string",
                    "description": "Keyword search in drill name and description"
                },
                "limit": {
                    "type": "integer",
                    "default": 10,
                    "description": "Number of results to return"
                }
            }
        }
    },
    {
        "name": "generate_practice_plan",
        "description": "Generate a complete AI-powered practice plan for a team. Uses the drill library, team roster, team systems, and hockey glossary to create a structured plan from warm-up through cool-down.",
        "input_schema": {
            "type": "object",
            "properties": {
                "team_name": {
                    "type": "string",
                    "description": "The team to build the practice plan for"
                },
                "duration_minutes": {
                    "type": "integer",
                    "default": 90,
                    "description": "Total practice duration in minutes (60, 75, 90, or 120)"
                },
                "focus_areas": {
                    "type": "array",
                    "items": {"type": "string"},
                    "description": "Focus areas: skating, passing, shooting, puck_handling, offensive_systems, defensive_systems, checking, special_teams, conditioning, compete_level, transition, battle_drills"
                },
                "age_level": {
                    "type": "string",
                    "default": "JUNIOR_COLLEGE_PRO",
                    "description": "Age level: U8, U10, U12, U14, U16_U18, JUNIOR_COLLEGE_PRO"
                },
                "notes": {
                    "type": "string",
                    "description": "Additional coaching notes or specific requests"
                }
            },
            "required": ["team_name"]
        }
    },
    {
        "name": "get_player_recent_form",
        "description": "Get a player's recent game-by-game performance, streaks, and season progression trends. Shows last N games with per-game stats, totals, averages, point/goal streaks, and season-over-season trajectory.",
        "input_schema": {
            "type": "object",
            "properties": {
                "player_name": {
                    "type": "string",
                    "description": "Player's full or partial name"
                },
                "last_n": {
                    "type": "integer",
                    "default": 5,
                    "description": "Number of recent games to analyze (1-20, default 5)"
                }
            },
            "required": ["player_name"]
        }
    },
    # ── Phase 2 Tools ────────────────────────────────────────────
    {
        "name": "get_team_context",
        "description": "Get a team's current line combinations, roster overview by position, and tactical systems (forecheck, breakout, defensive scheme). Use when asked about lines, roster, or team systems.",
        "input_schema": {
            "type": "object",
            "properties": {
                "team_name": {
                    "type": "string",
                    "description": "Team name (e.g. 'Chatham Maroons')"
                }
            },
            "required": ["team_name"]
        }
    },
    {
        "name": "get_game_context",
        "description": "Get live league standings, upcoming games, and recent scores from HockeyTech. Use when asked about standings, schedule, upcoming opponents, or recent results.",
        "input_schema": {
            "type": "object",
            "properties": {
                "league": {
                    "type": "string",
                    "default": "gojhl",
                    "description": "League code (gojhl, ohl, ojhl)"
                },
                "team_name": {
                    "type": "string",
                    "description": "Optional team name to highlight in standings"
                }
            }
        }
    },
    {
        "name": "get_coaching_prep",
        "description": "Get active game plans and series strategies. Shows upcoming game prep, opponent analysis, matchups, talking points, and active series with scores and adjustments. Use when asked about game prep, tonight's plan, or series status.",
        "input_schema": {
            "type": "object",
            "properties": {
                "team_name": {
                    "type": "string",
                    "description": "Filter by team name"
                },
                "opponent": {
                    "type": "string",
                    "description": "Filter by opponent name"
                }
            }
        }
    },
    {
        "name": "search_scout_notes",
        "description": "Search scouting observations across the organization. Filter by player name, tags (e.g. 'speed', 'compete', 'shot'), or note type. Use when asked about what scouts have observed or noted about a player.",
        "input_schema": {
            "type": "object",
            "properties": {
                "player_name": {
                    "type": "string",
                    "description": "Player name (partial match)"
                },
                "tag": {
                    "type": "string",
                    "description": "Filter by tag (e.g. 'speed', 'skating', 'compete', 'leadership')"
                },
                "note_type": {
                    "type": "string",
                    "description": "Filter by type",
                    "enum": ["general", "game", "workout", "interview"]
                },
                "limit": {
                    "type": "integer",
                    "default": 10,
                    "description": "Max notes to return (1-25)"
                }
            }
        }
    },
    {
        "name": "get_scouting_list",
        "description": "Get the user's active scouting watchlist — players they are tracking with priority levels, target reasons, and notes. Use when asked about who they're watching or scouting priorities.",
        "input_schema": {
            "type": "object",
            "properties": {
                "priority": {
                    "type": "string",
                    "description": "Filter by priority level",
                    "enum": ["high", "medium", "low"]
                },
                "limit": {
                    "type": "integer",
                    "default": 10,
                    "description": "Max items to return (1-20)"
                }
            }
        }
    },
    {
        "name": "diagnose_player_struggles",
        "description": "Analyze what's going wrong with a struggling player. Compares recent form (last N games) to season averages, identifies declining metrics, streak data, and provides a structured diagnostic with potential causes. Use when someone asks 'what's wrong with [player]', 'why is [player] struggling', or 'diagnose [player]'.",
        "input_schema": {
            "type": "object",
            "properties": {
                "player_name": {
                    "type": "string",
                    "description": "Player's full or partial name"
                },
                "look_back_games": {
                    "type": "integer",
                    "default": 10,
                    "description": "Number of recent games to analyze (5-20)"
                }
            },
            "required": ["player_name"]
        }
    }
]


# ── Bench Talk Tool Execution Functions ──────────────────────────

def _pt_query_players(params: dict, org_id: str) -> tuple[dict, dict]:
    """Search players with filters, adapted to actual ProspectX DB schema.
    Returns (tool_result, entity_refs) tuple."""
    conn = get_db()
    try:
        query = """
            SELECT p.id, p.first_name, p.last_name, p.position, p.current_team,
                   p.current_league, p.dob, p.shoots, p.height_cm, p.weight_kg,
                   p.archetype, p.image_url,
                   ps.gp, ps.g, ps.a, ps.p, ps.plus_minus, ps.pim,
                   ps.shots, ps.shooting_pct,
                   CASE WHEN ps.gp > 0 THEN ROUND(CAST(ps.p AS REAL) / ps.gp, 2) ELSE 0 END as ppg
            FROM players p
            LEFT JOIN player_stats ps ON p.id = ps.player_id AND ps.stat_type = 'season'
            WHERE p.org_id = ?
        """
        sql_params: list = [org_id]

        if params.get("position"):
            pos = params["position"].upper()
            if pos == "F":
                query += " AND p.position IN ('C', 'LW', 'RW', 'F')"
            else:
                query += " AND p.position = ?"
                sql_params.append(pos)

        if params.get("league"):
            query += " AND p.current_league LIKE ?"
            sql_params.append(f"%{params['league']}%")

        if params.get("team"):
            query += " AND p.current_team LIKE ?"
            sql_params.append(f"%{params['team']}%")

        if params.get("name"):
            query += " AND (p.first_name || ' ' || p.last_name) LIKE ?"
            sql_params.append(f"%{params['name']}%")

        if params.get("min_gp"):
            query += " AND ps.gp >= ?"
            sql_params.append(params["min_gp"])

        query += " ORDER BY ps.p DESC NULLS LAST"
        query += f" LIMIT {min(int(params.get('limit', 10)), 50)}"

        rows = conn.execute(query, sql_params).fetchall()
        players = []
        player_ids = []
        for r in rows:
            players.append({
                "id": r["id"],
                "name": f"{r['first_name']} {r['last_name']}",
                "position": r["position"],
                "team": r["current_team"],
                "league": r["current_league"],
                "dob": r["dob"],
                "shoots": r["shoots"],
                "archetype": r["archetype"],
                "gp": r["gp"],
                "g": r["g"],
                "a": r["a"],
                "p": r["p"],
                "plus_minus": r["plus_minus"],
                "pim": r["pim"],
                "ppg": r["ppg"],
                "shooting_pct": r["shooting_pct"],
            })
            player_ids.append(r["id"])

        return {"players": players, "count": len(players)}, {"player_ids": player_ids}
    finally:
        conn.close()


def _pt_get_player_intelligence(params: dict, org_id: str) -> tuple[dict, dict]:
    """Get ProspectX Intelligence profile for a player.
    Returns (tool_result, entity_refs) tuple."""
    conn = get_db()
    try:
        name = params["player_name"]
        # Find the player
        player = conn.execute("""
            SELECT p.id, p.first_name, p.last_name, p.position, p.current_team,
                   p.current_league, p.dob, p.shoots, p.height_cm, p.weight_kg,
                   p.archetype, p.image_url
            FROM players p
            WHERE p.org_id = ? AND (p.first_name || ' ' || p.last_name) LIKE ?
            LIMIT 1
        """, (org_id, f"%{name}%")).fetchone()

        if not player:
            return {"error": f"Player '{name}' not found in database."}, {}

        pid = player["id"]

        # Get stats
        stats = conn.execute("""
            SELECT gp, g, a, p, plus_minus, pim, shots, shooting_pct, toi_seconds, season
            FROM player_stats
            WHERE player_id = ? AND stat_type = 'season'
            ORDER BY season DESC LIMIT 1
        """, (pid,)).fetchone()

        # Get intelligence
        intel = conn.execute("""
            SELECT archetype, archetype_confidence, overall_grade, offensive_grade,
                   defensive_grade, skating_grade, hockey_iq_grade, compete_grade,
                   summary, strengths, development_areas, comparable_players,
                   stat_signature, tags, projection, trigger
            FROM player_intelligence
            WHERE player_id = ?
            ORDER BY version DESC LIMIT 1
        """, (pid,)).fetchone()

        # Get scout notes count
        notes_count = conn.execute(
            "SELECT COUNT(*) as cnt FROM scout_notes WHERE player_id = ?", (pid,)
        ).fetchone()["cnt"]

        result = {
            "player": {
                "name": f"{player['first_name']} {player['last_name']}",
                "position": player["position"],
                "team": player["current_team"],
                "league": player["current_league"],
                "dob": player["dob"],
                "shoots": player["shoots"],
                "height_cm": player["height_cm"],
                "weight_kg": player["weight_kg"],
            },
            "stats": None,
            "intelligence": None,
            "scout_notes_count": notes_count,
        }

        if stats:
            result["stats"] = {
                "season": stats["season"],
                "gp": stats["gp"], "g": stats["g"], "a": stats["a"], "p": stats["p"],
                "plus_minus": stats["plus_minus"], "pim": stats["pim"],
                "shots": stats["shots"], "shooting_pct": stats["shooting_pct"],
                "ppg": round(stats["p"] / stats["gp"], 2) if stats["gp"] else 0,
            }

        if intel:
            result["intelligence"] = {
                "archetype": intel["archetype"],
                "overall_grade": intel["overall_grade"],
                "offensive_grade": intel["offensive_grade"],
                "defensive_grade": intel["defensive_grade"],
                "skating_grade": intel["skating_grade"],
                "hockey_iq_grade": intel["hockey_iq_grade"],
                "compete_grade": intel["compete_grade"],
                "summary": intel["summary"],
                "strengths": json.loads(intel["strengths"]) if intel["strengths"] else [],
                "development_areas": json.loads(intel["development_areas"]) if intel["development_areas"] else [],
                "comparable_players": json.loads(intel["comparable_players"]) if intel["comparable_players"] else [],
                "projection": intel["projection"],
                "stat_signature": json.loads(intel["stat_signature"]) if intel["stat_signature"] else None,
            }

        # Add progression context from player_stats_history
        try:
            hist_rows = conn.execute("""
                SELECT psh.season, psh.gp, psh.g, psh.a, psh.p, psh.plus_minus
                FROM player_stats_history psh
                INNER JOIN (
                    SELECT season, MAX(date_recorded) as max_date
                    FROM player_stats_history WHERE player_id = ? GROUP BY season
                ) latest ON psh.season = latest.season AND psh.date_recorded = latest.max_date
                WHERE psh.player_id = ?
                ORDER BY psh.season ASC
            """, (pid, pid)).fetchall()
            if hist_rows:
                seasons = []
                for hr in hist_rows:
                    hd = dict(hr)
                    gp = hd.get("gp", 0) or 0
                    hd["ppg_rate"] = round((hd.get("p", 0) or 0) / gp, 2) if gp > 0 else 0.0
                    seasons.append(hd)
                result["progression"] = seasons
        except Exception:
            pass

        return result, {"player_ids": [pid]}
    finally:
        conn.close()


def _pt_compare_players(params: dict, org_id: str) -> tuple[dict, dict]:
    """Compare two players side-by-side.
    Returns (tool_result, entity_refs) tuple."""
    p1_result, p1_refs = _pt_get_player_intelligence({"player_name": params["player1_name"]}, org_id)
    p2_result, p2_refs = _pt_get_player_intelligence({"player_name": params["player2_name"]}, org_id)

    combined_player_ids = p1_refs.get("player_ids", []) + p2_refs.get("player_ids", [])

    if "error" in p1_result or "error" in p2_result:
        errors = []
        if "error" in p1_result:
            errors.append(p1_result["error"])
        if "error" in p2_result:
            errors.append(p2_result["error"])
        return {"error": " | ".join(errors)}, {"player_ids": combined_player_ids}

    return {"player1": p1_result, "player2": p2_result}, {"player_ids": combined_player_ids}


def _pt_start_report(params: dict, org_id: str, user_id: str) -> tuple[dict, dict]:
    """Queue and trigger a report for generation via Bench Talk.
    Creates the report record and launches background generation.
    Returns (tool_result, entity_refs) tuple."""
    conn = get_db()
    try:
        name = params["player_name"]
        player = conn.execute("""
            SELECT id, first_name, last_name FROM players
            WHERE org_id = ? AND (first_name || ' ' || last_name) LIKE ?
            LIMIT 1
        """, (org_id, f"%{name}%")).fetchone()

        if not player:
            return {"error": f"Player '{name}' not found."}, {}

        report_type = params["report_type"]
        player_id = player["id"]
        player_name = f"{player['first_name']} {player['last_name']}"

        # Get template
        template = conn.execute(
            "SELECT * FROM report_templates WHERE report_type = ? AND (org_id = ? OR is_global = 1) LIMIT 1",
            (report_type, org_id),
        ).fetchone()

        if not template:
            return {"error": f"Report template '{report_type}' not found."}, {"player_ids": [player_id]}

        report_id = gen_id()
        title = f"{template['template_name']} — {player_name}"

        conn.execute("""
            INSERT INTO reports (id, org_id, player_id, template_id, report_type, title, status, created_by, created_at)
            VALUES (?, ?, ?, ?, ?, ?, 'processing', ?, ?)
        """, (report_id, org_id, player_id, template["id"], report_type, title, user_id, now_iso()))
        conn.commit()

        # Launch background generation thread
        thread = threading.Thread(
            target=_pt_background_generate_report,
            args=(report_id, org_id, user_id, player_id, report_type, title),
            daemon=True,
        )
        thread.start()

        return {
            "report_id": report_id,
            "status": "processing",
            "message": f"Report '{title}' is now generating. You can view its progress at /reports/{report_id}. Generation typically takes 30-60 seconds.",
        }, {"player_ids": [player_id], "report_ids": [report_id]}
    finally:
        conn.close()


def _pt_background_generate_report(report_id: str, org_id: str, user_id: str, player_id: str, report_type: str, title: str):
    """Background thread: generate a report using Claude API.
    Uses the same logic as the /reports/generate endpoint."""
    conn = get_db()
    start_time = time.perf_counter()
    try:
        client = get_anthropic_client()
        if not client:
            conn.execute(
                "UPDATE reports SET status = 'failed', error_message = ?, generated_at = ? WHERE id = ?",
                ("No Anthropic API key configured.", now_iso(), report_id),
            )
            conn.commit()
            return

        # Gather player data
        player_row = conn.execute("SELECT * FROM players WHERE id = ? AND org_id = ?", (player_id, org_id)).fetchone()
        if not player_row:
            conn.execute(
                "UPDATE reports SET status = 'failed', error_message = 'Player not found', generated_at = ? WHERE id = ?",
                (now_iso(), report_id),
            )
            conn.commit()
            return

        player = _player_from_row(player_row)

        def _row_get(row, key, default=None):
            try:
                return row[key]
            except (IndexError, KeyError):
                return default

        # Gather stats
        stats_rows = conn.execute(
            "SELECT * FROM player_stats WHERE player_id = ? ORDER BY season DESC, created_at DESC",
            (player_id,),
        ).fetchall()
        stats_list = []
        for sr in stats_rows:
            stat_entry = {
                "season": sr["season"], "stat_type": sr["stat_type"],
                "gp": sr["gp"], "g": sr["g"], "a": sr["a"], "p": sr["p"],
                "plus_minus": sr["plus_minus"], "pim": sr["pim"],
                "shots": sr["shots"], "sog": sr["sog"],
                "shooting_pct": sr["shooting_pct"],
                "toi_seconds": sr["toi_seconds"],
            }
            ext_raw = _row_get(sr, "extended_stats")
            if ext_raw:
                try:
                    stat_entry["extended_stats"] = json.loads(ext_raw)
                except Exception:
                    pass
            stats_list.append(stat_entry)

        # Gather scout notes
        notes_rows = conn.execute(
            "SELECT * FROM scout_notes WHERE player_id = ? ORDER BY created_at DESC LIMIT 20",
            (player_id,),
        ).fetchall()
        notes_list = [{"note_text": n["note_text"], "note_type": n["note_type"], "tags": json.loads(n["tags"]) if n["tags"] else [], "created_at": n["created_at"]} for n in notes_rows]

        # Gather intelligence
        intel_row = conn.execute(
            "SELECT * FROM player_intelligence WHERE player_id = ? ORDER BY version DESC LIMIT 1",
            (player_id,),
        ).fetchone()
        intel_data = None
        if intel_row:
            intel_data = {
                "archetype": intel_row["archetype"],
                "overall_grade": intel_row["overall_grade"],
                "summary": intel_row["summary"],
                "strengths": json.loads(intel_row["strengths"]) if intel_row["strengths"] else [],
                "development_areas": json.loads(intel_row["development_areas"]) if intel_row["development_areas"] else [],
            }

        # Get template
        template = conn.execute(
            "SELECT * FROM report_templates WHERE report_type = ? AND (org_id = ? OR is_global = 1) LIMIT 1",
            (report_type, org_id),
        ).fetchone()

        llm_model = "claude-sonnet-4-20250514"
        player_name = f"{player['first_name']} {player['last_name']}"
        report_type_name = template['template_name'] if template else report_type

        # ── Gather additional data (team system, lines, progression, recent form) ──
        team_system_data = None
        if player.get("current_team"):
            ts_row = conn.execute(
                "SELECT * FROM team_systems WHERE team_name = ? AND org_id = ? ORDER BY created_at DESC LIMIT 1",
                (player["current_team"], org_id),
            ).fetchone()
            if ts_row:
                team_system_data = {
                    "team_name": ts_row["team_name"],
                    "forecheck": ts_row["forecheck"], "dz_coverage": ts_row["dz_coverage"],
                    "oz_setup": ts_row["oz_setup"], "pp_formation": ts_row["pp_formation"],
                    "pk_formation": ts_row["pk_formation"], "breakout": ts_row["breakout"],
                    "notes": ts_row["notes"],
                }

        # Line combinations
        line_data = []
        if player.get("current_team"):
            line_rows = conn.execute(
                "SELECT * FROM line_combinations WHERE team_name = ? AND org_id = ? ORDER BY line_type, line_number",
                (player["current_team"], org_id),
            ).fetchall()
            for lr in line_rows:
                slots = json.loads(lr["slots"]) if lr.get("slots") else {}
                if player_name in str(slots):
                    line_data.append({"line_type": lr["line_type"], "line_number": lr["line_number"], "slots": slots})

        # Historical progression
        progression = []
        for sr in stats_rows:
            if sr.get("gp") and sr["gp"] > 0:
                progression.append({
                    "season": sr["season"], "gp": sr["gp"], "g": sr["g"], "a": sr["a"], "p": sr["p"],
                    "ppg_rate": round(sr["p"] / sr["gp"], 2) if sr["p"] and sr["gp"] else None,
                })

        # Recent form (last 10 games)
        recent_form = []
        try:
            game_rows = conn.execute(
                "SELECT * FROM player_game_stats WHERE player_id = ? ORDER BY game_date DESC LIMIT 10",
                (player_id,),
            ).fetchall()
            for gr in game_rows:
                recent_form.append({
                    "game_date": gr["game_date"],
                    "opponent": gr["opponent"] if "opponent" in gr.keys() else None,
                    "g": gr["goals"], "a": gr["assists"], "p": gr["points"],
                    "plus_minus": gr["plus_minus"] if "plus_minus" in gr.keys() else None,
                    "pim": gr["pim"] if "pim" in gr.keys() else None,
                })
        except Exception:
            pass  # Table may not exist or be empty

        # ── Build system prompt (aligned with main report generator) ──
        bg_system_prompt = f"""You are ProspectX, an elite hockey scouting intelligence engine. Generate a professional-grade {report_type_name} report.

Your report must be:
- Data-driven: Reference specific stats (GP, G, A, P, +/-, PIM, S%, etc.) when available
- Tactically literate: Use real hockey language — forecheck roles, transition play, gap control, cycle game
- Professionally formatted: Use ALL_CAPS_WITH_UNDERSCORES section headers (e.g., EXECUTIVE_SUMMARY, KEY_NUMBERS, STRENGTHS, DEVELOPMENT_AREAS, SYSTEM_FIT, PROJECTION, BOTTOM_LINE)
- Specific to position: Tailor analysis to the player's position
- Honest and balanced: Give an accurate, scout-grade assessment
- Hockey vernacular: Use authentic hockey language (sniper, grinder, wheels, puck-moving D, etc.)

PROSPECT GRADING SCALE (include an overall grade):
  A = Top-Line / Franchise | A- = Top-6 / Top-4 | B+ = Middle-6 | B = Bottom-6 / Depth
  B- = NHL Fringe / AHL Top | C+ = AHL Regular | C = AHL Depth | D = Junior/College | NR = Not Rated

Include "Overall Grade: X" in EXECUTIVE_SUMMARY and BOTTOM_LINE.

EVIDENCE DISCIPLINE:
- Label inferences: "INFERENCE — [reasoning]"
- Missing data: write "DATA NOT AVAILABLE" when it materially affects analysis
- End EXECUTIVE_SUMMARY and BOTTOM_LINE with: "Confidence: HIGH | MED | LOW — [reason]"

Format each section header on its own line in ALL_CAPS_WITH_UNDERSCORES format, followed by content.
Do NOT use === delimiters. Do NOT use markdown code blocks or formatting."""

        # Append template-specific instructions if rich prompt exists
        bg_tpl_prompt = template["prompt_text"] if template else None
        if bg_tpl_prompt and len(bg_tpl_prompt) > 200:
            bg_system_prompt += f"""

TEMPLATE-SPECIFIC INSTRUCTIONS FOR {report_type_name.upper()}:
{bg_tpl_prompt}"""

        # ── Build input data ──
        input_data = {
            "player": player,
            "stats": stats_list,
            "scout_notes": notes_list,
            "intelligence": intel_data,
        }
        if team_system_data:
            input_data["team_system"] = team_system_data
        if line_data:
            input_data["line_combinations"] = line_data
        if progression:
            input_data["historical_progression"] = progression
        if recent_form:
            input_data["recent_form_last_10"] = recent_form

        user_prompt = f"Generate a {report_type_name} for the following player. Here is ALL available data:\n\n" + json.dumps(input_data, indent=2, default=str)

        response = client.messages.create(
            model=llm_model,
            max_tokens=8000,
            system=bg_system_prompt,
            messages=[{"role": "user", "content": user_prompt}],
        )

        output_text = "".join(b.text for b in response.content if b.type == "text")
        tokens_used = response.usage.input_tokens + response.usage.output_tokens
        generation_time = int((time.perf_counter() - start_time) * 1000)

        # ── Validate and optionally repair the report output (mode-aware) ──
        _bg_mode = resolve_mode(template_slug=report_type)
        output_text, validation_warnings = _validate_and_repair_report(
            output_text, report_type, client, llm_model, mode=_bg_mode
        )
        if validation_warnings:
            for w in validation_warnings:
                logger.warning("BG Report %s validation: %s", report_id, w)
                if w.startswith("SOFT:") or w.startswith("HARD:"):
                    _log_error_to_db("POST", f"/bench-talk/report/{report_id}", 200, f"BG Report validation: {w}")

        # Score quality
        quality = _score_report_quality(output_text, report_type, mode=_bg_mode)
        conn.execute("""
            UPDATE reports SET
                status = 'complete', title = ?, output_text = ?,
                llm_model = ?, llm_tokens = ?, generation_time_ms = ?,
                generated_at = ?, quality_score = ?, quality_details = ?
            WHERE id = ?
        """, (title, output_text, llm_model, tokens_used, generation_time, now_iso(),
              quality["score"], json.dumps(quality), report_id))
        conn.commit()
        logger.info("Bench Talk background report %s generated in %dms, quality=%.1f", report_id, generation_time, quality["score"])

    except Exception as e:
        logger.exception("Bench Talk background report generation failed for %s", report_id)
        try:
            conn.execute(
                "UPDATE reports SET status = 'failed', error_message = ?, generated_at = ? WHERE id = ?",
                (str(e)[:500], now_iso(), report_id),
            )
            conn.commit()
        except Exception:
            pass
    finally:
        conn.close()


def _pt_league_leaders(params: dict, org_id: str) -> tuple[dict, dict]:
    """Get league leaders by stat category.
    Returns (tool_result, entity_refs) tuple."""
    conn = get_db()
    try:
        league = params["league"]
        stat = params.get("stat", "points")
        limit = min(int(params.get("limit", 10)), 50)

        stat_col_map = {
            "goals": "ps.g",
            "assists": "ps.a",
            "points": "ps.p",
            "ppg": "CASE WHEN ps.gp > 0 THEN CAST(ps.p AS REAL) / ps.gp ELSE 0 END",
        }
        order_col = stat_col_map.get(stat, "ps.p")

        query = f"""
            SELECT p.id, p.first_name, p.last_name, p.position, p.current_team,
                   ps.gp, ps.g, ps.a, ps.p, ps.plus_minus, ps.pim,
                   CASE WHEN ps.gp > 0 THEN ROUND(CAST(ps.p AS REAL) / ps.gp, 2) ELSE 0 END as ppg
            FROM players p
            JOIN player_stats ps ON p.id = ps.player_id AND ps.stat_type = 'season'
            WHERE p.org_id = ? AND p.current_league LIKE ?
        """
        sql_params: list = [org_id, f"%{league}%"]

        if params.get("position"):
            query += " AND p.position = ?"
            sql_params.append(params["position"].upper())

        query += f" ORDER BY {order_col} DESC LIMIT ?"
        sql_params.append(limit)

        rows = conn.execute(query, sql_params).fetchall()
        leaders = []
        player_ids = []
        for i, r in enumerate(rows, 1):
            leaders.append({
                "rank": i,
                "name": f"{r['first_name']} {r['last_name']}",
                "position": r["position"],
                "team": r["current_team"],
                "gp": r["gp"], "g": r["g"], "a": r["a"], "p": r["p"],
                "plus_minus": r["plus_minus"], "ppg": r["ppg"],
            })
            player_ids.append(r["id"])

        return {"league": league, "stat": stat, "leaders": leaders, "count": len(leaders)}, {"player_ids": player_ids}
    finally:
        conn.close()


def _pt_query_drills(params: dict, org_id: str) -> tuple[dict, dict]:
    """Search the drill library. Returns matching drills."""
    conn = get_db()
    try:
        where = ["(org_id IS NULL OR org_id = ?)"]
        sql_params: list = [org_id]

        category = params.get("category")
        if category:
            where.append("category = ?")
            sql_params.append(category)
        age_level = params.get("age_level")
        if age_level:
            where.append("age_levels LIKE ?")
            sql_params.append(f'%"{age_level}"%')
        tags = params.get("tags")
        if tags:
            for tag in tags.split(","):
                tag = tag.strip()
                if tag:
                    where.append("tags LIKE ?")
                    sql_params.append(f'%"{tag}"%')
        intensity = params.get("intensity")
        if intensity:
            where.append("intensity = ?")
            sql_params.append(intensity)
        search = params.get("search")
        if search:
            where.append("(LOWER(name) LIKE ? OR LOWER(description) LIKE ?)")
            sql_params.extend([f"%{search.lower()}%", f"%{search.lower()}%"])

        limit = min(params.get("limit", 10), 20)
        sql_params.append(limit)
        rows = conn.execute(
            f"SELECT * FROM drills WHERE {' AND '.join(where)} ORDER BY category, name LIMIT ?", sql_params
        ).fetchall()

        drills = []
        for r in rows:
            d = _drill_row_to_dict(r)
            drills.append({
                "id": d["id"], "name": d["name"], "category": d["category"],
                "description": d["description"][:200],
                "coaching_points": (d.get("coaching_points") or "")[:150],
                "duration_minutes": d["duration_minutes"],
                "ice_surface": d["ice_surface"], "intensity": d["intensity"],
                "age_levels": d["age_levels"], "tags": d["tags"],
                "concept_id": d.get("concept_id"),
            })
        return {"drills": drills, "count": len(drills)}, {}
    finally:
        conn.close()


def _pt_generate_practice_plan(params: dict, org_id: str, user_id: str) -> tuple[dict, dict]:
    """Generate a practice plan via the AI generation endpoint logic."""
    import asyncio
    conn = get_db()
    try:
        body = PracticePlanGenerateRequest(
            team_name=params["team_name"],
            duration_minutes=params.get("duration_minutes", 90),
            focus_areas=params.get("focus_areas", []),
            age_level=params.get("age_level", "JUNIOR_COLLEGE_PRO"),
            notes=params.get("notes"),
        )

        # Reuse the same logic as the generate endpoint — build a synchronous version
        roster_rows = conn.execute(
            "SELECT first_name, last_name, position, shoots FROM players WHERE org_id = ? AND LOWER(current_team) = LOWER(?)",
            (org_id, body.team_name)
        ).fetchall()
        roster_summary = [f"{r['first_name']} {r['last_name']} ({r['position']}, {r['shoots'] or '?'})" for r in roster_rows]

        team_system = None
        ts_row = conn.execute(
            "SELECT * FROM team_systems WHERE org_id = ? AND LOWER(team_name) = LOWER(?)", (org_id, body.team_name)
        ).fetchone()
        if ts_row:
            team_system = dict(ts_row)

        drill_rows = conn.execute(
            "SELECT * FROM drills WHERE (org_id IS NULL OR org_id = ?) AND age_levels LIKE ? ORDER BY category, name",
            (org_id, f'%"{body.age_level}"%')
        ).fetchall()
        available_drills = []
        for dr in drill_rows:
            d = _drill_row_to_dict(dr)
            available_drills.append({
                "id": d["id"], "name": d["name"], "category": d["category"],
                "description": d["description"][:150], "duration_minutes": d["duration_minutes"],
                "ice_surface": d["ice_surface"], "intensity": d["intensity"],
                "skill_focus": d["skill_focus"], "concept_id": d.get("concept_id"), "tags": d["tags"],
            })

        focus_str = ", ".join(body.focus_areas) if body.focus_areas else "general skills"

        # Build a simple mock plan (Bench Talk will present it nicely)
        warmup = [d for d in available_drills if d["category"] == "warm_up"][:1]
        skills = [d for d in available_drills if d["category"] in ("passing", "shooting", "skating")][:2]
        systems = [d for d in available_drills if d["category"] in ("offensive", "defensive", "transition")][:2]
        games = [d for d in available_drills if d["category"] in ("small_area_games", "battle")][:1]
        cond = [d for d in available_drills if d["category"] == "conditioning"][:1]

        def _mk_phase(phase, label, dl, dur):
            return {"phase": phase, "phase_label": label, "duration_minutes": dur,
                    "drills": [{"drill_id": d["id"], "drill_name": d["name"],
                                "duration_minutes": d["duration_minutes"],
                                "coaching_notes": f"Focus on {focus_str}."} for d in dl]}

        plan_data = {
            "title": f"Practice Plan: {focus_str.title()} — {body.team_name}",
            "phases": [
                _mk_phase("warm_up", "Warm Up", warmup, 10),
                _mk_phase("skill_work", "Skill Work", skills, 25),
                _mk_phase("systems", "Team Systems", systems, 20),
                _mk_phase("scrimmage", "Game Situations", games, 15),
                _mk_phase("conditioning", "Conditioning", cond, 10),
                {"phase": "cool_down", "phase_label": "Cool Down", "duration_minutes": 5,
                 "drills": [{"drill_id": None, "drill_name": "Easy skate and stretch",
                             "duration_minutes": 5, "coaching_notes": "Light skate, static stretching."}]},
            ],
            "coaching_summary": f"Practice focused on {focus_str} for the {body.team_name}."
        }

        # Try AI generation
        client = get_anthropic_client()
        if client:
            try:
                system_prompt = """You are ProspectX Practice Plan Intelligence. Generate a structured practice plan in JSON format.
Return ONLY valid JSON with phases: warm_up, skill_work, systems, scrimmage, conditioning, cool_down.
Each phase has drills from the provided library with drill_id, drill_name, duration_minutes, coaching_notes.
AGE-LEVEL RULES: U8=maximum fun, games, no systems/tactics, constant puck touches. U10=intro skills, still game-heavy. U12=begin team concepts, basic forecheck/breakout. U14+=full systems, PP/PK, tactical depth. Do NOT assign systems drills to U8/U10."""

                user_prompt = f"""Generate a {body.duration_minutes}-minute practice for {body.team_name}.
Age: {body.age_level}. Focus: {focus_str}. {f'Notes: {body.notes}' if body.notes else ''}
Roster: {len(roster_summary)} players. {f'System: {team_system.get("forecheck","N/A")} forecheck' if team_system else ''}
Available drills: {json.dumps(available_drills[:40], indent=1)}"""

                msg = client.messages.create(
                    model="claude-sonnet-4-20250514", max_tokens=3000,
                    system=system_prompt,
                    messages=[{"role": "user", "content": user_prompt}],
                )
                resp = msg.content[0].text.strip()
                if resp.startswith("```"):
                    resp = resp.split("```")[1]
                    if resp.startswith("json"):
                        resp = resp[4:]
                plan_data = json.loads(resp)
            except Exception as e:
                logger.error("Bench Talk practice plan generation error: %s", e)

        # Save the plan
        plan_id = gen_id()
        now = datetime.utcnow().isoformat()
        title = plan_data.get("title", f"Practice Plan — {body.team_name}")
        conn.execute("""
            INSERT INTO practice_plans (id, org_id, user_id, team_name, title, age_level,
                duration_minutes, focus_areas, plan_data, notes, status, created_at, updated_at)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, 'draft', ?, ?)
        """, (plan_id, org_id, user_id, body.team_name, title, body.age_level,
              body.duration_minutes, json.dumps(body.focus_areas), json.dumps(plan_data), body.notes, now, now))

        for phase_data in plan_data.get("phases", []):
            for i, drill_entry in enumerate(phase_data.get("drills", [])):
                drill_id = drill_entry.get("drill_id")
                if drill_id:
                    conn.execute("""
                        INSERT INTO practice_plan_drills (id, practice_plan_id, drill_id, phase,
                            sequence_order, duration_minutes, coaching_notes)
                        VALUES (?, ?, ?, ?, ?, ?, ?)
                    """, (gen_id(), plan_id, drill_id, phase_data["phase"], i,
                          drill_entry.get("duration_minutes", 10), drill_entry.get("coaching_notes")))
        conn.commit()

        return {
            "plan_id": plan_id,
            "title": title,
            "duration_minutes": body.duration_minutes,
            "phases": len(plan_data.get("phases", [])),
            "coaching_summary": plan_data.get("coaching_summary", ""),
            "message": f"Practice plan created! View it at /practice-plans/{plan_id}"
        }, {"practice_plan_id": plan_id}
    finally:
        conn.close()


def _pt_get_player_recent_form(params: dict, org_id: str) -> tuple[dict, dict]:
    """Get a player's recent game-by-game performance, streaks, and progression."""
    player_name = params.get("player_name", "")
    last_n = min(max(params.get("last_n", 5), 1), 20)
    conn = get_db()
    try:
        # Find the player
        player = conn.execute("""
            SELECT id, first_name, last_name, position, current_team, current_league
            FROM players
            WHERE org_id = ? AND (first_name || ' ' || last_name) LIKE ?
            ORDER BY CASE WHEN (first_name || ' ' || last_name) = ? THEN 0 ELSE 1 END
            LIMIT 1
        """, (org_id, f"%{player_name}%", player_name)).fetchone()

        if not player:
            return {"error": f"No player found matching '{player_name}'"}, {}

        pid = player["id"]
        pname = f"{player['first_name']} {player['last_name']}"
        refs = {"player_id": pid, "player_name": pname}

        result = {
            "player": pname,
            "position": player["position"],
            "team": player["current_team"],
            "league": player["current_league"],
        }

        # ── Recent games ──
        game_rows = conn.execute("""
            SELECT game_date, opponent, home_away, goals, assists, points,
                   plus_minus, pim, shots, ppg, shg, gwg
            FROM player_game_stats
            WHERE player_id = ?
            ORDER BY game_date DESC
            LIMIT ?
        """, (pid, last_n)).fetchall()

        source = "hockeytech"
        if not game_rows:
            # Fallback to player_stats game rows
            game_rows = conn.execute("""
                SELECT created_at as game_date, g as goals, a as assists, p as points,
                       plus_minus, pim, shots
                FROM player_stats
                WHERE player_id = ? AND stat_type = 'game'
                ORDER BY created_at DESC
                LIMIT ?
            """, (pid, last_n)).fetchall()
            source = "instat" if game_rows else "none"

        if game_rows:
            games = [dict(r) for r in game_rows]
            n = len(games)
            total_g = sum(g.get("goals", 0) or 0 for g in games)
            total_a = sum(g.get("assists", 0) or 0 for g in games)
            total_p = sum(g.get("points", 0) or 0 for g in games)

            # Point streak
            pt_streak = 0
            for g in games:
                if (g.get("points", 0) or 0) > 0:
                    pt_streak += 1
                else:
                    break
            streak_str = f"{pt_streak}-game point streak" if pt_streak >= 2 else ("Point in last game" if pt_streak == 1 else "No active point streak")

            # Goal streak
            goal_streak = 0
            for g in games:
                if (g.get("goals", 0) or 0) > 0:
                    goal_streak += 1
                else:
                    break

            result["recent_form"] = {
                "games": games,
                "games_found": n,
                "totals": {"g": total_g, "a": total_a, "p": total_p},
                "averages": {"gpg": round(total_g / n, 2), "apg": round(total_a / n, 2), "ppg": round(total_p / n, 2)},
                "point_streak": streak_str,
                "goal_streak": f"{goal_streak}-game goal streak" if goal_streak >= 2 else None,
                "source": source,
            }
        else:
            result["recent_form"] = {"games_found": 0, "message": "No game-by-game data available. Sync game logs first."}

        # ── Season progression ──
        hist_rows = conn.execute("""
            SELECT psh.season, psh.gp, psh.g, psh.a, psh.p, psh.plus_minus, psh.pim,
                   psh.ppg, psh.shots, psh.shooting_pct, psh.league, psh.team_name
            FROM player_stats_history psh
            INNER JOIN (
                SELECT season, MAX(date_recorded) as max_date
                FROM player_stats_history
                WHERE player_id = ?
                GROUP BY season
            ) latest ON psh.season = latest.season AND psh.date_recorded = latest.max_date
            WHERE psh.player_id = ?
            ORDER BY psh.season ASC
        """, (pid, pid)).fetchall()

        if hist_rows:
            seasons = []
            for hr in hist_rows:
                hd = dict(hr)
                gp = hd.get("gp", 0) or 0
                hd["ppg_rate"] = round((hd.get("p", 0) or 0) / gp, 2) if gp > 0 else 0.0
                seasons.append(hd)
            result["progression"] = {
                "seasons": seasons,
                "trend": "improving" if len(seasons) >= 2 and seasons[-1]["ppg_rate"] > seasons[-2]["ppg_rate"] + 0.1
                         else ("declining" if len(seasons) >= 2 and seasons[-1]["ppg_rate"] < seasons[-2]["ppg_rate"] - 0.1
                         else "stable" if len(seasons) >= 2 else "single_season"),
            }
        else:
            # Fallback to player_stats season rows
            season_rows = conn.execute("""
                SELECT season, gp, g, a, p, plus_minus, pim, shots, shooting_pct
                FROM player_stats
                WHERE player_id = ? AND stat_type = 'season'
                ORDER BY season ASC
            """, (pid,)).fetchall()
            if season_rows:
                seasons = []
                for sr in season_rows:
                    sd = dict(sr)
                    gp = sd.get("gp", 0) or 0
                    sd["ppg_rate"] = round((sd.get("p", 0) or 0) / gp, 2) if gp > 0 else 0.0
                    seasons.append(sd)
                result["progression"] = {"seasons": seasons, "trend": "data_from_current_stats"}

        return result, refs
    finally:
        conn.close()


def _pt_get_team_context(params: dict, org_id: str) -> tuple[dict, dict]:
    """Get a team's lines, roster overview, and tactical systems."""
    team_name = params.get("team_name", "")
    conn = get_db()
    try:
        # Roster count by position
        pos_rows = conn.execute("""
            SELECT position, COUNT(*) as cnt FROM players
            WHERE org_id = ? AND current_team = ? AND (is_deleted = 0 OR is_deleted IS NULL)
            GROUP BY position ORDER BY cnt DESC
        """, (org_id, team_name)).fetchall()
        positions = {r["position"]: r["cnt"] for r in pos_rows if r["position"]}
        roster_count = sum(positions.values())

        # Line combinations
        line_rows = conn.execute("""
            SELECT line_type, line_label, line_order, player_names, player_refs,
                   toi_seconds, goals_for, goals_against, data_source
            FROM line_combinations
            WHERE org_id = ? AND team_name = ?
            ORDER BY line_order ASC
        """, (org_id, team_name)).fetchall()

        lines = []
        for lr in line_rows:
            line_data = {
                "type": lr["line_type"],
                "label": lr["line_label"] or lr["line_type"],
                "players": lr["player_names"],
                "toi_seconds": lr["toi_seconds"] or 0,
                "goals_for": lr["goals_for"] or 0,
                "goals_against": lr["goals_against"] or 0,
                "source": lr["data_source"] or "manual",
            }
            # Parse player_refs JSON for names
            if lr["player_refs"]:
                try:
                    refs = json.loads(lr["player_refs"])
                    line_data["player_details"] = [
                        {"name": r.get("name", ""), "position": r.get("position", "")}
                        for r in refs if isinstance(r, dict)
                    ]
                except (json.JSONDecodeError, TypeError):
                    pass
            lines.append(line_data)

        # Team systems
        sys_row = conn.execute("""
            SELECT forecheck, dz_structure, oz_setup, pp_formation, pk_formation,
                   neutral_zone, breakout, pace, physicality, offensive_style, notes
            FROM team_systems WHERE org_id = ? AND team_name = ?
            ORDER BY created_at DESC LIMIT 1
        """, (org_id, team_name)).fetchone()

        systems = {}
        if sys_row:
            systems = {k: sys_row[k] for k in sys_row.keys() if sys_row[k]}

        return {
            "team": team_name,
            "roster_count": roster_count,
            "positions": positions,
            "lines": lines,
            "systems": systems,
        }, {}
    finally:
        conn.close()


async def _pt_get_game_context(params: dict, org_id: str) -> tuple[dict, dict]:
    """Get live standings, upcoming games, and recent scores from HockeyTech."""
    league = params.get("league", "gojhl")
    team_name = params.get("team_name")

    try:
        client = HockeyTechClient(league)
    except ValueError:
        return {"error": f"Unknown league: {league}"}, {}

    try:
        season_id = await client.get_current_season_id()
        if not season_id:
            return {"error": "No current season found"}, {}

        # Get standings
        raw_standings = await client.get_standings(season_id)
        standings = []
        team_rank = None
        for i, s in enumerate(raw_standings[:15], 1):
            entry = {
                "rank": i,
                "team": s.get("team_name", s.get("name", "")),
                "gp": s.get("games_played", 0),
                "wins": s.get("wins", 0),
                "losses": s.get("losses", 0),
                "otl": s.get("ot_losses", 0),
                "points": s.get("points", 0),
            }
            standings.append(entry)
            if team_name and team_name.lower() in entry["team"].lower():
                team_rank = entry

        # Get scorebar (recent + upcoming)
        raw_games = await client.get_scorebar(days_back=3, days_ahead=7)
        upcoming = []
        recent = []
        now_str = datetime.now().strftime("%Y-%m-%d")
        for g in raw_games:
            game_date = g.get("date_with_day", g.get("game_date", ""))
            game_info = {
                "date": game_date,
                "home": g.get("home_team", ""),
                "away": g.get("visiting_team", g.get("away_team", "")),
                "home_score": g.get("home_goal_count", g.get("home_score")),
                "away_score": g.get("visiting_goal_count", g.get("away_score")),
                "status": g.get("game_status", g.get("status", "")),
            }
            # Categorize as upcoming or recent
            raw_date = g.get("game_date", "")
            if raw_date >= now_str and game_info["status"] in ("", "Not Started", None, "0"):
                upcoming.append(game_info)
            else:
                recent.append(game_info)

        result = {
            "league": league,
            "standings": standings[:10],
            "upcoming_games": upcoming[:5],
            "recent_results": recent[:5],
        }
        if team_rank:
            result["your_team"] = team_rank

        return result, {}
    except Exception as e:
        logger.warning("get_game_context error: %s", str(e))
        return {"error": f"Failed to fetch game context: {str(e)}"}, {}


def _pt_get_coaching_prep(params: dict, org_id: str) -> tuple[dict, dict]:
    """Get active game plans and series strategies."""
    team_name = params.get("team_name")
    opponent = params.get("opponent")
    conn = get_db()
    try:
        # Active/draft game plans
        gp_query = "SELECT * FROM game_plans WHERE org_id = ? AND status IN ('active', 'draft')"
        gp_params_list: list = [org_id]
        if team_name:
            gp_query += " AND team_name LIKE ?"
            gp_params_list.append(f"%{team_name}%")
        if opponent:
            gp_query += " AND opponent_team_name LIKE ?"
            gp_params_list.append(f"%{opponent}%")
        gp_query += " ORDER BY game_date ASC LIMIT 5"

        gp_rows = conn.execute(gp_query, gp_params_list).fetchall()
        game_plans = []
        for gp in gp_rows:
            plan = {
                "id": gp["id"],
                "team": gp["team_name"],
                "opponent": gp["opponent_team_name"],
                "date": gp["game_date"],
                "session_type": gp["session_type"] if "session_type" in gp.keys() else "pre_game",
                "status": gp["status"],
                "opponent_analysis": gp["opponent_analysis"] or "",
                "our_strategy": gp["our_strategy"] or "",
                "keys_to_game": gp["keys_to_game"] or "",
                "special_teams_plan": gp["special_teams_plan"] or "",
            }
            # Parse JSON fields
            for json_field in ["matchups", "talking_points"]:
                val = gp[json_field] if json_field in gp.keys() else None
                if val:
                    try:
                        plan[json_field] = json.loads(val)
                    except (json.JSONDecodeError, TypeError):
                        plan[json_field] = val
            game_plans.append(plan)

        # Active series
        sp_query = "SELECT * FROM series_plans WHERE org_id = ? AND status = 'active'"
        sp_params_list: list = [org_id]
        if team_name:
            sp_query += " AND team_name LIKE ?"
            sp_params_list.append(f"%{team_name}%")
        sp_query += " LIMIT 5"

        sp_rows = conn.execute(sp_query, sp_params_list).fetchall()
        active_series = []
        for sp in sp_rows:
            series = {
                "id": sp["id"],
                "team": sp["team_name"],
                "opponent": sp["opponent_team_name"],
                "series_name": sp["series_name"],
                "format": sp["series_format"],
                "score": sp["current_score"],
                "status": sp["status"],
            }
            for json_field in ["working_strategies", "needs_adjustment", "game_notes"]:
                val = sp[json_field] if json_field in sp.keys() else None
                if val:
                    try:
                        series[json_field] = json.loads(val)
                    except (json.JSONDecodeError, TypeError):
                        series[json_field] = val
            active_series.append(series)

        return {
            "game_plans": game_plans,
            "active_series": active_series,
            "total_plans": len(game_plans),
            "total_series": len(active_series),
        }, {}
    finally:
        conn.close()


def _pt_search_scout_notes(params: dict, org_id: str, user_id: str) -> tuple[dict, dict]:
    """Search scouting observations with filters."""
    player_name = params.get("player_name")
    tag_filter = params.get("tag")
    note_type = params.get("note_type")
    limit = min(params.get("limit", 10), 25)

    conn = get_db()
    try:
        query = """
            SELECT sn.note_text, sn.note_type, sn.tags, sn.is_private, sn.created_at,
                   p.id as player_id, p.first_name, p.last_name, p.position, p.current_team,
                   u.first_name as scout_first, u.last_name as scout_last
            FROM scout_notes sn
            JOIN players p ON sn.player_id = p.id
            JOIN users u ON sn.scout_id = u.id
            WHERE sn.org_id = ? AND (p.is_deleted = 0 OR p.is_deleted IS NULL)
              AND (sn.is_private = 0 OR sn.scout_id = ?)
        """
        query_params: list = [org_id, user_id]

        if player_name:
            query += " AND (p.first_name || ' ' || p.last_name LIKE ?)"
            query_params.append(f"%{player_name}%")
        if note_type:
            query += " AND sn.note_type = ?"
            query_params.append(note_type)

        query += " ORDER BY sn.created_at DESC LIMIT ?"
        query_params.append(limit * 3 if tag_filter else limit)  # over-fetch if filtering by tag

        rows = conn.execute(query, query_params).fetchall()

        notes = []
        player_ids = set()
        for r in rows:
            tags = []
            try:
                tags = json.loads(r["tags"]) if r["tags"] else []
            except (json.JSONDecodeError, TypeError):
                pass

            # Tag filter (post-query since tags are JSON)
            if tag_filter and not any(tag_filter.lower() in t.lower() for t in tags):
                continue

            player_ids.add(r["player_id"])
            notes.append({
                "player_name": f"{r['first_name']} {r['last_name']}",
                "position": r["position"],
                "team": r["current_team"],
                "scout": f"{r['scout_first']} {r['scout_last']}",
                "note": r["note_text"],
                "type": r["note_type"],
                "tags": tags,
                "date": r["created_at"][:10] if r["created_at"] else "",
            })
            if len(notes) >= limit:
                break

        return {
            "notes": notes,
            "total_found": len(notes),
        }, {"player_ids": list(player_ids)}
    finally:
        conn.close()


def _pt_get_scouting_list(params: dict, org_id: str, user_id: str) -> tuple[dict, dict]:
    """Get the user's active scouting watchlist."""
    priority = params.get("priority")
    limit = min(params.get("limit", 10), 20)

    conn = get_db()
    try:
        query = """
            SELECT sl.priority, sl.target_reason, sl.scout_notes, sl.tags, sl.created_at,
                   p.id as player_id, p.first_name, p.last_name, p.position,
                   p.current_team, p.current_league
            FROM scouting_list sl
            JOIN players p ON sl.player_id = p.id
            WHERE sl.user_id = ? AND sl.org_id = ? AND sl.is_active = 1
              AND (p.is_deleted = 0 OR p.is_deleted IS NULL)
        """
        query_params: list = [user_id, org_id]

        if priority:
            query += " AND sl.priority = ?"
            query_params.append(priority)

        query += " ORDER BY CASE sl.priority WHEN 'high' THEN 1 WHEN 'medium' THEN 2 WHEN 'low' THEN 3 END, sl.created_at DESC LIMIT ?"
        query_params.append(limit)

        rows = conn.execute(query, query_params).fetchall()

        items = []
        player_ids = set()
        for r in rows:
            player_ids.add(r["player_id"])
            tags = []
            try:
                tags = json.loads(r["tags"]) if r["tags"] else []
            except (json.JSONDecodeError, TypeError):
                pass
            items.append({
                "player_name": f"{r['first_name']} {r['last_name']}",
                "position": r["position"],
                "team": r["current_team"],
                "league": r["current_league"],
                "priority": r["priority"],
                "target_reason": r["target_reason"],
                "notes": r["scout_notes"],
                "tags": tags,
                "added_date": r["created_at"][:10] if r["created_at"] else "",
            })

        return {
            "items": items,
            "total": len(items),
        }, {"player_ids": list(player_ids)}
    finally:
        conn.close()


def _pt_diagnose_player_struggles(params: dict, org_id: str) -> tuple[dict, dict]:
    """Diagnose a struggling player by comparing recent form to season averages."""
    player_name = params.get("player_name", "")
    look_back = min(max(params.get("look_back_games", 10), 5), 20)

    conn = get_db()
    try:
        # Find the player by fuzzy name match
        name_parts = player_name.strip().split()
        if len(name_parts) >= 2:
            rows = conn.execute(
                """SELECT id, first_name, last_name, position, current_team, current_league
                   FROM players WHERE org_id = ? AND (is_deleted = 0 OR is_deleted IS NULL)
                   AND LOWER(first_name) LIKE ? AND LOWER(last_name) LIKE ?
                   LIMIT 5""",
                (org_id, f"%{name_parts[0].lower()}%", f"%{name_parts[-1].lower()}%"),
            ).fetchall()
        else:
            rows = conn.execute(
                """SELECT id, first_name, last_name, position, current_team, current_league
                   FROM players WHERE org_id = ? AND (is_deleted = 0 OR is_deleted IS NULL)
                   AND (LOWER(first_name) LIKE ? OR LOWER(last_name) LIKE ?)
                   LIMIT 5""",
                (org_id, f"%{player_name.lower()}%", f"%{player_name.lower()}%"),
            ).fetchall()

        if not rows:
            return {"error": f"No player found matching '{player_name}'"}, {}

        player = dict(rows[0])
        pid = player["id"]
        pname = f"{player['first_name']} {player['last_name']}"

        # Get season stats
        season_stats = conn.execute(
            "SELECT * FROM player_stats WHERE player_id = ? ORDER BY season DESC LIMIT 3", (pid,)
        ).fetchall()

        # Get recent game log
        game_log = conn.execute(
            "SELECT * FROM game_stats WHERE player_id = ? ORDER BY game_date DESC LIMIT ?",
            (pid, look_back),
        ).fetchall()

        # Get player intelligence
        intel = conn.execute(
            "SELECT archetype, overall_grade, offensive_grade, defensive_grade, skating_grade, hockey_iq_grade, compete_grade, strengths, development_areas FROM player_intelligence WHERE player_id = ? ORDER BY created_at DESC LIMIT 1",
            (pid,),
        ).fetchone()

        # Get recent scout notes
        notes = conn.execute(
            "SELECT note_text, note_type, created_at FROM scout_notes WHERE player_id = ? AND org_id = ? ORDER BY created_at DESC LIMIT 5",
            (pid, org_id),
        ).fetchall()

        # Build diagnostic
        diagnostic = {
            "player": pname,
            "position": player["position"],
            "team": player["current_team"],
            "league": player["current_league"],
        }

        # Season averages
        if season_stats:
            s = dict(season_stats[0])
            gp = s.get("gp", 0) or 0
            if gp > 0:
                diagnostic["season_averages"] = {
                    "gp": gp, "g": s.get("g", 0), "a": s.get("a", 0), "p": s.get("p", 0),
                    "gpg": round((s.get("g", 0) or 0) / gp, 3),
                    "apg": round((s.get("a", 0) or 0) / gp, 3),
                    "ppg": round((s.get("p", 0) or 0) / gp, 3),
                    "plus_minus": s.get("plus_minus"),
                    "pim": s.get("pim"),
                    "shots": s.get("shots"),
                    "shooting_pct": s.get("shooting_pct"),
                }

        # Recent form
        if game_log:
            recent_g = sum(g.get("goals", 0) or 0 for g in [dict(gl) for gl in game_log])
            recent_a = sum(g.get("assists", 0) or 0 for g in [dict(gl) for gl in game_log])
            recent_p = sum(g.get("points", 0) or 0 for g in [dict(gl) for gl in game_log])
            recent_pm = sum(g.get("plus_minus", 0) or 0 for g in [dict(gl) for gl in game_log])
            n_games = len(game_log)

            diagnostic["recent_form"] = {
                "games_analyzed": n_games,
                "goals": recent_g, "assists": recent_a, "points": recent_p,
                "gpg": round(recent_g / n_games, 3),
                "apg": round(recent_a / n_games, 3),
                "ppg": round(recent_p / n_games, 3),
                "plus_minus": recent_pm,
            }

            # Streak analysis
            goal_drought = 0
            point_drought = 0
            for gl in [dict(g) for g in game_log]:
                if (gl.get("goals", 0) or 0) == 0:
                    goal_drought += 1
                else:
                    break
            for gl in [dict(g) for g in game_log]:
                if (gl.get("points", 0) or 0) == 0:
                    point_drought += 1
                else:
                    break

            diagnostic["streaks"] = {
                "current_goal_drought": goal_drought,
                "current_point_drought": point_drought,
            }

            # Declining metrics
            declining = []
            stable = []
            if "season_averages" in diagnostic:
                sa = diagnostic["season_averages"]
                rf = diagnostic["recent_form"]
                for metric, label in [("gpg", "Goals per game"), ("apg", "Assists per game"), ("ppg", "Points per game")]:
                    if sa.get(metric, 0) > 0:
                        delta = rf[metric] - sa[metric]
                        pct = (delta / sa[metric]) * 100 if sa[metric] > 0 else 0
                        entry = {"metric": label, "season": sa[metric], "recent": rf[metric], "delta_pct": round(pct, 1)}
                        if pct < -15:
                            declining.append(entry)
                        else:
                            stable.append(entry)

                diagnostic["declining_metrics"] = declining
                diagnostic["stable_metrics"] = stable

            # Potential causes
            causes = []
            if goal_drought >= 5:
                causes.append("Extended goal drought (5+ games) — may indicate shooting slump or reduced shot volume")
            if point_drought >= 4:
                causes.append("Extended point drought (4+ games) — could signal deployment changes or linemate issues")
            if declining:
                for d in declining:
                    if d["delta_pct"] < -30:
                        causes.append(f"Significant decline in {d['metric']} ({d['delta_pct']}% vs season avg)")
            if not causes:
                causes.append("No significant statistical decline detected — struggles may be perception-based or related to unmeasured factors (effort, positioning, confidence)")
            diagnostic["potential_causes"] = causes

        else:
            diagnostic["recent_form"] = {"error": "No game-by-game data available for trend analysis"}

        # Scout note flags
        if notes:
            note_flags = []
            for n in notes:
                nd = dict(n)
                if nd.get("note_type") in ("concern", "development"):
                    note_flags.append({
                        "type": nd["note_type"],
                        "note": nd["note_text"][:200],
                        "date": nd["created_at"][:10] if nd["created_at"] else "",
                    })
            diagnostic["scout_note_flags"] = note_flags

        # Intelligence context
        if intel:
            diagnostic["intelligence"] = {
                "archetype": intel["archetype"],
                "overall_grade": intel["overall_grade"],
                "strengths": intel["strengths"],
                "development_areas": intel["development_areas"],
            }

        return diagnostic, {"player_ids": [pid]}
    finally:
        conn.close()


async def _execute_bench_talk_tool(tool_name: str, tool_input: dict, org_id: str, user_id: str) -> tuple[dict, dict]:
    """Route Bench Talk tool calls to the appropriate function.
    Returns (tool_result, entity_refs) tuple."""
    if tool_name == "query_players":
        return _pt_query_players(tool_input, org_id)
    elif tool_name == "get_player_intelligence":
        return _pt_get_player_intelligence(tool_input, org_id)
    elif tool_name == "compare_players":
        return _pt_compare_players(tool_input, org_id)
    elif tool_name == "start_report_generation":
        return _pt_start_report(tool_input, org_id, user_id)
    elif tool_name == "league_leaders":
        return _pt_league_leaders(tool_input, org_id)
    elif tool_name == "query_drills":
        return _pt_query_drills(tool_input, org_id)
    elif tool_name == "generate_practice_plan":
        return _pt_generate_practice_plan(tool_input, org_id, user_id)
    elif tool_name == "get_player_recent_form":
        return _pt_get_player_recent_form(tool_input, org_id)
    # Phase 2 tools
    elif tool_name == "get_team_context":
        return _pt_get_team_context(tool_input, org_id)
    elif tool_name == "get_game_context":
        return await _pt_get_game_context(tool_input, org_id)
    elif tool_name == "get_coaching_prep":
        return _pt_get_coaching_prep(tool_input, org_id)
    elif tool_name == "search_scout_notes":
        return _pt_search_scout_notes(tool_input, org_id, user_id)
    elif tool_name == "get_scouting_list":
        return _pt_get_scouting_list(tool_input, org_id, user_id)
    elif tool_name == "diagnose_player_struggles":
        return _pt_diagnose_player_struggles(tool_input, org_id)
    else:
        return {"error": f"Unknown tool: {tool_name}"}, {}


# ── Bench Talk Pydantic Models ───────────────────────────────────

class BenchTalkMessageRequest(BaseModel):
    message: str = Field(..., min_length=1, max_length=4000)
    mode: Optional[str] = None  # PXI mode override for this message

class BenchTalkFeedbackRequest(BaseModel):
    message_id: str
    rating: str = Field(..., pattern="^(positive|negative)$")
    feedback_text: Optional[str] = None


# ── Bench Talk Endpoints ──────────────────────────────────────────

@app.post("/bench-talk/conversations")
async def create_bench_talk_conversation(token_data: dict = Depends(verify_token)):
    """Create a new Bench Talk conversation."""
    user_id = token_data["user_id"]
    org_id = token_data["org_id"]
    conn_check = get_db()
    _check_email_verified(user_id, conn_check)
    conn_check.close()
    conv_id = gen_id()
    conn = get_db()
    try:
        conn.execute("""
            INSERT INTO bench_talk_conversations (id, user_id, org_id, title, created_at, updated_at)
            VALUES (?, ?, ?, 'New Conversation', ?, ?)
        """, (conv_id, user_id, org_id, now_iso(), now_iso()))
        conn.commit()
        return {"conversation_id": conv_id, "title": "New Conversation"}
    finally:
        conn.close()


@app.get("/bench-talk/conversations")
async def list_bench_talk_conversations(token_data: dict = Depends(verify_token)):
    """List all Bench Talk conversations for the current user."""
    user_id = token_data["user_id"]
    conn = get_db()
    try:
        rows = conn.execute("""
            SELECT c.id, c.title, c.created_at, c.updated_at,
                   (SELECT content FROM bench_talk_messages
                    WHERE conversation_id = c.id
                    ORDER BY created_at DESC LIMIT 1) as last_message,
                   (SELECT COUNT(*) FROM bench_talk_messages
                    WHERE conversation_id = c.id) as message_count
            FROM bench_talk_conversations c
            WHERE c.user_id = ?
            ORDER BY c.updated_at DESC
        """, (user_id,)).fetchall()
        return [dict(r) for r in rows]
    finally:
        conn.close()


@app.get("/bench-talk/conversations/{conversation_id}")
async def get_bench_talk_conversation(conversation_id: str, token_data: dict = Depends(verify_token)):
    """Get all messages in a Bench Talk conversation."""
    user_id = token_data["user_id"]
    conn = get_db()
    try:
        conv = conn.execute(
            "SELECT * FROM bench_talk_conversations WHERE id = ? AND user_id = ?",
            (conversation_id, user_id)
        ).fetchone()
        if not conv:
            raise HTTPException(status_code=404, detail="Conversation not found")

        messages = conn.execute("""
            SELECT id, role, content, metadata, tokens_used, created_at
            FROM bench_talk_messages
            WHERE conversation_id = ?
            ORDER BY created_at ASC
        """, (conversation_id,)).fetchall()

        return {
            "conversation": dict(conv),
            "messages": [dict(m) for m in messages]
        }
    finally:
        conn.close()


@app.delete("/bench-talk/conversations/{conversation_id}")
async def delete_bench_talk_conversation(conversation_id: str, token_data: dict = Depends(verify_token)):
    """Delete a Bench Talk conversation and its messages."""
    user_id = token_data["user_id"]
    conn = get_db()
    try:
        conv = conn.execute(
            "SELECT id FROM bench_talk_conversations WHERE id = ? AND user_id = ?",
            (conversation_id, user_id)
        ).fetchone()
        if not conv:
            raise HTTPException(status_code=404, detail="Conversation not found")

        conn.execute("DELETE FROM bench_talk_feedback WHERE message_id IN (SELECT id FROM bench_talk_messages WHERE conversation_id = ?)", (conversation_id,))
        conn.execute("DELETE FROM bench_talk_messages WHERE conversation_id = ?", (conversation_id,))
        conn.execute("DELETE FROM bench_talk_conversations WHERE id = ?", (conversation_id,))
        conn.commit()
        return {"deleted": True}
    finally:
        conn.close()


@app.post("/bench-talk/conversations/{conversation_id}/messages")
async def send_bench_talk_message(
    conversation_id: str,
    req: BenchTalkMessageRequest,
    token_data: dict = Depends(verify_token)
):
    """Send a message to Bench Talk and get an AI response."""
    user_id = token_data["user_id"]
    org_id = token_data["org_id"]

    # ── Email verification check ──
    ev_conn = get_db()
    _check_email_verified(user_id, ev_conn)
    ev_conn.close()

    # ── Usage limit check ──
    usage_conn = get_db()
    try:
        _check_tier_limit(user_id, "bench_talks", usage_conn)
    finally:
        usage_conn.close()

    conn = get_db()
    try:
        # Verify conversation ownership
        conv = conn.execute(
            "SELECT id FROM bench_talk_conversations WHERE id = ? AND user_id = ?",
            (conversation_id, user_id)
        ).fetchone()
        if not conv:
            raise HTTPException(status_code=404, detail="Conversation not found")

        # Look up user's hockey role and name for personalization
        user_row = conn.execute(
            "SELECT first_name, hockey_role FROM users WHERE id = ?", (user_id,)
        ).fetchone()
        user_first_name = user_row["first_name"] if user_row and user_row["first_name"] else "there"
        hockey_role = user_row["hockey_role"] if user_row and user_row["hockey_role"] else "scout"

        # ── PXI Mode Resolution ──
        # Priority: req.mode → conversation stored mode → user hockey_role fallback
        conv_row = conn.execute(
            "SELECT mode FROM bench_talk_conversations WHERE id = ?", (conversation_id,)
        ).fetchone()
        conv_mode = conv_row["mode"] if conv_row and conv_row["mode"] else None

        # If user sent a mode override, persist it to the conversation
        if req.mode and req.mode in VALID_MODES:
            conn.execute(
                "UPDATE bench_talk_conversations SET mode = ?, updated_at = ? WHERE id = ?",
                (req.mode, now_iso(), conversation_id)
            )
            conn.commit()
            resolved_mode = req.mode
        else:
            resolved_mode = resolve_mode(hockey_role, conv_mode, None)

        # Labels for personalization (expanded for all 10 modes)
        hockey_role_labels = {
            "scout": "Scout",
            "gm": "General Manager",
            "coach": "Head Coach",
            "player": "Player",
            "parent": "Hockey Parent",
            "analyst": "Analyst",
            "agent": "Agent",
            "skill_coach": "Skill Coach",
            "mental_coach": "Mental Coach",
            "broadcast": "Broadcaster",
            "producer": "Producer",
        }
        hockey_role_label = hockey_role_labels.get(resolved_mode, "Scout")
        # Use PXI mode block instead of old role instructions
        role_instructions = PXI_MODE_BLOCKS.get(resolved_mode, PXI_MODE_BLOCKS.get("scout", ""))

        # Save user message
        user_msg_id = gen_id()
        conn.execute("""
            INSERT INTO bench_talk_messages (id, conversation_id, role, content, created_at)
            VALUES (?, ?, 'user', ?, ?)
        """, (user_msg_id, conversation_id, req.message, now_iso()))
        conn.commit()

        # Check if Anthropic API key is available
        client = get_anthropic_client()
        if not client:
            # Mock mode — no API key
            mock_response = (
                "Hey! I'm Bench Talk, your ProspectX hockey conversation engine. I'm currently running in demo mode "
                "because no Anthropic API key is configured.\n\n"
                "Once connected, I can help you:\n"
                "- **Search players** by position, league, team, or stats\n"
                "- **Get ProspectX Intelligence** grades and archetypes\n"
                "- **Compare players** side-by-side\n"
                "- **Generate reports** (19 professional templates)\n"
                "- **Find league leaders** in any stat category\n\n"
                "To enable full Bench Talk intelligence, add your Anthropic API key to `backend/.env`."
            )
            asst_msg_id = gen_id()
            conn.execute("""
                INSERT INTO bench_talk_messages (id, conversation_id, role, content, tokens_used, created_at)
                VALUES (?, ?, 'assistant', ?, 0, ?)
            """, (asst_msg_id, conversation_id, mock_response, now_iso()))
            # Update conversation title on first message
            msg_count = conn.execute(
                "SELECT COUNT(*) as cnt FROM bench_talk_messages WHERE conversation_id = ?",
                (conversation_id,)
            ).fetchone()["cnt"]
            if msg_count <= 2:
                title = req.message[:50] + ("..." if len(req.message) > 50 else "")
                conn.execute(
                    "UPDATE bench_talk_conversations SET title = ?, updated_at = ? WHERE id = ?",
                    (title, now_iso(), conversation_id)
                )
            conn.commit()

            # Increment usage
            _increment_usage(user_id, "bench_talk", conversation_id, org_id, conn)

            return {
                "message": {
                    "id": asst_msg_id,
                    "role": "assistant",
                    "content": mock_response,
                    "tokens_used": 0,
                    "created_at": now_iso(),
                }
            }

        # ── Real Claude API call ──
        # Get conversation history
        history = conn.execute("""
            SELECT role, content FROM bench_talk_messages
            WHERE conversation_id = ?
            ORDER BY created_at ASC
        """, (conversation_id,)).fetchall()

        messages = [{"role": h["role"], "content": h["content"]} for h in history]

        # Conversation management: prevent token overflow on long conversations
        MAX_MESSAGES = 20  # ~10 turns of back-and-forth
        if len(messages) > MAX_MESSAGES:
            # Keep first message for topic context + last (MAX_MESSAGES - 1) messages
            messages = [messages[0]] + messages[-(MAX_MESSAGES - 1):]

        system_prompt = BENCH_TALK_SYSTEM_PROMPT.format(
            current_date=datetime.now().strftime("%B %d, %Y"),
            user_first_name=user_first_name,
            hockey_role_label=hockey_role_label,
            role_instructions=role_instructions,
        )

        # First API call (may include tool use)
        response = client.messages.create(
            model="claude-sonnet-4-20250514",
            max_tokens=4096,
            system=system_prompt,
            messages=messages,
            tools=BENCH_TALK_TOOLS,
        )

        assistant_text = ""
        tool_results_for_api = []
        tokens_used = response.usage.input_tokens + response.usage.output_tokens
        referenced_players: set = set()
        referenced_reports: set = set()

        for block in response.content:
            if block.type == "text":
                assistant_text += block.text
            elif block.type == "tool_use":
                tool_result, entity_refs = await _execute_bench_talk_tool(block.name, block.input, org_id, user_id)
                referenced_players.update(entity_refs.get("player_ids", []))
                referenced_reports.update(entity_refs.get("report_ids", []))
                tool_results_for_api.append({
                    "type": "tool_result",
                    "tool_use_id": block.id,
                    "content": json.dumps(tool_result, default=str),
                })

        # If tools were called, send results back for final response
        if tool_results_for_api:
            messages.append({"role": "assistant", "content": response.content})
            messages.append({"role": "user", "content": tool_results_for_api})

            final_response = client.messages.create(
                model="claude-sonnet-4-20250514",
                max_tokens=4096,
                system=system_prompt,
                messages=messages,
            )

            assistant_text = "".join(
                b.text for b in final_response.content if b.type == "text"
            )
            tokens_used += final_response.usage.input_tokens + final_response.usage.output_tokens

        # Save assistant message
        asst_msg_id = gen_id()
        metadata = json.dumps({
            "tool_calls": len(tool_results_for_api),
            "tokens": tokens_used,
            "player_ids": list(referenced_players),
            "report_ids": list(referenced_reports),
        })
        conn.execute("""
            INSERT INTO bench_talk_messages (id, conversation_id, role, content, metadata, tokens_used, created_at)
            VALUES (?, ?, 'assistant', ?, ?, ?, ?)
        """, (asst_msg_id, conversation_id, assistant_text, metadata, tokens_used, now_iso()))

        # Update conversation title on first exchange
        msg_count = conn.execute(
            "SELECT COUNT(*) as cnt FROM bench_talk_messages WHERE conversation_id = ?",
            (conversation_id,)
        ).fetchone()["cnt"]
        if msg_count <= 2:
            title = req.message[:50] + ("..." if len(req.message) > 50 else "")
            conn.execute(
                "UPDATE bench_talk_conversations SET title = ?, updated_at = ? WHERE id = ?",
                (title, now_iso(), conversation_id)
            )
        else:
            conn.execute(
                "UPDATE bench_talk_conversations SET updated_at = ? WHERE id = ?",
                (now_iso(), conversation_id)
            )
        conn.commit()

        # Increment usage after successful response
        _increment_usage(user_id, "bench_talk", conversation_id, org_id, conn)

        return {
            "message": {
                "id": asst_msg_id,
                "role": "assistant",
                "content": assistant_text,
                "metadata": metadata,
                "tokens_used": tokens_used,
                "created_at": now_iso(),
            }
        }
    except HTTPException:
        raise
    except Exception as e:
        logger.exception("Bench Talk message error for conversation %s", conversation_id)
        # Still save an error response so the conversation doesn't break
        try:
            error_msg_id = gen_id()
            error_text = "I'm sorry, I encountered an error processing your request. Please try again."
            conn.execute("""
                INSERT INTO bench_talk_messages (id, conversation_id, role, content, tokens_used, created_at)
                VALUES (?, ?, 'assistant', ?, 0, ?)
            """, (error_msg_id, conversation_id, error_text, now_iso()))
            conn.commit()
        except Exception:
            pass
        return {
            "message": {
                "id": error_msg_id if 'error_msg_id' in dir() else gen_id(),
                "role": "assistant",
                "content": "I'm sorry, I encountered an error processing your request. Please try again.",
                "tokens_used": 0,
                "created_at": now_iso(),
            }
        }
    finally:
        conn.close()


@app.post("/bench-talk/feedback")
async def submit_bench_talk_feedback(req: BenchTalkFeedbackRequest, token_data: dict = Depends(verify_token)):
    """Submit thumbs up/down feedback on a Bench Talk response."""
    user_id = token_data["user_id"]
    org_id = token_data["org_id"]
    conn = get_db()
    _check_email_verified(user_id, conn)
    try:
        feedback_id = gen_id()
        conn.execute("""
            INSERT INTO bench_talk_feedback (id, message_id, user_id, org_id, rating, feedback_text, created_at)
            VALUES (?, ?, ?, ?, ?, ?, ?)
        """, (feedback_id, req.message_id, user_id, org_id, req.rating, req.feedback_text, now_iso()))
        conn.commit()
        return {"success": True, "feedback_id": feedback_id}
    finally:
        conn.close()


# ============================================================
# PXI MODE ENDPOINTS
# ============================================================

@app.get("/pxi/modes")
async def get_pxi_modes():
    """Return the list of available PXI modes with metadata."""
    return PXI_MODES


@app.get("/pxi/wiring")
async def get_pxi_wiring():
    """Return the mode→template wiring table for frontend auto-selection."""
    return MODE_TEMPLATE_WIRING


class ModeUpdateRequest(BaseModel):
    mode: str = Field(..., pattern="^(scout|coach|analyst|gm|agent|parent|skill_coach|mental_coach|broadcast|producer)$")


@app.put("/bench-talk/conversations/{conversation_id}/mode")
async def update_conversation_mode(
    conversation_id: str,
    req: ModeUpdateRequest,
    token_data: dict = Depends(verify_token)
):
    """Update the PXI mode for a Bench Talk conversation."""
    user_id = token_data["user_id"]
    conn = get_db()
    try:
        conv = conn.execute(
            "SELECT id FROM bench_talk_conversations WHERE id = ? AND user_id = ?",
            (conversation_id, user_id)
        ).fetchone()
        if not conv:
            raise HTTPException(status_code=404, detail="Conversation not found")
        conn.execute(
            "UPDATE bench_talk_conversations SET mode = ?, updated_at = ? WHERE id = ?",
            (req.mode, now_iso(), conversation_id)
        )
        conn.commit()
        return {"conversation_id": conversation_id, "mode": req.mode}
    finally:
        conn.close()


# ============================================================
# REPORT SHARE ENDPOINTS
# ============================================================

@app.post("/reports/{report_id}/share")
async def create_report_share_link(report_id: str, token_data: dict = Depends(verify_token)):
    """Generate a shareable link for a report. Requires auth + same org to view."""
    user_id = token_data["user_id"]
    org_id = token_data["org_id"]
    conn = get_db()
    try:
        report = conn.execute(
            "SELECT id, org_id, share_token FROM reports WHERE id = ? AND org_id = ?",
            (report_id, org_id)
        ).fetchone()
        if not report:
            raise HTTPException(status_code=404, detail="Report not found")

        # Reuse existing token or generate new one
        share_token = report["share_token"] if report["share_token"] else str(uuid.uuid4())[:12]
        if not report["share_token"]:
            conn.execute(
                "UPDATE reports SET share_token = ? WHERE id = ?",
                (share_token, report_id)
            )
            conn.commit()

        frontend_url = os.getenv("FRONTEND_URL", "http://localhost:3000")
        share_url = f"{frontend_url}/reports/shared/{share_token}"
        return {"share_token": share_token, "share_url": share_url}
    finally:
        conn.close()


@app.get("/reports/shared/{share_token}")
async def get_shared_report(share_token: str, token_data: dict = Depends(verify_token)):
    """View a shared report. Requires auth + same org as report owner."""
    org_id = token_data["org_id"]
    conn = get_db()
    try:
        report = conn.execute(
            "SELECT * FROM reports WHERE share_token = ?", (share_token,)
        ).fetchone()
        if not report:
            raise HTTPException(status_code=404, detail="Shared report not found")
        if report["org_id"] != org_id:
            raise HTTPException(status_code=403, detail="Access denied — different organization")
        return dict(report)
    finally:
        conn.close()


@app.put("/reports/{report_id}/team-share")
async def toggle_team_share(report_id: str, token_data: dict = Depends(verify_token)):
    """Toggle org-wide sharing for a report."""
    org_id = token_data["org_id"]
    conn = get_db()
    try:
        report = conn.execute(
            "SELECT id, shared_with_org FROM reports WHERE id = ? AND org_id = ?",
            (report_id, org_id)
        ).fetchone()
        if not report:
            raise HTTPException(status_code=404, detail="Report not found")

        new_val = 0 if report["shared_with_org"] else 1
        conn.execute(
            "UPDATE reports SET shared_with_org = ? WHERE id = ?",
            (new_val, report_id)
        )
        conn.commit()
        return {"report_id": report_id, "shared_with_org": bool(new_val)}
    finally:
        conn.close()


@app.get("/reports/shared-with-me")
async def get_reports_shared_with_me(token_data: dict = Depends(verify_token)):
    """List reports shared with the user's org (excluding their own)."""
    user_id = token_data["user_id"]
    org_id = token_data["org_id"]
    conn = get_db()
    try:
        rows = conn.execute(
            """SELECT id, report_type, title, status, generated_at, created_at, player_id, team_name
               FROM reports
               WHERE org_id = ? AND shared_with_org = 1 AND (created_by != ? OR created_by IS NULL)
               ORDER BY created_at DESC LIMIT 50""",
            (org_id, user_id)
        ).fetchall()
        return [dict(r) for r in rows]
    finally:
        conn.close()


@app.get("/reports/{report_id}/quality")
async def get_report_quality(report_id: str, token_data: dict = Depends(verify_token)):
    """Get quality score and breakdown for a report."""
    org_id = token_data["org_id"]
    conn = get_db()
    try:
        row = conn.execute(
            "SELECT quality_score, quality_details FROM reports WHERE id = ? AND org_id = ?",
            (report_id, org_id)
        ).fetchone()
        if not row:
            raise HTTPException(status_code=404, detail="Report not found")
        quality_details = json.loads(row["quality_details"]) if row["quality_details"] else None
        return {
            "report_id": report_id,
            "quality_score": row["quality_score"],
            "details": quality_details,
        }
    finally:
        conn.close()


# ── DOCX Export ──────────────────────────────────────────────────

def _generate_report_docx(report_row: dict, player_row: dict = None) -> io.BytesIO:
    """Generate a Word document from a report row.
    Returns a BytesIO object containing the .docx file."""
    doc = DocxDocument()

    # ── Page margins ──
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # ── Color constants ──
    navy = RGBColor(0x0F, 0x2A, 0x3D)
    teal = RGBColor(0x18, 0xB3, 0xA6)
    orange = RGBColor(0xF3, 0x6F, 0x21)
    dark_gray = RGBColor(0x33, 0x33, 0x33)

    # ── Brand Header ──
    brand_para = doc.add_paragraph()
    brand_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_p = brand_para.add_run("Prospect")
    run_p.font.size = Pt(20)
    run_p.font.color.rgb = teal
    run_p.bold = True
    run_x = brand_para.add_run("X")
    run_x.font.size = Pt(20)
    run_x.font.color.rgb = orange
    run_x.bold = True
    run_sub = brand_para.add_run("  Intelligence Report")
    run_sub.font.size = Pt(12)
    run_sub.font.color.rgb = navy

    # ── Report Title ──
    report_type = report_row.get("report_type", "report")
    from types import SimpleNamespace  # noqa: E402
    label = report_type.replace("_", " ").title()
    # Try to get a better label
    try:
        from pxi_prompt_core import MODE_TEMPLATE_WIRING
    except ImportError:
        pass

    subject_name = ""
    if player_row:
        subject_name = f"{player_row.get('first_name', '')} {player_row.get('last_name', '')}".strip()
    elif report_row.get("team_name"):
        subject_name = report_row["team_name"]

    title_text = f"{subject_name} — {label}" if subject_name else label
    title_para = doc.add_heading(title_text, level=1)
    for run in title_para.runs:
        run.font.color.rgb = navy

    # ── Meta info ──
    gen_date = report_row.get("generated_at", report_row.get("created_at", ""))
    if gen_date:
        try:
            dt = datetime.fromisoformat(gen_date.replace("Z", "+00:00"))
            gen_date = dt.strftime("%B %d, %Y")
        except Exception:
            pass
    meta_parts = []
    if gen_date:
        meta_parts.append(f"Generated: {gen_date}")
    if report_row.get("overall_grade"):
        meta_parts.append(f"Grade: {report_row['overall_grade']}")
    quality_score = report_row.get("quality_score")
    if quality_score is not None:
        meta_parts.append(f"Quality Score: {quality_score:.0f}/100")
    if meta_parts:
        meta_para = doc.add_paragraph(" | ".join(meta_parts))
        meta_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in meta_para.runs:
            run.font.size = Pt(9)
            run.font.color.rgb = dark_gray
            run.italic = True

    doc.add_paragraph("")  # spacer

    # ── Parse report output into sections ──
    output_text = report_row.get("output_text", "")
    if not output_text:
        doc.add_paragraph("No report content available.")
    else:
        # Split on ALL_CAPS section headers (same pattern as frontend)
        section_pattern = re.compile(r'^([A-Z][A-Z0-9_]{2,}(?:\s+[A-Z0-9_]+)*)(?:\s*[:—\-]?\s*)(.*)', re.MULTILINE)
        lines = output_text.split("\n")
        current_heading = None
        current_lines: list = []

        def _flush_section():
            if current_heading:
                # Add section heading
                heading_label = current_heading.replace("_", " ").title()
                h = doc.add_heading(heading_label, level=2)
                for run in h.runs:
                    run.font.color.rgb = teal
            body = "\n".join(current_lines).strip()
            if not body:
                return
            # Process body content
            for para_text in body.split("\n"):
                para_text = para_text.strip()
                if not para_text:
                    continue
                # Detect bullet points
                is_bullet = para_text.startswith(("- ", "• ", "* ", "– "))
                if is_bullet:
                    bullet_text = para_text.lstrip("-•*– ").strip()
                    p = doc.add_paragraph(bullet_text, style="List Bullet")
                    for run in p.runs:
                        run.font.size = Pt(10)
                        run.font.color.rgb = dark_gray
                else:
                    p = doc.add_paragraph(para_text)
                    for run in p.runs:
                        run.font.size = Pt(10)
                        run.font.color.rgb = dark_gray

        for line in lines:
            match = section_pattern.match(line.strip())
            if match and len(match.group(1)) >= 4:
                # Flush previous section
                _flush_section()
                current_heading = match.group(1)
                current_lines = []
                # If there's text after the header on the same line
                rest = match.group(2).strip()
                if rest:
                    current_lines.append(rest)
            else:
                current_lines.append(line)

        # Flush last section
        _flush_section()

    # ── Footer ──
    doc.add_paragraph("")
    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_f = footer_para.add_run("Generated by ProspectX Intelligence Platform")
    run_f.font.size = Pt(8)
    run_f.font.color.rgb = dark_gray
    run_f.italic = True

    # ── Write to BytesIO ──
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


@app.get("/reports/{report_id}/export/docx")
async def export_report_docx(report_id: str, token_data: dict = Depends(verify_token)):
    """Export a report as a Word DOCX document."""
    if not _docx_available:
        raise HTTPException(status_code=501, detail="DOCX export not available — python-docx not installed")

    org_id = token_data["org_id"]
    conn = get_db()
    try:
        report_row = conn.execute(
            "SELECT * FROM reports WHERE id = ? AND org_id = ?",
            (report_id, org_id)
        ).fetchone()
        if not report_row:
            raise HTTPException(status_code=404, detail="Report not found")
        report = dict(report_row)

        # Get player info if this is a player report
        player_row = None
        if report.get("player_id"):
            p = conn.execute(
                "SELECT first_name, last_name, position, team FROM players WHERE id = ?",
                (report["player_id"],)
            ).fetchone()
            if p:
                player_row = dict(p)

        docx_buf = _generate_report_docx(report, player_row)

        # Build filename
        subject = ""
        if player_row:
            subject = f"{player_row.get('first_name', '')}_{player_row.get('last_name', '')}"
        elif report.get("team_name"):
            subject = report["team_name"].replace(" ", "_")
        report_type = report.get("report_type", "report")
        filename = f"ProspectX_{subject}_{report_type}.docx" if subject else f"ProspectX_{report_type}.docx"
        # Sanitize filename
        filename = re.sub(r'[^\w\-.]', '_', filename)

        return StreamingResponse(
            docx_buf,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f'attachment; filename="{filename}"'}
        )
    finally:
        conn.close()


# ── Email Share Preview ──────────────────────────────────────────

@app.post("/reports/{report_id}/email-preview")
async def email_preview(report_id: str, token_data: dict = Depends(verify_token)):
    """Generate pre-formatted email content for sharing a report via email.
    Auto-generates a share token if one doesn't exist."""
    org_id = token_data["org_id"]
    conn = get_db()
    try:
        report_row = conn.execute(
            "SELECT * FROM reports WHERE id = ? AND org_id = ?",
            (report_id, org_id)
        ).fetchone()
        if not report_row:
            raise HTTPException(status_code=404, detail="Report not found")
        report = dict(report_row)

        # Auto-generate share token if not set
        share_token = report.get("share_token")
        if not share_token:
            share_token = str(uuid.uuid4())[:12]
            conn.execute(
                "UPDATE reports SET share_token = ? WHERE id = ?",
                (share_token, report_id)
            )
            conn.commit()

        # Build share URL
        frontend_url = os.getenv("FRONTEND_URL", "http://localhost:3000")
        share_url = f"{frontend_url}/reports/{report_id}?token={share_token}"

        # Get player info for subject line
        subject_name = ""
        if report.get("player_id"):
            p = conn.execute(
                "SELECT first_name, last_name FROM players WHERE id = ?",
                (report["player_id"],)
            ).fetchone()
            if p:
                subject_name = f"{p['first_name']} {p['last_name']}"
        elif report.get("team_name"):
            subject_name = report["team_name"]

        report_type = report.get("report_type", "report").replace("_", " ").title()

        # Build email subject
        if subject_name:
            email_subject = f"ProspectX Report: {subject_name} - {report_type}"
        else:
            email_subject = f"ProspectX Report: {report_type}"

        # Build email body — concise summary
        grade = report.get("overall_grade", "")
        gen_date = report.get("generated_at", "")
        if gen_date:
            try:
                dt = datetime.fromisoformat(gen_date.replace("Z", "+00:00"))
                gen_date = dt.strftime("%B %d, %Y")
            except Exception:
                pass

        body_lines = [f"I'm sharing a ProspectX Intelligence report with you."]
        body_lines.append("")
        if subject_name:
            body_lines.append(f"Subject: {subject_name}")
        body_lines.append(f"Report Type: {report_type}")
        if grade:
            body_lines.append(f"Overall Grade: {grade}")
        if gen_date:
            body_lines.append(f"Generated: {gen_date}")
        body_lines.append("")
        body_lines.append(f"View the full report here:")
        body_lines.append(share_url)
        body_lines.append("")
        body_lines.append("---")
        body_lines.append("Sent via ProspectX Intelligence Platform")

        return {
            "subject": email_subject,
            "body": "\n".join(body_lines),
            "share_url": share_url,
            "share_token": share_token,
        }
    finally:
        conn.close()


@app.get("/bench-talk/suggestions")
async def get_bench_talk_suggestions(token_data: dict = Depends(verify_token)):
    """Get personalized suggestion prompts for Bench Talk.

    Combines 2 role-based suggestions with up to 4 activity-based suggestions
    pulled from the user's real data (recent notes, active game plans, series, etc.).
    """
    user_id = token_data["user_id"]
    org_id = token_data["org_id"]
    conn = get_db()

    try:
        user_row = conn.execute("SELECT hockey_role, first_name FROM users WHERE id = ?", (user_id,)).fetchone()
        hockey_role = user_row["hockey_role"] if user_row and user_row["hockey_role"] else "scout"

        # ── 2 core role-based suggestions (always shown) ──────────────
        role_suggestions = {
            "scout": [
                {"text": "Show me GOHL scoring leaders", "icon": "trophy"},
                {"text": "Compare two players side by side", "icon": "compare"},
            ],
            "gm": [
                {"text": "Show me the top scorers in the GOHL", "icon": "trophy"},
                {"text": "Compare two players for a trade target", "icon": "compare"},
            ],
            "coach": [
                {"text": "Who are my top performers this season?", "icon": "trophy"},
                {"text": "Help me build line combinations", "icon": "search"},
            ],
            "player": [
                {"text": "How do my stats compare to the league?", "icon": "compare"},
                {"text": "Show me the league leaders in my position", "icon": "trophy"},
            ],
            "parent": [
                {"text": "Help me understand my kid's stats", "icon": "search"},
                {"text": "What do scouts look for in a player?", "icon": "trophy"},
            ],
        }
        suggestions = list(role_suggestions.get(hockey_role, role_suggestions["scout"]))

        # ── Activity-based suggestions (up to 4) ─────────────────────
        activity = []

        # 1. Last scouted player (most recent note)
        last_note = conn.execute("""
            SELECT p.first_name, p.last_name, p.current_team
            FROM scout_notes sn JOIN players p ON sn.player_id = p.id
            WHERE sn.scout_id = ? AND sn.org_id = ? AND (p.is_deleted = 0 OR p.is_deleted IS NULL)
            ORDER BY sn.created_at DESC LIMIT 1
        """, (user_id, org_id)).fetchone()
        if last_note:
            name = f"{last_note['first_name']} {last_note['last_name']}"
            activity.append({"text": f"Tell me about {name}'s development", "icon": "search"})

        # 2. Active game plan
        active_gp = conn.execute("""
            SELECT opponent FROM game_plans
            WHERE org_id = ? AND status = 'active' ORDER BY date ASC LIMIT 1
        """, (org_id,)).fetchone()
        if active_gp and active_gp["opponent"]:
            activity.append({"text": f"Help me prep for {active_gp['opponent']}", "icon": "shield"})

        # 3. Active series
        active_series = conn.execute("""
            SELECT series_name, opponent FROM series_plans
            WHERE org_id = ? AND status = 'active' LIMIT 1
        """, (org_id,)).fetchone()
        if active_series:
            opp = active_series["opponent"] or active_series["series_name"]
            activity.append({"text": f"Update me on the {opp} series", "icon": "shield"})

        # 4. User's most-scouted team
        top_team = conn.execute("""
            SELECT p.current_team, COUNT(*) as cnt FROM scout_notes sn
            JOIN players p ON sn.player_id = p.id
            WHERE sn.scout_id = ? AND p.current_team IS NOT NULL
              AND (p.is_deleted = 0 OR p.is_deleted IS NULL)
            GROUP BY p.current_team ORDER BY cnt DESC LIMIT 1
        """, (user_id,)).fetchone()
        if top_team and top_team["current_team"]:
            activity.append({"text": f"How are the {top_team['current_team']} performing?", "icon": "trophy"})

        # 5. Recently created player (fallback if no notes)
        if not last_note:
            recent_player = conn.execute("""
                SELECT first_name, last_name FROM players
                WHERE created_by = ? AND org_id = ? AND (is_deleted = 0 OR is_deleted IS NULL)
                ORDER BY created_at DESC LIMIT 1
            """, (user_id, org_id)).fetchone()
            if recent_player:
                name = f"{recent_player['first_name']} {recent_player['last_name']}"
                activity.append({"text": f"Generate a report for {name}", "icon": "file"})

        # 6. Cold-start suggestions (if very little activity)
        counts = conn.execute("""
            SELECT
                (SELECT COUNT(*) FROM reports WHERE org_id = ?) as report_count,
                (SELECT COUNT(*) FROM scout_notes WHERE scout_id = ? AND org_id = ?) as note_count
        """, (org_id, user_id, org_id)).fetchone()

        if counts["report_count"] == 0:
            activity.append({"text": "Generate your first scouting report", "icon": "file"})
        if counts["note_count"] == 0 and not any("scout note" in a["text"].lower() for a in activity):
            activity.append({"text": "How do I write better scout notes?", "icon": "file"})

        # Combine: 2 role + up to 4 activity (no duplicates)
        seen_texts = {s["text"] for s in suggestions}
        for item in activity[:4]:
            if item["text"] not in seen_texts:
                suggestions.append(item)
                seen_texts.add(item["text"])
                if len(suggestions) >= 6:
                    break

    finally:
        conn.close()

    return {"suggestions": suggestions}


class BenchTalkContextRequest(BaseModel):
    player_ids: list[str] = []
    report_ids: list[str] = []


# ── Calendar & Schedule ────────────────────────────────────
class CalendarEventCreate(BaseModel):
    type: str = "OTHER"
    title: str
    description: Optional[str] = None
    start_time: str  # ISO 8601
    end_time: Optional[str] = None
    timezone: str = "America/Toronto"
    location: Optional[str] = None
    league_name: Optional[str] = None
    opponent_name: Optional[str] = None
    is_home: Optional[bool] = None
    team_id: Optional[str] = None
    player_id: Optional[str] = None
    visibility: str = "ORG"


class CalendarEventUpdate(BaseModel):
    type: Optional[str] = None
    title: Optional[str] = None
    description: Optional[str] = None
    start_time: Optional[str] = None
    end_time: Optional[str] = None
    timezone: Optional[str] = None
    location: Optional[str] = None
    league_name: Optional[str] = None
    opponent_name: Optional[str] = None
    is_home: Optional[bool] = None
    team_id: Optional[str] = None
    player_id: Optional[str] = None
    visibility: Optional[str] = None


class CalendarFeedCreate(BaseModel):
    label: str
    provider: str = "ICAL_GENERIC"
    url: str
    team_id: Optional[str] = None


# ── Messaging ──────────────────────────────────────────────
class SendMessageRequest(BaseModel):
    conversation_id: Optional[str] = None
    recipient_id: Optional[str] = None
    content: str


class ContactRequestCreate(BaseModel):
    target_player_id: str
    message: Optional[str] = None


class ContactRequestResolve(BaseModel):
    status: str  # "approved" or "denied"


class BlockUserRequest(BaseModel):
    blocked_id: str
    reason: Optional[str] = None


@app.post("/bench-talk/context")
async def get_bench_talk_context(req: BenchTalkContextRequest, token_data: dict = Depends(verify_token)):
    """Fetch full Player and Report objects for Bench Talk sidebar context display."""
    user_id = token_data["user_id"]
    org_id = token_data["org_id"]
    conn = get_db()
    _check_email_verified(user_id, conn)
    try:
        players = []
        if req.player_ids:
            placeholders = ",".join("?" for _ in req.player_ids)
            rows = conn.execute(f"""
                SELECT p.id, p.org_id, p.first_name, p.last_name, p.dob, p.position,
                       p.shoots, p.height_cm, p.weight_kg, p.current_team, p.current_league,
                       p.archetype, p.image_url, p.birth_year, p.age_group,
                       p.draft_eligible_year, p.league_tier, p.created_at,
                       ps.gp, ps.g, ps.a, ps.p,
                       CASE WHEN ps.gp > 0 THEN ROUND(CAST(ps.p AS REAL) / ps.gp, 2) ELSE 0 END as ppg
                FROM players p
                LEFT JOIN player_stats ps ON p.id = ps.player_id AND ps.stat_type = 'season'
                WHERE p.id IN ({placeholders}) AND p.org_id = ?
            """, (*req.player_ids, org_id)).fetchall()
            for r in rows:
                p = dict(r)
                p["passports"] = []
                p["notes"] = None
                p["tags"] = []
                players.append(p)

        reports = []
        if req.report_ids:
            placeholders = ",".join("?" for _ in req.report_ids)
            rows = conn.execute(f"""
                SELECT id, org_id, player_id, team_name, report_type, title,
                       status, generated_at, llm_model, llm_tokens, created_at
                FROM reports
                WHERE id IN ({placeholders}) AND org_id = ?
            """, (*req.report_ids, org_id)).fetchall()
            for r in rows:
                rpt = dict(r)
                rpt["output_json"] = None
                rpt["output_text"] = None
                rpt["error_message"] = None
                reports.append(rpt)

        return {"players": players, "reports": reports}
    finally:
        conn.close()


# ============================================================
# GAME PLANS
# ============================================================

@app.post("/game-plans")
async def create_game_plan(
    request: dict = Body(...),
    token_data: dict = Depends(verify_token),
):
    """Create a new game plan."""
    org_id = token_data["org_id"]
    user_id = token_data["user_id"]
    conn = get_db()
    _check_tier_permission(user_id, "can_create_game_plans", conn)

    plan_id = str(uuid.uuid4())
    import json as _json
    conn.execute("""
        INSERT INTO game_plans (id, org_id, user_id, team_name, opponent_team_name,
            game_date, opponent_analysis, our_strategy, matchups, special_teams_plan,
            keys_to_game, lines_snapshot, status, session_type, talking_points,
            forecheck, breakout, defensive_system, what_worked, what_didnt_work,
            game_result, game_score, created_at, updated_at)
        VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, CURRENT_TIMESTAMP, CURRENT_TIMESTAMP)
    """, (plan_id, org_id, user_id,
          request.get("team_name", ""), request.get("opponent_team_name", ""),
          request.get("game_date"), request.get("opponent_analysis", ""),
          request.get("our_strategy", ""),
          _json.dumps(request.get("matchups", {})),
          request.get("special_teams_plan", ""),
          request.get("keys_to_game", ""),
          _json.dumps(request.get("lines_snapshot", {})),
          request.get("status", "draft"),
          request.get("session_type", "pre_game"),
          _json.dumps(request.get("talking_points", {})),
          request.get("forecheck", ""),
          request.get("breakout", ""),
          request.get("defensive_system", ""),
          request.get("what_worked", ""),
          request.get("what_didnt_work", ""),
          request.get("game_result", ""),
          request.get("game_score", "")))
    conn.commit()
    row = conn.execute("SELECT * FROM game_plans WHERE id = ?", (plan_id,)).fetchone()
    conn.close()
    return dict(row)


@app.get("/game-plans")
async def list_game_plans(
    team: Optional[str] = None,
    status: Optional[str] = None,
    session_type: Optional[str] = None,
    limit: int = Query(default=50, ge=1, le=200),
    skip: int = Query(default=0, ge=0),
    token_data: dict = Depends(verify_token),
):
    """List game plans (Chalk Talk sessions) for the organization."""
    org_id = token_data["org_id"]
    conn = get_db()
    query = "SELECT * FROM game_plans WHERE org_id = ?"
    params: list = [org_id]
    if team:
        query += " AND (LOWER(team_name) = LOWER(?) OR LOWER(opponent_team_name) = LOWER(?))"
        params.extend([team, team])
    if status:
        query += " AND status = ?"
        params.append(status)
    if session_type:
        query += " AND session_type = ?"
        params.append(session_type)
    query += " ORDER BY game_date DESC, created_at DESC LIMIT ? OFFSET ?"
    params.extend([limit, skip])
    rows = conn.execute(query, params).fetchall()
    conn.close()
    return [dict(r) for r in rows]


@app.get("/game-plans/{plan_id}")
async def get_game_plan(plan_id: str, token_data: dict = Depends(verify_token)):
    """Get a single game plan by ID."""
    org_id = token_data["org_id"]
    conn = get_db()
    row = conn.execute("SELECT * FROM game_plans WHERE id = ? AND org_id = ?", (plan_id, org_id)).fetchone()
    conn.close()
    if not row:
        raise HTTPException(status_code=404, detail="Game plan not found")
    return dict(row)


@app.put("/game-plans/{plan_id}")
async def update_game_plan(
    plan_id: str,
    request: dict = Body(...),
    token_data: dict = Depends(verify_token),
):
    """Update a game plan."""
    org_id = token_data["org_id"]
    conn = get_db()
    existing = conn.execute("SELECT * FROM game_plans WHERE id = ? AND org_id = ?", (plan_id, org_id)).fetchone()
    if not existing:
        conn.close()
        raise HTTPException(status_code=404, detail="Game plan not found")

    import json as _json
    allowed = {"team_name", "opponent_team_name", "game_date", "opponent_analysis",
               "our_strategy", "special_teams_plan", "keys_to_game", "status",
               "session_type", "forecheck", "breakout", "defensive_system",
               "what_worked", "what_didnt_work", "game_result", "game_score"}
    json_fields = {"matchups", "lines_snapshot", "talking_points"}

    sets = ["updated_at = CURRENT_TIMESTAMP"]
    params: list = []
    for field, value in request.items():
        if field in allowed:
            sets.append(f"{field} = ?")
            params.append(value)
        elif field in json_fields:
            sets.append(f"{field} = ?")
            params.append(_json.dumps(value) if not isinstance(value, str) else value)

    if len(sets) == 1:
        conn.close()
        raise HTTPException(status_code=400, detail="No valid fields to update")

    params.append(plan_id)
    conn.execute(f"UPDATE game_plans SET {', '.join(sets)} WHERE id = ?", params)
    conn.commit()
    row = conn.execute("SELECT * FROM game_plans WHERE id = ?", (plan_id,)).fetchone()
    conn.close()
    return dict(row)


@app.delete("/game-plans/{plan_id}")
async def delete_game_plan(plan_id: str, token_data: dict = Depends(verify_token)):
    """Delete a game plan."""
    org_id = token_data["org_id"]
    conn = get_db()
    result = conn.execute("DELETE FROM game_plans WHERE id = ? AND org_id = ?", (plan_id, org_id))
    conn.commit()
    conn.close()
    if result.rowcount == 0:
        raise HTTPException(status_code=404, detail="Game plan not found")
    return {"status": "deleted", "plan_id": plan_id}


# ============================================================
# SERIES PLANS
# ============================================================

@app.post("/series")
async def create_series_plan(
    request: dict = Body(...),
    token_data: dict = Depends(verify_token),
):
    """Create a new series plan."""
    org_id = token_data["org_id"]
    user_id = token_data["user_id"]
    conn = get_db()
    _check_tier_permission(user_id, "can_create_series", conn)

    series_id = str(uuid.uuid4())
    import json as _json
    conn.execute("""
        INSERT INTO series_plans (id, org_id, user_id, team_name, opponent_team_name,
            series_name, series_format, current_score, game_notes, working_strategies,
            needs_adjustment, status, opponent_systems, key_players_dossier,
            matchup_plan, adjustments, momentum_log, created_at, updated_at)
        VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, CURRENT_TIMESTAMP, CURRENT_TIMESTAMP)
    """, (series_id, org_id, user_id,
          request.get("team_name", ""), request.get("opponent_team_name", ""),
          request.get("series_name", ""), request.get("series_format", "best_of_7"),
          request.get("current_score", "0-0"),
          _json.dumps(request.get("game_notes", [])),
          _json.dumps(request.get("working_strategies", [])),
          _json.dumps(request.get("needs_adjustment", [])),
          request.get("status", "active"),
          _json.dumps(request.get("opponent_systems", {})),
          _json.dumps(request.get("key_players_dossier", [])),
          _json.dumps(request.get("matchup_plan", {})),
          _json.dumps(request.get("adjustments", [])),
          _json.dumps(request.get("momentum_log", []))))
    conn.commit()
    row = conn.execute("SELECT * FROM series_plans WHERE id = ?", (series_id,)).fetchone()
    conn.close()
    return dict(row)


@app.get("/series")
async def list_series_plans(
    status: Optional[str] = None,
    token_data: dict = Depends(verify_token),
):
    """List series plans for the organization."""
    org_id = token_data["org_id"]
    conn = get_db()
    query = "SELECT * FROM series_plans WHERE org_id = ?"
    params: list = [org_id]
    if status:
        query += " AND status = ?"
        params.append(status)
    query += " ORDER BY created_at DESC"
    rows = conn.execute(query, params).fetchall()
    conn.close()
    return [dict(r) for r in rows]


@app.get("/series/{series_id}")
async def get_series_plan(series_id: str, token_data: dict = Depends(verify_token)):
    """Get a single series plan."""
    org_id = token_data["org_id"]
    conn = get_db()
    row = conn.execute("SELECT * FROM series_plans WHERE id = ? AND org_id = ?", (series_id, org_id)).fetchone()
    conn.close()
    if not row:
        raise HTTPException(status_code=404, detail="Series plan not found")
    return dict(row)


@app.put("/series/{series_id}")
async def update_series_plan(
    series_id: str,
    request: dict = Body(...),
    token_data: dict = Depends(verify_token),
):
    """Update a series plan."""
    org_id = token_data["org_id"]
    conn = get_db()
    existing = conn.execute("SELECT * FROM series_plans WHERE id = ? AND org_id = ?", (series_id, org_id)).fetchone()
    if not existing:
        conn.close()
        raise HTTPException(status_code=404, detail="Series plan not found")

    import json as _json
    allowed = {"team_name", "opponent_team_name", "series_name", "series_format",
               "current_score", "status"}
    json_fields = {"game_notes", "working_strategies", "needs_adjustment",
                   "opponent_systems", "key_players_dossier", "matchup_plan",
                   "adjustments", "momentum_log"}

    sets = ["updated_at = CURRENT_TIMESTAMP"]
    params: list = []
    for field, value in request.items():
        if field in allowed:
            sets.append(f"{field} = ?")
            params.append(value)
        elif field in json_fields:
            sets.append(f"{field} = ?")
            params.append(_json.dumps(value) if not isinstance(value, str) else value)

    if len(sets) == 1:
        conn.close()
        raise HTTPException(status_code=400, detail="No valid fields to update")

    params.append(series_id)
    conn.execute(f"UPDATE series_plans SET {', '.join(sets)} WHERE id = ?", params)
    conn.commit()
    row = conn.execute("SELECT * FROM series_plans WHERE id = ?", (series_id,)).fetchone()
    conn.close()
    return dict(row)


@app.delete("/series/{series_id}")
async def delete_series_plan(series_id: str, token_data: dict = Depends(verify_token)):
    """Delete a series plan."""
    org_id = token_data["org_id"]
    conn = get_db()
    result = conn.execute("DELETE FROM series_plans WHERE id = ? AND org_id = ?", (series_id, org_id))
    conn.commit()
    conn.close()
    if result.rowcount == 0:
        raise HTTPException(status_code=404, detail="Series plan not found")
    return {"status": "deleted", "series_id": series_id}


# ============================================================
# SCOUTING LIST
# ============================================================


@app.post("/scouting-list")
async def add_to_scouting_list(
    request: dict = Body(...),
    token_data: dict = Depends(verify_token),
):
    """Add a player to the scouting list."""
    org_id = token_data["org_id"]
    user_id = token_data["user_id"]
    conn = get_db()
    _check_tier_permission(user_id, "can_use_scouting_list", conn)

    player_id = request.get("player_id")
    if not player_id:
        conn.close()
        raise HTTPException(status_code=400, detail="player_id is required")

    # Check player exists
    player = conn.execute("SELECT id FROM players WHERE id = ?", (player_id,)).fetchone()
    if not player:
        conn.close()
        raise HTTPException(status_code=404, detail="Player not found")

    # Check if already on list
    existing = conn.execute(
        "SELECT id FROM scouting_list WHERE org_id = ? AND user_id = ? AND player_id = ?",
        (org_id, user_id, player_id)
    ).fetchone()
    if existing:
        conn.close()
        raise HTTPException(status_code=409, detail="Player already on scouting list")

    # Check list limit
    tier_config = _check_tier_permission(user_id, "can_use_scouting_list", conn)
    max_list = tier_config.get("max_scouting_list", 0)
    if max_list > 0:
        current_count = conn.execute(
            "SELECT COUNT(*) FROM scouting_list WHERE org_id = ? AND user_id = ? AND is_active = 1",
            (org_id, user_id)
        ).fetchone()[0]
        if current_count >= max_list:
            conn.close()
            raise HTTPException(status_code=403, detail=f"Scouting list limit reached ({max_list}). Upgrade for more.")

    import json as _json
    item_id = str(uuid.uuid4())
    conn.execute("""
        INSERT INTO scouting_list (id, org_id, user_id, player_id, priority,
            target_reason, scout_notes, tags, is_active, list_order,
            created_at, updated_at)
        VALUES (?, ?, ?, ?, ?, ?, ?, ?, 1, 0, CURRENT_TIMESTAMP, CURRENT_TIMESTAMP)
    """, (item_id, org_id, user_id, player_id,
          request.get("priority", "medium"),
          request.get("target_reason", ""),
          request.get("scout_notes", ""),
          _json.dumps(request.get("tags", []))))
    conn.commit()
    row = conn.execute("""
        SELECT sl.*, p.first_name, p.last_name, p.position, p.current_team,
               p.current_league, p.image_url
        FROM scouting_list sl
        JOIN players p ON sl.player_id = p.id
        WHERE sl.id = ?
    """, (item_id,)).fetchone()
    conn.close()
    return dict(row)


@app.get("/scouting-list")
async def list_scouting_list(
    priority: Optional[str] = None,
    is_active: Optional[int] = Query(default=1),
    search: Optional[str] = None,
    limit: int = Query(default=50, ge=1, le=200),
    skip: int = Query(default=0, ge=0),
    token_data: dict = Depends(verify_token),
):
    """List scouting list entries with player info."""
    org_id = token_data["org_id"]
    user_id = token_data["user_id"]
    conn = get_db()
    query = """
        SELECT sl.*, p.first_name, p.last_name, p.position, p.current_team,
               p.current_league, p.image_url
        FROM scouting_list sl
        JOIN players p ON sl.player_id = p.id
        WHERE sl.org_id = ? AND sl.user_id = ?
    """
    params: list = [org_id, user_id]
    if is_active is not None:
        query += " AND sl.is_active = ?"
        params.append(is_active)
    if priority:
        query += " AND sl.priority = ?"
        params.append(priority)
    if search:
        query += " AND (LOWER(p.first_name || ' ' || p.last_name) LIKE LOWER(?))"
        params.append(f"%{search}%")
    query += " ORDER BY CASE sl.priority WHEN 'high' THEN 1 WHEN 'medium' THEN 2 WHEN 'low' THEN 3 END, sl.created_at DESC"
    query += " LIMIT ? OFFSET ?"
    params.extend([limit, skip])
    rows = conn.execute(query, params).fetchall()
    conn.close()
    return [dict(r) for r in rows]


@app.get("/scouting-list/{item_id}")
async def get_scouting_list_item(item_id: str, token_data: dict = Depends(verify_token)):
    """Get a single scouting list entry."""
    org_id = token_data["org_id"]
    user_id = token_data["user_id"]
    conn = get_db()
    row = conn.execute("""
        SELECT sl.*, p.first_name, p.last_name, p.position, p.current_team,
               p.current_league, p.image_url
        FROM scouting_list sl
        JOIN players p ON sl.player_id = p.id
        WHERE sl.id = ? AND sl.org_id = ? AND sl.user_id = ?
    """, (item_id, org_id, user_id)).fetchone()
    conn.close()
    if not row:
        raise HTTPException(status_code=404, detail="Scouting list entry not found")
    return dict(row)


@app.put("/scouting-list/{item_id}")
async def update_scouting_list_item(
    item_id: str,
    request: dict = Body(...),
    token_data: dict = Depends(verify_token),
):
    """Update a scouting list entry."""
    org_id = token_data["org_id"]
    user_id = token_data["user_id"]
    conn = get_db()
    existing = conn.execute(
        "SELECT id FROM scouting_list WHERE id = ? AND org_id = ? AND user_id = ?",
        (item_id, org_id, user_id)
    ).fetchone()
    if not existing:
        conn.close()
        raise HTTPException(status_code=404, detail="Scouting list entry not found")

    import json as _json
    allowed = {"priority", "target_reason", "scout_notes", "is_active", "list_order"}
    json_fields = {"tags"}

    sets = ["updated_at = CURRENT_TIMESTAMP"]
    params: list = []
    for field, value in request.items():
        if field in allowed:
            sets.append(f"{field} = ?")
            params.append(value)
        elif field in json_fields:
            sets.append(f"{field} = ?")
            params.append(_json.dumps(value) if not isinstance(value, str) else value)

    if len(sets) == 1:
        conn.close()
        raise HTTPException(status_code=400, detail="No valid fields to update")

    params.append(item_id)
    conn.execute(f"UPDATE scouting_list SET {', '.join(sets)} WHERE id = ?", params)
    conn.commit()
    row = conn.execute("""
        SELECT sl.*, p.first_name, p.last_name, p.position, p.current_team,
               p.current_league, p.image_url
        FROM scouting_list sl
        JOIN players p ON sl.player_id = p.id
        WHERE sl.id = ?
    """, (item_id,)).fetchone()
    conn.close()
    return dict(row)


@app.delete("/scouting-list/{item_id}")
async def remove_from_scouting_list(item_id: str, token_data: dict = Depends(verify_token)):
    """Remove a player from the scouting list."""
    org_id = token_data["org_id"]
    user_id = token_data["user_id"]
    conn = get_db()
    result = conn.execute(
        "DELETE FROM scouting_list WHERE id = ? AND org_id = ? AND user_id = ?",
        (item_id, org_id, user_id)
    )
    conn.commit()
    conn.close()
    if result.rowcount == 0:
        raise HTTPException(status_code=404, detail="Scouting list entry not found")
    return {"status": "removed", "item_id": item_id}


@app.post("/scouting-list/{item_id}/view")
async def track_scouting_view(item_id: str, token_data: dict = Depends(verify_token)):
    """Track a view of a scouting list entry."""
    org_id = token_data["org_id"]
    user_id = token_data["user_id"]
    conn = get_db()
    result = conn.execute("""
        UPDATE scouting_list
        SET times_viewed = times_viewed + 1, last_viewed = CURRENT_TIMESTAMP
        WHERE id = ? AND org_id = ? AND user_id = ?
    """, (item_id, org_id, user_id))
    conn.commit()
    conn.close()
    if result.rowcount == 0:
        raise HTTPException(status_code=404, detail="Scouting list entry not found")
    return {"status": "viewed"}


# ============================================================
# MY DATA
# ============================================================

@app.get("/my-data/summary")
async def my_data_summary(token_data: dict = Depends(verify_token)):
    """Get aggregate counts for the current user's contributions."""
    org_id = token_data["org_id"]
    user_id = token_data["user_id"]
    conn = get_db()

    players_created = conn.execute(
        "SELECT COUNT(*) FROM players WHERE org_id = ? AND created_by = ? AND (is_deleted = 0 OR is_deleted IS NULL)",
        (org_id, user_id)
    ).fetchone()[0]

    uploads = conn.execute(
        "SELECT COUNT(*) FROM import_jobs WHERE org_id = ? AND user_id = ?",
        (org_id, user_id)
    ).fetchone()[0]

    corrections = conn.execute(
        "SELECT COUNT(*) FROM player_corrections WHERE org_id = ? AND user_id = ?",
        (org_id, user_id)
    ).fetchone()[0]

    corrections_approved = conn.execute(
        "SELECT COUNT(*) FROM player_corrections WHERE org_id = ? AND user_id = ? AND status = 'approved'",
        (org_id, user_id)
    ).fetchone()[0]

    reports = conn.execute(
        "SELECT COUNT(*) FROM reports WHERE org_id = ? AND user_id = ?",
        (org_id, user_id)
    ).fetchone()[0]

    notes = conn.execute(
        "SELECT COUNT(*) FROM scout_notes WHERE org_id = ? AND user_id = ?",
        (org_id, user_id)
    ).fetchone()[0]

    conn.close()
    return {
        "players_created": players_created,
        "uploads": uploads,
        "corrections_submitted": corrections,
        "corrections_approved": corrections_approved,
        "reports_generated": reports,
        "notes_created": notes,
    }


@app.get("/my-data/uploads")
async def my_data_uploads(token_data: dict = Depends(verify_token)):
    """Get upload/import history for the current user."""
    org_id = token_data["org_id"]
    user_id = token_data["user_id"]
    conn = get_db()
    rows = conn.execute("""
        SELECT * FROM import_jobs WHERE org_id = ? AND user_id = ?
        ORDER BY created_at DESC LIMIT 50
    """, (org_id, user_id)).fetchall()
    conn.close()
    return [dict(r) for r in rows]


@app.get("/my-data/players")
async def my_data_players(token_data: dict = Depends(verify_token)):
    """Get players created by the current user."""
    org_id = token_data["org_id"]
    user_id = token_data["user_id"]
    conn = get_db()
    rows = conn.execute("""
        SELECT id, first_name, last_name, position, current_team, current_league, created_at
        FROM players WHERE org_id = ? AND created_by = ? AND (is_deleted = 0 OR is_deleted IS NULL)
        ORDER BY created_at DESC LIMIT 100
    """, (org_id, user_id)).fetchall()
    conn.close()
    return [dict(r) for r in rows]


@app.get("/my-data/corrections")
async def my_data_corrections(token_data: dict = Depends(verify_token)):
    """Get corrections submitted by the current user."""
    org_id = token_data["org_id"]
    user_id = token_data["user_id"]
    conn = get_db()
    rows = conn.execute("""
        SELECT c.*, p.first_name, p.last_name
        FROM player_corrections c
        LEFT JOIN players p ON c.player_id = p.id
        WHERE c.org_id = ? AND c.user_id = ?
        ORDER BY c.created_at DESC LIMIT 100
    """, (org_id, user_id)).fetchall()
    conn.close()
    return [dict(r) for r in rows]


# ============================================================
# CALENDAR & SCHEDULE
# ============================================================


def _sync_ical_feed(conn, feed_id: str, org_id: str) -> dict:
    """Fetch and parse an iCal feed, upsert events into the events table."""
    try:
        from icalendar import Calendar as ICalendar
    except ImportError:
        return {"synced_count": 0, "errors": ["icalendar library not installed"]}

    feed = conn.execute("SELECT * FROM calendar_feeds WHERE id = ? AND org_id = ?", (feed_id, org_id)).fetchone()
    if not feed:
        return {"synced_count": 0, "errors": ["Feed not found"]}

    feed_url = feed["url"]
    errors = []
    synced = 0

    try:
        resp = httpx.get(feed_url, timeout=30.0, follow_redirects=True)
        resp.raise_for_status()
        ical_text = resp.text
    except Exception as e:
        error_msg = f"Failed to fetch feed: {str(e)}"
        conn.execute("UPDATE calendar_feeds SET sync_error = ?, updated_at = datetime('now') WHERE id = ?", (error_msg, feed_id))
        conn.commit()
        return {"synced_count": 0, "errors": [error_msg]}

    try:
        cal = ICalendar.from_ical(ical_text)
    except Exception as e:
        error_msg = f"Failed to parse iCal: {str(e)}"
        conn.execute("UPDATE calendar_feeds SET sync_error = ?, updated_at = datetime('now') WHERE id = ?", (error_msg, feed_id))
        conn.commit()
        return {"synced_count": 0, "errors": [error_msg]}

    for component in cal.walk():
        if component.name != "VEVENT":
            continue
        try:
            uid = str(component.get("UID", ""))
            summary = str(component.get("SUMMARY", "Untitled Event"))
            description = str(component.get("DESCRIPTION", "")) if component.get("DESCRIPTION") else None
            location = str(component.get("LOCATION", "")) if component.get("LOCATION") else None

            dtstart = component.get("DTSTART")
            dtend = component.get("DTEND")
            start_dt = dtstart.dt if dtstart else None
            end_dt = dtend.dt if dtend else None

            if not start_dt:
                continue

            # Convert date to datetime if needed
            if hasattr(start_dt, "isoformat"):
                start_str = start_dt.isoformat()
            else:
                start_str = str(start_dt)

            end_str = None
            if end_dt:
                if hasattr(end_dt, "isoformat"):
                    end_str = end_dt.isoformat()
                else:
                    end_str = str(end_dt)

            # Timezone
            tz = "America/Toronto"
            if dtstart and hasattr(dtstart, "params") and "TZID" in dtstart.params:
                tz = str(dtstart.params["TZID"])

            # Heuristics for event type and opponent
            event_type = "GAME"
            opponent = None
            is_home = None
            summary_lower = summary.lower()

            if "practice" in summary_lower or "skate" in summary_lower:
                event_type = "PRACTICE"
            elif "tournament" in summary_lower or "tourney" in summary_lower:
                event_type = "TOURNAMENT"
            elif "showcase" in summary_lower:
                event_type = "SHOWCASE"
            elif "meeting" in summary_lower:
                event_type = "MEETING"
            elif "deadline" in summary_lower:
                event_type = "DEADLINE"

            # Try to extract opponent from "vs" or "@"
            if " vs " in summary or " vs. " in summary:
                parts = re.split(r"\s+vs\.?\s+", summary, maxsplit=1)
                if len(parts) == 2:
                    opponent = parts[1].strip()
                    is_home = 1
            elif " @ " in summary or " at " in summary_lower:
                parts = re.split(r"\s+[@]\s+|\s+at\s+", summary, maxsplit=1, flags=re.IGNORECASE)
                if len(parts) == 2:
                    opponent = parts[1].strip()
                    is_home = 0

            # Upsert: check if exists by source_external_id
            existing = conn.execute(
                "SELECT id FROM events WHERE org_id = ? AND source = 'ICAL' AND source_external_id = ?",
                (org_id, uid)
            ).fetchone()

            if existing:
                conn.execute("""
                    UPDATE events SET title = ?, start_time = ?, end_time = ?, location = ?,
                    description = ?, type = ?, opponent_name = ?, is_home = ?, timezone = ?,
                    updated_at = datetime('now')
                    WHERE id = ?
                """, (summary, start_str, end_str, location, description, event_type,
                      opponent, is_home, tz, existing["id"]))
            else:
                event_id = str(uuid.uuid4())
                conn.execute("""
                    INSERT INTO events (id, org_id, team_id, feed_id, type, source, source_external_id,
                    title, description, start_time, end_time, timezone, location,
                    opponent_name, is_home, visibility)
                    VALUES (?, ?, ?, ?, ?, 'ICAL', ?, ?, ?, ?, ?, ?, ?, ?, ?, 'ORG')
                """, (event_id, org_id, feed["team_id"], feed_id, event_type, uid,
                      summary, description, start_str, end_str, tz, location,
                      opponent, is_home))
            synced += 1
        except Exception as e:
            errors.append(f"Event parse error: {str(e)}")

    # Update feed metadata
    event_count = conn.execute("SELECT COUNT(*) FROM events WHERE feed_id = ?", (feed_id,)).fetchone()[0]
    conn.execute("""
        UPDATE calendar_feeds SET last_sync_at = datetime('now'), event_count = ?,
        sync_error = NULL, updated_at = datetime('now')
        WHERE id = ?
    """, (event_count, feed_id))
    conn.commit()

    return {"synced_count": synced, "errors": errors}


@app.get("/api/calendar/events")
async def get_calendar_events(
    token_data: dict = Depends(verify_token),
    from_date: Optional[str] = Query(None, alias="from"),
    to_date: Optional[str] = Query(None, alias="to"),
    team_id: Optional[str] = None,
    player_id: Optional[str] = None,
    event_type: Optional[str] = Query(None, alias="type"),
):
    """Get calendar events for the user's org, filtered by params."""
    org_id = token_data["org_id"]
    conn = get_db()

    query = "SELECT * FROM events WHERE org_id = ?"
    params: list = [org_id]

    if from_date:
        query += " AND start_time >= ?"
        params.append(from_date)
    if to_date:
        query += " AND start_time <= ?"
        params.append(to_date + "T23:59:59")
    if team_id:
        query += " AND team_id = ?"
        params.append(team_id)
    if player_id:
        query += " AND player_id = ?"
        params.append(player_id)
    if event_type:
        query += " AND type = ?"
        params.append(event_type)

    query += " ORDER BY start_time ASC"
    rows = conn.execute(query, params).fetchall()
    conn.close()
    return [dict(r) for r in rows]


@app.post("/api/calendar/events")
async def create_calendar_event(
    body: CalendarEventCreate,
    token_data: dict = Depends(verify_token),
):
    """Create a manual calendar event."""
    org_id = token_data["org_id"]
    user_id = token_data["user_id"]
    event_id = str(uuid.uuid4())
    conn = get_db()
    conn.execute("""
        INSERT INTO events (id, org_id, team_id, player_id, created_by_user_id, type, source,
        title, description, start_time, end_time, timezone, location,
        league_name, opponent_name, is_home, visibility)
        VALUES (?, ?, ?, ?, ?, ?, 'MANUAL', ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
    """, (event_id, org_id, body.team_id, body.player_id, user_id, body.type,
          body.title, body.description, body.start_time, body.end_time, body.timezone,
          body.location, body.league_name, body.opponent_name,
          1 if body.is_home is True else (0 if body.is_home is False else None),
          body.visibility))
    conn.commit()
    row = conn.execute("SELECT * FROM events WHERE id = ?", (event_id,)).fetchone()
    conn.close()
    return dict(row)


@app.put("/api/calendar/events/{event_id}")
async def update_calendar_event(
    event_id: str,
    body: CalendarEventUpdate,
    token_data: dict = Depends(verify_token),
):
    """Update a manual calendar event. Only creator or admin can edit. Only manual events."""
    org_id = token_data["org_id"]
    user_id = token_data["user_id"]
    conn = get_db()
    event = conn.execute("SELECT * FROM events WHERE id = ? AND org_id = ?", (event_id, org_id)).fetchone()
    if not event:
        conn.close()
        raise HTTPException(status_code=404, detail="Event not found")
    if event["source"] != "MANUAL":
        conn.close()
        raise HTTPException(status_code=400, detail="Only manual events can be edited")
    if event["created_by_user_id"] != user_id:
        role = token_data.get("role", "")
        if role != "admin":
            conn.close()
            raise HTTPException(status_code=403, detail="Only the creator or admin can edit this event")

    updates = {}
    for field in ["type", "title", "description", "start_time", "end_time", "timezone",
                  "location", "league_name", "opponent_name", "team_id", "player_id", "visibility"]:
        val = getattr(body, field, None)
        if val is not None:
            updates[field] = val
    if body.is_home is not None:
        updates["is_home"] = 1 if body.is_home else 0

    if updates:
        set_clause = ", ".join(f"{k} = ?" for k in updates)
        vals = list(updates.values())
        vals.append(event_id)
        conn.execute(f"UPDATE events SET {set_clause}, updated_at = datetime('now') WHERE id = ?", vals)
        conn.commit()

    row = conn.execute("SELECT * FROM events WHERE id = ?", (event_id,)).fetchone()
    conn.close()
    return dict(row)


@app.delete("/api/calendar/events/{event_id}")
async def delete_calendar_event(
    event_id: str,
    token_data: dict = Depends(verify_token),
):
    """Delete a manual calendar event."""
    org_id = token_data["org_id"]
    user_id = token_data["user_id"]
    conn = get_db()
    event = conn.execute("SELECT * FROM events WHERE id = ? AND org_id = ?", (event_id, org_id)).fetchone()
    if not event:
        conn.close()
        raise HTTPException(status_code=404, detail="Event not found")
    if event["source"] != "MANUAL":
        conn.close()
        raise HTTPException(status_code=400, detail="Only manual events can be deleted")
    if event["created_by_user_id"] != user_id:
        role = token_data.get("role", "")
        if role != "admin":
            conn.close()
            raise HTTPException(status_code=403, detail="Only the creator or admin can delete this event")

    conn.execute("DELETE FROM events WHERE id = ?", (event_id,))
    conn.commit()
    conn.close()
    return {"detail": "Event deleted"}


@app.get("/api/calendar/feeds")
async def get_calendar_feeds(token_data: dict = Depends(verify_token)):
    """Get all calendar feeds for the user's org."""
    org_id = token_data["org_id"]
    conn = get_db()
    rows = conn.execute("SELECT * FROM calendar_feeds WHERE org_id = ? ORDER BY created_at DESC", (org_id,)).fetchall()
    conn.close()
    return [dict(r) for r in rows]


@app.post("/api/calendar/feeds")
async def create_calendar_feed(
    body: CalendarFeedCreate,
    token_data: dict = Depends(verify_token),
):
    """Create a calendar feed and trigger initial sync."""
    org_id = token_data["org_id"]
    feed_id = str(uuid.uuid4())
    conn = get_db()
    conn.execute("""
        INSERT INTO calendar_feeds (id, org_id, team_id, label, provider, url)
        VALUES (?, ?, ?, ?, ?, ?)
    """, (feed_id, org_id, body.team_id, body.label, body.provider, body.url))
    conn.commit()

    # Trigger initial sync
    sync_result = _sync_ical_feed(conn, feed_id, org_id)

    row = conn.execute("SELECT * FROM calendar_feeds WHERE id = ?", (feed_id,)).fetchone()
    conn.close()
    result = dict(row)
    result["sync_result"] = sync_result
    return result


@app.post("/api/calendar/feeds/{feed_id}/sync")
async def sync_calendar_feed(
    feed_id: str,
    token_data: dict = Depends(verify_token),
):
    """Manually trigger sync for a specific feed."""
    org_id = token_data["org_id"]
    conn = get_db()
    feed = conn.execute("SELECT * FROM calendar_feeds WHERE id = ? AND org_id = ?", (feed_id, org_id)).fetchone()
    if not feed:
        conn.close()
        raise HTTPException(status_code=404, detail="Feed not found")

    result = _sync_ical_feed(conn, feed_id, org_id)
    conn.close()
    return result


@app.delete("/api/calendar/feeds/{feed_id}")
async def delete_calendar_feed(
    feed_id: str,
    token_data: dict = Depends(verify_token),
):
    """Delete a calendar feed and all its events."""
    org_id = token_data["org_id"]
    conn = get_db()
    feed = conn.execute("SELECT * FROM calendar_feeds WHERE id = ? AND org_id = ?", (feed_id, org_id)).fetchone()
    if not feed:
        conn.close()
        raise HTTPException(status_code=404, detail="Feed not found")

    conn.execute("DELETE FROM events WHERE feed_id = ? AND org_id = ?", (feed_id, org_id))
    conn.execute("DELETE FROM calendar_feeds WHERE id = ?", (feed_id,))
    conn.commit()
    conn.close()
    return {"detail": "Feed and associated events deleted"}


# ============================================================
# MESSAGING + PARENTAL APPROVAL
# ============================================================


# ── Safety Rules Engine ────────────────────────────────────
MINOR_ROLES = {"player"}
INTERNAL_ROLES = {"coach", "gm", "admin"}


def _check_messaging_permission(conn, sender_id: str, recipient_id: str, sender_org_id: str) -> dict:
    """Check if sender can message recipient. Returns {allowed, reason, parent_id}."""
    # Check blocks
    block = conn.execute(
        "SELECT id FROM message_blocks WHERE (blocker_id = ? AND blocked_id = ?) OR (blocker_id = ? AND blocked_id = ?)",
        (recipient_id, sender_id, sender_id, recipient_id)
    ).fetchone()
    if block:
        return {"allowed": False, "reason": "blocked"}

    # Get sender + recipient user info
    sender = conn.execute("SELECT id, org_id, hockey_role, email FROM users WHERE id = ?", (sender_id,)).fetchone()
    recipient = conn.execute("SELECT id, org_id, hockey_role, email, linked_player_id FROM users WHERE id = ?", (recipient_id,)).fetchone()

    if not sender or not recipient:
        return {"allowed": False, "reason": "not_allowed"}

    sender_role = (sender["hockey_role"] or "").lower()
    recipient_role = (recipient["hockey_role"] or "").lower()
    same_org = sender["org_id"] == recipient["org_id"]

    # Player-to-player messaging blocked in v1
    if sender_role in MINOR_ROLES and recipient_role in MINOR_ROLES:
        return {"allowed": False, "reason": "not_allowed"}

    # Same org: free messaging
    if same_org:
        return {"allowed": True}

    # External: check if recipient is a minor (player role = minor in v1)
    if recipient_role in MINOR_ROLES:
        # Find parent linked to this player
        parent = conn.execute(
            "SELECT id FROM users WHERE linked_player_id = ? AND hockey_role = 'parent'",
            (recipient_id,)
        ).fetchone()
        if not parent:
            # Try by linked_player_id matching player's associated player record
            # Check if recipient has a linked player in the players table
            player_rec = conn.execute(
                "SELECT id FROM players WHERE id = ? OR (LOWER(first_name || ' ' || last_name) IN (SELECT LOWER(first_name || ' ' || last_name) FROM users WHERE id = ?))",
                (recipient_id, recipient_id)
            ).fetchone()
            if player_rec:
                parent = conn.execute(
                    "SELECT id FROM users WHERE linked_player_id = ?",
                    (player_rec["id"],)
                ).fetchone()
        if parent:
            # Check if there's already an approved contact request
            approved = conn.execute(
                "SELECT id FROM contact_requests WHERE requester_id = ? AND target_player_id = ? AND status = 'approved'",
                (sender_id, recipient_id)
            ).fetchone()
            if approved:
                return {"allowed": True}
            return {"allowed": False, "reason": "requires_approval", "parent_id": parent["id"]}
        return {"allowed": False, "reason": "no_parent_linked"}

    # External messaging to parent role — check if sender role needs approval
    if recipient_role == "parent":
        # Scouts/agents messaging parents externally needs approval via the player
        return {"allowed": True}  # Parent can decide to respond or block

    # Default: allow for adult-to-adult cross-org
    return {"allowed": True}


def _hydrate_participants(conn, participant_ids: list) -> list:
    """Hydrate user IDs into participant objects."""
    participants = []
    for uid in participant_ids:
        user = conn.execute(
            "SELECT id, first_name, last_name, hockey_role, org_id FROM users WHERE id = ?",
            (uid,)
        ).fetchone()
        if user:
            org = conn.execute("SELECT name FROM organizations WHERE id = ?", (user["org_id"],)).fetchone()
            participants.append({
                "user_id": user["id"],
                "name": f"{user['first_name'] or ''} {user['last_name'] or ''}".strip() or user["id"][:8],
                "role": user["hockey_role"] or "unknown",
                "org_name": org["name"] if org else None,
                "is_verified": True,
            })
    return participants


@app.get("/api/messages/conversations")
async def get_conversations(token_data: dict = Depends(verify_token)):
    """Get all conversations for the authenticated user."""
    user_id = token_data["user_id"]
    conn = get_db()

    # Get user info (check if parent)
    user = conn.execute("SELECT hockey_role, linked_player_id FROM users WHERE id = ?", (user_id,)).fetchone()
    user_role = (user["hockey_role"] or "").lower() if user else ""
    linked_player_id = user["linked_player_id"] if user else None

    # Find conversations where user is a participant
    rows = conn.execute("SELECT * FROM msg_conversations ORDER BY updated_at DESC").fetchall()
    result = []

    for row in rows:
        try:
            pids = json.loads(row["participant_ids"])
        except (json.JSONDecodeError, TypeError):
            continue

        # User is participant OR parent with linked player in conversation
        is_participant = user_id in pids
        is_parent_observer = (user_role == "parent" and linked_player_id and linked_player_id in pids)

        if not is_participant and not is_parent_observer:
            continue

        conv = dict(row)
        conv["participant_ids"] = pids
        conv["participants"] = _hydrate_participants(conn, pids)

        # Last message
        last_msg = conn.execute(
            "SELECT m.*, u.first_name, u.last_name FROM msg_messages m LEFT JOIN users u ON m.sender_id = u.id WHERE m.conversation_id = ? ORDER BY m.sent_at DESC LIMIT 1",
            (row["id"],)
        ).fetchone()
        if last_msg:
            conv["last_message"] = {
                "id": last_msg["id"],
                "conversation_id": last_msg["conversation_id"],
                "sender_id": last_msg["sender_id"],
                "sender_name": f"{last_msg['first_name'] or ''} {last_msg['last_name'] or ''}".strip(),
                "content": last_msg["content"],
                "sent_at": last_msg["sent_at"],
                "read_at": last_msg["read_at"],
                "is_system_message": bool(last_msg["is_system_message"]),
            }

        # Unread count
        unread = conn.execute(
            "SELECT COUNT(*) FROM msg_messages WHERE conversation_id = ? AND sender_id != ? AND read_at IS NULL",
            (row["id"], user_id)
        ).fetchone()[0]
        conv["unread_count"] = unread

        result.append(conv)

    conn.close()
    return result


@app.post("/api/messages/send")
async def send_message(body: SendMessageRequest, token_data: dict = Depends(verify_token)):
    """Send a message. Creates conversation if needed."""
    user_id = token_data["user_id"]
    org_id = token_data["org_id"]
    conn = get_db()

    if not body.content or not body.content.strip():
        conn.close()
        raise HTTPException(status_code=400, detail="Message cannot be empty")

    conversation_id = body.conversation_id

    if conversation_id:
        # Verify user is participant
        conv = conn.execute("SELECT * FROM msg_conversations WHERE id = ?", (conversation_id,)).fetchone()
        if not conv:
            conn.close()
            raise HTTPException(status_code=404, detail="Conversation not found")
        pids = json.loads(conv["participant_ids"])
        if user_id not in pids:
            conn.close()
            raise HTTPException(status_code=403, detail="Not a participant")
        if conv["status"] == "blocked":
            conn.close()
            raise HTTPException(status_code=403, detail="This conversation has been blocked")
        if conv["status"] == "pending_approval":
            conn.close()
            raise HTTPException(status_code=403, detail="Waiting for parental approval")
    elif body.recipient_id:
        # Check safety rules
        perm = _check_messaging_permission(conn, user_id, body.recipient_id, org_id)
        if not perm["allowed"]:
            conn.close()
            if perm["reason"] == "requires_approval":
                raise HTTPException(status_code=403, detail=json.dumps({
                    "error": "approval_required",
                    "parent_id": perm.get("parent_id"),
                    "message": "This player is a minor. A contact request must be sent to their parent for approval."
                }))
            elif perm["reason"] == "blocked":
                raise HTTPException(status_code=403, detail="You have been blocked by this user")
            elif perm["reason"] == "no_parent_linked":
                raise HTTPException(status_code=403, detail="No parent account linked to this player")
            else:
                raise HTTPException(status_code=403, detail="Messaging not allowed")

        # Find or create conversation
        existing = conn.execute("SELECT * FROM msg_conversations WHERE status = 'active'").fetchall()
        conversation_id = None
        for ex in existing:
            try:
                ex_pids = json.loads(ex["participant_ids"])
            except (json.JSONDecodeError, TypeError):
                continue
            if set(ex_pids) == {user_id, body.recipient_id}:
                conversation_id = ex["id"]
                break

        if not conversation_id:
            conversation_id = str(uuid.uuid4())
            participants = json.dumps([user_id, body.recipient_id])
            conv_org = org_id if conn.execute("SELECT org_id FROM users WHERE id = ?", (body.recipient_id,)).fetchone()["org_id"] == org_id else None
            conn.execute(
                "INSERT INTO msg_conversations (id, org_id, participant_ids, status) VALUES (?, ?, ?, 'active')",
                (conversation_id, conv_org, participants)
            )
    else:
        conn.close()
        raise HTTPException(status_code=400, detail="Either conversation_id or recipient_id is required")

    # Create message
    msg_id = str(uuid.uuid4())
    conn.execute(
        "INSERT INTO msg_messages (id, conversation_id, sender_id, content) VALUES (?, ?, ?, ?)",
        (msg_id, conversation_id, user_id, body.content.strip())
    )
    conn.execute("UPDATE msg_conversations SET updated_at = datetime('now') WHERE id = ?", (conversation_id,))
    conn.commit()

    # Get sender name
    sender = conn.execute("SELECT first_name, last_name FROM users WHERE id = ?", (user_id,)).fetchone()
    sender_name = f"{sender['first_name'] or ''} {sender['last_name'] or ''}".strip() if sender else ""

    msg = conn.execute("SELECT * FROM msg_messages WHERE id = ?", (msg_id,)).fetchone()
    conn.close()

    return {
        "id": msg["id"],
        "conversation_id": msg["conversation_id"],
        "sender_id": msg["sender_id"],
        "sender_name": sender_name,
        "content": msg["content"],
        "sent_at": msg["sent_at"],
        "read_at": msg["read_at"],
        "is_system_message": False,
    }


@app.get("/api/messages/{conversation_id}")
async def get_messages(conversation_id: str, token_data: dict = Depends(verify_token)):
    """Get messages for a conversation. Marks unread as read."""
    user_id = token_data["user_id"]
    conn = get_db()

    conv = conn.execute("SELECT * FROM msg_conversations WHERE id = ?", (conversation_id,)).fetchone()
    if not conv:
        conn.close()
        raise HTTPException(status_code=404, detail="Conversation not found")

    pids = json.loads(conv["participant_ids"])

    # Check access: participant or parent observer
    user = conn.execute("SELECT hockey_role, linked_player_id FROM users WHERE id = ?", (user_id,)).fetchone()
    user_role = (user["hockey_role"] or "").lower() if user else ""
    linked_player_id = user["linked_player_id"] if user else None
    is_participant = user_id in pids
    is_parent_observer = (user_role == "parent" and linked_player_id and linked_player_id in pids)

    if not is_participant and not is_parent_observer:
        conn.close()
        raise HTTPException(status_code=403, detail="Access denied")

    # Mark unread messages as read (only if direct participant, not parent observer)
    if is_participant:
        conn.execute(
            "UPDATE msg_messages SET read_at = datetime('now') WHERE conversation_id = ? AND sender_id != ? AND read_at IS NULL",
            (conversation_id, user_id)
        )
        conn.commit()

    # Fetch messages
    rows = conn.execute(
        "SELECT m.*, u.first_name, u.last_name FROM msg_messages m LEFT JOIN users u ON m.sender_id = u.id WHERE m.conversation_id = ? ORDER BY m.sent_at ASC",
        (conversation_id,)
    ).fetchall()

    messages = []
    for r in rows:
        messages.append({
            "id": r["id"],
            "conversation_id": r["conversation_id"],
            "sender_id": r["sender_id"],
            "sender_name": f"{r['first_name'] or ''} {r['last_name'] or ''}".strip(),
            "content": r["content"],
            "sent_at": r["sent_at"],
            "read_at": r["read_at"],
            "is_system_message": bool(r["is_system_message"]),
        })

    conn.close()
    return {
        "conversation": dict(conv) | {"participant_ids": pids, "participants": _hydrate_participants(conn if not conn else get_db(), pids)},
        "messages": messages,
        "is_parent_observer": is_parent_observer,
    }


@app.post("/api/messages/contact-request")
async def create_contact_request(body: ContactRequestCreate, token_data: dict = Depends(verify_token)):
    """Create a contact request for messaging a minor."""
    user_id = token_data["user_id"]
    org_id = token_data["org_id"]
    conn = get_db()

    # Get requester info
    requester = conn.execute("SELECT first_name, last_name, hockey_role FROM users WHERE id = ?", (user_id,)).fetchone()
    if not requester:
        conn.close()
        raise HTTPException(status_code=404, detail="User not found")

    org = conn.execute("SELECT name FROM organizations WHERE id = ?", (org_id,)).fetchone()
    requester_name = f"{requester['first_name'] or ''} {requester['last_name'] or ''}".strip()
    requester_role = requester["hockey_role"] or "unknown"
    requester_org = org["name"] if org else ""

    # Find parent for target player
    parent = conn.execute(
        "SELECT id FROM users WHERE linked_player_id = ? AND LOWER(hockey_role) = 'parent'",
        (body.target_player_id,)
    ).fetchone()
    if not parent:
        conn.close()
        raise HTTPException(status_code=400, detail="No parent account linked to this player. Contact request cannot be sent.")

    parent_id = parent["id"]

    # Check: blocked?
    block = conn.execute(
        "SELECT id FROM message_blocks WHERE blocker_id = ? AND blocked_id = ?",
        (parent_id, user_id)
    ).fetchone()
    if block:
        conn.close()
        raise HTTPException(status_code=403, detail="You have been blocked by this player's parent")

    # Check: pending request already?
    pending = conn.execute(
        "SELECT id FROM contact_requests WHERE requester_id = ? AND target_player_id = ? AND status = 'pending'",
        (user_id, body.target_player_id)
    ).fetchone()
    if pending:
        conn.close()
        raise HTTPException(status_code=400, detail="You already have a pending contact request for this player")

    # Check: cooldown from previous denial?
    denied = conn.execute(
        "SELECT cooldown_until FROM contact_requests WHERE requester_id = ? AND target_player_id = ? AND status = 'denied' ORDER BY resolved_at DESC LIMIT 1",
        (user_id, body.target_player_id)
    ).fetchone()
    if denied and denied["cooldown_until"]:
        now_str = datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M:%S")
        if now_str < denied["cooldown_until"]:
            conn.close()
            raise HTTPException(status_code=403, detail=f"Contact request denied. You can try again after {denied['cooldown_until'][:10]}")

    # Create request
    req_id = str(uuid.uuid4())
    conn.execute(
        "INSERT INTO contact_requests (id, requester_id, requester_name, requester_role, requester_org, target_player_id, parent_id, message) VALUES (?, ?, ?, ?, ?, ?, ?, ?)",
        (req_id, user_id, requester_name, requester_role, requester_org, body.target_player_id, parent_id, body.message)
    )
    conn.commit()

    row = conn.execute("SELECT * FROM contact_requests WHERE id = ?", (req_id,)).fetchone()
    conn.close()
    result = dict(row)

    # Add target player name
    pconn = get_db()
    player = pconn.execute("SELECT first_name, last_name FROM players WHERE id = ?", (body.target_player_id,)).fetchone()
    pconn.close()
    if player:
        result["target_player_name"] = f"{player['first_name'] or ''} {player['last_name'] or ''}".strip()

    return result


@app.get("/api/messages/contact-requests")
async def get_contact_requests(token_data: dict = Depends(verify_token)):
    """Get pending contact requests (for parents and admins)."""
    user_id = token_data["user_id"]
    conn = get_db()

    user = conn.execute("SELECT hockey_role FROM users WHERE id = ?", (user_id,)).fetchone()
    role = (user["hockey_role"] or "").lower() if user else ""

    if role == "parent":
        rows = conn.execute(
            "SELECT cr.*, p.first_name, p.last_name FROM contact_requests cr LEFT JOIN players p ON cr.target_player_id = p.id WHERE cr.parent_id = ? ORDER BY cr.requested_at DESC",
            (user_id,)
        ).fetchall()
    elif role in ("admin", "gm"):
        org_id = token_data["org_id"]
        rows = conn.execute(
            "SELECT cr.*, p.first_name, p.last_name FROM contact_requests cr LEFT JOIN players p ON cr.target_player_id = p.id LEFT JOIN users u ON cr.parent_id = u.id WHERE u.org_id = ? ORDER BY cr.requested_at DESC",
            (org_id,)
        ).fetchall()
    else:
        conn.close()
        return []

    result = []
    for r in rows:
        d = dict(r)
        if r["first_name"]:
            d["target_player_name"] = f"{r['first_name'] or ''} {r['last_name'] or ''}".strip()
        result.append(d)

    conn.close()
    return result


@app.put("/api/messages/contact-requests/{request_id}")
async def resolve_contact_request(request_id: str, body: ContactRequestResolve, token_data: dict = Depends(verify_token)):
    """Approve or deny a contact request. Parent or admin only."""
    user_id = token_data["user_id"]
    conn = get_db()

    req = conn.execute("SELECT * FROM contact_requests WHERE id = ?", (request_id,)).fetchone()
    if not req:
        conn.close()
        raise HTTPException(status_code=404, detail="Contact request not found")

    # Verify parent or admin
    if req["parent_id"] != user_id:
        user = conn.execute("SELECT hockey_role FROM users WHERE id = ?", (user_id,)).fetchone()
        role = (user["hockey_role"] or "").lower() if user else ""
        if role not in ("admin", "gm"):
            conn.close()
            raise HTTPException(status_code=403, detail="Only the parent or admin can resolve this request")

    if body.status not in ("approved", "denied"):
        conn.close()
        raise HTTPException(status_code=400, detail="Status must be 'approved' or 'denied'")

    now_str = datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M:%S")

    if body.status == "approved":
        # Create conversation between requester and target player
        conv_id = str(uuid.uuid4())
        participants = json.dumps([req["requester_id"], req["target_player_id"]])
        conn.execute(
            "INSERT INTO msg_conversations (id, org_id, participant_ids, status) VALUES (?, NULL, ?, 'active')",
            (conv_id, participants)
        )

        # Add system message
        parent_user = conn.execute("SELECT first_name, last_name FROM users WHERE id = ?", (req["parent_id"],)).fetchone()
        parent_name = f"{parent_user['first_name'] or ''} {parent_user['last_name'] or ''}".strip() if parent_user else "Parent"
        sys_msg_id = str(uuid.uuid4())
        conn.execute(
            "INSERT INTO msg_messages (id, conversation_id, sender_id, content, is_system_message) VALUES (?, ?, ?, ?, 1)",
            (sys_msg_id, conv_id, req["parent_id"], f"Contact approved by {parent_name}. You may now communicate.")
        )

        conn.execute(
            "UPDATE contact_requests SET status = 'approved', resolved_at = ? WHERE id = ?",
            (now_str, request_id)
        )
    else:
        # Denied — set 30-day cooldown
        cooldown = (datetime.now(timezone.utc) + timedelta(days=30)).strftime("%Y-%m-%d %H:%M:%S")
        conn.execute(
            "UPDATE contact_requests SET status = 'denied', resolved_at = ?, cooldown_until = ? WHERE id = ?",
            (now_str, cooldown, request_id)
        )

    conn.commit()
    row = conn.execute("SELECT * FROM contact_requests WHERE id = ?", (request_id,)).fetchone()
    conn.close()
    return dict(row)


@app.post("/api/messages/block")
async def block_user(body: BlockUserRequest, token_data: dict = Depends(verify_token)):
    """Block a user from messaging."""
    user_id = token_data["user_id"]
    conn = get_db()

    # Check not blocking self
    if body.blocked_id == user_id:
        conn.close()
        raise HTTPException(status_code=400, detail="Cannot block yourself")

    # Create block
    block_id = str(uuid.uuid4())
    try:
        conn.execute(
            "INSERT INTO message_blocks (id, blocker_id, blocked_id, reason) VALUES (?, ?, ?, ?)",
            (block_id, user_id, body.blocked_id, body.reason)
        )
    except sqlite3.IntegrityError:
        conn.close()
        return {"detail": "User already blocked"}

    # Block any active conversations between them
    convs = conn.execute("SELECT * FROM msg_conversations WHERE status = 'active'").fetchall()
    for c in convs:
        try:
            pids = json.loads(c["participant_ids"])
        except (json.JSONDecodeError, TypeError):
            continue
        if user_id in pids and body.blocked_id in pids:
            conn.execute("UPDATE msg_conversations SET status = 'blocked', updated_at = datetime('now') WHERE id = ?", (c["id"],))

    conn.commit()
    conn.close()
    return {"detail": "User blocked"}


@app.get("/api/messages/unread-count")
async def get_unread_count(token_data: dict = Depends(verify_token)):
    """Get total unread message count for nav badge."""
    user_id = token_data["user_id"]
    conn = get_db()

    # Get all conversations user is in
    convs = conn.execute("SELECT * FROM msg_conversations WHERE status = 'active'").fetchall()
    total_unread = 0
    for c in convs:
        try:
            pids = json.loads(c["participant_ids"])
        except (json.JSONDecodeError, TypeError):
            continue
        if user_id in pids:
            unread = conn.execute(
                "SELECT COUNT(*) FROM msg_messages WHERE conversation_id = ? AND sender_id != ? AND read_at IS NULL",
                (c["id"], user_id)
            ).fetchone()[0]
            total_unread += unread

    conn.close()
    return {"unread_count": total_unread}


# ============================================================
# MAIN
# ============================================================

if __name__ == "__main__":
    import uvicorn
    port = int(os.getenv("PORT", "8000"))
    logger.info("Starting ProspectX API (SQLite) on port %d", port)
    logger.info("Database: %s", DB_FILE)
    logger.info("API Docs: http://localhost:%d/docs", port)
    uvicorn.run(app, host="0.0.0.0", port=port)